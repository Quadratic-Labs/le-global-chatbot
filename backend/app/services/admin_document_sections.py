"""
Admin section editing - ORDER 8A architecture.

The CURRENT DOCX is the unique source of truth: there is no separate
"admin override" layer anymore. Edit and Add physically mutate a copy
of the source DOCX first, validate it by a full reparse, and only then
atomically apply both the targeted OpenSearch chunks and the new
current source file - never an override layered on top of an
unmodified DOCX, and OpenSearch is never the source of truth for
reconstructing editable text (it is checked only as a post-write
invariant).
"""
from __future__ import annotations
import os
import re
import tempfile
from collections.abc import Sequence
from pathlib import Path
from typing import Any
from docx import Document
from opensearchpy import OpenSearch
from app.core.country_registry import canonical_country_name
from app.models.admin_document_sections import AdminDocumentSectionAddResponse, AdminDocumentSectionDeleteResponse, AdminDocumentSectionListResponse, AdminDocumentSectionResponse, AdminDocumentSectionSummary, AdminDocumentSectionUpdateResponse
from app.services.admin_document_lifecycle import AdminDocumentLifecycleError, AdminDocumentRollbackError, _ensure_no_country_conflict, _get_document_metadata, _required_string, _validate_document_id
from app.services.document_section_state import mark_admin_modified, section_id_for_legal_topic
from app.services.country_lock import DEFAULT_LOCK_TIMEOUT_SECONDS, country_lock
from app.services.docx_parser import TopicLocation, locate_top_level_topics, parse_docx_sections
from app.services.document_chunk_builder import DocumentMetadata, build_document_chunks, metadata_from_content, resolve_effective_legal_topic, validate_docx_format
from app.services.document_indexer import DEFAULT_BULK_CHUNK_SIZE, DocumentIndexingError, _delete_chunks_except, _restore_section_snapshot, _snapshot_document_chunks, replace_document_section_chunks
from app.services.document_mutation import LegalTopicAlreadyExistsError, insert_top_level_topic, normalize_topic_title, remove_top_level_topic, rename_top_level_topic, replace_top_level_topic
from app.services.document_source_resolver import DocumentSourceConflictError, resolve_document_source_path
from app.clients.opensearch import get_opensearch_client
from app.services.document_chunk_builder import DEFAULT_MAX_CHARS, split_parsed_sections
_SLUG_PATTERN = re.compile('[^a-z0-9]+')
_MAX_SECTION_TITLE_LENGTH = 100

class AdminDocumentSectionNotFoundError(LookupError):
    """Raised when a requested section does not exist in the document."""

    def __init__(self, *, document_id: str, section_id: str) -> None:
        self.document_id = document_id
        self.section_id = section_id
        super().__init__(f'Section {section_id!r} was not found in document {document_id!r}.')

    def to_detail(self) -> dict[str, object]:
        return {'code': 'document_section_not_found', 'message': str(self), 'operation': 'section_update', 'document_id': self.document_id, 'section_id': self.section_id}

class AdminDocumentSectionInvalidError(ValueError):
    """Raised when the submitted section content cannot be accepted."""

    def __init__(self, *, message: str) -> None:
        super().__init__(message)

    def to_detail(self) -> dict[str, object]:
        return {'code': 'document_section_invalid', 'message': str(self), 'operation': 'section_update'}

class AdminDocumentSectionUpdateFailedError(RuntimeError):
    """Raised when a section could not be saved safely."""

    def __init__(self, *, message: str) -> None:
        super().__init__(message)

    def to_detail(self) -> dict[str, object]:
        return {'code': 'document_section_update_failed', 'message': str(self), 'operation': 'section_update'}

class AdminDocumentSectionAlreadyExistsError(ValueError):
    """
    Raised when a new or renamed section's title collides with an
    existing topic (mission "ORDER 8G-A" extended this from Add-only
    to also cover Rename - operation reflects whichever caused it).
    """

    def __init__(self, *, title: str, operation: str='section_add') -> None:
        self.title = title
        self.operation = operation
        super().__init__(f'A section already exists with this title: {title!r}.')

    def to_detail(self) -> dict[str, object]:
        return {'code': 'section_already_exists', 'message': str(self), 'operation': self.operation, 'title': self.title}

class AdminDocumentSectionPositionError(ValueError):
    """Raised when an Add position reference is missing or ambiguous."""

    def __init__(self, *, message: str) -> None:
        super().__init__(message)

    def to_detail(self) -> dict[str, object]:
        return {'code': 'document_section_invalid_position', 'message': str(self), 'operation': 'section_add'}

class AdminDocumentSectionLastRemainingError(ValueError):
    """
    Raised when deleting a section would leave the document with no
    usable top-level legal section left (mission "ORDER 8G-A", section
    6/7).
    """

    def __init__(self, *, document_id: str, section_id: str) -> None:
        self.document_id = document_id
        self.section_id = section_id
        super().__init__('This section cannot be deleted because it is the only remaining section in this document.')

    def to_detail(self) -> dict[str, object]:
        return {'code': 'section_is_last_remaining', 'message': str(self), 'operation': 'section_delete', 'document_id': self.document_id, 'section_id': self.section_id}

def _resolve_current_source_path(*, document_metadata: dict[str, Any], source_directory: Path) -> Path:
    country_code = _required_string(document_metadata, 'country_code')
    source_filename = _required_string(document_metadata, 'source_filename')
    try:
        resolved = resolve_document_source_path(source_root=source_directory, country_code=country_code, source_filename=source_filename)
    except DocumentSourceConflictError as error:
        raise AdminDocumentLifecycleError('Multiple distinct source files resolve for this document.') from error
    if resolved.path is None:
        raise AdminDocumentLifecycleError('The source DOCX backing this document is missing.')
    return resolved.path

def _current_topics(*, source_path: Path, country: str) -> list[TopicLocation]:
    """The CURRENT DOCX's own real top-level topics, right now."""
    document = Document(source_path)
    return locate_top_level_topics(document, country=country)

def _effective_legal_topic(parsed_section: Any, country: str) -> str | None:
    """The effective legal_topic of one parsed section, or None for
    overview/front-matter content that is not an editable section."""
    section = parsed_section.section.strip()
    try:
        document_type, legal_topic = resolve_effective_legal_topic(parsed_section=parsed_section, section=section, country=country)
    except Exception:
        return None
    if document_type != 'comparator':
        return None
    return legal_topic

def list_effective_sections(*, document_id: str, source_directory: Path, client: OpenSearch | None=None) -> AdminDocumentSectionListResponse:
    """
    Every top-level legal topic that REALLY exists in the current
    DOCX right now - known taxonomy topic or admin-added custom one
    alike - never a fixed list of the 11 historical topics, and never
    derived from OpenSearch (ORDER 8A, section 6).
    """
    document_id = _validate_document_id(document_id)
    opensearch_client = client or get_opensearch_client()
    document_metadata = _get_document_metadata(document_id=document_id, client=opensearch_client)
    country_code = _required_string(document_metadata, 'country_code')
    country = canonical_country_name(country_code)
    source_path = _resolve_current_source_path(document_metadata=document_metadata, source_directory=source_directory)
    topics = _current_topics(source_path=source_path, country=country)
    return AdminDocumentSectionListResponse(document_id=document_id, sections=[AdminDocumentSectionSummary(section_id=section_id_for_legal_topic(topic.legal_topic), legal_topic=topic.legal_topic) for topic in topics])

def get_effective_section(*, document_id: str, section_id: str, source_directory: Path, client: OpenSearch | None=None) -> AdminDocumentSectionResponse:
    """
    The current content of one section, extracted structurally from
    the CURRENT DOCX - never from OpenSearch, never from a persisted
    override (ORDER 8A, section 6).
    """
    document_id = _validate_document_id(document_id)
    opensearch_client = client or get_opensearch_client()
    document_metadata = _get_document_metadata(document_id=document_id, client=opensearch_client)
    country_code = _required_string(document_metadata, 'country_code')
    country = canonical_country_name(country_code)
    source_path = _resolve_current_source_path(document_metadata=document_metadata, source_directory=source_directory)
    topics = _current_topics(source_path=source_path, country=country)
    matching_topic = next((topic for topic in topics if section_id_for_legal_topic(topic.legal_topic) == section_id), None)
    if matching_topic is None:
        raise AdminDocumentSectionNotFoundError(document_id=document_id, section_id=section_id)
    parsed_sections = parse_docx_sections(source_path, country=country)
    matching_content = [parsed_section.content for parsed_section in parsed_sections if _effective_legal_topic(parsed_section, country) == matching_topic.legal_topic]
    content = '\n\n'.join(matching_content)
    return AdminDocumentSectionResponse(document_id=document_id, country_code=country_code, country_name=country, section_id=section_id, legal_topic=matching_topic.legal_topic, content=content)

def _make_temp_docx_path(source_path: Path) -> Path:
    """
    A temporary DOCX path in the SAME directory as source_path, so the
    later atomic replace (os.replace) is a same-filesystem rename,
    never a cross-device copy (ORDER 8A, section 15, step 4).
    """
    file_descriptor, temporary_path_str = tempfile.mkstemp(prefix=f'.{source_path.stem}-', suffix='.tmp.docx', dir=source_path.parent)
    os.close(file_descriptor)
    return Path(temporary_path_str)

def _fsync_path(path: Path) -> None:
    """fsync a file or directory by path, for durability after a write."""
    flags = os.O_RDONLY if path.is_dir() else os.O_RDWR
    file_descriptor = os.open(path, flags)
    try:
        os.fsync(file_descriptor)
    finally:
        os.close(file_descriptor)

def _build_metadata(*, document_metadata: dict[str, Any], detected: DocumentMetadata) -> DocumentMetadata:
    """
    Combine the freshly re-detected DOCX metadata (country/country_code
    from content) with the identity metadata the file keeps on disk
    (its own filename, reference year) - the file's storage identity
    never changes just because its content was edited.
    """
    reference_year = document_metadata.get('reference_year')
    return DocumentMetadata(country=detected.country, country_code=detected.country_code, reference_year=int(reference_year) if isinstance(reference_year, int) else detected.reference_year, language='en', source_filename=_required_string(document_metadata, 'source_filename'), source_format=detected.source_format)

def _validate_mutated_docx(*, temp_path: Path, country: str, expected_country_code: str, document_metadata: dict[str, Any]) -> DocumentMetadata:
    """
    ORDER 8A section 14: the mandatory reparse-validation every DOCX
    mutation must pass before anything is allowed to touch OpenSearch
    or the real source file. Returns the metadata to build chunks
    with. Raises AdminDocumentSectionUpdateFailedError on any failure
    - callers must guarantee zero mutation happened before this point.
    """
    try:
        validate_docx_format(temp_path)
    except Exception as error:
        raise AdminDocumentSectionUpdateFailedError(message='The updated document is not a valid DOCX file - no change was saved.') from error
    try:
        detected = metadata_from_content(temp_path, country_code=expected_country_code, language='en')
    except Exception as error:
        raise AdminDocumentSectionUpdateFailedError(message="The updated document's country could not be re-validated - no change was saved.") from error
    if detected.country_code.strip().upper() != expected_country_code.strip().upper():
        raise AdminDocumentSectionUpdateFailedError(message="The edit would change this document's country - no change was saved.")
    return _build_metadata(document_metadata=document_metadata, detected=detected)
_WHITESPACE_PATTERN = re.compile('\\s+')

def _normalize_for_comparison(text: str) -> str:
    return _WHITESPACE_PATTERN.sub(' ', text).strip()

def _build_topic_chunks(*, temp_path: Path, country: str, legal_topic: str, metadata: DocumentMetadata) -> list[Any]:
    """Build chunks for exactly one legal_topic from a (temp) DOCX."""
    parsed_sections = parse_docx_sections(temp_path, country=country)
    matching = [parsed_section for parsed_section in parsed_sections if _effective_legal_topic(parsed_section, country) == legal_topic]
    split_sections = split_parsed_sections(matching, max_chars=DEFAULT_MAX_CHARS)
    return build_document_chunks(split_sections, metadata)

def update_effective_section(*, document_id: str, section_id: str, new_content: str, new_title: str | None=None, source_directory: Path, client: OpenSearch | None=None, lock_timeout_seconds: float=DEFAULT_LOCK_TIMEOUT_SECONDS) -> AdminDocumentSectionUpdateResponse:
    """
    Edit an existing section: the CURRENT DOCX is really modified
    (ORDER 8A, section 8), validated by a full reparse (section 14),
    and only then applied atomically to both OpenSearch (targeted to
    this one legal_topic only) and the real source file (section 15).

    Mission "ORDER 8G-A", section 2/3 - new_title extends this same
    transaction with an optional rename: a rename is only a rename
    when the EFFECTIVE (normalized) title actually changes - an
    omitted title, or one that only differs by whitespace/casing, is
    the pre-existing content-only Edit path, byte-for-byte unchanged.
    A genuine rename additionally re-validates the reparsed document's
    topic set (old title gone, new title exactly once, topic count and
    every other topic unchanged) before touching OpenSearch, and rolls
    back through _rollback_rename (restore old chunks, remove any
    new-topic leftovers, restore the source DOCX) on any failure.
    """
    document_id = _validate_document_id(document_id)
    trimmed_content = new_content.strip()
    if not trimmed_content:
        raise AdminDocumentSectionInvalidError(message='Section content must not be empty.')
    opensearch_client = client or get_opensearch_client()
    preliminary_metadata = _get_document_metadata(document_id=document_id, client=opensearch_client)
    country_code_for_lock = _required_string(preliminary_metadata, 'country_code')
    with country_lock(source_directory, country_code_for_lock, timeout_seconds=lock_timeout_seconds):
        document_metadata = _get_document_metadata(document_id=document_id, client=opensearch_client)
        country_code = _required_string(document_metadata, 'country_code')
        if country_code != country_code_for_lock:
            raise AdminDocumentLifecycleError("The document's country changed during the operation.")
        _ensure_no_country_conflict(country_code=country_code, client=opensearch_client, operation='section_update')
        country = canonical_country_name(country_code)
        source_path = _resolve_current_source_path(document_metadata=document_metadata, source_directory=source_directory)
        existing_topics = _current_topics(source_path=source_path, country=country)
        matching_topic = next((topic for topic in existing_topics if section_id_for_legal_topic(topic.legal_topic) == section_id), None)
        if matching_topic is None:
            raise AdminDocumentSectionNotFoundError(document_id=document_id, section_id=section_id)
        old_legal_topic = matching_topic.legal_topic
        trimmed_new_title = new_title.strip() if new_title else None
        is_rename = trimmed_new_title is not None and normalize_topic_title(trimmed_new_title) != normalize_topic_title(old_legal_topic)
        if is_rename:
            trimmed_new_title = _validate_new_section_title(title=trimmed_new_title, existing_topics=[topic for topic in existing_topics if topic.legal_topic != old_legal_topic], operation='section_update')
        target_legal_topic = trimmed_new_title if is_rename else old_legal_topic
        original_bytes = source_path.read_bytes()
        try:
            pre_edit_snapshot = [item for item in _snapshot_document_chunks(client=opensearch_client, document_id=document_id) if item['_source'].get('legal_topic') == old_legal_topic]
        except DocumentIndexingError as error:
            raise AdminDocumentSectionUpdateFailedError(message="The section's current content could not be read from the search index before saving.") from error
        temp_path = _make_temp_docx_path(source_path)
        try:
            try:
                if is_rename:
                    rename_top_level_topic(file_path=source_path, output_path=temp_path, country=country, legal_topic=old_legal_topic, new_title=trimmed_new_title, new_content=trimmed_content)
                else:
                    replace_top_level_topic(file_path=source_path, output_path=temp_path, country=country, legal_topic=old_legal_topic, new_content=trimmed_content)
                _fsync_path(temp_path)
            except Exception as error:
                raise AdminDocumentSectionUpdateFailedError(message=f'The section could not be edited: {error}') from error
            metadata = _validate_mutated_docx(temp_path=temp_path, country=country, expected_country_code=country_code, document_metadata=document_metadata)
            if is_rename:
                try:
                    reparsed_topics = locate_top_level_topics(Document(temp_path), country=country)
                except Exception as error:
                    raise AdminDocumentSectionUpdateFailedError(message=f'The updated document could not be re-parsed - no change was saved: {error}') from error
                reparsed_names = [topic.legal_topic for topic in reparsed_topics]
                if old_legal_topic in reparsed_names:
                    raise AdminDocumentSectionUpdateFailedError(message="The section's old title is still present after renaming - no change was saved.")
                if reparsed_names.count(target_legal_topic) != 1:
                    raise AdminDocumentSectionUpdateFailedError(message="The section's new title could not be uniquely identified after renaming - no change was saved.")
                if len(reparsed_names) != len(existing_topics):
                    raise AdminDocumentSectionUpdateFailedError(message="Renaming this section changed the document's overall section structure - no change was saved.")
                if len(set(reparsed_names)) != len(reparsed_names):
                    raise AdminDocumentSectionUpdateFailedError(message='Renaming this section produced a duplicate section title - no change was saved.')
            try:
                new_chunks = _build_topic_chunks(temp_path=temp_path, country=country, legal_topic=target_legal_topic, metadata=metadata)
            except Exception as error:
                raise AdminDocumentSectionUpdateFailedError(message=f'The updated document could not be re-parsed - no change was saved: {error}') from error
            actual_topics = {chunk.legal_topic for chunk in new_chunks}
            if actual_topics != {target_legal_topic}:
                raise AdminDocumentSectionUpdateFailedError(message='The edited content did not resolve back to the same section - no change was saved.')
            actual_content = '\n\n'.join((chunk.content for chunk in new_chunks))
            if _normalize_for_comparison(actual_content) != _normalize_for_comparison(trimmed_content):
                raise AdminDocumentSectionUpdateFailedError(message='The saved content did not match what was submitted - no change was saved.')
            try:
                indexing_result = replace_document_section_chunks(new_chunks, target_legal_topic, client=opensearch_client)
            except DocumentIndexingError as error:
                raise AdminDocumentSectionUpdateFailedError(message='The section could not be saved to the search index.') from error
            if is_rename:
                try:
                    _delete_chunks_except(client=opensearch_client, filters=[{'term': {'document_id': document_id}}, {'term': {'legal_topic': old_legal_topic}}], keep_chunk_ids=[], context=f'document {document_id!r} rename cleanup')
                except DocumentIndexingError as error:
                    _rollback_rename(opensearch_client=opensearch_client, document_id=document_id, old_legal_topic=old_legal_topic, new_legal_topic=target_legal_topic, pre_rename_snapshot=pre_edit_snapshot, error=error, context='section rename')
            try:
                os.replace(temp_path, source_path)
                _fsync_path(source_path.parent)
            except OSError as replace_error:
                if is_rename:
                    _rollback_rename(opensearch_client=opensearch_client, document_id=document_id, old_legal_topic=old_legal_topic, new_legal_topic=target_legal_topic, pre_rename_snapshot=pre_edit_snapshot, error=replace_error, context='section rename')
                else:
                    _rollback_after_source_replace_failure(opensearch_client=opensearch_client, document_id=document_id, legal_topic=old_legal_topic, pre_mutation_snapshot=pre_edit_snapshot, replace_error=replace_error, context='section')
            temp_path = None
            try:
                _verify_section_invariant(source_path=source_path, country=country, legal_topic=target_legal_topic, metadata=metadata, document_id=document_id, client=opensearch_client)
                if is_rename:
                    _verify_topic_absent(document_id=document_id, legal_topic=old_legal_topic, client=opensearch_client)
            except AdminDocumentSectionUpdateFailedError as verify_error:
                if is_rename:
                    _rollback_rename(opensearch_client=opensearch_client, document_id=document_id, old_legal_topic=old_legal_topic, new_legal_topic=target_legal_topic, pre_rename_snapshot=pre_edit_snapshot, error=verify_error, context='section rename', source_path=source_path, original_source_bytes=original_bytes)
                else:
                    _rollback_after_invariant_failure(opensearch_client=opensearch_client, document_id=document_id, legal_topic=old_legal_topic, pre_mutation_snapshot=pre_edit_snapshot, verify_error=verify_error, context='section', source_path=source_path, original_source_bytes=original_bytes)
            mark_admin_modified(source_directory, document_id)
            return AdminDocumentSectionUpdateResponse(document_id=document_id, section_id=section_id_for_legal_topic(target_legal_topic), legal_topic=target_legal_topic, indexed_chunks=indexing_result.indexed_chunks)
        finally:
            if temp_path is not None:
                temp_path.unlink(missing_ok=True)

def _verify_section_invariant(*, source_path: Path, country: str, legal_topic: str, metadata: DocumentMetadata, document_id: str, client: OpenSearch) -> None:
    """
    ORDER 8A section 18: after success, deterministically verify that
    the target topic as parsed from the now-current DOCX matches what
    is actually indexed for it - derived from the parser/chunk builder,
    never from a raw textarea-text comparison.
    """
    try:
        expected_chunks = _build_topic_chunks(temp_path=source_path, country=country, legal_topic=legal_topic, metadata=metadata)
    except Exception as error:
        raise AdminDocumentSectionUpdateFailedError(message=f'The post-write invariant check could not re-parse the current document: {error}') from error
    try:
        indexed = [item for item in _snapshot_document_chunks(client=client, document_id=document_id) if item['_source'].get('legal_topic') == legal_topic]
    except DocumentIndexingError as error:
        raise AdminDocumentSectionUpdateFailedError(message='The post-write invariant check could not read the search index.') from error
    expected_hashes = sorted((chunk.content_hash for chunk in expected_chunks))
    indexed_hashes = sorted((item['_source'].get('content_hash') for item in indexed))
    if expected_hashes != indexed_hashes:
        raise AdminDocumentSectionUpdateFailedError(message="Post-write invariant check failed: the search index does not match the current document's own content for this section.")

def _verify_topic_absent(*, document_id: str, legal_topic: str, client: OpenSearch) -> None:
    """
    ORDER 8G-A - after a rename or delete, deterministically verify no
    chunk is left behind under a legal_topic that should no longer
    exist for this document (the old title after a rename; the
    deleted section's title after a delete).
    """
    try:
        remaining = [item for item in _snapshot_document_chunks(client=client, document_id=document_id) if item['_source'].get('legal_topic') == legal_topic]
    except DocumentIndexingError as error:
        raise AdminDocumentSectionUpdateFailedError(message='The post-write invariant check could not read the search index.') from error
    if remaining:
        raise AdminDocumentSectionUpdateFailedError(message='Post-write invariant check failed: chunks are still present in the search index under a topic that should no longer exist.')

def _rollback_after_source_replace_failure(*, opensearch_client: OpenSearch, document_id: str, legal_topic: str, pre_mutation_snapshot: list[dict[str, Any]], replace_error: OSError, context: str) -> None:
    """
    ORDER 8A section 16, case E: the atomic source replace failed
    AFTER OpenSearch already succeeded - roll OpenSearch back to
    exactly its pre-mutation snapshot (the old source is already
    intact, since os.replace never partially applies). Shared by
    Edit and Add - only the snapshot/topic and error message differ.
    """
    try:
        _restore_section_snapshot(client=opensearch_client, document_id=document_id, legal_topic=legal_topic, snapshot=pre_mutation_snapshot, bulk_chunk_size=DEFAULT_BULK_CHUNK_SIZE)
    except Exception as rollback_error:
        raise AdminDocumentRollbackError(f'The {context} could not be saved to the source document, and rolling the search index back afterward also failed - manual recovery is required.') from rollback_error
    raise AdminDocumentSectionUpdateFailedError(message=f'The {context} could not be saved to the source document; the search index was rolled back to its previous content.') from replace_error

def _rollback_after_invariant_failure(*, opensearch_client: OpenSearch, document_id: str, legal_topic: str, pre_mutation_snapshot: list[dict[str, Any]], verify_error: AdminDocumentSectionUpdateFailedError, context: str, source_path: Path | None=None, original_source_bytes: bytes | None=None) -> None:
    """
    ORDER 8A section 16, case F: failure after the source was already
    replaced - restore the original source exactly (when a pre-
    mutation copy was captured) AND restore OpenSearch exactly. Shared
    by Edit and Add.
    """
    source_restored = True
    if source_path is not None and original_source_bytes is not None:
        try:
            source_path.write_bytes(original_source_bytes)
            _fsync_path(source_path)
            _fsync_path(source_path.parent)
        except OSError:
            source_restored = False
    index_restored = True
    try:
        _restore_section_snapshot(client=opensearch_client, document_id=document_id, legal_topic=legal_topic, snapshot=pre_mutation_snapshot, bulk_chunk_size=DEFAULT_BULK_CHUNK_SIZE)
    except Exception:
        index_restored = False
    if not source_restored or not index_restored:
        raise AdminDocumentRollbackError(f'The post-write invariant check failed for the {context}, and the rollback afterward was incomplete (source ' + ('restored' if source_restored else 'NOT restored') + ', index ' + ('restored' if index_restored else 'NOT restored') + ') - manual recovery is required.') from verify_error
    raise verify_error

def _rollback_rename(*, opensearch_client: OpenSearch, document_id: str, old_legal_topic: str, new_legal_topic: str, pre_rename_snapshot: list[dict[str, Any]], error: Exception, context: str, source_path: Path | None=None, original_source_bytes: bytes | None=None) -> None:
    """
    ORDER 8G-A, section 3 - roll a failed section rename back to its
    exact pre-rename state: restore the OLD topic's chunks from its
    snapshot, remove any NEW-topic chunks a partially-completed rename
    already indexed, and (once the source DOCX had already been
    replaced) restore its original bytes too.

    source_path/original_source_bytes are only passed once the source
    file has actually been replaced already (the invariant-failure
    case) - for an earlier failure (chunk cleanup, or the atomic
    replace itself) the source file is untouched and needs no restore.
    """
    source_restored = True
    if source_path is not None and original_source_bytes is not None:
        try:
            source_path.write_bytes(original_source_bytes)
            _fsync_path(source_path)
            _fsync_path(source_path.parent)
        except OSError:
            source_restored = False
    old_restored = True
    try:
        _restore_section_snapshot(client=opensearch_client, document_id=document_id, legal_topic=old_legal_topic, snapshot=pre_rename_snapshot, bulk_chunk_size=DEFAULT_BULK_CHUNK_SIZE)
    except Exception:
        old_restored = False
    new_cleaned_up = True
    try:
        _delete_chunks_except(client=opensearch_client, filters=[{'term': {'document_id': document_id}}, {'term': {'legal_topic': new_legal_topic}}], keep_chunk_ids=[], context=f'document {document_id!r} rename rollback cleanup')
    except Exception:
        new_cleaned_up = False
    if not source_restored or not old_restored or (not new_cleaned_up):
        raise AdminDocumentRollbackError(f'The {context} failed, and the rollback afterward was incomplete (source ' + ('restored' if source_restored else 'NOT restored') + ', old section ' + ('restored' if old_restored else 'NOT restored') + ', new section ' + ('cleaned up' if new_cleaned_up else 'NOT cleaned up') + ') - manual recovery is required.') from error
    if isinstance(error, AdminDocumentSectionUpdateFailedError):
        raise error
    if isinstance(error, OSError):
        raise AdminDocumentSectionUpdateFailedError(message=f'The {context} could not be saved to the source document; the search index was rolled back to its previous content.') from error
    raise AdminDocumentSectionUpdateFailedError(message=f'The {context} could not be completed; the search index and source document were rolled back to their previous content.') from error

def _validate_new_section_title(*, title: str, existing_topics: Sequence[TopicLocation], operation: str='section_add') -> str:
    trimmed = ' '.join(title.split())
    if not trimmed:
        raise AdminDocumentSectionInvalidError(message='Section title must not be empty.')
    if len(trimmed) > _MAX_SECTION_TITLE_LENGTH:
        raise AdminDocumentSectionInvalidError(message=f'Section title must be at most {_MAX_SECTION_TITLE_LENGTH} characters.')
    normalized_new = normalize_topic_title(trimmed)
    for topic in existing_topics:
        if normalize_topic_title(topic.legal_topic) == normalized_new or section_id_for_legal_topic(topic.legal_topic) == section_id_for_legal_topic(trimmed):
            raise AdminDocumentSectionAlreadyExistsError(title=trimmed, operation=operation)
    return trimmed

def _resolve_add_position(*, position: str, existing_topics: Sequence[TopicLocation], document_id: str) -> str:
    normalized = position.strip()
    if normalized in ('beginning', 'end'):
        return normalized
    if normalized.startswith('after:'):
        target_section_id = normalized[len('after:'):]
        matching = next((topic for topic in existing_topics if section_id_for_legal_topic(topic.legal_topic) == target_section_id), None)
        if matching is None:
            raise AdminDocumentSectionNotFoundError(document_id=document_id, section_id=target_section_id)
        return f'after:{matching.legal_topic}'
    raise AdminDocumentSectionPositionError(message=f"Unsupported position: {position!r}. Use 'beginning', 'end', or 'after:<section_id>'.")

def add_new_section(*, document_id: str, title: str, content: str, position: str, source_directory: Path, client: OpenSearch | None=None, lock_timeout_seconds: float=DEFAULT_LOCK_TIMEOUT_SECONDS) -> AdminDocumentSectionAddResponse:
    """
    Add a brand-new top-level legal topic: a real new heading is added
    to the current DOCX at the requested position, validated by a full
    reparse, and only then applied atomically (ORDER 8A, sections 9-18).
    """
    document_id = _validate_document_id(document_id)
    trimmed_content = content.strip()
    if not trimmed_content:
        raise AdminDocumentSectionInvalidError(message='Section content must not be empty.')
    opensearch_client = client or get_opensearch_client()
    preliminary_metadata = _get_document_metadata(document_id=document_id, client=opensearch_client)
    country_code_for_lock = _required_string(preliminary_metadata, 'country_code')
    with country_lock(source_directory, country_code_for_lock, timeout_seconds=lock_timeout_seconds):
        document_metadata = _get_document_metadata(document_id=document_id, client=opensearch_client)
        country_code = _required_string(document_metadata, 'country_code')
        if country_code != country_code_for_lock:
            raise AdminDocumentLifecycleError("The document's country changed during the operation.")
        _ensure_no_country_conflict(country_code=country_code, client=opensearch_client, operation='section_add')
        country = canonical_country_name(country_code)
        source_path = _resolve_current_source_path(document_metadata=document_metadata, source_directory=source_directory)
        existing_topics = _current_topics(source_path=source_path, country=country)
        trimmed_title = _validate_new_section_title(title=title, existing_topics=existing_topics)
        resolved_position = _resolve_add_position(position=position, existing_topics=existing_topics, document_id=document_id)
        existing_legal_topics = {topic.legal_topic for topic in existing_topics}
        original_bytes = source_path.read_bytes()
        temp_path = _make_temp_docx_path(source_path)
        try:
            try:
                insert_top_level_topic(file_path=source_path, output_path=temp_path, country=country, title=trimmed_title, content=trimmed_content, position=resolved_position)
                _fsync_path(temp_path)
            except LegalTopicAlreadyExistsError as error:
                raise AdminDocumentSectionAlreadyExistsError(title=trimmed_title) from error
            except Exception as error:
                raise AdminDocumentSectionUpdateFailedError(message=f'The section could not be added: {error}') from error
            metadata = _validate_mutated_docx(temp_path=temp_path, country=country, expected_country_code=country_code, document_metadata=document_metadata)
            try:
                new_topics = locate_top_level_topics(Document(temp_path), country=country)
            except Exception as error:
                raise AdminDocumentSectionUpdateFailedError(message=f'The updated document could not be re-parsed - no change was saved: {error}') from error
            added_topics = [topic for topic in new_topics if topic.legal_topic not in existing_legal_topics]
            if len(added_topics) != 1:
                raise AdminDocumentSectionUpdateFailedError(message='The new section could not be uniquely identified after reparsing - no change was saved.')
            legal_topic = added_topics[0].legal_topic
            new_topic_names = [topic.legal_topic for topic in new_topics]
            new_index = new_topic_names.index(legal_topic)
            if resolved_position == 'beginning':
                position_ok = new_index == 0
            elif resolved_position == 'end':
                position_ok = new_index == len(new_topic_names) - 1
            else:
                after_topic = resolved_position[len('after:'):]
                position_ok = new_index > 0 and new_topic_names[new_index - 1] == after_topic
            if not position_ok:
                raise AdminDocumentSectionUpdateFailedError(message='The new section was not placed at the requested position - no change was saved.')
            try:
                new_chunks = _build_topic_chunks(temp_path=temp_path, country=country, legal_topic=legal_topic, metadata=metadata)
            except Exception as error:
                raise AdminDocumentSectionUpdateFailedError(message=f'The updated document could not be re-parsed - no change was saved: {error}') from error
            if {chunk.legal_topic for chunk in new_chunks} != {legal_topic}:
                raise AdminDocumentSectionUpdateFailedError(message="The new section's content did not resolve to a single topic - no change was saved.")
            try:
                pre_add_snapshot = [item for item in _snapshot_document_chunks(client=opensearch_client, document_id=document_id) if item['_source'].get('legal_topic') == legal_topic]
            except DocumentIndexingError as error:
                raise AdminDocumentSectionUpdateFailedError(message='The search index could not be read before adding the section.') from error
            try:
                indexing_result = replace_document_section_chunks(new_chunks, legal_topic, client=opensearch_client)
            except DocumentIndexingError as error:
                raise AdminDocumentSectionUpdateFailedError(message='The new section could not be saved to the search index.') from error
            try:
                os.replace(temp_path, source_path)
                _fsync_path(source_path.parent)
            except OSError as replace_error:
                _rollback_after_source_replace_failure(opensearch_client=opensearch_client, document_id=document_id, legal_topic=legal_topic, pre_mutation_snapshot=pre_add_snapshot, replace_error=replace_error, context='new section')
            temp_path = None
            try:
                _verify_section_invariant(source_path=source_path, country=country, legal_topic=legal_topic, metadata=metadata, document_id=document_id, client=opensearch_client)
            except AdminDocumentSectionUpdateFailedError as verify_error:
                _rollback_after_invariant_failure(opensearch_client=opensearch_client, document_id=document_id, legal_topic=legal_topic, pre_mutation_snapshot=pre_add_snapshot, verify_error=verify_error, context='new section', source_path=source_path, original_source_bytes=original_bytes)
            mark_admin_modified(source_directory, document_id)
            return AdminDocumentSectionAddResponse(document_id=document_id, section_id=section_id_for_legal_topic(legal_topic), legal_topic=legal_topic, indexed_chunks=indexing_result.indexed_chunks)
        finally:
            if temp_path is not None:
                temp_path.unlink(missing_ok=True)

def delete_section(*, document_id: str, section_id: str, source_directory: Path, client: OpenSearch | None=None, lock_timeout_seconds: float=DEFAULT_LOCK_TIMEOUT_SECONDS) -> AdminDocumentSectionDeleteResponse:
    """
    Permanently remove one top-level legal section - heading and all
    of its body content - from the current DOCX (mission "ORDER 8G-A",
    section 7).

    Mirrors update_effective_section's exact transaction shape: lock,
    snapshot, mutate a temp copy, reparse-validate, apply to
    OpenSearch, atomically replace the source, verify - reusing the
    same shared rollback helpers, since a delete is exactly a rename-
    to-nothing for one already-known legal_topic (no new chunks are
    ever built). Blocks deleting the document's last remaining usable
    section.
    """
    document_id = _validate_document_id(document_id)
    opensearch_client = client or get_opensearch_client()
    preliminary_metadata = _get_document_metadata(document_id=document_id, client=opensearch_client)
    country_code_for_lock = _required_string(preliminary_metadata, 'country_code')
    with country_lock(source_directory, country_code_for_lock, timeout_seconds=lock_timeout_seconds):
        document_metadata = _get_document_metadata(document_id=document_id, client=opensearch_client)
        country_code = _required_string(document_metadata, 'country_code')
        if country_code != country_code_for_lock:
            raise AdminDocumentLifecycleError("The document's country changed during the operation.")
        _ensure_no_country_conflict(country_code=country_code, client=opensearch_client, operation='section_delete')
        country = canonical_country_name(country_code)
        source_path = _resolve_current_source_path(document_metadata=document_metadata, source_directory=source_directory)
        existing_topics = _current_topics(source_path=source_path, country=country)
        matching_topic = next((topic for topic in existing_topics if section_id_for_legal_topic(topic.legal_topic) == section_id), None)
        if matching_topic is None:
            raise AdminDocumentSectionNotFoundError(document_id=document_id, section_id=section_id)
        if len(existing_topics) <= 1:
            raise AdminDocumentSectionLastRemainingError(document_id=document_id, section_id=section_id)
        legal_topic = matching_topic.legal_topic
        other_topics_before = {topic.legal_topic for topic in existing_topics if topic.legal_topic != legal_topic}
        original_bytes = source_path.read_bytes()
        try:
            pre_delete_snapshot = [item for item in _snapshot_document_chunks(client=opensearch_client, document_id=document_id) if item['_source'].get('legal_topic') == legal_topic]
        except DocumentIndexingError as error:
            raise AdminDocumentSectionUpdateFailedError(message="The section's current content could not be read from the search index before deleting.") from error
        temp_path = _make_temp_docx_path(source_path)
        try:
            try:
                remove_top_level_topic(file_path=source_path, output_path=temp_path, country=country, legal_topic=legal_topic)
                _fsync_path(temp_path)
            except Exception as error:
                raise AdminDocumentSectionUpdateFailedError(message=f'The section could not be deleted: {error}') from error
            _validate_mutated_docx(temp_path=temp_path, country=country, expected_country_code=country_code, document_metadata=document_metadata)
            try:
                reparsed_topics = locate_top_level_topics(Document(temp_path), country=country)
            except Exception as error:
                raise AdminDocumentSectionUpdateFailedError(message=f'The updated document could not be re-parsed - no change was saved: {error}') from error
            reparsed_names = {topic.legal_topic for topic in reparsed_topics}
            if legal_topic in reparsed_names:
                raise AdminDocumentSectionUpdateFailedError(message='The section is still present after deleting - no change was saved.')
            if reparsed_names != other_topics_before:
                raise AdminDocumentSectionUpdateFailedError(message='Deleting this section changed other sections in the document - no change was saved.')
            try:
                _delete_chunks_except(client=opensearch_client, filters=[{'term': {'document_id': document_id}}, {'term': {'legal_topic': legal_topic}}], keep_chunk_ids=[], context=f'document {document_id!r} section delete')
            except DocumentIndexingError as error:
                try:
                    _restore_section_snapshot(client=opensearch_client, document_id=document_id, legal_topic=legal_topic, snapshot=pre_delete_snapshot, bulk_chunk_size=DEFAULT_BULK_CHUNK_SIZE)
                except Exception as rollback_error:
                    raise AdminDocumentRollbackError('The section could not be removed from the search index, and restoring it afterward also failed - manual recovery is required.') from rollback_error
                raise AdminDocumentSectionUpdateFailedError(message='The section could not be removed from the search index - no change was saved.') from error
            try:
                os.replace(temp_path, source_path)
                _fsync_path(source_path.parent)
            except OSError as replace_error:
                _rollback_after_source_replace_failure(opensearch_client=opensearch_client, document_id=document_id, legal_topic=legal_topic, pre_mutation_snapshot=pre_delete_snapshot, replace_error=replace_error, context='section deletion')
            temp_path = None
            try:
                _verify_topic_absent(document_id=document_id, legal_topic=legal_topic, client=opensearch_client)
            except AdminDocumentSectionUpdateFailedError as verify_error:
                _rollback_after_invariant_failure(opensearch_client=opensearch_client, document_id=document_id, legal_topic=legal_topic, pre_mutation_snapshot=pre_delete_snapshot, verify_error=verify_error, context='section deletion', source_path=source_path, original_source_bytes=original_bytes)
            mark_admin_modified(source_directory, document_id)
            return AdminDocumentSectionDeleteResponse(document_id=document_id, section_id=section_id, legal_topic=legal_topic)
        finally:
            if temp_path is not None:
                temp_path.unlink(missing_ok=True)
