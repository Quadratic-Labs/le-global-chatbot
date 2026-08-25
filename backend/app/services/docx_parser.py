from __future__ import annotations

import re
import zipfile
from collections.abc import Iterator, Sequence
from dataclasses import dataclass
from html import unescape
from pathlib import Path
from typing import IO, Final, TypeAlias

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from app.core.legal_taxonomy import (
    get_canonical_legal_topic,
    is_overview_section,
)
from app.core.subsection_taxonomy import (
    get_canonical_subsection,
    get_subsection_topic_override,
)


_HEADING_STYLE_PATTERN = re.compile(
    r"^(?:heading|titre)\s*([1-9])$",
    re.IGNORECASE,
)

_TOPIC_NUMBER_PREFIX_PATTERN = re.compile(
    r"^\s*"
    r"(?:[|¦=]+\s*)?"
    r"(?:\d{1,2}|[IVX]{1,6})\s*[.)]\s*",
)

_LEADING_DECORATION_PATTERN = re.compile(
    r"^\s*[|¦=]+\s*"
)

_TRAILING_DECORATION_PATTERN = re.compile(
    r"\s*[|¦=]+\s*$"
)

_IGNORED_PARAGRAPH_TEXTS = frozenset(
    {
        "|",
        "¦",
        "=",
    }
)

_FALSE_XML_VALUES = frozenset(
    {
        "0",
        "false",
        "off",
    }
)

_MAX_CUSTOM_TOPIC_HEADING_LENGTH: Final[int] = 100

_SENTENCE_TERMINATOR_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"[.;,:]$"
)

# ORDER 8A-C: a dedicated Word paragraph style is the deterministic,
# DOCX-native marker for an ADMIN-added top-level legal topic - see
# document_mutation.py, which creates/reuses this exact style name
# whenever it inserts a new section. Recognizing it never depends on
# the document's own native heading convention (Heading 1, bold-only,
# numbered, or anything else), on the country, on document position,
# or on the visible title text - only on this one style being applied.
ADMIN_SECTION_STYLE_NAME: Final[str] = "LE Global Admin Section Heading"

BlockItem: TypeAlias = Paragraph | Table


@dataclass(frozen=True, slots=True)
class ParsedSection:
    """A structured section extracted from a DOCX document."""

    section: str
    subsection: str | None
    content: str

    # True only for a top-level legal topic recognized outside the
    # fixed LEGAL_TOPICS taxonomy (an admin-added custom section). Lets
    # the chunk builder accept it as its own legal_topic instead of
    # raising UnknownLegalTopicError, without weakening that check for
    # any other, genuinely-unrecognized heading.
    is_custom_legal_topic: bool = False


def _normalize_text(
    value: str,
) -> str:
    """Normalize whitespace while keeping readable text."""

    return " ".join(
        value
        .replace("\xa0", " ")
        .split()
    )


def _clean_structural_label(
    value: str,
) -> str:
    """Remove decorative separators surrounding a title."""

    normalized = _normalize_text(
        value
    )

    normalized = _LEADING_DECORATION_PATTERN.sub(
        "",
        normalized,
    )

    normalized = _TRAILING_DECORATION_PATTERN.sub(
        "",
        normalized,
    )

    return normalized.strip()


def _has_alnum(
    text: str,
) -> bool:
    """Return whether text contains useful alphanumeric content."""

    return any(
        character.isalnum()
        for character in text
    )


def _get_heading_level(
    paragraph: Paragraph,
) -> int | None:
    """Return the effective Word heading level."""

    style_name = (
        paragraph.style.name
        if paragraph.style is not None
        else ""
    )

    style_match = _HEADING_STYLE_PATTERN.match(
        style_name.strip()
    )

    if style_match:
        return int(
            style_match.group(1)
        )

    paragraph_properties = paragraph._p.pPr

    if (
        paragraph_properties is not None
        and paragraph_properties.outlineLvl is not None
    ):
        return (
            int(
                paragraph_properties.outlineLvl.val
            )
            + 1
        )

    return None


def _has_numbering(
    paragraph: Paragraph,
) -> bool:
    """Return whether the paragraph belongs to a Word list."""

    paragraph_properties = paragraph._p.pPr

    return (
        paragraph_properties is not None
        and paragraph_properties.numPr is not None
    )


def _has_topic_number_prefix(
    text: str,
) -> bool:
    """Return whether text begins with a legal-topic number."""

    return (
        _TOPIC_NUMBER_PREFIX_PATTERN.match(
            text
        )
        is not None
    )


def _has_explicit_bold_text(
    paragraph: Paragraph,
) -> bool:
    """Return whether visible text is explicitly formatted in bold."""

    return any(
        run.bold is True
        for run in paragraph.runs
        if run.text.strip()
    )


def _is_false_xml_boolean(
    value: str | None,
) -> bool:
    """Interpret an explicit false Word XML boolean value."""

    return (
        value is not None
        and value.lower() in _FALSE_XML_VALUES
    )


def _is_explicitly_unbolded(
    paragraph: Paragraph,
) -> bool:
    """
    Detect a paragraph whose bold formatting was explicitly disabled.

    Some source documents use heading styles for ordinary body text
    and override their inherited bold formatting.
    """

    paragraph_properties = paragraph._p.pPr

    if paragraph_properties is not None:
        run_properties = paragraph_properties.find(
            qn("w:rPr")
        )

        if run_properties is not None:
            bold_property = run_properties.find(
                qn("w:b")
            )

            if (
                bold_property is not None
                and _is_false_xml_boolean(
                    bold_property.get(
                        qn("w:val")
                    )
                )
            ):
                return True

    visible_runs = [
        run
        for run in paragraph.runs
        if run.text.strip()
    ]

    return (
        bool(visible_runs)
        and all(
            run.bold is False
            for run in visible_runs
        )
    )


def _has_main_section_signal(
    paragraph: Paragraph,
    text: str,
    heading_level: int | None,
) -> bool:
    """Return whether a legal-topic paragraph is structurally a title."""

    return any(
        (
            heading_level == 1,
            _has_numbering(
                paragraph
            ),
            _has_topic_number_prefix(
                text
            ),
            _has_explicit_bold_text(
                paragraph
            ),
        )
    )


def _get_main_legal_topic(
    paragraph: Paragraph,
    text: str,
    heading_level: int | None,
    country: str | None,
) -> str | None:
    """Return the canonical topic when a main section is detected."""

    legal_topic = get_canonical_legal_topic(
        section=text,
        country=country,
    )

    if legal_topic is None:
        return None

    if not _has_main_section_signal(
        paragraph=paragraph,
        text=text,
        heading_level=heading_level,
    ):
        return None

    return legal_topic


def is_admin_section_heading(
    paragraph: Paragraph,
) -> bool:
    """
    Return whether a paragraph carries the dedicated ADMIN-added
    top-level-section marker style (ORDER 8A-C, section 2).

    This is the sole, deterministic identity check for an admin-
    created section: it never depends on the document's own native
    heading convention, on country, on document position, or on the
    visible title text - only on this exact style name. Nothing in
    any real historical L&E document uses this project-specific style
    name, so it can never collide with genuine native content.
    """

    style = paragraph.style

    return (
        style is not None
        and style.name == ADMIN_SECTION_STYLE_NAME
    )


_TOPIC_SIGNAL_HEADING_LEVEL_1: Final[str] = "heading_level_1"
_TOPIC_SIGNAL_NUMBERING: Final[str] = "numbering"
_TOPIC_SIGNAL_PREFIX: Final[str] = "topic_number_prefix"
_TOPIC_SIGNAL_BOLD: Final[str] = "explicit_bold"


def _main_section_signals(
    paragraph: Paragraph,
    text: str,
    heading_level: int | None,
) -> frozenset[str]:
    """Return which individual main-section structural signals hold."""

    signals: set[str] = set()

    if heading_level == 1:
        signals.add(
            _TOPIC_SIGNAL_HEADING_LEVEL_1
        )

    if _has_numbering(
        paragraph
    ):
        signals.add(
            _TOPIC_SIGNAL_NUMBERING
        )

    if _has_topic_number_prefix(
        text
    ):
        signals.add(
            _TOPIC_SIGNAL_PREFIX
        )

    if _has_explicit_bold_text(
        paragraph
    ):
        signals.add(
            _TOPIC_SIGNAL_BOLD
        )

    return frozenset(
        signals
    )


def _learn_custom_topic_signal_requirement(
    document: DocxDocument,
    country: str | None,
) -> frozenset[str] | None:
    """
    Learn which structural signals this document's own confirmed legal
    topics consistently carry, so a custom (non-taxonomy) heading is
    only ever recognized against the SAME evidentiary bar - never a
    weaker one.

    Returns None (never accept a custom topic) when the document has
    no confirmed topic to learn from, or when its confirmed topics do
    not consistently use a real Word Heading 1 / outline-level-0
    paragraph. Bold text, list numbering, and textual "N." prefixes are
    also used pervasively by ordinary subsections and body emphasis in
    these documents, so on their own they are too ambiguous to promote
    an unrecognized heading to a legal topic - only Heading 1 is a
    structural marker reserved for real top-level sections. Documents
    whose own topics rely solely on those weaker signals (no Heading 1
    at all) do not support custom top-level topics.
    """

    if country is None:
        return None

    confirmed_signal_sets: list[frozenset[str]] = []

    for block_item in _iter_block_items(
        document
    ):
        if not isinstance(
            block_item,
            Paragraph,
        ):
            continue

        text = _normalize_text(
            block_item.text
        )

        if (
            not text
            or _is_ignored_text(
                text
            )
        ):
            continue

        heading_level = _get_heading_level(
            block_item
        )

        legal_topic = _get_main_legal_topic(
            paragraph=block_item,
            text=text,
            heading_level=heading_level,
            country=country,
        )

        if legal_topic is None:
            continue

        confirmed_signal_sets.append(
            _main_section_signals(
                paragraph=block_item,
                text=text,
                heading_level=heading_level,
            )
        )

    if not confirmed_signal_sets:
        return None

    required = confirmed_signal_sets[0]

    for signals in confirmed_signal_sets[1:]:
        required = required & signals

    if _TOPIC_SIGNAL_HEADING_LEVEL_1 not in required:
        return None

    return required


def _looks_like_topic_title(
    label: str,
) -> bool:
    """
    Return whether text reads like a section title rather than a
    sentence or list-item fragment.

    A real legal-topic heading is a short title case/sentence case
    phrase (e.g. "Hiring Practices", or "12. Remote Working" with the
    document's own numbering convention); it never starts with a
    lowercase word (once any leading number prefix is set aside) or
    ends with sentence punctuation the way an enumerated list item or
    clause does (e.g. "the Corporations Act;").
    """

    if not label:
        return False

    unprefixed = _TOPIC_NUMBER_PREFIX_PATTERN.sub(
        "",
        label,
    ) or label

    if not unprefixed[0].isupper():
        return False

    if _SENTENCE_TERMINATOR_PATTERN.search(
        label
    ):
        return False

    return True


def _get_custom_legal_topic(
    paragraph: Paragraph,
    text: str,
    heading_level: int | None,
    country: str | None,
    past_front_matter: bool,
    required_signals: frozenset[str] | None,
) -> str | None:
    """
    Return a custom (non-taxonomy) legal-topic title, when justified.

    An admin can add a brand-new top-level topic that has no entry in
    the fixed LEGAL_TOPICS taxonomy (e.g. "Remote Working"). Recognizing
    it generically requires evidence at least as strong as whatever
    this specific document's OWN confirmed topics already carry
    (required_signals, from _learn_custom_topic_signal_requirement) -
    a single ambiguous signal such as bold text is not enough on its
    own, since ordinary subsections also use bold text.

    Only considered once the document's own front matter has been left
    behind (its overview or first main heading has already been seen),
    so a title page or introductory heading is never mistaken for a
    legal topic.
    """

    if country is None:
        return None

    if not past_front_matter:
        return None

    if required_signals is None:
        return None

    candidate_signals = _main_section_signals(
        paragraph=paragraph,
        text=text,
        heading_level=heading_level,
    )

    if not required_signals.issubset(
        candidate_signals
    ):
        return None

    label = _clean_structural_label(
        text
    )

    # A real topic heading is a short title, not a full sentence - this
    # also matches the "reasonable length" convention new custom
    # section titles must already follow (see admin ADD validation).
    # Excludes the rare list item that happens to share the document's
    # own heading signal combination by formatting accident.
    if len(label) > _MAX_CUSTOM_TOPIC_HEADING_LENGTH:
        return None

    if not _looks_like_topic_title(
        label
    ):
        return None

    return label


def _get_subsection_topic_override(
    paragraph: Paragraph,
    text: str,
) -> str | None:
    """
    Return a one-off legal-topic override for specific known headings.

    See SUBSECTION_TOPIC_OVERRIDES in subsection_taxonomy.py for why
    this table exists: some content-topic mismatches cannot be
    detected from DOCX structure alone.
    """

    if not _has_explicit_bold_text(
        paragraph
    ):
        return None

    return get_subsection_topic_override(
        text
    )


def _is_overview_heading(
    paragraph: Paragraph,
    text: str,
    heading_level: int | None,
    country: str | None,
) -> bool:
    """Return whether the paragraph starts the document overview."""

    if not is_overview_section(
        section=text,
        country=country,
    ):
        return False

    return any(
        (
            heading_level == 1,
            _has_explicit_bold_text(
                paragraph
            ),
        )
    )


def _is_generic_main_heading(
    paragraph: Paragraph,
    heading_level: int | None,
    country: str | None,
) -> bool:
    """
    Accept Heading 1 in generic mode.

    Strict L&E parsing is activated whenever country is supplied.
    """

    if country is not None:
        return False

    if heading_level != 1:
        return False

    if _is_explicitly_unbolded(
        paragraph
    ):
        return False

    return True


def _get_subsection_label(
    paragraph: Paragraph,
    text: str,
    heading_level: int | None,
    parent_topic: str | None,
    country: str | None,
) -> str | None:
    """
    Return a subsection label when the evidence is reliable.

    Generic mode accepts normal Heading 2 paragraphs.

    Strict L&E mode requires:

    - a match in the controlled subsection taxonomy; and
    - Heading 2, Heading 4, or explicit bold formatting.

    Unknown bold paragraphs remain part of the legal content.
    """

    if _is_explicitly_unbolded(
        paragraph
    ):
        return None

    if country is None:
        if heading_level != 2:
            return None

        if _has_numbering(
            paragraph
        ):
            return None

        return _clean_structural_label(
            text
        )

    canonical_subsection = get_canonical_subsection(
        parent_topic=parent_topic,
        subsection=text,
    )

    if canonical_subsection is None:
        return None

    has_structural_signal = any(
        (
            heading_level == 2,
            heading_level == 4,
            _has_explicit_bold_text(
                paragraph
            ),
        )
    )

    if not has_structural_signal:
        return None

    return canonical_subsection


def _is_ignored_text(
    text: str,
) -> bool:
    """Return whether text is only a decorative separator."""

    return text in _IGNORED_PARAGRAPH_TEXTS


def _iter_block_items(
    document: DocxDocument,
) -> Iterator[BlockItem]:
    """Yield paragraphs and tables in their real document order."""

    for child in document.element.body.iterchildren():
        if isinstance(
            child,
            CT_P,
        ):
            yield Paragraph(
                child,
                document,
            )

        elif isinstance(
            child,
            CT_Tbl,
        ):
            yield Table(
                child,
                document,
            )


def _serialize_table(
    table: Table,
) -> str:
    """Convert a meaningful Word table to readable plain text."""

    serialized_rows: list[str] = []

    for row in table.rows:
        cells = [
            _normalize_text(
                cell.text
            )
            for cell in row.cells
        ]

        if not any(
            cells
        ):
            continue

        serialized_row = " | ".join(
            cells
        ).strip()

        if not _has_alnum(
            serialized_row
        ):
            continue

        serialized_rows.append(
            serialized_row
        )

    return "\n".join(
        serialized_rows
    )


def parse_docx_sections(
    file_path: Path,
    country: str | None = None,
) -> list[ParsedSection]:
    """
    Parse a DOCX file into structured sections.

    Strict L&E mode is used when country is provided:

    - main sections must match the legal taxonomy;
    - subsections must match the controlled subsection taxonomy;
    - Heading 2, Heading 4, and bold legacy labels are supported;
    - unknown formatted paragraphs remain content.

    Generic mode is used without country:

    - Heading 1 creates a section;
    - Heading 2 creates a subsection.
    """

    file_path = file_path.resolve()

    if not file_path.is_file():
        raise FileNotFoundError(
            f"DOCX file not found: {file_path}"
        )

    if file_path.suffix.lower() != ".docx":
        raise ValueError(
            "Unsupported document format: "
            f"{file_path.suffix}"
        )

    document = Document(
        file_path
    )

    parsed_sections: list[ParsedSection] = []

    current_section = (
        f"Employment Law Overview {country}"
        if country
        else "General"
    )

    current_legal_topic: str | None = None
    current_subsection: str | None = None
    current_is_custom_legal_topic = False

    # Becomes True once the document's own overview or first main
    # heading has been seen, so a title-page or introductory heading
    # is never mistaken for a custom (non-taxonomy) legal topic.
    past_front_matter = False

    # The structural signal combination this document's own confirmed
    # topics consistently carry (or None if there is none to learn),
    # used as the evidentiary bar a custom heading must also clear.
    required_custom_topic_signals = _learn_custom_topic_signal_requirement(
        document=document,
        country=country,
    )

    # Set while inside a one-off SUBSECTION_TOPIC_OVERRIDES block: holds
    # the (section, legal_topic) to restore as soon as the next heading
    # of any kind is found, so the override never permanently changes
    # what topic subsequent subsections of the enclosing section
    # resolve against.
    pending_topic_override: (
        tuple[str, str | None, bool] | None
    ) = None

    content_buffer: list[str] = []

    def flush_content() -> None:
        content = "\n\n".join(
            content_buffer
        ).strip()

        if content and _has_alnum(
            content
        ):
            parsed_sections.append(
                ParsedSection(
                    section=current_section,
                    subsection=current_subsection,
                    content=content,
                    is_custom_legal_topic=current_is_custom_legal_topic,
                )
            )

        content_buffer.clear()

    for block_item in _iter_block_items(
        document
    ):
        if isinstance(
            block_item,
            Paragraph,
        ):
            text = _normalize_text(
                block_item.text
            )

            if (
                not text
                or _is_ignored_text(
                    text
                )
            ):
                continue

            heading_level = _get_heading_level(
                block_item
            )

            if is_admin_section_heading(
                block_item
            ):
                flush_content()

                current_section = _clean_structural_label(
                    text
                )

                current_legal_topic = current_section
                current_subsection = None
                current_is_custom_legal_topic = True
                pending_topic_override = None
                past_front_matter = True
                continue

            topic_override = _get_subsection_topic_override(
                paragraph=block_item,
                text=text,
            )

            if topic_override is not None:
                flush_content()

                if pending_topic_override is None:
                    pending_topic_override = (
                        current_section,
                        current_legal_topic,
                        current_is_custom_legal_topic,
                    )

                current_section = _clean_structural_label(
                    text
                )

                current_legal_topic = topic_override
                current_subsection = None
                current_is_custom_legal_topic = False
                continue

            legal_topic = _get_main_legal_topic(
                paragraph=block_item,
                text=text,
                heading_level=heading_level,
                country=country,
            )

            if legal_topic is not None:
                flush_content()

                current_section = _clean_structural_label(
                    text
                )

                current_legal_topic = legal_topic
                current_subsection = None
                current_is_custom_legal_topic = False
                pending_topic_override = None
                past_front_matter = True
                continue

            if _is_overview_heading(
                paragraph=block_item,
                text=text,
                heading_level=heading_level,
                country=country,
            ):
                flush_content()

                current_section = _clean_structural_label(
                    text
                )

                current_legal_topic = None
                current_subsection = None
                current_is_custom_legal_topic = False
                pending_topic_override = None
                past_front_matter = True
                continue

            if _is_generic_main_heading(
                paragraph=block_item,
                heading_level=heading_level,
                country=country,
            ):
                flush_content()

                current_section = _clean_structural_label(
                    text
                )

                current_legal_topic = None
                current_subsection = None
                current_is_custom_legal_topic = False
                pending_topic_override = None
                past_front_matter = True
                continue

            custom_legal_topic = _get_custom_legal_topic(
                paragraph=block_item,
                text=text,
                heading_level=heading_level,
                country=country,
                past_front_matter=past_front_matter,
                required_signals=required_custom_topic_signals,
            )

            if custom_legal_topic is not None:
                flush_content()

                current_section = _clean_structural_label(
                    text
                )

                current_legal_topic = custom_legal_topic
                current_subsection = None
                current_is_custom_legal_topic = True
                pending_topic_override = None
                continue

            subsection_parent_topic = (
                pending_topic_override[1]
                if pending_topic_override is not None
                else current_legal_topic
            )

            subsection_label = _get_subsection_label(
                paragraph=block_item,
                text=text,
                heading_level=heading_level,
                parent_topic=subsection_parent_topic,
                country=country,
            )

            if subsection_label is not None:
                flush_content()

                if pending_topic_override is not None:
                    (
                        current_section,
                        current_legal_topic,
                        current_is_custom_legal_topic,
                    ) = pending_topic_override

                    pending_topic_override = None

                current_subsection = subsection_label
                continue

            content_buffer.append(
                text
            )

            continue

        table_content = _serialize_table(
            block_item
        )

        if table_content:
            content_buffer.append(
                table_content
            )

    flush_content()

    return parsed_sections


@dataclass(frozen=True, slots=True)
class TopicLocation:
    """
    One top-level legal topic's raw position within a DOCX body.

    For mutation purposes only (Edit/Add a section) - not a substitute
    for parse_docx_sections. Subsections are not tracked here: a
    mutation always replaces a topic's content in full, and the normal
    parser re-derives subsection metadata from the result afterwards.
    """

    legal_topic: str
    is_custom_legal_topic: bool
    heading_element: CT_P
    body_elements: tuple[CT_P | CT_Tbl, ...]


def locate_top_level_topics(
    document: DocxDocument,
    country: str,
) -> list[TopicLocation]:
    """
    Find each top-level legal-topic heading (canonical or custom) in an
    already-open DOCX document and the raw body elements between it
    and the next top-level heading-like paragraph (main topic,
    overview, or generic-main heading).

    Takes the same in-memory document object a caller intends to
    mutate (rather than a file path), so the returned elements belong
    to the exact tree the caller will modify and save - never a
    separate, throwaway parse of the same file.

    Reuses the same heading-classification rules as
    parse_docx_sections, so a topic's boundary here always matches
    what the parser would extract as that topic's content.
    """

    required_custom_topic_signals = _learn_custom_topic_signal_requirement(
        document=document,
        country=country,
    )

    locations: list[TopicLocation] = []

    current_heading_element: CT_P | None = None
    current_legal_topic: str | None = None
    current_is_custom = False
    current_body_elements: list[CT_P | CT_Tbl] = []
    past_front_matter = False

    def flush_topic() -> None:
        if (
            current_heading_element is not None
            and current_legal_topic is not None
        ):
            locations.append(
                TopicLocation(
                    legal_topic=current_legal_topic,
                    is_custom_legal_topic=current_is_custom,
                    heading_element=current_heading_element,
                    body_elements=tuple(
                        current_body_elements
                    ),
                )
            )

    for block_item in _iter_block_items(
        document
    ):
        if isinstance(
            block_item,
            Paragraph,
        ):
            text = _normalize_text(
                block_item.text
            )

            if (
                not text
                or _is_ignored_text(
                    text
                )
            ):
                if current_legal_topic is not None:
                    current_body_elements.append(
                        block_item._p
                    )

                continue

            heading_level = _get_heading_level(
                block_item
            )

            if is_admin_section_heading(
                block_item
            ):
                flush_topic()

                current_heading_element = block_item._p
                current_legal_topic = _clean_structural_label(
                    text
                )
                current_is_custom = True
                current_body_elements = []
                past_front_matter = True
                continue

            legal_topic = _get_main_legal_topic(
                paragraph=block_item,
                text=text,
                heading_level=heading_level,
                country=country,
            )

            if legal_topic is not None:
                flush_topic()

                current_heading_element = block_item._p
                current_legal_topic = legal_topic
                current_is_custom = False
                current_body_elements = []
                past_front_matter = True
                continue

            if _is_overview_heading(
                paragraph=block_item,
                text=text,
                heading_level=heading_level,
                country=country,
            ):
                flush_topic()

                current_heading_element = None
                current_legal_topic = None
                current_body_elements = []
                past_front_matter = True
                continue

            if _is_generic_main_heading(
                paragraph=block_item,
                heading_level=heading_level,
                country=country,
            ):
                flush_topic()

                current_heading_element = None
                current_legal_topic = None
                current_body_elements = []
                past_front_matter = True
                continue

            custom_legal_topic = _get_custom_legal_topic(
                paragraph=block_item,
                text=text,
                heading_level=heading_level,
                country=country,
                past_front_matter=past_front_matter,
                required_signals=required_custom_topic_signals,
            )

            if custom_legal_topic is not None:
                flush_topic()

                current_heading_element = block_item._p
                current_legal_topic = custom_legal_topic
                current_is_custom = True
                current_body_elements = []
                continue

            # Subsections, one-off topic overrides, and plain content
            # all stay part of the current topic's own body elements.
            if current_legal_topic is not None:
                current_body_elements.append(
                    block_item._p
                )

            continue

        if current_legal_topic is not None:
            current_body_elements.append(
                block_item._tbl
            )

    flush_topic()

    return locations


# --- Contact-card extraction -------------------------------------------
#
# Word text boxes (w:drawing > wps:txbx > w:txbxContent) are anchored
# drawings, not part of the document body's main paragraph/table flow
# that parse_docx_sections walks - so the firm/office and "CONTACT
# PERSON" cards every L&E Global document carries in its introduction
# are invisible to it. These functions are a separate, additive reader
# for that one container, used only to build one dedicated Contact
# chunk per document; they never change how normal legal sections are
# parsed or chunked.

_CONTACT_PERSON_MARKERS: Final[frozenset[str]] = frozenset(
    {
        "contact person",
        "contact persons",
    }
)

_EMAIL_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}"
)

_WEBSITE_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"(?:https?://\S+|www\.\S+)",
    re.IGNORECASE,
)

_PHONE_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"\+?\d[\d\s().-]{5,}\d"
)

_MISSING_WORD_BOUNDARY_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"(?<=[a-z0-9])(?=[A-Z])"
)

_REPEATED_COMMA_PATTERN: Final[re.Pattern[str]] = re.compile(
    r",\s*,+"
)

_TEXT_BOX_CONTENT_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"<w:txbxContent>(.*?)</w:txbxContent>",
    re.DOTALL,
)

_TEXT_BOX_PARAGRAPH_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"<w:p[ >].*?</w:p>",
    re.DOTALL,
)

# Matches either a run's text (captured in group 1) or a line break,
# in document order, so a break between two runs can be turned into a
# separator instead of silently disappearing between two <w:t> tags.
_TEXT_BOX_TOKEN_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"<w:t[^>]*>(.*?)</w:t>|<w:br\s*/>",
    re.DOTALL,
)


@dataclass(frozen=True, slots=True)
class ExtractedContact:
    """One validated L&E Global contact, extracted from a source DOCX."""

    member_firm: str | None = None
    contact_person: str | None = None
    email: str | None = None
    phone: str | None = None
    address: str | None = None
    website: str | None = None

    def has_any_field(self) -> bool:
        """Return whether at least one field was actually found."""

        return any(
            (
                self.member_firm,
                self.contact_person,
                self.email,
                self.phone,
                self.address,
                self.website,
            )
        )


def _clean_text_box_line(
    value: str,
) -> str:
    """
    Repair the two artifacts of reconstructing one visual line from
    separate XML runs and line breaks:

    - a missing space where two visually wrapped runs were joined
      without one, for example "BuildingEC3A" -> "Building EC3A",
      since Word text boxes routinely rely on the box width to wrap a
      line rather than an explicit break;
    - a doubled comma where a run's own trailing punctuation meets the
      comma this reader inserts for a line break, for example
      "Castro, " + break + "Ono" -> "Castro, , Ono" -> "Castro, Ono".
    """

    without_missing_boundaries = (
        _MISSING_WORD_BOUNDARY_PATTERN.sub(
            " ",
            value,
        )
    )

    return _REPEATED_COMMA_PATTERN.sub(
        ",",
        without_missing_boundaries,
    )


def extract_text_box_blocks(
    file_path: Path,
) -> list[list[str]]:
    """
    Extract every text box's non-empty paragraph lines, in order.

    Compatibility fallbacks (mc:Fallback) duplicate the same text box
    content for older Word versions, so a block identical to the one
    immediately before it is treated as that duplicate and dropped.
    """

    try:
        with zipfile.ZipFile(file_path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode(
                "utf-8",
                errors="ignore",
            )

    except (
        KeyError,
        OSError,
        zipfile.BadZipFile,
    ):
        return []

    blocks: list[list[str]] = []
    previous_lines: list[str] | None = None

    for raw_block in _TEXT_BOX_CONTENT_PATTERN.findall(
        document_xml
    ):
        lines: list[str] = []

        for paragraph_xml in _TEXT_BOX_PARAGRAPH_PATTERN.findall(
            raw_block
        ):
            tokens = [
                match.group(1)
                if match.group(1) is not None
                else ", "
                for match in _TEXT_BOX_TOKEN_PATTERN.finditer(
                    paragraph_xml
                )
            ]

            line = _clean_text_box_line(
                _normalize_text(
                    unescape(
                        "".join(tokens)
                    )
                )
            )

            if line:
                lines.append(line)

        if lines and lines == previous_lines:
            continue

        if lines:
            blocks.append(lines)
            previous_lines = lines

    return blocks


def _select_contact_phone(
    block: Sequence[str],
) -> str | None:
    """
    Select the most plausible telephone number from a firm card.

    A postal code can also satisfy the deliberately permissive
    _PHONE_PATTERN. Searching the whole joined block and accepting
    its first match therefore makes an address such as

        100-0011 Tokyo, +81 355 012 111

    incorrectly select 100-0011 as the telephone.

    Evaluate every match independently instead. Prefer candidates
    with at least eight digits, then explicit international numbers,
    then a candidate occupying its whole visual line, and finally
    the candidate containing the most digits.

    If only a shorter candidate exists it is still preserved, so this
    remains backwards-compatible with unusual short telephone formats.
    """

    candidates: list[
        tuple[
            tuple[int, int, int, int],
            str,
        ]
    ] = []

    for line in block[1:]:
        normalized_line = _normalize_text(line)

        for match in _PHONE_PATTERN.finditer(line):
            candidate = match.group(0).strip()

            if not candidate:
                continue

            digit_count = sum(
                character.isdigit()
                for character in candidate
            )

            score = (
                int(digit_count >= 8),
                int(candidate.startswith("+")),
                int(
                    _normalize_text(candidate)
                    == normalized_line
                ),
                digit_count,
            )

            candidates.append(
                (
                    score,
                    candidate,
                )
            )

    if not candidates:
        return None

    return max(
        candidates,
        key=lambda item: item[0],
    )[1]


def _remove_contact_phone_from_address_line(
    line: str,
    phone: str | None,
) -> str:
    """
    Remove the selected telephone from an address line while keeping
    the rest of that line.

    Some DOCX contact cards visually place the telephone at the end
    of the same text-box paragraph as the postal address.
    """

    if not phone or phone not in line:
        return line

    cleaned = line.replace(
        phone,
        "",
        1,
    )

    cleaned = _normalize_text(cleaned)
    cleaned = _REPEATED_COMMA_PATTERN.sub(
        ",",
        cleaned,
    )

    return cleaned.strip(" ,")



def parse_contact_blocks(
    blocks: Sequence[Sequence[str]],
    country: str | None = None,
) -> list[ExtractedContact]:
    """
    Classify already-extracted text-box blocks into contact entries.

    Pure text-processing function - takes plain lines, not a DOCX
    file - so it can be exercised directly with a synthetic block
    structure. Two block shapes are recognized generically, without
    any per-document or per-country assumption:

    - a "CONTACT PERSON" (or "CONTACT PERSONS") block: a line matching
      that marker, followed by a name and one or more emails;
    - a firm/office block: any other block containing an email or a
      phone-like pattern (a bare website mention alone is not enough,
      since a generic site link can appear elsewhere in a document
      without being part of a contact card). Its first line is taken
      as the member firm name; any other line becomes the address,
      except one that repeats the supplied country name, since that
      is expected to come from validated document metadata instead.

    Firm/office blocks and CONTACT PERSON blocks are collected
    separately, in document order, then paired position-by-position:
    the first firm block with the first CONTACT PERSON block found
    anywhere in the document, and so on - since different source
    documents lay these two cards out in different relative order. A
    document with several such pairs yields several contacts, in
    order. A block of either kind with no counterpart is still
    reported using only the fields it actually has - never inventing
    the other side. Exact duplicate entries (the same document's
    compatibility fallback markup surfacing as a second, near-identical
    block) are collapsed to one.
    """

    normalized_country = (
        _normalize_text(
            country
        ).casefold()
        if country
        else None
    )

    firm_entries: list[dict[str, str | None]] = []
    person_entries: list[
        tuple[str | None, str | None]
    ] = []

    for block in blocks:
        marker_index = next(
            (
                index
                for index, line in enumerate(block)
                if line.strip().casefold()
                in _CONTACT_PERSON_MARKERS
            ),
            None,
        )

        if marker_index is not None:
            emails: list[str] = []
            name_lines: list[str] = []

            for line in block[marker_index + 1:]:
                email_match = _EMAIL_PATTERN.search(
                    line
                )

                if email_match:
                    emails.append(
                        email_match.group(0)
                    )
                    continue

                name_lines.append(
                    line
                )

            person_entries.append(
                (
                    (
                        name_lines[0]
                        if name_lines
                        else None
                    ),
                    (
                        ", ".join(emails)
                        if emails
                        else None
                    ),
                )
            )

            continue

        joined = " ".join(
            block
        )

        phone = _select_contact_phone(
            block
        )

        website_match = _WEBSITE_PATTERN.search(
            joined
        )

        if not (
            _EMAIL_PATTERN.search(joined)
            or phone
        ):
            continue

        address_lines: list[str] = []

        for line in block[1:]:
            if (
                website_match
                and line == website_match.group(0)
            ):
                continue

            if (
                normalized_country is not None
                and line.casefold()
                == normalized_country
            ):
                continue

            address_line = (
                _remove_contact_phone_from_address_line(
                    line,
                    phone,
                )
            )

            if address_line:
                address_lines.append(
                    address_line
                )

        firm_entries.append(
            {
                "member_firm": (
                    block[0]
                    if block
                    else None
                ),
                "phone": phone,
                "website": (
                    website_match.group(0)
                    if website_match
                    else None
                ),
                "address": (
                    ", ".join(
                        address_lines
                    )
                    if address_lines
                    else None
                ),
            }
        )

    # Collapse near-duplicate compatibility blocks that the raw-line
    # dedup in extract_text_box_blocks missed because they differ in
    # some incidental XML detail (for example two different internal
    # hyperlink relationship ids for the same displayed email), so
    # they do not each consume a slot in the position-based pairing
    # below.
    deduplicated_person_entries: list[
        tuple[str | None, str | None]
    ] = []
    seen_person_keys: set[
        tuple[str | None, str | None]
    ] = set()

    for person_entry in person_entries:
        if person_entry in seen_person_keys:
            continue

        seen_person_keys.add(
            person_entry
        )

        deduplicated_person_entries.append(
            person_entry
        )

    person_entries = deduplicated_person_entries

    deduplicated_firm_entries: list[
        dict[str, str | None]
    ] = []
    seen_firm_keys: set[
        tuple[str | None, ...]
    ] = set()

    for firm_entry in firm_entries:
        firm_key = tuple(
            firm_entry.get(field_name)
            for field_name in (
                "member_firm",
                "phone",
                "website",
                "address",
            )
        )

        if firm_key in seen_firm_keys:
            continue

        seen_firm_keys.add(
            firm_key
        )

        deduplicated_firm_entries.append(
            firm_entry
        )

    firm_entries = deduplicated_firm_entries

    pair_count = max(
        len(firm_entries),
        len(person_entries),
    )

    contacts: list[ExtractedContact] = []
    seen_contacts: set[
        tuple[str | None, str | None, str | None]
    ] = set()

    for index in range(pair_count):
        firm_fields = (
            firm_entries[index]
            if index < len(firm_entries)
            else {}
        )

        contact_person, email = (
            person_entries[index]
            if index < len(person_entries)
            else (None, None)
        )

        contact = ExtractedContact(
            member_firm=firm_fields.get(
                "member_firm"
            ),
            contact_person=contact_person,
            email=email,
            phone=firm_fields.get(
                "phone"
            ),
            address=firm_fields.get(
                "address"
            ),
            website=firm_fields.get(
                "website"
            ),
        )

        if not contact.has_any_field():
            continue

        dedup_key = (
            contact.member_firm,
            contact.contact_person,
            contact.email,
        )

        if dedup_key in seen_contacts:
            continue

        seen_contacts.add(
            dedup_key
        )

        contacts.append(
            contact
        )

    return contacts



def _extract_standalone_contact_emails(
    line: str,
) -> list[str]:
    """
    Return emails only when the visual paragraph consists solely of
    one or more email addresses.

    This deliberately rejects emails embedded in ordinary prose, such
    as a project-level L&E Global POC sentence.
    """

    normalized = _normalize_text(
        line
    )

    if not normalized:
        return []

    candidate = re.sub(
        r"(?i)\bmailto:",
        "",
        normalized,
    )

    matches = [
        match.group(0)
        for match in _EMAIL_PATTERN.finditer(
            candidate
        )
    ]

    if not matches:
        return []

    remainder = _EMAIL_PATTERN.sub(
        "",
        candidate,
    )

    remainder = re.sub(
        r"[\s,;|/]+",
        "",
        remainder,
    )

    if remainder:
        return []

    return matches


def _extract_standalone_contact_phone(
    line: str,
) -> str | None:
    """
    Return a telephone only when the whole paragraph is phone-like.

    Legal prose containing a number therefore cannot become contact
    data merely because the permissive phone regex finds a match.
    """

    normalized = _normalize_text(
        line
    )

    if not normalized:
        return None

    candidates: list[
        tuple[int, int, str]
    ] = []

    for match in _PHONE_PATTERN.finditer(
        normalized
    ):
        candidate = match.group(0).strip()

        digit_count = sum(
            character.isdigit()
            for character in candidate
        )

        if digit_count < 8:
            continue

        remainder = (
            normalized[:match.start()]
            + normalized[match.end():]
        )

        remainder = re.sub(
            r"[\s,;:|/()\-]+",
            "",
            remainder,
        )

        if remainder:
            continue

        candidates.append(
            (
                int(candidate.startswith("+")),
                digit_count,
                candidate,
            )
        )

    if not candidates:
        return None

    return max(
        candidates,
        key=lambda item: (
            item[0],
            item[1],
        ),
    )[2]


def _plain_contact_role_and_firm(
    line: str,
) -> tuple[str, str] | None:
    """
    Parse a high-confidence role/firm paragraph.

    Examples:
        Partners, Flichy Grangé Avocats
        Partner, Example Employment Law

    A firm name on its own is intentionally insufficient.
    """

    normalized = _normalize_text(
        line
    )

    if not normalized:
        return None

    match = re.match(
        (
            r"^(?P<role>"
            r"(?:managing\s+|senior\s+)?partners?"
            r"|attorneys?"
            r"|lawyers?"
            r"|counsel"
            r"|associates?"
            r")"
            r"\s*[,:\-]\s*"
            r"(?P<firm>.+?)\s*$"
        ),
        normalized,
        flags=re.IGNORECASE,
    )

    if match is None:
        return None

    firm = _normalize_text(
        match.group("firm")
    )

    if not firm:
        return None

    return (
        _normalize_text(
            match.group("role")
        ),
        firm,
    )


def _looks_like_plain_contact_person(
    line: str,
) -> bool:
    """
    Conservative validation for the person/persons paragraph that
    precedes a role + firm line.
    """

    normalized = _normalize_text(
        line
    )

    if not normalized:
        return False

    if len(normalized) > 120:
        return False

    if (
        _EMAIL_PATTERN.search(normalized)
        or _WEBSITE_PATTERN.search(normalized)
        or _PHONE_PATTERN.search(normalized)
    ):
        return False

    lowered = normalized.casefold()

    if lowered.startswith(
        (
            "for ",
            "please ",
            "this ",
            "the ",
            "your ",
            "contact ",
            "www.",
            "http://",
            "https://",
        )
    ):
        return False

    words = normalized.split()

    if not 2 <= len(words) <= 12:
        return False

    alphabetic_words = [
        word
        for word in words
        if any(
            character.isalpha()
            for character in word
        )
    ]

    return len(
        alphabetic_words
    ) >= 2


def _iter_plain_paragraph_contact_blocks(
    document,
) -> Iterator[tuple[int, int, ExtractedContact]]:
    """
    Shared core for _extract_plain_paragraph_contacts (text extraction)
    and find_plain_paragraph_contact_block_bounds
    (contact_document_area.py's canonicalizer, which needs to know
    exactly which body paragraphs to remove and where to insert the
    canonical table in their place - a document whose legacy member-
    firm contact lives in ordinary body paragraphs, like France's, has
    no floating shape for the existing carrier-removal logic to find).

    This is intentionally NOT a general scan for every email, phone,
    URL, or firm name in the document.

    A contact requires a local, coherent structure:

        person/persons
        role, member firm
        standalone email(s)
        optional standalone phone

    That high-confidence shape prevents ordinary legal references,
    government URLs, project-level POCs, and a bare authoring-firm
    name from becoming contact records.

    Yields (first_paragraph_index, last_paragraph_index,
    ExtractedContact) per detected block, in document order -
    first_paragraph_index is the person-line paragraph's own index
    into document.paragraphs, last_paragraph_index is the phone line's
    index if one was found, else the last email line's index.
    """

    indexed_lines = [
        (original_index, normalized)
        for original_index, paragraph in enumerate(document.paragraphs)
        if (normalized := _normalize_text(paragraph.text))
    ]
    lines = [text for _, text in indexed_lines]

    seen: set[tuple[str | None, str | None, str | None]] = set()

    index = 0

    while index < len(lines):
        emails = _extract_standalone_contact_emails(lines[index])

        if not emails:
            index += 1
            continue

        first_email_index = index
        all_emails = list(emails)

        index += 1

        while index < len(lines):
            additional = _extract_standalone_contact_emails(lines[index])

            if not additional:
                break

            all_emails.extend(additional)

            index += 1

        last_email_index = index - 1

        if first_email_index < 2:
            continue

        role_line = lines[first_email_index - 1]
        person_line = lines[first_email_index - 2]

        role_and_firm = _plain_contact_role_and_firm(role_line)

        if role_and_firm is None:
            continue

        if not _looks_like_plain_contact_person(person_line):
            continue

        # Explicit defence against project-level/global POC regions.
        nearby_start = max(0, first_email_index - 4)

        nearby = " ".join(
            lines[nearby_start:first_email_index + 1]
        ).casefold()

        if "l&e global poc" in nearby or "l & e global poc" in nearby:
            continue

        _, member_firm = role_and_firm

        phone = None
        phone_index: int | None = None

        if index < len(lines):
            phone = _extract_standalone_contact_phone(lines[index])

            if phone is not None:
                phone_index = index

        contact = ExtractedContact(
            member_firm=member_firm,
            contact_person=person_line,
            email=", ".join(all_emails),
            phone=phone,
            address=None,
            website=None,
        )

        key = (
            contact.member_firm,
            contact.contact_person,
            contact.email,
        )

        if key in seen:
            continue

        seen.add(key)

        first_paragraph_index = indexed_lines[first_email_index - 2][0]
        last_line_index = (
            phone_index if phone_index is not None else last_email_index
        )
        last_paragraph_index = indexed_lines[last_line_index][0]

        yield first_paragraph_index, last_paragraph_index, contact


def _extract_plain_paragraph_contacts(
    file_path: Path,
) -> list[ExtractedContact]:
    """
    Conservative fallback for organically-uploaded DOCX files whose
    member-firm contact is stored in ordinary body paragraphs rather
    than Word text boxes - see _iter_plain_paragraph_contact_blocks for
    the detection rules.

    Existing text-box contacts never reach this function:
    extract_contacts_from_docx() returns those first.
    """

    try:
        from docx import Document as WordDocument

        document = WordDocument(
            file_path
        )

    except Exception:
        # Contact extraction is optional metadata. Preserve the legacy
        # behaviour of returning no contacts when this fallback cannot
        # read the DOCX, rather than changing the document's legal
        # parsing/error boundary.
        return []

    return [
        contact
        for _, _, contact in _iter_plain_paragraph_contact_blocks(document)
    ]


def find_plain_paragraph_contact_block_bounds(
    document,
) -> tuple[int, int] | None:
    """
    The (first_paragraph_index, last_paragraph_index) body-paragraph
    bounds of the FIRST plain-paragraph legacy contact block this
    document's own organic content contains (see
    _iter_plain_paragraph_contact_blocks) - both inclusive, indexing
    into document.paragraphs. None if this document has no such block
    at all (e.g. its legacy contact area is a floating shape instead,
    or it has no legacy contact area at all).
    """

    for first_index, last_index, _ in _iter_plain_paragraph_contact_blocks(
        document
    ):
        return (first_index, last_index)

    return None


_PERSON_SPLIT_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"\s*,\s*|\s+and\s+", re.IGNORECASE
)


def split_combined_legacy_contact(
    contact: ExtractedContact,
) -> list[ExtractedContact] | None:
    """
    Split a legacy contact whose contact_person names multiple people
    ("Caroline Scherrmann and Florence Bacquet") into one
    ExtractedContact per person, mapped 1:1 in source order onto that
    same contact's own comma-joined email addresses - ONLY when doing
    so is unambiguous: the number of name segments must exactly equal
    the number of email addresses (at least two of each), and every
    resulting name segment must itself independently look like a real
    person's name. member_firm/phone/address/website are copied to
    every split contact unchanged - this only ever separates the
    person/email pairing, never the shared firm-side fields.

    Returns None (never guesses) whenever contact_person or email is
    missing, the counts do not match, or any name segment fails the
    person-name check - a firm name that happens to contain "and"
    (e.g. "Smith and Jones LLP") is never at risk here regardless,
    since only contact_person is ever split, never member_firm. In
    every None case the caller must keep the original combined
    representation rather than fabricate an incorrect split.
    """

    if not contact.contact_person or not contact.email:
        return None

    name_segments = [
        segment.strip()
        for segment in _PERSON_SPLIT_PATTERN.split(contact.contact_person)
        if segment.strip()
    ]

    email_segments = [
        segment.strip()
        for segment in contact.email.split(",")
        if segment.strip()
    ]

    if len(name_segments) < 2 or len(name_segments) != len(email_segments):
        return None

    if not all(
        _looks_like_plain_contact_person(name) for name in name_segments
    ):
        return None

    if not all(
        _EMAIL_PATTERN.fullmatch(email) for email in email_segments
    ):
        return None

    return [
        ExtractedContact(
            member_firm=contact.member_firm,
            contact_person=name,
            email=email,
            phone=contact.phone,
            address=contact.address,
            website=contact.website,
        )
        for name, email in zip(name_segments, email_segments)
    ]


CONTACT_TABLE_HIDDEN_MARKER: Final[str] = "LE-GLOBAL-CONTACT-TABLE-V1"

# One hidden tag per left-cell field, written as its own zero-size
# hidden run immediately before that field's own visible run, in the
# SAME paragraph (mission "FERMER LE CRUD CONTACTS AVANT STREAMING").
#
# The canonical table is a format WE OWN - both the writer
# (contact_document_area.py's _fill_text_cell) and this reader - so its
# own round-trip has no reason to depend on content heuristics at all.
# Before this, member_firm/address/phone/website were told apart by
# which line SCORED strongest for phone-shapedness/URL-shapedness
# (_classify_canonical_firm_lines below): correct for every real
# corpus document tested, but not deterministic in the general case -
# a contact whose address happens to contain a stronger phone-shaped
# substring than its own actual phone value (confirmed directly: an
# adversarial add/edit round-trip for Brazil/Romania/Singapore-shaped
# data swapped address and phone, because the address's embedded
# "+55 115 086 5000" outscored the dedicated phone line "04038-904")
# could still be misclassified. A canonical table WE write can instead
# just say what each line is, unambiguously, and never need to guess.
#
# Legacy heuristic parsing (organically-authored documents this system
# never wrote) is entirely unaffected - the tags below are only ever
# looked for INSIDE a table already identified by
# CONTACT_TABLE_HIDDEN_MARKER, never in ordinary body paragraphs or
# floating text boxes. A canonical table already written by an earlier
# version of this code (before these tags existed) has none of them -
# _extract_canonical_table_contacts falls back to the same content-
# based classification for that row, so an already-canonicalized real
# document is never broken by this change.
#
# contact_person/email (the RIGHT cell) were originally left untagged,
# still read via _EMAIL_PATTERN.findall() - reasonable at the time
# since they were never the source of an observed field-shift bug, but
# not actually deterministic: _EMAIL_PATTERN's local-part character
# class is ASCII-only ([A-Za-z0-9._%+-]), so a real, legitimate email
# with an accented local part (confirmed directly: "françois@...")
# gets silently truncated to whatever ASCII suffix follows the
# accented character, corrupting both email AND (since the truncated
# prefix then fails to match as an email at all) misclassifying it as
# a second contact_person line instead. Tagging these two fields the
# same way as the LEFT cell removes the last content-dependent field
# in a canonical row.
_FIELD_TAG_MEMBER_FIRM: Final[str] = "LE-GLOBAL-FIELD-FIRM-V1:"
_FIELD_TAG_ADDRESS: Final[str] = "LE-GLOBAL-FIELD-ADDRESS-V1:"
_FIELD_TAG_PHONE: Final[str] = "LE-GLOBAL-FIELD-PHONE-V1:"
_FIELD_TAG_WEBSITE: Final[str] = "LE-GLOBAL-FIELD-WEBSITE-V1:"
_FIELD_TAG_CONTACT_PERSON: Final[str] = "LE-GLOBAL-FIELD-PERSON-V1:"
_FIELD_TAG_EMAIL: Final[str] = "LE-GLOBAL-FIELD-EMAIL-V1:"

_FIELD_TAGS_BY_NAME: Final[dict[str, str]] = {
    "member_firm": _FIELD_TAG_MEMBER_FIRM,
    "address": _FIELD_TAG_ADDRESS,
    "phone": _FIELD_TAG_PHONE,
    "website": _FIELD_TAG_WEBSITE,
    "contact_person": _FIELD_TAG_CONTACT_PERSON,
    "email": _FIELD_TAG_EMAIL,
}
_FIELD_NAMES_BY_TAG: Final[dict[str, str]] = {
    tag: name for name, tag in _FIELD_TAGS_BY_NAME.items()
}


def _extract_tagged_canonical_fields(
    lines: Sequence[str],
) -> dict[str, str] | None:
    """
    Deterministically recover this canonical row's fields from one
    cell's lines when EVERY line carries one of the hidden field tags
    above (a table written by the current writer) - used for both the
    left cell (member_firm/address/phone/website) and the right cell
    (contact_person/email). Returns None - never a partial result -
    the instant any line lacks a recognized tag, so the caller falls
    back to content-based classification for that cell as a whole (an
    older canonical table, written before these tags existed).
    """

    fields: dict[str, str] = {}

    for line in lines:
        matched_tag = next(
            (tag for tag in _FIELD_NAMES_BY_TAG if line.startswith(tag)),
            None,
        )

        if matched_tag is None:
            return None

        field_name = _FIELD_NAMES_BY_TAG[matched_tag]
        fields[field_name] = _normalize_text(line[len(matched_tag):])

    return fields


def _score_phone_candidate(line: str) -> tuple[int, int, int, int]:
    """The best _PHONE_PATTERN match score for one line, using the
    same scoring shape as _select_contact_phone - reused here to
    identify WHICH line of a canonical table's firm cell is the phone
    line, never to extract a phone substring out of it. A line with no
    match at all scores the sentinel (0, 0, 0, 0)."""

    normalized_line = _normalize_text(line)
    best = (0, 0, 0, 0)

    for match in _PHONE_PATTERN.finditer(line):
        candidate = match.group(0).strip()

        if not candidate:
            continue

        digit_count = sum(character.isdigit() for character in candidate)

        score = (
            int(digit_count >= 8),
            int(candidate.startswith("+")),
            int(_normalize_text(candidate) == normalized_line),
            digit_count,
        )

        if score > best:
            best = score

    return best


def _classify_canonical_firm_lines(
    firm_lines: Sequence[str],
) -> tuple[str | None, str | None, str | None, list[int]]:
    """
    Recover member_firm/phone/website from one canonical table row's
    left-cell lines BY CONTENT, not by fixed position - the writer
    (_fill_text_cell) always emits member_firm/address/phone/website
    in that order but SKIPS whichever fields are empty, so a contact
    with no member_firm at all has its phone (or whichever field comes
    next) land on line 0. Reading line 0 as "always the firm name"
    would misclassify that value.

    - phone: the line with the strongest _PHONE_PATTERN signal, if
      any, compared across EVERY line (including line 0). Its ENTIRE
      text becomes the phone value - the writer never shares a line
      between two fields, so any trailing annotation an Admin typed as
      part of the phone value (e.g. "+1 555 0100 (mobile)") is
      preserved verbatim rather than leaking into address.
    - website: the one remaining line that is, IN ITS ENTIRETY, just a
      URL - exactly how the writer places contact.website alone on its
      own paragraph. An address sentence that merely MENTIONS a URL
      (e.g. "...please see www.example.com.") keeps that mention as
      ordinary prose, untouched: fullmatch fails for a line that has
      anything else on it, by construction.
    - member_firm: line 0, unless line 0 was itself claimed above as
      the phone or website line (member_firm was empty and that
      field's value landed there instead).

    Returns (member_firm, phone, website, remaining_line_indices) -
    the last being every index not claimed by any of the three, for
    the caller to treat as address candidates.
    """

    phone_index: int | None = None
    best_score = (0, 0, 0, 0)

    for index, line in enumerate(firm_lines):
        score = _score_phone_candidate(line)

        if score > best_score:
            best_score = score
            phone_index = index

    phone = firm_lines[phone_index] if phone_index is not None else None

    website_index: int | None = None
    website: str | None = None

    for index, line in enumerate(firm_lines):
        if index == phone_index:
            continue

        match = _WEBSITE_PATTERN.fullmatch(line)

        if match:
            website_index = index
            website = match.group(0)
            break

    member_firm_index: int | None = 0 if firm_lines else None

    if member_firm_index in (phone_index, website_index):
        member_firm_index = None

    member_firm = (
        firm_lines[member_firm_index]
        if member_firm_index is not None
        else None
    )

    remaining_indices = [
        index
        for index in range(len(firm_lines))
        if index not in (member_firm_index, phone_index, website_index)
    ]

    return member_firm, phone, website, remaining_indices


def _extract_canonical_table_contacts(
    file_path: Path,
    country: str | None = None,
) -> list[ExtractedContact] | None:
    """
    Read the Admin-managed canonical contact table - a standard,
    borderless, in-flow Word table this system itself writes into the
    persisted source (one row per contact: a LEFT cell with that
    contact's own member firm/address/phone/website, a RIGHT cell with
    its own inline photo followed by its own CONTACT PERSON/name/email)
    - identified unambiguously by a hidden marker paragraph in the
    first row's left cell, never by guessing which table in the
    document is the contact one.

    Returns None (never an empty list) when no such table is found, so
    extract_contacts_from_docx() falls back to the legacy text-box
    parser - an empty list is reserved for "the table is present and
    genuinely lists zero contacts".

    Each row's two cells are classified DIRECTLY using the same field
    rules parse_contact_blocks() applies to the legacy text-box format
    (member firm/address/phone/website in one cell, the CONTACT PERSON
    marker/name/email in the other) - but per-row, never through parse_
    contact_blocks() itself: that function collects firm and person
    blocks into two SEPARATE lists and deduplicates each before pairing
    them back up position-by-position, a step designed for the legacy
    format's own Choice/Fallback VML duplication artifact. Two DIFFERENT
    canonical-table rows that happen to share the exact same contact
    person and email (the "duplicates explicitly allowed" contract
    add_contact() itself documents) would otherwise collapse into one
    during that dedup, silently losing a genuine contact. Each row's
    own firm+person lines are already correctly grouped by construction
    here, so no cross-row collection, dedup, or pairing is needed at
    all.
    """

    from docx import Document as WordDocument

    try:
        document = WordDocument(file_path)
    except Exception:
        return None

    marker_table = None

    for table in document.tables:
        if not table.rows:
            continue

        first_cell_text = table.rows[0].cells[0].text

        if CONTACT_TABLE_HIDDEN_MARKER in first_cell_text:
            marker_table = table
            break

    if marker_table is None:
        return None

    normalized_country = (
        _normalize_text(country).casefold() if country else None
    )

    contacts: list[ExtractedContact] = []

    for row in marker_table.rows:
        firm_lines = [
            _normalize_text(paragraph.text)
            for paragraph in row.cells[0].paragraphs
            if _normalize_text(paragraph.text)
            and CONTACT_TABLE_HIDDEN_MARKER not in paragraph.text
        ]
        person_lines = [
            _normalize_text(paragraph.text)
            for paragraph in row.cells[1].paragraphs
            if _normalize_text(paragraph.text)
            and paragraph.text.strip().casefold()
            not in _CONTACT_PERSON_MARKERS
        ]

        if not firm_lines and not person_lines:
            continue

        tagged_fields = _extract_tagged_canonical_fields(firm_lines)

        if tagged_fields is not None:
            # Every line in this row carries its own hidden field tag
            # (written by the current code) - no content heuristic
            # needed or consulted at all, so a value that happens to
            # look like a different field (a postal code that looks
            # phone-shaped, a phone number embedded in an address) can
            # never be misclassified.
            member_firm = tagged_fields.get("member_firm")
            phone = tagged_fields.get("phone")
            website = tagged_fields.get("website")
            address_lines = (
                [tagged_fields["address"]] if "address" in tagged_fields else []
            )
        else:
            # No tags at all - a canonical table written before this
            # mechanism existed (an already-canonicalized real
            # document, e.g. from a live Admin action predating this
            # fix). Fall back to the same content-based classification
            # this reader has always used, so that table keeps reading
            # exactly as it did before.
            member_firm, phone, website, address_line_indices = (
                _classify_canonical_firm_lines(firm_lines)
            )

            address_lines = []

            for index in address_line_indices:
                line = firm_lines[index]

                if (
                    normalized_country is not None
                    and line.casefold() == normalized_country
                ):
                    continue

                address_lines.append(line)

        tagged_person_fields = _extract_tagged_canonical_fields(
            person_lines
        )

        if tagged_person_fields is not None:
            # Every line in this cell carries its own hidden field tag
            # - no content heuristic (nor _EMAIL_PATTERN, which is
            # ASCII-local-part-only and can corrupt an accented email
            # such as "françois@...") is consulted at all.
            contact_person = tagged_person_fields.get("contact_person")
            email = tagged_person_fields.get("email")
        else:
            # No tags at all - a canonical table written before this
            # mechanism existed. Fall back to the same content-based
            # classification this reader has always used.
            emails: list[str] = []
            name_lines: list[str] = []

            for line in person_lines:
                # A contact can legitimately have more than one email
                # (e.g. two named individuals sharing a firm entry),
                # and the writer side (_fill_photo_and_person_cell)
                # renders them all as ONE comma-joined line, exactly
                # mirroring ContactState's own single `email` string
                # field - so a single line can itself contain multiple
                # addresses. findall() (not search()) is required to
                # recover all of them, in the order they were written,
                # or a round-trip silently drops every email after the
                # first (the real France incident:
                # "scherrmann@flichy.com, bacquet@flichy.com" on one
                # line lost the second address and the canonical
                # rebuild's own round-trip check correctly rejected the
                # mutation rather than persist it).
                line_emails = _EMAIL_PATTERN.findall(line)

                if line_emails:
                    emails.extend(line_emails)
                else:
                    name_lines.append(line)

            contact_person = name_lines[0] if name_lines else None
            email = ", ".join(emails) if emails else None

        contact = ExtractedContact(
            member_firm=member_firm,
            contact_person=contact_person,
            email=email,
            phone=phone,
            address=", ".join(address_lines) if address_lines else None,
            website=website,
        )

        if contact.has_any_field():
            contacts.append(contact)

    return contacts


def extract_contacts_from_docx(
    file_path: Path,
    country: str | None = None,
) -> list[ExtractedContact]:
    """
    Extract every validated contact card from one source DOCX.

    Checks for the canonical Admin-managed table first (present only
    on a document this mission's mechanism has itself rebuilt), then
    the deterministic hidden-marker block (mission "ORDER 8G-B2.1") -
    present only on a DOCX this system itself materialized for
    Download - and falls back to the ordinary text-box-based legacy
    parser for every real, organically-uploaded document, where
    neither marker is ever present.
    """

    canonical_table_contacts = _extract_canonical_table_contacts(
        file_path, country=country
    )

    if canonical_table_contacts is not None:
        return canonical_table_contacts

    deterministic_contacts = extract_deterministic_contact_blocks(
        file_path
    )

    if deterministic_contacts is not None:
        return deterministic_contacts

    legacy_contacts = parse_contact_blocks(
        extract_text_box_blocks(
            file_path
        ),
        country=country,
    )

    if legacy_contacts:
        return legacy_contacts

    return _extract_plain_paragraph_contacts(
        file_path
    )


def build_contact_chunk_content(
    contacts: Sequence[ExtractedContact],
) -> str:
    """Build the structured text of one Contact-subsection chunk."""

    entries: list[str] = []

    for contact in contacts:
        lines: list[str] = []

        if contact.member_firm:
            lines.append(
                f"Member firm: {contact.member_firm}"
            )

        if contact.contact_person:
            lines.append(
                f"Contact person: {contact.contact_person}"
            )

        if contact.email:
            lines.append(
                f"Email: {contact.email}"
            )

        if contact.phone:
            lines.append(
                f"Phone: {contact.phone}"
            )

        if contact.address:
            lines.append(
                f"Address: {contact.address}"
            )

        if contact.website:
            lines.append(
                f"Website: {contact.website}"
            )

        if lines:
            entries.append(
                "\n".join(
                    lines
                )
            )

    return "\n\n".join(
        entries
    )


# --- Deterministic Contact block (mission "ORDER 8G-B2.1" family) ----
#
# A Download of the effective (current) DOCX materializes structured
# Contact state INSIDE the document's own real Contact text box - the
# same visual card a human already looks at (mission "ORDER 8G-B2.1R2"
# - two earlier attempts placed the content correctly for
# extract_contacts_from_docx() but not for a human eye: v0.8.1 appended
# it at the document's end, v0.8.2 moved it to plain paragraphs near
# the title but outside the actual visual card). No visible technical
# marker is written anywhere in the document; instead, one hidden
# (w:vanish) marker paragraph is the first line inside the box - never
# rendered by Word, but still literally present in <w:t> text for this
# function's own regex-based extraction to find. Deliberately
# distinctive so it can never coincidentally match real, organically
# authored legal document text.

CONTACT_BOX_HIDDEN_MARKER: Final[str] = (
    "LE-GLOBAL-CONTACT-BLOCK-V3"
)

# Closes a hidden-marker block inserted directly among ordinary body
# paragraphs (a document with no pre-existing Contact text box at all
# - see app.services.document_contact_materializer's in-flow-paragraph
# fallback). A floating text box's own bounds already delimit its
# content unambiguously, so this end marker is only ever needed - and
# only ever looked for - when the block lives among plain body
# paragraphs, where consecutive <w:p> elements alone can't tell the
# inserted block apart from the document's own surrounding content.
# Hidden the same way as the start marker: present in <w:t> text for
# parsing, never rendered by Word.
CONTACT_BOX_HIDDEN_END_MARKER: Final[str] = (
    "LE-GLOBAL-CONTACT-BLOCK-END-V3"
)

DETERMINISTIC_NO_CONTACTS_LINE: Final[str] = (
    "No L&E Global contact is currently configured."
)

_DETERMINISTIC_FIELD_LABELS: Final[
    tuple[tuple[str, str], ...]
] = (
    ("Member firm: ", "member_firm"),
    ("Contact person: ", "contact_person"),
    ("Email: ", "email"),
    ("Phone: ", "phone"),
    ("Address: ", "address"),
    ("Website: ", "website"),
)


def _parse_deterministic_contact_entry(
    lines: Sequence[str],
) -> ExtractedContact:
    """Parse one blank-line-delimited group of "Label: value" lines
    (the exact inverse of build_contact_chunk_content's own per-contact
    formatting) into one ExtractedContact."""

    fields: dict[str, str] = {}

    for line in lines:
        for label, field_name in _DETERMINISTIC_FIELD_LABELS:
            if line.startswith(label):
                fields[field_name] = line[len(label):].strip()
                break

    return ExtractedContact(**fields)


def _parse_contact_lines_after_marker(
    remaining: Sequence[str],
) -> list[ExtractedContact]:
    """Turn the blank-line-delimited lines following a hidden marker
    into contacts - shared by both places the marker can appear (a
    floating text box's own content, or plain in-flow body paragraphs
    for a document with no pre-existing Contact box)."""

    if remaining and remaining[0].strip() == DETERMINISTIC_NO_CONTACTS_LINE:
        return []

    contacts: list[ExtractedContact] = []
    current_lines: list[str] = []

    def flush_current_entry() -> None:
        if current_lines:
            contacts.append(
                _parse_deterministic_contact_entry(current_lines)
            )

    for text in remaining:
        stripped = text.strip()

        if stripped == "":
            flush_current_entry()
            current_lines = []
            continue

        current_lines.append(stripped)

    flush_current_entry()

    return contacts


def extract_deterministic_contact_blocks(
    source: Path | IO[bytes],
) -> list[ExtractedContact] | None:
    """
    Recognize the hidden-marker Contact block a materialized Download
    DOCX carries - either one <w:txbxContent> block (the document's own
    real Contact text box, reused and resized in place by
    app.services.document_contact_materializer) or, for a document with
    no pre-existing Contact box at all, a run of plain in-flow body
    paragraphs bounded by a hidden start and end marker - whose first
    line is the hidden marker.

    Returns None (never an empty list) when no such block is found, so
    callers fall back to the ordinary text-box-based legacy parser - an
    empty list is reserved for "the marker is present and explicitly
    states zero contacts". Reads document.xml directly (the same way
    extract_text_box_blocks() does) rather than via python-docx's
    paragraph API, since a floating shape's own content never shows up
    in python-docx's paragraph API at all.
    """

    try:
        with zipfile.ZipFile(source) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode(
                "utf-8",
                errors="ignore",
            )

    except (
        KeyError,
        OSError,
        zipfile.BadZipFile,
    ):
        return None

    for match in _TEXT_BOX_CONTENT_PATTERN.finditer(document_xml):
        lines = _text_box_lines_preserving_blanks(match.group(1))

        if not lines or lines[0].strip() != CONTACT_BOX_HIDDEN_MARKER:
            continue

        return _parse_contact_lines_after_marker(lines[1:])

    # No floating Contact box carries the marker - look for it among
    # plain body paragraphs instead (the in-flow fallback used for a
    # document with no pre-existing Contact box, e.g. an originally
    # zero-contact source). Unlike a text box's own bounds, consecutive
    # <w:p> elements alone can't tell an inserted block apart from the
    # document's own surrounding paragraphs, so this path is bounded by
    # an explicit hidden end marker instead.
    paragraph_lines = [
        _single_paragraph_line(match.group(0))
        for match in _TEXT_BOX_PARAGRAPH_PATTERN.finditer(document_xml)
    ]

    for index, line in enumerate(paragraph_lines):
        if line.strip() != CONTACT_BOX_HIDDEN_MARKER:
            continue

        block_lines: list[str] = []

        for later_line in paragraph_lines[index + 1:]:
            if later_line.strip() == CONTACT_BOX_HIDDEN_END_MARKER:
                return _parse_contact_lines_after_marker(block_lines)

            block_lines.append(later_line)

        # End marker never found - malformed, fall through to the
        # legacy parser rather than guessing where the block ends.
        break

    return None


def _single_paragraph_line(paragraph_xml: str) -> str:
    """Extract the one line of text one already-isolated <w:p>...</w:p>
    fragment represents - the same per-paragraph token-to-text logic
    _text_box_lines_preserving_blanks applies to each paragraph it
    finds, factored out so a single paragraph (e.g. one already matched
    by _TEXT_BOX_PARAGRAPH_PATTERN elsewhere) can be converted without
    re-running paragraph-boundary detection on it."""

    tokens = [
        match.group(1)
        if match.group(1) is not None
        else ", "
        for match in _TEXT_BOX_TOKEN_PATTERN.finditer(
            paragraph_xml
        )
    ]

    return _clean_text_box_line(
        _normalize_text(
            unescape(
                "".join(tokens)
            )
        )
    )


def _text_box_lines_preserving_blanks(
    raw_block_xml: str,
) -> list[str]:
    """Like extract_text_box_blocks()'s own per-paragraph line
    extraction, but keeps one "" entry per blank paragraph instead of
    dropping it - the hidden-marker Contact block relies on a blank
    paragraph as the delimiter between one contact's fields and the
    next."""

    return [
        _single_paragraph_line(paragraph_xml)
        for paragraph_xml in _TEXT_BOX_PARAGRAPH_PATTERN.findall(
            raw_block_xml
        )
    ]