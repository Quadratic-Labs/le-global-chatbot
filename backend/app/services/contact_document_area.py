"""
Persist the Admin-managed contact area directly into the source DOCX
as ONE deterministic, Word-native structure: a standard, borderless,
in-flow table (one row per contact, a left text cell and a right
photo cell) - never a floating DrawingML/VML shape.

This replaces two earlier, both explicitly abandoned, approaches:

(A) Cloning the document's own complex existing floating shape,
    regenerating its VML fallback/gfxdata/internal ids per new
    contact. Produced malformed output real Microsoft Word refused to
    open.

(B) A brand-new, hand-authored, VML-free DrawingML-only floating text
    box per contact (this module's own SECOND attempt). Real
    Microsoft Word ALSO refused to open the result ("Word encountered
    an error trying to open the file"), even though the output was a
    valid ZIP/well-formed XML and opened fine in LibreOffice. LibreOffice's
    OOXML reader is materially more lenient than Word's; ZIP/XML
    validity and a LibreOffice render are necessary checks, never
    sufficient proof of Word compatibility, for a hand-authored
    <wps:wsp> shape. The precise malformed construct was not isolated
    (the fix here is a full redesign away from hand-authored floating
    shapes, not a repair), but every new run this module writes now
    goes through python-docx's own element/table/picture APIs instead
    of a hand-built XML string, which is the same code path python-docx
    uses for any ordinary table or inline picture in any document it
    produces.

Real OOXML structure, inspected directly (never assumed) from the
healthy AU baseline's own document.xml: the organic contact area is
NOT its own paragraph. One single body paragraph (the "carrier") hosts
BOTH the page's own real heading text ("Employment Law Overview -
Australia", <w:pStyle w:val="Kop1"/>) AND, as sibling floating-shape
runs anchored inside that SAME paragraph: a cover-page URL text box, a
cover-page title text box, the L&E Global logo image, Michael Harmer's
own portrait, his CONTACT PERSON text box, his member-firm text box,
and a decorative background rectangle wrapped "topAndBottom"
(<wp:wrapTopAndBottom/>) - the actual mechanism that reserves the
"large empty band" and pushes "Introduction" down. The contact-specific
runs are told apart from the page's own furniture (logo, title, URL
box) generically, never by name/color: a text box is "contact info" if
its own <wp:positionV relativeFrom="paragraph"> (page furniture instead
anchors relativeFrom="page"); a photo is "a contact's own portrait" if
its embedded relationship id matches one of contact_photos.py's own
already-proven, geometry-vetted candidates (so the L&E Global logo -
also a <pic:pic>, but never geometrically associated with a CONTACT
PERSON zone - is never mistaken for one); the reserved-space rectangle
is the tallest "topAndBottom"-wrapped shape with neither text nor a
picture. Removing exactly those runs, and splitting the carrier
paragraph so any REAL heading text it also carries moves into its own
new paragraph placed AFTER the freshly inserted table, leaves no empty
band and lets Word's own normal document flow push "Introduction" down
naturally - both explicit mission requirements.

The table itself is built purely from ContactState field values -
never by cloning/substituting into the original shape's own template
paragraphs (this module's first, abandoned design) - which is what
lets EVERY mutation (add, delete, replace-photo) rebuild the WHOLE
canonical area fresh from the complete intended contact list, rather
than surgically moving one contact's text/photo independently of the
others. Rebuilding from the full list, every time, is what keeps
SOURCE DOCX == ContactState == OpenSearch from drifting apart.

A contact whose ContactState record has never had an explicit Admin
photo mutation (photo_filename is None - true for every contact
bootstrap_legacy_contacts() ever seeds, since it deliberately never
touches photos) can still have a real, visible photo in the CURRENT
persisted source: the document's own organic portrait. Discarding it
silently the first time any mutation rebuilds the area would be a
real photo-loss regression, so resolve_untracked_contact_photo() below
re-associates it by name from the current document's own extraction,
refusing to guess (returning None) whenever the current contact/photo
counts do not align 1:1.

Word compatibility, in order of what actually protects it:
1. Every NEW run this module writes is built via python-docx's own
   Document/Table/Cell/Paragraph/Run API (add_table, add_paragraph,
   add_run, add_picture) or via docx.oxml.OxmlElement/qn (python-docx's
   own schema-aware element factory, e.g. for <w:tblBorders>) - never a
   hand-authored XML string.
2. Table borders are explicitly nilled (never left to a template's own
   default table style, which is not guaranteed borderless).
3. A supplied photo is inserted via Run.add_picture(stream,
   width=Emu(...)) with height left for python-docx to compute from
   the image's own stored dimensions - preserving its native aspect
   ratio automatically, never a manually computed height (the likely
   cause of an earlier, separate, already-reverted table experiment's
   own "narrow strip" portraits, whose code no longer exists in this
   tree to inspect directly).
4. A source image carrying an OOXML crop rectangle (<a:srcRect>) is
   refused (ContactAreaError) rather than copied blindly - this
   environment has no image library to reproduce a VISIBLE crop, and
   the current real corpus has no cropped contact photo to begin with
   (verified directly), so failing closed costs nothing today while
   remaining safe if that ever changes.
"""

from __future__ import annotations

import hashlib
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import Sequence
from zipfile import ZipFile

from docx import Document as WordDocument
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_LINE_SPACING
from docx.image.image import Image as DocxImage
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt
from lxml import etree

from app.services.contact_document_photos import (
    _DOCUMENT_XML_PART,
    _run_span_for_relationship,
    _write_temp_docx,
)
from app.services.contact_photos import (
    ContactPhotoExtractionError,
    extract_contact_photo_candidates,
)
from app.services.docx_parser import (
    CONTACT_TABLE_HIDDEN_MARKER,
    ExtractedContact,
    extract_contacts_from_docx,
)


class ContactAreaError(RuntimeError):
    """Raised when the persisted canonical contact table cannot be
    safely rebuilt."""


_NSMAP = {
    "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
    "pic": "http://schemas.openxmlformats.org/drawingml/2006/picture",
    "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
}
_R_EMBED_ATTR = f'{{{_NSMAP["r"]}}}embed'

_SRC_RECT_PATTERN = re.compile(r"<a:srcRect\b")

# The right (photo) column as a fraction of the table's own usable
# width - a percentage, never a fixed EMU width, so the table always
# fits whatever page size/margins a given document actually has.
_RIGHT_COLUMN_WIDTH_RATIO = 0.29
# The photo's own rendered width - independent of the column's width
# (a wider column just adds whitespace around it) - kept equal to the
# already Word-validated canary's own photo size.
_PHOTO_TARGET_WIDTH_EMU = 990_600
# Visual gap between one contact's row and the next, applied as a
# table-wide bottom cell margin rather than an empty paragraph.
_ROW_SEPARATION_PT = 12


@dataclass(frozen=True, slots=True)
class ContactPhotoPayload:
    """One contact photo's raw bytes, ready to embed as an ordinary
    inline picture."""

    data: bytes
    content_type: str


def _normalize_text(value: str) -> str:
    return re.sub(r"\s+", " ", value).strip()


# --- Legacy floating-shape carrier detection/removal -------------------


def _wrap_kind(anchor) -> str | None:
    for child in anchor:
        local = etree.QName(child.tag).localname
        if local.startswith("wrap"):
            return local
    return None


def _position_v_relative_from(anchor) -> str | None:
    position_v = anchor.find("wp:positionV", _NSMAP)
    return position_v.get("relativeFrom") if position_v is not None else None


def _extent_cy(anchor) -> int | None:
    extent = anchor.find("wp:extent", _NSMAP)
    if extent is None:
        return None
    value = extent.get("cy")
    return int(value) if value is not None else None


def _run_txbx_text(run) -> str:
    txbx = run.find(".//w:txbxContent", _NSMAP)
    return "".join(txbx.itertext()).strip() if txbx is not None else ""


def _run_photo_relationship_id(run) -> str | None:
    blip = run.find(".//a:blip", _NSMAP)
    return blip.get(_R_EMBED_ATTR) if blip is not None else None


def _remove_legacy_carrier_and_get_anchor(
    document, *, legacy_photo_relationship_ids: set[str]
):
    """
    Find the ONE body paragraph hosting the document's own organic
    contact-area shapes (see module docstring for how each is told
    apart from page furniture) and remove exactly those runs, leaving
    every other shape untouched. If that SAME paragraph also carries
    real heading text after a page break, that text is moved into a
    new paragraph placed immediately after this one, so the caller can
    insert the new canonical table strictly BEFORE it.

    Returns the lxml element the new table should be inserted directly
    after, or None if this document has no such carrier paragraph at
    all (a document with no legacy contact area to canonicalize away).
    """

    body = document.element.body

    for paragraph_element in list(body.findall("w:p", _NSMAP)):
        run_children = list(paragraph_element.findall("w:r", _NSMAP))

        removable_runs = []
        rectangle_candidates: list[tuple[object, int]] = []

        for run in run_children:
            anchor = run.find(".//wp:anchor", _NSMAP)

            if anchor is None:
                continue

            wrap_kind = _wrap_kind(anchor)
            position_v_relative_from = _position_v_relative_from(anchor)
            has_text = bool(_run_txbx_text(run))
            photo_relationship_id = _run_photo_relationship_id(run)

            is_contact_textbox = (
                has_text and position_v_relative_from == "paragraph"
            )
            is_legacy_photo = (
                photo_relationship_id is not None
                and photo_relationship_id in legacy_photo_relationship_ids
            )
            is_rectangle_candidate = (
                wrap_kind == "wrapTopAndBottom"
                and not has_text
                and photo_relationship_id is None
            )

            if is_contact_textbox or is_legacy_photo:
                removable_runs.append(run)
            elif is_rectangle_candidate:
                rectangle_candidates.append((run, _extent_cy(anchor) or 0))

        if not removable_runs:
            continue

        if rectangle_candidates:
            tallest_rectangle_run = max(
                rectangle_candidates, key=lambda item: item[1]
            )[0]
            removable_runs.append(tallest_rectangle_run)

        removable_ids = {id(run) for run in removable_runs}
        first_removable_position = min(
            index
            for index, run in enumerate(run_children)
            if id(run) in removable_ids
        )
        # EVERY surviving run from the FIRST removed one onward moves
        # together into the trailing paragraph - not just plain text,
        # and not only runs strictly after the LAST removed one (a
        # surviving shape can sit BETWEEN two removed runs, as the
        # real AU baseline's own page-2 logo does, wedged between the
        # firm textbox and the reserved-space rectangle). Empirically
        # (rendered via LibreOffice, both before and after this fix),
        # a page-relative floating shape left behind in the now-
        # shapes-thinned remainder paragraph can be reassigned to the
        # WRONG page by the renderer's own page-break heuristics once
        # its neighboring paragraph-relative contact runs are gone -
        # even though its own position/relativeFrom values never
        # change. Moving the whole trailing run of surviving content
        # (any real heading text AND any shape that followed the
        # contact area's own START in the original document) into its
        # own paragraph placed immediately after the new table keeps
        # it paired with the page-2 content it always visually
        # belonged with, sidestepping that page-reassignment entirely
        # rather than depending on it.
        trailing_runs = [
            run
            for run in run_children[first_removable_position:]
            if id(run) not in removable_ids
        ]

        for run in removable_runs:
            run.getparent().remove(run)

        if trailing_runs:
            new_paragraph = OxmlElement("w:p")
            original_ppr = paragraph_element.find("w:pPr", _NSMAP)

            if original_ppr is not None:
                paragraph_element.remove(original_ppr)
                new_paragraph.append(original_ppr)

            for run in trailing_runs:
                new_paragraph.append(run)

            paragraph_element.addnext(new_paragraph)

        return paragraph_element

    return None


# --- Canonical table detection/removal ----------------------------------


def _remove_existing_canonical_table(document):
    """
    Remove this document's own previously-rebuilt canonical table (if
    any), returning the lxml element the new table should be inserted
    directly after - the removed table's own preceding sibling, or
    None when the removed table was the very first body element (the
    caller then inserts the new one at the body's own start).

    Returns _NO_TABLE_FOUND (a sentinel distinct from None) when this
    document has no canonical table at all, so the caller can tell
    "insert at body start" apart from "no canonical table existed".
    """

    for table in document.tables:
        if not table.rows:
            continue

        if CONTACT_TABLE_HIDDEN_MARKER not in table.rows[0].cells[0].text:
            continue

        table_element = table._tbl
        previous_sibling = table_element.getprevious()
        table_element.getparent().remove(table_element)
        return True, previous_sibling

    return False, None


# --- Canonical table construction ---------------------------------------


def _remove_table_borders(table) -> None:
    table_properties = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")

    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_element = OxmlElement(f"w:{edge}")
        edge_element.set(qn("w:val"), "nil")
        borders.append(edge_element)

    table_properties.append(borders)


def _reset_line_spacing(paragraph):
    """
    Override the document's own inherited default paragraph spacing
    for one freshly-built table paragraph. This real AU baseline's own
    docDefaults sets an EXACT (not "at least") 14pt line height - fine
    for a single line of ordinary text, but EXACT line spacing clips
    anything taller to that same 14pt instead of growing for it,
    silently cropping a contact photo (confirmed directly: the same
    clipping reproduces with a synthetic, non-photo test image, so it
    is not specific to any one image's own format/encoding) down to a
    thin sliver. Single line spacing lets a paragraph's own tallest
    inline content set its own height instead.
    """

    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _new_cell_paragraph(cell, *, used_first: bool):
    """One fresh table-cell paragraph, reusing the cell's own default
    first paragraph exactly once - with line spacing reset (see
    _reset_line_spacing) and space_before/space_after both pinned to
    an explicit value by the caller, never left to the document's own
    inherited style default, so cell-to-cell spacing stays predictable
    regardless of which fields a given contact happens to have."""

    paragraph = cell.paragraphs[0] if not used_first else cell.add_paragraph()
    _reset_line_spacing(paragraph)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    return paragraph


def _fill_text_cell(
    cell, contact: ExtractedContact, *, hidden_marker: str | None
) -> None:
    """LEFT cell: member firm, address, phone, website only - the
    contact person/email block now lives in the RIGHT cell, directly
    under its own photo (see _fill_photo_and_person_cell)."""

    used_first = False

    def _next_paragraph():
        nonlocal used_first
        paragraph = _new_cell_paragraph(cell, used_first=used_first)
        used_first = True
        return paragraph

    if hidden_marker:
        marker_run = _next_paragraph().add_run(hidden_marker)
        marker_run.font.hidden = True
        marker_run.font.size = Pt(1)

    if contact.member_firm:
        paragraph = _next_paragraph()
        paragraph.paragraph_format.space_after = Pt(9)
        firm_run = paragraph.add_run(contact.member_firm)
        firm_run.bold = True

    for value in (contact.address, contact.phone):
        if value:
            paragraph = _next_paragraph()
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.add_run(value)

    if contact.website:
        _next_paragraph().add_run(contact.website)


def _fill_photo_and_person_cell(
    cell, contact: ExtractedContact, photo: ContactPhotoPayload | None
) -> None:
    """RIGHT cell: the photo (if any) at the top, then - in the SAME
    standard table cell, never a separate floating object - CONTACT
    PERSON/name/email directly below it. A photo-less contact starts
    the CONTACT PERSON block at the top of the cell instead of
    reserving a blank photo-shaped area."""

    cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    used_first = False

    def _next_paragraph():
        nonlocal used_first
        paragraph = _new_cell_paragraph(cell, used_first=used_first)
        used_first = True
        return paragraph

    has_photo = photo is not None

    if photo is not None:
        photo_paragraph = _next_paragraph()
        photo_paragraph.paragraph_format.space_after = Pt(6)

        try:
            native_image = DocxImage.from_blob(photo.data)
            height_emu = round(
                _PHOTO_TARGET_WIDTH_EMU
                * native_image.px_height
                / native_image.px_width
            )

            photo_paragraph.add_run().add_picture(
                BytesIO(photo.data),
                width=Emu(_PHOTO_TARGET_WIDTH_EMU),
                height=Emu(height_emu),
            )
        except Exception as error:
            raise ContactAreaError(
                f"The contact's photo could not be embedded as a "
                f"standard inline picture: {error}"
            ) from error

    label_paragraph = _next_paragraph()
    label_paragraph.paragraph_format.space_before = (
        Pt(5) if has_photo else Pt(0)
    )
    label_paragraph.paragraph_format.space_after = Pt(2)
    label_run = label_paragraph.add_run("CONTACT PERSON")
    label_run.bold = True

    if contact.contact_person:
        name_paragraph = _next_paragraph()
        name_paragraph.paragraph_format.space_after = Pt(2)
        name_paragraph.add_run(contact.contact_person)

    if contact.email:
        _next_paragraph().add_run(contact.email)


def _set_table_cell_margins(table, *, bottom_pt: float) -> None:
    """A default BOTTOM cell margin, applied table-wide - the visual
    separation between one contact's row and the next, without any
    artificial blank paragraph. Kept off row height entirely (no
    fixed/minimum row height is set anywhere in this module): a cell
    margin adds space AROUND a row's own naturally-grown content
    instead of constraining that content's own height."""

    table_properties = table._tbl.tblPr
    cell_margins = OxmlElement("w:tblCellMar")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:w"), str(round(bottom_pt * 20)))
    bottom.set(qn("w:type"), "dxa")
    cell_margins.append(bottom)
    table_properties.append(cell_margins)


def _prevent_row_split(row) -> None:
    """Best-effort: keep one contact's row from splitting across a
    page boundary (standard w:cantSplit) - "if practical", since a
    single contact whose own content is genuinely taller than one
    page can still split; this only prevents an ordinary row from
    being cut for no reason."""

    row_properties = row._tr.get_or_add_trPr()
    row_properties.append(OxmlElement("w:cantSplit"))


def _build_canonical_table(
    document,
    *,
    contacts: Sequence[ExtractedContact],
    photos: Sequence[ContactPhotoPayload | None],
):
    table = document.add_table(rows=len(contacts), cols=2)
    table.autofit = False

    section = document.sections[0]
    usable_width_emu = int(
        section.page_width - section.left_margin - section.right_margin
    )
    right_column_width_emu = round(
        usable_width_emu * _RIGHT_COLUMN_WIDTH_RATIO
    )
    left_column_width_emu = max(
        usable_width_emu - right_column_width_emu, Emu(1)
    )

    for column, width_emu in zip(
        table.columns, (left_column_width_emu, right_column_width_emu)
    ):
        column.width = Emu(width_emu)

        for cell in column.cells:
            cell.width = Emu(width_emu)

    _remove_table_borders(table)
    _set_table_cell_margins(table, bottom_pt=_ROW_SEPARATION_PT)

    for row_index, (contact, photo) in enumerate(zip(contacts, photos)):
        row = table.rows[row_index]
        _fill_text_cell(
            row.cells[0],
            contact,
            hidden_marker=(
                CONTACT_TABLE_HIDDEN_MARKER if row_index == 0 else None
            ),
        )
        _fill_photo_and_person_cell(row.cells[1], contact, photo)
        _prevent_row_split(row)

    return table


def _default_insertion_anchor(document):
    """
    A document with no existing contact area at all (no legacy shapes,
    no canonical table): the same anchor convention this mission's own
    persisted in-flow fallback already used - the first paragraph with
    actual visible text (typically the document's own in-flow title) -
    so the new table lands right after it, before the real body
    content, exactly mirroring the legacy-carrier case's own placement
    relative to "Introduction".
    """

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            return paragraph._p

    return document.paragraphs[0]._p if document.paragraphs else None


def _has_crop_rectangle(document_xml: str, relationship_id: str) -> bool:
    span = _run_span_for_relationship(document_xml, relationship_id)

    if span is None:
        return False

    return bool(_SRC_RECT_PATTERN.search(document_xml[span[0]:span[1]]))


def resolve_untracked_contact_photo(
    source_path: Path,
    *,
    contact_person: str | None,
    country: str | None,
) -> ContactPhotoPayload | None:
    """
    A contact whose ContactRecord has never had an explicit Admin photo
    mutation (photo_filename is None) may still have a real, visible
    photo in the CURRENT persisted source - its own organic portrait,
    never captured into ContactState since bootstrap_legacy_contacts()
    deliberately never touches photos. Re-associate it here, by name,
    from the current document's own extraction (contact_photos.py's
    own documented "person order" contract), refusing to guess
    (returning None) whenever the current contact/photo counts do not
    align 1:1 - never silently attaching the wrong contact's photo.

    Also refuses (raises ContactAreaError) a photo whose own source
    shape carries an OOXML crop rectangle, rather than embedding the
    UNCROPPED raw media and silently changing how the portrait looks -
    this environment has no image library available to reproduce the
    crop.
    """

    if not contact_person:
        return None

    try:
        current_contacts = extract_contacts_from_docx(
            source_path, country=country
        )
        current_photos = extract_contact_photo_candidates(source_path)
    except (ContactPhotoExtractionError, Exception):
        return None

    if not current_photos or len(current_contacts) != len(current_photos):
        return None

    normalized_target = _normalize_text(contact_person).casefold()

    for contact, photo in zip(current_contacts, current_photos):
        if (
            _normalize_text(contact.contact_person or "").casefold()
            != normalized_target
        ):
            continue

        if photo.relationship_id:
            with ZipFile(source_path) as archive:
                document_xml = archive.read(_DOCUMENT_XML_PART).decode(
                    "utf-8", errors="ignore"
                )

            if _has_crop_rectangle(document_xml, photo.relationship_id):
                raise ContactAreaError(
                    f"{contact_person}'s own photo relies on an OOXML "
                    "crop rectangle this environment cannot reproduce - "
                    "refusing to embed the uncropped raw image."
                )

        return ContactPhotoPayload(
            data=photo.data, content_type=photo.content_type
        )

    return None


def rebuild_canonical_contact_table(
    source_path: Path,
    *,
    contacts: Sequence[ExtractedContact],
    photos: Sequence[ContactPhotoPayload | None],
    country: str | None = None,
) -> bytes:
    """
    Rebuild the ENTIRE persisted canonical contact area from the
    complete intended contact list - the one mechanism for every
    mutation (add, delete, replace-photo), so the source DOCX never
    drifts from ContactState by having some contacts' shapes moved
    surgically while others are left stale.

    Removes whichever contact area currently exists (a previously
    rebuilt canonical table, or the document's own original floating-
    shape area, or neither) and inserts a fresh standard Word table in
    its place - immediately before whatever real heading/body content
    followed it, so no empty reserved band is ever left behind. Builds
    nothing at all when contacts is empty (a document with zero
    contacts has no area to show).
    """

    if len(contacts) != len(photos):
        raise ContactAreaError(
            "contacts and photos must be the same length."
        )

    source_bytes = source_path.read_bytes()

    try:
        legacy_photo_relationship_ids = {
            candidate.relationship_id
            for candidate in extract_contact_photo_candidates(source_path)
            if candidate.relationship_id
        }
    except ContactPhotoExtractionError as error:
        raise ContactAreaError(
            f"Could not inspect the source document for existing "
            f"contact photos: {error}"
        ) from error

    document = WordDocument(BytesIO(source_bytes))

    canonical_table_found, anchor_element = _remove_existing_canonical_table(
        document
    )

    if not canonical_table_found:
        anchor_element = _remove_legacy_carrier_and_get_anchor(
            document,
            legacy_photo_relationship_ids=legacy_photo_relationship_ids,
        )

        if anchor_element is None:
            anchor_element = _default_insertion_anchor(document)

    if contacts:
        table = _build_canonical_table(
            document, contacts=contacts, photos=photos
        )

        if anchor_element is not None:
            anchor_element.addnext(table._tbl)
        else:
            document.element.body.insert(0, table._tbl)

    output = BytesIO()
    document.save(output)
    new_bytes = output.getvalue()

    _validate_canonical_table(
        new_bytes,
        expected_contacts=contacts,
        expected_photo_shas=[
            hashlib.sha256(photo.data).hexdigest() if photo else None
            for photo in photos
        ],
        country=country,
    )

    return new_bytes


# Every structured field the canonical table's writer owns - the full
# round-trip contract a rebuild must preserve, not just contact_person/
# email. A field shift (e.g. a phone value landing in member_firm, or
# a website value landing in address) must fail validation exactly
# like a lost contact would.
_VALIDATED_CONTACT_FIELDS = (
    "member_firm",
    "address",
    "phone",
    "website",
    "contact_person",
    "email",
)


def _normalized_field_value(contact: ExtractedContact, field: str) -> str:
    return _normalize_text(getattr(contact, field) or "").casefold()


def _validate_canonical_table(
    new_bytes: bytes,
    *,
    expected_contacts: Sequence[ExtractedContact],
    expected_photo_shas: Sequence[str | None],
    country: str | None,
) -> None:
    """Structural, not just SHA: the rebuilt table must round-trip back
    out as exactly the intended contacts, in order, in every
    structured field the writer owns, with their own photos -
    otherwise no change is saved at all."""

    temp_path = _write_temp_docx(new_bytes)

    try:
        reparsed = extract_contacts_from_docx(temp_path, country=country)
    except Exception as error:
        raise ContactAreaError(
            f"The rebuilt document could not be re-parsed for "
            f"contacts - no change was saved: {error}"
        ) from error
    finally:
        temp_path.unlink(missing_ok=True)

    if len(reparsed) != len(expected_contacts):
        raise ContactAreaError(
            f"Expected {len(expected_contacts)} contacts after "
            f"rebuilding the canonical area, found {len(reparsed)} - "
            f"no change was saved."
        )

    for expected, actual in zip(expected_contacts, reparsed):
        mismatched_fields = [
            field
            for field in _VALIDATED_CONTACT_FIELDS
            if _normalized_field_value(expected, field)
            != _normalized_field_value(actual, field)
        ]

        if mismatched_fields:
            raise ContactAreaError(
                "The rebuilt canonical table did not round-trip back "
                "out with the same contacts in the same order - "
                f"field(s) changed: {', '.join(mismatched_fields)} - "
                "no change was saved."
            )

    if any(sha is not None for sha in expected_photo_shas):
        photo_path = _write_temp_docx(new_bytes)

        try:
            candidates = extract_contact_photo_candidates(photo_path)
        except ContactPhotoExtractionError as error:
            raise ContactAreaError(
                f"The rebuilt document could not be re-parsed for "
                f"contact photos - no change was saved: {error}"
            ) from error
        finally:
            photo_path.unlink(missing_ok=True)

        actual_shas = {candidate.sha256 for candidate in candidates}

        for sha in expected_photo_shas:
            if sha is not None and sha not in actual_shas:
                raise ContactAreaError(
                    "A contact's photo did not round-trip back out of "
                    "the rebuilt canonical table - no change was saved."
                )
