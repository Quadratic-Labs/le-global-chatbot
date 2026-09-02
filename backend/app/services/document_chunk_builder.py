"""Consolidated service document_chunk_builder.py; includes former docx_country_marker.py."""
from __future__ import annotations
import zipfile
from dataclasses import dataclass
from pathlib import Path
from xml.etree import ElementTree as ET
from app.core.country_registry import UnknownCountryCodeError, canonical_country_name, normalize_country_code
import hashlib
import re
from collections import defaultdict
from collections.abc import Sequence
from typing import Final
from docx import Document
from app.core.country_registry import resolve_country
from app.core.country_registry import UnknownCountryNameError
from app.core.legal_taxonomy import get_canonical_legal_topic, is_overview_section
from app.core.subsection_taxonomy import get_subsection_topic_override
from app.models.document import DocumentChunk
from app.services.docx_parser import ExtractedContact, ParsedSection, build_contact_chunk_content, extract_contacts_from_docx, extract_text_box_blocks, parse_docx_sections
from app.services.docx_parser import ParsedSection
CUSTOM_PROPERTIES_PARTNAME = 'docProps/custom.xml'
_CONTENT_TYPES_PARTNAME = '[Content_Types].xml'
_ROOT_RELS_PARTNAME = '_rels/.rels'
CUSTOM_PROPERTIES_CONTENT_TYPE = 'application/vnd.openxmlformats-officedocument.custom-properties+xml'
CUSTOM_PROPERTIES_RELATIONSHIP_TYPE = 'http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties'
_CUSTOM_PROPS_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/custom-properties'
_VT_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes'
_CT_NS = 'http://schemas.openxmlformats.org/package/2006/content-types'
_RELS_NS = 'http://schemas.openxmlformats.org/package/2006/relationships'
_CUSTOM_PROPERTY_FMTID = '{D5CDD505-2E9C-101B-9397-08002B2CF9AE}'
_FIRST_CUSTOM_PROPERTY_PID = 2
COUNTRY_CODE_PROPERTY_NAME = 'LE Global Country Code'
COUNTRY_NAME_PROPERTY_NAME = 'LE Global Country Name'
_DETERMINISTIC_ZIP_DATE_TIME = (1980, 1, 1, 0, 0, 0)
for _prefix, _uri in (('', _CUSTOM_PROPS_NS), ('vt', _VT_NS)):
    ET.register_namespace(_prefix, _uri)

@dataclass(frozen=True, slots=True)
class CountryMarker:
    """One validated, DOCX-native country marker."""
    country_code: str
    country_name: str

class InvalidCountryMarkerValueError(ValueError):
    """Raised when a caller asks to write an unrecognized country code."""

def _property_element(parent: ET.Element, *, pid: int, name: str, value: str) -> ET.Element:
    element = ET.SubElement(parent, f'{{{_CUSTOM_PROPS_NS}}}property', {'fmtid': _CUSTOM_PROPERTY_FMTID, 'pid': str(pid), 'name': name})
    lpwstr = ET.SubElement(element, f'{{{_VT_NS}}}lpwstr')
    lpwstr.text = value
    return element

def _parse_existing_custom_properties(custom_xml: bytes | None) -> list[tuple[int, str, str]]:
    """
    Return every existing custom property as (pid, name, value.

    Deliberately preserves properties this module does not own (any
    custom property a Word user may have already added) - only the two
    LE Global-owned property names are ever touched by _upsert below.
    Malformed or unreadable existing content.xml is treated as absent
    rather than raised, so a document with a foreign/corrupt custom
    properties part still safely gets a fresh, well-formed one.
    """
    if not custom_xml:
        return []
    try:
        root = ET.fromstring(custom_xml)
    except ET.ParseError:
        return []
    properties: list[tuple[int, str, str]] = []
    for element in root.findall(f'{{{_CUSTOM_PROPS_NS}}}property'):
        name = element.get('name')
        pid_raw = element.get('pid')
        if not name or pid_raw is None:
            continue
        try:
            pid = int(pid_raw)
        except ValueError:
            continue
        value_element = element.find(f'{{{_VT_NS}}}lpwstr')
        value = value_element.text if value_element is not None and value_element.text is not None else ''
        properties.append((pid, name, value))
    return properties

def _serialize_custom_properties(properties: list[tuple[int, str, str]]) -> bytes:
    """Build docProps/custom.xml deterministically from a property list."""
    root = ET.Element(f'{{{_CUSTOM_PROPS_NS}}}Properties')
    for pid, name, value in sorted(properties, key=lambda item: item[0]):
        _property_element(root, pid=pid, name=name, value=value)
    body = ET.tostring(root, encoding='unicode')
    return f'<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n{body}'.encode('utf-8')

def _upsert_custom_properties(custom_xml: bytes | None, updates: dict[str, str]) -> bytes:
    """
    Return docProps/custom.xml content with `updates` applied.

    An existing LE Global property (matched by name) keeps its pid and
    only has its value replaced, so re-writing the same marker never
    disturbs any other property's identity. A brand-new property is
    assigned the next unused pid, in the fixed iteration order of
    `updates` - so writing the same set of updates to an empty part
    always assigns the same pids.
    """
    existing = _parse_existing_custom_properties(custom_xml)
    by_name = {name: (pid, value) for pid, name, value in existing}
    other_properties = [(pid, name, value) for pid, name, value in existing if name not in updates]
    used_pids = {pid for pid, _, _ in existing}
    next_pid = max(used_pids, default=_FIRST_CUSTOM_PROPERTY_PID - 1) + 1
    next_pid = max(next_pid, _FIRST_CUSTOM_PROPERTY_PID)
    owned_properties: list[tuple[int, str, str]] = []
    for name, value in updates.items():
        existing_entry = by_name.get(name)
        if existing_entry is not None:
            pid, _ = existing_entry
        else:
            pid = next_pid
            next_pid += 1
        owned_properties.append((pid, name, value))
    return _serialize_custom_properties(other_properties + owned_properties)

def _ensure_content_type_override(content_types_xml: bytes) -> bytes:
    """Add the custom-properties Override, if not already declared."""
    root = ET.fromstring(content_types_xml)
    for override in root.findall(f'{{{_CT_NS}}}Override'):
        if override.get('PartName') == f'/{CUSTOM_PROPERTIES_PARTNAME}':
            return content_types_xml
    ET.SubElement(root, f'{{{_CT_NS}}}Override', {'PartName': f'/{CUSTOM_PROPERTIES_PARTNAME}', 'ContentType': CUSTOM_PROPERTIES_CONTENT_TYPE})
    return ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n' + ET.tostring(root, encoding='unicode')).encode('utf-8')

def _ensure_root_relationship(rels_xml: bytes) -> bytes:
    """Add the custom-properties relationship, if not already present."""
    root = ET.fromstring(rels_xml)
    existing_ids: set[int] = set()
    for relationship in root.findall(f'{{{_RELS_NS}}}Relationship'):
        if relationship.get('Type') == CUSTOM_PROPERTIES_RELATIONSHIP_TYPE:
            return rels_xml
        raw_id = relationship.get('Id', '')
        if raw_id.startswith('rId') and raw_id[3:].isdigit():
            existing_ids.add(int(raw_id[3:]))
    next_id = max(existing_ids, default=0) + 1
    ET.SubElement(root, f'{{{_RELS_NS}}}Relationship', {'Id': f'rId{next_id}', 'Type': CUSTOM_PROPERTIES_RELATIONSHIP_TYPE, 'Target': CUSTOM_PROPERTIES_PARTNAME})
    return ('<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n' + ET.tostring(root, encoding='unicode')).encode('utf-8')

def read_country_marker(document_path: Path) -> CountryMarker | None:
    """
    Read and validate one DOCX-native country marker, or None.

    Returns None - never raises - for a document with no marker part,
    an unreadable/malformed one, or one whose country code does not
    resolve through the country registry: an invalid or unrecognized
    marker is always safely ignored, exactly like a document with no
    marker at all (mission section 11) - never trusted as arbitrary
    text, never a bypass of the confirmation this marker only persists.
    """
    try:
        with zipfile.ZipFile(document_path) as archive:
            if CUSTOM_PROPERTIES_PARTNAME not in archive.namelist():
                return None
            custom_xml = archive.read(CUSTOM_PROPERTIES_PARTNAME)
    except (OSError, zipfile.BadZipFile, KeyError):
        return None
    properties = {name: value for _, name, value in _parse_existing_custom_properties(custom_xml)}
    raw_code = properties.get(COUNTRY_CODE_PROPERTY_NAME)
    if not raw_code or not raw_code.strip():
        return None
    try:
        normalized_code = normalize_country_code(raw_code)
    except UnknownCountryCodeError:
        return None
    country_name = (properties.get(COUNTRY_NAME_PROPERTY_NAME) or '').strip() or canonical_country_name(normalized_code)
    return CountryMarker(country_code=normalized_code, country_name=country_name)

def write_country_marker(source_path: Path, destination_path: Path, *, country_code: str, country_name: str) -> None:
    """
    Write `source_path`'s content to `destination_path`, with a DOCX-
    native country marker for (country_code, country_name) upserted.

    Every other part - body, styles, headers/footers, media, any
    unrelated custom property or relationship - is preserved unchanged.
    Deterministic: the same (source bytes, country_code, country_name)
    always produces the same destination bytes, so repeated writes (or
    a later independent re-parse) never disagree with each other.

    Raises InvalidCountryMarkerValueError for a country_code the
    registry does not recognize - this function never persists an
    unvalidated value.
    """
    try:
        normalized_code = normalize_country_code(country_code)
    except UnknownCountryCodeError as error:
        raise InvalidCountryMarkerValueError(f'Cannot persist an unrecognized country code: {country_code!r}.') from error
    normalized_name = country_name.strip()
    if not normalized_name:
        raise InvalidCountryMarkerValueError('country_name must not be empty.')
    with zipfile.ZipFile(source_path) as source_zip:
        names = [info.filename for info in source_zip.infolist()]
        contents = {name: source_zip.read(name) for name in names}
    contents[CUSTOM_PROPERTIES_PARTNAME] = _upsert_custom_properties(contents.get(CUSTOM_PROPERTIES_PARTNAME), {COUNTRY_CODE_PROPERTY_NAME: normalized_code, COUNTRY_NAME_PROPERTY_NAME: normalized_name})
    contents[_CONTENT_TYPES_PARTNAME] = _ensure_content_type_override(contents[_CONTENT_TYPES_PARTNAME])
    contents[_ROOT_RELS_PARTNAME] = _ensure_root_relationship(contents[_ROOT_RELS_PARTNAME])
    ordered_names = list(names)
    for extra_name in (CUSTOM_PROPERTIES_PARTNAME, _CONTENT_TYPES_PARTNAME, _ROOT_RELS_PARTNAME):
        if extra_name not in ordered_names:
            ordered_names.append(extra_name)
    destination_path.parent.mkdir(parents=True, exist_ok=True)
    with zipfile.ZipFile(destination_path, 'w', zipfile.ZIP_DEFLATED) as destination_zip:
        for name in ordered_names:
            info = zipfile.ZipInfo(filename=name, date_time=_DETERMINISTIC_ZIP_DATE_TIME)
            info.compress_type = zipfile.ZIP_DEFLATED
            destination_zip.writestr(info, contents[name])
CONTACT_SUBSECTION: Final[str] = 'Contact'
DOCUMENT_FAMILY: Final[str] = 'employment-law-overview'
_TITLE_LINE_PATTERNS: Final[tuple[re.Pattern[str], ...]] = (re.compile('^Labour and Employment Law in\\s+(?P<country>.+?)(?:\\s+(?P<year>(?:19|20)\\d{2}))?$', re.IGNORECASE), re.compile('^Employment Law Overview(?:\\s*-\\s*|\\s+)(?P<country>.+?)(?:\\s+(?P<year>(?:19|20)\\d{2}))?$', re.IGNORECASE))
_BARE_FAMILY_HEADING_PATTERN: Final[re.Pattern[str]] = re.compile('^(?:Labour and Employment Law|Employment Law Overview)\\s*/?\\s*$', re.IGNORECASE)
_FAMILY_HEADING_TOKEN_PATTERN: Final[re.Pattern[str]] = re.compile('Labour and Employment Law|Employment Law Overview', re.IGNORECASE)
_TITLE_SCAN_LIMIT: Final[int] = 100
_TITLE_SCAN_CHARACTER_LIMIT: Final[int] = 12000
_DASH_NORMALIZATION_PATTERN: Final[re.Pattern[str]] = re.compile('[‐‑‒–—―]')

def _normalize_front_matter_line(value: str) -> str:
    """Normalize whitespace and dash variants in one front-matter line."""
    return ' '.join(_DASH_NORMALIZATION_PATTERN.sub('-', value).split())

class UnknownLegalTopicError(ValueError):
    """Raised when a section is outside the approved taxonomy."""

class InvalidDocxFormatError(ValueError):
    """Raised when a file is not a genuinely valid, parseable DOCX."""

class UndeterminableDocumentCountryError(ValueError):
    """
    Raised when no supported country can be identified from the
    document's own title/cover content - never from its filename.
    """

class AmbiguousDocumentCountryError(ValueError):
    """
    Raised when more than one distinct country is found in the
    document's title/cover content, with no clear primary one.
    """

@dataclass(frozen=True, slots=True)
class DocumentMetadata:
    """Metadata shared by all chunks from one source document."""
    country: str
    country_code: str
    reference_year: int | None
    language: str
    source_filename: str
    source_format: str = 'docx'

def _sha256(value: str) -> str:
    """Return a deterministic SHA-256 hexadecimal digest."""
    return hashlib.sha256(value.encode('utf-8')).hexdigest()

def storage_filename_for_country(country_code: str) -> str:
    """
    The internal, collision-free on-disk storage name for one
    country's active document under this pipeline's one document
    family - never the user-supplied filename.

    Keying storage by this instead of the original filename (mission
    "CONTINUATION PATCH 0.4.3", section 10) is what makes two
    completely unrelated uploads that happen to share a filename
    (e.g. both named "final.docx") safe: each country still gets its
    own, stable storage path, so neither can silently overwrite the
    other's stored source file on disk. The user's own original
    filename is preserved separately, purely as display/audit
    metadata (DocumentChunk.source_filename) - never as a path.
    """
    return f'{country_code.strip().upper()}.docx'

def validate_docx_format(file_path: Path) -> None:
    """
    Confirm a file is a genuinely valid, parseable DOCX archive -
    the .docx extension alone proves nothing (mission "CONTINUATION
    PATCH 0.4.3", section 6). Never proceeds to content parsing,
    OpenSearch, or any storage change when this fails.
    """
    if not zipfile.is_zipfile(file_path):
        raise InvalidDocxFormatError('The uploaded file is not a valid DOCX document.')
    try:
        with zipfile.ZipFile(file_path) as archive:
            names = set(archive.namelist())
    except zipfile.BadZipFile as error:
        raise InvalidDocxFormatError('The uploaded file is not a valid DOCX document.') from error
    if '[Content_Types].xml' not in names or 'word/document.xml' not in names:
        raise InvalidDocxFormatError('The uploaded file is not a valid DOCX document.')
    try:
        Document(file_path)
    except Exception as error:
        raise InvalidDocxFormatError('The uploaded file is not a valid DOCX document.') from error
_DRAWINGML_TEXT_PATTERN: Final[re.Pattern[str]] = re.compile('<a:t[^>]*>([^<]*)</a:t>')
_WORD_TEXT_RUN_PATTERN: Final[re.Pattern[str]] = re.compile('<w:t[^>]*>([^<]*)</w:t>')
_HEADER_PART_NAME_PATTERN: Final[re.Pattern[str]] = re.compile('^word/header\\d+\\.xml$')

def _extract_drawingml_texts(file_path: Path) -> list[str]:
    """
    Extract DrawingML (SmartArt/WordArt) run text from document.xml -
    a:t elements live in a different XML namespace than a text box's
    own w:t runs, and are otherwise invisible to python-docx and to
    extract_text_box_blocks alike (mission "HOTFIX 0.4.4", section 3).
    """
    try:
        with zipfile.ZipFile(file_path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')
    except (KeyError, OSError, zipfile.BadZipFile):
        return []
    return [text for text in _DRAWINGML_TEXT_PATTERN.findall(document_xml) if text.strip()]

def _extract_header_texts(file_path: Path) -> list[str]:
    """
    Extract each header part's own paragraph text, in file order -
    some real documents carry the country name in a running header
    rather than the document body (mission "HOTFIX 0.4.4", section 4).
    """
    texts: list[str] = []
    try:
        with zipfile.ZipFile(file_path) as archive:
            header_names = sorted((name for name in archive.namelist() if _HEADER_PART_NAME_PATTERN.match(name)))
            for name in header_names:
                header_xml = archive.read(name).decode('utf-8', errors='ignore')
                joined = ''.join(_WORD_TEXT_RUN_PATTERN.findall(header_xml)).strip()
                if joined:
                    texts.append(joined)
    except (OSError, zipfile.BadZipFile):
        return []
    return texts

def _leading_front_matter_blocks(file_path: Path, limit: int=_TITLE_SCAN_LIMIT, character_limit: int=_TITLE_SCAN_CHARACTER_LIMIT) -> list[str]:
    """
    The document's own front matter, never its full body (mission
    "HOTFIX 0.4.4", section 4): legal documents may name many other
    countries deep in their own text, which must never be mistaken
    for the document's own country.

    Word text boxes (cover-page titles, firm/contact cards) are
    anchored drawings, invisible to python-docx's own paragraph
    iteration - real L&E Global documents routinely put their actual
    "Employment Law Overview <Country>" cover title there rather than
    in a normal paragraph, so those lines are read first, followed by
    the docProps title/subject (when set) and then the leading body
    paragraphs. Capped at `limit` blocks or `character_limit`
    cumulative characters, whichever is reached first.
    """
    blocks: list[str] = []
    total_characters = 0

    def add(text: str) -> bool:
        """Append one candidate line; return False once capped."""
        nonlocal total_characters
        stripped = text.strip()
        if not stripped:
            return True
        blocks.append(stripped)
        total_characters += len(stripped)
        return len(blocks) < limit and total_characters < character_limit
    for text_box_lines in extract_text_box_blocks(file_path):
        for line in text_box_lines:
            if not add(line):
                return blocks
    for header_text in _extract_header_texts(file_path):
        if not add(header_text):
            return blocks
    for drawingml_text in _extract_drawingml_texts(file_path):
        if not add(drawingml_text):
            return blocks
    document = Document(file_path)
    for title_like in (document.core_properties.title, document.core_properties.subject):
        if title_like and (not add(title_like)):
            return blocks
    for paragraph in document.paragraphs:
        if not add(paragraph.text):
            return blocks
    return blocks

def _match_title_line(line: str, next_line: str | None) -> tuple[str, int | None] | None:
    """
    One candidate (raw country token, optional year) from a single
    title/cover line, or from that line paired with the one right
    after it - covers a one-line cover ("Employment Law Overview
    Canada 2026", "Labour and Employment Law / Canada") and a
    two-line cover in either order: a bare "Labour and Employment
    Law" heading immediately followed by a "Canada" line, or a bare
    "France" / "Portugal" country line immediately followed by a
    heading line (some real covers put the country first).
    """
    line = _normalize_front_matter_line(line)
    next_line = _normalize_front_matter_line(next_line) if next_line is not None else None
    for pattern in _TITLE_LINE_PATTERNS:
        match = pattern.fullmatch(line)
        if match is None:
            continue
        year_value = match.group('year')
        return (match.group('country').strip(), int(year_value) if year_value is not None else None)
    if '/' in line:
        prefix, _, suffix = line.rpartition('/')
        if _BARE_FAMILY_HEADING_PATTERN.fullmatch(f'{prefix.strip()} /'):
            candidate = suffix.strip()
            if candidate:
                return (candidate, None)
    if _BARE_FAMILY_HEADING_PATTERN.fullmatch(line) and next_line and (not _BARE_FAMILY_HEADING_PATTERN.fullmatch(next_line)):
        return (next_line.strip(), None)
    if next_line and _FAMILY_HEADING_TOKEN_PATTERN.search(next_line) and (not _FAMILY_HEADING_TOKEN_PATTERN.search(line)) and line:
        return (line.strip(), None)
    return None

def _detect_year_only_from_content(file_path: Path) -> int | None:
    """
    Best-effort reference_year from the document's own title/cover
    content, independent of whether a country resolves there at all.

    Used only when a DOCX-native country marker (see
    docx_country_marker.py) already settled the country - the marker
    takes priority for country (mission "ORDER 8E-A1", section 11),
    but a title line's year is still worth capturing rather than
    losing it just because country resolution is skipped.
    """
    lines = _leading_front_matter_blocks(file_path)
    for index, line in enumerate(lines):
        next_line = lines[index + 1] if index + 1 < len(lines) else None
        candidate = _match_title_line(line, next_line)
        if candidate is None:
            continue
        _, year = candidate
        if year is not None:
            return year
    return None

def _detect_country_and_year_from_content(file_path: Path, country_code: str | None=None) -> tuple[str, str, int | None]:
    """
    Resolve (country, country_code, reference_year) from the
    document's own title/cover content - never from its filename.

    A valid DOCX-native country marker (see docx_country_marker.py)
    always takes priority over content detection (mission "ORDER
    8E-A1", section 11): an Admin's manually-selected country, once
    persisted into the DOCX itself, must keep being recognized on
    every future Download/Reindex/re-upload, even if the document's
    own content still can't be resolved to a country on its own. The
    marker only ever settles *which* country was detected - it is
    never a bypass of the confirmation the upload flow always
    requires (see admin_document_replacement.py).

    Only the leading front matter is scanned (see
    _leading_front_matter_blocks). Candidates are deduplicated by
    resolved country *code*, not raw text, so trivially different
    phrasings of the same country never look ambiguous. More than
    one distinct country code found there - a genuinely ambiguous
    cover - is refused rather than guessed at (section 7).
    """
    marker = read_country_marker(file_path)
    if marker is not None:
        country, resolved_code = resolve_country(raw_country=marker.country_name, country_code=country_code)
        return (country, resolved_code, _detect_year_only_from_content(file_path))
    lines = _leading_front_matter_blocks(file_path)
    resolved: list[tuple[str, str, int | None]] = []
    for index, line in enumerate(lines):
        next_line = lines[index + 1] if index + 1 < len(lines) else None
        candidate = _match_title_line(line, next_line)
        if candidate is None:
            continue
        raw_country, year = candidate
        try:
            country, resolved_code = resolve_country(raw_country=raw_country, country_code=country_code)
        except UnknownCountryNameError:
            continue
        resolved.append((country, resolved_code, year))
    if not resolved:
        raise UndeterminableDocumentCountryError('Unable to identify a supported country from the document content.')
    distinct_codes = {code for _, code, _ in resolved}
    if len(distinct_codes) > 1:
        raise AmbiguousDocumentCountryError('Unable to determine a unique document country from the document content.')
    country, resolved_code, _ = resolved[0]
    reference_year = next((year for _, _, year in resolved if year is not None), None)
    return (country, resolved_code, reference_year)

def _validate_metadata(metadata: DocumentMetadata) -> None:
    """Validate metadata before creating chunks."""
    if not metadata.country.strip():
        raise ValueError('country must not be empty')
    country_code = metadata.country_code.strip()
    if len(country_code) != 2 or not country_code.isalpha():
        raise ValueError('country_code must contain exactly two alphabetical characters')
    if not metadata.language.strip():
        raise ValueError('language must not be empty')
    if not metadata.source_filename.strip():
        raise ValueError('source_filename must not be empty')
    if not metadata.source_format.strip():
        raise ValueError('source_format must not be empty')
    if metadata.reference_year is not None and (not 1900 <= metadata.reference_year <= 2100):
        raise ValueError('reference_year must be between 1900 and 2100')

def _build_document_id(metadata: DocumentMetadata) -> str:
    """
    Build a stable identifier for one source document.

    Identity is country_code + document family + language only -
    never reference_year, never source_filename (mission
    "CONTINUATION PATCH 0.4.3", section 11). A new upload for a
    country that is already active, of any year and under any
    filename, resolves to this exact same document_id, so it
    replaces the previous version through the existing replace
    mechanism instead of creating a second, parallel document.
    reference_year and source_filename remain per-chunk *version*
    metadata (see DocumentChunk) - they simply never affect which
    document a chunk belongs to.
    """
    identity = '\x1f'.join(('document-v2', DOCUMENT_FAMILY, metadata.country_code.strip().upper(), metadata.language.strip().casefold()))
    return f'doc_{_sha256(identity)}'

def _build_chunk_id(document_id: str, document_type: str, legal_topic: str | None, section: str, subsection: str | None, occurrence: int) -> str:
    """Build a stable chunk identifier from its structural path."""
    identity = '\x1f'.join(('chunk-v1', document_id, document_type, legal_topic or '', section.casefold(), (subsection or '').casefold(), str(occurrence)))
    return f'chunk_{_sha256(identity)}'

def metadata_from_content(file_path: Path, country_code: str | None=None, language: str='en') -> DocumentMetadata:
    """
    Extract canonical metadata from the document's own content - the
    filename is never consulted for country, year, or document type
    (mission "CONTINUATION PATCH 0.4.3", section 5): it is preserved
    only as source_filename, for display/audit, never as a source of
    truth for anything else.

    The country token is resolved through the central country
    registry, exactly as before - only *where* it is read from has
    changed, from the filename to the document's title/cover content.

    An explicit country code is optional. When supplied, it must
    correspond to the country found in the document's own content.
    """
    country, resolved_country_code, reference_year = _detect_country_and_year_from_content(file_path, country_code=country_code)
    source_format = file_path.suffix.lstrip('.').lower()
    return DocumentMetadata(country=country, country_code=resolved_country_code, reference_year=reference_year, language=language.strip().lower(), source_filename=file_path.name, source_format=source_format or 'docx')

def resolve_effective_legal_topic(*, parsed_section: ParsedSection, section: str, country: str) -> tuple[str, str | None]:
    """
    Return (document_type, legal_topic) for one parsed section - the
    single rule every caller that needs a section's effective topic
    must share (the chunk builder, and the admin section-editing
    service that reads the current DOCX as authority), so a section
    is never classified two different ways in two different places.

    Raises UnknownLegalTopicError for a heading-like section that
    resolves to neither the fixed taxonomy, a one-off subsection
    override, nor a parser-confirmed custom top-level topic - a real
    parser bug, never silently swallowed.
    """
    if is_overview_section(section=section, country=country):
        return ('overview', None)
    legal_topic = get_canonical_legal_topic(section=section, country=country) or get_subsection_topic_override(section)
    if legal_topic is None and parsed_section.is_custom_legal_topic:
        legal_topic = section
    if legal_topic is None:
        raise UnknownLegalTopicError(f'Unknown legal topic detected. Section: {section!r}. The document was not indexed because the topic is outside the approved taxonomy.')
    return ('comparator', legal_topic)

def build_document_chunks(parsed_sections: Sequence[ParsedSection], metadata: DocumentMetadata) -> list[DocumentChunk]:
    """Enrich parsed sections into indexable legal chunks."""
    _validate_metadata(metadata)
    country = metadata.country.strip()
    country_code = metadata.country_code.strip().upper()
    language = metadata.language.strip().lower()
    source_filename = metadata.source_filename.strip()
    source_format = metadata.source_format.strip().lower()
    document_id = _build_document_id(metadata)
    path_occurrences: defaultdict[tuple[str, str, str, str], int] = defaultdict(int)
    chunks: list[DocumentChunk] = []
    for parsed_section in parsed_sections:
        section = parsed_section.section.strip()
        subsection = parsed_section.subsection.strip() if parsed_section.subsection else None
        content = parsed_section.content.strip()
        if not content:
            continue
        document_type, legal_topic = resolve_effective_legal_topic(parsed_section=parsed_section, section=section, country=country)
        path_key = (document_type, legal_topic or '', section.casefold(), (subsection or '').casefold())
        path_occurrences[path_key] += 1
        occurrence = path_occurrences[path_key]
        chunk_id = _build_chunk_id(document_id=document_id, document_type=document_type, legal_topic=legal_topic, section=section, subsection=subsection, occurrence=occurrence)
        chunks.append(DocumentChunk(document_id=document_id, chunk_id=chunk_id, country=country, country_code=country_code, legal_topic=legal_topic, document_type=document_type, language=language, section=section, subsection=subsection, content=content, source_filename=source_filename, source_format=source_format, content_hash=_sha256(content), reference_year=metadata.reference_year))
    return chunks

def build_contact_chunk_for_contacts(contacts: Sequence[ExtractedContact], metadata: DocumentMetadata) -> DocumentChunk | None:
    """
    Build the one Contact-subsection DocumentChunk representing every
    contact in `contacts`, or None when there is nothing to index.

    Public and reusable: the ONE shared formatter/builder behind every
    Contact-chunk-producing path (initial DOCX parsing at upload time,
    below; Admin Contact CRUD synchronization; and Reindex when a
    structured contact state already exists - see admin_contacts.py) -
    never a second, separately maintained implementation (mission
    "ORDER 8G-B1", section 8). Reuses the same document_id/chunk_id
    scheme as every other chunk of this document, so it lives in the
    same OpenSearch mapping with no new field.
    """
    if not contacts:
        return None
    content = build_contact_chunk_content(contacts)
    if not content:
        return None
    country = metadata.country.strip()
    country_code = metadata.country_code.strip().upper()
    document_id = _build_document_id(metadata)
    section = f'Employment Law Overview {country}'
    chunk_id = _build_chunk_id(document_id=document_id, document_type='overview', legal_topic=None, section=section, subsection=CONTACT_SUBSECTION, occurrence=1)
    return DocumentChunk(document_id=document_id, chunk_id=chunk_id, country=country, country_code=country_code, legal_topic=None, document_type='overview', language=metadata.language.strip().lower(), section=section, subsection=CONTACT_SUBSECTION, content=content, source_filename=metadata.source_filename.strip(), source_format=metadata.source_format.strip().lower(), content_hash=_sha256(content), reference_year=metadata.reference_year)

def _build_contact_chunk(file_path: Path, metadata: DocumentMetadata) -> DocumentChunk | None:
    """
    Build one Contact-subsection chunk from a source DOCX, if it has
    a validated contact card.

    Returns None when no contact could be extracted, rather than
    indexing an empty placeholder.
    """
    contacts = extract_contacts_from_docx(file_path, country=metadata.country)
    return build_contact_chunk_for_contacts(contacts, metadata)

def build_document_chunks_from_docx(file_path: Path, country_code: str | None=None, language: str='en') -> list[DocumentChunk]:
    """Parse and enrich one L&E DOCX document."""
    validate_docx_format(file_path)
    metadata = metadata_from_content(file_path=file_path, country_code=country_code, language=language)
    parsed_sections = split_parsed_sections(parsed_sections=parse_docx_sections(file_path=file_path, country=metadata.country), max_chars=6000)
    chunks = build_document_chunks(parsed_sections=parsed_sections, metadata=metadata)
    contact_chunk = _build_contact_chunk(file_path=file_path, metadata=metadata)
    if contact_chunk is not None:
        chunks.append(contact_chunk)
    return chunks
DEFAULT_MAX_CHARS: Final[int] = 6000
_PARAGRAPH_BREAK_PATTERN: Final[re.Pattern[str]] = re.compile('\\n\\s*\\n+')
_SENTENCE_BREAK_PATTERN: Final[re.Pattern[str]] = re.compile('(?<=[.!?])\\s+(?=[A-Z0-9À-ÖØ-Þ\\"\'“‘(])')

class SectionSplitterError(ValueError):
    """Raised when a parsed section cannot be safely split."""

def _validate_max_chars(max_chars: int) -> None:
    """Validate the configured hard character limit."""
    if max_chars < 100:
        raise SectionSplitterError('max_chars must be at least 100.')

def _split_by_words(text: str, max_chars: int) -> list[str]:
    """
    Split text at whitespace boundaries.

    A single token longer than the hard limit, such as an abnormally
    long URL, is sliced as a final fallback.
    """
    words = text.split()
    if not words:
        return []
    chunks: list[str] = []
    current_words: list[str] = []

    def flush_current() -> None:
        if not current_words:
            return
        chunks.append(' '.join(current_words))
        current_words.clear()
    for word in words:
        if len(word) > max_chars:
            flush_current()
            for start_index in range(0, len(word), max_chars):
                chunks.append(word[start_index:start_index + max_chars])
            continue
        candidate = ' '.join((*current_words, word))
        if current_words and len(candidate) > max_chars:
            flush_current()
            current_words.append(word)
        else:
            current_words.append(word)
    flush_current()
    return chunks

def _pack_units(units: Sequence[str], separator: str, max_chars: int) -> list[str]:
    """Pack already bounded units without exceeding the hard limit."""
    chunks: list[str] = []
    current_units: list[str] = []

    def flush_current() -> None:
        if not current_units:
            return
        chunks.append(separator.join(current_units))
        current_units.clear()
    for raw_unit in units:
        unit = raw_unit.strip()
        if not unit:
            continue
        if len(unit) > max_chars:
            raise SectionSplitterError('Internal splitter error: an oversized unit reached the packing stage.')
        candidate = separator.join((*current_units, unit))
        if current_units and len(candidate) > max_chars:
            flush_current()
            current_units.append(unit)
        else:
            current_units.append(unit)
    flush_current()
    return chunks

def _split_oversized_paragraph(paragraph: str, max_chars: int) -> list[str]:
    """
    Split one oversized paragraph.

    Sentence boundaries are preferred. Whitespace splitting is used
    when a sentence itself exceeds the hard limit.
    """
    sentences = [sentence.strip() for sentence in _SENTENCE_BREAK_PATTERN.split(paragraph) if sentence.strip()]
    if len(sentences) <= 1:
        return _split_by_words(text=paragraph, max_chars=max_chars)
    bounded_sentences: list[str] = []
    for sentence in sentences:
        if len(sentence) <= max_chars:
            bounded_sentences.append(sentence)
        else:
            bounded_sentences.extend(_split_by_words(text=sentence, max_chars=max_chars))
    return _pack_units(units=bounded_sentences, separator=' ', max_chars=max_chars)

def split_text(text: str, max_chars: int=DEFAULT_MAX_CHARS) -> list[str]:
    """
    Split text deterministically while preserving its logical order.

    Priority:

    1. paragraph boundaries;
    2. sentence boundaries;
    3. whitespace boundaries;
    4. hard slicing for exceptionally long individual tokens.
    """
    _validate_max_chars(max_chars)
    normalized_text = text.strip()
    if not normalized_text:
        return []
    if len(normalized_text) <= max_chars:
        return [normalized_text]
    paragraphs = [paragraph.strip() for paragraph in _PARAGRAPH_BREAK_PATTERN.split(normalized_text) if paragraph.strip()]
    chunks: list[str] = []
    current_paragraphs: list[str] = []

    def flush_current_paragraphs() -> None:
        if not current_paragraphs:
            return
        chunks.append('\n\n'.join(current_paragraphs))
        current_paragraphs.clear()
    for paragraph in paragraphs:
        if len(paragraph) > max_chars:
            flush_current_paragraphs()
            chunks.extend(_split_oversized_paragraph(paragraph=paragraph, max_chars=max_chars))
            continue
        candidate = '\n\n'.join((*current_paragraphs, paragraph))
        if current_paragraphs and len(candidate) > max_chars:
            flush_current_paragraphs()
            current_paragraphs.append(paragraph)
        else:
            current_paragraphs.append(paragraph)
    flush_current_paragraphs()
    if any((len(chunk) > max_chars for chunk in chunks)):
        raise SectionSplitterError(f'The splitter produced a chunk above the configured limit of {max_chars} characters.')
    return chunks

def split_parsed_sections(parsed_sections: Sequence[ParsedSection], max_chars: int=DEFAULT_MAX_CHARS) -> list[ParsedSection]:
    """
    Split only oversized ParsedSection objects.

    The legal section and subsection labels remain unchanged. Multiple
    pieces with the same structural path are later assigned distinct,
    deterministic occurrences by document_chunk_builder.py.
    """
    _validate_max_chars(max_chars)
    split_sections: list[ParsedSection] = []
    for parsed_section in parsed_sections:
        content_parts = split_text(text=parsed_section.content, max_chars=max_chars)
        for content_part in content_parts:
            split_sections.append(ParsedSection(section=parsed_section.section, subsection=parsed_section.subsection, content=content_part, is_custom_legal_topic=parsed_section.is_custom_legal_topic))
    return split_sections
