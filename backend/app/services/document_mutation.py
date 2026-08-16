"""
Reusable primitives for mutating a copy of a source DOCX's top-level
legal topics.

This is the physical-edit layer behind ORDER 8A's "current DOCX is the
unique source of truth" contract: replacing an existing topic's content
or inserting a brand-new one. Callers are responsible for the actual
transaction (locking, snapshotting, temp-file placement, fsync,
reparse-validation, atomic replace, rollback) - this module only knows
how to produce a correctly-structured DOCX given the requested edit.

Built entirely on python-docx / the existing docx_parser heading rules
- there is no second parser implementation here, and no topic name is
ever hardcoded.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Final

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.text.paragraph import CT_P
from docx.shared import Pt
from docx.styles.style import BaseStyle
from docx.text.paragraph import Paragraph

from app.services.docx_parser import (
    ADMIN_SECTION_STYLE_NAME,
    TopicLocation,
    locate_top_level_topics,
)

_BLANK_LINE_PATTERN: Final[re.Pattern[str]] = re.compile(r"\n{2,}")

_PREFIX_NUMBER_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"^\s*(\d{1,2})\s*[.)]\s*"
)

_DEFAULT_ADMIN_SECTION_FONT_SIZE: Final[Pt] = Pt(14)

_ADMIN_SECTION_OUTLINE_LEVEL: Final[int] = 0


class DocxMutationError(RuntimeError):
    """Base error for a DOCX mutation that could not be performed safely."""


class LegalTopicNotFoundError(DocxMutationError):
    """The target legal topic is not present exactly once in the DOCX."""

    def __init__(self, legal_topic: str) -> None:
        super().__init__(
            f"Legal topic not found (or not unique): {legal_topic!r}"
        )
        self.legal_topic = legal_topic


class LegalTopicAlreadyExistsError(DocxMutationError):
    """A new topic's title collides with an existing top-level topic."""

    def __init__(self, title: str) -> None:
        super().__init__(
            f"A top-level topic already exists with this title: {title!r}"
        )
        self.title = title


class InvalidSectionPositionError(DocxMutationError):
    """An Add position reference does not resolve unambiguously."""


class NoAnchorTopicError(DocxMutationError):
    """There is no existing top-level topic to model a new heading on."""


def normalize_topic_title(
    title: str,
) -> str:
    """
    Normalize a topic title for duplicate-comparison purposes: trim,
    collapse internal whitespace, casefold, and ignore any leading
    manually-typed numbering prefix (e.g. "01. ") so "Hiring Practices"
    and "01.  Hiring   Practices" compare as the same title.
    """

    collapsed = " ".join(
        title.split()
    )

    unprefixed = _PREFIX_NUMBER_PATTERN.sub(
        "",
        collapsed,
    ).strip()

    return (unprefixed or collapsed).casefold()


def _split_into_paragraph_blocks(
    text: str,
) -> list[str]:
    """
    Split plain text into paragraph blocks on blank lines (two or more
    consecutive newlines), preserving Unicode content and any single
    embedded newlines within a block as soft line breaks.
    """

    normalized = (
        text
        .replace("\r\n", "\n")
        .replace("\r", "\n")
    )

    blocks = _BLANK_LINE_PATTERN.split(
        normalized
    )

    stripped = [
        block.strip("\n")
        for block in blocks
    ]

    non_empty = [
        block
        for block in stripped
        if block.strip()
    ]

    return non_empty or [""]


def _new_body_paragraph_elements(
    document: DocxDocument,
    text: str,
) -> list[CT_P]:
    """
    Build detached body-paragraph XML elements from plain text: one
    paragraph per blank-line-delimited block, with any single embedded
    newline preserved as a soft line break within that same paragraph.
    Never interprets the text as markup.
    """

    elements: list[CT_P] = []

    for block in _split_into_paragraph_blocks(text):
        paragraph = document.add_paragraph()

        lines = block.split("\n")

        for index, line in enumerate(lines):
            if index > 0:
                paragraph.add_run().add_break()

            if line:
                paragraph.add_run(line)

        document.element.body.remove(paragraph._p)
        elements.append(paragraph._p)

    return elements


def _find_topic(
    locations: list[TopicLocation],
    legal_topic: str,
) -> TopicLocation:
    matches = [
        location
        for location in locations
        if location.legal_topic == legal_topic
    ]

    if len(matches) != 1:
        raise LegalTopicNotFoundError(
            legal_topic
        )

    return matches[0]


def replace_top_level_topic(
    file_path: Path,
    output_path: Path,
    country: str,
    legal_topic: str,
    new_content: str,
) -> None:
    """
    Replace ALL content of an existing top-level legal topic with new
    plain-text content, and save the result to output_path.

    The topic's own heading paragraph is left completely unchanged;
    only the body elements between it and the next top-level heading
    are removed and replaced. Everything else in the document (front
    matter, other topics, country identity, title, media, headers and
    footers, document properties) is untouched.
    """

    file_path = file_path.resolve()

    document = Document(
        file_path
    )

    locations = locate_top_level_topics(
        document,
        country=country,
    )

    target = _find_topic(
        locations,
        legal_topic,
    )

    new_elements = _new_body_paragraph_elements(
        document,
        new_content,
    )

    anchor = target.heading_element

    for old_element in target.body_elements:
        old_element.getparent().remove(
            old_element
        )

    for new_element in new_elements:
        anchor.addnext(
            new_element
        )
        anchor = new_element

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        output_path
    )


def rename_top_level_topic(
    file_path: Path,
    output_path: Path,
    country: str,
    legal_topic: str,
    new_title: str,
    new_content: str | None = None,
) -> None:
    """
    Rename an existing top-level legal topic's heading text - and,
    optionally, replace its content too - and save the result to
    output_path.

    The heading paragraph's own style/formatting is left completely
    unchanged; only its visible run text is replaced. This is what
    lets a renamed CANONICAL heading still be recognized as a valid
    (now custom-labelled) top-level topic afterward by the same
    structural signals the document's other canonical headings
    already carry (see docx_parser._learn_custom_topic_signal_
    requirement) - never by reassigning it a different style. A
    previously admin-added heading simply stays on
    ADMIN_SECTION_STYLE_NAME, exactly as before.
    """

    file_path = file_path.resolve()

    document = Document(
        file_path
    )

    locations = locate_top_level_topics(
        document,
        country=country,
    )

    target = _find_topic(
        locations,
        legal_topic,
    )

    heading_paragraph = Paragraph(
        target.heading_element,
        document,
    )

    for run in list(heading_paragraph.runs):
        run._r.getparent().remove(
            run._r
        )

    heading_paragraph.add_run(
        new_title
    )

    if new_content is not None:
        new_elements = _new_body_paragraph_elements(
            document,
            new_content,
        )

        anchor = target.heading_element

        for old_element in target.body_elements:
            old_element.getparent().remove(
                old_element
            )

        for new_element in new_elements:
            anchor.addnext(
                new_element
            )
            anchor = new_element

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        output_path
    )


def remove_top_level_topic(
    file_path: Path,
    output_path: Path,
    country: str,
    legal_topic: str,
) -> None:
    """
    Remove an existing top-level legal topic entirely - its heading
    and every body element (paragraphs, tables) between it and the
    next top-level heading - and save the result to output_path.

    Uses the exact same structural boundary (TopicLocation) that Edit
    already relies on, so no separate boundary-detection logic exists
    for Delete. Everything else in the document - other sections,
    their order, front matter, styles - is left completely untouched.
    """

    file_path = file_path.resolve()

    document = Document(
        file_path
    )

    locations = locate_top_level_topics(
        document,
        country=country,
    )

    target = _find_topic(
        locations,
        legal_topic,
    )

    for element in (
        target.heading_element,
        *target.body_elements,
    ):
        element.getparent().remove(
            element
        )

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        output_path
    )


def _derive_admin_section_font(
    document: DocxDocument,
    reference_paragraph: Paragraph | None,
) -> tuple[bool, Pt]:
    """
    Best-effort visual formatting for a brand-new admin section
    heading (ORDER 8A-C, section 3): "reasonable" bold/size derived
    from the nearest real top-level heading when safe, falling back to
    the document's own built-in Heading 1 style, falling back to a
    restrained standard heading appearance. This is cosmetic only -
    the marker STYLE NAME is the sole, deterministic identity a
    reparse relies on, never this formatting.
    """

    if reference_paragraph is not None:
        style_font = (
            reference_paragraph.style.font
            if reference_paragraph.style is not None
            else None
        )

        if style_font is not None and (
            style_font.bold is not None
            or style_font.size is not None
        ):
            return (
                style_font.bold if style_font.bold is not None else True,
                style_font.size or _DEFAULT_ADMIN_SECTION_FONT_SIZE,
            )

        for run in reference_paragraph.runs:
            if run.text.strip():
                return (
                    run.font.bold if run.font.bold is not None else True,
                    run.font.size or _DEFAULT_ADMIN_SECTION_FONT_SIZE,
                )

    try:
        heading_one = document.styles["Heading 1"]

    except KeyError:
        return True, _DEFAULT_ADMIN_SECTION_FONT_SIZE

    return (
        heading_one.font.bold if heading_one.font.bold is not None else True,
        heading_one.font.size or _DEFAULT_ADMIN_SECTION_FONT_SIZE,
    )


def _get_or_create_admin_section_style(
    document: DocxDocument,
    reference_paragraph: Paragraph | None,
) -> BaseStyle:
    """
    Get or create the dedicated ADMIN-added-section marker style
    (ORDER 8A-C, section 2) - a real, persistent Word paragraph style
    embedded in the DOCX itself, no external state of any kind. Reused
    as-is on every subsequent Add in the same document; created once,
    with "reasonable" visual formatting, the first time.
    """

    styles = document.styles

    try:
        return styles[ADMIN_SECTION_STYLE_NAME]

    except KeyError:
        pass

    style = styles.add_style(
        ADMIN_SECTION_STYLE_NAME,
        WD_STYLE_TYPE.PARAGRAPH,
    )

    bold, size = _derive_admin_section_font(
        document,
        reference_paragraph,
    )

    style.font.bold = bold
    style.font.size = size

    return style


def _set_outline_level(
    paragraph: Paragraph,
    level: int,
) -> None:
    """
    Mark a paragraph as top-level in Word's own outline (navigation
    pane / table of contents) - cosmetic, never relied on for parsing
    identity (the marker style name is).
    """

    paragraph_properties = paragraph._p.get_or_add_pPr()
    paragraph_properties.get_or_add_outlineLvl().val = level


def insert_top_level_topic(
    file_path: Path,
    output_path: Path,
    country: str,
    title: str,
    content: str,
    position: str,
) -> None:
    """
    Insert a brand-new top-level legal topic into a DOCX, and save the
    result to output_path.

    position is one of "beginning", "end", or "after:<legal_topic>".
    The new heading uses the dedicated ADMIN-section marker style
    (ORDER 8A-C, section 2), created once per document and reused
    afterward - this works identically regardless of whether the
    surrounding document uses Heading 1, bold-only headings, or any
    other native convention, since the parser's own recognition of
    this style never depends on the document's native structure.
    """

    file_path = file_path.resolve()

    document = Document(
        file_path
    )

    locations = locate_top_level_topics(
        document,
        country=country,
    )

    if not locations:
        raise NoAnchorTopicError(
            "No existing top-level topic to model a new heading on."
        )

    normalized_title = normalize_topic_title(
        title
    )

    for location in locations:
        if (
            normalize_topic_title(
                location.legal_topic
            )
            == normalized_title
        ):
            raise LegalTopicAlreadyExistsError(
                title
            )

    if position == "beginning":
        anchor_location = locations[0]
        insert_after_element = None
        insert_before_element = anchor_location.heading_element

    elif position == "end":
        anchor_location = locations[-1]
        trailing_elements = anchor_location.body_elements
        insert_after_element = (
            trailing_elements[-1]
            if trailing_elements
            else anchor_location.heading_element
        )
        insert_before_element = None

    elif position.startswith(
        "after:"
    ):
        target_topic = position[
            len("after:") :
        ]

        anchor_location = _find_topic(
            locations,
            target_topic,
        )

        trailing_elements = anchor_location.body_elements
        insert_after_element = (
            trailing_elements[-1]
            if trailing_elements
            else anchor_location.heading_element
        )
        insert_before_element = None

    else:
        raise InvalidSectionPositionError(
            f"Unsupported position: {position!r}"
        )

    anchor_paragraph = Paragraph(
        anchor_location.heading_element,
        document,
    )

    admin_section_style = _get_or_create_admin_section_style(
        document,
        reference_paragraph=anchor_paragraph,
    )

    new_heading_paragraph = document.add_paragraph(
        title,
        style=admin_section_style.name,
    )

    _set_outline_level(
        new_heading_paragraph,
        _ADMIN_SECTION_OUTLINE_LEVEL,
    )

    new_heading_element = new_heading_paragraph._p

    document.element.body.remove(
        new_heading_element
    )

    new_content_elements = _new_body_paragraph_elements(
        document,
        content,
    )

    if insert_before_element is not None:
        insert_before_element.addprevious(
            new_heading_element
        )
        anchor = new_heading_element

        for element in new_content_elements:
            anchor.addnext(
                element
            )
            anchor = element

    else:
        assert insert_after_element is not None

        anchor = insert_after_element
        anchor.addnext(
            new_heading_element
        )
        anchor = new_heading_element

        for element in new_content_elements:
            anchor.addnext(
                element
            )
            anchor = element

    output_path.parent.mkdir(
        parents=True,
        exist_ok=True,
    )

    document.save(
        output_path
    )
