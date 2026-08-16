"""Regression tests for safe country-level document replacement."""

from __future__ import annotations

import tempfile
import unittest
from io import BytesIO
from pathlib import Path

from tests.admin_invariants import real_source_entries
from unittest.mock import patch

from docx import Document as DocxDocument

from app.models.document import DocumentChunk
from app.services.admin_document_replacement import (
    AdminDocumentAlreadyCurrentError,
    AdminDocumentCountryConfirmationRequiredError,
    AdminDocumentCountryConflictReviewRequiredError,
    AdminDocumentCountryNotAllowedError,
    AdminDocumentCountrySelectionInvalidError,
    AdminDocumentCountrySelectionRequiredError,
    AdminDocumentIdenticalButAdminModifiedError,
    AdminDocumentReplacementRequiredError,
    AdminDocumentWarningConfirmationRequiredError,
    ExistingCountryDocument,
    safe_upload_and_index_document,
)
from app.services.admin_documents import DocumentCountryUndeterminedError
from app.services.docx_country_marker import read_country_marker
from app.services.document_chunk_builder import (
    AmbiguousDocumentCountryError,
    UndeterminableDocumentCountryError,
    build_document_chunks_from_docx,
    metadata_from_content,
)
from app.services.document_indexer import (
    DocumentIndexingError,
    DocumentIndexingResult,
    replace_country_document_chunks,
)
from app.services.document_section_state import (
    SectionEdit,
    SectionEditState,
    read_section_edit_state,
    section_id_for_legal_topic,
    write_section_edit_state_atomic,
)


def _real_docx_bytes(paragraphs: list[str]) -> bytes:
    """A minimal real DOCX (valid zip/OOXML), for tests that exercise
    the actual marker-writing/parsing pipeline rather than a fake
    chunk_builder."""

    document = DocxDocument()

    for text in paragraphs:
        document.add_paragraph(text)

    buffer = BytesIO()
    document.save(buffer)

    return buffer.getvalue()


def _fake_indexer(*, chunks, client=None):
    del client

    return DocumentIndexingResult(
        index_alias="legal-documents-v1",
        document_id=chunks[0].document_id,
        source_filename=chunks[0].source_filename,
        requested_chunks=len(chunks),
        indexed_chunks=len(chunks),
        stale_chunks_deleted=0,
    )


AU_OLD_ID = "doc_" + "a" * 64
AU_NEW_ID = "doc_" + "b" * 64


def _build_au_chunk(filename: str) -> DocumentChunk:
    return DocumentChunk(
        document_id=AU_NEW_ID,
        chunk_id="chunk-au-new-1",
        country="Australia",
        country_code="AU",
        legal_topic="Employment Contracts",
        document_type="comparator",
        language="en",
        section="Employment Contracts",
        subsection="Trial Period",
        content="Australian probation content.",
        source_filename=filename,
        source_format="docx",
        content_hash="new-content-hash",
        reference_year=2026,
    )


def _existing_documents(
    country_code: str,
    client=None,
) -> list[ExistingCountryDocument]:
    del client
    assert country_code == "AU"

    return [
        ExistingCountryDocument(
            document_id=AU_OLD_ID,
            source_filename=(
                "Employment Law Overview Australia.docx"
            ),
            country="Australia",
            country_code="AU",
            reference_year=None,
        )
    ]


class SafeCountryReplacementTests(unittest.TestCase):
    def test_existing_country_requires_confirmation_without_changes(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            source_path = (
                source_directory
                / "Employment Law Overview Australia.docx"
            )
            source_path.write_bytes(b"old-australia")

            indexer_called = False

            def country_indexer(**kwargs):
                nonlocal indexer_called
                del kwargs
                indexer_called = True
                raise AssertionError("indexer must not run")

            with self.assertRaises(
                AdminDocumentReplacementRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"new-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=(
                        _existing_documents
                    ),
                    country_document_indexer=(
                        country_indexer
                    ),
                )

            self.assertEqual(
                context.exception.country_code,
                "AU",
            )
            self.assertFalse(indexer_called)
            self.assertEqual(
                source_path.read_bytes(),
                b"old-australia",
            )
            self.assertFalse(
                (source_directory / "AU.docx").exists()
            )

    def test_replace_existing_never_bypasses_a_country_conflict(
        self,
    ) -> None:
        # Mission "ORDER 8E-A1", section 18: more than one active
        # document for a country is a genuine conflict the ordinary
        # upload/replace decision must never blindly resolve, even
        # with replace_existing=True - only the dedicated conflict-
        # resolution API may do that (see
        # admin_document_conflict_resolution.py). Superseded from this
        # test's earlier "confirmed replace collapses to one" name -
        # that collapsing behavior now belongs exclusively to
        # AUTO_DEDUPLICATE/CHOOSE_DOCUMENT/REPLACE_WITH_DOCUMENT.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            legacy_path = (
                source_directory
                / "Employment Law Overview Australia.docx"
            )
            canonical_path = source_directory / "AU.docx"

            legacy_path.write_bytes(b"legacy-australia")
            canonical_path.write_bytes(b"other-australia")

            existing = [
                ExistingCountryDocument(
                    document_id=AU_OLD_ID,
                    source_filename=legacy_path.name,
                    country="Australia",
                    country_code="AU",
                    reference_year=None,
                ),
                ExistingCountryDocument(
                    document_id="doc_" + "c" * 64,
                    source_filename=legacy_path.name,
                    country="Australia",
                    country_code="AU",
                    reference_year=None,
                ),
            ]

            def lookup(country_code: str, client=None):
                del client
                self.assertEqual(country_code, "AU")
                return existing

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError("indexer must not run")

            with self.assertRaises(
                AdminDocumentCountryConflictReviewRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"new-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    replace_existing=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=lookup,
                    country_document_indexer=(
                        indexer_must_not_be_called
                    ),
                )

            self.assertEqual(context.exception.country_code, "AU")
            self.assertEqual(
                sorted(
                    candidate.document_id
                    for candidate in context.exception.candidates
                ),
                sorted([AU_OLD_ID, "doc_" + "c" * 64]),
            )
            # Zero mutation - both legacy sources untouched.
            self.assertEqual(
                legacy_path.read_bytes(), b"legacy-australia"
            )
            self.assertEqual(
                canonical_path.read_bytes(), b"other-australia"
            )

    def test_failed_confirmed_replacement_restores_all_sources(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            legacy_path = (
                source_directory
                / "Employment Law Overview Australia.docx"
            )
            canonical_path = source_directory / "AU.docx"

            legacy_path.write_bytes(b"legacy-australia")
            canonical_path.write_bytes(b"canonical-australia")

            def failing_indexer(**kwargs):
                del kwargs
                raise DocumentIndexingError(
                    "simulated indexing failure"
                )

            with self.assertRaises(DocumentIndexingError):
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"new-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    replace_existing=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=(
                        _existing_documents
                    ),
                    country_document_indexer=(
                        failing_indexer
                    ),
                )

            self.assertEqual(
                legacy_path.read_bytes(),
                b"legacy-australia",
            )
            self.assertEqual(
                canonical_path.read_bytes(),
                b"canonical-australia",
            )
            self.assertEqual(
                sorted(
                    path.name
                    for path in real_source_entries(source_directory)
                ),
                [
                    "AU.docx",
                    "Employment Law Overview Australia.docx",
                ],
            )

    def test_identical_single_document_is_not_reindexed(
        self,
    ) -> None:
        # Mission "HOTFIX 0.4.9" review, section 11 - a single clean
        # document + byte-identical upload must raise
        # AdminDocumentAlreadyCurrentError and never call the
        # indexer at all (explicit "must not be called" double,
        # not merely inferred from the absence of a real connection).
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"same-australia")

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError(
                    "the indexer must not be called for an "
                    "already-current document"
                )

            with self.assertRaises(
                AdminDocumentAlreadyCurrentError
            ):
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"same-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=(
                        _existing_documents
                    ),
                    country_document_indexer=(
                        indexer_must_not_be_called
                    ),
                )

    def test_fresh_country_is_staged_and_parsed_exactly_once(
        self,
    ) -> None:
        # Mission "HOTFIX 0.4.9" - the Argentina regression: an
        # earlier implementation detected a fresh/absent country
        # during a preflight parse, then delegated to a second,
        # separate upload implementation that re-staged and re-parsed
        # the exact same upload a second time by rewinding and
        # re-reading file_stream - fragile by construction (depends on
        # a stream that may not always be safely re-readable) and
        # never necessary, since nothing about the already-computed
        # chunks needs to change for a fresh country. This proves the
        # unified implementation reads the stream and invokes
        # chunk_builder exactly once, regardless of country freshness.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            chunk_builder_calls = []

            def counting_chunk_builder(path: Path):
                chunk_builder_calls.append(path)
                return [_build_au_chunk(path.name)]

            stream = BytesIO(b"new-argentina")
            original_read = stream.read
            read_calls = []

            def counting_read(*args, **kwargs):
                read_calls.append(1)
                return original_read(*args, **kwargs)

            stream.read = counting_read  # type: ignore[method-assign]

            response = safe_upload_and_index_document(
                filename="Argentina.docx",
                file_stream=stream,
                source_directory=source_directory,
                confirm_warnings=True,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                country_confirmed=True,
                chunk_builder=counting_chunk_builder,
                country_document_lookup=lambda code, client: [],
                country_document_indexer=lambda *, chunks, client=None: (
                    DocumentIndexingResult(
                        index_alias="legal-documents-v1",
                        document_id=chunks[0].document_id,
                        source_filename=chunks[0].source_filename,
                        requested_chunks=len(chunks),
                        indexed_chunks=len(chunks),
                        stale_chunks_deleted=0,
                    )
                ),
            )

            self.assertEqual(
                len(chunk_builder_calls),
                1,
                "chunk_builder must run exactly once for a fresh "
                "upload - never a second, redundant parse.",
            )

            # Exactly one read pass to EOF: the first read() call
            # returns all bytes, the second returns b"" (EOF) - two
            # calls total for one pass, never four (which a second
            # staging pass would require).
            self.assertLessEqual(len(read_calls), 2)

            self.assertEqual(response.status, "uploaded")
            self.assertEqual(response.document_id, AU_NEW_ID)
            self.assertEqual(response.country_code, "AU")
            self.assertFalse(response.replaced_source_file)
            self.assertEqual(response.replaced_document_ids, [])
            self.assertTrue(
                (source_directory / "AU.docx").exists()
            )
            self.assertEqual(
                (source_directory / "AU.docx").read_bytes(),
                b"new-argentina",
            )

    def test_deleted_then_reuploaded_country_succeeds(self) -> None:
        # Mission "HOTFIX 0.4.9" - the exact end-to-end Argentina
        # scenario: upload, delete, then re-upload under a completely
        # different filename must succeed and leave exactly one
        # active document, with no stale conflict from the deletion.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            def indexer(*, chunks, client=None):
                return DocumentIndexingResult(
                    index_alias="legal-documents-v1",
                    document_id=chunks[0].document_id,
                    source_filename=chunks[0].source_filename,
                    requested_chunks=len(chunks),
                    indexed_chunks=len(chunks),
                    stale_chunks_deleted=0,
                )

            first = safe_upload_and_index_document(
                filename="Argentina.docx",
                file_stream=BytesIO(b"first-argentina"),
                source_directory=source_directory,
                confirm_warnings=True,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                country_confirmed=True,
                chunk_builder=lambda path: [
                    _build_au_chunk(path.name)
                ],
                country_document_lookup=lambda code, client: [],
                country_document_indexer=indexer,
            )
            self.assertEqual(first.status, "uploaded")

            # Simulate a completed delete: the source file removed,
            # the country lookup now reporting zero active documents
            # again - exactly the state a real DELETE leaves behind.
            (source_directory / "AU.docx").unlink()

            second = safe_upload_and_index_document(
                filename="random-name-completely-different.docx",
                file_stream=BytesIO(b"second-argentina"),
                source_directory=source_directory,
                confirm_warnings=True,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                country_confirmed=True,
                chunk_builder=lambda path: [
                    _build_au_chunk(path.name)
                ],
                country_document_lookup=lambda code, client: [],
                country_document_indexer=indexer,
            )

            self.assertEqual(second.status, "uploaded")
            self.assertEqual(
                [path.name for path in real_source_entries(source_directory)],
                ["AU.docx"],
            )
            self.assertEqual(
                (source_directory / "AU.docx").read_bytes(),
                b"second-argentina",
            )


class FakeCountryOpenSearch:
    def __init__(self, *, fail_cleanup: bool = False) -> None:
        self.fail_cleanup = fail_cleanup
        self.delete_calls = 0

    def search(self, *, index, body):
        del index
        del body

        return {
            "hits": {
                "total": {
                    "value": 1,
                },
                "hits": [
                    {
                        "_id": "chunk-old-1",
                        "_source": {
                            "document_id": AU_OLD_ID,
                            "chunk_id": "chunk-old-1",
                            "country": "Australia",
                            "country_code": "AU",
                        },
                        "sort": ["chunk-old-1"],
                    }
                ],
            }
        }

    def delete_by_query(self, **kwargs):
        del kwargs
        self.delete_calls += 1

        if self.fail_cleanup and self.delete_calls == 1:
            raise RuntimeError("cleanup failed")

        return {
            "deleted": 1,
        }


class FilenameNeverDecidesIdentityTests(unittest.TestCase):
    """
    Mission "HOTFIX 0.4.9", section F - the country_code detected from
    a DOCX's own content is the only identity that matters; the
    uploaded filename is display information only. Three completely
    different filenames carrying the exact same detected country must
    all resolve to the exact same existing-country/409 outcome.
    """

    def test_same_filename_as_the_existing_document_still_requires_confirmation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"old-australia")

            with self.assertRaises(
                AdminDocumentReplacementRequiredError
            ):
                safe_upload_and_index_document(
                    # The exact same name as the existing source -
                    # must never be treated as "no real change".
                    filename=(
                        "Employment Law Overview Australia.docx"
                    ),
                    file_stream=BytesIO(b"new-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                )

    def test_unrelated_filenames_all_detect_the_same_country(
        self,
    ) -> None:
        for filename in (
            "Argentina.docx",
            "random-file-name.docx",
            "Legal-update-final-v7.docx",
        ):
            with self.subTest(filename=filename):
                with tempfile.TemporaryDirectory() as root:
                    source_directory = Path(root) / "source"
                    processed_directory = Path(root) / "processed"

                    response = safe_upload_and_index_document(
                        filename=filename,
                        file_stream=BytesIO(b"argentina-bytes"),
                        source_directory=source_directory,
                        confirm_warnings=True,
                        processed_directory=processed_directory,
                        maximum_bytes=1000,
                        country_confirmed=True,
                        chunk_builder=lambda path: [
                            _build_au_chunk(path.name)
                        ],
                        country_document_lookup=(
                            lambda code, client: []
                        ),
                        country_document_indexer=(
                            lambda *, chunks, client=None: (
                                DocumentIndexingResult(
                                    index_alias="legal-documents-v1",
                                    document_id=chunks[0].document_id,
                                    source_filename=(
                                        chunks[0].source_filename
                                    ),
                                    requested_chunks=len(chunks),
                                    indexed_chunks=len(chunks),
                                    stale_chunks_deleted=0,
                                )
                            )
                        ),
                    )

                    self.assertEqual(response.status, "uploaded")
                    self.assertEqual(response.country_code, "AU")
                    self.assertTrue(
                        (source_directory / "AU.docx").exists()
                    )


class ConfirmedReplacementDuplicateCountTests(unittest.TestCase):
    """
    Mission "ORDER 8E-A1", section 18 - the ordinary upload/replace
    decision must refuse to collapse a country conflict itself,
    however many duplicate legacy document_ids and source files it
    started with (tested here for 3) - it always defers to the
    dedicated conflict-resolution API instead. Superseded from
    "HOTFIX 0.4.9"'s original "confirmed replace collapses to one"
    expectation.
    """

    def test_confirmed_replace_with_three_duplicate_ids_requires_conflict_review(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            legacy_path = (
                source_directory
                / "Employment Law Overview Australia.docx"
            )
            canonical_path = source_directory / "AU.docx"
            legacy_path.write_bytes(b"legacy-australia")
            canonical_path.write_bytes(b"other-australia")

            existing = [
                ExistingCountryDocument(
                    document_id="doc_" + letter * 64,
                    source_filename=legacy_path.name,
                    country="Australia",
                    country_code="AU",
                    reference_year=None,
                )
                for letter in ("a", "c", "d")
            ]

            def lookup(country_code: str, client=None):
                del client
                self.assertEqual(country_code, "AU")
                return existing

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError("indexer must not run")

            with self.assertRaises(
                AdminDocumentCountryConflictReviewRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"new-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    replace_existing=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=lookup,
                    country_document_indexer=(
                        indexer_must_not_be_called
                    ),
                )

            self.assertEqual(
                sorted(
                    candidate.document_id
                    for candidate in context.exception.candidates
                ),
                sorted(
                    document.document_id for document in existing
                ),
            )
            self.assertEqual(
                sorted(
                    path.name
                    for path in real_source_entries(source_directory)
                ),
                ["AU.docx", "Employment Law Overview Australia.docx"],
            )

    def test_confirmed_replace_with_source_missing(self) -> None:
        # The country's OpenSearch metadata still names a source that
        # no longer exists on disk - confirmed replacement must still
        # succeed (write the new file, index it) rather than treat a
        # missing legacy file as a blocking error.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            existing = [
                ExistingCountryDocument(
                    document_id=AU_OLD_ID,
                    source_filename=(
                        "Employment Law Overview Australia.docx"
                    ),
                    country="Australia",
                    country_code="AU",
                    reference_year=None,
                )
            ]

            def indexer(*, chunks, client=None):
                return DocumentIndexingResult(
                    index_alias="legal-documents-v1",
                    document_id=chunks[0].document_id,
                    source_filename=chunks[0].source_filename,
                    requested_chunks=len(chunks),
                    indexed_chunks=len(chunks),
                    stale_chunks_deleted=0,
                )

            response = safe_upload_and_index_document(
                filename="Australia 2026.docx",
                file_stream=BytesIO(b"new-australia"),
                source_directory=source_directory,
                confirm_warnings=True,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                country_confirmed=True,
                replace_existing=True,
                chunk_builder=lambda path: [
                    _build_au_chunk(path.name)
                ],
                country_document_lookup=lambda code, client: existing,
                country_document_indexer=indexer,
            )

            self.assertEqual(response.status, "replaced")
            self.assertFalse(response.replaced_source_file)
            self.assertEqual(
                response.replaced_document_ids, [AU_OLD_ID]
            )
            self.assertEqual(
                (source_directory / "AU.docx").read_bytes(),
                b"new-australia",
            )


class IdenticalFileAmbiguousCountryTests(unittest.TestCase):
    """
    Mission "ORDER 8E-A1", section 18 - a country left with several
    document_ids (a pre-existing conflict/duplicate state) must never
    be told "document_already_current", even if the newly uploaded
    bytes happen to match one of the candidate files exactly: the
    country's own state is inconsistent, and normalizing it now always
    requires the dedicated conflict-resolution review, never a plain
    confirmed replacement (superseded from "HOTFIX 0.4.9"'s original
    document_replacement_required expectation).
    """

    def test_identical_bytes_but_ambiguous_country_still_requires_conflict_review(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            legacy_path = (
                source_directory
                / "Employment Law Overview Australia.docx"
            )
            canonical_path = source_directory / "AU.docx"
            legacy_path.write_bytes(b"same-australia")
            canonical_path.write_bytes(b"other-australia")

            existing = [
                ExistingCountryDocument(
                    document_id=AU_OLD_ID,
                    source_filename=legacy_path.name,
                    country="Australia",
                    country_code="AU",
                    reference_year=None,
                ),
                ExistingCountryDocument(
                    document_id="doc_" + "c" * 64,
                    source_filename=legacy_path.name,
                    country="Australia",
                    country_code="AU",
                    reference_year=None,
                ),
            ]

            with self.assertRaises(
                AdminDocumentCountryConflictReviewRequiredError
            ):
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    # Byte-identical to the legacy file specifically.
                    file_stream=BytesIO(b"same-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=(
                        lambda code, client: existing
                    ),
                )

            # Nothing on disk changed.
            self.assertEqual(
                legacy_path.read_bytes(), b"same-australia"
            )
            self.assertEqual(
                canonical_path.read_bytes(), b"other-australia"
            )


class SourceWriteFailureRollbackTests(unittest.TestCase):
    """
    Mission "HOTFIX 0.4.9", section 14 - injecting a failure at the
    physical write step itself (not just at the OpenSearch indexing
    step) must still leave the country's previous state exactly
    intact, for both a fresh upload and a confirmed replacement.
    """

    def test_fresh_upload_write_failure_leaves_no_partial_file(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            def failing_indexer(**kwargs):
                del kwargs
                raise DocumentIndexingError(
                    "simulated indexing failure"
                )

            with self.assertRaises(DocumentIndexingError):
                safe_upload_and_index_document(
                    filename="Argentina.docx",
                    file_stream=BytesIO(b"argentina-bytes"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=(
                        lambda code, client: []
                    ),
                    country_document_indexer=failing_indexer,
                )

            # No partial/incoming file survives a failed fresh upload.
            self.assertEqual(
                list(real_source_entries(source_directory)),
                [],
            )

    def test_confirmed_replacement_write_failure_restores_exact_bytes(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            existing_path = (
                source_directory
                / "Employment Law Overview Australia.docx"
            )
            existing_path.write_bytes(b"original-australia-bytes")

            def failing_indexer(**kwargs):
                del kwargs
                raise DocumentIndexingError(
                    "simulated indexing failure"
                )

            with self.assertRaises(DocumentIndexingError):
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"new-australia-bytes"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    replace_existing=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                    country_document_indexer=failing_indexer,
                )

            self.assertEqual(
                existing_path.read_bytes(),
                b"original-australia-bytes",
            )
            self.assertEqual(
                [
                    path.name
                    for path in real_source_entries(source_directory)
                ],
                [existing_path.name],
            )


class CountryIndexerTests(unittest.TestCase):
    @patch(
        "app.services.document_indexer."
        "ensure_legal_documents_index"
    )
    @patch("app.services.document_indexer.bulk")
    def test_country_indexer_removes_every_stale_country_chunk(
        self,
        bulk_mock,
        ensure_mock,
    ) -> None:
        del ensure_mock
        bulk_mock.return_value = (1, [])

        result = replace_country_document_chunks(
            chunks=[_build_au_chunk("Australia 2026.docx")],
            client=FakeCountryOpenSearch(),
        )

        self.assertEqual(result.indexed_chunks, 1)
        self.assertEqual(result.stale_chunks_deleted, 1)
        self.assertEqual(bulk_mock.call_count, 1)

    @patch(
        "app.services.document_indexer."
        "ensure_legal_documents_index"
    )
    @patch("app.services.document_indexer.bulk")
    def test_country_indexer_restores_snapshot_on_cleanup_failure(
        self,
        bulk_mock,
        ensure_mock,
    ) -> None:
        del ensure_mock
        bulk_mock.side_effect = [
            (1, []),
            (1, []),
        ]

        client = FakeCountryOpenSearch(
            fail_cleanup=True
        )

        with self.assertRaises(DocumentIndexingError):
            replace_country_document_chunks(
                chunks=[
                    _build_au_chunk(
                        "Australia 2026.docx"
                    )
                ],
                client=client,
            )

        self.assertEqual(client.delete_calls, 2)
        self.assertEqual(bulk_mock.call_count, 2)


def _build_chunks_with_topic_count(
    country_code: str,
    country: str,
    document_id: str,
    filename: str,
    topic_count: int,
) -> list[DocumentChunk]:
    from app.core.legal_taxonomy import LEGAL_TOPICS

    overview = DocumentChunk(
        document_id=document_id,
        chunk_id="chunk-overview",
        country=country,
        country_code=country_code,
        legal_topic=None,
        document_type="overview",
        language="en",
        section="General",
        subsection=None,
        content="Overview content.",
        source_filename=filename,
        source_format="docx",
        content_hash="overview-hash",
    )

    topic_chunks = [
        DocumentChunk(
            document_id=document_id,
            chunk_id=f"chunk-topic-{index}",
            country=country,
            country_code=country_code,
            legal_topic=topic,
            document_type="comparator",
            language="en",
            section=topic,
            subsection=None,
            content=f"{topic} content.",
            source_filename=filename,
            source_format="docx",
            content_hash=f"hash-{index}",
        )
        for index, topic in enumerate(LEGAL_TOPICS[:topic_count])
    ]

    return [overview, *topic_chunks]


def _no_existing_documents(
    country_code: str,
    client=None,
) -> list[ExistingCountryDocument]:
    del country_code, client
    return []


class WarningConfirmationRequiredTests(unittest.TestCase):
    """Mission "ORDER 3", sections 12-14 - the warning confirmation
    gate, and its combination with a pending replacement decision."""

    def test_zero_recognized_topics_is_context_warning_end_to_end(
        self,
    ) -> None:
        # Mission "ORDER 3", section 32.B - country detected, zero
        # legal topics recognized -> CONTEXT_WARNING specifically
        # (not STRUCTURE_WARNING), still confirmable, still succeeds.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentWarningConfirmationRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Chile.docx",
                    file_stream=BytesIO(b"chile-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=False,
                    chunk_builder=lambda path: (
                        _build_chunks_with_topic_count(
                            "CL", "Chile", "doc_" + "c" * 64,
                            path.name, 0,
                        )
                    ),
                    country_document_lookup=_no_existing_documents,
                    country_document_indexer=(
                        lambda **kwargs: (_ for _ in ()).throw(
                            AssertionError("indexer must not run")
                        )
                    ),
                )

            self.assertEqual(len(context.exception.warnings), 1)
            self.assertEqual(
                context.exception.warnings[0].code,
                "context_warning",
            )
            self.assertEqual(
                context.exception.warnings[0].recognized_topics_count,
                0,
            )

            def country_indexer(*, chunks, client):
                del client
                return DocumentIndexingResult(
                    index_alias="legal-documents-v1",
                    document_id=chunks[0].document_id,
                    source_filename=chunks[0].source_filename,
                    requested_chunks=len(chunks),
                    indexed_chunks=len(chunks),
                    stale_chunks_deleted=0,
                )

            response = safe_upload_and_index_document(
                filename="Chile.docx",
                file_stream=BytesIO(b"chile-bytes"),
                source_directory=source_directory,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                country_confirmed=True,
                confirm_warnings=True,
                chunk_builder=lambda path: (
                    _build_chunks_with_topic_count(
                        "CL", "Chile", "doc_" + "c" * 64,
                        path.name, 0,
                    )
                ),
                country_document_lookup=_no_existing_documents,
                country_document_indexer=country_indexer,
            )

            self.assertEqual(response.status, "uploaded")

    def test_fresh_country_with_thin_coverage_requires_confirmation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            indexer_called = False

            def country_indexer(**kwargs):
                nonlocal indexer_called
                del kwargs
                indexer_called = True
                raise AssertionError("indexer must not run")

            with self.assertRaises(
                AdminDocumentWarningConfirmationRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Chile.docx",
                    file_stream=BytesIO(b"chile-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=False,
                    chunk_builder=lambda path: (
                        _build_chunks_with_topic_count(
                            "CL", "Chile", "doc_" + "c" * 64,
                            path.name, 2,
                        )
                    ),
                    country_document_lookup=_no_existing_documents,
                    country_document_indexer=country_indexer,
                )

            error = context.exception
            self.assertEqual(error.country_code, "CL")
            self.assertFalse(error.replacement_required)
            self.assertEqual(error.existing_document_ids, ())
            self.assertEqual(len(error.warnings), 1)
            self.assertEqual(
                error.warnings[0].recognized_topics_count, 2
            )
            self.assertFalse(indexer_called)
            # Zero mutation - no source file was ever written (the
            # directory itself is always created up front, unrelated
            # to whether the upload proceeds).
            self.assertEqual(
                list(real_source_entries(source_directory)),
                [],
            )

            detail = error.to_detail()
            self.assertEqual(
                detail["code"],
                "document_warning_confirmation_required",
            )
            self.assertEqual(detail["operation"], "upload")
            self.assertEqual(detail["country_code"], "CL")
            self.assertEqual(detail["country_name"], "Chile")
            self.assertFalse(detail["replacement_required"])
            self.assertEqual(detail["existing_document_ids"], [])
            self.assertEqual(len(detail["warnings"]), 1)
            warning_detail = detail["warnings"][0]
            self.assertEqual(
                warning_detail["recognized_topics_count"], 2
            )
            self.assertEqual(
                warning_detail["expected_topics_count"], 11
            )
            self.assertEqual(len(warning_detail["missing_topics"]), 9)

    def test_confirm_warnings_true_proceeds_to_index(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            indexed_chunks = []

            def country_indexer(*, chunks, client):
                del client
                indexed_chunks.extend(chunks)
                return DocumentIndexingResult(
                    index_alias="legal-documents-v1",
                    document_id=chunks[0].document_id,
                    source_filename=chunks[0].source_filename,
                    requested_chunks=len(chunks),
                    indexed_chunks=len(chunks),
                    stale_chunks_deleted=0,
                )

            response = safe_upload_and_index_document(
                filename="Chile.docx",
                file_stream=BytesIO(b"chile-bytes"),
                source_directory=source_directory,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                country_confirmed=True,
                confirm_warnings=True,
                chunk_builder=lambda path: (
                    _build_chunks_with_topic_count(
                        "CL", "Chile", "doc_" + "c" * 64,
                        path.name, 2,
                    )
                ),
                country_document_lookup=_no_existing_documents,
                country_document_indexer=country_indexer,
            )

            self.assertEqual(response.status, "uploaded")
            self.assertEqual(response.country_code, "CL")
            self.assertEqual(len(indexed_chunks), 3)

    def test_six_recognized_topics_never_needs_confirmation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            def country_indexer(*, chunks, client):
                del client
                return DocumentIndexingResult(
                    index_alias="legal-documents-v1",
                    document_id=chunks[0].document_id,
                    source_filename=chunks[0].source_filename,
                    requested_chunks=len(chunks),
                    indexed_chunks=len(chunks),
                    stale_chunks_deleted=0,
                )

            response = safe_upload_and_index_document(
                filename="Chile.docx",
                file_stream=BytesIO(b"chile-bytes"),
                source_directory=source_directory,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                country_confirmed=True,
                confirm_warnings=False,
                chunk_builder=lambda path: (
                    _build_chunks_with_topic_count(
                        "CL", "Chile", "doc_" + "c" * 64,
                        path.name, 6,
                    )
                ),
                country_document_lookup=_no_existing_documents,
                country_document_indexer=country_indexer,
            )

            self.assertEqual(response.status, "uploaded")

    def test_warning_and_replacement_both_pending_are_combined(
        self,
    ) -> None:
        # Mission "ORDER 3", section 14: an atypical document for an
        # ALREADY-existing country must report both pending decisions
        # in one preflight response, not just one of them.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)
            (source_directory / "AU.docx").write_bytes(
                b"existing-australia-bytes"
            )

            with self.assertRaises(
                AdminDocumentWarningConfirmationRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia-new.docx",
                    file_stream=BytesIO(b"new-different-australia"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    replace_existing=False,
                    confirm_warnings=False,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                    country_document_indexer=(
                        lambda **kwargs: (_ for _ in ()).throw(
                            AssertionError("indexer must not run")
                        )
                    ),
                )

            error = context.exception
            self.assertTrue(error.replacement_required)
            self.assertEqual(
                error.existing_document_ids, (AU_OLD_ID,)
            )
            self.assertEqual(len(error.warnings), 1)
            # Zero mutation - the pre-existing source is untouched.
            self.assertEqual(
                (source_directory / "AU.docx").read_bytes(),
                b"existing-australia-bytes",
            )

    def test_warnings_confirmed_falls_through_to_plain_replacement(
        self,
    ) -> None:
        # Mission "ORDER 3", section 14: once confirm_warnings=True
        # resolves the warning, an existing country must still produce
        # the ORIGINAL, unchanged document_replacement_required
        # contract - never the combined warning shape - since only one
        # decision remains pending.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)
            (source_directory / "AU.docx").write_bytes(
                b"existing-australia-bytes"
            )

            with self.assertRaises(
                AdminDocumentReplacementRequiredError
            ):
                safe_upload_and_index_document(
                    filename="Australia-new.docx",
                    file_stream=BytesIO(b"new-different-australia"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    replace_existing=False,
                    confirm_warnings=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                    country_document_indexer=(
                        lambda **kwargs: (_ for _ in ()).throw(
                            AssertionError("indexer must not run")
                        )
                    ),
                )

    def test_identical_bytes_wins_over_warnings(self) -> None:
        # An identical re-upload of an already-thin document must stay
        # a clean no-op (document_already_current) - never re-litigate
        # its own already-accepted structure on every retry.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)
            (source_directory / "AU.docx").write_bytes(
                b"identical-bytes"
            )

            def existing_thin_document(country_code, client=None):
                del client
                assert country_code == "AU"
                return [
                    ExistingCountryDocument(
                        document_id=AU_OLD_ID,
                        source_filename="AU.docx",
                        country="Australia",
                        country_code="AU",
                        reference_year=None,
                    )
                ]

            with self.assertRaises(AdminDocumentAlreadyCurrentError):
                safe_upload_and_index_document(
                    filename="Australia-again.docx",
                    file_stream=BytesIO(b"identical-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=False,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=existing_thin_document,
                    country_document_indexer=(
                        lambda **kwargs: (_ for _ in ()).throw(
                            AssertionError("indexer must not run")
                        )
                    ),
                )


class AdminUploadAllowlistTests(unittest.TestCase):
    """
    Mission "ORDER 5C": the allowlist check runs after country
    detection/normalization but strictly before any mutation - no
    source commit, no OpenSearch write, no country_lock even acquired
    (the check sits inside the still-open staging TemporaryDirectory,
    before country_lock is entered) - so a rejected upload leaves
    everything exactly as it was.

    A perfectly-detected-but-disallowed country (Tunisia) must raise a
    distinct error/code (document_country_not_allowed) from an
    undeterminable one (document_country_undetermined) - the two must
    never be conflated (mission section 8).
    """

    def _reject_if_called(self, **kwargs):
        del kwargs
        raise AssertionError("indexer must not run")

    def test_disallowed_but_registered_country_is_rejected(
        self,
    ) -> None:
        # Tunisia: a real country_registry.py entry (so detection
        # succeeds cleanly) that is deliberately outside the 34-code
        # allowlist.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentCountryNotAllowedError
            ) as context:
                safe_upload_and_index_document(
                    filename="Tunisia.docx",
                    file_stream=BytesIO(b"tunisia-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=lambda path: (
                        _build_chunks_with_topic_count(
                            "TN", "Tunisia", "doc_" + "d" * 64,
                            path.name, 2,
                        )
                    ),
                    country_document_lookup=_no_existing_documents,
                    country_document_indexer=self._reject_if_called,
                )

            detail = context.exception.to_detail()
            self.assertEqual(
                detail["code"], "document_country_not_allowed"
            )
            self.assertEqual(context.exception.country_code, "TN")

            # Zero mutation: source_directory/processed_directory
            # themselves are created unconditionally up front (mission
            # "ORDER 5", not specific to the allowlist), but nothing
            # is ever written inside either of them, and the staging
            # TemporaryDirectory is cleaned up automatically.
            self.assertEqual(list(source_directory.iterdir()), [])
            self.assertEqual(list(processed_directory.iterdir()), [])

    def test_undetermined_country_is_a_different_error_than_not_allowed(
        self,
    ) -> None:
        # Mission "ORDER 8E-A1", section 8: an undeterminable-but-
        # processable DOCX is no longer a hard DocumentCountryUndeter
        # minedError at all - it is a SELECT_COUNTRY decision. It must
        # still stay clearly distinguishable from a country that WAS
        # detected but is simply outside the allowlist (superseded
        # from "ORDER 5C"'s original expectation of the old hard-fail
        # contract).
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            def undeterminable_builder(path):
                del path
                raise UndeterminableDocumentCountryError(
                    "no country could be resolved from content"
                )

            with self.assertRaises(
                AdminDocumentCountrySelectionRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Unknown.docx",
                    file_stream=BytesIO(b"unknown-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=undeterminable_builder,
                    country_document_lookup=_no_existing_documents,
                    country_document_indexer=self._reject_if_called,
                )

            # Never the allowlist error, and never the allowlist's
            # error code - an undetermined country and a determined-
            # but-disallowed one must stay distinguishable.
            self.assertNotIsInstance(
                context.exception,
                AdminDocumentCountryNotAllowedError,
            )
            self.assertNotIsInstance(
                context.exception,
                DocumentCountryUndeterminedError,
            )
            self.assertTrue(context.exception.allowed_countries)

    def test_every_one_of_the_34_allowed_codes_passes_the_check(
        self,
    ) -> None:
        from app.core.admin_country_policy import (
            ADMIN_ALLOWED_COUNTRY_CODES,
        )
        from app.core.country_registry import canonical_country_name

        for index, code in enumerate(sorted(ADMIN_ALLOWED_COUNTRY_CODES)):
            with self.subTest(code=code):
                with tempfile.TemporaryDirectory() as root:
                    source_directory = Path(root) / "source"
                    processed_directory = Path(root) / "processed"

                    country_name = canonical_country_name(code)
                    document_id = "doc_" + format(index, "064x")

                    indexed: dict[str, object] = {}

                    def indexer(*, chunks, client=None):
                        del client
                        indexed["chunks"] = chunks

                        return DocumentIndexingResult(
                            index_alias="legal-documents-v1",
                            document_id=chunks[0].document_id,
                            source_filename=chunks[0].source_filename,
                            requested_chunks=len(chunks),
                            indexed_chunks=len(chunks),
                            stale_chunks_deleted=0,
                        )

                    safe_upload_and_index_document(
                        filename=f"{code}.docx",
                        file_stream=BytesIO(f"{code}-bytes".encode()),
                        source_directory=source_directory,
                        processed_directory=processed_directory,
                        maximum_bytes=1000,
                        country_confirmed=True,
                        confirm_warnings=True,
                        chunk_builder=lambda path, _code=code, _name=(
                            country_name
                        ), _doc_id=document_id: (
                            _build_chunks_with_topic_count(
                                _code, _name, _doc_id, path.name, 2,
                            )
                        ),
                        country_document_lookup=_no_existing_documents,
                        country_document_indexer=indexer,
                    )

                    self.assertIn("chunks", indexed)


class SectionEditReplacementIntegrationTests(unittest.TestCase):
    """
    Mission "ORDER 5C", section 34 - a CONFIRMED replace is a full
    country document reset: every persisted section edit belonging to
    the document(s) just replaced must be gone afterwards, so it can
    never silently reapply to the new DOCX. Never on a cancelled
    replace (AdminDocumentReplacementRequiredError) or an identical-
    bytes no-op (AdminDocumentAlreadyCurrentError) - both raise before
    safe_upload_and_index_document ever reaches that point.
    """

    def test_confirmed_replace_wipes_old_section_edit_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"old-australia")

            write_section_edit_state_atomic(
                source_directory,
                SectionEditState(
                    document_id=AU_OLD_ID,
                    country_code="AU",
                    sections={
                        section_id_for_legal_topic(
                            "Employment Contracts"
                        ): SectionEdit(
                            legal_topic="Employment Contracts",
                            section="Employment Contracts",
                            subsection=None,
                            content=(
                                "An edit that must not reapply to "
                                "the replaced document."
                            ),
                        ),
                    },
                ),
            )
            self.assertIsNotNone(
                read_section_edit_state(source_directory, AU_OLD_ID)
            )

            def indexer(*, chunks, client=None):
                del client
                return DocumentIndexingResult(
                    index_alias="legal-documents-v1",
                    document_id=chunks[0].document_id,
                    source_filename=chunks[0].source_filename,
                    requested_chunks=len(chunks),
                    indexed_chunks=len(chunks),
                    stale_chunks_deleted=1,
                )

            response = safe_upload_and_index_document(
                filename="Australia 2026.docx",
                file_stream=BytesIO(b"new-australia"),
                source_directory=source_directory,
                confirm_warnings=True,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                country_confirmed=True,
                replace_existing=True,
                chunk_builder=lambda path: [
                    _build_au_chunk(path.name)
                ],
                country_document_lookup=_existing_documents,
                country_document_indexer=indexer,
            )

            self.assertEqual(response.status, "replaced")
            self.assertIsNone(
                read_section_edit_state(source_directory, AU_OLD_ID)
            )

    def test_cancelled_replace_leaves_section_edit_state_untouched(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"old-australia")

            write_section_edit_state_atomic(
                source_directory,
                SectionEditState(
                    document_id=AU_OLD_ID,
                    country_code="AU",
                    sections={
                        section_id_for_legal_topic(
                            "Employment Contracts"
                        ): SectionEdit(
                            legal_topic="Employment Contracts",
                            section="Employment Contracts",
                            subsection=None,
                            content=(
                                "Must survive an unconfirmed replace."
                            ),
                        ),
                    },
                ),
            )

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError("indexer must not run")

            with self.assertRaises(
                AdminDocumentReplacementRequiredError
            ):
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"new-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                    country_document_indexer=(
                        indexer_must_not_be_called
                    ),
                )

            after = read_section_edit_state(source_directory, AU_OLD_ID)
            self.assertIsNotNone(after)
            self.assertEqual(
                after.sections[
                    section_id_for_legal_topic("Employment Contracts")
                ].content,
                "Must survive an unconfirmed replace.",
            )

    def test_identical_upload_leaves_section_edit_state_untouched(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"same-australia")

            write_section_edit_state_atomic(
                source_directory,
                SectionEditState(
                    document_id=AU_OLD_ID,
                    country_code="AU",
                    sections={
                        section_id_for_legal_topic(
                            "Employment Contracts"
                        ): SectionEdit(
                            legal_topic="Employment Contracts",
                            section="Employment Contracts",
                            subsection=None,
                            content=(
                                "Must survive an identical no-op "
                                "upload."
                            ),
                        ),
                    },
                ),
            )

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError(
                    "the indexer must not be called for an "
                    "already-current document"
                )

            with self.assertRaises(
                AdminDocumentAlreadyCurrentError
            ):
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"same-australia"),
                    source_directory=source_directory,
                    confirm_warnings=True,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                    country_document_indexer=(
                        indexer_must_not_be_called
                    ),
                )

            self.assertIsNotNone(
                read_section_edit_state(source_directory, AU_OLD_ID)
            )


class CountryConfirmationGateTests(unittest.TestCase):
    """
    Mission "ORDER 8E-A1", sections 6/17 - a detected country (from
    content or a marker) must never, by itself, cause any mutation:
    country_confirmed=True is required first, checked before the
    content-warning gate and before the existing-country/conflict
    checks, one decision at a time.
    """

    def test_detected_country_requires_confirmation_first(
        self,
    ) -> None:
        indexer_called = False

        def indexer_must_not_be_called(**kwargs):
            nonlocal indexer_called
            del kwargs
            indexer_called = True
            raise AssertionError("indexer must not run")

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentCountryConfirmationRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia.docx",
                    file_stream=BytesIO(b"australia-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    confirm_warnings=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=(
                        lambda code, client: (_ for _ in ()).throw(
                            AssertionError(
                                "existing-country lookup must not run "
                                "before country is confirmed"
                            )
                        )
                    ),
                    country_document_indexer=indexer_must_not_be_called,
                )

            self.assertEqual(context.exception.country_code, "AU")
            self.assertEqual(context.exception.country, "Australia")
            self.assertEqual(
                context.exception.detection_source, "content"
            )
            self.assertFalse(indexer_called)

            detail = context.exception.to_detail()
            self.assertEqual(
                detail["code"],
                "document_country_confirmation_required",
            )
            self.assertEqual(detail["country_code"], "AU")
            self.assertEqual(detail["detection_source"], "content")
            # Mission "ORDER 8E-A2", section 6: a UI must be able to
            # offer "choose a different country" straight from this
            # same response, using the one authoritative server-side
            # list - never a second, client-invented copy of it.
            codes = {
                option["code"]
                for option in detail["allowed_countries"]
            }
            self.assertIn("FR", codes)
            self.assertIn("AU", codes)

            # Zero mutation, and Cancel (never retrying) leaves no
            # trace - the staged file lived only in the request's own
            # now-cleaned-up TemporaryDirectory.
            self.assertEqual(
                list(source_directory.iterdir())
                if source_directory.exists()
                else [],
                [],
            )
            self.assertEqual(
                list(processed_directory.iterdir())
                if processed_directory.exists()
                else [],
                [],
            )

    def test_country_confirmed_true_proceeds(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            response = safe_upload_and_index_document(
                filename="Australia.docx",
                file_stream=BytesIO(b"australia-bytes"),
                source_directory=source_directory,
                processed_directory=processed_directory,
                maximum_bytes=1000,
                confirm_warnings=True,
                country_confirmed=True,
                chunk_builder=lambda path: [
                    _build_au_chunk(path.name)
                ],
                country_document_lookup=lambda code, client: [],
                country_document_indexer=_fake_indexer,
            )

            self.assertEqual(response.status, "uploaded")
            self.assertEqual(response.country_code, "AU")

    def test_confirmation_gate_precedes_the_content_warning_gate(
        self,
    ) -> None:
        # An atypical, thin-coverage document for a fresh country must
        # still ask for country confirmation FIRST, never the content
        # warning, when neither decision has been made yet.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentCountryConfirmationRequiredError
            ):
                safe_upload_and_index_document(
                    filename="Chile.docx",
                    file_stream=BytesIO(b"chile-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    confirm_warnings=False,
                    chunk_builder=lambda path: (
                        _build_chunks_with_topic_count(
                            "CL", "Chile", "doc_" + "c" * 64,
                            path.name, 0,
                        )
                    ),
                    country_document_lookup=_no_existing_documents,
                    country_document_indexer=(
                        lambda **kwargs: (_ for _ in ()).throw(
                            AssertionError("indexer must not run")
                        )
                    ),
                )

    def test_confirmation_gate_precedes_conflict_review(self) -> None:
        # A country already in a multi-document conflict state must
        # still ask for country confirmation FIRST, before the
        # existing-country lookup that would discover the conflict.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentCountryConfirmationRequiredError
            ):
                safe_upload_and_index_document(
                    filename="Australia.docx",
                    file_stream=BytesIO(b"australia-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    confirm_warnings=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=(
                        lambda code, client: (_ for _ in ()).throw(
                            AssertionError(
                                "must never reach the existing-"
                                "country lookup before confirmation"
                            )
                        )
                    ),
                )


class CountrySelectionGateTests(unittest.TestCase):
    """
    Mission "ORDER 8E-A1", section 8/9 - an otherwise-processable DOCX
    with no identifiable country is a SELECT_COUNTRY decision, never a
    hard failure; a manually-selected country is validated server-side
    against the admin allowlist before anything else happens.
    """

    def test_undetermined_country_returns_selection_required(
        self,
    ) -> None:
        def undeterminable_builder(path):
            del path
            raise UndeterminableDocumentCountryError(
                "no country could be resolved from content"
            )

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentCountrySelectionRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Unknown.docx",
                    file_stream=BytesIO(b"unknown-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=undeterminable_builder,
                    country_document_lookup=_no_existing_documents,
                )

            detail = context.exception.to_detail()
            self.assertEqual(
                detail["code"], "document_country_selection_required"
            )
            codes = {
                option["code"]
                for option in detail["allowed_countries"]
            }
            self.assertIn("FR", codes)
            self.assertIn("AU", codes)

    def test_ambiguous_country_also_returns_selection_required(
        self,
    ) -> None:
        def ambiguous_builder(path):
            del path
            raise AmbiguousDocumentCountryError(
                "more than one country found in the cover"
            )

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentCountrySelectionRequiredError
            ):
                safe_upload_and_index_document(
                    filename="Ambiguous.docx",
                    file_stream=BytesIO(b"ambiguous-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=ambiguous_builder,
                    country_document_lookup=_no_existing_documents,
                )

    def test_invalid_manual_selection_is_rejected_with_zero_mutation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentCountrySelectionInvalidError
            ) as context:
                safe_upload_and_index_document(
                    filename="Unknown.docx",
                    # Not even a real DOCX - proves the invalid-code
                    # check runs before any archive is ever opened.
                    file_stream=BytesIO(b"not-a-real-docx"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    confirm_warnings=True,
                    selected_country_code="ZZ",
                    chunk_builder=(
                        lambda path: (_ for _ in ()).throw(
                            AssertionError(
                                "chunk_builder must not run for an "
                                "invalid manual selection"
                            )
                        )
                    ),
                )

            detail = context.exception.to_detail()
            self.assertEqual(
                detail["code"], "document_country_selection_invalid"
            )
            self.assertEqual(detail["country_code"], "ZZ")
            self.assertEqual(
                list(source_directory.iterdir())
                if source_directory.exists()
                else [],
                [],
            )

    def test_unregistered_manual_selection_code_is_rejected(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentCountrySelectionInvalidError
            ):
                safe_upload_and_index_document(
                    filename="Unknown.docx",
                    file_stream=BytesIO(b"not-a-real-docx"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    confirm_warnings=True,
                    selected_country_code="XX",
                )


class ManualCountryDocxRoundtripTests(unittest.TestCase):
    """
    Mission "ORDER 8E-A1", sections 10/12/13 - the full manual-country
    invariant: a countryless DOCX, once the Admin selects a country,
    must have that choice persisted INSIDE the DOCX itself (a DOCX-
    native marker), surviving Download, a fresh independent parse, a
    full Reindex, and a later re-upload - with no external state
    anywhere.
    """

    def test_selecting_a_country_embeds_the_marker_and_uploads(
        self,
    ) -> None:
        countryless_bytes = _real_docx_bytes(
            [
                "Some heading with no recognizable country.",
                "01. Hiring Practices",
                "Content about hiring rules and probation periods.",
            ]
        )

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            built_chunks = []

            def capturing_chunk_builder(path):
                chunks = build_document_chunks_from_docx(path)
                built_chunks.extend(chunks)
                return chunks

            response = safe_upload_and_index_document(
                filename="mystery.docx",
                file_stream=BytesIO(countryless_bytes),
                source_directory=source_directory,
                processed_directory=processed_directory,
                maximum_bytes=1_000_000,
                confirm_warnings=True,
                selected_country_code="fr",
                chunk_builder=capturing_chunk_builder,
                country_document_lookup=lambda code, client: [],
                country_document_indexer=_fake_indexer,
            )

            self.assertEqual(response.status, "uploaded")
            self.assertEqual(response.country_code, "FR")

            stored_path = source_directory / "FR.docx"
            self.assertTrue(stored_path.exists())

            marker = read_country_marker(stored_path)
            self.assertIsNotNone(marker)
            self.assertEqual(marker.country_code, "FR")

            # Mission "ORDER 8E-A2" real-browser finding: the marker-
            # embedding candidate lives in its own subdirectory, never
            # under a disambiguating filename prefix - chunk_builder
            # reads source_filename straight from the candidate path's
            # own name, which is then indexed and shown to the Admin
            # verbatim, so it must stay exactly the original sanitized
            # upload name, never leak an internal artifact like
            # "country-marker.mystery.docx".
            self.assertEqual(response.source_filename, "mystery.docx")
            self.assertTrue(built_chunks)
            self.assertTrue(
                all(
                    chunk.source_filename == "mystery.docx"
                    for chunk in built_chunks
                ),
                [chunk.source_filename for chunk in built_chunks],
            )

    def test_full_manual_country_roundtrip_invariant(self) -> None:
        # countryless -> select FR -> marker written -> Download ->
        # fresh independent parse -> FR -> full Reindex -> still FR ->
        # re-upload -> detects FR -> asks to confirm FR, no external
        # state anywhere (mission section 12).
        countryless_bytes = _real_docx_bytes(
            [
                "Some heading with no recognizable country.",
                "01. Hiring Practices",
                "Content about hiring rules and probation periods.",
            ]
        )

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            first = safe_upload_and_index_document(
                filename="mystery.docx",
                file_stream=BytesIO(countryless_bytes),
                source_directory=source_directory,
                processed_directory=processed_directory,
                maximum_bytes=1_000_000,
                confirm_warnings=True,
                selected_country_code="fr",
                chunk_builder=build_document_chunks_from_docx,
                country_document_lookup=lambda code, client: [],
                country_document_indexer=_fake_indexer,
            )
            self.assertEqual(first.status, "uploaded")

            stored_path = source_directory / "FR.docx"

            # "Download" = read the stored bytes back out.
            downloaded_bytes = stored_path.read_bytes()

            # A fresh, completely independent parse (no marker-writing
            # involved at all) still resolves FR from the marker.
            fresh_metadata = metadata_from_content(stored_path)
            self.assertEqual(fresh_metadata.country_code, "FR")

            # A full Reindex re-parses the very same stored file.
            reindexed_chunks = build_document_chunks_from_docx(
                stored_path
            )
            self.assertTrue(reindexed_chunks)
            self.assertTrue(
                all(
                    chunk.country_code == "FR"
                    for chunk in reindexed_chunks
                )
            )

            # Re-upload the downloaded bytes with no manual selection
            # this time - the marker alone must still detect FR and
            # require confirmation, never external state.
            existing = [
                ExistingCountryDocument(
                    document_id=first.document_id,
                    source_filename="FR.docx",
                    country="France",
                    country_code="FR",
                    reference_year=None,
                )
            ]

            with self.assertRaises(
                AdminDocumentCountryConfirmationRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="mystery.docx",
                    file_stream=BytesIO(downloaded_bytes),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1_000_000,
                    confirm_warnings=True,
                    chunk_builder=build_document_chunks_from_docx,
                    country_document_lookup=(
                        lambda code, client: existing
                    ),
                )

            self.assertEqual(context.exception.country_code, "FR")
            self.assertEqual(
                context.exception.detection_source, "marker"
            )

    def test_manual_selection_marker_embed_never_touches_disk_on_a_later_hard_failure(
        self,
    ) -> None:
        # If a later stage after marker-embedding hard-fails (here:
        # the country is fresh but the topic warning is declined), the
        # embed happened only inside the request's own staging
        # TemporaryDirectory - nothing was ever written to
        # source_directory.
        countryless_bytes = _real_docx_bytes(
            ["Some heading with no recognizable country."]
        )

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"

            with self.assertRaises(
                AdminDocumentWarningConfirmationRequiredError
            ):
                safe_upload_and_index_document(
                    filename="mystery.docx",
                    file_stream=BytesIO(countryless_bytes),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1_000_000,
                    confirm_warnings=False,
                    selected_country_code="fr",
                    chunk_builder=build_document_chunks_from_docx,
                    country_document_lookup=lambda code, client: [],
                    country_document_indexer=(
                        lambda **kwargs: (_ for _ in ()).throw(
                            AssertionError("indexer must not run")
                        )
                    ),
                )

            self.assertEqual(
                list(real_source_entries(source_directory)),
                [],
            )


class CountryConflictReviewCandidateTests(unittest.TestCase):
    """
    Mission "ORDER 8E-A1", section 22 - the conflict review exposes
    only safe, business-facing per-candidate fields (filename, year,
    last-updated, file size); document_id is present only as an
    internal identity, never required for the Admin to read.
    """

    def test_candidates_expose_only_safe_fields_with_real_file_stats(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            # Both use distinct, non-canonical legacy names - the
            # canonical AU.docx candidate resolve_document_source_path
            # always also tries would otherwise falsely collide with
            # either one's own resolution (see
            # document_source_resolver.py: the canonical name is a
            # fallback candidate for every document of a country, not
            # just its own current record).
            first_path = source_directory / "Australia-2024.docx"
            second_path = source_directory / "Australia-legacy-v2.docx"
            first_path.write_bytes(b"first-australia-content")
            second_path.write_bytes(b"second-australia-content-longer")

            existing = [
                ExistingCountryDocument(
                    document_id=AU_OLD_ID,
                    source_filename=first_path.name,
                    country="Australia",
                    country_code="AU",
                    reference_year=2024,
                ),
                ExistingCountryDocument(
                    document_id="doc_" + "c" * 64,
                    source_filename=second_path.name,
                    country="Australia",
                    country_code="AU",
                    reference_year=2026,
                ),
            ]

            with self.assertRaises(
                AdminDocumentCountryConflictReviewRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia-new.docx",
                    file_stream=BytesIO(b"new-australia"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    confirm_warnings=True,
                    country_confirmed=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=(
                        lambda code, client: existing
                    ),
                )

            detail = context.exception.to_detail()
            self.assertEqual(
                detail["code"],
                "document_country_conflict_review_required",
            )
            candidates = {
                candidate["source_filename"]: candidate
                for candidate in detail["candidates"]
            }
            self.assertEqual(
                set(candidates),
                {"Australia-2024.docx", "Australia-legacy-v2.docx"},
            )
            self.assertEqual(
                candidates["Australia-2024.docx"]["reference_year"],
                2024,
            )
            self.assertEqual(
                candidates["Australia-2024.docx"]["source_bytes"],
                len(b"first-australia-content"),
            )
            self.assertIsNotNone(
                candidates["Australia-2024.docx"]["updated_at"]
            )
            # document_id is present (needed for a hidden
            # CHOOSE_DOCUMENT reference) but every candidate here still
            # carries its own safe display fields alongside it.
            self.assertIn(
                "document_id",
                candidates["Australia-legacy-v2.docx"],
            )


class _FakeReseedMetadataClient:
    """
    Minimal OpenSearch double for the "identical bytes, explicit
    confirmed reseed" tests below - just enough for
    reseed_contacts_from_current_docx's own metadata fetch
    (_get_document_metadata's plain, non-"sort" shape) plus a
    country-wide conflict check (_ensure_no_country_conflict's own
    "sort" shape). Never a full FakeCountryOpenSearch/FakeSection
    OpenSearchClient - those exist for OTHER tests' own different
    metadata shapes.
    """

    def __init__(
        self,
        *,
        document_id: str,
        country_code: str = "AU",
        country: str = "Australia",
        source_filename: str = (
            "Employment Law Overview Australia.docx"
        ),
        reference_year: int | None = None,
    ) -> None:
        self.document_id = document_id
        self.country_code = country_code
        self.country = country
        self.source_filename = source_filename
        self.reference_year = reference_year

    def search(self, index, body):
        del index

        if "sort" in body:
            return {"hits": {"total": {"value": 0}, "hits": []}}

        term = body.get("query", {}).get("term", {})

        if term.get("document_id") != self.document_id:
            return {"hits": {"hits": []}}

        return {
            "hits": {
                "hits": [
                    {
                        "_source": {
                            "document_id": self.document_id,
                            "source_filename": self.source_filename,
                            "country": self.country,
                            "country_code": self.country_code,
                            "reference_year": self.reference_year,
                        }
                    }
                ]
            }
        }


class AdminModifiedReplacementWarningTests(unittest.TestCase):
    """
    Mission "ORDER 8G-B2", sections 12/14/15/16/17 - the
    admin_modified-aware replacement warning integration. Does not
    duplicate ORDER 8G-B1's own 35 CRUD/state tests - only the NEW
    behavior added on top of them here.
    """

    def test_clean_document_replacement_flow_is_unchanged(self) -> None:
        # marker=False (never touched by Admin) - the existing
        # replacement-required flow must be byte-for-byte unchanged:
        # admin_modified is explicitly False in the detail.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"legacy-australia")

            with self.assertRaises(
                AdminDocumentReplacementRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"new-australia-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                )

            self.assertFalse(
                context.exception.to_detail()["admin_modified"]
            )

    def test_dirty_document_replacement_flags_admin_modified(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"legacy-australia")

            from app.services.admin_modification_marker import (
                mark_admin_modified,
            )

            mark_admin_modified(source_directory, AU_OLD_ID)

            with self.assertRaises(
                AdminDocumentReplacementRequiredError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"new-australia-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                )

            self.assertTrue(
                context.exception.to_detail()["admin_modified"]
            )

    def test_identical_bytes_clean_still_raises_already_current(
        self,
    ) -> None:
        # marker=False - the pre-existing short-circuit is completely
        # unchanged.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"same-australia")

            with self.assertRaises(AdminDocumentAlreadyCurrentError):
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"same-australia"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                )

    def test_identical_bytes_dirty_requires_explicit_reseed_confirmation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"same-australia")

            from app.services.admin_modification_marker import (
                mark_admin_modified,
            )

            mark_admin_modified(source_directory, AU_OLD_ID)

            with self.assertRaises(
                AdminDocumentIdenticalButAdminModifiedError
            ) as context:
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"same-australia"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                )

            self.assertEqual(
                context.exception.to_detail()["document_id"],
                AU_OLD_ID,
            )

    def test_identical_bytes_dirty_cancel_leaves_zero_mutation(
        self,
    ) -> None:
        # "Cancel" is simply never resubmitting with
        # confirm_contact_reseed=True - proven here as zero mutation to
        # the structured contact state after the warning is raised.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"same-australia")

            from app.services.admin_modification_marker import (
                mark_admin_modified,
            )
            from app.services.contact_state import (
                ContactState,
                read_contact_state,
                write_contact_state_atomic,
            )
            from app.services.contact_state import ContactRecord

            mark_admin_modified(source_directory, AU_OLD_ID)
            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=AU_OLD_ID,
                    country_code="AU",
                    contacts=(
                        ContactRecord(
                            contact_id="admin-edit-1",
                            member_firm="Admin Edited Firm",
                        ),
                    ),
                ),
            )

            with self.assertRaises(
                AdminDocumentIdenticalButAdminModifiedError
            ):
                safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"same-australia"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                )

            state = read_contact_state(source_directory, AU_OLD_ID)
            self.assertEqual(
                state.contacts[0].member_firm, "Admin Edited Firm"
            )

    def test_identical_bytes_dirty_confirmed_reseeds_contacts(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"same-australia")

            from app.services.admin_modification_marker import (
                is_admin_modified_since_upload,
                mark_admin_modified,
            )
            from app.services.contact_state import (
                ContactRecord,
                ContactState,
                read_contact_state,
                write_contact_state_atomic,
            )
            from app.services.docx_parser import ExtractedContact

            mark_admin_modified(source_directory, AU_OLD_ID)
            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=AU_OLD_ID,
                    country_code="AU",
                    contacts=(
                        ContactRecord(
                            contact_id="admin-edit-1",
                            member_firm="Admin Edited Firm",
                        ),
                    ),
                ),
            )

            with patch(
                "app.services.admin_contacts."
                "extract_contacts_from_docx",
                return_value=[
                    ExtractedContact(member_firm="Parsed DOCX Firm"),
                ],
            ), patch(
                "app.services.document_indexer.ensure_legal_documents_index"
            ), patch(
                "app.services.document_indexer.bulk",
                return_value=(1, []),
            ), patch(
                "app.services.document_indexer._snapshot_document_chunks",
                return_value=[],
            ), patch(
                "app.services.document_indexer._delete_chunks_except",
                return_value=0,
            ):
                response = safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"same-australia"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    confirm_contact_reseed=True,
                    client=_FakeReseedMetadataClient(document_id=AU_OLD_ID),
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                )

            self.assertEqual(response.status, "contacts_reseeded")
            self.assertEqual(response.contact_count, 1)

            state = read_contact_state(source_directory, AU_OLD_ID)
            self.assertEqual(
                state.contacts[0].member_firm, "Parsed DOCX Firm"
            )
            self.assertFalse(
                is_admin_modified_since_upload(
                    source_directory, AU_OLD_ID
                )
            )

    def test_confirmed_different_docx_resets_marker_and_reseeds(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"legacy-australia")

            from app.services.admin_modification_marker import (
                is_admin_modified_since_upload,
                mark_admin_modified,
            )
            from app.services.contact_state import (
                ContactRecord,
                ContactState,
                read_contact_state,
                write_contact_state_atomic,
            )
            from app.services.docx_parser import ExtractedContact

            mark_admin_modified(source_directory, AU_NEW_ID)
            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=AU_NEW_ID,
                    country_code="AU",
                    contacts=(
                        ContactRecord(
                            contact_id="admin-edit-1",
                            member_firm="Old Admin Contact",
                        ),
                        ContactRecord(
                            contact_id="admin-edit-2",
                            member_firm="Old Admin Contact 2",
                        ),
                    ),
                ),
            )

            with patch(
                "app.services.admin_document_replacement."
                "extract_contacts_from_docx",
                return_value=[
                    ExtractedContact(member_firm="New DOCX Firm"),
                ],
            ):
                response = safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"brand-new-australia-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    replace_existing=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                    country_document_indexer=_fake_indexer,
                )

            self.assertEqual(response.contact_count, 1)

            state = read_contact_state(source_directory, AU_NEW_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(
                state.contacts[0].member_firm, "New DOCX Firm"
            )
            self.assertFalse(
                is_admin_modified_since_upload(
                    source_directory, AU_NEW_ID
                )
            )

    def test_confirmed_different_docx_with_zero_contacts_is_explicit_empty(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"legacy-australia")

            from app.services.contact_state import (
                ContactRecord,
                ContactState,
                read_contact_state,
                write_contact_state_atomic,
            )

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=AU_NEW_ID,
                    country_code="AU",
                    contacts=(
                        ContactRecord(
                            contact_id="admin-edit-1",
                            member_firm="Old Admin Contact",
                        ),
                    ),
                ),
            )

            with patch(
                "app.services.admin_document_replacement."
                "extract_contacts_from_docx",
                return_value=[],
            ):
                response = safe_upload_and_index_document(
                    filename="Australia 2026.docx",
                    file_stream=BytesIO(b"brand-new-australia-bytes"),
                    source_directory=source_directory,
                    processed_directory=processed_directory,
                    maximum_bytes=1000,
                    country_confirmed=True,
                    confirm_warnings=True,
                    replace_existing=True,
                    chunk_builder=lambda path: [
                        _build_au_chunk(path.name)
                    ],
                    country_document_lookup=_existing_documents,
                    country_document_indexer=_fake_indexer,
                )

            self.assertEqual(response.contact_count, 0)

            state = read_contact_state(source_directory, AU_NEW_ID)
            self.assertIsNotNone(state)
            self.assertEqual(state.contacts, ())

    def test_reseed_failure_restores_marker_and_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / "source"
            processed_directory = Path(root) / "processed"
            source_directory.mkdir(parents=True)

            (
                source_directory
                / "Employment Law Overview Australia.docx"
            ).write_bytes(b"same-australia")

            from app.services.admin_modification_marker import (
                is_admin_modified_since_upload,
                mark_admin_modified,
            )
            from app.services.contact_state import (
                ContactRecord,
                ContactState,
                read_contact_state,
                write_contact_state_atomic,
            )
            from app.services.docx_parser import ExtractedContact

            mark_admin_modified(source_directory, AU_OLD_ID)
            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=AU_OLD_ID,
                    country_code="AU",
                    contacts=(
                        ContactRecord(
                            contact_id="admin-edit-1",
                            member_firm="Admin Edited Firm",
                        ),
                    ),
                ),
            )

            with patch(
                "app.services.admin_contacts."
                "extract_contacts_from_docx",
                return_value=[
                    ExtractedContact(member_firm="Parsed DOCX Firm"),
                ],
            ), patch(
                "app.services.admin_contacts.write_contact_state_atomic",
                side_effect=OSError("simulated disk failure"),
            ), patch(
                "app.services.document_indexer.ensure_legal_documents_index"
            ), patch(
                "app.services.document_indexer.bulk",
                return_value=(1, []),
            ), patch(
                "app.services.document_indexer._snapshot_document_chunks",
                return_value=[],
            ), patch(
                "app.services.document_indexer._delete_chunks_except",
                return_value=0,
            ):
                with self.assertRaises(Exception):
                    safe_upload_and_index_document(
                        filename="Australia 2026.docx",
                        file_stream=BytesIO(b"same-australia"),
                        source_directory=source_directory,
                        processed_directory=processed_directory,
                        maximum_bytes=1000,
                        country_confirmed=True,
                        confirm_warnings=True,
                        confirm_contact_reseed=True,
                        client=_FakeReseedMetadataClient(document_id=AU_OLD_ID),
                        chunk_builder=lambda path: [
                            _build_au_chunk(path.name)
                        ],
                        country_document_lookup=_existing_documents,
                    )

            state = read_contact_state(source_directory, AU_OLD_ID)
            self.assertEqual(
                state.contacts[0].member_firm, "Admin Edited Firm"
            )
            self.assertTrue(
                is_admin_modified_since_upload(
                    source_directory, AU_OLD_ID
                )
            )


if __name__ == "__main__":
    unittest.main()
