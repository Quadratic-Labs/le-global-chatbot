import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from app.core.country_registry import (
    CountryMetadataMismatchError,
)
from app.core.legal_taxonomy import (
    get_canonical_legal_topic,
    normalize_topic,
)
from app.services.document_chunk_builder import (
    AmbiguousDocumentCountryError,
    DocumentMetadata,
    InvalidDocxFormatError,
    UndeterminableDocumentCountryError,
    UnknownLegalTopicError,
    build_document_chunks,
    build_document_chunks_from_docx,
    metadata_from_content,
    validate_docx_format,
)
from app.services.docx_parser import ParsedSection


def _build_docx(
    directory: Path,
    title_lines: list[str],
    body_paragraphs: list[str] | None = None,
    filename: str = "document.docx",
) -> Path:
    """
    A minimal real DOCX whose leading paragraphs are exactly
    `title_lines` - the title/cover area metadata_from_content scans
    - followed by any extra body content, saved under an arbitrary
    filename (never itself a source of metadata).
    """

    document = Document()

    for line in title_lines:
        document.add_paragraph(line)

    for paragraph in body_paragraphs or []:
        document.add_paragraph(paragraph)

    file_path = directory / filename
    document.save(file_path)

    return file_path


_DOCUMENT_XML_PATH = "word/document.xml"


def _inject_body_xml_fragment(
    file_path: Path,
    xml_fragment: str,
) -> None:
    """
    Insert a raw XML fragment as the first child of <w:body> in an
    existing, already-saved DOCX - the only way to build a fixture
    containing structures python-docx has no high-level API for
    (a real Word text box's w:txbxContent, DrawingML a:t runs), since
    both are anchored drawings rather than ordinary paragraph content.
    """

    with zipfile.ZipFile(file_path) as archive:
        entries = {
            name: archive.read(name)
            for name in archive.namelist()
        }

    document_xml = entries[_DOCUMENT_XML_PATH].decode("utf-8")

    marker = "<w:body>"
    insertion_point = document_xml.index(marker) + len(marker)

    entries[_DOCUMENT_XML_PATH] = (
        document_xml[:insertion_point]
        + xml_fragment
        + document_xml[insertion_point:]
    ).encode("utf-8")

    with zipfile.ZipFile(file_path, "w") as archive:
        for name, data in entries.items():
            archive.writestr(name, data)


def _build_docx_with_text_box_title(
    directory: Path,
    text_box_text: str,
    body_paragraphs: list[str] | None = None,
    filename: str = "document.docx",
) -> Path:
    """
    A minimal real DOCX whose only title-shaped text lives inside a
    Word text box (w:txbxContent) - invisible to python-docx's own
    paragraph iteration, exactly like 8 of the 10 real production
    documents this mission adds support for.
    """

    document = Document()

    for paragraph in body_paragraphs or ["Introduction"]:
        document.add_paragraph(paragraph)

    file_path = directory / filename
    document.save(file_path)

    _inject_body_xml_fragment(
        file_path,
        (
            "<w:p><w:txbxContent><w:p><w:r><w:t>"
            f"{text_box_text}"
            "</w:t></w:r></w:p></w:txbxContent></w:p>"
        ),
    )

    return file_path


def _build_docx_with_drawingml_title(
    directory: Path,
    drawingml_text: str,
    body_paragraphs: list[str] | None = None,
    filename: str = "document.docx",
) -> Path:
    """
    A minimal real DOCX whose only title-shaped text is a DrawingML
    (SmartArt/WordArt) a:t run - a third XML namespace, distinct from
    both ordinary paragraphs and text boxes.
    """

    document = Document()

    for paragraph in body_paragraphs or ["Introduction"]:
        document.add_paragraph(paragraph)

    file_path = directory / filename
    document.save(file_path)

    _inject_body_xml_fragment(
        file_path,
        (
            '<w:p><w:r><a:t xmlns:a="http://schemas.openxmlformats.org'
            '/drawingml/2006/main">'
            f"{drawingml_text}"
            "</a:t></w:r></w:p>"
        ),
    )

    return file_path


def _build_docx_with_header_title(
    directory: Path,
    header_text: str,
    body_paragraphs: list[str] | None = None,
    filename: str = "document.docx",
) -> Path:
    """A minimal real DOCX whose only title-shaped text is in a header."""

    document = Document()
    document.sections[0].header.paragraphs[0].text = header_text

    for paragraph in body_paragraphs or ["Introduction"]:
        document.add_paragraph(paragraph)

    file_path = directory / filename
    document.save(file_path)

    return file_path


class LegalTaxonomyTests(unittest.TestCase):
    def test_normalizes_numbered_topic(self) -> None:
        self.assertEqual(
            normalize_topic(
                section="01. Hiring Practices",
                country="Spain",
            ),
            "Hiring Practices",
        )

    def test_removes_country_suffix(self) -> None:
        self.assertEqual(
            normalize_topic(
                section=(
                    "06. Social Media "
                    "and Data Privacy in Spain"
                ),
                country="Spain",
            ),
            "Social Media and Data Privacy",
        )

    def test_returns_canonical_topic(self) -> None:
        self.assertEqual(
            get_canonical_legal_topic(
                section=(
                    "07. Termination "
                    "of Employment Contracts"
                ),
                country="Spain",
            ),
            "Termination of Employment Contracts",
        )

    def test_roman_numeral_prefix_is_stripped(self) -> None:
        self.assertEqual(
            get_canonical_legal_topic(
                section="II.  Hiring Practices",
                country="Taiwan",
            ),
            "Hiring Practices",
        )

    def test_wages_and_work_hours_is_a_working_conditions_alias(
        self,
    ) -> None:
        # The real USA document's own section 3 heading - a genuine
        # phrasing variant found in production content.
        self.assertEqual(
            get_canonical_legal_topic(
                section="03. Wages and Work Hours",
                country="United States",
            ),
            "Working Conditions",
        )

    def test_termination_of_employment_without_contracts_is_an_alias(
        self,
    ) -> None:
        # The real Philippines document's own bold heading - "Contracts"
        # is dropped.
        self.assertEqual(
            get_canonical_legal_topic(
                section="Termination of Employment",
                country="Philippines",
            ),
            "Termination of Employment Contracts",
        )


class DocumentChunkBuilderTests(unittest.TestCase):
    def setUp(self) -> None:
        self.metadata = DocumentMetadata(
            country="Spain",
            country_code="ES",
            reference_year=2026,
            language="en",
            source_filename=(
                "Labour and Employment Law "
                "in Spain 2026.docx"
            ),
        )

    def test_builds_overview_and_comparator_chunks(
        self,
    ) -> None:
        parsed_sections = [
            ParsedSection(
                section="Employment Law Overview Spain",
                subsection="Introduction",
                content="Overview content.",
            ),
            ParsedSection(
                section="01. Hiring Practices",
                subsection=(
                    "Requirement for Foreign "
                    "Employees to Work"
                ),
                content="Hiring content.",
            ),
        ]

        chunks = build_document_chunks(
            parsed_sections=parsed_sections,
            metadata=self.metadata,
        )

        self.assertEqual(
            len(chunks),
            2,
        )

        overview_chunk = chunks[0]

        self.assertEqual(
            overview_chunk.document_type,
            "overview",
        )

        self.assertIsNone(
            overview_chunk.legal_topic
        )

        comparator_chunk = chunks[1]

        self.assertEqual(
            comparator_chunk.document_type,
            "comparator",
        )

        self.assertEqual(
            comparator_chunk.legal_topic,
            "Hiring Practices",
        )

        self.assertEqual(
            comparator_chunk.country,
            "Spain",
        )

        self.assertEqual(
            comparator_chunk.country_code,
            "ES",
        )

        self.assertEqual(
            comparator_chunk.reference_year,
            2026,
        )

    def test_ids_are_deterministic(self) -> None:
        original_sections = [
            ParsedSection(
                section="02. Employment Contracts",
                subsection="Notice Period",
                content="Original legal content.",
            )
        ]

        updated_sections = [
            ParsedSection(
                section="02. Employment Contracts",
                subsection="Notice Period",
                content="Updated legal content.",
            )
        ]

        original_chunks = build_document_chunks(
            parsed_sections=original_sections,
            metadata=self.metadata,
        )

        repeated_chunks = build_document_chunks(
            parsed_sections=original_sections,
            metadata=self.metadata,
        )

        updated_chunks = build_document_chunks(
            parsed_sections=updated_sections,
            metadata=self.metadata,
        )

        self.assertEqual(
            original_chunks[0].document_id,
            repeated_chunks[0].document_id,
        )

        self.assertEqual(
            original_chunks[0].chunk_id,
            repeated_chunks[0].chunk_id,
        )

        self.assertEqual(
            original_chunks[0].chunk_id,
            updated_chunks[0].chunk_id,
        )

        self.assertNotEqual(
            original_chunks[0].content_hash,
            updated_chunks[0].content_hash,
        )

    def test_rejects_unknown_legal_topic(self) -> None:
        parsed_sections = [
            ParsedSection(
                section="12. Imaginary Legal Topic",
                subsection="Unknown subsection",
                content="Content that must not be indexed.",
            )
        ]

        with self.assertRaises(
            UnknownLegalTopicError
        ):
            build_document_chunks(
                parsed_sections=parsed_sections,
                metadata=self.metadata,
            )

    def test_extracts_metadata_from_content(
        self,
    ) -> None:
        # Mission "CONTINUATION PATCH 0.4.3": metadata now comes from
        # the document's own title/cover content, never its filename
        # - the arbitrary filename here proves that directly.
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Labour and Employment Law in Spain 2026"],
                filename="final.docx",
            )

            metadata = metadata_from_content(
                file_path=file_path,
                country_code="es",
            )

        self.assertEqual(metadata.country, "Spain")
        self.assertEqual(metadata.country_code, "ES")
        self.assertEqual(metadata.reference_year, 2026)
        self.assertEqual(metadata.language, "en")
        self.assertEqual(metadata.source_filename, "final.docx")

    def test_rejects_content_with_no_identifiable_country(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Some random legal memo with no title structure."],
            )

            with self.assertRaises(
                UndeterminableDocumentCountryError
            ):
                metadata_from_content(file_path=file_path)

    def test_title_with_leading_definite_article_resolves_country(
        self,
    ) -> None:
        # Mission "ORDER 2": Czech Republic's real production source
        # reads "Labour and employment law in the Czech Republic" -
        # with the article - and the registry had no "the Czech
        # Republic" alias, so this raised
        # UndeterminableDocumentCountryError on Reindex (masked by the
        # router into a generic 502). Reproduces the exact title
        # phrasing, not just the country name in isolation.
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                [
                    "Labour and employment law in the Czech Republic"
                ],
            )

            metadata = metadata_from_content(
                file_path=file_path
            )

        self.assertEqual(metadata.country, "Czech Republic")
        self.assertEqual(metadata.country_code, "CZ")

    def test_uk_content_uses_canonical_country_name(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Labour and Employment Law in UK 2026"],
            )

            metadata = metadata_from_content(
                file_path=file_path,
                country_code="GB",
            )

        self.assertEqual(metadata.country, "United Kingdom")
        self.assertEqual(metadata.country_code, "GB")
        self.assertEqual(metadata.reference_year, 2026)

    def test_infers_country_code_from_content(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Labour and Employment Law in Sweden 2026"],
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Sweden")
        self.assertEqual(metadata.country_code, "SE")
        self.assertEqual(metadata.reference_year, 2026)

    def test_accepts_overview_title_without_year(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Australia"],
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Australia")
        self.assertEqual(metadata.country_code, "AU")
        self.assertIsNone(metadata.reference_year)

    def test_an_arbitrary_filename_never_affects_detection(
        self,
    ) -> None:
        # The filename names a different country than the content -
        # the content must always win, and the filename is preserved
        # only as display metadata (mission section 5).
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Canada 2026"],
                filename="Spain-template-used-for-Canada.docx",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Canada")
        self.assertEqual(metadata.country_code, "CA")
        self.assertEqual(
            metadata.source_filename,
            "Spain-template-used-for-Canada.docx",
        )

    def test_rejects_content_country_code_mismatch(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Labour and Employment Law in Spain 2026"],
            )

            with self.assertRaises(
                CountryMetadataMismatchError
            ):
                metadata_from_content(
                    file_path=file_path,
                    country_code="GB",
                )


class ContentMetadataFixtureTests(unittest.TestCase):
    """
    Mission "CONTINUATION PATCH 0.4.3", section 15 - the 8 mandatory
    filename/content/expected-outcome fixture cases A-H.
    """

    def test_case_a_plain_filename_with_title_country_and_year(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Canada 2026"],
                filename="final.docx",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Canada")
        self.assertEqual(metadata.country_code, "CA")
        self.assertEqual(metadata.reference_year, 2026)

    def test_case_b_filename_names_a_different_country_than_content(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Labour and Employment Law in Canada 2025"],
                filename="Spain-final.docx",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Canada")
        self.assertEqual(metadata.country_code, "CA")
        self.assertEqual(metadata.reference_year, 2025)
        # No conflict is ever raised - the filename is never even
        # consulted, so "Spain" in it has no bearing on the result.
        self.assertEqual(metadata.source_filename, "Spain-final.docx")

    def test_case_c_edited_replacement_filename_with_matching_content(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Canada 2026"],
                filename=(
                    "Canada_2026-04-15-Employment-Law-"
                    "Overview-EDITED.docx"
                ),
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Canada")
        self.assertEqual(metadata.country_code, "CA")
        self.assertEqual(metadata.reference_year, 2026)

    def test_case_d_generic_filename_with_no_year_in_content(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Spain"],
                filename="document_received_from_client.docx",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Spain")
        self.assertEqual(metadata.country_code, "ES")
        self.assertIsNone(metadata.reference_year)

    def test_case_e_filename_names_a_country_but_content_names_none(
        self,
    ) -> None:
        # The filename "Peru-2026.docx" must never be used as a
        # fallback - content with no identifiable country is refused
        # outright, never silently resolved to "Peru".
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Some random legal memo with no title structure."],
                filename="Peru-2026.docx",
            )

            with self.assertRaises(
                UndeterminableDocumentCountryError
            ):
                metadata_from_content(file_path=file_path)

    def test_case_f_ambiguous_cover_naming_two_countries_is_refused(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                [
                    "Employment Law Overview Spain 2026",
                    "Employment Law Overview Canada 2026",
                ],
            )

            with self.assertRaises(
                AmbiguousDocumentCountryError
            ):
                metadata_from_content(file_path=file_path)

    def test_case_g_body_mentions_of_other_countries_are_ignored(
        self,
    ) -> None:
        # Countries named only in ordinary body prose (never shaped
        # as a title line) must never be treated as candidates - only
        # the one clear cover country is selected.
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Canada 2026"],
                body_paragraphs=[
                    "This overview also references comparable rules "
                    "in France, Germany, and Peru for context.",
                    "See also the equivalent Spain and Australia "
                    "frameworks discussed elsewhere.",
                ],
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Canada")
        self.assertEqual(metadata.country_code, "CA")

    def test_case_h_content_year_wins_over_filename_year(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Canada 2025"],
                filename="Canada-2026-report.docx",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country_code, "CA")
        self.assertEqual(metadata.reference_year, 2025)


def _valid_docx_entries(directory: Path) -> dict[str, bytes]:
    """The raw ZIP entries of one real, minimal, valid DOCX."""

    valid_path = directory / "valid-reference.docx"
    document = Document()
    document.add_paragraph("Employment Law Overview Canada 2026")
    document.save(valid_path)

    with zipfile.ZipFile(valid_path) as archive:
        return {name: archive.read(name) for name in archive.namelist()}


class InvalidDocxFormatTests(unittest.TestCase):
    """
    Mission "CONTINUATION PATCH 0.4.3", section 16 - the .docx
    extension alone proves nothing: each of these must be refused by
    validate_docx_format before any content parsing is attempted.
    """

    def test_renamed_text_file_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = Path(directory) / "renamed.docx"
            file_path.write_bytes(
                b"This is just plain text, not a docx at all."
            )

            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_invalid_zip_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = Path(directory) / "invalid-zip.docx"
            # Starts with the ZIP local-file-header magic bytes but has
            # no valid central directory - not a real archive.
            file_path.write_bytes(b"PK\x03\x04" + b"\x00" * 50)

            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_empty_file_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = Path(directory) / "empty.docx"
            file_path.write_bytes(b"")

            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_zip_missing_content_types_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            entries = _valid_docx_entries(Path(directory))
            file_path = Path(directory) / "missing-content-types.docx"

            with zipfile.ZipFile(file_path, "w") as archive:
                for name, data in entries.items():
                    if name == "[Content_Types].xml":
                        continue
                    archive.writestr(name, data)

            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_zip_missing_document_xml_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            entries = _valid_docx_entries(Path(directory))
            file_path = Path(directory) / "missing-document-xml.docx"

            with zipfile.ZipFile(file_path, "w") as archive:
                for name, data in entries.items():
                    if name == "word/document.xml":
                        continue
                    archive.writestr(name, data)

            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_corrupted_document_xml_is_rejected(self) -> None:
        # Both required entries are present by name - only their
        # content is broken, so this exercises python-docx's own
        # Document() openability check specifically.
        with TemporaryDirectory() as directory:
            entries = dict(_valid_docx_entries(Path(directory)))
            entries["word/document.xml"] = b"<this is not valid xml"
            file_path = Path(directory) / "corrupted.docx"

            with zipfile.ZipFile(file_path, "w") as archive:
                for name, data in entries.items():
                    archive.writestr(name, data)

            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)


class RealDocumentStructureTests(unittest.TestCase):
    """
    Mission "HOTFIX 0.4.4", Mission 2/2, section 11 - the real-world
    document structures the original 17-document corpus never used:
    title in a text box, DrawingML, header, a document with no
    standard cover, roman-numeral topic headings, decomposed Unicode,
    a law's own year ignored, filename ignored, and an ambiguous
    cover - each proven with a minimal synthetic fixture, mirroring
    the 10 real production documents this mission validates.
    """

    def test_title_in_a_text_box_is_found(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx_with_text_box_title(
                Path(directory),
                "Employment Law Overview Chile",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Chile")
        self.assertEqual(metadata.country_code, "CL")

    def test_title_in_a_text_box_with_en_dash_separator(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx_with_text_box_title(
                Path(directory),
                "Employment Law Overview – Taiwan",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Taiwan")
        self.assertEqual(metadata.country_code, "TW")

    def test_title_in_drawingml_is_found(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx_with_drawingml_title(
                Path(directory),
                "Employment Law Overview Ireland",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Ireland")
        self.assertEqual(metadata.country_code, "IE")

    def test_title_in_a_header_is_found(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx_with_header_title(
                Path(directory),
                "Employment Law Overview Colombia",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Colombia")
        self.assertEqual(metadata.country_code, "CO")

    def test_all_caps_title_with_no_separator_is_found(self) -> None:
        # "EMPLOYMENT LAW OVERVIEW PHILIPPINES" - the real Philippines
        # document's own exact cover-paragraph shape.
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["EMPLOYMENT LAW OVERVIEW PHILIPPINES"],
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Philippines")
        self.assertEqual(metadata.country_code, "PH")

    def test_reversed_cover_country_then_plural_heading_with_year_range(
        self,
    ) -> None:
        # The real France document's own cover: a bare country name,
        # then a heading with a plural "Overviews" and a year range -
        # the reverse order and wording of the usual two-line cover.
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["FRANCE", "EMPLOYMENT LAW OVERVIEWS 2025 - 2026"],
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "France")
        self.assertEqual(metadata.country_code, "FR")

    def test_reversed_cover_country_then_heading_with_template_boilerplate(
        self,
    ) -> None:
        # The real Portugal document's own cover: a bare country
        # name, then a heading wrapped in unrelated boilerplate words
        # on both sides ("Country-Specific ... Template").
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                [
                    "PORTUGAL",
                    "COUNTRY-SPECIFIC EMPLOYMENT LAW OVERVIEWS 2026 TEMPLATE",
                ],
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Portugal")
        self.assertEqual(metadata.country_code, "PT")

    def test_reversed_cover_fallback_never_silently_misdetects_an_unrelated_country(
        self,
    ) -> None:
        # The reversed-order fallback's "next line merely contains
        # the family heading phrase" check is deliberately loose - an
        # unrelated bare country name sitting directly before some
        # other line that happens to contain "Employment Law
        # Overview" (e.g. a real cover's own one-line title for the
        # actual country) must never be silently accepted as *the*
        # document's country. The pre-existing ambiguity refusal
        # (more than one distinct country code found in the front
        # matter) is what has to catch this - proven here rather than
        # merely assumed.
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Spain", "Employment Law Overview Canada"],
            )

            with self.assertRaises(
                AmbiguousDocumentCountryError
            ):
                metadata_from_content(file_path=file_path)

    def test_document_with_no_standard_cover_is_refused(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Table of Contents", "Some unrelated preamble text."],
            )

            with self.assertRaises(
                UndeterminableDocumentCountryError
            ):
                metadata_from_content(file_path=file_path)

    def test_decomposed_unicode_turkiye_is_recognized(self) -> None:
        # "Turkiye" typed with a combining diaeresis (u + U+0308)
        # instead of the precomposed "ü" (U+00FC) must still resolve,
        # via NFKC normalization.
        decomposed_turkiye = "Türkiye"

        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                [f"Employment Law Overview {decomposed_turkiye}"],
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Türkiye")
        self.assertEqual(metadata.country_code, "TR")

    def test_a_laws_own_year_is_never_used_as_reference_year(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Chile"],
                body_paragraphs=[
                    "The Labour Code was originally enacted in 1984 "
                    "and has been amended many times since.",
                ],
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country_code, "CL")
        self.assertIsNone(metadata.reference_year)

    def test_filename_country_and_year_are_ignored(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Colombia"],
                filename="Chile_2020_FINAL-EDITED.docx",
            )

            metadata = metadata_from_content(file_path=file_path)

        self.assertEqual(metadata.country, "Colombia")
        self.assertEqual(metadata.country_code, "CO")
        self.assertIsNone(metadata.reference_year)
        self.assertEqual(
            metadata.source_filename,
            "Chile_2020_FINAL-EDITED.docx",
        )

    def test_ambiguous_cover_among_new_countries_is_refused(
        self,
    ) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                [
                    "Employment Law Overview Chile",
                    "Employment Law Overview Colombia",
                ],
            )

            with self.assertRaises(
                AmbiguousDocumentCountryError
            ):
                metadata_from_content(file_path=file_path)

    def test_roman_numeral_topic_heading_is_recognized(self) -> None:
        # Taiwan's own real structure: roman-numeral, all-caps main
        # headings ("II.  HIRING PRACTICES") rather than the
        # "01. Hiring Practices" convention the original corpus uses.
        # ("I. GENERAL OVERVIEW" itself is exercised by the end-to-end
        # test below instead: build_document_chunks validates every
        # explicit ParsedSection strictly, while the real parser
        # simply absorbs an unrecognized overview-shaped heading into
        # whichever section is already open - never promoting it to
        # its own ParsedSection at all.)
        parsed_sections = [
            ParsedSection(
                section="II.  Hiring Practices",
                subsection="1. Introduction",
                content="Hiring content.",
            ),
        ]

        metadata = DocumentMetadata(
            country="Taiwan",
            country_code="TW",
            reference_year=None,
            language="en",
            source_filename="taiwan.docx",
        )

        chunks = build_document_chunks(
            parsed_sections=parsed_sections,
            metadata=metadata,
        )

        topics = {
            chunk.legal_topic
            for chunk in chunks
            if chunk.legal_topic is not None
        }

        self.assertIn("Hiring Practices", topics)

    def test_roman_numeral_heading_recognized_end_to_end_from_real_docx(
        self,
    ) -> None:
        # Same Taiwan-shaped structure, this time parsed from an
        # actual DOCX (plain, non-bold, non-heading-styled paragraphs,
        # exactly like the real Taiwan document) through the full
        # build_document_chunks_from_docx pipeline.
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview Taiwan"],
                body_paragraphs=[
                    "I. GENERAL OVERVIEW",
                    "1. Introduction",
                    "Some overview content.",
                    "II.  Hiring Practices",
                    "Hiring content for Taiwan.",
                ],
            )

            chunks = build_document_chunks_from_docx(
                file_path=file_path,
                country_code="TW",
                language="en",
            )

        topics = {
            chunk.legal_topic
            for chunk in chunks
            if chunk.legal_topic is not None
        }

        self.assertIn("Hiring Practices", topics)

    def test_usa_jurisdiction_suffix_heading_text_is_unchanged(
        self,
    ) -> None:
        # Mission "HOTFIX 0.4.4", final targeted correction, test 8:
        # the topic-matching normalization used to recognize "in the
        # USA" must never alter the section heading actually stored
        # on the chunk.
        with TemporaryDirectory() as directory:
            file_path = _build_docx(
                Path(directory),
                ["Employment Law Overview United States"],
                body_paragraphs=[
                    "06. Social Media and Data Privacy in the USA",
                    "Content about social media rules.",
                ],
            )

            chunks = build_document_chunks_from_docx(
                file_path=file_path,
                country_code="US",
                language="en",
            )

        matching_chunks = [
            chunk
            for chunk in chunks
            if chunk.legal_topic == "Social Media and Data Privacy"
        ]

        self.assertTrue(matching_chunks)
        self.assertEqual(
            matching_chunks[0].section,
            "06. Social Media and Data Privacy in the USA",
        )


if __name__ == "__main__":
    unittest.main()