"""
Tests for the DOCX-native country marker (mission "ORDER 8E-A1",
sections 10/37) - a standard OOXML custom document property, read and
written without any python-docx-native support (confirmed absent in
1.2.0), preserving every other part of the archive and producing
byte-identical output for byte-identical logical input.
"""

from __future__ import annotations

import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from app.services.docx_country_marker import (
    CountryMarker,
    InvalidCountryMarkerValueError,
    read_country_marker,
    write_country_marker,
)


def _build_docx(directory: Path, paragraphs: list[str]) -> Path:
    document = Document()

    for text in paragraphs:
        document.add_paragraph(text)

    path = directory / "document.docx"
    document.save(path)

    return path


class ReadWithNoMarkerTests(unittest.TestCase):
    def test_document_with_no_custom_properties_part_reads_as_none(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            path = _build_docx(Path(root), ["No marker here."])

            self.assertIsNone(read_country_marker(path))


class WriteAndReadRoundtripTests(unittest.TestCase):
    def test_write_then_read_returns_the_same_marker(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Some content."])
            destination = Path(root) / "marked.docx"

            write_country_marker(
                source,
                destination,
                country_code="fr",
                country_name="France",
            )

            marker = read_country_marker(destination)

            self.assertEqual(
                marker,
                CountryMarker(
                    country_code="FR",
                    country_name="France",
                ),
            )

    def test_marker_survives_a_full_python_docx_load_and_save(
        self,
    ) -> None:
        # The Admin's own section Edit/Add workflow loads the current
        # DOCX via python-docx, mutates it, and saves it back out - the
        # marker must never be silently dropped by that round-trip.
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Some content."])
            marked = Path(root) / "marked.docx"

            write_country_marker(
                source,
                marked,
                country_code="de",
                country_name="Germany",
            )

            document = Document(marked)
            document.add_paragraph("An unrelated edit.")
            edited = Path(root) / "edited.docx"
            document.save(edited)

            self.assertEqual(
                read_country_marker(edited),
                CountryMarker(
                    country_code="DE",
                    country_name="Germany",
                ),
            )

    def test_marked_document_is_still_a_valid_docx(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Body paragraph one."])
            marked = Path(root) / "marked.docx"

            write_country_marker(
                source,
                marked,
                country_code="jp",
                country_name="Japan",
            )

            reopened = Document(marked)

            self.assertEqual(
                [p.text for p in reopened.paragraphs],
                ["Body paragraph one."],
            )


class DeterminismTests(unittest.TestCase):
    def test_writing_the_same_marker_twice_is_byte_identical(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Content."])
            first = Path(root) / "first.docx"
            second = Path(root) / "second.docx"

            write_country_marker(
                source, first, country_code="ca", country_name="Canada"
            )
            write_country_marker(
                source, second, country_code="ca", country_name="Canada"
            )

            self.assertEqual(
                first.read_bytes(), second.read_bytes()
            )

    def test_re_embedding_on_an_already_marked_file_is_idempotent(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Content."])
            once = Path(root) / "once.docx"
            twice = Path(root) / "twice.docx"

            write_country_marker(
                source, once, country_code="ca", country_name="Canada"
            )
            write_country_marker(
                once, twice, country_code="ca", country_name="Canada"
            )

            self.assertEqual(
                once.read_bytes(), twice.read_bytes()
            )

    def test_changing_the_country_updates_the_value_not_the_pid(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Content."])
            first = Path(root) / "first.docx"
            changed = Path(root) / "changed.docx"

            write_country_marker(
                source, first, country_code="fr", country_name="France"
            )
            write_country_marker(
                first,
                changed,
                country_code="de",
                country_name="Germany",
            )

            with zipfile.ZipFile(first) as archive:
                first_xml = archive.read("docProps/custom.xml")

            with zipfile.ZipFile(changed) as archive:
                changed_xml = archive.read("docProps/custom.xml")

            # Same pid="2" in both - only the value changed.
            self.assertIn(b'pid="2"', first_xml)
            self.assertIn(b'pid="2"', changed_xml)
            self.assertEqual(
                read_country_marker(changed),
                CountryMarker(
                    country_code="DE", country_name="Germany"
                ),
            )


class InvalidMarkerTests(unittest.TestCase):
    def test_write_rejects_an_unrecognized_country_code(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Content."])

            with self.assertRaises(InvalidCountryMarkerValueError):
                write_country_marker(
                    source,
                    Path(root) / "invalid.docx",
                    country_code="ZZ",
                    country_name="Nowhere",
                )

    def test_read_safely_ignores_a_corrupted_marker_value(
        self,
    ) -> None:
        # A garbage value in an otherwise well-formed custom.xml (from
        # a foreign tool, or file corruption) must never be trusted as
        # a real country - it is silently treated as absent, exactly
        # like a document with no marker at all.
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Content."])
            marked = Path(root) / "marked.docx"
            write_country_marker(
                source, marked, country_code="fr", country_name="France"
            )

            with zipfile.ZipFile(marked) as archive:
                contents = {
                    name: archive.read(name)
                    for name in archive.namelist()
                }

            contents["docProps/custom.xml"] = contents[
                "docProps/custom.xml"
            ].replace(b">FR<", b">NOT-A-COUNTRY<")

            corrupted = Path(root) / "corrupted.docx"
            with zipfile.ZipFile(
                corrupted, "w", zipfile.ZIP_DEFLATED
            ) as archive:
                for name, data in contents.items():
                    archive.writestr(name, data)

            self.assertIsNone(read_country_marker(corrupted))

    def test_write_rejects_an_empty_country_name(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Content."])

            with self.assertRaises(InvalidCountryMarkerValueError):
                write_country_marker(
                    source,
                    Path(root) / "invalid.docx",
                    country_code="FR",
                    country_name="   ",
                )


class ForeignCustomPropertiesPreservedTests(unittest.TestCase):
    def test_a_pre_existing_unrelated_custom_property_survives(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _build_docx(Path(root), ["Content."])
            with_foreign_property = Path(root) / "foreign.docx"

            with zipfile.ZipFile(source) as archive:
                contents = {
                    name: archive.read(name)
                    for name in archive.namelist()
                }

            custom_xml = (
                '<?xml version="1.0" encoding="UTF-8" '
                'standalone="yes"?>\r\n'
                '<Properties xmlns="http://schemas.openxmlformats.org'
                '/officeDocument/2006/custom-properties" '
                'xmlns:vt="http://schemas.openxmlformats.org/'
                'officeDocument/2006/docPropsVTypes">'
                '<property fmtid="{D5CDD505-2E9C-101B-9397-'
                '08002B2CF9AE}" pid="2" name="SomeOtherTool">'
                "<vt:lpwstr>keep-me</vt:lpwstr></property>"
                "</Properties>"
            ).encode("utf-8")
            contents["docProps/custom.xml"] = custom_xml
            contents["[Content_Types].xml"] = contents[
                "[Content_Types].xml"
            ].replace(
                b"</Types>",
                (
                    b'<Override PartName="/docProps/custom.xml" '
                    b'ContentType="application/vnd.openxmlformats-'
                    b'officedocument.custom-properties+xml"/></Types>'
                ),
            )
            contents["_rels/.rels"] = contents["_rels/.rels"].replace(
                b"</Relationships>",
                (
                    b'<Relationship Id="rId99" Type="http://schemas.'
                    b"openxmlformats.org/officeDocument/2006/"
                    b'relationships/custom-properties" '
                    b'Target="docProps/custom.xml"/></Relationships>'
                ),
            )

            with zipfile.ZipFile(
                with_foreign_property, "w", zipfile.ZIP_DEFLATED
            ) as archive:
                for name, data in contents.items():
                    archive.writestr(name, data)

            marked = Path(root) / "marked-with-foreign.docx"
            write_country_marker(
                with_foreign_property,
                marked,
                country_code="pt",
                country_name="Portugal",
            )

            with zipfile.ZipFile(marked) as archive:
                final_custom_xml = archive.read(
                    "docProps/custom.xml"
                )

            self.assertIn(b"SomeOtherTool", final_custom_xml)
            self.assertIn(b"keep-me", final_custom_xml)
            self.assertEqual(
                read_country_marker(marked),
                CountryMarker(
                    country_code="PT", country_name="Portugal"
                ),
            )
            # python-docx must still be able to open it cleanly.
            Document(marked)


if __name__ == "__main__":
    unittest.main()
