"""
Real ASGI/HTTP-transport tests for the admin document endpoints
(mission "HOTFIX 0.4.9" review 2, sections 11-15).

backend/tests/test_admin_documents_router_integration.py calls the
FastAPI route functions directly - useful, but not a transport-level
test: it never exercises ApiProtectionMiddleware, FastAPI's own
request routing, or its automatic multipart/RequestValidationError
handling. This file does, using a minimal, hand-rolled ASGI client
(scope/receive/send, driven by asyncio.run) rather than adding httpx
as a new dependency - it is not installed in this project
(fastapi.testclient.TestClient hard-requires it), and the mission
explicitly asks not to add a heavy dependency just for this. The
approach here is exactly what httpx.ASGITransport itself does
internally: call the real ASGI application callable directly with a
synthetic scope, never a real socket, never a real port, never real
production settings/services.

Every test patches get_settings()/get_opensearch_client()/
get_redis_client() (all @lru_cache) to point at temp directories, a
fake in-memory OpenSearch, and a fake Redis that always allows the
request through - never a real cluster, never a real port, never
production credentials.
"""

from __future__ import annotations

import asyncio
import io
import json
import os
import tempfile
import unittest
import uuid
from pathlib import Path
from typing import Any
from unittest.mock import patch

from docx import Document

from app.core.config import get_settings


TEST_API_ACCESS_KEY = "test-only-asgi-api-access-key"
TEST_ADMIN_API_KEY = "test-only-asgi-admin-api-key"

_SETTINGS_ENV_KEYS = (
    "OPENSEARCH_URL",
    "OPENSEARCH_PASSWORD",
    "REDIS_URL",
    "DOCUMENT_SOURCE_DIR",
    "DOCUMENT_PROCESSED_DIR",
    "ADMIN_API_KEY",
    "API_ACCESS_KEY",
)


class ASGIResponse:
    """A minimal stand-in for httpx.Response - just what these tests
    need: status_code, header lookup, raw body, and .json()."""

    def __init__(
        self,
        status_code: int,
        headers: list[tuple[bytes, bytes]],
        body: bytes,
    ) -> None:
        self.status_code = status_code
        self._headers = headers
        self.content = body

    def header(self, name: str) -> str | None:
        target = name.lower().encode()

        for key, value in self._headers:
            if key.lower() == target:
                return value.decode()

        return None

    def json(self) -> Any:
        return json.loads(self.content)


async def _asgi_request(
    app: Any,
    method: str,
    path: str,
    headers: dict[str, str] | None = None,
    body: bytes = b"",
) -> ASGIResponse:
    """
    Drive the real ASGI application callable directly with a
    synthetic in-memory scope/receive/send - the same protocol
    httpx.ASGITransport itself implements, without needing httpx.
    """

    request_headers = [
        (key.lower().encode(), value.encode())
        for key, value in (headers or {}).items()
    ]

    scope = {
        "type": "http",
        "asgi": {"version": "3.0", "spec_version": "2.3"},
        "http_version": "1.1",
        "method": method,
        "path": path,
        "raw_path": path.encode(),
        "query_string": b"",
        "root_path": "",
        "headers": request_headers,
        "client": ("127.0.0.1", 12345),
        "server": ("127.0.0.1", 80),
        "scheme": "http",
        "state": {},
    }

    body_delivered = {"done": False}

    async def receive() -> dict[str, Any]:
        if body_delivered["done"]:
            return {"type": "http.disconnect"}

        body_delivered["done"] = True

        return {
            "type": "http.request",
            "body": body,
            "more_body": False,
        }

    collected: dict[str, Any] = {
        "status": None,
        "headers": [],
        "body": b"",
    }

    async def send(message: dict[str, Any]) -> None:
        if message["type"] == "http.response.start":
            collected["status"] = message["status"]
            collected["headers"] = message["headers"]

        elif message["type"] == "http.response.body":
            collected["body"] += message.get("body", b"")

    await app(scope, receive, send)

    return ASGIResponse(
        collected["status"],
        collected["headers"],
        collected["body"],
    )


def asgi_request(
    app: Any,
    method: str,
    path: str,
    headers: dict[str, str] | None = None,
    body: bytes = b"",
) -> ASGIResponse:
    return asyncio.run(
        _asgi_request(app, method, path, headers=headers, body=body)
    )


def build_multipart_body(
    *,
    file_field: str,
    filename: str,
    file_content: bytes,
    extra_fields: dict[str, str] | None = None,
    include_file: bool = True,
) -> tuple[bytes, str]:
    """Build a real multipart/form-data body byte-for-byte, the same
    shape a browser or the WordPress PHP proxy would send."""

    boundary = f"----ASGITestBoundary{uuid.uuid4().hex}"
    lines: list[bytes] = []

    for name, value in (extra_fields or {}).items():
        lines.append(f"--{boundary}".encode())
        lines.append(
            f'Content-Disposition: form-data; name="{name}"'.encode()
        )
        lines.append(b"")
        lines.append(value.encode())

    if include_file:
        lines.append(f"--{boundary}".encode())
        lines.append(
            (
                f'Content-Disposition: form-data; name="{file_field}"; '
                f'filename="{filename}"'
            ).encode()
        )
        lines.append(
            b"Content-Type: application/vnd.openxmlformats-"
            b"officedocument.wordprocessingml.document"
        )
        lines.append(b"")
        lines.append(file_content)

    lines.append(f"--{boundary}--".encode())
    lines.append(b"")

    body = b"\r\n".join(lines)
    content_type = f"multipart/form-data; boundary={boundary}"

    return body, content_type


def _build_minimal_docx_bytes(country_line: str) -> bytes:
    document = Document()
    document.add_paragraph(country_line)
    document.add_paragraph("I. GENERAL OVERVIEW")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Overview content.")
    document.add_paragraph("II. Hiring Practices")
    document.add_paragraph("Hiring content.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_CL_DOCX_BYTES = _build_minimal_docx_bytes(
    "Labour and Employment Law in Chile 2026"
)


class FakeOpenSearch:
    """Chunk-granular in-memory OpenSearch double - see
    test_admin_documents_router_integration.py's FakeOpenSearch for
    the same, already-established shape."""

    def __init__(self) -> None:
        self.chunks: dict[str, dict[str, Any]] = {}

    def add(
        self,
        *,
        document_id: str,
        country_code: str,
        source_filename: str,
        chunk_id: str | None = None,
    ) -> None:
        resolved_chunk_id = chunk_id or f"{document_id}-chunk-0"
        self.chunks[resolved_chunk_id] = {
            "document_id": document_id,
            "chunk_id": resolved_chunk_id,
            "country_code": country_code,
            "source_filename": source_filename,
            "country": country_code,
            "reference_year": 2026,
        }

    def document_ids_for(self, country_code: str) -> set[str]:
        return {
            c["document_id"]
            for c in self.chunks.values()
            if c["country_code"] == country_code
        }

    def search(self, *, index: str, body: dict[str, Any]) -> dict[str, Any]:
        if "aggs" in body and "query" not in body:
            # The admin catalog listing (GET /api/v1/admin/documents)
            # - a terms aggregation over document_id with a top_hits
            # sub-aggregation, never a query.term shape.
            by_document: dict[str, list[dict[str, Any]]] = {}
            for chunk in self.chunks.values():
                by_document.setdefault(
                    chunk["document_id"], []
                ).append(chunk)

            return {
                "aggregations": {
                    "documents": {
                        "buckets": [
                            {
                                "key": document_id,
                                "doc_count": len(chunks),
                                "metadata": {
                                    "hits": {
                                        "hits": [
                                            {
                                                "_source": {
                                                    "document_id": (
                                                        document_id
                                                    ),
                                                    "source_filename": (
                                                        chunks[0][
                                                            "source_filename"
                                                        ]
                                                    ),
                                                    "country": (
                                                        chunks[0][
                                                            "country"
                                                        ]
                                                    ),
                                                    "country_code": (
                                                        chunks[0][
                                                            "country_code"
                                                        ]
                                                    ),
                                                    "language": "en",
                                                    "document_type": (
                                                        "comparator"
                                                    ),
                                                    "reference_year": (
                                                        chunks[0][
                                                            "reference_year"
                                                        ]
                                                    ),
                                                }
                                            }
                                        ]
                                    }
                                },
                            }
                            for document_id, chunks in (
                                by_document.items()
                            )
                        ]
                    }
                }
            }

        term = body["query"].get("term") or {}

        if "country_code" in term:
            # One hit per real chunk (never deduplicated by
            # document_id) - matches real OpenSearch's own per-chunk
            # granularity; callers that want distinct documents
            # already deduplicate client-side (mission "ORDER 3B",
            # section 6 - keep this fake's shape faithful to what
            # test_admin_documents_router_integration.py's own
            # FakeOpenSearch already documents).
            hits = [
                c for c in self.chunks.values()
                if c["country_code"] == term["country_code"]
            ]
        elif "document_id" in term:
            hits = [
                c for c in self.chunks.values()
                if c["document_id"] == term["document_id"]
            ]
        else:
            hits = list(self.chunks.values())

        return {
            "hits": {
                "total": {"value": len(hits)},
                "hits": [
                    {
                        "_id": h["chunk_id"],
                        "_source": h,
                        # Mission "ORDER 3B": real OpenSearch 3.7
                        # includes a "sort" array whenever the request
                        # itself carries a "sort" clause, as
                        # _fetch_all_chunks's search_after pagination
                        # always does.
                        "sort": [h["chunk_id"]],
                    }
                    for h in sorted(
                        hits, key=lambda c: c["chunk_id"]
                    )
                ],
            }
        }

    def delete_by_query(self, **kwargs: Any) -> dict[str, Any]:
        del kwargs
        return {"deleted": 0}


def _bulk_fake(fake: FakeOpenSearch, actions: Any, **kwargs: Any):
    del kwargs
    action_list = list(actions)

    for action in action_list:
        doc = action["_source"]
        fake.add(
            document_id=doc["document_id"],
            chunk_id=doc["chunk_id"],
            country_code=doc["country_code"],
            source_filename=doc["source_filename"],
        )

    return (len(action_list), [])


class FakeRedis:
    """Always allows the request through - a real rate-limit backend
    is unrelated to what these tests verify."""

    def eval(self, script: str, numkeys: int, *args: Any) -> list[int]:
        del script, numkeys, args

        return [1, 60]


class AdminAsgiTestCase(unittest.TestCase):
    """Fresh temp dirs, fake OpenSearch, fake Redis, real test API/
    admin keys - the real FastAPI ASGI app, imported fresh under
    these settings."""

    def setUp(self) -> None:
        self._tempdir = tempfile.TemporaryDirectory()
        root = Path(self._tempdir.name)
        self.source_dir = root / "source"
        self.processed_dir = root / "processed"
        self.source_dir.mkdir(parents=True)
        self.processed_dir.mkdir(parents=True)

        self._original_env = {
            key: os.environ.get(key) for key in _SETTINGS_ENV_KEYS
        }
        os.environ["OPENSEARCH_URL"] = "http://unused-in-asgi-tests:9200"
        os.environ["OPENSEARCH_PASSWORD"] = "unused"
        os.environ["REDIS_URL"] = "redis://unused-in-asgi-tests:6379/0"
        os.environ["DOCUMENT_SOURCE_DIR"] = str(self.source_dir)
        os.environ["DOCUMENT_PROCESSED_DIR"] = str(self.processed_dir)
        os.environ["ADMIN_API_KEY"] = TEST_ADMIN_API_KEY
        os.environ["API_ACCESS_KEY"] = TEST_API_ACCESS_KEY
        get_settings.cache_clear()

        self.fake_opensearch = FakeOpenSearch()
        self.fake_redis = FakeRedis()

        self._patches = [
            patch(
                "app.services.admin_document_replacement."
                "get_opensearch_client",
                return_value=self.fake_opensearch,
            ),
            patch(
                "app.services.admin_document_lifecycle."
                "get_opensearch_client",
                return_value=self.fake_opensearch,
            ),
            patch(
                "app.services.admin_documents.get_opensearch_client",
                return_value=self.fake_opensearch,
            ),
            patch(
                "app.services.document_indexer.get_opensearch_client",
                return_value=self.fake_opensearch,
            ),
            patch(
                "app.services.document_indexer.bulk",
                side_effect=(
                    lambda client, actions, **kw: _bulk_fake(
                        self.fake_opensearch, actions, **kw
                    )
                ),
            ),
            patch(
                "app.services.document_indexer."
                "ensure_legal_documents_index"
            ),
            patch(
                "app.middleware.api_protection.get_redis_client",
                return_value=self.fake_redis,
            ),
        ]
        for one_patch in self._patches:
            one_patch.start()

        # Imported only after settings/patches are in place, and
        # freshly re-imported per test - app.main constructs its
        # FastAPI() instance and registers middleware at import time.
        import importlib

        import app.main as main_module

        importlib.reload(main_module)
        self.app = main_module.app

    def tearDown(self) -> None:
        for one_patch in reversed(self._patches):
            one_patch.stop()

        for key, value in self._original_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value

        get_settings.cache_clear()
        self._tempdir.cleanup()

    def _auth_headers(self) -> dict[str, str]:
        return {
            "X-API-Key": TEST_API_ACCESS_KEY,
            "X-Admin-Key": TEST_ADMIN_API_KEY,
        }


class FreshUploadAsgiTests(AdminAsgiTestCase):
    def test_asgi_fresh_multipart_upload_returns_201(self) -> None:
        body, content_type = build_multipart_body(
            file_field="file",
            filename="Chile.docx",
            file_content=_CL_DOCX_BYTES,
            extra_fields={
                "replace_existing": "false",
                "confirm_warnings": "true",
                "country_confirmed": "true",
            },
        )

        response = asgi_request(
            self.app,
            "POST",
            "/api/v1/admin/documents",
            headers={
                **self._auth_headers(),
                "Content-Type": content_type,
                "Content-Length": str(len(body)),
            },
            body=body,
        )

        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload["status"], "uploaded")
        self.assertEqual(payload["country_code"], "CL")
        self.assertEqual(
            len(self.fake_opensearch.document_ids_for("CL")), 1
        )


class ExistingCountryAsgiTests(AdminAsgiTestCase):
    def test_asgi_existing_country_returns_structured_409(self) -> None:
        self.fake_opensearch.add(
            document_id="doc_" + "a" * 64,
            country_code="CL",
            source_filename="Chile.docx",
        )
        (self.source_dir / "CL.docx").write_bytes(b"existing-cl-bytes")

        before_ids = set(
            self.fake_opensearch.document_ids_for("CL")
        )

        body, content_type = build_multipart_body(
            file_field="file",
            filename="Chile.docx",
            file_content=_CL_DOCX_BYTES,
            extra_fields={
                "replace_existing": "false",
                "confirm_warnings": "true",
                "country_confirmed": "true",
            },
        )

        response = asgi_request(
            self.app,
            "POST",
            "/api/v1/admin/documents",
            headers={
                **self._auth_headers(),
                "Content-Type": content_type,
                "Content-Length": str(len(body)),
            },
            body=body,
        )

        self.assertEqual(response.status_code, 409)
        payload = response.json()
        self.assertEqual(
            payload["detail"]["code"],
            "document_replacement_required",
        )
        self.assertEqual(payload["detail"]["country_code"], "CL")

        # Zero mutation.
        self.assertEqual(
            set(self.fake_opensearch.document_ids_for("CL")),
            before_ids,
        )
        self.assertEqual(
            (self.source_dir / "CL.docx").read_bytes(),
            b"existing-cl-bytes",
        )


class FastApiValidationAsgiTests(AdminAsgiTestCase):
    def test_asgi_missing_file_returns_fastapi_detail_list(
        self,
    ) -> None:
        # No manually-constructed HTTPException(422) anywhere in this
        # path - this is FastAPI's own automatic RequestValidationError
        # for a required multipart field that was never sent, the
        # exact mechanism behind the WordPress fallback-message defect
        # this mission's earlier phase fixed.
        body, content_type = build_multipart_body(
            file_field="file",
            filename="unused.docx",
            file_content=b"unused",
            extra_fields={
                "replace_existing": "false",
                "confirm_warnings": "true",
            },
            include_file=False,
        )

        response = asgi_request(
            self.app,
            "POST",
            "/api/v1/admin/documents",
            headers={
                **self._auth_headers(),
                "Content-Type": content_type,
                "Content-Length": str(len(body)),
            },
            body=body,
        )

        self.assertEqual(response.status_code, 422)
        payload = response.json()
        self.assertIsInstance(payload["detail"], list)
        self.assertGreater(len(payload["detail"]), 0)
        self.assertIn("msg", payload["detail"][0])
        self.assertTrue(
            isinstance(payload["detail"][0]["msg"], str)
            and payload["detail"][0]["msg"]
        )


class AuthAsgiTests(AdminAsgiTestCase):
    def test_asgi_missing_api_key_returns_401(self) -> None:
        response = asgi_request(
            self.app,
            "GET",
            "/api/v1/admin/documents",
            headers={},
        )

        self.assertEqual(response.status_code, 401)

    def test_asgi_wrong_api_key_returns_401(self) -> None:
        response = asgi_request(
            self.app,
            "GET",
            "/api/v1/admin/documents",
            headers={"X-API-Key": "wrong-key"},
        )

        self.assertEqual(response.status_code, 401)

    def test_asgi_correct_api_key_but_missing_admin_key_returns_401(
        self,
    ) -> None:
        response = asgi_request(
            self.app,
            "GET",
            "/api/v1/admin/documents",
            headers={"X-API-Key": TEST_API_ACCESS_KEY},
        )

        self.assertEqual(response.status_code, 401)

    def test_asgi_correct_api_key_but_wrong_admin_key_returns_401(
        self,
    ) -> None:
        response = asgi_request(
            self.app,
            "GET",
            "/api/v1/admin/documents",
            headers={
                "X-API-Key": TEST_API_ACCESS_KEY,
                "X-Admin-Key": "wrong-admin-key",
            },
        )

        self.assertEqual(response.status_code, 401)

    def test_asgi_both_keys_correct_reaches_the_route(self) -> None:
        response = asgi_request(
            self.app,
            "GET",
            "/api/v1/admin/documents",
            headers=self._auth_headers(),
        )

        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload["total"], 0)


if __name__ == "__main__":
    unittest.main()
