"""
Tests for the DOCX mechanics behind Admin-managed contacts: the
canonical, Word-native in-flow contact table (the design that replaced
both an earlier hand-authored floating-textbox design - rejected by real
Microsoft Word - and the original two-box-cloning design before it), the
older floating-shape photo mutation primitives still used to detect
whether a document's Contact area is in that OLD format, raw contact-
photo extraction from a DOCX's own image relationships, and
contact/photo association by geometry or position.

Mirrors the established real-corpus convention throughout this suite:
exercise the REAL corpus documents (temp copies, never the files
themselves - see tests.support.corpus_paths/tests.support.documents),
skipping gracefully when the corpus is unavailable, and proving
structural correctness (XML validity, no dangling relationships, no new
floating shapes) rather than merely checking that strings exist
somewhere in document.xml.
"""

from __future__ import annotations

import hashlib
import io
import re
import tempfile
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from unittest.mock import patch

from docx import Document
from docx import Document as WordDocument
from lxml import etree

from app.services import contact_document_area as _contact_document_area
from app.services.contact_document_area import (
    ContactAreaError,
    ContactPhotoPayload,
    rebuild_canonical_contact_table,
    resolve_untracked_contact_photo,
)
from app.services.contact_document_photos import (
    ContactDocumentPhotoError,
    add_contact_photo_to_document,
    add_new_contact_photo_to_document,
    remove_contact_photo_from_document,
    replace_contact_photo_in_document,
)
from app.services.contact_people import associate_contact_photos
from app.services.contact_photos import (
    ContactPhotoCandidate,
    extract_contact_photo_candidates,
)
from app.services.docx_parser import (
    CONTACT_TABLE_HIDDEN_MARKER,
    ExtractedContact,
    extract_contacts_from_docx,
    split_combined_legacy_contact,
)
from app.services.document_contact_materializer import (
    _find_all_contact_runs,
    _is_contact_related_block,
)

from tests.support.corpus_paths import resolve_source_root
from tests.support.documents import (
    make_png,
    require_corpus_copy,
    skip_if_already_canonicalized,
)

SOURCE_ROOT = resolve_source_root()


def _sha(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()


_VALID_PNG = make_png(183, 234, (200, 50, 50))


def _make_webp_vp8l(width: int, height: int) -> bytes:
    """A minimal, real, valid lossless (VP8L) WebP image - the format
    Pillow's default Image.save(format="WEBP") produces for a synthetic
    in-memory image, and the one this environment's own _WebpImageHeader
    parser (contact_document_area.py) must recognize for python-docx to
    embed it at all (python-docx 1.2.0 ships no WebP support whatsoever -
    confirmed directly)."""

    import struct

    packed = ((height - 1) << 14) | (width - 1)
    payload = (
        bytes([0x2F]) + packed.to_bytes(4, "little") + bytes([0x88, 0x88, 0x08])
    )
    chunk = b"VP8L" + struct.pack("<I", len(payload)) + payload
    if len(chunk) % 2:
        chunk += b"\x00"
    riff_payload = b"WEBP" + chunk
    return b"RIFF" + struct.pack("<I", len(riff_payload)) + riff_payload


def _make_webp_vp8(width: int, height: int) -> bytes:
    """The other real-world WebP sub-format (simple lossy VP8) - a
    different byte layout than VP8L, commonly produced by a browser or
    screenshot tool."""

    import struct

    frame_tag = b"\x00\x00\x00"
    start_code = b"\x9d\x01\x2a"
    width_bytes = (width & 0x3FFF).to_bytes(2, "little")
    height_bytes = (height & 0x3FFF).to_bytes(2, "little")
    payload = frame_tag + start_code + width_bytes + height_bytes + b"\x00" * 4
    chunk = b"VP8 " + struct.pack("<I", len(payload)) + payload
    if len(chunk) % 2:
        chunk += b"\x00"
    riff_payload = b"WEBP" + chunk
    return b"RIFF" + struct.pack("<I", len(riff_payload)) + riff_payload


_VALID_WEBP_LOSSLESS = _make_webp_vp8l(120, 160)
_VALID_WEBP_LOSSY = _make_webp_vp8(120, 160)

_WP_NS = "{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}"
_W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_PIC_NS = "{http://schemas.openxmlformats.org/drawingml/2006/picture}"


# =========================================================================
# CANONICAL CONTACT TABLE (contact_document_area.py)
# =========================================================================


class ContactDocumentAreaTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)

    def _require_copy(self, filename: str) -> Path:
        return require_corpus_copy(
            self, SOURCE_ROOT, filename, Path(self.temp.name)
        )

    def _structural_checks(self, path: Path) -> None:
        """Zip integrity, XML well-formedness, no dangling
        relationships, no [trash]/ parts, and no NEW floating contact
        shapes (wp:anchor) - the canonical table's own content is
        exclusively ordinary inline pictures/paragraphs."""

        with zipfile.ZipFile(path) as archive:
            self.assertIsNone(archive.testzip())
            names = set(archive.namelist())
            self.assertFalse(
                [n for n in names if n.startswith("[trash]")],
                "no [trash]/ parts may be left behind",
            )
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")
            rels_xml = archive.read(
                "word/_rels/document.xml.rels"
            ).decode("utf-8")
            content_types_xml = archive.read(
                "[Content_Types].xml"
            ).decode("utf-8")

        for label, content in (
            ("document.xml", document_xml),
            ("rels", rels_xml),
            ("content types", content_types_xml),
        ):
            try:
                ET.fromstring(content)
            except ET.ParseError as error:
                self.fail(f"{label} is not well-formed XML: {error}")

        referenced_ids = set(
            re.findall(r'r:(?:id|embed)="([^"]+)"', document_xml)
        )
        declared_ids = set(re.findall(r'Id="([^"]+)"', rels_xml))
        self.assertEqual(
            set(), referenced_ids - declared_ids,
            "every r:id/r:embed referenced in document.xml must be "
            "declared in the relationships part",
        )

        # Re-opening via python-docx is itself a strong well-formedness
        # signal (the same parser Word's own OOXML reader family is
        # closely related to, unlike a bare lxml.fromstring check).
        WordDocument(path)

    def _decorative_top_and_bottom_wrapped_shapes(
        self, document_xml: str
    ) -> int:
        """Count floating shapes wrapped topAndBottom that carry
        neither real text nor a picture - contact_document_area.py's
        own removal target for the legacy reserved-space rectangle
        ("the tallest topAndBottom-wrapped shape with neither text nor
        a picture", per its module docstring). Checking this semantic
        invariant directly - rather than hardcoding an expected shape
        count - stays correct whether or not a given source document
        still happens to carry a decorative rectangle: real organic
        content drifts over time (a page's own title text box, for
        example, is also wrapped topAndBottom, has real text, and must
        never be mistaken for a removal candidate)."""

        root = ET.fromstring(document_xml)
        decorative = 0

        for anchor in root.iter(f"{_WP_NS}anchor"):
            if anchor.find(f"{_WP_NS}wrapTopAndBottom") is None:
                continue

            has_text = any(
                (node.text or "").strip()
                for node in anchor.iter(f"{_W_NS}t")
            )
            has_picture = anchor.find(f".//{_PIC_NS}pic") is not None

            if not has_text and not has_picture:
                decorative += 1

        return decorative

    def _table_xml(self, document_xml: str) -> str:
        marker_index = document_xml.find(CONTACT_TABLE_HIDDEN_MARKER)
        self.assertNotEqual(
            -1, marker_index,
            "the canonical table's own hidden marker must be present",
        )
        table_start = document_xml.rfind("<w:tbl>", 0, marker_index)
        table_end = document_xml.find("</w:tbl>", marker_index)
        self.assertNotEqual(-1, table_start)
        self.assertNotEqual(-1, table_end)
        return document_xml[table_start:table_end + len("</w:tbl>")]

    def test_add_first_contact_with_photo_produces_coherent_block(
        self,
    ) -> None:
        """Canonicalizing AU's own organic contact area leaves no large
        empty band (Introduction follows the new table directly) and
        the sole contact keeps its own photo.

        The "no empty band" check asserts the semantic contract
        directly (no decorative, text-less/picture-less
        topAndBottom-wrapped shape survives, and no NEW one is
        introduced) rather than a hardcoded shape count: AU's own real
        content has drifted since this test was written and its
        current baseline no longer carries a separate removable
        reserved-space rectangle at all (verified directly against the
        real corpus - AU's only topAndBottom-wrapped shape today is
        its own title text box, which legitimately survives
        unchanged)."""

        path = self._require_copy("AU.docx")

        with zipfile.ZipFile(path) as archive:
            original_document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")
        original_wrap_count = original_document_xml.count(
            "wrapTopAndBottom"
        )

        original_contacts = extract_contacts_from_docx(
            path, country="Australia"
        )
        self.assertEqual(1, len(original_contacts))
        michael = original_contacts[0]

        michael_photo = resolve_untracked_contact_photo(
            path, contact_person=michael.contact_person, country="Australia"
        )
        self.assertIsNotNone(michael_photo)

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(michael,),
            photos=(michael_photo,),
            country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")

        # The canonical TABLE ITSELF contains no floating shapes at
        # all - only ordinary inline content. Page furniture unrelated
        # to the contact area (URL box, title box, logo) legitimately
        # keeps its own pre-existing floating shapes elsewhere in the
        # document, so this checks the table's own XML specifically,
        # never a whole-document count.
        table_xml = self._table_xml(document_xml)
        self.assertEqual(
            0, table_xml.count("<wp:anchor"),
            "the canonical table must contain no floating shapes at "
            "all - only ordinary inline content",
        )

        # No large reserved band: no decorative, text-less/picture-less
        # topAndBottom-wrapped shape (the legacy reserved-space
        # rectangle) survives canonicalization - whether or not this
        # particular source document still happened to carry one.
        # Real content wrapped the same way (a title text box) is
        # never a removal candidate and must be untouched.
        self.assertEqual(
            0,
            self._decorative_top_and_bottom_wrapped_shapes(document_xml),
            "no decorative, text-less, picture-less topAndBottom "
            "rectangle may survive canonicalization",
        )
        # Each real shape contributes exactly one "wrapTopAndBottom"
        # substring match (its own DrawingML Choice branch only - the
        # VML Fallback branch spells it differently,
        # type="topAndBottom"), so canonicalization must never
        # introduce a NEW one.
        self.assertLessEqual(
            document_xml.count("wrapTopAndBottom"),
            original_wrap_count,
            "canonicalization must never introduce a new "
            "topAndBottom-wrapped shape",
        )

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(1, len(reparsed))
        self.assertEqual("Michael Harmer", reparsed[0].contact_person)

        photos = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(photos))
        self.assertEqual(michael_photo.data, photos[0].data)

    def test_two_contacts_two_distinct_photos(self) -> None:
        """Two contacts, each with its own distinct photo, round-trip
        as two separate contacts with two separate photos - never
        merged or cross-associated."""

        path = self._require_copy("AU.docx")

        michael = extract_contacts_from_docx(path, country="Australia")[0]
        michael_photo = resolve_untracked_contact_photo(
            path, contact_person=michael.contact_person, country="Australia"
        )

        jane = ExtractedContact(
            member_firm="Second Firm Pty Ltd",
            contact_person="Jane Secondary",
            email="jane.secondary@secondfirm.com.au",
            phone="+61 2 9000 0000",
            address="100 Test Street, Level 5",
            website="www.secondfirm.com.au",
        )
        jane_photo = ContactPhotoPayload(
            data=_VALID_PNG, content_type="image/png"
        )

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(michael, jane),
            photos=(michael_photo, jane_photo),
            country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(2, len(reparsed))
        self.assertEqual("Michael Harmer", reparsed[0].contact_person)
        self.assertEqual("Jane Secondary", reparsed[1].contact_person)

        photos = extract_contact_photo_candidates(path)
        shas = {p.sha256 for p in photos}
        self.assertEqual(2, len(shas))
        self.assertIn(michael_photo and _sha(michael_photo.data), shas)
        self.assertIn(_sha(jane_photo.data), shas)

    def test_three_contacts_all_preserved_in_order(self) -> None:
        """Three contacts round-trip in the same order, with the same
        field values."""

        path = self._require_copy("AU.docx")
        michael = extract_contacts_from_docx(path, country="Australia")[0]

        jane = ExtractedContact(
            member_firm="Second Firm Pty Ltd",
            contact_person="Jane Secondary",
            email="jane.secondary@secondfirm.com.au",
        )
        priya = ExtractedContact(
            member_firm="Third Firm Pty Ltd",
            contact_person="Priya Third",
            email="priya.third@thirdfirm.com.au",
        )

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(michael, jane, priya),
            photos=(None, None, None),
            country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(3, len(reparsed))
        self.assertEqual(
            ["Michael Harmer", "Jane Secondary", "Priya Third"],
            [c.contact_person for c in reparsed],
        )

    def test_contact_without_photo_produces_coherent_block(self) -> None:
        """A contact with no photo at all still produces a coherent
        block - an empty right cell, never a missing row or misaligned
        text."""

        path = self._require_copy("AU.docx")
        michael = extract_contacts_from_docx(path, country="Australia")[0]

        jane = ExtractedContact(
            member_firm="Second Firm Pty Ltd",
            contact_person="Jane Secondary",
            email="jane.secondary@secondfirm.com.au",
        )

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(michael, jane),
            photos=(None, None),
            country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(2, len(reparsed))

        photos = extract_contact_photo_candidates(path)
        self.assertEqual(0, len(photos))

    def test_delete_contact_b_leaves_a_intact(self) -> None:
        """Rebuilding from a contact list with B removed leaves B's
        text/photo gone and A's text/photo untouched."""

        path = self._require_copy("AU.docx")
        michael = extract_contacts_from_docx(path, country="Australia")[0]
        michael_photo = resolve_untracked_contact_photo(
            path, contact_person=michael.contact_person, country="Australia"
        )

        jane = ExtractedContact(
            member_firm="Second Firm Pty Ltd",
            contact_person="Jane Secondary",
            email="jane.secondary@secondfirm.com.au",
        )
        jane_photo = ContactPhotoPayload(
            data=_VALID_PNG, content_type="image/png"
        )

        with_both = rebuild_canonical_contact_table(
            path,
            contacts=(michael, jane),
            photos=(michael_photo, jane_photo),
            country="Australia",
        )
        path.write_bytes(with_both)

        # Now rebuild again from ONLY Michael - simulating Jane's
        # deletion the same way admin_contacts.py's unified
        # synchronization does (rebuild from the full surviving list).
        after_delete = rebuild_canonical_contact_table(
            path,
            contacts=(michael,),
            photos=(michael_photo,),
            country="Australia",
        )
        path.write_bytes(after_delete)
        self._structural_checks(path)

        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")

        self.assertNotIn("Jane Secondary", document_xml)
        self.assertNotIn("jane.secondary@secondfirm.com.au", document_xml)
        self.assertIn("Michael Harmer", document_xml)

        remaining_photos = {
            p.sha256 for p in extract_contact_photo_candidates(path)
        }
        self.assertNotIn(_sha(jane_photo.data), remaining_photos)
        self.assertIn(_sha(michael_photo.data), remaining_photos)

    def test_deleting_every_contact_leaves_no_table_and_no_band(
        self,
    ) -> None:
        """Rebuilding with zero contacts removes the canonical area
        entirely - no marker-only leftover table, no empty band."""

        path = self._require_copy("AU.docx")
        michael = extract_contacts_from_docx(path, country="Australia")[0]

        with_contact = rebuild_canonical_contact_table(
            path, contacts=(michael,), photos=(None,), country="Australia",
        )
        path.write_bytes(with_contact)

        emptied = rebuild_canonical_contact_table(
            path, contacts=(), photos=(), country="Australia",
        )
        path.write_bytes(emptied)
        self._structural_checks(path)

        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")

        self.assertNotIn(CONTACT_TABLE_HIDDEN_MARKER, document_xml)
        self.assertNotIn("<w:tbl>", document_xml)
        self.assertEqual(0, len(extract_contacts_from_docx(path)))

    def test_legal_content_unchanged(self) -> None:
        """Real legal section text (well past the contact area) is
        byte-identical after the rebuild."""

        path = self._require_copy("AU.docx")

        with zipfile.ZipFile(path) as archive:
            before = archive.read("word/document.xml").decode("utf-8")

        introduction_index = before.find("Introduction")
        self.assertNotEqual(-1, introduction_index)
        legal_tail_before = before[introduction_index:]

        michael = extract_contacts_from_docx(path, country="Australia")[0]
        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(michael,), photos=(None,), country="Australia",
        )
        path.write_bytes(new_bytes)

        with zipfile.ZipFile(path) as archive:
            after = archive.read("word/document.xml").decode("utf-8")

        introduction_index_after = after.find("Introduction")
        self.assertNotEqual(-1, introduction_index_after)
        legal_tail_after = after[introduction_index_after:]

        self.assertEqual(
            legal_tail_before, legal_tail_after,
            "legal content from Introduction onward must be "
            "byte-for-byte unchanged",
        )

    def test_no_matching_contact_area_still_produces_canonical_table(
        self,
    ) -> None:
        """A document with no legacy floating contact area at all
        (Portugal) still gets a real canonical table - never left
        ContactState-only."""

        path = self._require_copy("PT.docx")

        new_contact = ExtractedContact(
            member_firm="Someone New Lda",
            contact_person="Someone New",
            email="new@example.test",
            phone="+351 21 000 0000",
            address="Rua Nova 1",
            website="www.newfirm.test",
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(new_contact,), photos=(None,), country="Portugal",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Portugal")
        self.assertEqual(1, len(reparsed))
        self.assertEqual("Someone New", reparsed[0].contact_person)

    def test_second_rebuild_replaces_rather_than_duplicates_table(
        self,
    ) -> None:
        """Rebuilding twice in a row replaces the one canonical table
        rather than stacking a second one alongside it."""

        path = self._require_copy("PT.docx")

        first = ExtractedContact(
            member_firm="First Firm Lda",
            contact_person="First Person",
            email="first@example.test",
        )
        second_contact = ExtractedContact(
            member_firm="Second Firm Lda",
            contact_person="Second Person",
            email="second@example.test",
        )

        path.write_bytes(
            rebuild_canonical_contact_table(
                path, contacts=(first,), photos=(None,), country="Portugal",
            )
        )
        path.write_bytes(
            rebuild_canonical_contact_table(
                path,
                contacts=(first, second_contact),
                photos=(None, None),
                country="Portugal",
            )
        )
        self._structural_checks(path)

        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")

        self.assertEqual(1, document_xml.count("<w:tbl>"))
        self.assertEqual(
            1, document_xml.count(CONTACT_TABLE_HIDDEN_MARKER)
        )

        reparsed = extract_contacts_from_docx(path, country="Portugal")
        self.assertEqual(2, len(reparsed))

    # =====================================================================
    # Real Argentina/Canada incidents predate this module's current
    # single-rebuild design (see its own module docstring: two earlier,
    # abandoned approaches). Real Canada (still un-canonicalized, verified
    # directly) proves the CURRENT mechanism already removes a real legacy
    # floating-shape carrier correctly. The invariant below is the
    # general, non-country-specific backstop: even if some OTHER
    # country's own legacy shape geometry ever defeats
    # _remove_legacy_carrier_and_get_anchor's detection, the mutation
    # fails loudly instead of silently persisting two areas.
    # =====================================================================

    def test_canada_legacy_floating_carrier_add_produces_single_zone(
        self,
    ) -> None:
        """Real Canada: Robert Bayne's own original contact area is a
        legacy floating-shape carrier, still un-canonicalized. Adding a
        second contact must leave EXACTLY one canonical table, with the
        legacy carrier entirely gone - never both side by side."""

        path = self._require_copy("CA.docx")

        baseline = extract_contacts_from_docx(path, country="Canada")
        self.assertEqual(1, len(baseline))
        self.assertEqual("Robert Bayne", baseline[0].contact_person)

        new_contact = ExtractedContact(
            member_firm="CRUD Regression LLP",
            contact_person="Regression Contact",
            email="contact@crud-regression.example",
            phone="+33 1 23 45 67 89",
            address="12 Test Street, Floor 3, 75008 Paris",
            website="www.crud-regression.example",
        )

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(*baseline, new_contact),
            photos=(None, None),
            country="Canada",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")

        self.assertEqual(
            1, document_xml.count(CONTACT_TABLE_HIDDEN_MARKER)
        )
        self.assertEqual(1, document_xml.count("<w:tbl>"))

        reparsed = extract_contacts_from_docx(path, country="Canada")
        self.assertEqual(2, len(reparsed))
        self.assertEqual("Robert Bayne", reparsed[0].contact_person)
        self.assertEqual(
            "Regression Contact", reparsed[1].contact_person
        )

    def test_single_zone_invariant_rejects_undetected_legacy_carrier(
        self,
    ) -> None:
        """Adversarial: simulate a legacy floating-shape carrier that
        detection failed to remove (a canonical table inserted
        alongside it, exactly as if _remove_legacy_carrier_and_get_
        anchor had missed it) - the invariant must refuse to persist
        this, rather than silently producing the historical Argentina/
        Canada two-block defect for whatever country's shape geometry
        someday defeats detection."""

        path = self._require_copy("CA.docx")
        source_bytes = path.read_bytes()

        legacy_photo_relationship_ids = {
            candidate.relationship_id
            for candidate in extract_contact_photo_candidates(path)
            if candidate.relationship_id
        }

        document = WordDocument(path)

        new_contact = ExtractedContact(
            member_firm="Undetected Carrier Regression LLP",
            contact_person="Undetected Carrier Contact",
            email="undetected@example.test",
        )

        # Deliberately skip carrier removal - this is the exact defect
        # shape the invariant must catch: a fresh canonical table
        # inserted while the real legacy carrier is still there.
        table = _contact_document_area._build_canonical_table(
            document, contacts=(new_contact,), photos=(None,)
        )
        document.element.body.insert(0, table._tbl)

        output = io.BytesIO()
        document.save(output)

        with self.assertRaises(ContactAreaError) as raised:
            _contact_document_area._assert_single_contact_zone(
                output.getvalue(),
                expected_table_count=1,
                legacy_photo_relationship_ids=(
                    legacy_photo_relationship_ids
                ),
            )

        self.assertIn("floating-shape", str(raised.exception))
        self.assertEqual(
            source_bytes,
            path.read_bytes(),
            "the real corpus file must never be touched by this test",
        )

    def test_single_zone_invariant_rejects_duplicated_canonical_table(
        self,
    ) -> None:
        """Adversarial: two canonical tables in the same document (a
        table-count detection miss) must also be refused, not just a
        leftover legacy shape."""

        path = self._require_copy("AU.docx")

        document = WordDocument(path)

        extra_contact = ExtractedContact(
            member_firm="Duplicate Table Regression LLP",
            contact_person="Duplicate Table Contact",
            email="duplicate@example.test",
        )

        extra_table = _contact_document_area._build_canonical_table(
            document, contacts=(extra_contact,), photos=(None,)
        )
        document.element.body.insert(0, extra_table._tbl)

        output = io.BytesIO()
        document.save(output)

        with self.assertRaises(ContactAreaError) as raised:
            _contact_document_area._assert_single_contact_zone(
                output.getvalue(),
                expected_table_count=1,
                legacy_photo_relationship_ids=set(),
            )

        self.assertIn("Expected exactly 1", str(raised.exception))

    def test_photo_with_crop_rectangle_is_refused_not_guessed(
        self,
    ) -> None:
        """resolve_untracked_contact_photo refuses (rather than
        blindly embeds) a legacy photo whose own source shape carries
        an OOXML crop rectangle this environment cannot reproduce."""

        path = self._require_copy("AU.docx")

        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")

        # Inject a synthetic crop rectangle into Michael's own photo
        # shape to simulate a cropped legacy photo - this baseline
        # itself has none (verified separately), so this test
        # fabricates the one condition it needs to exercise.
        patched_xml = document_xml.replace(
            '<a:blip r:embed="rId10"/>',
            '<a:blip r:embed="rId10"/><a:srcRect l="1000" t="1000" '
            'r="1000" b="1000"/>',
            1,
        )
        self.assertNotEqual(
            document_xml, patched_xml,
            "sanity: the replacement must actually match something",
        )

        from app.services.contact_document_photos import _rewrite_zip

        patched_bytes = _rewrite_zip(
            path.read_bytes(),
            replacements={"word/document.xml": patched_xml},
        )
        path.write_bytes(patched_bytes)

        with self.assertRaises(ContactAreaError):
            resolve_untracked_contact_photo(
                path, contact_person="Michael Harmer", country="Australia"
            )

    def test_bare_srcrect_with_no_offsets_is_not_treated_as_a_crop(
        self,
    ) -> None:
        """A bare <a:srcRect/> (no l/t/r/b attributes at all - OOXML
        defaults each to 0) declares no visible crop whatsoever and
        must NOT be refused - only a genuinely non-zero offset should
        be. Real Indonesia corpus content carries exactly this shape on
        its own contact portrait (see
        test_indonesia_untracked_photo_with_noop_crop_resolves for the
        real-corpus regression); this test isolates the same defect
        with a synthetic AU fixture so it does not depend on Indonesia's
        own content remaining unchanged."""

        path = self._require_copy("AU.docx")

        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8")

        patched_xml = document_xml.replace(
            '<a:blip r:embed="rId10"/>',
            '<a:blip r:embed="rId10"/><a:srcRect/>',
            1,
        )
        self.assertNotEqual(
            document_xml, patched_xml,
            "sanity: the replacement must actually match something",
        )

        from app.services.contact_document_photos import _rewrite_zip

        patched_bytes = _rewrite_zip(
            path.read_bytes(),
            replacements={"word/document.xml": patched_xml},
        )
        path.write_bytes(patched_bytes)

        photo = resolve_untracked_contact_photo(
            path, contact_person="Michael Harmer", country="Australia"
        )
        self.assertIsNotNone(photo)

    def test_indonesia_untracked_photo_with_noop_crop_resolves(
        self,
    ) -> None:
        """Real-corpus regression: Indonesia's own contact portrait
        (Marshall Situmorang) carries a bare, all-zero <a:srcRect/> -
        a Word-emitted no-op, not an actual crop. Before the offset-
        aware fix, resolve_untracked_contact_photo refused it outright,
        which would block ANY first Contact CRUD mutation on this
        document (Add/Edit/Delete all need to preserve an untracked
        contact's existing photo via this same path)."""

        path = self._require_copy("ID.docx")

        photo = resolve_untracked_contact_photo(
            path, contact_person="Marshall Situmorang", country="Indonesia"
        )
        self.assertIsNotNone(photo)

    def test_each_of_the_six_fields_individually_empty_round_trips(
        self,
    ) -> None:
        """All six fields are individually optional. address/website-
        both-empty (France's own real shape) and member_firm-empty
        (IE/IN's own real shape) are already covered elsewhere; this
        closes the remaining combinations - contact_person and email
        individually empty - at the actual DOCX round-trip layer, not
        just the AdminContactWriteRequest validation layer (see
        AdminContactWriteRequestOptionalFieldTests in
        test_admin_contacts.py for that one)."""

        path = self._require_copy("AU.docx")

        full_fields = dict(
            member_firm="Complete Fields LLP",
            address="1 Complete Street",
            phone="+1 555 000 1111",
            website="www.complete-fields.example",
            contact_person="Complete Person",
            email="complete@example.test",
        )

        for empty_field in full_fields:
            with self.subTest(empty_field=empty_field):
                fields = dict(full_fields)
                fields[empty_field] = None
                contact = ExtractedContact(**fields)

                new_bytes = rebuild_canonical_contact_table(
                    path, contacts=(contact,), photos=(None,), country="Australia",
                )
                # rebuild_canonical_contact_table already re-parses and
                # raises ContactAreaError on any mismatch internally -
                # reaching this point without an exception IS the
                # positive assertion. Persist and re-verify explicitly
                # too, for a clear failure message per field.
                path.write_bytes(new_bytes)
                reparsed = extract_contacts_from_docx(path, country="Australia")
                self.assertEqual(1, len(reparsed))

                for field_name, expected_value in fields.items():
                    self.assertEqual(
                        expected_value,
                        getattr(reparsed[0], field_name),
                        f"field {field_name!r} did not round-trip when "
                        f"{empty_field!r} was empty",
                    )

    def test_multiple_contacts_with_different_missing_fields_and_a_duplicate(
        self,
    ) -> None:
        """Several contacts with DIFFERENT missing-field shapes,
        content deliberately chosen to fool a content-based classifier
        (a postal code, a phone-like substring embedded in an address,
        an apartment/unit number, a scheme-less URL, accented/unicode
        names and an apostrophe), and an exact duplicate pair - all in
        the SAME canonical table, at once. No field may shift into
        another contact's slot, and duplicates must survive unmerged."""

        path = self._require_copy("AU.docx")

        firm_only = ExtractedContact(member_firm="Firm Only LLP")
        person_email_only = ExtractedContact(
            contact_person="Person Only", email="person-only@example.test"
        )
        fooling_shapes = ExtractedContact(
            member_firm="Müller & O'Connell Associés",
            address=(
                "Apt #4B, 12 Rue de l'Église, 75008 Paris, "
                "+33 1 23 45 67 89"
            ),
            phone="75008",
            website="fooling-shapes.example",
            contact_person="François O'Connell-Müller",
            email="françois@fooling-shapes.example",
        )
        duplicate_a = ExtractedContact(
            member_firm="Duplicate LLP",
            contact_person="Duplicate Person",
            email="duplicate@example.test",
        )
        duplicate_b = ExtractedContact(
            member_firm="Duplicate LLP",
            contact_person="Duplicate Person",
            email="duplicate@example.test",
        )

        contacts = (
            firm_only,
            person_email_only,
            fooling_shapes,
            duplicate_a,
            duplicate_b,
        )
        photos = (None,) * len(contacts)

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=contacts, photos=photos, country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(list(contacts), reparsed)

    def test_accented_email_local_part_never_corrupted(self) -> None:
        """Real bug found while building the item C/D/E combinatorial
        test above: _EMAIL_PATTERN's local-part character class is
        ASCII-only, so an email with an accented local part (a
        genuinely legitimate address, e.g. a French admin's own name)
        was silently truncated on read-back, and the truncated
        fragment then failed to even look like an email, spilling into
        contact_person instead. contact_person/email are now tagged
        the same deterministic way as the left cell's four fields, so
        this no longer depends on _EMAIL_PATTERN at all for a
        canonical table this code itself wrote."""

        path = self._require_copy("AU.docx")

        contact = ExtractedContact(
            contact_person="François Müller",
            email="françois.müller@example.test",
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(contact,), photos=(None,), country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(1, len(reparsed))
        self.assertEqual("François Müller", reparsed[0].contact_person)
        self.assertEqual(
            "françois.müller@example.test", reparsed[0].email
        )

    def test_webp_contact_photo_can_be_embedded(self) -> None:
        """Real bug, reproduced directly against pre-fix code: python-
        docx 1.2.0 has no WebP header parser at all, so ANY Add+Photo/
        Edit+Photo/Replace-Photo mutation with a WebP file (a format
        this codebase already accepts and advertises - admin_contact_
        photos.py's own "Only JPEG, PNG and WebP images are accepted")
        failed with a raw UnrecognizedImageError wrapped as
        ContactAreaError. contact_document_area.py registers a small,
        dependency-free WebP header parser into python-docx's own
        image factory; this proves both real-world WebP sub-formats -
        lossless VP8L (what Pillow produces by default) and lossy
        simple VP8 (a different byte layout, common from a browser or
        screenshot tool) - now embed and round-trip correctly, with no
        image library added."""

        path = self._require_copy("AU.docx")
        baseline = extract_contacts_from_docx(path, country="Australia")

        webp_variants = {
            "lossless (VP8L)": _VALID_WEBP_LOSSLESS,
            "lossy (VP8)": _VALID_WEBP_LOSSY,
        }

        for variant_name, webp_bytes in webp_variants.items():
            with self.subTest(variant=variant_name):
                new_contact = ExtractedContact(
                    member_firm="WebP Test Firm",
                    contact_person=f"WebP Test Person ({variant_name})",
                    email="webp@example.test",
                )
                photo = ContactPhotoPayload(
                    data=webp_bytes, content_type="image/webp"
                )

                new_bytes = rebuild_canonical_contact_table(
                    path,
                    contacts=(*baseline, new_contact),
                    photos=(*([None] * len(baseline)), photo),
                    country="Australia",
                )
                path.write_bytes(new_bytes)
                self._structural_checks(path)

                candidates = extract_contact_photo_candidates(path)
                actual_shas = {c.sha256 for c in candidates}
                self.assertIn(_sha(webp_bytes), actual_shas)

                reparsed = extract_contacts_from_docx(
                    path, country="Australia"
                )
                self.assertEqual(len(baseline) + 1, len(reparsed))
                self.assertEqual(
                    f"WebP Test Person ({variant_name})",
                    reparsed[-1].contact_person,
                )

    def test_multiple_emails_on_one_line_round_trip_preserved(self) -> None:
        """Regression for the real France defect found by the
        full-corpus mutation test: FR's actual contact carries two
        email addresses on a single comma-joined line
        ("scherrmann@flichy.com, bacquet@flichy.com"), because
        ExtractedContact/ContactState model email as one string field
        and the writer renders it verbatim as one line. The canonical
        table reader used to recover only the first address via a
        single regex .search(); the round-trip validator correctly
        caught the loss and refused the rebuild with a
        ContactAreaError, but the fix (using .findall() to recover
        every address on the line) must let the exact real France
        contact round-trip cleanly with both emails intact, in
        order."""

        path = self._require_copy("FR.docx")

        flichy = ExtractedContact(
            member_firm="Flichy Grangé Avocats",
            contact_person="Caroline Scherrmann and Florence Bacquet",
            email="scherrmann@flichy.com, bacquet@flichy.com",
            phone="+33 1 56 62 30 00",
        )

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(flichy,),
            photos=(None,),
            country="France",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="France")
        self.assertEqual(1, len(reparsed))
        self.assertEqual(
            "scherrmann@flichy.com, bacquet@flichy.com",
            reparsed[0].email,
        )

    def test_empty_firm_cell_phone_addition_round_trips(self) -> None:
        """Real cases: IE's real contact (Aoife Bradley) and IN's real
        contact (Avik Biswas) both have member_firm/address/phone/
        website all empty. Giving each a phone value must land it as
        phone - never misread as member_firm (which the writer would
        otherwise never have written in the first place, since it was
        empty to begin with). Confirmed against both real documents,
        since the fix is architectural (never specific to either
        country)."""

        cases = {
            "Ireland": ("IE.docx", "Ireland", "+353 1 234 5678"),
            "India": ("IN.docx", "India", "+91 11 4567 8900"),
        }

        for label, (filename, country, new_phone) in cases.items():
            with self.subTest(country=label):
                path = self._require_copy(filename)
                baseline = extract_contacts_from_docx(path, country=country)
                self.assertEqual(1, len(baseline))
                self.assertIsNone(baseline[0].member_firm)
                self.assertIsNone(baseline[0].phone)

                with_phone = ExtractedContact(
                    contact_person=baseline[0].contact_person,
                    email=baseline[0].email,
                    phone=new_phone,
                )

                new_bytes = rebuild_canonical_contact_table(
                    path,
                    contacts=(with_phone,),
                    photos=(None,),
                    country=country,
                )
                path.write_bytes(new_bytes)
                self._structural_checks(path)

                reparsed = extract_contacts_from_docx(path, country=country)
                self.assertEqual(1, len(reparsed))
                self.assertIsNone(reparsed[0].member_firm)
                self.assertEqual(new_phone, reparsed[0].phone)

    def test_us_noop_rebuild_preserves_embedded_url_and_website(self) -> None:
        """Quirk C, real case: US's real address is a prose sentence
        that mentions "www.jacksonlewis.com." (with its own
        sentence-ending period), and website is the separate, clean
        "www.jacksonlewis.com". A no-op rebuild (identical fields, no
        edit at all) must reproduce BOTH exactly - the embedded
        mention must never be mistaken for the dedicated website
        line, and the dedicated website line must never pick up the
        sentence's trailing punctuation."""

        path = self._require_copy("US.docx")
        baseline = extract_contacts_from_docx(path, country="United States")

        if (
            len(baseline) != 1
            or not baseline[0].address
            or "please see www.jacksonlewis.com." not in baseline[0].address
        ):
            # A real Admin has since used the live Contact CRUD
            # feature against production US.docx (confirmed directly:
            # its own ContactState now shows zero contacts, deleted
            # via a genuine Add/Edit/Delete action, real-world content
            # drift unrelated to this fix) - the fix itself already
            # has permanent, corpus-independent coverage in
            # test_docx_parser.py's
            # test_website_mentioned_inside_address_prose_is_not_extracted,
            # which does not depend on live corpus content.
            self.skipTest(
                "US.docx's real content no longer has the exact "
                "embedded-URL contact this test assumes (real corpus "
                "content has drifted since this test was written) - "
                "see test_docx_parser.ClassifyCanonicalFirmLinesTests."
                "test_website_mentioned_inside_address_prose_is_not_extracted "
                "for permanent, corpus-independent coverage of this "
                "exact fix"
            )

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=tuple(baseline),
            photos=(None,),
            country="United States",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="United States")
        self.assertEqual(1, len(reparsed))
        self.assertEqual(baseline[0].address, reparsed[0].address)
        self.assertEqual("www.jacksonlewis.com", reparsed[0].website)

    def test_phone_with_trailing_annotation_round_trips(self) -> None:
        """Quirk A, real case: AU's real phone value, given a
        trailing annotation an Admin might legitimately type (e.g.
        "(mobile)"), must round-trip as ONE whole phone value -
        never split, with the annotation leaking into address."""

        path = self._require_copy("AU.docx")
        baseline = extract_contacts_from_docx(path, country="Australia")[0]

        annotated = ExtractedContact(
            member_firm=baseline.member_firm,
            contact_person=baseline.contact_person,
            email=baseline.email,
            phone=f"{baseline.phone} (mobile)",
            address=baseline.address,
            website=baseline.website,
        )

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(annotated,),
            photos=(None,),
            country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(1, len(reparsed))
        self.assertEqual(f"{baseline.phone} (mobile)", reparsed[0].phone)
        self.assertEqual(baseline.address, reparsed[0].address)

    def test_full_field_validator_rejects_a_field_shift(self) -> None:
        """The strengthened _validate_canonical_table() must reject a
        rebuild whose re-parsed contact shows a field-shift (a phone
        value landing in member_firm) - not just a contact_person/
        email mismatch. Simulated via a patched extract_contacts_
        from_docx, since the real reader no longer produces this
        shift on its own (that's the fix)."""

        path = self._require_copy("AU.docx")
        baseline = extract_contacts_from_docx(path, country="Australia")[0]

        shifted = ExtractedContact(
            member_firm=baseline.phone,
            contact_person=baseline.contact_person,
            email=baseline.email,
            phone=None,
            address=baseline.address,
            website=baseline.website,
        )

        with patch(
            "app.services.contact_document_area.extract_contacts_from_docx",
            return_value=[shifted],
        ):
            with self.assertRaises(ContactAreaError) as raised:
                rebuild_canonical_contact_table(
                    path,
                    contacts=(baseline,),
                    photos=(None,),
                    country="Australia",
                )

        message = str(raised.exception)
        self.assertIn("member_firm", message)
        self.assertIn("phone", message)

    # =====================================================================
    # FRANCE: legacy plain-paragraph contact area - France's real legacy
    # contact lives in ordinary body paragraphs near the END of the
    # document, immediately before "YOUR L&E GLOBAL POC", not a floating
    # shape near the start like AU/BE.
    # =====================================================================

    def _france_split_contacts(self) -> tuple[ExtractedContact, ExtractedContact]:
        path = self._require_copy("FR.docx")
        document = WordDocument(path)

        if any(
            table.rows
            and CONTACT_TABLE_HIDDEN_MARKER in table.rows[0].cells[0].text
            for table in document.tables
        ):
            # A real Admin has since used the live Contact CRUD
            # feature against production FR.docx (confirmed directly:
            # it now has a canonical contact table with Caroline/
            # Florence already split, plus real test contacts added
            # live - real-world content drift, and in this specific
            # case a genuine, welcome confirmation that the location
            # fix and multi-person split work correctly in practice).
            # The legacy-split scenario these tests exercise can no
            # longer be reproduced against the real corpus - see
            # FranceSyntheticContactCrudTests below for permanent,
            # drift-immune coverage of the same Add/Edit/Delete/photo/
            # clear-field behaviors.
            self.skipTest(
                "FR.docx has since been canonicalized by real Admin "
                "usage (real corpus content has drifted since this "
                "test was written) - see "
                "FranceSyntheticContactCrudTests for permanent, "
                "corpus-independent coverage of the same scenarios"
            )

        combined = extract_contacts_from_docx(path, country="France")[0]
        split = split_combined_legacy_contact(combined)
        self.assertIsNotNone(split)
        self.assertEqual(2, len(split))
        return split[0], split[1]

    def test_france_canonical_table_replaces_legacy_block_before_poc(
        self,
    ) -> None:
        """The canonical table must land at the EXACT location the
        legacy plain-paragraph block occupied - before "YOUR L&E
        GLOBAL POC" - never the document start, with the old legacy
        text fully removed and the POC/Jessica Stout block untouched."""

        path = self._require_copy("FR.docx")
        caroline, florence = self._france_split_contacts()

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(caroline, florence),
            photos=(None, None),
            country="France",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        document = WordDocument(path)
        full_text = "\n".join(p.text for p in document.paragraphs)

        self.assertEqual(1, len(document.tables))
        self.assertEqual(
            0,
            full_text.count("Caroline Scherrmann and Florence Bacquet"),
            "the old legacy combined-name paragraph must be gone",
        )
        self.assertEqual(
            1, full_text.count("YOUR L&E GLOBAL POC"),
            "the POC heading must survive exactly once",
        )
        self.assertEqual(
            1, full_text.count("Jessica Stout"),
            "Jessica Stout must survive - she is not a member-firm "
            "contact and must never be removed",
        )

        # The table must appear structurally BEFORE the POC paragraph,
        # never at the document start.
        body_children = list(document.element.body)
        table_position = next(
            i for i, child in enumerate(body_children)
            if child.tag.endswith("}tbl")
        )
        poc_position = next(
            i for i, child in enumerate(body_children)
            if child.tag.endswith("}p")
            and "YOUR L&E GLOBAL POC" in "".join(
                node.text or "" for node in child.iter()
                if node.tag.endswith("}t")
            )
        )
        self.assertLess(
            table_position, poc_position,
            "the canonical table must be inserted before the POC "
            "block, not at the document's start",
        )
        self.assertGreater(
            table_position, 10,
            "the table must not have landed at the very start of the "
            "document (the _default_insertion_anchor fallback must "
            "never fire when a real legacy contact area exists)",
        )

        reparsed = extract_contacts_from_docx(path, country="France")
        self.assertEqual(
            ["Caroline Scherrmann", "Florence Bacquet"],
            [c.contact_person for c in reparsed],
        )

    def test_france_add_temporary_contact_then_delete_round_trips(
        self,
    ) -> None:
        """Caroline/Florence baseline -> add a temporary third contact
        -> full-field round-trip valid, in order -> delete the
        temporary contact -> back to exactly Caroline/Florence, with no
        ContactAreaError at any stage (the internal round-trip
        validator is never weakened or bypassed - the data is made to
        satisfy it)."""

        path = self._require_copy("FR.docx")
        caroline, florence = self._france_split_contacts()
        temporary = ExtractedContact(
            member_firm="Temp Firm",
            contact_person="Temporary Person",
            email="temp@example.com",
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(caroline, florence), photos=(None, None),
            country="France",
        )
        path.write_bytes(new_bytes)

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(caroline, florence, temporary),
            photos=(None, None, None), country="France",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="France")
        self.assertEqual(
            ["Caroline Scherrmann", "Florence Bacquet", "Temporary Person"],
            [c.contact_person for c in reparsed],
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(caroline, florence), photos=(None, None),
            country="France",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="France")
        self.assertEqual(
            ["Caroline Scherrmann", "Florence Bacquet"],
            [c.contact_person for c in reparsed],
        )

    def test_france_photo_stays_with_the_correct_split_contact(
        self,
    ) -> None:
        """A photo attached to only one of the two split France
        contacts must never migrate to the other."""

        path = self._require_copy("FR.docx")
        caroline, florence = self._france_split_contacts()
        caroline_photo = ContactPhotoPayload(
            data=_VALID_PNG, content_type="image/png"
        )

        new_bytes = rebuild_canonical_contact_table(
            path,
            contacts=(caroline, florence),
            photos=(caroline_photo, None),
            country="France",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        photos = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(photos))
        self.assertEqual(_sha(caroline_photo.data), photos[0].sha256)

        reparsed = extract_contacts_from_docx(path, country="France")
        self.assertEqual(
            ["Caroline Scherrmann", "Florence Bacquet"],
            [c.contact_person for c in reparsed],
        )

    def test_france_clearing_website_persists_as_cleared(self) -> None:
        """Clearing a previously-set field must round-trip as actually
        cleared, never silently retaining the stale value."""

        path = self._require_copy("FR.docx")
        caroline, florence = self._france_split_contacts()

        with_website = ExtractedContact(
            member_firm=caroline.member_firm,
            contact_person=caroline.contact_person,
            email=caroline.email,
            phone=caroline.phone,
            website="www.example.com",
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(with_website, florence), photos=(None, None),
            country="France",
        )
        path.write_bytes(new_bytes)
        reparsed = extract_contacts_from_docx(path, country="France")
        self.assertEqual("www.example.com", reparsed[0].website)

        cleared = ExtractedContact(
            member_firm=caroline.member_firm,
            contact_person=caroline.contact_person,
            email=caroline.email,
            phone=caroline.phone,
            website=None,
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(cleared, florence), photos=(None, None),
            country="France",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="France")
        self.assertIsNone(
            reparsed[0].website,
            "a cleared field must round-trip as cleared, not stale",
        )

    def test_belgium_combined_contact_splits_consistently_like_france(
        self,
    ) -> None:
        """Verified directly against the real corpus rather than
        assumed: Belgium's own real legacy contact ALSO names two
        people sharing one comma-joined email string ("Chris van Olmen
        and Nicolas Simon"), the identical shape to France's. The
        multi-person split is a general, country-agnostic rule (never a
        France-specific exception), so it correctly and safely applies
        here too - this is the consistency the split logic must have,
        not a special case to avoid. What must genuinely stay
        unaffected is the ALREADY-validated canonical table mechanism's
        own handling of an explicitly-given, already-distinct
        multi-contact list: split_combined_legacy_contact is only ever
        invoked once, at legacy bootstrap time - never during an
        ordinary rebuild/update round-trip - so two contacts already
        split (by bootstrap or given directly, as here) keep
        round-tripping as two distinct, unmerged contacts."""

        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )
        document = WordDocument(path)
        if any(
            table.rows
            and CONTACT_TABLE_HIDDEN_MARKER in table.rows[0].cells[0].text
            for table in document.tables
        ):
            self.skipTest(
                "Belgium's document has since been canonicalized by "
                "real Admin usage (real corpus content has drifted "
                "since this test was written) - its combined legacy "
                "contact no longer exists in raw form to split"
            )

        baseline = extract_contacts_from_docx(path, country="Belgium")
        self.assertEqual(
            1, len(baseline),
            "Belgium's real legacy contact is currently one combined "
            "record, the same shape as France's, prior to any split",
        )

        split = split_combined_legacy_contact(baseline[0])
        self.assertIsNotNone(split)
        self.assertEqual(
            ["Chris van Olmen", "Nicolas Simon"],
            [c.contact_person for c in split],
        )

        # Once split (mirroring what bootstrap_legacy_contacts would
        # persist into ContactState), the canonical table mechanism
        # must keep them distinct across rebuilds, never re-merging or
        # re-splitting them.
        new_bytes = rebuild_canonical_contact_table(
            path, contacts=tuple(split),
            photos=tuple(None for _ in split), country="Belgium",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Belgium")
        self.assertEqual(
            ["Chris van Olmen", "Nicolas Simon"],
            [c.contact_person for c in reparsed],
        )

        for contact in reparsed:
            self.assertIsNone(
                split_combined_legacy_contact(contact),
                f"{contact.contact_person!r} is already a single "
                "person and must never be split further",
            )

    # =====================================================================
    # CANONICAL FIELD DETERMINISM: the canonical table is a format this
    # system owns (both writer and reader), so its own round-trip must be
    # fully deterministic - never dependent on content heuristics that can
    # misclassify a value shaped like a different field (a postal code
    # that looks phone-shaped, a real phone number embedded inside an
    # address). Confirmed live: a global Add+Photo dry-run over the real
    # corpus produced this EXACT swap for Brazil/Romania/Singapore-shaped
    # data, because the address field's own embedded phone number
    # outscored the shorter, dedicated (but genuinely intended) phone
    # value under the old content-based scorer. Fixed via a hidden
    # per-field tag written into the canonical table itself
    # (_FIELD_TAG_*, docx_parser.py) - legacy heuristic parsing of
    # organically-authored documents is completely untouched; an OLDER,
    # untagged canonical table (already written by a previous version of
    # this code) still falls back to the same content-based
    # classification as before.
    # =====================================================================

    def test_adversarial_address_phone_never_swaps(self) -> None:
        """The identical field-tag mechanism, confirmed independently
        against three real, differently-shaped adversarial addresses -
        never a country-specific patch."""

        cases = {
            "Brazil": (
                "Labour and Employment Law in Brazil 2026.docx",
                ExtractedContact(
                    member_firm="Tozzini Freire",
                    contact_person="Gabriela Lima",
                    email="glima@tozzinifreire.com.br",
                    address=(
                        "Rua Borges Lagoa, 1328, 04038-904 São Paulo, "
                        "+55 115 086 5000"
                    ),
                    phone="04038-904",
                ),
            ),
            "Romania": (
                "Labour and Employment Law in Romania 2026.docx",
                ExtractedContact(
                    member_firm="Volonciu & Associates",
                    contact_person="Magda Volonciu",
                    email="magdavolonciu@volonciu.ro",
                    address=(
                        "No. 35 Alexandru Constantinescu Street, 2nd "
                        "Floor, 011471 1st District Bucharest, "
                        "+40 372 755 699"
                    ),
                    phone="011471 1",
                ),
            ),
            "Singapore": (
                "Employment Law Overview Singapore 2026.docx",
                ExtractedContact(
                    member_firm="Clyde & Co Clasis",
                    contact_person="Thomas Choo",
                    email="thomas.choo@clydeco.com",
                    address=(
                        "12 Marina Boulevard | , Marina Bay Financial "
                        "Centre Tower 3 | #30 - 03, 018982 Singapore, "
                        "+65 654 465 00"
                    ),
                    phone="30 - 03",
                ),
            ),
        }

        for country, (filename, adversarial) in cases.items():
            with self.subTest(country=country):
                path = self._require_copy(filename)

                new_bytes = rebuild_canonical_contact_table(
                    path, contacts=(adversarial,), photos=(None,),
                    country=country,
                )
                path.write_bytes(new_bytes)
                self._structural_checks(path)

                reparsed = extract_contacts_from_docx(path, country=country)
                self.assertEqual(1, len(reparsed))
                self.assertEqual(adversarial.address, reparsed[0].address)
                self.assertEqual(adversarial.phone, reparsed[0].phone)

    def test_address_never_mistaken_for_website(self) -> None:
        """An address value that happens to contain a URL-shaped
        fragment must never be reclassified as the website field, and
        vice-versa - the tag alone decides, never content shape."""

        path = self._require_copy("AU.docx")

        adversarial = ExtractedContact(
            member_firm="Test Firm",
            contact_person="Test Person",
            email="test@example.com",
            address="123 www.looks-like-a-site.com Street",
            website="not-a-url-shaped-value",
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(adversarial,), photos=(None,),
            country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(1, len(reparsed))
        self.assertEqual(adversarial.address, reparsed[0].address)
        self.assertEqual(adversarial.website, reparsed[0].website)

    def test_full_six_field_canonical_round_trip_is_deterministic(
        self,
    ) -> None:
        """Every one of the six fields, including deliberately
        adversarial/ambiguous shapes, must round-trip byte-for-byte
        through a canonical rebuild - the general contract this whole
        mechanism exists to guarantee, not just the reported
        countries."""

        path = self._require_copy("AU.docx")

        adversarial = ExtractedContact(
            member_firm="123-456-7890",
            contact_person="Test Person",
            email="test@example.com",
            phone="www.not-really-a-website.com",
            address="+1 555 000 0000 is not the phone field here",
            website="04038-904",
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(adversarial,), photos=(None,),
            country="Australia",
        )
        path.write_bytes(new_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(1, len(reparsed))
        self.assertEqual(adversarial, reparsed[0])

    def test_website_without_www_prefix_never_leaks_into_address(
        self,
    ) -> None:
        """Explains a real, previously-observed divergence: a global
        dry-run using a "clean" (already www.-prefixed) website value
        passed, while a real Admin typing a website WITHOUT that
        prefix in the live browser (e.g. "zhonglun.com" instead of
        "www.zhonglun.com" - a completely natural human edit) hit
        "field(s) changed: address, website", because the old website
        detector required a www./http:// prefix to recognize a line as
        the dedicated website field at all; without it, the value fell
        through into address instead, and website came back None.
        Confirmed by reproducing this exact error message against the
        real China and Spain contacts on the pre-fix code - the
        identical architectural fix, not a country-specific patch."""

        cases = {
            "China": ("CN.docx", "China"),
            "Spain": (
                "Labour and Employment Law in Spain 2026.docx", "Spain",
            ),
        }

        for label, (filename, country) in cases.items():
            with self.subTest(country=label):
                path = self._require_copy(filename)
                baseline = extract_contacts_from_docx(
                    path, country=country
                )[0]

                edited = ExtractedContact(
                    member_firm=baseline.member_firm,
                    contact_person=baseline.contact_person,
                    email=baseline.email,
                    phone=baseline.phone,
                    address=baseline.address,
                    website=baseline.website.replace("www.", ""),
                )

                new_bytes = rebuild_canonical_contact_table(
                    path, contacts=(edited,), photos=(None,),
                    country=country,
                )
                path.write_bytes(new_bytes)
                self._structural_checks(path)

                reparsed = extract_contacts_from_docx(
                    path, country=country
                )
                self.assertEqual(1, len(reparsed))
                self.assertEqual(edited.address, reparsed[0].address)
                self.assertEqual(edited.website, reparsed[0].website)

    def test_old_untagged_canonical_table_still_reads_correctly(
        self,
    ) -> None:
        """An already-canonicalized real document, written by a
        version of this code before the field tags existed (verified
        directly: AR.docx's real production canonical table has none
        of them), must keep reading correctly via the same content-
        based fallback classification this reader has always used -
        this mechanism never breaks an already-deployed canonical
        table, it only makes NEW ones fully deterministic."""

        path = self._require_copy("AU.docx")

        contact = ExtractedContact(
            member_firm="Untagged Firm",
            contact_person="Untagged Person",
            email="untagged@example.com",
            phone="+1 555 010 0000",
            address="1 Untagged Street",
            website="www.untagged.example.com",
        )

        new_bytes = rebuild_canonical_contact_table(
            path, contacts=(contact,), photos=(None,),
            country="Australia",
        )
        path.write_bytes(new_bytes)

        # Simulate an OLDER canonical table by stripping the hidden
        # field tags back out, leaving the marker and visible values
        # exactly as an untagged table always looked.
        from app.services.docx_parser import (
            _FIELD_TAG_ADDRESS,
            _FIELD_TAG_CONTACT_PERSON,
            _FIELD_TAG_EMAIL,
            _FIELD_TAG_MEMBER_FIRM,
            _FIELD_TAG_PHONE,
            _FIELD_TAG_WEBSITE,
        )

        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read("word/document.xml").decode("utf-8")

        for tag in (
            _FIELD_TAG_MEMBER_FIRM, _FIELD_TAG_ADDRESS,
            _FIELD_TAG_PHONE, _FIELD_TAG_WEBSITE,
            _FIELD_TAG_CONTACT_PERSON, _FIELD_TAG_EMAIL,
        ):
            self.assertIn(
                tag, document_xml,
                "sanity: the tag must actually be present to strip",
            )
            document_xml = document_xml.replace(tag, "")

        from app.services.contact_document_photos import _rewrite_zip

        untagged_bytes = _rewrite_zip(
            path.read_bytes(),
            replacements={"word/document.xml": document_xml},
        )
        path.write_bytes(untagged_bytes)
        self._structural_checks(path)

        reparsed = extract_contacts_from_docx(path, country="Australia")
        self.assertEqual(1, len(reparsed))
        self.assertEqual(contact, reparsed[0])


def _build_synthetic_france_shaped_docx(path: Path) -> None:
    """
    A minimal, from-scratch DOCX matching France's own real structure
    exactly enough to exercise find_plain_paragraph_contact_block_
    bounds / split_combined_legacy_contact / rebuild_canonical_
    contact_table the same way real FR.docx used to, before real Admin
    usage canonicalized it in production (real corpus content drift -
    see ContactDocumentAreaTests._france_split_contacts).

    Shape: title -> padding body paragraphs (legal-looking filler, so
    the eventual canonical table lands well past the document start,
    matching the real "table must not land at the very start" check) ->
    the legacy plain-paragraph contact block (person / role+firm /
    comma-joined emails / phone) -> "YOUR L&E GLOBAL POC" -> Jessica
    Stout (the project-level POC, never a member-firm contact) ->
    disclaimer.
    """

    document = WordDocument()
    document.add_paragraph("Employment Law Overview - France")

    for index in range(12):
        document.add_paragraph(
            f"Filler legal paragraph number {index} with enough text "
            "to look like real body content."
        )

    document.add_paragraph("Caroline Scherrmann and Florence Bacquet")
    document.add_paragraph("Partners, Flichy Grangé Avocats")
    document.add_paragraph("scherrmann@flichy.com, bacquet@flichy.com")
    document.add_paragraph("+33 1 56 62 30 00")

    document.add_paragraph("YOUR L&E GLOBAL POC")
    document.add_paragraph("Jessica Stout")
    document.add_paragraph("jstout@leglobal.law")

    document.add_paragraph(
        "Disclaimer: this publication is for informational purposes "
        "only and does not constitute legal advice."
    )

    document.save(path)


class FranceSyntheticContactCrudTests(unittest.TestCase):
    """
    Permanent, corpus-independent coverage of France's own special
    contact shape - a synthetic fixture built fresh in setUp, immune to
    real corpus content drift (unlike ContactDocumentAreaTests._france_
    split_contacts, which now skips because real FR.docx has since been
    canonicalized by genuine Admin usage in production).
    """

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.path = Path(self.temp.name) / "france-synthetic.docx"
        _build_synthetic_france_shaped_docx(self.path)

    def _split_contacts(self) -> tuple[ExtractedContact, ExtractedContact]:
        combined = extract_contacts_from_docx(
            self.path, country="France"
        )[0]
        split = split_combined_legacy_contact(combined)
        self.assertIsNotNone(split)
        self.assertEqual(2, len(split))
        return split[0], split[1]

    def test_canonical_table_replaces_legacy_block_before_poc(
        self,
    ) -> None:
        caroline, florence = self._split_contacts()

        new_bytes = rebuild_canonical_contact_table(
            self.path,
            contacts=(caroline, florence),
            photos=(None, None),
            country="France",
        )
        self.path.write_bytes(new_bytes)

        document = WordDocument(self.path)
        full_text = "\n".join(p.text for p in document.paragraphs)

        self.assertEqual(1, len(document.tables))
        self.assertEqual(
            0,
            full_text.count("Caroline Scherrmann and Florence Bacquet"),
            "the old legacy combined-name paragraph must be gone",
        )
        self.assertEqual(1, full_text.count("YOUR L&E GLOBAL POC"))
        self.assertEqual(1, full_text.count("Jessica Stout"))

        body_children = list(document.element.body)
        table_position = next(
            i for i, child in enumerate(body_children)
            if child.tag.endswith("}tbl")
        )
        poc_position = next(
            i for i, child in enumerate(body_children)
            if child.tag.endswith("}p")
            and "YOUR L&E GLOBAL POC" in "".join(
                node.text or "" for node in child.iter()
                if node.tag.endswith("}t")
            )
        )
        self.assertLess(table_position, poc_position)
        self.assertGreater(table_position, 10)

        reparsed = extract_contacts_from_docx(self.path, country="France")
        self.assertEqual(
            ["Caroline Scherrmann", "Florence Bacquet"],
            [c.contact_person for c in reparsed],
        )

    def test_add_temporary_contact_then_delete_round_trips(self) -> None:
        caroline, florence = self._split_contacts()
        temporary = ExtractedContact(
            member_firm="Temp Firm",
            contact_person="Temporary Person",
            email="temp@example.com",
        )

        self.path.write_bytes(
            rebuild_canonical_contact_table(
                self.path, contacts=(caroline, florence),
                photos=(None, None), country="France",
            )
        )
        self.path.write_bytes(
            rebuild_canonical_contact_table(
                self.path, contacts=(caroline, florence, temporary),
                photos=(None, None, None), country="France",
            )
        )

        reparsed = extract_contacts_from_docx(self.path, country="France")
        self.assertEqual(
            ["Caroline Scherrmann", "Florence Bacquet", "Temporary Person"],
            [c.contact_person for c in reparsed],
        )

        self.path.write_bytes(
            rebuild_canonical_contact_table(
                self.path, contacts=(caroline, florence),
                photos=(None, None), country="France",
            )
        )

        reparsed = extract_contacts_from_docx(self.path, country="France")
        self.assertEqual(
            ["Caroline Scherrmann", "Florence Bacquet"],
            [c.contact_person for c in reparsed],
        )

    def test_photo_stays_with_the_correct_split_contact(self) -> None:
        caroline, florence = self._split_contacts()
        caroline_photo = ContactPhotoPayload(
            data=_VALID_PNG, content_type="image/png"
        )

        self.path.write_bytes(
            rebuild_canonical_contact_table(
                self.path,
                contacts=(caroline, florence),
                photos=(caroline_photo, None),
                country="France",
            )
        )

        photos = extract_contact_photo_candidates(self.path)
        self.assertEqual(1, len(photos))
        self.assertEqual(_sha(caroline_photo.data), photos[0].sha256)

        reparsed = extract_contacts_from_docx(self.path, country="France")
        self.assertEqual(
            ["Caroline Scherrmann", "Florence Bacquet"],
            [c.contact_person for c in reparsed],
        )

    def test_clearing_website_persists_as_cleared(self) -> None:
        caroline, florence = self._split_contacts()

        with_website = ExtractedContact(
            member_firm=caroline.member_firm,
            contact_person=caroline.contact_person,
            email=caroline.email,
            phone=caroline.phone,
            website="www.example.com",
        )

        self.path.write_bytes(
            rebuild_canonical_contact_table(
                self.path, contacts=(with_website, florence),
                photos=(None, None), country="France",
            )
        )
        reparsed = extract_contacts_from_docx(self.path, country="France")
        self.assertEqual("www.example.com", reparsed[0].website)

        cleared = ExtractedContact(
            member_firm=caroline.member_firm,
            contact_person=caroline.contact_person,
            email=caroline.email,
            phone=caroline.phone,
            website=None,
        )

        self.path.write_bytes(
            rebuild_canonical_contact_table(
                self.path, contacts=(cleared, florence),
                photos=(None, None), country="France",
            )
        )
        reparsed = extract_contacts_from_docx(self.path, country="France")
        self.assertIsNone(reparsed[0].website)


# =========================================================================
# OLD FLOATING-SHAPE PHOTO MUTATION PRIMITIVES (contact_document_photos.py)
#
# The mechanism the canonical contact table (above) replaced for the
# textual contact area; contact_document_photos.py's own low-level image
# mutation primitives (replace/remove/add-by-name/add-new-anchor) remain a
# real, independently importable part of this service module, and
# _is_contact_related_block/_find_all_contact_runs (bottom of this file)
# still detect whether a document's Contact area is in this OLD format.
# =========================================================================

# Degenerate (not real-dimensioned) fixtures are deliberately fine here -
# unlike contact_document_area.py's canonical table, these primitives
# never compute a photo's proportional height from its own header, so a
# minimal 1x1-scale placeholder round-trips exactly like a real photo
# would.
_LEGACY_VALID_JPEG = bytes.fromhex(
    "ffd8ffe000104a46494600010100000100010000ffd9"
)
_LEGACY_VALID_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000"
    "000108060000001f15c4890000000a49444154789c63"
    "60000002000100ffff03000006000557bfabd4000000"
    "0049454e44ae426082"
)


class RealCorpusContactDocumentPhotoTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)

    def _require_copy(self, filename: str) -> Path:
        return require_corpus_copy(
            self, SOURCE_ROOT, filename, Path(self.temp.name)
        )

    def _unrelated_parts(self, path: Path) -> dict[str, bytes]:
        """Every zip part except document.xml/rels/content-types -
        the parts a photo mutation must never touch."""

        skip = {
            "word/document.xml",
            "word/_rels/document.xml.rels",
            "[Content_Types].xml",
        }

        with zipfile.ZipFile(path) as archive:
            return {
                name: archive.read(name)
                for name in archive.namelist()
                if name not in skip
            }

    def test_belgium_replace_isolates_the_other_contacts_photo(
        self,
    ) -> None:
        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        original_candidates = extract_contact_photo_candidates(path)
        by_name = {
            c.source_filename: c for c in original_candidates
        }
        chris = by_name["image2.jpg"]
        chris_sha = chris.sha256
        nicolas_sha = by_name["image1.png"].sha256

        original_paragraphs = len(Document(str(path)).paragraphs)
        unrelated_before = self._unrelated_parts(path)
        del unrelated_before[chris.media_path]

        new_bytes = replace_contact_photo_in_document(
            path,
            target_sha256=chris_sha,
            new_data=_LEGACY_VALID_JPEG,
            new_content_type="image/jpeg",
        )

        out = Path(self.temp.name) / "belgium_replaced.docx"
        out.write_bytes(new_bytes)

        new_shas = {
            c.sha256 for c in extract_contact_photo_candidates(out)
        }

        self.assertIn(
            hashlib.sha256(_LEGACY_VALID_JPEG).hexdigest(), new_shas
        )
        self.assertNotIn(chris_sha, new_shas)
        self.assertIn(nicolas_sha, new_shas)

        self.assertEqual(
            original_paragraphs, len(Document(str(out)).paragraphs)
        )
        unrelated_after = self._unrelated_parts(out)
        del unrelated_after[chris.media_path]
        self.assertEqual(unrelated_before, unrelated_after)

    def test_belgium_remove_isolates_the_other_contacts_photo(
        self,
    ) -> None:
        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )
        skip_if_already_canonicalized(self, path)

        by_name = {
            c.source_filename: c.sha256
            for c in extract_contact_photo_candidates(path)
        }
        chris_sha = by_name["image2.jpg"]
        nicolas_sha = by_name["image1.png"]

        original_paragraphs = len(Document(str(path)).paragraphs)

        new_bytes = remove_contact_photo_from_document(
            path,
            target_sha256=nicolas_sha,
        )

        out = Path(self.temp.name) / "belgium_removed.docx"
        out.write_bytes(new_bytes)

        remaining = extract_contact_photo_candidates(out)

        self.assertEqual(1, len(remaining))
        self.assertEqual(chris_sha, remaining[0].sha256)
        self.assertEqual(
            original_paragraphs, len(Document(str(out)).paragraphs)
        )

    def test_belgium_add_into_a_shared_zone_fails_closed(
        self,
    ) -> None:
        """
        Chris and Nicolas share ONE combined "CONTACT PERSON" textbox
        zone, disambiguated only by their two photos' own separate
        geometry - never by two separate zones. Adding a photo for
        Nicolas after his own is removed has no safe, deterministic
        place to insert it (the zone still names Chris too), so this
        must fail closed rather than guess.
        """

        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        by_name = {
            c.source_filename: c.sha256
            for c in extract_contact_photo_candidates(path)
        }
        nicolas_sha = by_name["image1.png"]

        removed_bytes = remove_contact_photo_from_document(
            path,
            target_sha256=nicolas_sha,
        )
        removed_path = Path(self.temp.name) / "belgium_removed.docx"
        removed_path.write_bytes(removed_bytes)

        with self.assertRaises(ContactDocumentPhotoError):
            add_contact_photo_to_document(
                removed_path,
                contact_person="Nicolas Simon",
                new_data=_LEGACY_VALID_JPEG,
                new_content_type="image/jpeg",
                other_contact_persons=["Chris van Olmen"],
            )

    def test_germany_add_to_a_genuinely_photo_less_contact(
        self,
    ) -> None:
        path = self._require_copy("DE.docx")

        original_candidates = extract_contact_photo_candidates(path)
        self.assertEqual(
            0,
            len(original_candidates),
            "Tobias Pusch is expected to have no existing photo in "
            "the real corpus - if this fails, the corpus changed and "
            "this test's premise needs revisiting.",
        )

        original_paragraphs = len(Document(str(path)).paragraphs)

        new_bytes = add_contact_photo_to_document(
            path,
            contact_person="Tobias Pusch",
            new_data=_LEGACY_VALID_PNG,
            new_content_type="image/png",
            other_contact_persons=[],
        )

        out = Path(self.temp.name) / "de_added.docx"
        out.write_bytes(new_bytes)

        candidates = extract_contact_photo_candidates(out)

        self.assertEqual(1, len(candidates))
        self.assertEqual(
            hashlib.sha256(_LEGACY_VALID_PNG).hexdigest(),
            candidates[0].sha256,
        )
        self.assertEqual(
            original_paragraphs, len(Document(str(out)).paragraphs)
        )

    def test_argentina_replace_then_remove_single_contact(
        self,
    ) -> None:
        path = self._require_copy("AR.docx")

        original_candidates = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(original_candidates))
        target_sha = original_candidates[0].sha256

        original_paragraphs = len(Document(str(path)).paragraphs)

        replaced_bytes = replace_contact_photo_in_document(
            path,
            target_sha256=target_sha,
            new_data=_LEGACY_VALID_JPEG,
            new_content_type="image/jpeg",
        )
        replaced_path = Path(self.temp.name) / "ar_replaced.docx"
        replaced_path.write_bytes(replaced_bytes)

        replaced_candidates = extract_contact_photo_candidates(
            replaced_path
        )
        self.assertEqual(1, len(replaced_candidates))
        self.assertEqual(
            hashlib.sha256(_LEGACY_VALID_JPEG).hexdigest(),
            replaced_candidates[0].sha256,
        )
        self.assertEqual(
            original_paragraphs,
            len(Document(str(replaced_path)).paragraphs),
        )

        removed_bytes = remove_contact_photo_from_document(
            replaced_path,
            target_sha256=hashlib.sha256(_LEGACY_VALID_JPEG).hexdigest(),
        )
        removed_path = Path(self.temp.name) / "ar_removed.docx"
        removed_path.write_bytes(removed_bytes)

        self.assertEqual(
            0, len(extract_contact_photo_candidates(removed_path))
        )
        self.assertEqual(
            original_paragraphs,
            len(Document(str(removed_path)).paragraphs),
        )

    def test_replace_of_an_unlocatable_sha_fails_closed(
        self,
    ) -> None:
        path = self._require_copy("AR.docx")

        with self.assertRaises(ContactDocumentPhotoError):
            replace_contact_photo_in_document(
                path,
                target_sha256="0" * 64,
                new_data=_LEGACY_VALID_JPEG,
                new_content_type="image/jpeg",
            )

    def test_remove_of_an_unlocatable_sha_fails_closed(
        self,
    ) -> None:
        path = self._require_copy("AR.docx")

        with self.assertRaises(ContactDocumentPhotoError):
            remove_contact_photo_from_document(
                path,
                target_sha256="0" * 64,
            )

    def test_add_for_an_unknown_contact_person_fails_closed(
        self,
    ) -> None:
        path = self._require_copy("DE.docx")

        with self.assertRaises(ContactDocumentPhotoError):
            add_contact_photo_to_document(
                path,
                contact_person="Nobody Real",
                new_data=_LEGACY_VALID_PNG,
                new_content_type="image/png",
                other_contact_persons=[],
            )

    def test_replace_of_the_existing_zone_now_uses_true_geometry(
        self,
    ) -> None:
        """
        Regression guard for a fixed bug: adding a photo for an
        EXISTING named zone must genuinely geometrically overlap that
        zone (reason "GEOMETRY"), never merely happen to be the sole
        remaining portrait in the whole document (reason
        "UNIQUE_PORTRAIT") - the fallback only ever works by
        coincidence and breaks the moment the document has more than
        one photo total, exactly the scenario a brand-new second
        contact's photo introduces.
        """

        path = self._require_copy("DE.docx")

        new_bytes = add_contact_photo_to_document(
            path,
            contact_person="Tobias Pusch",
            new_data=_LEGACY_VALID_PNG,
            new_content_type="image/png",
            other_contact_persons=[],
        )

        out = Path(self.temp.name) / "de_geometry.docx"
        out.write_bytes(new_bytes)

        candidates = extract_contact_photo_candidates(out)
        self.assertEqual(1, len(candidates))
        self.assertEqual("GEOMETRY", candidates[0].reason)

    def test_add_new_contact_photo_anchors_to_the_largest_zone_alongside_an_existing_photo(
        self,
    ) -> None:
        """
        The core blocker scenario this primitive exists for: a
        document that ALREADY has one contact's photo, and a genuinely
        brand-new second contact (whose name cannot possibly appear
        anywhere yet) also gets a photo. Both must round-trip as
        independently valid, GEOMETRY-reasoned candidates - never
        relying on the "exactly one remaining portrait" fallback,
        which cannot disambiguate once there is more than one
        unassociated photo.
        """

        path = self._require_copy("AR.docx")
        skip_if_already_canonicalized(self, path)

        original = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(original))
        original_sha = original[0].sha256

        original_paragraphs = len(Document(str(path)).paragraphs)

        new_bytes = add_new_contact_photo_to_document(
            path,
            new_data=_LEGACY_VALID_JPEG,
            new_content_type="image/jpeg",
        )

        out = Path(self.temp.name) / "ar_new_contact.docx"
        out.write_bytes(new_bytes)

        candidates = extract_contact_photo_candidates(out)
        self.assertEqual(2, len(candidates))

        for candidate in candidates:
            self.assertEqual(
                "GEOMETRY",
                candidate.reason,
                "must be a real geometric match, never a fallback "
                "that only works by coincidence",
            )

        new_shas = {c.sha256 for c in candidates}
        self.assertIn(original_sha, new_shas)
        self.assertIn(
            hashlib.sha256(_LEGACY_VALID_JPEG).hexdigest(), new_shas
        )
        self.assertEqual(
            original_paragraphs, len(Document(str(out)).paragraphs)
        )

    def test_add_new_contact_photo_into_belgiums_shared_zone_isolates_existing_photos(
        self,
    ) -> None:
        """
        Belgium's shared "CONTACT PERSON" zone (Chris + Nicolas) is
        the document's only (and therefore largest) CONTACT PERSON
        zone - a brand-new third contact's photo anchors there too,
        and must never disturb either of the two existing contacts'
        own photos.
        """

        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )
        skip_if_already_canonicalized(self, path)

        original_shas = {
            c.sha256 for c in extract_contact_photo_candidates(path)
        }
        self.assertEqual(2, len(original_shas))

        new_bytes = add_new_contact_photo_to_document(
            path,
            new_data=_LEGACY_VALID_PNG,
            new_content_type="image/png",
        )

        out = Path(self.temp.name) / "belgium_new_contact.docx"
        out.write_bytes(new_bytes)

        candidates = extract_contact_photo_candidates(out)
        new_shas = {c.sha256 for c in candidates}

        self.assertEqual(3, len(new_shas))
        self.assertTrue(original_shas.issubset(new_shas))
        self.assertIn(
            hashlib.sha256(_LEGACY_VALID_PNG).hexdigest(), new_shas
        )

    def test_add_new_contact_photo_fails_closed_with_no_contact_zone_at_all(
        self,
    ) -> None:
        """
        A document with genuinely zero "CONTACT PERSON" zones (no
        structural home at all) must fail closed rather than insert
        the photo at an arbitrary position such as the document's end.
        """

        path = self._require_copy("PT.docx")

        self.assertEqual(0, len(extract_contact_photo_candidates(path)))

        with self.assertRaises(ContactDocumentPhotoError):
            add_new_contact_photo_to_document(
                path,
                new_data=_LEGACY_VALID_JPEG,
                new_content_type="image/jpeg",
            )


# =========================================================================
# RAW PHOTO EXTRACTION (contact_photos.py)
# =========================================================================


class RealCorpusContactPhotoTests(unittest.TestCase):

    def _require(self, filename: str) -> Path:
        path = SOURCE_ROOT / filename

        if not path.exists():
            self.skipTest(
                f"Real corpus source unavailable: {path}"
            )

        return path

    def test_belgium_has_exactly_two_contact_photos(self) -> None:
        path = self._require(
            "Labour and Employment Law in Belgium 2026.docx"
        )
        skip_if_already_canonicalized(self, path)

        photos = extract_contact_photo_candidates(path)

        self.assertEqual(2, len(photos))

        self.assertEqual(
            ["image2.jpg", "image1.png"],
            [photo.source_filename for photo in photos],
        )

        self.assertEqual(
            ["GEOMETRY", "GEOMETRY"],
            [photo.reason for photo in photos],
        )

        for photo in photos:
            self.assertTrue(photo.data)
            self.assertTrue(photo.sha256)
            self.assertTrue(photo.content_type.startswith("image/"))

    def test_ireland_rejects_the_wide_false_positive(self) -> None:
        path = self._require("IE.docx")

        photos = extract_contact_photo_candidates(path)

        self.assertEqual(1, len(photos))
        self.assertEqual(
            "image2.jpg",
            photos[0].source_filename,
        )

    def test_indonesia_rejects_pagoda_and_logo(self) -> None:
        path = self._require("ID.docx")

        photos = extract_contact_photo_candidates(path)

        self.assertEqual(1, len(photos))
        self.assertEqual(
            "image3.jpeg",
            photos[0].source_filename,
        )
        self.assertEqual(
            "UNIQUE_PORTRAIT",
            photos[0].reason,
        )

    def test_chile_has_no_contact_photo(self) -> None:
        path = self._require("CL.docx")

        self.assertEqual(
            [],
            extract_contact_photo_candidates(path),
        )

    def test_germany_has_no_contact_photo(self) -> None:
        path = self._require("DE.docx")

        self.assertEqual(
            [],
            extract_contact_photo_candidates(path),
        )

    def test_india_has_no_contact_photo(self) -> None:
        path = self._require("IN.docx")

        self.assertEqual(
            [],
            extract_contact_photo_candidates(path),
        )

    def test_france_has_no_contact_photo(self) -> None:
        path = self._require("FR.docx")
        skip_if_already_canonicalized(self, path)

        self.assertEqual(
            [],
            extract_contact_photo_candidates(path),
        )

    def test_result_is_deterministic(self) -> None:
        path = self._require(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        first = extract_contact_photo_candidates(path)
        second = extract_contact_photo_candidates(path)

        self.assertEqual(
            [
                (
                    item.source_filename,
                    item.sha256,
                    item.reason,
                )
                for item in first
            ],
            [
                (
                    item.source_filename,
                    item.sha256,
                    item.reason,
                )
                for item in second
            ],
        )

    def test_candidate_is_an_immutable_value_object(self) -> None:
        candidate = ContactPhotoCandidate(
            source_filename="portrait.jpg",
            content_type="image/jpeg",
            data=b"photo",
            sha256="abc",
            reason="GEOMETRY",
        )

        with self.assertRaises(
            (AttributeError, TypeError),
        ):
            candidate.source_filename = "changed.jpg"


# =========================================================================
# CONTACT/PHOTO ASSOCIATION (contact_people.py)
# =========================================================================


def _photo(name: str) -> ContactPhotoCandidate:
    return ContactPhotoCandidate(
        source_filename=name,
        content_type="image/jpeg",
        data=name.encode(),
        sha256="sha-" + name,
        reason="GEOMETRY",
    )


class ContactPhotoAssociationUnitTests(unittest.TestCase):

    def test_single_contact_single_photo_is_associated(self) -> None:
        contact = ExtractedContact(
            member_firm="Firm",
            contact_person="Jane Doe",
            email="jane@example.com",
            phone="+1 111",
            address="Address",
            website="example.com",
        )

        result = associate_contact_photos(
            [contact],
            [_photo("jane.jpg")],
        )

        self.assertEqual(1, len(result))
        self.assertEqual("Jane Doe", result[0].contact_person)
        self.assertEqual("jane@example.com", result[0].email)
        self.assertEqual(
            "jane.jpg",
            result[0].photo.source_filename,
        )

    def test_single_contact_without_photo_is_preserved(self) -> None:
        contact = ExtractedContact(
            member_firm="Firm",
            contact_person="Jane Doe",
            email="jane@example.com",
            phone="+1 111",
            address="Address",
            website="example.com",
        )

        result = associate_contact_photos(
            [contact],
            [],
        )

        self.assertEqual(1, len(result))
        self.assertEqual("Jane Doe", result[0].contact_person)
        self.assertIsNone(result[0].photo)

    def test_two_existing_contacts_two_photos_are_mapped_in_order(
        self,
    ) -> None:
        contacts = [
            ExtractedContact(
                member_firm="Firm",
                contact_person="Person One",
                email="one@example.com",
            ),
            ExtractedContact(
                member_firm="Firm",
                contact_person="Person Two",
                email="two@example.com",
            ),
        ]

        result = associate_contact_photos(
            contacts,
            [
                _photo("one.jpg"),
                _photo("two.jpg"),
            ],
        )

        self.assertEqual(2, len(result))

        self.assertEqual(
            ("Person One", "one@example.com", "one.jpg"),
            (
                result[0].contact_person,
                result[0].email,
                result[0].photo.source_filename,
            ),
        )

        self.assertEqual(
            ("Person Two", "two@example.com", "two.jpg"),
            (
                result[1].contact_person,
                result[1].email,
                result[1].photo.source_filename,
            ),
        )

    def test_combined_contact_is_split_only_when_photo_count_matches(
        self,
    ) -> None:
        contact = ExtractedContact(
            member_firm="Shared Firm",
            contact_person="Person One and Person Two",
            email="one@example.com, two@example.com",
            phone="+32 123",
            address="Shared address",
            website="firm.example",
        )

        result = associate_contact_photos(
            [contact],
            [
                _photo("one.jpg"),
                _photo("two.jpg"),
            ],
        )

        self.assertEqual(2, len(result))

        self.assertEqual("Person One", result[0].contact_person)
        self.assertEqual("one@example.com", result[0].email)
        self.assertEqual("one.jpg", result[0].photo.source_filename)

        self.assertEqual("Person Two", result[1].contact_person)
        self.assertEqual("two@example.com", result[1].email)
        self.assertEqual("two.jpg", result[1].photo.source_filename)

        # Shared firm-level fields are copied to each autonomous card.
        for item in result:
            self.assertEqual("Shared Firm", item.member_firm)
            self.assertEqual("+32 123", item.phone)
            self.assertEqual("Shared address", item.address)
            self.assertEqual("firm.example", item.website)

    def test_combined_contact_without_photos_is_not_split(self) -> None:
        contact = ExtractedContact(
            member_firm="Firm",
            contact_person="Person One and Person Two",
            email="one@example.com, two@example.com",
        )

        result = associate_contact_photos(
            [contact],
            [],
        )

        self.assertEqual(1, len(result))
        self.assertEqual(
            "Person One and Person Two",
            result[0].contact_person,
        )
        self.assertEqual(
            "one@example.com, two@example.com",
            result[0].email,
        )
        self.assertIsNone(result[0].photo)

    def test_mismatched_person_email_photo_counts_never_guess(self) -> None:
        contact = ExtractedContact(
            member_firm="Firm",
            contact_person="Person One and Person Two",
            email="one@example.com",
        )

        result = associate_contact_photos(
            [contact],
            [
                _photo("one.jpg"),
                _photo("two.jpg"),
            ],
        )

        self.assertEqual(1, len(result))
        self.assertEqual(
            "Person One and Person Two",
            result[0].contact_person,
        )
        self.assertEqual(
            "one@example.com",
            result[0].email,
        )
        self.assertIsNone(result[0].photo)

    def test_extra_photos_never_get_arbitrarily_assigned(self) -> None:
        contact = ExtractedContact(
            member_firm="Firm",
            contact_person="Jane Doe",
            email="jane@example.com",
        )

        result = associate_contact_photos(
            [contact],
            [
                _photo("a.jpg"),
                _photo("b.jpg"),
            ],
        )

        self.assertEqual(1, len(result))
        self.assertIsNone(result[0].photo)


class RealCorpusContactPersonPhotoTests(unittest.TestCase):

    def _require(self, filename: str) -> Path:
        path = SOURCE_ROOT / filename

        if not path.exists():
            self.skipTest(
                f"Real corpus source unavailable: {path}"
            )

        return path

    def test_belgium_becomes_two_individual_contacts(self) -> None:
        path = self._require(
            "Labour and Employment Law in Belgium 2026.docx"
        )
        skip_if_already_canonicalized(self, path)

        parsed = extract_contacts_from_docx(path)
        photos = extract_contact_photo_candidates(path)

        self.assertEqual(1, len(parsed))
        self.assertEqual(2, len(photos))

        result = associate_contact_photos(
            parsed,
            photos,
        )

        self.assertEqual(2, len(result))

        self.assertEqual(
            "Chris van Olmen",
            result[0].contact_person,
        )
        self.assertEqual(
            "chris.van.olmen@vow.be",
            result[0].email,
        )
        self.assertEqual(
            "image2.jpg",
            result[0].photo.source_filename,
        )

        self.assertEqual(
            "Nicolas Simon",
            result[1].contact_person,
        )
        self.assertEqual(
            "nicolas.simon@vow.be",
            result[1].email,
        )
        self.assertEqual(
            "image1.png",
            result[1].photo.source_filename,
        )

        for contact in result:
            self.assertEqual(
                "Van Olmen & Wynant",
                contact.member_firm,
            )
            self.assertEqual(
                "+32 264 405 11",
                contact.phone,
            )
            self.assertEqual(
                "www.vow.be",
                contact.website,
            )

    def test_france_remains_backward_compatible_without_photos(
        self,
    ) -> None:
        path = self._require("FR.docx")
        skip_if_already_canonicalized(self, path)

        parsed = extract_contacts_from_docx(path)
        photos = extract_contact_photo_candidates(path)

        self.assertEqual([], photos)

        result = associate_contact_photos(
            parsed,
            photos,
        )

        self.assertEqual(1, len(result))
        self.assertEqual(
            "Caroline Scherrmann and Florence Bacquet",
            result[0].contact_person,
        )
        self.assertEqual(
            "scherrmann@flichy.com, bacquet@flichy.com",
            result[0].email,
        )
        self.assertIsNone(result[0].photo)


# =========================================================================
# OLD FLOATING-SHAPE CONTACT-BOX DETECTION
# (document_contact_materializer.py)
#
# This module used to ALSO drive materialize_effective_docx() (a
# per-download ephemeral DOCX rebuild); that function was removed once
# every Admin mutation started persisting its effective result into the
# source DOCX atomically before returning (see
# test_admin_documents_router_integration.py's DownloadHttpContractTests
# and backend/integration_tests/docx_contact_mutation_matrix.py for that
# coverage). What remains here - _is_contact_related_block/
# _find_all_contact_runs - is still live: contact_document_photos.py uses
# it to detect whether a document's Contact area is still the OLD
# floating-shape format.
# =========================================================================


def _textbox_run_xml(
    lines: list[str],
    width_emu: int = 1000000,
    height_emu: int = 500000,
) -> str:
    inner_paragraphs = "".join(
        f"<w:p><w:r><w:t>{line}</w:t></w:r></w:p>" for line in lines
    )
    return (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<w:rPr><w:noProof/></w:rPr>"
        "<w:drawing><wp:anchor>"
        f'<wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        '<wp:docPr id="1" name="Box"/>'
        "<a:graphic>"
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp>"
        f"<wps:txbx><w:txbxContent>{inner_paragraphs}</w:txbxContent></wps:txbx>"
        "</wps:wsp></a:graphicData></a:graphic>"
        "</wp:anchor></w:drawing></w:r>"
    )


def _build_docx_with_contact_textbox_at(
    paragraphs: list[str],
    textbox_paragraph_index: int,
    lines: list[str] | None = None,
    width_emu: int = 1000000,
    height_emu: int = 500000,
) -> bytes:
    """
    A real, openable DOCX with a minimal but structurally faithful
    floating Contact text box (CONTACT PERSON marker + name + email,
    matching the real corpus's own <w:txbxContent> shape, wrapped in a
    DrawingML <w:drawing><wp:anchor> run so _extract_extent and the
    <w:r>/<w:txbxContent> nesting _find_run_span must see through both
    apply) injected into one specific paragraph.
    """

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)

    target_paragraph = document.paragraphs[textbox_paragraph_index]

    run_element = etree.fromstring(
        _textbox_run_xml(
            lines
            or ["CONTACT PERSON", "Test Person", "test.person@example.com"],
            width_emu=width_emu,
            height_emu=height_emu,
        ).encode("utf-8")
    )
    target_paragraph._p.append(run_element)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _document_xml_of(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
        return archive.read("word/document.xml").decode("utf-8")


class IsContactRelatedBlockTests(unittest.TestCase):
    def test_contact_person_block_is_classified_as_contact_related(
        self,
    ) -> None:
        self.assertTrue(
            _is_contact_related_block(["CONTACT PERSON", "Jane Doe"])
        )

    def test_firm_block_with_email_is_classified_as_contact_related(
        self,
    ) -> None:
        self.assertTrue(
            _is_contact_related_block(["SOME FIRM", "info@example.com"])
        )

    def test_firm_block_with_phone_is_classified_as_contact_related(
        self,
    ) -> None:
        self.assertTrue(
            _is_contact_related_block(["SOME FIRM", "+1 555 000 1234"])
        )

    def test_plain_branding_block_is_not_contact_related(self) -> None:
        self.assertFalse(_is_contact_related_block(["www.leglobal.law"]))

    def test_document_title_block_is_not_contact_related(self) -> None:
        self.assertFalse(
            _is_contact_related_block(
                ["Employment Law Overview - Testland"]
            )
        )


class FindAllContactRunsTests(unittest.TestCase):
    """
    Regression coverage for two bugs found and fixed together: (1)
    _find_run_span originally took the first "</w:r>" after a
    txbxContent's start, which is actually an INNER run nested inside
    the box's own paragraphs, not the enclosing floating-shape run: it
    must instead balance nested <w:r>/</w:r> tags. (2) a floating
    shape's Choice and Fallback branches each produce their own
    <w:txbxContent> regex match for the SAME enclosing run; the second
    occurrence must be recognized as already-covered by the first
    match's span, not re-resolved (which - before the fix - picked one
    of the first branch's own inner runs as a bogus second "run").
    """

    def test_finds_single_contact_run_with_its_extent(self) -> None:
        docx_bytes = _build_docx_with_contact_textbox_at(
            ["Title", "Body"],
            textbox_paragraph_index=0,
            width_emu=1234000,
            height_emu=567000,
        )

        runs = _find_all_contact_runs(_document_xml_of(docx_bytes))

        self.assertEqual(len(runs), 1)
        self.assertEqual(runs[0].width_emu, 1234000)
        self.assertEqual(runs[0].height_emu, 567000)

    def test_ignores_non_contact_textboxes(self) -> None:
        document = Document()
        document.add_paragraph("Title")
        document.add_paragraph("Body")

        branding_run = etree.fromstring(
            _textbox_run_xml(["www.leglobal.law"]).encode("utf-8")
        )
        document.paragraphs[0]._p.append(branding_run)

        buffer = io.BytesIO()
        document.save(buffer)

        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))
        self.assertEqual(runs, [])

    def test_two_separate_contact_runs_both_found_and_largest_is_selectable(
        self,
    ) -> None:
        document = Document()
        document.add_paragraph("Title")

        small_run = etree.fromstring(
            _textbox_run_xml(
                ["CONTACT PERSON", "Jane Doe", "jane@example.com"],
                width_emu=500000,
                height_emu=300000,
            ).encode("utf-8")
        )
        large_run = etree.fromstring(
            _textbox_run_xml(
                ["SOME FIRM", "firm@example.com"],
                width_emu=2000000,
                height_emu=1000000,
            ).encode("utf-8")
        )
        document.paragraphs[0]._p.append(small_run)
        document.paragraphs[0]._p.append(large_run)

        buffer = io.BytesIO()
        document.save(buffer)

        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))

        self.assertEqual(len(runs), 2)
        primary = max(runs, key=lambda run: run.width_emu * run.height_emu)
        self.assertEqual(primary.width_emu, 2000000)

    def test_choice_and_fallback_style_duplicate_txbxcontent_is_one_run(
        self,
    ) -> None:
        # Mirrors a real floating shape's mc:Choice/mc:Fallback pair:
        # the SAME visible box content appears twice inside one <w:r>,
        # each occurrence separately matched by the txbxContent regex.
        # Must resolve to exactly one run, not two.
        document = Document()
        document.add_paragraph("Title")

        inner = (
            "<w:p><w:r><w:t>CONTACT PERSON</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>Jane Doe</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>jane@example.com</w:t></w:r></w:p>"
        )
        duplicated_run_xml = (
            '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<mc:AlternateContent "
            'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
            '<mc:Choice Requires="wps">'
            f"<w:txbxContent>{inner}</w:txbxContent>"
            "</mc:Choice>"
            "<mc:Fallback>"
            f"<w:txbxContent>{inner}</w:txbxContent>"
            "</mc:Fallback>"
            "</mc:AlternateContent>"
            "</w:r>"
        )
        run_element = etree.fromstring(duplicated_run_xml.encode("utf-8"))
        document.paragraphs[0]._p.append(run_element)

        buffer = io.BytesIO()
        document.save(buffer)

        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))
        self.assertEqual(
            len(runs), 1,
            "Choice+Fallback duplicate content must resolve to one run, "
            "not one real run plus a bogus nested-run duplicate",
        )


if __name__ == "__main__":
    unittest.main()
