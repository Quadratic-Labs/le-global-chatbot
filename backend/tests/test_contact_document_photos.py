"""
Tests for real, persisted-source-DOCX contact photo synchronization
(mission "COMPLETE CONTACT PHOTO CRUD + DOCX SOURCE SYNCHRONIZATION").

Mirrors test_contact_photos.py's own established convention: exercise
the REAL corpus under /data/documents/source (the container's bind
mount for the real source tree - see integration_tests/README.md),
skipping gracefully when that path is unavailable (e.g. on a host
outside the container) rather than failing - never any synthetic
hand-built DOCX, and never a mutation of the real corpus file itself:
every mutation here runs against a throwaway temp copy.
"""

from __future__ import annotations

import hashlib
import tempfile
import unittest
import zipfile
from pathlib import Path

from docx import Document

from app.services.contact_document_photos import (
    ContactDocumentPhotoError,
    add_contact_photo_to_document,
    add_new_contact_photo_to_document,
    remove_contact_photo_from_document,
    replace_contact_photo_in_document,
)
from app.services.contact_photos import extract_contact_photo_candidates


SOURCE_ROOT = Path("/data/documents/source")

_VALID_JPEG = bytes.fromhex(
    "ffd8ffe000104a46494600010100000100010000ffd9"
)
_VALID_PNG = bytes.fromhex(
    "89504e470d0a1a0a0000000d49484452000000010000"
    "000108060000001f15c4890000000a49444154789c63"
    "60000002000100ffff03000006000557bfabd4000000"
    "0049454e44ae426082"
)


class RealCorpusContactDocumentPhotoTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)

    def _require_copy(self, filename: str) -> Path:
        source = SOURCE_ROOT / filename

        if not source.exists():
            self.skipTest(
                f"Real corpus source unavailable: {source}"
            )

        original_bytes = source.read_bytes()
        copy_path = Path(self.temp.name) / filename
        copy_path.write_bytes(original_bytes)

        # The caller's own temp copy must never be the thing that
        # proves nothing was mutated - always compare back against
        # the REAL file on disk too.
        self.addCleanup(
            lambda: self.assertEqual(
                original_bytes,
                source.read_bytes(),
                f"{source} was mutated by this test.",
            )
        )

        return copy_path

    def _unrelated_parts(self, path: Path) -> dict[str, bytes]:
        """Every zip part except document.xml/rels/content-types -
        the parts a photo mutation must never touch."""

        skip = {
            "word/document.xml",
            "word/_rels/document.xml.rels",
            "[Content_Types].xml",
        }

        with zipfile.ZipFile(path) as archive:
            return {
                name: archive.read(name)
                for name in archive.namelist()
                if name not in skip
            }

    def test_belgium_replace_isolates_the_other_contacts_photo(
        self,
    ) -> None:
        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        original_candidates = extract_contact_photo_candidates(path)
        by_name = {
            c.source_filename: c for c in original_candidates
        }
        chris = by_name["image2.jpg"]
        chris_sha = chris.sha256
        nicolas_sha = by_name["image1.png"].sha256

        original_paragraphs = len(Document(str(path)).paragraphs)
        unrelated_before = self._unrelated_parts(path)
        del unrelated_before[chris.media_path]

        new_bytes = replace_contact_photo_in_document(
            path,
            target_sha256=chris_sha,
            new_data=_VALID_JPEG,
            new_content_type="image/jpeg",
        )

        out = Path(self.temp.name) / "belgium_replaced.docx"
        out.write_bytes(new_bytes)

        new_shas = {
            c.sha256 for c in extract_contact_photo_candidates(out)
        }

        self.assertIn(
            hashlib.sha256(_VALID_JPEG).hexdigest(), new_shas
        )
        self.assertNotIn(chris_sha, new_shas)
        self.assertIn(nicolas_sha, new_shas)

        self.assertEqual(
            original_paragraphs, len(Document(str(out)).paragraphs)
        )
        unrelated_after = self._unrelated_parts(out)
        del unrelated_after[chris.media_path]
        self.assertEqual(unrelated_before, unrelated_after)

    def test_belgium_remove_isolates_the_other_contacts_photo(
        self,
    ) -> None:
        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        by_name = {
            c.source_filename: c.sha256
            for c in extract_contact_photo_candidates(path)
        }
        chris_sha = by_name["image2.jpg"]
        nicolas_sha = by_name["image1.png"]

        original_paragraphs = len(Document(str(path)).paragraphs)

        new_bytes = remove_contact_photo_from_document(
            path,
            target_sha256=nicolas_sha,
        )

        out = Path(self.temp.name) / "belgium_removed.docx"
        out.write_bytes(new_bytes)

        remaining = extract_contact_photo_candidates(out)

        self.assertEqual(1, len(remaining))
        self.assertEqual(chris_sha, remaining[0].sha256)
        self.assertEqual(
            original_paragraphs, len(Document(str(out)).paragraphs)
        )

    def test_belgium_add_into_a_shared_zone_fails_closed(
        self,
    ) -> None:
        """
        Chris and Nicolas share ONE combined "CONTACT PERSON" textbox
        zone, disambiguated only by their two photos' own separate
        geometry - never by two separate zones. Adding a photo for
        Nicolas after his own is removed has no safe, deterministic
        place to insert it (the zone still names Chris too), so this
        must fail closed rather than guess.
        """

        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        by_name = {
            c.source_filename: c.sha256
            for c in extract_contact_photo_candidates(path)
        }
        nicolas_sha = by_name["image1.png"]

        removed_bytes = remove_contact_photo_from_document(
            path,
            target_sha256=nicolas_sha,
        )
        removed_path = Path(self.temp.name) / "belgium_removed.docx"
        removed_path.write_bytes(removed_bytes)

        with self.assertRaises(ContactDocumentPhotoError):
            add_contact_photo_to_document(
                removed_path,
                contact_person="Nicolas Simon",
                new_data=_VALID_JPEG,
                new_content_type="image/jpeg",
                other_contact_persons=["Chris van Olmen"],
            )

    def test_germany_add_to_a_genuinely_photo_less_contact(
        self,
    ) -> None:
        path = self._require_copy("DE.docx")

        original_candidates = extract_contact_photo_candidates(path)
        self.assertEqual(
            0,
            len(original_candidates),
            "Tobias Pusch is expected to have no existing photo in "
            "the real corpus - if this fails, the corpus changed and "
            "this test's premise needs revisiting.",
        )

        original_paragraphs = len(Document(str(path)).paragraphs)

        new_bytes = add_contact_photo_to_document(
            path,
            contact_person="Tobias Pusch",
            new_data=_VALID_PNG,
            new_content_type="image/png",
            other_contact_persons=[],
        )

        out = Path(self.temp.name) / "de_added.docx"
        out.write_bytes(new_bytes)

        candidates = extract_contact_photo_candidates(out)

        self.assertEqual(1, len(candidates))
        self.assertEqual(
            hashlib.sha256(_VALID_PNG).hexdigest(),
            candidates[0].sha256,
        )
        self.assertEqual(
            original_paragraphs, len(Document(str(out)).paragraphs)
        )

    def test_argentina_replace_then_remove_single_contact(
        self,
    ) -> None:
        path = self._require_copy("AR.docx")

        original_candidates = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(original_candidates))
        target_sha = original_candidates[0].sha256

        original_paragraphs = len(Document(str(path)).paragraphs)

        replaced_bytes = replace_contact_photo_in_document(
            path,
            target_sha256=target_sha,
            new_data=_VALID_JPEG,
            new_content_type="image/jpeg",
        )
        replaced_path = Path(self.temp.name) / "ar_replaced.docx"
        replaced_path.write_bytes(replaced_bytes)

        replaced_candidates = extract_contact_photo_candidates(
            replaced_path
        )
        self.assertEqual(1, len(replaced_candidates))
        self.assertEqual(
            hashlib.sha256(_VALID_JPEG).hexdigest(),
            replaced_candidates[0].sha256,
        )
        self.assertEqual(
            original_paragraphs,
            len(Document(str(replaced_path)).paragraphs),
        )

        removed_bytes = remove_contact_photo_from_document(
            replaced_path,
            target_sha256=hashlib.sha256(_VALID_JPEG).hexdigest(),
        )
        removed_path = Path(self.temp.name) / "ar_removed.docx"
        removed_path.write_bytes(removed_bytes)

        self.assertEqual(
            0, len(extract_contact_photo_candidates(removed_path))
        )
        self.assertEqual(
            original_paragraphs,
            len(Document(str(removed_path)).paragraphs),
        )

    def test_replace_of_an_unlocatable_sha_fails_closed(
        self,
    ) -> None:
        path = self._require_copy("AR.docx")

        with self.assertRaises(ContactDocumentPhotoError):
            replace_contact_photo_in_document(
                path,
                target_sha256="0" * 64,
                new_data=_VALID_JPEG,
                new_content_type="image/jpeg",
            )

    def test_remove_of_an_unlocatable_sha_fails_closed(
        self,
    ) -> None:
        path = self._require_copy("AR.docx")

        with self.assertRaises(ContactDocumentPhotoError):
            remove_contact_photo_from_document(
                path,
                target_sha256="0" * 64,
            )

    def test_add_for_an_unknown_contact_person_fails_closed(
        self,
    ) -> None:
        path = self._require_copy("DE.docx")

        with self.assertRaises(ContactDocumentPhotoError):
            add_contact_photo_to_document(
                path,
                contact_person="Nobody Real",
                new_data=_VALID_PNG,
                new_content_type="image/png",
                other_contact_persons=[],
            )

    def test_replace_of_the_existing_zone_now_uses_true_geometry(
        self,
    ) -> None:
        """
        Regression guard for the exact bug the "FINAL BLOCKER" mission
        fixed: adding a photo for an EXISTING named zone must
        genuinely geometrically overlap that zone (reason "GEOMETRY"),
        never merely happen to be the sole remaining portrait in the
        whole document (reason "UNIQUE_PORTRAIT") - the fallback only
        ever works by coincidence and breaks the moment the document
        has more than one photo total, exactly the scenario a
        brand-new second contact's photo introduces.
        """

        path = self._require_copy("DE.docx")

        new_bytes = add_contact_photo_to_document(
            path,
            contact_person="Tobias Pusch",
            new_data=_VALID_PNG,
            new_content_type="image/png",
            other_contact_persons=[],
        )

        out = Path(self.temp.name) / "de_geometry.docx"
        out.write_bytes(new_bytes)

        candidates = extract_contact_photo_candidates(out)
        self.assertEqual(1, len(candidates))
        self.assertEqual("GEOMETRY", candidates[0].reason)

    def test_add_new_contact_photo_anchors_to_the_largest_zone_alongside_an_existing_photo(
        self,
    ) -> None:
        """
        The mission's own core blocker scenario: a document that
        ALREADY has one contact's photo, and a genuinely brand-new
        second contact (whose name cannot possibly appear anywhere
        yet) also gets a photo. Both must round-trip as independently
        valid, GEOMETRY-reasoned candidates - never relying on the
        "exactly one remaining portrait" fallback, which cannot
        disambiguate once there is more than one unassociated photo.
        """

        path = self._require_copy("AR.docx")

        original = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(original))
        original_sha = original[0].sha256

        original_paragraphs = len(Document(str(path)).paragraphs)

        new_bytes = add_new_contact_photo_to_document(
            path,
            new_data=_VALID_JPEG,
            new_content_type="image/jpeg",
        )

        out = Path(self.temp.name) / "ar_new_contact.docx"
        out.write_bytes(new_bytes)

        candidates = extract_contact_photo_candidates(out)
        self.assertEqual(2, len(candidates))

        for candidate in candidates:
            self.assertEqual(
                "GEOMETRY",
                candidate.reason,
                "must be a real geometric match, never a fallback "
                "that only works by coincidence",
            )

        new_shas = {c.sha256 for c in candidates}
        self.assertIn(original_sha, new_shas)
        self.assertIn(
            hashlib.sha256(_VALID_JPEG).hexdigest(), new_shas
        )
        self.assertEqual(
            original_paragraphs, len(Document(str(out)).paragraphs)
        )

    def test_add_new_contact_photo_into_belgiums_shared_zone_isolates_existing_photos(
        self,
    ) -> None:
        """
        Belgium's shared "CONTACT PERSON" zone (Chris + Nicolas) is
        the document's only (and therefore largest) CONTACT PERSON
        zone - a brand-new third contact's photo anchors there too,
        and must never disturb either of the two existing contacts'
        own photos.
        """

        path = self._require_copy(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        original_shas = {
            c.sha256 for c in extract_contact_photo_candidates(path)
        }
        self.assertEqual(2, len(original_shas))

        new_bytes = add_new_contact_photo_to_document(
            path,
            new_data=_VALID_PNG,
            new_content_type="image/png",
        )

        out = Path(self.temp.name) / "belgium_new_contact.docx"
        out.write_bytes(new_bytes)

        candidates = extract_contact_photo_candidates(out)
        new_shas = {c.sha256 for c in candidates}

        self.assertEqual(3, len(new_shas))
        self.assertTrue(original_shas.issubset(new_shas))
        self.assertIn(
            hashlib.sha256(_VALID_PNG).hexdigest(), new_shas
        )

    def test_add_new_contact_photo_fails_closed_with_no_contact_zone_at_all(
        self,
    ) -> None:
        """
        A document with genuinely zero "CONTACT PERSON" zones (no
        structural home at all) must fail closed rather than insert
        the photo at an arbitrary position such as the document's end
        (mission section 6).
        """

        path = self._require_copy("PT.docx")

        self.assertEqual(0, len(extract_contact_photo_candidates(path)))

        with self.assertRaises(ContactDocumentPhotoError):
            add_new_contact_photo_to_document(
                path,
                new_data=_VALID_JPEG,
                new_content_type="image/jpeg",
            )


if __name__ == "__main__":
    unittest.main()
