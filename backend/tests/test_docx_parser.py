import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.core.legal_taxonomy import get_canonical_legal_topic
from app.core.subsection_taxonomy import get_subsection_topic_override
from app.services.docx_parser import (
    ExtractedContact,
    build_contact_chunk_content,
    extract_contacts_from_docx,
    find_plain_paragraph_contact_block_bounds,
    parse_contact_blocks,
    parse_docx_sections,
    split_combined_legacy_contact,
    _classify_canonical_firm_lines,
)


def _mark_as_numbered(
    paragraph: Paragraph,
) -> None:
    """Mark a paragraph as a numbered-list item in Word XML."""

    paragraph_properties = (
        paragraph._p.get_or_add_pPr()
    )

    numbering_properties = OxmlElement(
        "w:numPr"
    )

    indentation_level = OxmlElement(
        "w:ilvl"
    )
    indentation_level.set(
        qn("w:val"),
        "0",
    )

    numbering_id = OxmlElement(
        "w:numId"
    )
    numbering_id.set(
        qn("w:val"),
        "1",
    )

    numbering_properties.append(
        indentation_level
    )
    numbering_properties.append(
        numbering_id
    )

    paragraph_properties.append(
        numbering_properties
    )


def _mark_as_explicitly_unbolded(
    paragraph: Paragraph,
) -> None:
    """Apply the malformed formatting found in some L&E DOCX files."""

    paragraph_properties = (
        paragraph._p.get_or_add_pPr()
    )

    run_properties = paragraph_properties.find(
        qn("w:rPr")
    )

    if run_properties is None:
        run_properties = OxmlElement(
            "w:rPr"
        )
        paragraph_properties.append(
            run_properties
        )

    bold_property = run_properties.find(
        qn("w:b")
    )

    if bold_property is None:
        bold_property = OxmlElement(
            "w:b"
        )
        run_properties.append(
            bold_property
        )

    bold_property.set(
        qn("w:val"),
        "0",
    )


class DocxParserTests(unittest.TestCase):
    def test_groups_content_by_headings(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "sample-legal-document.docx"
            )

            document = Document()

            document.add_paragraph(
                "Introductory legal information."
            )

            document.add_heading(
                "Employment contracts",
                level=1,
            )

            document.add_paragraph(
                "General information about employment contracts."
            )

            document.add_heading(
                "Probationary period",
                level=2,
            )

            document.add_paragraph(
                "The probationary period depends on local law."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path
            )

            self.assertEqual(
                len(sections),
                3,
            )

            self.assertEqual(
                sections[0].section,
                "General",
            )
            self.assertIsNone(
                sections[0].subsection
            )
            self.assertEqual(
                sections[0].content,
                "Introductory legal information.",
            )

            self.assertEqual(
                sections[1].section,
                "Employment contracts",
            )
            self.assertIsNone(
                sections[1].subsection
            )
            self.assertEqual(
                sections[1].content,
                (
                    "General information about "
                    "employment contracts."
                ),
            )

            self.assertEqual(
                sections[2].section,
                "Employment contracts",
            )
            self.assertEqual(
                sections[2].subsection,
                "Probationary period",
            )
            self.assertEqual(
                sections[2].content,
                (
                    "The probationary period "
                    "depends on local law."
                ),
            )

    def test_preserves_tables_in_document_order(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "legal-table-document.docx"
            )

            document = Document()

            document.add_heading(
                "Termination",
                level=1,
            )

            document.add_paragraph(
                "Rules applicable before the table."
            )

            table = document.add_table(
                rows=3,
                cols=3,
            )

            table.cell(
                0,
                0,
            ).text = "Country"

            table.cell(
                0,
                1,
            ).text = "Notice period"

            table.cell(
                0,
                2,
            ).text = "Written notice"

            table.cell(
                1,
                0,
            ).text = "France"

            table.cell(
                1,
                1,
            ).text = "30 days"

            table.cell(
                1,
                2,
            ).text = "Required"

            table.cell(
                2,
                0,
            ).text = "Belgium"

            table.cell(
                2,
                1,
            ).text = "45 days"

            table.cell(
                2,
                2,
            ).text = "Required"

            document.add_paragraph(
                "Additional rules after the table."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path
            )

            self.assertEqual(
                len(sections),
                1,
            )

            self.assertEqual(
                sections[0].content,
                (
                    "Rules applicable before the table.\n\n"
                    "Country | Notice period | Written notice\n"
                    "France | 30 days | Required\n"
                    "Belgium | 45 days | Required\n\n"
                    "Additional rules after the table."
                ),
            )

    def test_ignores_decorative_table_content(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "decorative-table-document.docx"
            )

            document = Document()

            document.add_heading(
                "Working Conditions",
                level=1,
            )

            document.add_table(
                rows=1,
                cols=2,
            )

            document.add_heading(
                "Salary",
                level=2,
            )

            document.add_paragraph(
                "Employees must receive their agreed salary."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path
            )

            self.assertEqual(
                len(sections),
                1,
            )

            self.assertEqual(
                sections[0].subsection,
                "Salary",
            )

            self.assertEqual(
                sections[0].content,
                "Employees must receive their agreed salary.",
            )

    def test_keeps_numbered_heading_as_content(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "numbered-heading-document.docx"
            )

            document = Document()

            document.add_heading(
                "Termination",
                level=1,
            )

            document.add_heading(
                "Remedies",
                level=2,
            )

            document.add_paragraph(
                "The employee may challenge the dismissal."
            )

            numbered_paragraph = document.add_paragraph(
                "Pregnant employees are protected."
            )
            numbered_paragraph.style = "Heading 2"

            _mark_as_numbered(
                numbered_paragraph
            )

            document.add_heading(
                "Whistleblower Laws",
                level=2,
            )

            document.add_paragraph(
                "Whistleblowers receive legal protection."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path
            )

            self.assertEqual(
                len(sections),
                2,
            )

            self.assertEqual(
                sections[0].subsection,
                "Remedies",
            )

            self.assertIn(
                "Pregnant employees are protected.",
                sections[0].content,
            )

            self.assertEqual(
                sections[1].subsection,
                "Whistleblower Laws",
            )

    def test_keeps_unbolded_headings_as_content(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "unbolded-heading-document.docx"
            )

            document = Document()

            document.add_heading(
                "07. Termination of Employment Contracts",
                level=1,
            )

            document.add_heading(
                "Remedies for Wrongful Termination",
                level=2,
            )

            document.add_paragraph(
                "The employee may challenge the dismissal."
            )

            false_heading = document.add_paragraph(
                (
                    "Objectively null and void termination: "
                    "protected situations include:"
                )
            )
            false_heading.style = "Heading 2"

            _mark_as_explicitly_unbolded(
                false_heading
            )

            numbered_item = document.add_paragraph(
                "Pregnant employees are protected."
            )
            numbered_item.style = "Heading 2"

            _mark_as_numbered(
                numbered_item
            )
            _mark_as_explicitly_unbolded(
                numbered_item
            )

            final_paragraph = document.add_paragraph(
                (
                    "This protection means that the dismissal "
                    "may be declared null and void."
                )
            )
            final_paragraph.style = "Heading 2"

            _mark_as_explicitly_unbolded(
                final_paragraph
            )

            document.add_heading(
                "Whistleblower Laws",
                level=2,
            )

            document.add_paragraph(
                "Whistleblowers receive legal protection."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path
            )

            self.assertEqual(
                len(sections),
                2,
            )

            self.assertEqual(
                sections[0].subsection,
                "Remedies for Wrongful Termination",
            )

            self.assertIn(
                "Objectively null and void termination:",
                sections[0].content,
            )

            self.assertIn(
                "Pregnant employees are protected.",
                sections[0].content,
            )

            self.assertIn(
                "This protection means that the dismissal",
                sections[0].content,
            )

            self.assertEqual(
                sections[1].subsection,
                "Whistleblower Laws",
            )

    def test_accepts_heading_ending_with_colon(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "heading-with-colon-document.docx"
            )

            document = Document()

            document.add_heading(
                "Employment Contracts",
                level=1,
            )

            document.add_heading(
                "Notice periods:",
                level=2,
            )

            document.add_paragraph(
                "The statutory notice period is fifteen days."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path
            )

            self.assertEqual(
                len(sections),
                1,
            )

            self.assertEqual(
                sections[0].subsection,
                "Notice periods:",
            )

            self.assertEqual(
                sections[0].content,
                "The statutory notice period is fifteen days.",
            )

    def test_detects_legacy_numbered_bold_topic(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "legacy-document.docx"
            )

            document = Document()

            document.add_paragraph(
                "Introduction content."
            )

            topic = document.add_paragraph(
                "Hiring practices in Greece"
            )
            topic.style = "List Paragraph"
            topic.runs[0].bold = True

            _mark_as_numbered(
                topic
            )

            document.add_paragraph(
                "Foreign employees require permission."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Greece",
            )

            self.assertEqual(
                len(sections),
                2,
            )

            self.assertEqual(
                sections[0].section,
                "Employment Law Overview Greece",
            )

            self.assertEqual(
                sections[0].content,
                "Introduction content.",
            )

            self.assertEqual(
                sections[1].section,
                "Hiring practices in Greece",
            )

            self.assertEqual(
                sections[1].content,
                "Foreign employees require permission.",
            )

    def test_detects_hybrid_bold_numbered_topic(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "hybrid-document.docx"
            )

            document = Document()

            topic = document.add_paragraph(
                "02. Employment Contracts"
            )
            topic.runs[0].bold = True

            document.add_paragraph(
                "Employment contracts may be open-ended."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Italy",
            )

            self.assertEqual(
                len(sections),
                1,
            )

            self.assertEqual(
                sections[0].section,
                "02. Employment Contracts",
            )

            self.assertEqual(
                sections[0].content,
                "Employment contracts may be open-ended.",
            )

    def test_keeps_unrecognized_heading_one_as_content(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "australia-heading-document.docx"
            )

            document = Document()

            document.add_heading(
                "07. Termination of Employment Contracts",
                level=1,
            )

            document.add_paragraph(
                "Termination content before the malformed heading."
            )

            false_heading = document.add_paragraph(
                "Whistleblowers currently have legal protections."
            )
            false_heading.style = "Heading 1"

            document.add_paragraph(
                "Additional whistleblower content."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Australia",
            )

            self.assertEqual(
                len(sections),
                1,
            )

            self.assertEqual(
                sections[0].section,
                "07. Termination of Employment Contracts",
            )

            self.assertIn(
                (
                    "Termination content before "
                    "the malformed heading."
                ),
                sections[0].content,
            )

            self.assertIn(
                "Whistleblowers currently have legal protections.",
                sections[0].content,
            )

            self.assertIn(
                "Additional whistleblower content.",
                sections[0].content,
            )

    def test_keeps_heading_four_as_content(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "heading-four-document.docx"
            )

            document = Document()

            document.add_heading(
                "09. Transfer of Undertakings",
                level=1,
            )

            body_paragraph = document.add_paragraph(
                (
                    "The transfer does not automatically "
                    "terminate employment."
                )
            )
            body_paragraph.style = "Heading 4"

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Brazil",
            )

            self.assertEqual(
                len(sections),
                1,
            )

            self.assertEqual(
                sections[0].section,
                "09. Transfer of Undertakings",
            )

            self.assertEqual(
                sections[0].content,
                (
                    "The transfer does not automatically "
                    "terminate employment."
                ),
            )

    def test_splits_working_conditions_subsections(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "working-conditions-document.docx"
            )

            document = Document()

            document.add_heading(
                "03. Working Conditions",
                level=1,
            )

            for title, content in (
                (
                    "Overtime",
                    "Overtime legal content.",
                ),
                (
                    "Work Hours Record",
                    "Working time recording content.",
                ),
                (
                    "Paid Leave",
                    "Paid leave legal content.",
                ),
            ):
                paragraph = document.add_paragraph(
                    title
                )
                paragraph.runs[0].bold = True

                document.add_paragraph(
                    content
                )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Spain",
            )

            self.assertEqual(
                [
                    section.subsection
                    for section in sections
                ],
                [
                    "Overtime",
                    "Work Hours Record",
                    "Paid Leave",
                ],
            )

            self.assertEqual(
                sections[0].content,
                "Overtime legal content.",
            )

    def test_notice_of_termination_and_redundancy_pay_starts_new_section(
        self,
    ) -> None:
        """
        Regression test for a real L&E document defect.

        In the Australia source document, a bold "Notice of Termination
        and Redundancy Pay" paragraph appears inside the "Working
        Conditions" section (after "Overtime"), using the same plain
        bold-"Normal"-style formatting as ordinary emphasis elsewhere in
        the document. Before the fix, this heading was not recognized
        by any taxonomy rule, so its content silently stayed folded
        into the preceding Overtime subsection under the wrong legal
        topic - making it unreachable for questions about redundancy
        pay, which are filtered by legal topic.
        """

        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "working-conditions-termination-document.docx"
            )

            document = Document()

            document.add_heading(
                "03. Working Conditions",
                level=1,
            )

            overtime_heading = document.add_paragraph(
                "Overtime"
            )
            overtime_heading.runs[0].bold = True

            document.add_paragraph(
                "Overtime is paid at a premium rate."
            )

            redundancy_heading = document.add_paragraph(
                "Notice of Termination and Redundancy Pay"
            )
            redundancy_heading.runs[0].bold = True

            document.add_paragraph(
                "An employee is entitled to redundancy pay "
                "calculated using the employee's base rate of pay."
            )

            paid_leave_heading = document.add_paragraph(
                "Paid Leave"
            )
            paid_leave_heading.runs[0].bold = True

            document.add_paragraph(
                "Employees accrue paid leave over each year "
                "of service."
            )

            document.save(
                file_path
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Australia",
            )

            overtime_section = next(
                section
                for section in sections
                if section.subsection == "Overtime"
            )

            self.assertNotIn(
                "redundancy pay",
                overtime_section.content.casefold(),
            )

            self.assertIn(
                "premium rate",
                overtime_section.content,
            )

            redundancy_section = next(
                (
                    section
                    for section in sections
                    if "redundancy" in section.content.casefold()
                ),
                None,
            )

            self.assertIsNotNone(
                redundancy_section
            )

            self.assertEqual(
                redundancy_section.section,
                "Notice of Termination and Redundancy Pay",
            )

            self.assertIsNone(
                redundancy_section.subsection
            )

            # The section text alone is not a canonical topic heading
            # (structure cannot tell it apart from ordinary bold
            # emphasis) - the override table is what supplies the
            # correct topic, exactly as document_chunk_builder does.
            self.assertIsNone(
                get_canonical_legal_topic(
                    section=redundancy_section.section,
                    country="Australia",
                )
            )

            self.assertEqual(
                get_subsection_topic_override(
                    redundancy_section.section
                ),
                "Termination of Employment Contracts",
            )

            self.assertIn(
                "base rate of pay",
                redundancy_section.content,
            )

            self.assertLess(
                sections.index(
                    overtime_section
                ),
                sections.index(
                    redundancy_section
                ),
            )

            # Regression guard: a legitimate Working Conditions
            # subsection following the override must not be swallowed
            # into the override's topic - the parser must revert to
            # the enclosing section/topic as soon as this next real
            # subsection is recognized.
            paid_leave_section = next(
                section
                for section in sections
                if section.subsection == "Paid Leave"
            )

            self.assertNotIn(
                "redundancy",
                paid_leave_section.content.casefold(),
            )

            self.assertIn(
                "accrue paid leave",
                paid_leave_section.content,
            )

            self.assertEqual(
                get_canonical_legal_topic(
                    section=paid_leave_section.section,
                    country="Australia",
                ),
                "Working Conditions",
            )

            self.assertLess(
                sections.index(
                    redundancy_section
                ),
                sections.index(
                    paid_leave_section
                ),
            )


class CustomTopicRecognitionTests(unittest.TestCase):
    """
    ORDER 8A, section 12 - the parser must recognize a brand-new,
    non-taxonomy top-level legal topic an admin adds, using only the
    document's own real structure, generically (no topic name ever
    hardcoded) - while never promoting front matter, an ordinary
    subsection, or a stray list item that happens to share one weak
    structural signal with the document's own real topics.
    """

    def _build(
        self,
        directory: Path,
        blocks: list[tuple[str, str, int | None]],
    ) -> Path:
        """
        blocks: (text, kind, heading_level) where kind is "heading" or
        "paragraph"; heading_level is the Word heading style level (or
        None for a plain paragraph).
        """

        file_path = directory / "sample.docx"
        document = Document()

        for text, kind, heading_level in blocks:
            if kind == "heading":
                document.add_heading(text, level=heading_level)
            else:
                document.add_paragraph(text)

        document.save(file_path)
        return file_path

    def test_custom_topic_recognized_after_a_confirmed_topic(
        self,
    ) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = self._build(
                Path(temp_dir),
                [
                    (
                        "Employment Law Overview United Kingdom",
                        "heading",
                        1,
                    ),
                    ("Hiring Practices", "heading", 1),
                    ("Real hiring content.", "paragraph", None),
                    ("Remote Working", "heading", 1),
                    (
                        "Employees may work remotely. MARKER.",
                        "paragraph",
                        None,
                    ),
                ],
            )

            sections = parse_docx_sections(
                file_path,
                country="United Kingdom",
            )

            custom = [
                section
                for section in sections
                if section.is_custom_legal_topic
            ]

            self.assertEqual(len(custom), 1)
            self.assertEqual(custom[0].section, "Remote Working")
            self.assertIn("MARKER", custom[0].content)

    def test_custom_topic_not_recognized_before_any_confirmed_topic(
        self,
    ) -> None:
        # A front-matter/introductory Heading 1 before the document's
        # own overview or first real topic must never become a fake
        # legal topic.
        with TemporaryDirectory() as temp_dir:
            file_path = self._build(
                Path(temp_dir),
                [
                    ("Introduction", "heading", 1),
                    (
                        "Some introductory front matter text.",
                        "paragraph",
                        None,
                    ),
                    (
                        "Employment Law Overview United Kingdom",
                        "heading",
                        1,
                    ),
                    ("Hiring Practices", "heading", 1),
                    ("Real hiring content.", "paragraph", None),
                ],
            )

            sections = parse_docx_sections(
                file_path,
                country="United Kingdom",
            )

            self.assertFalse(
                any(
                    section.is_custom_legal_topic
                    for section in sections
                )
            )
            self.assertNotIn(
                "Introduction",
                [section.section for section in sections],
            )

    def test_custom_topic_unsupported_when_real_topics_use_bold_only(
        self,
    ) -> None:
        # Real topics recognized only via explicit bold (no Heading 1
        # style, no numbering) give no reliable, unambiguous anchor to
        # hold a candidate heading to - ordinary bold subsections are
        # too common for that signal alone to be safe. Custom topics
        # are simply unsupported for this document shape.
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "sample.docx"
            document = Document()

            document.add_heading(
                "Employment Law Overview United Kingdom", level=1
            )

            bold_heading = document.add_paragraph()
            bold_heading.add_run("Hiring Practices").bold = True
            document.add_paragraph("Real hiring content.")

            bold_custom = document.add_paragraph()
            bold_custom.add_run("Remote Working").bold = True
            document.add_paragraph("Remote work content.")

            document.save(file_path)

            sections = parse_docx_sections(
                file_path,
                country="United Kingdom",
            )

            self.assertFalse(
                any(
                    section.is_custom_legal_topic
                    for section in sections
                )
            )

    def test_custom_topic_requires_the_same_signal_as_confirmed_topics(
        self,
    ) -> None:
        # Real topics here all carry BOTH Heading 1 AND list numbering
        # - a heading_level-1-only candidate (no numbering) must not
        # be promoted, even though it comes after a confirmed topic.
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "sample.docx"
            document = Document()

            document.add_heading(
                "Employment Law Overview United Kingdom", level=1
            )

            hiring_heading = document.add_heading(
                "Hiring Practices", level=1
            )
            _mark_as_numbered(hiring_heading)
            document.add_paragraph("Real hiring content.")

            document.add_heading(
                "Key Points", level=1
            )
            document.add_paragraph(
                "This is an ordinary sub-heading, not numbered like "
                "the real topics."
            )

            document.save(file_path)

            sections = parse_docx_sections(
                file_path,
                country="United Kingdom",
            )

            self.assertFalse(
                any(
                    section.is_custom_legal_topic
                    for section in sections
                )
            )
            self.assertNotIn(
                "Key Points",
                [section.section for section in sections],
            )

    def test_custom_topic_rejects_sentence_shaped_candidates(
        self,
    ) -> None:
        # A numbered list item that happens to share the document's
        # own (Heading 1 + numbering) signal, but reads like a clause
        # in an enumerated sentence (lowercase start, trailing
        # semicolon) rather than a title, must never be promoted.
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "sample.docx"
            document = Document()

            document.add_heading(
                "Employment Law Overview United Kingdom", level=1
            )

            hiring_heading = document.add_heading(
                "Hiring Practices", level=1
            )
            _mark_as_numbered(hiring_heading)
            document.add_paragraph("Real hiring content.")

            list_item_heading = document.add_heading(
                "the Corporations Act;", level=1
            )
            _mark_as_numbered(list_item_heading)
            document.add_paragraph("More content.")

            document.save(file_path)

            sections = parse_docx_sections(
                file_path,
                country="United Kingdom",
            )

            self.assertFalse(
                any(
                    section.is_custom_legal_topic
                    for section in sections
                )
            )

    def test_generic_mode_never_recognizes_custom_legal_topics(
        self,
    ) -> None:
        # Custom-topic recognition is a strict L&E (country-supplied)
        # concept only - generic mode's own Heading 1 handling is
        # unrelated and must never mark anything is_custom_legal_topic.
        with TemporaryDirectory() as temp_dir:
            file_path = self._build(
                Path(temp_dir),
                [
                    ("Some Section", "heading", 1),
                    ("Some content.", "paragraph", None),
                    ("Another Section", "heading", 1),
                    ("More content.", "paragraph", None),
                ],
            )

            sections = parse_docx_sections(file_path)

            self.assertFalse(
                any(
                    section.is_custom_legal_topic
                    for section in sections
                )
            )


class AdminSectionMarkerRecognitionTests(unittest.TestCase):
    """
    ORDER 8A-C - the dedicated ADMIN-section marker style
    (ADMIN_SECTION_STYLE_NAME) is the sole, deterministic identity a
    reparse relies on for an admin-added top-level topic: it must work
    regardless of the surrounding document's own native convention
    (Heading 1, bold-only, or anything else), regardless of country,
    and regardless of document position - unlike the older, structural-
    signal-based custom-topic heuristic this supersedes for anything
    the ADMIN Add feature itself creates.
    """

    def _add_marker_heading(
        self,
        document: Document,
        title: str,
    ) -> None:
        from docx.enum.style import WD_STYLE_TYPE

        from app.services.docx_parser import ADMIN_SECTION_STYLE_NAME

        try:
            style = document.styles[ADMIN_SECTION_STYLE_NAME]
        except KeyError:
            style = document.styles.add_style(
                ADMIN_SECTION_STYLE_NAME,
                WD_STYLE_TYPE.PARAGRAPH,
            )
            style.font.bold = True

        document.add_paragraph(title, style=style.name)

    def test_marker_recognized_in_bold_only_document(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "sample.docx"
            document = Document()

            document.add_paragraph(
                "Employment Law Overview United Kingdom"
            )

            hiring = document.add_paragraph()
            hiring.add_run("Hiring Practices").bold = True
            document.add_paragraph("Real hiring content.")

            self._add_marker_heading(document, "Remote Working")
            document.add_paragraph("Remote work content. MARKER.")

            document.save(file_path)

            sections = parse_docx_sections(
                file_path, country="United Kingdom"
            )

            custom = [s for s in sections if s.is_custom_legal_topic]
            self.assertEqual(len(custom), 1)
            self.assertEqual(custom[0].section, "Remote Working")
            self.assertIn("MARKER", custom[0].content)

            # native topic untouched
            hiring_sections = [
                s for s in sections if s.section == "Hiring Practices"
            ]
            self.assertEqual(len(hiring_sections), 1)
            self.assertEqual(
                hiring_sections[0].content, "Real hiring content."
            )

    def test_marker_recognized_with_no_native_topics_at_all(
        self,
    ) -> None:
        # The marker never depends on any structural-signal-learning
        # from the surrounding document - it must fire even when there
        # is nothing else to learn a signal from.
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "sample.docx"
            document = Document()

            document.add_paragraph("Just some plain front matter.")
            self._add_marker_heading(document, "Custom Only Section")
            document.add_paragraph("Custom-only content.")

            document.save(file_path)

            sections = parse_docx_sections(
                file_path, country="United Kingdom"
            )

            custom = [s for s in sections if s.is_custom_legal_topic]
            self.assertEqual(len(custom), 1)
            self.assertEqual(custom[0].section, "Custom Only Section")

    def test_marker_style_alone_never_appears_on_native_content(
        self,
    ) -> None:
        # Sanity check on the corpus-safety assumption the marker
        # design depends on: an ordinary document (no marker style
        # ever applied) never has any paragraph whose is_custom_legal_
        # topic could be confused with the marker mechanism.
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / "sample.docx"
            document = Document()

            document.add_heading(
                "Employment Law Overview United Kingdom", level=1
            )
            document.add_heading("Hiring Practices", level=1)
            document.add_paragraph("Real hiring content.")

            document.save(file_path)

            sections = parse_docx_sections(
                file_path, country="United Kingdom"
            )

            self.assertFalse(
                any(s.is_custom_legal_topic for s in sections)
            )


class ContactBlockParsingTests(unittest.TestCase):
    """
    Tests for parse_contact_blocks() against a synthetic paragraph
    structure - the shape _extract_text_box_blocks() would have
    produced from a real DOCX, without needing one.
    """

    def test_firm_then_contact_person_are_paired(
        self,
    ) -> None:
        blocks = [
            [
                "Example & Partners Advogados",
                "Freedonia",
                "1 Example Street, 6th floor, 00000 Sample City",
                "+00 000 000 00",
                "www.example-partners.test",
            ],
            [
                "CONTACT PERSON",
                "Alex Example",
                "alex@example-partners.test",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks,
            country="Freedonia",
        )

        self.assertEqual(
            len(contacts),
            1,
        )

        contact = contacts[0]

        self.assertEqual(
            contact.member_firm,
            "Example & Partners Advogados",
        )
        self.assertEqual(
            contact.contact_person,
            "Alex Example",
        )
        self.assertEqual(
            contact.email,
            "alex@example-partners.test",
        )
        self.assertEqual(
            contact.phone,
            "+00 000 000 00",
        )
        self.assertEqual(
            contact.website,
            "www.example-partners.test",
        )

        # The country line is excluded from the address since it must
        # come from validated document metadata instead.
        self.assertNotIn(
            "Freedonia",
            contact.address or "",
        )

    def test_contact_person_before_firm_block_are_still_paired(
        self,
    ) -> None:
        blocks = [
            [
                "CONTACT PERSON",
                "Nicolás Grandi",
                "ngrandi@allende.com",
            ],
            [
                "Allende & Brea",
                "Argentina",
                "Torre IRSA, Maipú 1300",
                "+54 114 318 9984",
                "www.allendebrea.com",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks,
            country="Argentina",
        )

        self.assertEqual(
            len(contacts),
            1,
        )

        contact = contacts[0]

        self.assertEqual(
            contact.member_firm,
            "Allende & Brea",
        )
        self.assertEqual(
            contact.contact_person,
            "Nicolás Grandi",
        )
        self.assertEqual(
            contact.email,
            "ngrandi@allende.com",
        )

    def test_bare_website_alone_is_not_a_firm_block(
        self,
    ) -> None:
        blocks = [
            [
                "www.leglobal.law",
            ],
            [
                "Some Firm",
                "Country",
                "123 Main Street",
                "+1 555 000 0000",
                "www.somefirm.example",
            ],
            [
                "CONTACT PERSON",
                "Jane Doe",
                "jane@somefirm.example",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks
        )

        self.assertEqual(
            len(contacts),
            1,
        )

        self.assertEqual(
            contacts[0].member_firm,
            "Some Firm",
        )

    def test_plural_contact_persons_marker_with_multiple_emails(
        self,
    ) -> None:
        blocks = [
            [
                "Van Olmen & Wynant",
                "Belgium",
                "Avenue Louise 221, 1050 Brussels",
                "+32 264 405 11",
                "www.vow.be",
            ],
            [
                "CONTACT PERSONS",
                "Chris van Olmen and Nicolas Simon",
                "chris.van.olmen@vow.be",
                "nicolas.simon@vow.be",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks,
            country="Belgium",
        )

        self.assertEqual(
            len(contacts),
            1,
        )

        contact = contacts[0]

        self.assertEqual(
            contact.contact_person,
            "Chris van Olmen and Nicolas Simon",
        )

        self.assertIn(
            "chris.van.olmen@vow.be",
            contact.email or "",
        )
        self.assertIn(
            "nicolas.simon@vow.be",
            contact.email or "",
        )

    def test_postal_code_before_real_phone_does_not_win(
        self,
    ) -> None:
        blocks = [
            [
                "Atsumi & Sakai",
                "Japan",
                (
                    "Fukoku Seimei Bldg., Reception: 16 F, "
                    "2-2-2 Uchisaiwaicho, Chiyoda-ku, "
                    "100-0011 Tokyo, +81 355 012 111"
                ),
                "www.aplaw.jp/en/",
            ],
            [
                "CONTACT PERSON",
                "Tatsuo Yamashima",
                "tatsuo.yamashima@aplaw.jp",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks,
            country="Japan",
        )

        self.assertEqual(
            len(contacts),
            1,
        )

        contact = contacts[0]

        self.assertEqual(
            contact.phone,
            "+81 355 012 111",
        )

        self.assertIn(
            "100-0011 Tokyo",
            contact.address or "",
        )

        self.assertNotIn(
            "+81 355 012 111",
            contact.address or "",
        )

        self.assertEqual(
            contact.member_firm,
            "Atsumi & Sakai",
        )

        self.assertEqual(
            contact.contact_person,
            "Tatsuo Yamashima",
        )


    def test_postal_code_line_before_local_phone_prefers_phone(
        self,
    ) -> None:
        blocks = [
            [
                "Example Firm",
                "100-0011 Tokyo",
                "03 5501 2111",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks
        )

        self.assertEqual(
            len(contacts),
            1,
        )

        self.assertEqual(
            contacts[0].phone,
            "03 5501 2111",
        )

        self.assertIn(
            "100-0011 Tokyo",
            contacts[0].address or "",
        )

        self.assertNotIn(
            "03 5501 2111",
            contacts[0].address or "",
        )


    def test_multiple_documents_worth_of_contacts_are_all_kept(
        self,
    ) -> None:
        blocks = [
            [
                "Firm One",
                "123 Street",
                "+1 555 111 1111",
            ],
            [
                "CONTACT PERSON",
                "Person One",
                "one@example.com",
            ],
            [
                "Firm Two",
                "456 Avenue",
                "+1 555 222 2222",
            ],
            [
                "CONTACT PERSON",
                "Person Two",
                "two@example.com",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks
        )

        self.assertEqual(
            len(contacts),
            2,
        )

        self.assertEqual(
            [
                contact.email
                for contact in contacts
            ],
            [
                "one@example.com",
                "two@example.com",
            ],
        )

    def test_no_text_box_blocks_returns_no_contacts(
        self,
    ) -> None:
        self.assertEqual(
            parse_contact_blocks(
                []
            ),
            [],
        )

    def test_unmatched_firm_block_reports_only_its_own_fields(
        self,
    ) -> None:
        blocks = [
            [
                "Only Firm",
                "42 Road",
                "+1 555 333 3333",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks
        )

        self.assertEqual(
            len(contacts),
            1,
        )

        self.assertEqual(
            contacts[0].member_firm,
            "Only Firm",
        )
        self.assertIsNone(
            contacts[0].contact_person
        )
        self.assertIsNone(
            contacts[0].email
        )

    def test_build_contact_chunk_content_omits_missing_fields(
        self,
    ) -> None:
        blocks = [
            [
                "CONTACT PERSON",
                "Jane Doe",
                "jane@example.com",
            ],
        ]

        contacts = parse_contact_blocks(
            blocks
        )

        content = build_contact_chunk_content(
            contacts
        )

        self.assertIn(
            "Contact person: Jane Doe",
            content,
        )
        self.assertIn(
            "Email: jane@example.com",
            content,
        )
        self.assertNotIn(
            "Member firm",
            content,
        )
        self.assertNotIn(
            "Phone",
            content,
        )
        self.assertNotIn(
            "Address",
            content,
        )
        self.assertNotIn(
            "Website",
            content,
        )




class PlainParagraphContactFallbackTests(unittest.TestCase):
    """Contract for contacts stored in ordinary DOCX paragraphs."""

    def test_france_plain_paragraph_contact_is_extracted(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "france-body-contact.docx"
            )

            document = Document()

            document.add_paragraph("FRANCE")
            document.add_paragraph(
                "EMPLOYMENT LAW OVERVIEWS 2025 - 2026"
            )
            document.add_paragraph(
                "FLICHY GRANGÉ AVOCATS"
            )

            document.add_heading(
                "I. GENERAL OVERVIEW",
                level=1,
            )
            document.add_paragraph(
                "Representative employment-law content."
            )

            document.add_paragraph(
                "Caroline Scherrmann and Florence Bacquet"
            )
            document.add_paragraph(
                "Partners, Flichy Grangé Avocats"
            )
            document.add_paragraph(
                "scherrmann@flichy.com"
            )
            document.add_paragraph(
                "bacquet@flichy.com"
            )
            document.add_paragraph(
                "+33 1 56 62 30 00"
            )

            document.add_paragraph(
                "YOUR L&E GLOBAL POC"
            )
            document.add_paragraph(
                (
                    "For all inquiries related to this project, "
                    "please contact Jessica Stout, International "
                    "Business Development Executive at L&E Global, "
                    "at jessica.stout@leglobal.law."
                )
            )

            document.save(file_path)

            contacts = extract_contacts_from_docx(
                file_path,
                country="France",
            )

            self.assertEqual(
                len(contacts),
                1,
            )

            contact = contacts[0]

            self.assertEqual(
                contact.member_firm,
                "Flichy Grangé Avocats",
            )
            self.assertEqual(
                contact.contact_person,
                (
                    "Caroline Scherrmann and "
                    "Florence Bacquet"
                ),
            )
            self.assertEqual(
                contact.email,
                (
                    "scherrmann@flichy.com, "
                    "bacquet@flichy.com"
                ),
            )
            self.assertEqual(
                contact.phone,
                "+33 1 56 62 30 00",
            )
            self.assertIsNone(contact.address)
            self.assertIsNone(contact.website)

            rendered = " ".join(
                value
                for value in (
                    contact.member_firm,
                    contact.contact_person,
                    contact.email,
                    contact.phone,
                )
                if value
            )

            self.assertNotIn(
                "Jessica Stout",
                rendered,
            )
            self.assertNotIn(
                "jessica.stout@leglobal.law",
                rendered,
            )

    def test_generic_plain_paragraph_contact_is_extracted(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "generic-body-contact.docx"
            )

            document = Document()

            document.add_paragraph("FREEDONIA")
            document.add_paragraph(
                "Example Employment Law"
            )

            document.add_heading(
                "Employment Contracts",
                level=1,
            )
            document.add_paragraph(
                "Representative legal information."
            )

            document.add_paragraph(
                "Alex Example and Sam Sample"
            )
            document.add_paragraph(
                "Partners, Example Employment Law"
            )
            document.add_paragraph(
                "alex@example-law.test"
            )
            document.add_paragraph(
                "sam@example-law.test"
            )
            document.add_paragraph(
                "+99 123 456 789"
            )

            document.save(file_path)

            contacts = extract_contacts_from_docx(
                file_path,
                country="Freedonia",
            )

            self.assertEqual(
                len(contacts),
                1,
            )

            contact = contacts[0]

            self.assertEqual(
                contact.member_firm,
                "Example Employment Law",
            )
            self.assertEqual(
                contact.contact_person,
                "Alex Example and Sam Sample",
            )
            self.assertEqual(
                contact.email,
                (
                    "alex@example-law.test, "
                    "sam@example-law.test"
                ),
            )
            self.assertEqual(
                contact.phone,
                "+99 123 456 789",
            )

    def test_project_poc_alone_is_not_member_firm_contact(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "project-poc-only.docx"
            )

            document = Document()

            document.add_paragraph(
                "YOUR L&E GLOBAL POC"
            )
            document.add_paragraph(
                (
                    "For all inquiries related to this project, "
                    "please contact Jessica Stout, International "
                    "Business Development Executive at L&E Global, "
                    "at jessica.stout@leglobal.law."
                )
            )

            document.save(file_path)

            self.assertEqual(
                extract_contacts_from_docx(
                    file_path,
                    country="France",
                ),
                [],
            )

    def test_firm_name_alone_is_not_a_contact(
        self,
    ) -> None:
        cases = (
            (
                "Portugal",
                "SRS LEGAL",
            ),
            (
                "Taiwan",
                "Lee and Li, Attorneys-at-Law",
            ),
        )

        for country, firm_name in cases:
            with self.subTest(country=country):
                with TemporaryDirectory() as temporary_directory:
                    file_path = (
                        Path(temporary_directory)
                        / f"{country.lower()}-firm-only.docx"
                    )

                    document = Document()
                    document.add_paragraph(
                        country.upper()
                    )
                    document.add_paragraph(
                        firm_name
                    )
                    document.add_heading(
                        "GENERAL OVERVIEW",
                        level=1,
                    )
                    document.add_paragraph(
                        "Representative legal information."
                    )

                    document.save(file_path)

                    self.assertEqual(
                        extract_contacts_from_docx(
                            file_path,
                            country=country,
                        ),
                        [],
                    )

    def test_legal_reference_coordinates_do_not_become_contact(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "legal-references-only.docx"
            )

            document = Document()

            document.add_paragraph(
                "PHILIPPINES"
            )
            document.add_heading(
                "Employment Benefits",
                level=1,
            )
            document.add_paragraph(
                (
                    "More information may be obtained from the "
                    "public labour authority at "
                    "www.labour-authority.example."
                )
            )
            document.add_paragraph(
                (
                    "The authority may also be reached at "
                    "+63 2 8123 4567 for public information."
                )
            )

            document.save(file_path)

            self.assertEqual(
                extract_contacts_from_docx(
                    file_path,
                    country="Philippines",
                ),
                [],
            )

    def test_existing_text_box_contact_keeps_priority(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = (
                Path(temporary_directory)
                / "legacy-textbox-priority.docx"
            )

            document = Document()

            document.add_paragraph(
                "Alternative Person"
            )
            document.add_paragraph(
                "Partner, Alternative Firm"
            )
            document.add_paragraph(
                "alternative@alternative.example"
            )
            document.add_paragraph(
                "+44 20 0000 0000"
            )

            document.save(file_path)

            legacy_blocks = [
                [
                    "Stable Firm",
                    "Testland",
                    "1 Existing Street",
                    "+1 555 111 2222",
                    "www.stable.example",
                ],
                [
                    "CONTACT PERSON",
                    "Stable Person",
                    "stable@stable.example",
                ],
            ]

            with patch(
                (
                    "app.services.docx_parser."
                    "extract_text_box_blocks"
                ),
                return_value=legacy_blocks,
            ):
                contacts = extract_contacts_from_docx(
                    file_path,
                    country="Testland",
                )

            self.assertEqual(
                len(contacts),
                1,
            )

            contact = contacts[0]

            self.assertEqual(
                contact.member_firm,
                "Stable Firm",
            )
            self.assertEqual(
                contact.contact_person,
                "Stable Person",
            )
            self.assertEqual(
                contact.email,
                "stable@stable.example",
            )
            self.assertEqual(
                contact.phone,
                "+1 555 111 2222",
            )


def _build_france_shaped_document(file_path: Path) -> None:
    """The exact real France source layout (mission "FINAL CONTACT CRUD
    CLOSURE"): legal content, then a legacy member-firm contact block
    naming two people sharing one firm/phone, immediately before the
    L&E Global POC block - never a synthetic table-only stand-in."""

    document = Document()

    document.add_paragraph("FRANCE")
    document.add_paragraph("EMPLOYMENT LAW OVERVIEWS 2025 - 2026")
    document.add_paragraph("FLICHY GRANGÉ AVOCATS")

    document.add_heading("I. GENERAL OVERVIEW", level=1)
    document.add_paragraph("Representative employment-law content.")

    document.add_paragraph("Caroline Scherrmann and Florence Bacquet")
    document.add_paragraph("Partners, Flichy Grangé Avocats")
    document.add_paragraph("scherrmann@flichy.com")
    document.add_paragraph("bacquet@flichy.com")
    document.add_paragraph("+33 1 56 62 30 00")

    document.add_paragraph("YOUR L&E GLOBAL POC")
    document.add_paragraph(
        (
            "For all inquiries related to this project, please "
            "contact Jessica Stout, International Business "
            "Development Executive at L&E Global, at "
            "jessica.stout@leglobal.law."
        )
    )
    document.add_paragraph("Disclaimer text follows here.")

    document.save(file_path)


class FindPlainParagraphContactBlockBoundsTests(unittest.TestCase):
    """Corpus-independent coverage of find_plain_paragraph_contact_
    block_bounds - the structural anchor-finder contact_document_
    area.py's canonicalizer uses so a France-style legacy contact
    area (ordinary body paragraphs, no floating shape) gets its
    canonical table replacement inserted at the SAME logical location,
    never silently falling back to the document's start."""

    def test_finds_the_exact_block_bounds_before_poc(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / "france.docx"
            _build_france_shaped_document(file_path)

            document = Document(file_path)
            bounds = find_plain_paragraph_contact_block_bounds(document)

            self.assertIsNotNone(bounds)
            first_index, last_index = bounds

            self.assertEqual(
                "Caroline Scherrmann and Florence Bacquet",
                document.paragraphs[first_index].text,
            )
            self.assertEqual(
                "+33 1 56 62 30 00",
                document.paragraphs[last_index].text,
            )

            # The paragraph immediately after the block must be the
            # POC heading, never legal content or the disclaimer -
            # the block's own bounds must not over- or under-reach.
            self.assertEqual(
                "YOUR L&E GLOBAL POC",
                document.paragraphs[last_index + 1].text,
            )

    def test_returns_none_when_no_plain_paragraph_block_exists(
        self,
    ) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / "no-contact.docx"
            document = Document()
            document.add_paragraph("Just an ordinary legal document.")
            document.add_paragraph("With no contact information at all.")
            document.save(file_path)

            reopened = Document(file_path)
            self.assertIsNone(
                find_plain_paragraph_contact_block_bounds(reopened)
            )


class SplitCombinedLegacyContactTests(unittest.TestCase):
    """Corpus-independent coverage of split_combined_legacy_contact -
    the narrow, deterministic normalization that turns a legacy
    contact naming multiple people into one ExtractedContact per
    person, only when the split is unambiguous."""

    def test_splits_two_people_sharing_one_firm_and_phone(self) -> None:
        combined = ExtractedContact(
            member_firm="Flichy Grangé Avocats",
            contact_person="Caroline Scherrmann and Florence Bacquet",
            email="scherrmann@flichy.com, bacquet@flichy.com",
            phone="+33 1 56 62 30 00",
        )

        split = split_combined_legacy_contact(combined)

        self.assertIsNotNone(split)
        self.assertEqual(2, len(split))

        self.assertEqual("Caroline Scherrmann", split[0].contact_person)
        self.assertEqual("scherrmann@flichy.com", split[0].email)
        self.assertEqual("Florence Bacquet", split[1].contact_person)
        self.assertEqual("bacquet@flichy.com", split[1].email)

        for contact in split:
            self.assertEqual("Flichy Grangé Avocats", contact.member_firm)
            self.assertEqual("+33 1 56 62 30 00", contact.phone)

    def test_returns_none_for_a_single_person_contact(self) -> None:
        single = ExtractedContact(
            member_firm="Some Firm",
            contact_person="Jane Doe",
            email="jane@example.com",
        )

        self.assertIsNone(split_combined_legacy_contact(single))

    def test_returns_none_when_person_and_email_counts_mismatch(
        self,
    ) -> None:
        mismatched = ExtractedContact(
            member_firm="Some Firm",
            contact_person="Jane Doe and John Roe",
            email="jane@example.com",
        )

        self.assertIsNone(split_combined_legacy_contact(mismatched))

    def test_never_splits_a_firm_name_containing_and(self) -> None:
        """A firm name that happens to contain "and" is never at
        risk: only contact_person is ever split, member_firm is
        always copied through unchanged."""

        firm_with_and = ExtractedContact(
            member_firm="Smith and Jones LLP",
            contact_person="Jane Doe",
            email="jane@example.com",
        )

        self.assertIsNone(split_combined_legacy_contact(firm_with_and))

    def test_splits_three_people(self) -> None:
        combined = ExtractedContact(
            member_firm="Big Firm LLP",
            contact_person="Alice Smith, Bob Jones and Carol White",
            email="alice@example.com, bob@example.com, carol@example.com",
        )

        split = split_combined_legacy_contact(combined)

        self.assertIsNotNone(split)
        self.assertEqual(3, len(split))
        self.assertEqual(
            ["Alice Smith", "Bob Jones", "Carol White"],
            [c.contact_person for c in split],
        )
        self.assertEqual(
            ["alice@example.com", "bob@example.com", "carol@example.com"],
            [c.email for c in split],
        )

    def test_returns_none_without_email(self) -> None:
        no_email = ExtractedContact(
            member_firm="Some Firm",
            contact_person="Jane Doe and John Roe",
            email=None,
        )

        self.assertIsNone(split_combined_legacy_contact(no_email))


class ClassifyCanonicalFirmLinesTests(unittest.TestCase):
    """
    Direct, corpus-independent coverage of
    _classify_canonical_firm_lines - the canonical table reader's
    field classifier, fixed to be semantic (content-based) rather
    than position-only. Corpus regressions for the real documents
    this was found against (IE/IN/US) live in
    test_contact_document_area.py; these tests pin the underlying
    classification rules precisely, for every input shape called out
    by the mission: clean/international/spaced/+-prefixed phones,
    a phone occupying the first line when member_firm is empty, a
    phone with a trailing annotation, and a website mentioned inside
    address prose vs. one on its own dedicated line.
    """

    def test_normal_case_all_four_fields_present(self) -> None:
        member_firm, phone, website, remaining = _classify_canonical_firm_lines(
            [
                "HARMERS WORKPLACE LAWYERS",
                "31 Market Street, Level 27 St Martins Tower, NSW 2000 Sydney",
                "+61 292 674 322",
                "WWW.HARMERS.COM.AU",
            ]
        )
        self.assertEqual("HARMERS WORKPLACE LAWYERS", member_firm)
        self.assertEqual("+61 292 674 322", phone)
        self.assertEqual("WWW.HARMERS.COM.AU", website)
        self.assertEqual(
            ["31 Market Street, Level 27 St Martins Tower, NSW 2000 Sydney"],
            [
                ["HARMERS WORKPLACE LAWYERS",
                 "31 Market Street, Level 27 St Martins Tower, NSW 2000 Sydney",
                 "+61 292 674 322", "WWW.HARMERS.COM.AU"][i]
                for i in remaining
            ],
        )

    def test_clean_phone(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(["+1 5551234567"])
        self.assertEqual("+1 5551234567", phone)

    def test_international_phone(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(["+81 355 012 111"])
        self.assertEqual("+81 355 012 111", phone)

    def test_phone_with_spaces(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(["+46 852 206 500"])
        self.assertEqual("+46 852 206 500", phone)

    def test_phone_with_plus_prefix(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(["+353 1 234 5678"])
        self.assertEqual("+353 1 234 5678", phone)

    def test_phone_without_plus_prefix(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(["1 212 545 4050"])
        self.assertEqual("1 212 545 4050", phone)

    def test_phone_as_first_firm_side_line_no_member_firm(self) -> None:
        """Quirk B: a contact whose member_firm/address/website are
        all empty has its phone land on line 0 - it must still be
        recognized as phone, never misread as member_firm."""

        member_firm, phone, website, remaining = _classify_canonical_firm_lines(
            ["+1 555 000 0000"]
        )
        self.assertIsNone(member_firm)
        self.assertEqual("+1 555 000 0000", phone)
        self.assertIsNone(website)
        self.assertEqual([], remaining)

    def test_phone_with_trailing_annotation_preserved_whole(self) -> None:
        """Quirk A: a phone value with a trailing annotation must be
        preserved in its ENTIRETY as the phone field - never split,
        with the remainder leaking into address."""

        member_firm, phone, website, remaining = _classify_canonical_firm_lines(
            [
                "Cederquist",
                "Hovslagargatan 3, SE-111 96 Stockholm",
                "+46 852 206 500 (updated)",
                "www.cederquist.se",
            ]
        )
        self.assertEqual("Cederquist", member_firm)
        self.assertEqual("+46 852 206 500 (updated)", phone)
        self.assertEqual("www.cederquist.se", website)
        self.assertEqual(
            ["Hovslagargatan 3, SE-111 96 Stockholm"],
            [
                [
                    "Cederquist",
                    "Hovslagargatan 3, SE-111 96 Stockholm",
                    "+46 852 206 500 (updated)",
                    "www.cederquist.se",
                ][i]
                for i in remaining
            ],
        )

    def test_phone_as_first_line_with_annotation_combined(self) -> None:
        """Quirk A and Quirk B together: no member_firm AND a phone
        annotation on the same, first line."""

        member_firm, phone, website, remaining = _classify_canonical_firm_lines(
            ["+353 1 234 5678 (mobile)"]
        )
        self.assertIsNone(member_firm)
        self.assertEqual("+353 1 234 5678 (mobile)", phone)
        self.assertEqual([], remaining)

    def test_website_mentioned_inside_address_prose_is_not_extracted(self) -> None:
        """Quirk C: an address sentence that merely MENTIONS a URL
        (with its own sentence-ending punctuation attached) must never
        be misread as the website field - only a line that is, in its
        entirety, just a URL is the dedicated website field."""

        member_firm, phone, website, remaining = _classify_canonical_firm_lines(
            [
                "Jackson Lewis PC",
                "USA, 666 Third Avenue, 29th Floor, 10017 New York, Jackson "
                "Lewis has over 60 offices throughout the USA. , For "
                "information, please see www.jacksonlewis.com.",
                "1 212 545 4050",
                "www.jacksonlewis.com",
            ]
        )
        self.assertEqual("Jackson Lewis PC", member_firm)
        self.assertEqual("1 212 545 4050", phone)
        self.assertEqual("www.jacksonlewis.com", website)
        remaining_lines = [
            [
                "Jackson Lewis PC",
                "USA, 666 Third Avenue, 29th Floor, 10017 New York, Jackson "
                "Lewis has over 60 offices throughout the USA. , For "
                "information, please see www.jacksonlewis.com.",
                "1 212 545 4050",
                "www.jacksonlewis.com",
            ][i]
            for i in remaining
        ]
        self.assertEqual(1, len(remaining_lines))
        self.assertIn("please see www.jacksonlewis.com.", remaining_lines[0])

    def test_no_phone_or_website_present(self) -> None:
        member_firm, phone, website, remaining = _classify_canonical_firm_lines(
            ["Some Firm Ltd", "123 Some Street"]
        )
        self.assertEqual("Some Firm Ltd", member_firm)
        self.assertIsNone(phone)
        self.assertIsNone(website)
        self.assertEqual([1], remaining)

    def test_empty_firm_lines(self) -> None:
        member_firm, phone, website, remaining = _classify_canonical_firm_lines([])
        self.assertIsNone(member_firm)
        self.assertIsNone(phone)
        self.assertIsNone(website)
        self.assertEqual([], remaining)


if __name__ == "__main__":
    unittest.main()