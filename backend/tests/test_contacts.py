"""Consolidated test module generated from validated domain owners."""

from __future__ import annotations



# ================================================================
# SOURCE: backend/tests/test_admin_contacts.py
# ================================================================

import contextlib
import hashlib
import shutil
import tempfile
import unittest
import zipfile
from contextlib import nullcontext
from pathlib import Path
from types import SimpleNamespace
from typing import Any
from unittest.mock import patch
from docx import Document
from app.routers.admin_contacts import router
from app.services import admin_contacts
from app.services.admin_contacts import AdminContactMutationFailedError, AdminContactNotFoundError, add_contact, apply_structured_contact_state_to_chunks, bootstrap_legacy_contacts, delete_contact, list_contacts, reseed_contact_state_from_parsed_contacts, reseed_contacts_from_current_docx, update_contact
from app.services.admin_contact_photos import AdminContactPhotoError, AdminContactPhotoNotFoundError, read_admin_contact_photo, remove_admin_contact_photo, replace_admin_contact_photo
from app.services.admin_document_lifecycle import AdminDocumentRollbackError, get_document_download, reindex_indexed_document
from app.services.admin_document_replacement import safe_upload_and_index_document
from app.services.admin_modification_marker import is_admin_modified_since_upload, mark_admin_modified
from app.services.contact_document_area import ContactPhotoPayload, rebuild_canonical_contact_table
from app.services.contact_photo_store import ContactPhotoStorageError, delete_contact_photo, read_contact_photo, write_contact_photo_atomic
from app.services.contact_photos import extract_contact_photo_candidates
from app.services.contact_state import ContactRecord, ContactState, ContactStateError, new_contact_id, read_contact_state, write_contact_state_atomic
from pydantic import ValidationError
from app.models.admin_contacts import AdminContactWriteRequest
from app.models.admin_documents import AdminDocumentListResponse, AdminDocumentSummary
from app.models.document import DocumentChunk
from app.services.docx_parser import CONTACT_TABLE_HIDDEN_MARKER, ExtractedContact, extract_contacts_from_docx
from app.services.document_indexer import DocumentIndexingResult
from tests.support.documents import resolve_source_root
from tests.support.documents import make_png as _make_png, require_corpus_copy, skip_if_already_canonicalized
_test_admin_contacts__SOURCE_ROOT = resolve_source_root()

def _real_document_id_for(country_code: str, language: str='en') -> str:
    """
    The one real, deterministic document_id
    build_contact_chunk_for_contacts (and every other chunk builder)
    would compute for this country_code/language - document_id is
    derived solely from country_code + DOCUMENT_FAMILY + language,
    never from anything a test happens to pick, so every fake chunk
    this test file seeds/asserts against must use this exact value,
    never an arbitrary placeholder string.
    """
    from app.services.document_chunk_builder import DocumentMetadata
    from app.services.docx_parser import ExtractedContact as _EC
    from app.services.document_chunk_builder import build_contact_chunk_for_contacts as _probe_builder
    probe_chunk = _probe_builder([_EC(member_firm='probe')], DocumentMetadata(country='Probe Country', country_code=country_code, reference_year=None, language=language, source_filename='probe.docx'))
    return probe_chunk.document_id
DOCUMENT_ID = _real_document_id_for('GB')
OTHER_DOCUMENT_ID = _real_document_id_for('FR')

def _make_valid_jpeg(width: int=183, height: int=234) -> bytes:
    """
    A minimal but REAL-SIZED, structurally complete baseline JPEG -
    SOI/APP0(JFIF)/SOF0/SOS/EOI with plausible width/height, never a
    degenerate single-pixel stub. python-docx's own image-header
    parser (used to compute a photo's proportional height in the
    canonical contact table) stops reading at the first SOS marker, so
    the entropy-coded scan data itself can be a single placeholder
    byte - only the headers before it need to be genuine.
    """
    import struct
    soi = b'\xff\xd8'
    app0_data = b'JFIF\x00\x01\x01\x00' + struct.pack('>HH', 96, 96) + b'\x00\x00'
    app0 = b'\xff\xe0' + struct.pack('>H', len(app0_data) + 2) + app0_data
    sof0_data = bytes([8]) + struct.pack('>HH', height, width) + bytes([1, 1, 17, 0])
    sof0 = b'\xff\xc0' + struct.pack('>H', len(sof0_data) + 2) + sof0_data
    sos_data = bytes([1, 1, 0, 0, 63, 0])
    sos = b'\xff\xda' + struct.pack('>H', len(sos_data) + 2) + sos_data
    return soi + app0 + sof0 + sos + b'\x00\xff\xd9'
_VALID_JPEG = _make_valid_jpeg()
_test_admin_contacts__VALID_PNG = _make_png(183, 234, (120, 80, 200))

def _write_request(*, member_firm: str='Example & Partners', contact_person: str='Alex Example', email: str='alex@example.test', phone: str='+1 555 0100', address: str='1 Example Street', website: str='www.example.test') -> AdminContactWriteRequest:
    return AdminContactWriteRequest(member_firm=member_firm, contact_person=contact_person, email=email, phone=phone, address=address, website=website)

def _full_contact_record(**overrides: Any) -> ContactRecord:
    defaults = dict(contact_id=new_contact_id(), member_firm='Example & Partners', contact_person='Alex Example', email='alex@example.test', phone='+1 555 0100', address='1 Example Street', website='www.example.test')
    defaults.update(overrides)
    return ContactRecord(**defaults)

class FakeContactOpenSearchClient:
    """
    OpenSearch test double for admin_contacts.py.

    Stateful over self.chunks (chunk_id -> source dict) so a mutation's
    real effect can be asserted afterwards, not merely inferred from
    call arguments - mirrors FakeSectionOpenSearchClient's own design
    in test_admin_document_sections.py exactly, generalized to filter
    on subsection.keyword (never legal_topic, which the Contact chunk
    always carries as None).
    """

    def __init__(self, *, document_id: str=DOCUMENT_ID, country_code: str='GB', country: str='United Kingdom', source_filename: str='GB.docx', reference_year: int | None=2026, chunks: dict[str, dict[str, Any]] | None=None, country_document_ids: list[str] | None=None, fail_delete_by_query_calls: int=0, delete_by_query_failure: Exception | None=None) -> None:
        self.document_id = document_id
        self.country_code = country_code
        self.country = country
        self.source_filename = source_filename
        self.reference_year = reference_year
        self.chunks: dict[str, dict[str, Any]] = dict(chunks or {})
        self.country_document_ids = country_document_ids if country_document_ids is not None else [document_id]
        self.fail_delete_by_query_calls = fail_delete_by_query_calls
        self.delete_by_query_failure = delete_by_query_failure
        self.delete_by_query_call_count = 0
        self.delete_by_query_calls: list[dict[str, Any]] = []

    def search(self, index: str, body: dict[str, Any]) -> dict[str, Any]:
        del index
        if 'sort' in body:
            term = body.get('query', {}).get('term', {})
            requested_document_id = term.get('document_id')
            requested_country_code = term.get('country_code')
            if requested_document_id is not None:
                matching_ids = sorted((chunk_id for chunk_id, chunk in self.chunks.items() if chunk['document_id'] == requested_document_id))
            elif requested_country_code is not None:
                return {'hits': {'total': {'value': len(self.country_document_ids)}, 'hits': [{'_id': f'doc-row-{doc_id}', '_source': {'document_id': doc_id, 'country_code': requested_country_code, 'country': self.country, 'source_filename': self.source_filename, 'reference_year': self.reference_year}, 'sort': [f'doc-row-{doc_id}']} for doc_id in self.country_document_ids]}}
            else:
                matching_ids = []
            return {'hits': {'total': {'value': len(matching_ids)}, 'hits': [{'_id': chunk_id, '_source': self.chunks[chunk_id], 'sort': [chunk_id]} for chunk_id in matching_ids]}}
        term = body.get('query', {}).get('term', {})
        requested_document_id = term.get('document_id')
        if requested_document_id is not None:
            if requested_document_id != self.document_id:
                return {'hits': {'hits': []}}
            return {'hits': {'hits': [{'_source': {'document_id': self.document_id, 'source_filename': self.source_filename, 'country': self.country, 'country_code': self.country_code, 'reference_year': self.reference_year}}]}}
        return {'hits': {'hits': []}}

    def delete_by_query(self, *, index: str, body: dict[str, Any], conflicts: str, refresh: bool) -> dict[str, Any]:
        del index, conflicts, refresh
        self.delete_by_query_call_count += 1
        if self.delete_by_query_call_count <= self.fail_delete_by_query_calls:
            raise self.delete_by_query_failure if self.delete_by_query_failure is not None else RuntimeError(f'simulated delete_by_query failure (call #{self.delete_by_query_call_count}).')
        query = body['query']
        filters = query['bool']['filter']
        document_id = next((clause['term']['document_id'] for clause in filters if 'document_id' in clause.get('term', {})))
        subsection = next((clause['term']['subsection.keyword'] for clause in filters if 'subsection.keyword' in clause.get('term', {})), None)
        keep_ids: set[str] = set()
        for clause in query['bool'].get('must_not', []):
            keep_ids.update(clause.get('terms', {}).get('chunk_id', []))
        to_delete = [chunk_id for chunk_id, chunk in self.chunks.items() if chunk['document_id'] == document_id and (subsection is None or chunk.get('subsection') == subsection) and (chunk_id not in keep_ids)]
        for chunk_id in to_delete:
            del self.chunks[chunk_id]
        self.delete_by_query_calls.append({'document_id': document_id, 'subsection': subsection, 'deleted': len(to_delete)})
        return {'deleted': len(to_delete), 'total': len(to_delete)}

def _bulk_writer(fake_client: FakeContactOpenSearchClient, *, fail_first_n_calls: int=0):
    call_count = {'n': 0}

    def fake_bulk(client, actions, **kwargs):
        del client, kwargs
        call_count['n'] += 1
        if call_count['n'] <= fail_first_n_calls:
            raise RuntimeError(f"simulated OpenSearch bulk failure (call #{call_count['n']})")
        action_list = list(actions)
        for action in action_list:
            fake_client.chunks[action['_id']] = dict(action['_source'])
        return (len(action_list), [])
    return fake_bulk

@contextlib.contextmanager
def _patched_indexer(fake_client: FakeContactOpenSearchClient, *, fail_bulk: bool=False):
    with patch('app.services.document_indexer.ensure_legal_documents_index'), patch('app.services.document_indexer.bulk', side_effect=_bulk_writer(fake_client, fail_first_n_calls=1 if fail_bulk else 0)):
        yield

def _seeded_contact_chunk(*, document_id: str=DOCUMENT_ID, country_code: str='GB', country: str='United Kingdom', source_filename: str='GB.docx', content: str='Member firm: Old Firm\nContact person: Old Person') -> dict[str, Any]:
    return {'document_id': document_id, 'country_code': country_code, 'country': country, 'source_filename': source_filename, 'reference_year': 2026, 'legal_topic': None, 'document_type': 'overview', 'language': 'en', 'section': f'Employment Law Overview {country}', 'subsection': 'Contact', 'content': content, 'content_hash': f'hash-{hash(content)}'}

def _seed_placeholder_source_docx(source_directory: Path, filename: str='GB.docx') -> None:
    """
    A minimal, valid, structurally-empty DOCX - just enough for
    resolve_document_source_path to find a real file on disk. Add/
    Delete Contact always rebuild the persisted source's canonical
    contact table now, regardless of whether this placeholder has any
    prior contact structure at all, so these tests exercise that
    rebuild too (harmlessly, against a document with no real legal
    content) alongside the ContactState/OpenSearch chunk mutation they
    were originally written for.
    """
    document = Document()
    document.add_paragraph('Placeholder document body.')
    document.save(str(source_directory / filename))

class ContactStateTests(unittest.TestCase):

    def test_absent_state_returns_none(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            self.assertIsNone(read_contact_state(Path(root), DOCUMENT_ID))

    def test_explicit_empty_state_is_not_none(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=()))
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertIsNotNone(state)
            self.assertEqual(state.contacts, ())

    def test_stable_ordering_round_trip(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            first = _full_contact_record(member_firm='Firm One')
            second = _full_contact_record(member_firm='Firm Two')
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=(first, second)))
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual([c.contact_id for c in state.contacts], [first.contact_id, second.contact_id])

    def test_write_is_atomic_no_partial_file_left_on_crash(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=(_full_contact_record(),)))
            state_dir = source_directory / '.admin-state' / 'contacts'
            leftover_temp_files = [path for path in state_dir.iterdir() if path.name.endswith('.json.tmp')]
            self.assertEqual(leftover_temp_files, [])

class ContactPhotoMetadataTests(unittest.TestCase):

    def test_legacy_contact_without_photo_fields_is_readable(self) -> None:
        record = ContactRecord.from_json_dict({'contact_id': 'contact-legacy', 'member_firm': 'Firm', 'contact_person': 'Jane Doe', 'email': 'jane@example.com', 'phone': '+1', 'address': 'Address', 'website': 'example.com'})
        self.assertIsNone(record.photo_filename)
        self.assertIsNone(record.photo_content_type)
        self.assertIsNone(record.photo_sha256)

    def test_photo_metadata_serializes_without_binary_data(self) -> None:
        digest = 'a' * 64
        record = ContactRecord(contact_id='contact-123', member_firm='Firm', contact_person='Jane Doe', email='jane@example.com', photo_filename=f'contact-123--{digest}.jpg', photo_content_type='image/jpeg', photo_sha256=digest)
        payload = record.to_json_dict()
        self.assertEqual(f'contact-123--{digest}.jpg', payload['photo_filename'])
        self.assertEqual('image/jpeg', payload['photo_content_type'])
        self.assertEqual(digest, payload['photo_sha256'])
        self.assertFalse(any((isinstance(value, (bytes, bytearray)) for value in payload.values())))

    def test_partial_photo_metadata_is_rejected(self) -> None:
        with self.assertRaises(ContactStateError):
            ContactRecord.from_json_dict({'contact_id': 'contact-123', 'photo_filename': 'photo.jpg', 'photo_content_type': None, 'photo_sha256': None})

class ContactPhotoStoreTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.source_directory = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_atomic_write_uses_contact_id_and_sha256(self) -> None:
        data = b'fake-jpeg-content'
        digest = hashlib.sha256(data).hexdigest()
        stored = write_contact_photo_atomic(self.source_directory, 'contact-123', data=data, content_type='image/jpeg')
        self.assertEqual(f'contact-123--{digest}.jpg', stored.filename)
        self.assertEqual('image/jpeg', stored.content_type)
        self.assertEqual(digest, stored.sha256)
        self.assertEqual(data, read_contact_photo(self.source_directory, stored.filename))

    def test_photo_is_stored_inside_admin_state(self) -> None:
        stored = write_contact_photo_atomic(self.source_directory, 'contact-123', data=b'photo', content_type='image/png')
        expected = self.source_directory / '.admin-state' / 'contact-photos' / stored.filename
        self.assertTrue(expected.is_file())

    def test_same_photo_write_is_idempotent(self) -> None:
        first = write_contact_photo_atomic(self.source_directory, 'contact-123', data=b'same-photo', content_type='image/jpeg')
        second = write_contact_photo_atomic(self.source_directory, 'contact-123', data=b'same-photo', content_type='image/jpeg')
        self.assertEqual(first, second)
        files = list((self.source_directory / '.admin-state' / 'contact-photos').iterdir())
        self.assertEqual(1, len(files))

    def test_failed_new_write_preserves_existing_photo(self) -> None:
        old = write_contact_photo_atomic(self.source_directory, 'contact-123', data=b'old-photo', content_type='image/jpeg')
        with patch('app.services.contact_photo_store.os.replace', side_effect=OSError('boom')):
            with self.assertRaises(ContactPhotoStorageError):
                write_contact_photo_atomic(self.source_directory, 'contact-123', data=b'new-photo', content_type='image/png')
        self.assertEqual(b'old-photo', read_contact_photo(self.source_directory, old.filename))
        store = self.source_directory / '.admin-state' / 'contact-photos'
        self.assertEqual([old.filename], sorted((p.name for p in store.iterdir() if p.is_file())))

    def test_unsupported_content_type_is_rejected(self) -> None:
        with self.assertRaises(ContactPhotoStorageError):
            write_contact_photo_atomic(self.source_directory, 'contact-123', data=b'photo', content_type='image/svg+xml')

    def test_empty_photo_is_rejected(self) -> None:
        with self.assertRaises(ContactPhotoStorageError):
            write_contact_photo_atomic(self.source_directory, 'contact-123', data=b'', content_type='image/jpeg')

    def test_path_traversal_filename_is_rejected(self) -> None:
        with self.assertRaises(ContactPhotoStorageError):
            read_contact_photo(self.source_directory, '../secret.jpg')
        with self.assertRaises(ContactPhotoStorageError):
            delete_contact_photo(self.source_directory, '../secret.jpg')

    def test_delete_is_safe_and_idempotent(self) -> None:
        stored = write_contact_photo_atomic(self.source_directory, 'contact-123', data=b'photo', content_type='image/webp')
        delete_contact_photo(self.source_directory, stored.filename)
        delete_contact_photo(self.source_directory, stored.filename)
        with self.assertRaises(ContactPhotoStorageError):
            read_contact_photo(self.source_directory, stored.filename)

class AdminModificationMarkerTests(unittest.TestCase):

    def test_absent_marker_defaults_false(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            self.assertFalse(is_admin_modified_since_upload(Path(root), DOCUMENT_ID))

    def test_mark_then_reset(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            mark_admin_modified(source_directory, DOCUMENT_ID)
            self.assertTrue(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

class ContactCrudTests(unittest.TestCase):

    def test_list_with_no_state_returns_empty(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()
            response = list_contacts(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(response.contacts, [])
            self.assertEqual(response.country_code, 'GB')

    def test_add_appends_with_fresh_stable_id(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                response = add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            self.assertTrue(response.contact_id)
            self.assertEqual(response.member_firm, 'Example & Partners')
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].contact_id, response.contact_id)
            self.assertTrue(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_add_syncs_the_opensearch_contact_chunk(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            contact_chunks = [chunk for chunk in client.chunks.values() if chunk.get('subsection') == 'Contact']
            self.assertEqual(len(contact_chunks), 1)
            self.assertIn('Example & Partners', contact_chunks[0]['content'])

    def test_zero_one_and_multiple_contacts(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            listing = list_contacts(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(listing.contacts, [])
            with _patched_indexer(client):
                first = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
            listing = list_contacts(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(len(listing.contacts), 1)
            with _patched_indexer(client):
                second = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm B'), source_directory=source_directory, client=client)
                third = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
            listing = list_contacts(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(len(listing.contacts), 3)
            self.assertEqual([c.contact_id for c in listing.contacts], [first.contact_id, second.contact_id, third.contact_id])

    def test_duplicate_contacts_have_distinct_ids(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                first = add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
                second = add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            self.assertEqual(first.member_firm, second.member_firm)
            self.assertEqual(first.email, second.email)
            self.assertNotEqual(first.contact_id, second.contact_id)

    def test_update_preserves_id_and_position(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                first = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
                second = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm B'), source_directory=source_directory, client=client)
                update_contact(document_id=DOCUMENT_ID, contact_id=second.contact_id, fields=_write_request(member_firm='Firm B Updated'), source_directory=source_directory, client=client)
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual([c.contact_id for c in state.contacts], [first.contact_id, second.contact_id])
            self.assertEqual(state.contacts[0].member_firm, 'Firm A')
            self.assertEqual(state.contacts[1].member_firm, 'Firm B Updated')

    def test_update_stale_contact_id_raises_not_found(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()
            with self.assertRaises(AdminContactNotFoundError):
                update_contact(document_id=DOCUMENT_ID, contact_id='does-not-exist', fields=_write_request(), source_directory=source_directory, client=client)

    def test_delete_removes_only_the_requested_contact(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                first = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
                second = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm B'), source_directory=source_directory, client=client)
                delete_contact(document_id=DOCUMENT_ID, contact_id=first.contact_id, source_directory=source_directory, client=client)
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].contact_id, second.contact_id)
            self.assertEqual(state.contacts[0].member_firm, 'Firm B')

    def test_delete_stale_contact_id_raises_not_found(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()
            with self.assertRaises(AdminContactNotFoundError):
                delete_contact(document_id=DOCUMENT_ID, contact_id='does-not-exist', source_directory=source_directory, client=client)

    def test_delete_last_contact_removes_stale_contact_chunk(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                only = add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
                self.assertEqual(len([c for c in client.chunks.values() if c.get('subsection') == 'Contact']), 1)
                delete_contact(document_id=DOCUMENT_ID, contact_id=only.contact_id, source_directory=source_directory, client=client)
            contact_chunks = [c for c in client.chunks.values() if c.get('subsection') == 'Contact']
            self.assertEqual(contact_chunks, [])

    def test_legal_topic_chunks_are_never_touched(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient(chunks={'legal-chunk-1': {'document_id': DOCUMENT_ID, 'country_code': 'GB', 'country': 'United Kingdom', 'source_filename': 'GB.docx', 'reference_year': 2026, 'legal_topic': 'Employment Contracts', 'document_type': 'overview', 'language': 'en', 'section': 'Employment Contracts', 'subsection': None, 'content': 'legal content', 'content_hash': 'hash-legal'}})
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            self.assertIn('legal-chunk-1', client.chunks)
            self.assertEqual(client.chunks['legal-chunk-1']['content'], 'legal content')

    def test_country_conflict_blocks_contact_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient(country_document_ids=[DOCUMENT_ID, OTHER_DOCUMENT_ID])
            with self.assertRaises(Exception):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)

    def test_list_exposes_has_photo_true_for_a_contact_with_a_photo(self) -> None:
        """
        The safe photo-presence contract Admin/View needs: the Admin
        list must never blindly attempt a photo fetch for every contact
        regardless of whether one exists - no pointless 404 image
        requests for contacts that never had a photo. The response
        exposes has_photo only - never photo_filename or any
        filesystem path.
        """
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=(ContactRecord(contact_id=new_contact_id(), member_firm='Example & Partners', contact_person='Jane Doe', email='jane@example.com', phone='+1 555 0000', address='1 Example Street', website='https://example.com', photo_filename='deadbeef.jpg', photo_content_type='image/jpeg', photo_sha256='a' * 64),)))
            listing = list_contacts(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(len(listing.contacts), 1)
            self.assertTrue(listing.contacts[0].has_photo)
            serialized = listing.contacts[0].model_dump()
            self.assertNotIn('photo_filename', serialized)
            self.assertNotIn('photo_content_type', serialized)
            self.assertNotIn('photo_sha256', serialized)

    def test_list_exposes_has_photo_false_for_a_contact_without_a_photo(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            listing = list_contacts(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(len(listing.contacts), 1)
            self.assertFalse(listing.contacts[0].has_photo)

class AdminContactPhotoTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)

    def _seed_photo_bearing_contact(self) -> FakeContactOpenSearchClient:
        """
        Argentina (AR.docx) - a real single contact with a real
        existing photo, so REPLACE/REMOVE exercise the actual
        DOCX-mutation code path, not a synthetic stand-in.
        """
        docx_path = require_corpus_copy(self, _test_admin_contacts__SOURCE_ROOT, 'AR.docx', self.root)
        candidates = extract_contact_photo_candidates(docx_path)
        assert len(candidates) == 1
        photo = candidates[0]
        stored = write_contact_photo_atomic(self.root, 'contact-test', data=photo.data, content_type=photo.content_type)
        write_contact_state_atomic(self.root, ContactState(document_id='doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', country_code='AR', contacts=(ContactRecord(contact_id='contact-test', member_firm='Allende & Brea', contact_person='Nicolás Grandi', email='ngrandi@allende.com', phone='+1 555 0100', address='Address', website='https://example.com', photo_filename=stored.filename, photo_content_type=stored.content_type, photo_sha256=stored.sha256),)))
        return FakeContactOpenSearchClient(document_id='doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', country_code='AR', country='Argentina', source_filename='AR.docx')

    def _seed_photo_less_contact(self) -> FakeContactOpenSearchClient:
        """
        Germany (DE.docx) - a real single contact (Tobias Pusch) who
        genuinely has no photo yet, so a PUT exercises the real
        ADD-into-the-document code path.
        """
        docx_path = require_corpus_copy(self, _test_admin_contacts__SOURCE_ROOT, 'DE.docx', self.root)
        assert extract_contact_photo_candidates(docx_path) == []
        write_contact_state_atomic(self.root, ContactState(document_id='doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', country_code='DE', contacts=(ContactRecord(contact_id='contact-test', member_firm='Pusch Wahlig Workplace Law', contact_person='Tobias Pusch', email='pusch@pwwl.de', phone='+1 555 0100', address='Address', website='https://example.com'),)))
        return FakeContactOpenSearchClient(document_id='doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', country_code='DE', country='Germany', source_filename='DE.docx')

    def _seed_belgium_two_contacts(self) -> FakeContactOpenSearchClient:
        """
        Belgium's real two-contact, two-photo document - to prove
        isolation holds at the FULL service layer (ContactState +
        photo store + source DOCX together), not merely at the raw
        DOCX-primitive level test_contact_documents.py already covers.
        """
        temp_copy = require_corpus_copy(self, _test_admin_contacts__SOURCE_ROOT, 'Labour and Employment Law in Belgium 2026.docx', self.root)
        docx_path = self.root / 'BE.docx'
        shutil.copyfile(temp_copy, docx_path)
        candidates = extract_contact_photo_candidates(docx_path)
        by_name = {c.source_filename: c for c in candidates}
        chris_photo = by_name['image2.jpg']
        nicolas_photo = by_name['image1.png']
        chris_stored = write_contact_photo_atomic(self.root, 'chris-id', data=chris_photo.data, content_type=chris_photo.content_type)
        nicolas_stored = write_contact_photo_atomic(self.root, 'nicolas-id', data=nicolas_photo.data, content_type=nicolas_photo.content_type)
        write_contact_state_atomic(self.root, ContactState(document_id='doc_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb', country_code='BE', contacts=(ContactRecord(contact_id='chris-id', member_firm='Van Olmen & Wynant', contact_person='Chris van Olmen', email='chris.van.olmen@vow.be', phone='+1 555 0100', address='Address', website='https://example.com', photo_filename=chris_stored.filename, photo_content_type=chris_stored.content_type, photo_sha256=chris_stored.sha256), ContactRecord(contact_id='nicolas-id', member_firm='Van Olmen & Wynant', contact_person='Nicolas Simon', email='nicolas.simon@vow.be', phone='+1 555 0100', address='Address', website='https://example.com', photo_filename=nicolas_stored.filename, photo_content_type=nicolas_stored.content_type, photo_sha256=nicolas_stored.sha256))))
        return FakeContactOpenSearchClient(document_id='doc_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb', country_code='BE', country='Belgium', source_filename='BE.docx')

    def test_belgium_two_contact_isolation_at_the_service_layer(self) -> None:
        """Mutating Chris's photo must never touch Nicolas's
        ContactState, photo file, or DOCX image - and vice versa."""
        client = self._seed_belgium_two_contacts()
        docx_path = self.root / 'BE.docx'
        doc_id = 'doc_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb'
        original_state = read_contact_state(self.root, doc_id)
        nicolas_before = next((c for c in original_state.contacts if c.contact_id == 'nicolas-id'))
        replace_admin_contact_photo(self.root, doc_id, 'chris-id', data=_VALID_JPEG, content_type='image/jpeg', client=client)
        state_after = read_contact_state(self.root, doc_id)
        nicolas_after = next((c for c in state_after.contacts if c.contact_id == 'nicolas-id'))
        self.assertEqual(nicolas_before, nicolas_after)
        docx_shas = {c.sha256 for c in extract_contact_photo_candidates(docx_path)}
        self.assertIn(nicolas_before.photo_sha256, docx_shas)
        self.assertEqual(2, len(docx_shas))
        self.assertTrue(remove_admin_contact_photo(self.root, doc_id, 'nicolas-id', client=client))
        final_state = read_contact_state(self.root, doc_id)
        chris_final = next((c for c in final_state.contacts if c.contact_id == 'chris-id'))
        self.assertIsNotNone(chris_final.photo_sha256)
        final_docx_shas = {c.sha256 for c in extract_contact_photo_candidates(docx_path)}
        self.assertEqual(1, len(final_docx_shas))
        self.assertIn(chris_final.photo_sha256, final_docx_shas)

    def test_add_photo_for_a_contact_whose_name_matches_nothing_in_the_document(self) -> None:
        """
        A brand-new contact's name will usually have no matching
        "CONTACT PERSON" zone in the document at all (that IS the
        common case for genuinely adding someone - their name cannot
        possibly already appear anywhere). The canonical-table
        mechanism this codebase now uses embeds the new photo into a
        freshly rebuilt table regardless - it never depends on the
        mutated contact's own name matching anything already in the
        document (only the OLDER floating-shape primitives in
        test_contact_documents.py's own
        RealCorpusContactDocumentPhotoTests still need that kind of
        name/zone matching, for a document not yet canonicalized).
        """
        client = self._seed_photo_less_contact()
        docx_path = self.root / 'DE.docx'
        doc_id = 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa'
        write_contact_state_atomic(self.root, ContactState(document_id=doc_id, country_code='DE', contacts=(ContactRecord(contact_id='contact-test', member_firm='Someone New GmbH', contact_person='Someone New', email='new@example.test', phone='+1 555 0100', address='Address', website='https://example.com'),)))
        photo = replace_admin_contact_photo(self.root, doc_id, 'contact-test', data=_VALID_JPEG, content_type='image/jpeg', client=client)
        docx_shas = {c.sha256 for c in extract_contact_photo_candidates(docx_path)}
        self.assertEqual(1, len(docx_shas))
        self.assertIn(photo.sha256, docx_shas)
        state = read_contact_state(self.root, doc_id)
        self.assertEqual(photo.sha256, state.contacts[0].photo_sha256)
        download = get_document_download(document_id=doc_id, source_directory=self.root, client=client)
        downloaded_shas = {c.sha256 for c in extract_contact_photo_candidates(download.path)}
        self.assertIn(photo.sha256, downloaded_shas, "the downloaded DOCX must contain the new contact's photo, not merely ContactState")

    def _seed_zero_zone_country(self) -> FakeContactOpenSearchClient:
        """
        Portugal (PT.docx) - a real document with genuinely ZERO
        "CONTACT PERSON" zones anywhere. The canonical contact table
        mechanism rebuilds the whole contact area from ContactState
        regardless of prior structure, so a photo mutation here now
        succeeds (synchronizing a fresh table into the source) rather
        than failing closed the way the old floating-shape-only
        mechanism had to.
        """
        docx_path = require_corpus_copy(self, _test_admin_contacts__SOURCE_ROOT, 'PT.docx', self.root)
        assert extract_contact_photo_candidates(docx_path) == []
        write_contact_state_atomic(self.root, ContactState(document_id='doc_cccccccccccccccccccccccccccccccccccccccccccccccccccccccccccccc', country_code='PT', contacts=(ContactRecord(contact_id='contact-test', member_firm='Someone New Lda', contact_person='Someone New', email='new@example.test', phone='+1 555 0100', address='Address', website='https://example.com'),)))
        return FakeContactOpenSearchClient(document_id='doc_cccccccccccccccccccccccccccccccccccccccccccccccccccccccccccccc', country_code='PT', country='Portugal', source_filename='PT.docx')

    def test_new_contact_photo_for_a_zero_zone_document_still_syncs(self) -> None:
        """
        A document with no contact area at all still gets the new
        photo synchronized into a freshly rebuilt canonical table -
        the source DOCX is never left unsynchronized just because it
        had no prior contact structure.
        """
        client = self._seed_zero_zone_country()
        doc_id = 'doc_cccccccccccccccccccccccccccccccccccccccccccccccccccccccccccccc'
        photo = replace_admin_contact_photo(self.root, doc_id, 'contact-test', data=_VALID_JPEG, content_type='image/jpeg', client=client)
        docx_path = self.root / 'PT.docx'
        docx_candidates = extract_contact_photo_candidates(docx_path)
        self.assertEqual(1, len(docx_candidates))
        self.assertEqual(photo.sha256, docx_candidates[0].sha256)
        state = read_contact_state(self.root, doc_id)
        self.assertEqual(photo.sha256, state.contacts[0].photo_sha256)

    def test_replace_read_remove_syncs_the_source_docx(self) -> None:
        client = self._seed_photo_bearing_contact()
        docx_path = self.root / 'AR.docx'
        photo = replace_admin_contact_photo(self.root, 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', 'contact-test', data=_VALID_JPEG, content_type='image/jpeg', client=client)
        self.assertEqual('image/jpeg', photo.content_type)
        docx_shas = {c.sha256 for c in extract_contact_photo_candidates(docx_path)}
        self.assertIn(photo.sha256, docx_shas)
        self.assertEqual(1, len(docx_shas))
        loaded = read_admin_contact_photo(self.root, 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', 'contact-test')
        self.assertEqual(_VALID_JPEG, loaded.data)
        self.assertTrue(remove_admin_contact_photo(self.root, 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', 'contact-test', client=client))
        with self.assertRaises(AdminContactPhotoNotFoundError):
            read_admin_contact_photo(self.root, 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', 'contact-test')
        self.assertEqual([], extract_contact_photo_candidates(docx_path))

    def test_replace_failure_leaves_the_source_docx_and_state_unchanged(self) -> None:
        """
        A ContactState photo SHA that no longer matches anything in
        the source DOCX (simulating drift/corruption) is never
        consulted for a REPLACE - the canonical table is rebuilt fresh
        from the newly-uploaded bytes directly, so a stale/corrupt
        prior reference cannot block a legitimate new upload. This is
        a deliberate improvement over the old floating-shape mechanism
        (which had to locate the OLD sha in the document before it
        could replace it): rebuilding the whole area from ContactState
        every time eliminates this entire class of drift-driven
        failure.
        """
        client = self._seed_photo_bearing_contact()
        write_contact_state_atomic(self.root, ContactState(document_id='doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', country_code='AR', contacts=(ContactRecord(contact_id='contact-test', member_firm='Allende & Brea', contact_person='Nicolás Grandi', email='ngrandi@allende.com', phone='+1 555 0100', address='Address', website='https://example.com', photo_filename='contact-test--deadbeef.jpg', photo_content_type='image/jpeg', photo_sha256='0' * 64),)))
        photo = replace_admin_contact_photo(self.root, 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', 'contact-test', data=_VALID_JPEG, content_type='image/jpeg', client=client)
        docx_path = self.root / 'AR.docx'
        docx_candidates = extract_contact_photo_candidates(docx_path)
        self.assertEqual(1, len(docx_candidates))
        self.assertEqual(photo.sha256, docx_candidates[0].sha256)
        state = read_contact_state(self.root, 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa')
        self.assertEqual(photo.sha256, state.contacts[0].photo_sha256)

    def test_downloaded_docx_reflects_a_replaced_photo(self) -> None:
        """
        After a mutation, the SAME backend document-download path
        Admin uses must reflect it - not merely ContactState.
        get_document_download() is exactly the function backing GET
        .../download. Exercised here against a REAL corpus document
        (Argentina), complementing DownloadByteStabilityTests' own
        byte-identical proof above (which uses a synthetic placeholder
        source) with the same guarantee against a document carrying
        real pre-existing structure.
        """
        client = self._seed_photo_bearing_contact()
        photo = replace_admin_contact_photo(self.root, 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', 'contact-test', data=_VALID_JPEG, content_type='image/jpeg', client=client)
        download = get_document_download(document_id='doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', source_directory=self.root, client=client)
        downloaded_shas = {c.sha256 for c in extract_contact_photo_candidates(download.path)}
        self.assertIn(photo.sha256, downloaded_shas)
        self.assertEqual(1, len(downloaded_shas))

    def test_invalid_photo_payloads_are_rejected(self) -> None:
        invalid_payloads = {'not_an_image': (b'not-an-image', 'image/jpeg'), 'declared_type_does_not_match_content': (b'\x89PNG\r\n\x1a\nfake', 'image/jpeg')}
        for case, (data, content_type) in invalid_payloads.items():
            with self.subTest(case=case):
                client = self._seed_photo_bearing_contact()
                with self.assertRaises(AdminContactPhotoError):
                    replace_admin_contact_photo(self.root, 'doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa', 'contact-test', data=data, content_type=content_type, client=client)

class AdminContactPhotoRouteContractTests(unittest.TestCase):
    """
    The WordPress Admin proxy (class-le-global-chatbot-admin.php) builds
    every contact/photo URL from the exact same DOCUMENTS_PATH constant
    and route shape - these tests protect that wiring directly, since a
    mismatch here 404s silently in the browser (the proven root cause of
    a real "Admin View/Edit shows no photo thumbnail" defect: the
    browser's <img> gets a 404 and its onerror handler silently removes
    it).
    """

    def test_photo_routes_keep_admin_security_dependencies(self):
        normal = None
        photo_routes = []
        for route in router.routes:
            methods = getattr(route, 'methods', set())
            if route.path.endswith('/{document_id}/contacts') and 'GET' in methods:
                normal = route
            if route.path.endswith('/{document_id}/contacts/{contact_id}/photo'):
                photo_routes.append(route)
        self.assertIsNotNone(normal)
        self.assertEqual(3, len(photo_routes))
        normal_dependencies = len(normal.dependant.dependencies)
        self.assertGreater(normal_dependencies, 0)
        for route in photo_routes:
            self.assertGreaterEqual(len(route.dependant.dependencies), normal_dependencies)

    def test_photo_route_paths_share_the_documents_prefix(self):
        """
        WordPress builds every contact photo URL as DOCUMENTS_PATH +
        "/" + document_id + "/contacts/" + contact_id + "/photo",
        where DOCUMENTS_PATH is the exact same "/api/v1/admin/documents"
        constant used to build the list/add/update/delete contact URLs.
        The three photo routes must therefore share that same prefix -
        if they don't, every request WordPress sends 404s.
        """
        contacts_list_route = next((route for route in router.routes if route.path.endswith('/{document_id}/contacts') and 'GET' in route.methods))
        documents_prefix = contacts_list_route.path.removesuffix('/{document_id}/contacts')
        photo_routes = [route for route in router.routes if route.path.endswith('/{document_id}/contacts/{contact_id}/photo')]
        self.assertEqual(3, len(photo_routes))
        for route in photo_routes:
            self.assertTrue(route.path.startswith(documents_prefix + '/'), f"{route.path!r} does not share the {documents_prefix!r} prefix WordPress's DOCUMENTS_PATH constant assumes every Admin contact route (including photo routes) uses")

class DownloadByteStabilityTests(unittest.TestCase):

    def _sha(self, path: Path) -> str:
        return hashlib.sha256(path.read_bytes()).hexdigest()

    def test_persisted_source_sha_equals_download_sha(self) -> None:
        """A: sha256(persisted source) == sha256(download body)."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(download.path, persisted_path)
            self.assertEqual(self._sha(download.path), self._sha(persisted_path))

    def test_ten_consecutive_downloads_are_byte_identical(self) -> None:
        """B: 10 consecutive downloads of an unchanged document all
        have exactly the same SHA256."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            hashes = {self._sha(get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client).path) for _ in range(10)}
            self.assertEqual(len(hashes), 1)

    def test_download_changes_nothing(self) -> None:
        """C: download changes neither the source file (mtime/size)
        nor ContactState nor the OpenSearch chunk set."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            stat_before = persisted_path.stat()
            state_before = read_contact_state(source_directory, DOCUMENT_ID)
            chunks_before = dict(client.chunks)
            get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            stat_after = persisted_path.stat()
            state_after = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(stat_before.st_mtime_ns, stat_after.st_mtime_ns)
            self.assertEqual(stat_before.st_size, stat_after.st_size)
            self.assertEqual(state_before, state_after)
            self.assertEqual(chunks_before, client.chunks)

    def test_download_never_reaches_docx_writing_code(self) -> None:
        """D: download must never invoke materialize_effective_docx,
        python-docx's Document.save, or the canonical contact
        rebuild - patch all three to explode, and prove download
        still succeeds untouched."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            persisted_bytes = (source_directory / 'GB.docx').read_bytes()
            with patch('app.services.contact_document_area.rebuild_canonical_contact_table', side_effect=AssertionError('download must never rebuild the canonical table')), patch('docx.document.Document.save', side_effect=AssertionError('download must never call Document.save')):
                download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
                downloaded_bytes = download.path.read_bytes()
        self.assertEqual(downloaded_bytes, persisted_bytes)

    def test_source_equals_download_after_add_with_photo(self) -> None:
        """E: after Contact Add + photo, persisted source == downloaded
        bytes."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                contact = add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
                replace_admin_contact_photo(source_directory, DOCUMENT_ID, contact.contact_id, data=_make_png(64, 64, (10, 20, 30)), content_type='image/png', client=client)
            persisted_path = source_directory / 'GB.docx'
            download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(self._sha(download.path), self._sha(persisted_path))

    def test_source_equals_download_after_update(self) -> None:
        """F: after Contact Update, persisted source == downloaded
        bytes."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                contact = add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
                update_contact(document_id=DOCUMENT_ID, contact_id=contact.contact_id, fields=_write_request(member_firm='Updated Firm'), source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(self._sha(download.path), self._sha(persisted_path))

    def test_source_equals_download_after_photo_replacement(self) -> None:
        """G: after photo replacement, persisted source == downloaded
        bytes."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                contact = add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
                replace_admin_contact_photo(source_directory, DOCUMENT_ID, contact.contact_id, data=_make_png(64, 64, (10, 20, 30)), content_type='image/png', client=client)
                replace_admin_contact_photo(source_directory, DOCUMENT_ID, contact.contact_id, data=_make_png(32, 96, (200, 100, 50)), content_type='image/png', client=client)
            persisted_path = source_directory / 'GB.docx'
            download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(self._sha(download.path), self._sha(persisted_path))

    def test_source_equals_download_after_delete(self) -> None:
        """H: after Contact Delete, persisted source == downloaded
        bytes."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                first = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm B'), source_directory=source_directory, client=client)
                delete_contact(document_id=DOCUMENT_ID, contact_id=first.contact_id, source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(self._sha(download.path), self._sha(persisted_path))

class UpdateContactPersistsToSourceTests(unittest.TestCase):
    """
    Regression coverage for a real bug: update_contact() used to only
    write ContactState/OpenSearch, never calling
    _synchronize_source_document() the way add_contact()/
    delete_contact() both already did - so a real Admin "Update
    Contact" text edit reported success and the ContactState/list
    endpoints reflected the new value, but the actual persisted (and
    therefore downloadable) source DOCX silently kept serving the OLD
    value until some unrelated later mutation happened to trigger a
    fresh rebuild. This was masked before download became a pure byte
    read (see DownloadByteStabilityTests above) - the OLD download path
    used to call materialize_effective_docx() on every GET, which
    rebuilt fresh from ContactState regardless, hiding the gap.

    test_source_equals_download_after_update above only proves
    download reads exactly what's on disk - it does NOT prove the
    disk copy actually reflects the update, since download and the
    persisted file are the same bytes by construction even when
    update_contact never wrote anything new. These tests inspect the
    actual DOCX content instead of only comparing two reads of
    whatever is already there.
    """

    def _sha(self, path: Path) -> str:
        return hashlib.sha256(path.read_bytes()).hexdigest()

    def test_update_persists_new_value_into_source_docx_and_download(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                contact = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
                update_contact(document_id=DOCUMENT_ID, contact_id=contact.contact_id, fields=_write_request(member_firm='Firm B'), source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            persisted_contacts = extract_contacts_from_docx(persisted_path, country=None)
            self.assertEqual(1, len(persisted_contacts))
            self.assertEqual('Firm B', persisted_contacts[0].member_firm)
            self.assertNotEqual('Firm A', persisted_contacts[0].member_firm, 'the old value must no longer be the effective contact value in the persisted document')
            download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(self._sha(download.path), self._sha(persisted_path))
            downloaded_contacts = extract_contacts_from_docx(download.path, country=None)
            self.assertEqual(1, len(downloaded_contacts))
            self.assertEqual('Firm B', downloaded_contacts[0].member_firm)

    def test_repeated_updates_persist_the_final_value(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                contact = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
                update_contact(document_id=DOCUMENT_ID, contact_id=contact.contact_id, fields=_write_request(member_firm='Firm B'), source_directory=source_directory, client=client)
                update_contact(document_id=DOCUMENT_ID, contact_id=contact.contact_id, fields=_write_request(member_firm='Firm C'), source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            persisted_contacts = extract_contacts_from_docx(persisted_path, country=None)
            self.assertEqual(1, len(persisted_contacts))
            self.assertEqual('Firm C', persisted_contacts[0].member_firm)
            download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            downloaded_contacts = extract_contacts_from_docx(download.path, country=None)
            self.assertEqual('Firm C', downloaded_contacts[0].member_firm)

    def test_update_persists_when_contact_has_no_existing_photo(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                contact = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A', phone='+1 555 0100'), source_directory=source_directory, client=client)
                update_contact(document_id=DOCUMENT_ID, contact_id=contact.contact_id, fields=_write_request(member_firm='Firm A', phone='+1 555 9999'), source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            persisted_contacts = extract_contacts_from_docx(persisted_path, country=None)
            self.assertEqual(1, len(persisted_contacts))
            self.assertEqual('+1 555 9999', persisted_contacts[0].phone)
            photos = extract_contact_photo_candidates(persisted_path)
            self.assertEqual(0, len(photos))

    def test_update_persists_when_contact_already_has_a_photo(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                contact = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
                replace_admin_contact_photo(source_directory, DOCUMENT_ID, contact.contact_id, data=_make_png(64, 64, (10, 20, 30)), content_type='image/png', client=client)
                update_contact(document_id=DOCUMENT_ID, contact_id=contact.contact_id, fields=_write_request(member_firm='Firm A Updated'), source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            persisted_contacts = extract_contacts_from_docx(persisted_path, country=None)
            self.assertEqual(1, len(persisted_contacts))
            self.assertEqual('Firm A Updated', persisted_contacts[0].member_firm)
            photos = extract_contact_photo_candidates(persisted_path)
            self.assertEqual(1, len(photos), "a text-only update must not drop the contact's already-attached photo")
            download = get_document_download(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(self._sha(download.path), self._sha(persisted_path))

class ContactRollbackTests(unittest.TestCase):

    def test_index_failure_restores_previous_state_and_marker(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                first = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
            previous_marker = is_admin_modified_since_upload(source_directory, DOCUMENT_ID)
            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(AdminContactMutationFailedError):
                    add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm B'), source_directory=source_directory, client=client)
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].contact_id, first.contact_id)
            self.assertEqual(is_admin_modified_since_upload(source_directory, DOCUMENT_ID), previous_marker)

    def test_state_write_failure_leaves_index_and_marker_untouched(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client), patch('app.services.admin_contacts.write_contact_state_atomic', side_effect=OSError('disk full')):
                with self.assertRaises(AdminContactMutationFailedError):
                    add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            self.assertIsNone(read_contact_state(source_directory, DOCUMENT_ID))
            self.assertEqual(client.chunks, {})
            self.assertFalse(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_update_index_failure_restores_previous_source_and_state(self) -> None:
        """Mirrors test_index_failure_restores_previous_state_and_marker
        above, for update_contact()'s own new source-DOCX
        synchronization: if the ContactState/index commit fails AFTER
        the DOCX has already been rebuilt with the new value, the DOCX
        must be restored to its exact prior bytes, not left holding
        the new value while ContactState still reports the old one."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                contact = add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Firm A'), source_directory=source_directory, client=client)
            persisted_path = source_directory / 'GB.docx'
            bytes_before = persisted_path.read_bytes()
            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(AdminContactMutationFailedError):
                    update_contact(document_id=DOCUMENT_ID, contact_id=contact.contact_id, fields=_write_request(member_firm='Firm B'), source_directory=source_directory, client=client)
            self.assertEqual(bytes_before, persisted_path.read_bytes())
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(state.contacts[0].member_firm, 'Firm A')

    def test_rollback_itself_failing_raises_rollback_error(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)
            with _patched_indexer(client, fail_bulk=True):
                with patch('app.services.admin_contacts.write_contact_state_atomic', side_effect=[None, OSError('rollback also fails')]):
                    with self.assertRaises(AdminDocumentRollbackError):
                        add_contact(document_id=DOCUMENT_ID, fields=_write_request(), source_directory=source_directory, client=client)

class ContactPhotoTransactionTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.source_directory = Path(self.temp.name)
        self.document_id = 'doc-1'
        self.country_code = 'BE'
        self.old_photo = write_contact_photo_atomic(self.source_directory, 'contact-old', data=b'old-photo', content_type='image/jpeg')
        self.old_record = ContactRecord(contact_id='contact-old', member_firm='Old Firm', contact_person='Old Person', email='old@example.com', photo_filename=self.old_photo.filename, photo_content_type=self.old_photo.content_type, photo_sha256=self.old_photo.sha256)
        write_contact_state_atomic(self.source_directory, ContactState(document_id=self.document_id, country_code=self.country_code, contacts=(self.old_record,)))
        self.new_photo = write_contact_photo_atomic(self.source_directory, 'contact-new', data=b'new-photo', content_type='image/png')
        self.new_record = ContactRecord(contact_id='contact-new', member_firm='New Firm', contact_person='New Person', email='new@example.com', photo_filename=self.new_photo.filename, photo_content_type=self.new_photo.content_type, photo_sha256=self.new_photo.sha256)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def _photo_path(self, filename: str) -> Path:
        return self.source_directory / '.admin-state' / 'contact-photos' / filename

    def _common_patches(self):
        return (patch.object(admin_contacts, '_document_metadata_for_chunks', return_value={}), patch.object(admin_contacts, 'build_contact_chunk_for_contacts', return_value=None), patch.object(admin_contacts, 'is_admin_modified_since_upload', return_value=True), patch.object(admin_contacts, 'write_admin_modified_marker'))

    def _apply(self) -> None:
        admin_contacts._apply_contact_state_change(document_id=self.document_id, country_code=self.country_code, source_directory=self.source_directory, new_contacts=(self.new_record,), document_metadata={}, client=object(), reset_marker=True)

    def _assert_rolled_back(self) -> None:
        state = read_contact_state(self.source_directory, self.document_id)
        self.assertIsNotNone(state)
        self.assertEqual(['contact-old'], [item.contact_id for item in state.contacts])
        self.assertTrue(self._photo_path(self.old_photo.filename).is_file())
        self.assertFalse(self._photo_path(self.new_photo.filename).exists())

    def test_success_keeps_new_photo_and_removes_superseded_photo(self) -> None:
        metadata_patch, chunk_patch, marker_read_patch, marker_write_patch = self._common_patches()
        with metadata_patch, chunk_patch, marker_read_patch, marker_write_patch, patch.object(admin_contacts, 'replace_document_contact_chunk'), patch.object(admin_contacts, 'reset_admin_modified'):
            self._apply()
        state = read_contact_state(self.source_directory, self.document_id)
        self.assertEqual(['contact-new'], [item.contact_id for item in state.contacts])
        self.assertTrue(self._photo_path(self.new_photo.filename).is_file())
        self.assertFalse(self._photo_path(self.old_photo.filename).exists())

    def test_various_failure_points_roll_back_the_new_photo(self) -> None:
        """Whichever step of _apply_contact_state_change fails first -
        the OpenSearch chunk sync, the admin-modified marker, or the
        ContactState sidecar write itself - the new contact's own
        just-written photo file must be rolled back and the prior
        contact/photo restored, never left half-applied."""
        metadata_patch, chunk_patch, marker_read_patch, marker_write_patch = self._common_patches()
        real_write = admin_contacts.write_contact_state_atomic
        call_counts = {'write': 0, 'replace': 0}

        def fail_first_write(source_directory, state):
            call_counts['write'] += 1
            if call_counts['write'] == 1:
                raise OSError('sidecar boom')
            return real_write(source_directory, state)

        def fail_first_replace(**kwargs):
            call_counts['replace'] += 1
            if call_counts['replace'] == 1:
                raise RuntimeError('opensearch boom')
        scenarios = {'opensearch_failure': [patch.object(admin_contacts, 'replace_document_contact_chunk', side_effect=fail_first_replace)], 'marker_failure': [patch.object(admin_contacts, 'replace_document_contact_chunk'), patch.object(admin_contacts, 'reset_admin_modified', side_effect=OSError('marker boom'))], 'sidecar_failure': [patch.object(admin_contacts, 'write_contact_state_atomic', side_effect=fail_first_write), patch.object(admin_contacts, 'replace_document_contact_chunk')]}
        for case_name, extra_patches in scenarios.items():
            with self.subTest(case=case_name):
                call_counts['write'] = 0
                call_counts['replace'] = 0
                with contextlib.ExitStack() as stack:
                    stack.enter_context(metadata_patch)
                    stack.enter_context(chunk_patch)
                    stack.enter_context(marker_read_patch)
                    stack.enter_context(marker_write_patch)
                    for extra_patch in extra_patches:
                        stack.enter_context(extra_patch)
                    with self.assertRaises(admin_contacts.AdminContactMutationFailedError):
                        self._apply()
                self._assert_rolled_back()

class ContactPhotoCrudPreservationTests(unittest.TestCase):
    """
    A unit-level check of update_contact()'s own field-construction
    logic - _apply_contact_state_change is mocked away entirely, so
    this proves update_contact() itself always carries the existing
    photo_filename/photo_content_type/photo_sha256 forward onto the
    updated record, independent of the full DOCX-synchronization
    integration UpdateContactPersistsToSourceTests proves above.
    """

    def test_business_update_preserves_existing_photo_metadata(self) -> None:
        with tempfile.TemporaryDirectory() as temp:
            source_directory = Path(temp)
            _seed_placeholder_source_docx(source_directory, 'BE.docx')
            document_id = 'doc_' + 'a' * 64
            contact_id = 'contact-123'
            photo = write_contact_photo_atomic(source_directory, contact_id, data=_make_png(64, 64, (10, 20, 30)), content_type='image/png')
            original = ContactRecord(contact_id=contact_id, member_firm='Firm', contact_person='Jane Doe', email='jane@example.com', phone='+32 OLD', address='Old address', website='example.com', photo_filename=photo.filename, photo_content_type=photo.content_type, photo_sha256=photo.sha256)
            write_contact_state_atomic(source_directory, ContactState(document_id=document_id, country_code='BE', contacts=(original,)))
            fields = AdminContactWriteRequest(member_firm='Firm', contact_person='Jane Doe', email='jane@example.com', phone='+32 111 0200', address='New address', website='www.example.com')
            with patch.object(admin_contacts, '_get_document_metadata', return_value={'country_code': 'BE'}), patch.object(admin_contacts, '_load_country_code_and_metadata', return_value=('BE', {'country': 'Belgium', 'source_filename': 'BE.docx'})), patch.object(admin_contacts, 'country_lock', return_value=nullcontext()), patch.object(admin_contacts, '_apply_contact_state_change') as apply_mock:
                admin_contacts.update_contact(document_id=document_id, contact_id=contact_id, fields=fields, source_directory=source_directory, client=object())
            new_contacts = apply_mock.call_args.kwargs['new_contacts']
            self.assertEqual(1, len(new_contacts))
            updated = new_contacts[0]
            self.assertEqual('+32 111 0200', updated.phone)
            self.assertEqual('New address', updated.address)
            self.assertEqual(contact_id, updated.contact_id)
            self.assertEqual(photo.filename, updated.photo_filename)
            self.assertEqual(photo.content_type, updated.photo_content_type)
            self.assertEqual(photo.sha256, updated.photo_sha256)

def _fake_document_lister(documents: list[AdminDocumentSummary]):

    def lister(*, source_directory: Path, client) -> AdminDocumentListResponse:
        del source_directory, client
        return AdminDocumentListResponse(total=len(documents), documents=documents)
    return lister

def _summary(*, document_id: str, country_code: str, source_filename: str) -> AdminDocumentSummary:
    return AdminDocumentSummary(document_id=document_id, source_filename=source_filename, country=country_code, country_code=country_code, language='en', document_type='overview', reference_year=2026, chunk_count=1, source_file_present=True, status='indexed')

class ContactPhotoDeleteSyncTests(unittest.TestCase):
    """
    Deleting a contact must remove ONLY that contact's own photo from
    the persisted source DOCX, using the SAME deterministic SHA-256
    association already proven at the raw-primitive layer in
    test_contact_documents.py - never position, filename, or a second,
    unrelated image-deletion implementation.
    """

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.document_id = _real_document_id_for('AU')

    def _require_copy(self, filename: str) -> Path:
        return require_corpus_copy(self, _test_admin_contacts__SOURCE_ROOT, filename, self.root)

    def _seed_two_contacts_with_photos(self) -> tuple[str, str]:
        """
        A real AU.docx with Michael Harmer's own real photo, plus a
        second contact added the same way Add Contact would (via
        rebuild_canonical_contact_table), each contact's own photo
        embedded for real, independently. Returns (michael_sha256,
        second_sha256).
        """
        docx_path = self._require_copy('AU.docx')
        michael_photo = extract_contact_photo_candidates(docx_path)[0]
        second_photo_data = _make_png(183, 234, (80, 120, 200))
        jane_contact = ExtractedContact(member_firm='Second Firm Pty Ltd', contact_person='Jane Secondary', email='jane.secondary@secondfirm.com.au', phone='+61 2 9000 0000', address='100 Test Street, Level 5', website='www.secondfirm.com.au')
        new_bytes = rebuild_canonical_contact_table(docx_path, contacts=(ExtractedContact(member_firm='HARMERS WORKPLACE LAWYERS', contact_person='Michael Harmer', email='michael.harmer@harmers.com.au', phone='+61 292 674 322', address='31 Market Street, Level 27 St Martins Tower', website='www.harmers.com.au'), jane_contact), photos=(ContactPhotoPayload(data=michael_photo.data, content_type=michael_photo.content_type), ContactPhotoPayload(data=second_photo_data, content_type='image/png')), country='Australia')
        docx_path.write_bytes(new_bytes)
        michael_stored = write_contact_photo_atomic(self.root, 'michael-id', data=michael_photo.data, content_type=michael_photo.content_type)
        second_stored = write_contact_photo_atomic(self.root, 'jane-id', data=second_photo_data, content_type='image/png')
        write_contact_state_atomic(self.root, ContactState(document_id=self.document_id, country_code='AU', contacts=(ContactRecord(contact_id='michael-id', member_firm='HARMERS WORKPLACE LAWYERS', contact_person='Michael Harmer', email='michael.harmer@harmers.com.au', phone='+61 292 674 322', address='31 Market Street, Level 27 St Martins Tower', website='www.harmers.com.au', photo_filename=michael_stored.filename, photo_content_type=michael_stored.content_type, photo_sha256=michael_stored.sha256), ContactRecord(contact_id='jane-id', member_firm='Second Firm Pty Ltd', contact_person='Jane Secondary', email='jane.secondary@secondfirm.com.au', phone='+61 2 9000 0000', address='100 Test Street, Level 5', website='www.secondfirm.com.au', photo_filename=second_stored.filename, photo_content_type=second_stored.content_type, photo_sha256=second_stored.sha256))))
        return (michael_stored.sha256, second_stored.sha256)

    def test_delete_contact_with_photo_removes_only_its_own_image(self) -> None:
        docx_path = self.root / 'AU.docx'
        michael_sha, jane_sha = self._seed_two_contacts_with_photos()
        client = FakeContactOpenSearchClient(document_id=self.document_id, country_code='AU', country='Australia', source_filename='AU.docx')
        with _patched_indexer(client):
            delete_contact(document_id=self.document_id, contact_id='jane-id', source_directory=self.root, client=client)
        remaining_shas = {c.sha256 for c in extract_contact_photo_candidates(docx_path)}
        self.assertIn(michael_sha, remaining_shas, "the other contact's own photo must remain byte/functionally intact")
        self.assertNotIn(jane_sha, remaining_shas, 'no stale photo of the deleted contact may remain anywhere in the DOCX')
        self.assertEqual(1, len(remaining_shas))
        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(1, len(state.contacts))
        self.assertEqual('michael-id', state.contacts[0].contact_id)

    def test_delete_mutation_failure_restores_source_docx_byte_for_byte(self) -> None:
        docx_path = self.root / 'AU.docx'
        self._seed_two_contacts_with_photos()
        original_bytes = docx_path.read_bytes()
        client = FakeContactOpenSearchClient(document_id=self.document_id, country_code='AU', country='Australia', source_filename='AU.docx')
        with self.assertRaises(AdminContactMutationFailedError):
            with _patched_indexer(client, fail_bulk=True):
                delete_contact(document_id=self.document_id, contact_id='jane-id', source_directory=self.root, client=client)
        self.assertEqual(original_bytes, docx_path.read_bytes(), 'the source DOCX must be restored to its exact original bytes when the ContactState/index commit fails after the photo was already removed')
        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(2, len(state.contacts))

    def test_delete_contact_without_a_photo_still_removes_its_text(self) -> None:
        """A contact with no photo_sha256 at all must still have its
        own TEXT block removed from the persisted source DOCX -
        deleting a contact must never leave its name/email/firm behind
        just because it never had a photo."""
        docx_path = self._require_copy('AU.docx')
        write_contact_state_atomic(self.root, ContactState(document_id=self.document_id, country_code='AU', contacts=(ContactRecord(contact_id='michael-id', member_firm='HARMERS WORKPLACE LAWYERS', contact_person='Michael Harmer', email='michael.harmer@harmers.com.au', phone='+61 292 674 322', address='31 Market Street', website='www.harmers.com.au'),)))
        client = FakeContactOpenSearchClient(document_id=self.document_id, country_code='AU', country='Australia', source_filename='AU.docx')
        with _patched_indexer(client):
            delete_contact(document_id=self.document_id, contact_id='michael-id', source_directory=self.root, client=client)
        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')
        self.assertNotIn('Michael Harmer', document_xml)
        self.assertNotIn('michael.harmer@harmers.com.au', document_xml)
        remaining = extract_contacts_from_docx(docx_path, country='Australia')
        self.assertEqual(0, len(remaining))

    def _seed_contact_a_only(self, docx_path: Path) -> str:
        """Contact A (Michael Harmer) exactly as AU.docx's own native
        organic block already reads, with his own real embedded photo
        registered in ContactState - the document's starting state
        before an Admin-added contact B exists at all."""
        michael_photo = extract_contact_photo_candidates(docx_path)[0]
        michael_stored = write_contact_photo_atomic(self.root, 'michael-id', data=michael_photo.data, content_type=michael_photo.content_type)
        write_contact_state_atomic(self.root, ContactState(document_id=self.document_id, country_code='AU', contacts=(ContactRecord(contact_id='michael-id', member_firm='HARMERS WORKPLACE LAWYERS', contact_person='Michael Harmer', email='michael.harmer@harmers.com.au', phone='+61 292 674 322', address='31 Market Street', website='www.harmers.com.au', photo_filename=michael_stored.filename, photo_content_type=michael_stored.content_type, photo_sha256=michael_stored.sha256),)))
        return michael_stored.sha256

    def _add_contact_b_with_photo_via_real_admin_flow(self, client: FakeContactOpenSearchClient) -> tuple[str, bytes]:
        """Exercises the REAL, two-call Admin surface exactly as an
        operator would use it - add_contact() (text only) followed by
        replace_admin_contact_photo() (attaches the photo, rebuilding
        the canonical table again) - never the lower-level
        rebuild_canonical_contact_table() primitive directly.
        Returns (contact_b_id, photo_b_bytes)."""
        jane_photo_data = _make_png(183, 234, (80, 120, 200))
        with _patched_indexer(client):
            response = add_contact(document_id=self.document_id, fields=_write_request(member_firm='Second Firm Pty Ltd', contact_person='Jane Secondary', email='jane.secondary@secondfirm.com.au', phone='+61 2 9000 0000', address='100 Test Street, Level 5', website='www.secondfirm.com.au'), source_directory=self.root, client=client)
        with _patched_indexer(client):
            replace_admin_contact_photo(self.root, self.document_id, response.contact_id, data=jane_photo_data, content_type='image/png', client=client)
        return (response.contact_id, jane_photo_data)

    def test_admin_added_contact_b_fully_removed_from_source_on_delete(self) -> None:
        """The user's mandatory focused test: ADD contact B + photo B
        (via the real two-call Admin surface, onto a document that
        already carries contact A natively), then DELETE contact B,
        then prove directly against the PERSISTED source DOCX that
        contact B's text and photo are gone while contact A's text and
        photo remain - byte-level proof, not ContactState alone."""
        docx_path = self._require_copy('AU.docx')
        michael_sha = self._seed_contact_a_only(docx_path)
        client = FakeContactOpenSearchClient(document_id=self.document_id, country_code='AU', country='Australia', source_filename='AU.docx')
        contact_b_id, jane_photo_data = self._add_contact_b_with_photo_via_real_admin_flow(client)
        jane_sha = hashlib.sha256(jane_photo_data).hexdigest()
        with zipfile.ZipFile(docx_path) as archive:
            pre_delete_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')
        self.assertIn('Jane Secondary', pre_delete_xml)
        pre_delete_shas = {c.sha256 for c in extract_contact_photo_candidates(docx_path)}
        self.assertIn(jane_sha, pre_delete_shas)
        with _patched_indexer(client):
            delete_contact(document_id=self.document_id, contact_id=contact_b_id, source_directory=self.root, client=client)
        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')
        remaining_shas = {c.sha256 for c in extract_contact_photo_candidates(docx_path)}
        contact_b_text_in_source = 'Jane Secondary' in document_xml or 'jane.secondary@secondfirm.com.au' in document_xml
        contact_b_photo_in_source = jane_sha in remaining_shas
        contact_a_text_in_source = 'Michael Harmer' in document_xml and 'michael.harmer@harmers.com.au' in document_xml
        contact_a_photo_in_source = michael_sha in remaining_shas
        print(f"CONTACT_B_TEXT_IN_SOURCE={('YES' if contact_b_text_in_source else 'NO')}")
        print(f"CONTACT_B_PHOTO_IN_SOURCE={('YES' if contact_b_photo_in_source else 'NO')}")
        print(f"CONTACT_A_TEXT_IN_SOURCE={('YES' if contact_a_text_in_source else 'NO')}")
        print(f"CONTACT_A_PHOTO_IN_SOURCE={('YES' if contact_a_photo_in_source else 'NO')}")
        self.assertFalse(contact_b_text_in_source, "deleted contact B's own text must not survive in the persisted source DOCX")
        self.assertFalse(contact_b_photo_in_source, "deleted contact B's own photo must not survive in the persisted source DOCX")
        self.assertTrue(contact_a_text_in_source, "surviving contact A's own text must remain untouched")
        self.assertTrue(contact_a_photo_in_source, "surviving contact A's own photo must remain untouched")
        remaining_contacts = extract_contacts_from_docx(docx_path, country='Australia')
        self.assertEqual(1, len(remaining_contacts))
        self.assertEqual('Michael Harmer', remaining_contacts[0].contact_person)

    def test_admin_added_contact_b_delete_failure_restores_source_byte_for_byte(self) -> None:
        """Same admin-added contact B scenario, but the ContactState/
        index commit fails after the source DOCX has already been
        rewritten to remove contact B's text+photo - the source must
        be restored to its EXACT prior bytes (with contact B still
        present), never left half-mutated."""
        docx_path = self._require_copy('AU.docx')
        self._seed_contact_a_only(docx_path)
        client = FakeContactOpenSearchClient(document_id=self.document_id, country_code='AU', country='Australia', source_filename='AU.docx')
        contact_b_id, _ = self._add_contact_b_with_photo_via_real_admin_flow(client)
        original_bytes = docx_path.read_bytes()
        with self.assertRaises(AdminContactMutationFailedError):
            with _patched_indexer(client, fail_bulk=True):
                delete_contact(document_id=self.document_id, contact_id=contact_b_id, source_directory=self.root, client=client)
        self.assertEqual(original_bytes, docx_path.read_bytes(), "the source DOCX must be restored to its exact original bytes (contact B's text and photo both still present) when the ContactState/index commit fails after the source rewrite already succeeded")
        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(2, len(state.contacts))
        self.assertIn(contact_b_id, {c.contact_id for c in state.contacts})

class AddContactFallbackSynchronizationTests(unittest.TestCase):
    """
    A document with no detected contact area at all (PT.docx - proven
    structure-less by test_contact_documents.py's own
    test_no_matching_contact_area_still_produces_canonical_table) must
    never leave a successful Add Contact as ContactState-only.
    add_contact() must fall back to contact_document_area.rebuild_
    canonical_contact_table()'s own no-existing-area handling
    (_default_insertion_anchor) and commit that to the source - SOURCE
    DOCX == CONTACT STATE == OPENSEARCH holds even here.
    """

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.document_id = _real_document_id_for('PT')

    def _require_copy(self, filename: str) -> Path:
        return require_corpus_copy(self, _test_admin_contacts__SOURCE_ROOT, filename, self.root)

    def _client(self) -> FakeContactOpenSearchClient:
        return FakeContactOpenSearchClient(document_id=self.document_id, country_code='PT', country='Portugal', source_filename='PT.docx')

    def test_add_contact_without_structural_area_still_synchronizes_source(self) -> None:
        docx_path = self._require_copy('PT.docx')
        original_bytes = docx_path.read_bytes()
        client = self._client()
        with _patched_indexer(client):
            add_contact(document_id=self.document_id, fields=_write_request(member_firm='Someone New Lda', contact_person='Someone New', email='new@example.test', phone='+351 21 000 0000', address='Rua Nova 1', website='www.newfirm.test'), source_directory=self.root, client=client)
        new_bytes = docx_path.read_bytes()
        self.assertNotEqual(original_bytes, new_bytes, 'a successful Add Contact must always rewrite the persisted source DOCX, even when the document has no two-box contact area to clone the style of - ContactState alone is never sufficient')
        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')
        self.assertIn('Someone New', document_xml)
        self.assertIn('new@example.test', document_xml)
        self.assertIn(CONTACT_TABLE_HIDDEN_MARKER, document_xml)
        self.assertIn('<w:tbl>', document_xml)
        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(1, len(state.contacts))

    def test_second_add_replaces_rather_than_duplicates_the_fallback_block(self) -> None:
        """Two sequential Add Contact calls against the same
        structure-less document must each leave the source
        synchronized with the FULL current contact list - never
        stacking a second, stale fallback block alongside the first
        (which would silently resurrect a deleted/superseded
        rendering of the earlier contact)."""
        docx_path = self._require_copy('PT.docx')
        client = self._client()
        with _patched_indexer(client):
            add_contact(document_id=self.document_id, fields=_write_request(member_firm='First Firm Lda', contact_person='First Person', email='first@example.test'), source_directory=self.root, client=client)
        with _patched_indexer(client):
            add_contact(document_id=self.document_id, fields=_write_request(member_firm='Second Firm Lda', contact_person='Second Person', email='second@example.test'), source_directory=self.root, client=client)
        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8', errors='ignore')
        self.assertIn('First Person', document_xml)
        self.assertIn('Second Person', document_xml)
        self.assertEqual(1, document_xml.count(CONTACT_TABLE_HIDDEN_MARKER), 'a second Add Contact must replace the one canonical table with a fresh rendering of the full contact list, never append a second one')
        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(2, len(state.contacts))

class ContactBootstrapTests(unittest.TestCase):

    def test_dry_run_does_not_write_anything(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'GB.docx').write_bytes(b'docx')
            summaries = [_summary(document_id=DOCUMENT_ID, country_code='GB', source_filename='GB.docx')]
            with patch('app.services.admin_contacts.extract_contacts_from_docx', return_value=[ExtractedContact(member_firm='Firm A')]):
                report = bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=True, document_lister=_fake_document_lister(summaries))
            self.assertTrue(report.dry_run)
            self.assertEqual(report.documents_seen, 1)
            self.assertEqual(report.contacts_seeded, 1)
            self.assertIsNone(read_contact_state(source_directory, DOCUMENT_ID))

    def test_wet_run_seeds_state_for_new_documents(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'GB.docx').write_bytes(b'docx')
            summaries = [_summary(document_id=DOCUMENT_ID, country_code='GB', source_filename='GB.docx')]
            with patch('app.services.admin_contacts.extract_contacts_from_docx', return_value=[ExtractedContact(member_firm='Firm A')]):
                report = bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=False, document_lister=_fake_document_lister(summaries))
            self.assertEqual(report.contacts_seeded, 1)
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].member_firm, 'Firm A')

    def test_zero_contact_document_gets_explicit_empty_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'FR.docx').write_bytes(b'docx')
            summaries = [_summary(document_id=OTHER_DOCUMENT_ID, country_code='FR', source_filename='FR.docx')]
            with patch('app.services.admin_contacts.extract_contacts_from_docx', return_value=[]):
                report = bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=False, document_lister=_fake_document_lister(summaries))
            self.assertEqual(report.zero_contact_documents, 1)
            self.assertEqual(report.contacts_seeded, 0)
            state = read_contact_state(source_directory, OTHER_DOCUMENT_ID)
            self.assertIsNotNone(state)
            self.assertEqual(state.contacts, ())

    def test_existing_state_is_never_overwritten(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'GB.docx').write_bytes(b'docx')
            existing_record = _full_contact_record(member_firm='Admin Edited Firm')
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=(existing_record,)))
            summaries = [_summary(document_id=DOCUMENT_ID, country_code='GB', source_filename='GB.docx')]
            with patch('app.services.admin_contacts.extract_contacts_from_docx', return_value=[ExtractedContact(member_firm='Stale DOCX Firm')]):
                report = bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=False, document_lister=_fake_document_lister(summaries))
            self.assertEqual(report.documents_skipped_existing_state, 1)
            self.assertEqual(report.contacts_seeded, 0)
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(state.contacts[0].member_firm, 'Admin Edited Firm')

    def test_legacy_incomplete_contacts_preserved_exactly(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'GB.docx').write_bytes(b'docx')
            summaries = [_summary(document_id=DOCUMENT_ID, country_code='GB', source_filename='GB.docx')]
            with patch('app.services.admin_contacts.extract_contacts_from_docx', return_value=[ExtractedContact(contact_person='Alex Example', email='alex@example.test')]):
                bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=False, document_lister=_fake_document_lister(summaries))
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertIsNone(state.contacts[0].member_firm)
            self.assertIsNone(state.contacts[0].phone)
            self.assertEqual(state.contacts[0].contact_person, 'Alex Example')

class AdminContactWriteRequestOptionalFieldTests(unittest.TestCase):
    """Section 11/12: every one of the six business fields is
    individually optional; the only validation is a cross-field "at
    least one field has a value" rule - a real member-firm contact
    (France's own Caroline Scherrmann) can genuinely have address/
    website empty, and that must never be rejected."""

    def test_every_field_individually_empty_is_accepted(self) -> None:
        for field in ('member_firm', 'contact_person', 'email', 'phone', 'address', 'website'):
            with self.subTest(field=field):
                fields = {'member_firm': 'Firm', 'contact_person': 'Person', 'email': 'person@example.com', 'phone': '+1 555 0100', 'address': '1 Example Street', 'website': 'www.example.com'}
                fields[field] = ''
                request = AdminContactWriteRequest(**fields)
                self.assertEqual('', getattr(request, field))

    def test_website_and_address_both_empty_is_accepted(self) -> None:
        """The exact real France shape."""
        request = AdminContactWriteRequest(member_firm='Flichy Grangé Avocats', contact_person='Caroline Scherrmann', email='scherrmann@flichy.com', phone='+33 1 56 62 30 00', address='', website='')
        self.assertEqual('', request.address)
        self.assertEqual('', request.website)

    def test_all_six_fields_blank_is_rejected(self) -> None:
        with self.assertRaises(ValidationError):
            AdminContactWriteRequest(member_firm='', contact_person='', email='', phone='', address='', website='')

    def test_all_six_fields_whitespace_only_is_rejected(self) -> None:
        with self.assertRaises(ValidationError):
            AdminContactWriteRequest(member_firm='   ', contact_person='', email='  ', phone='', address='', website='')

    def test_a_single_filled_field_is_sufficient(self) -> None:
        request = AdminContactWriteRequest(member_firm='', contact_person='Solo Person', email='', phone='', address='', website='')
        self.assertEqual('Solo Person', request.contact_person)

class FranceLegacyBootstrapSplitTests(unittest.TestCase):
    """Sections 7-9: France's real legacy contact - a single combined
    record naming two people - splits into two Admin-managed
    ContactRecords during bootstrap, with stable ids and Jessica Stout
    (the project-level L&E Global POC, never a member-firm contact)
    correctly excluded."""

    def _require_copy(self, filename: str) -> Path:
        source = _test_admin_contacts__SOURCE_ROOT / filename
        if not source.exists():
            self.skipTest(f'Real corpus source unavailable: {source}')
        document = Document(str(source))
        if any((table.rows and CONTACT_TABLE_HIDDEN_MARKER in table.rows[0].cells[0].text for table in document.tables)):
            self.skipTest(f'{filename} has since been canonicalized by real Admin usage (real corpus content has drifted since this test was written) - its legacy combined contact no longer exists in raw form to bootstrap from')
        return source

    def test_france_bootstrap_yields_two_stable_contacts_without_jessica(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            real_source = self._require_copy('FR.docx')
            (source_directory / 'FR.docx').write_bytes(real_source.read_bytes())
            document_id = 'doc_' + 'f' * 64
            summaries = [_summary(document_id=document_id, country_code='FR', source_filename='FR.docx')]
            report = bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=False, document_lister=_fake_document_lister(summaries))
            self.assertEqual(2, report.contacts_seeded)
            state = read_contact_state(source_directory, document_id)
            self.assertEqual(2, len(state.contacts))
            names = [c.contact_person for c in state.contacts]
            self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet'], names)
            self.assertNotIn('Jessica Stout', names)
            ids = {c.contact_id for c in state.contacts}
            self.assertEqual(2, len(ids), 'each split contact must get its own id')
            for contact in state.contacts:
                self.assertEqual('Flichy Grangé Avocats', contact.member_firm)
                self.assertEqual('+33 1 56 62 30 00', contact.phone)

    def test_repeated_bootstrap_does_not_change_ids_or_duplicate(self) -> None:
        """bootstrap_legacy_contacts never overwrites an existing
        sidecar (see its own docstring) - re-running it after France's
        split has already been persisted must be a complete no-op,
        never regenerating ids or duplicating the two contacts."""
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            real_source = self._require_copy('FR.docx')
            (source_directory / 'FR.docx').write_bytes(real_source.read_bytes())
            document_id = 'doc_' + 'f' * 64
            summaries = [_summary(document_id=document_id, country_code='FR', source_filename='FR.docx')]
            lister = _fake_document_lister(summaries)
            bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=False, document_lister=lister)
            first_ids = [c.contact_id for c in read_contact_state(source_directory, document_id).contacts]
            report = bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=False, document_lister=lister)
            self.assertEqual(1, report.documents_skipped_existing_state)
            self.assertEqual(0, report.contacts_seeded)
            second_state = read_contact_state(source_directory, document_id)
            second_ids = [c.contact_id for c in second_state.contacts]
            self.assertEqual(first_ids, second_ids)
            self.assertEqual(2, len(second_state.contacts))

class ContactBootstrapIsolatedCorpusTests(unittest.TestCase):
    """
    Mirrors a real production audit's own real-corpus baseline: 24
    documents, 22 contacts, 2 zero-contact documents (FR, PT) - here
    reproduced against a synthetic, isolated 24-document catalog (never
    the real production corpus), proving the bootstrap facility
    produces exactly this shape given that input.
    """

    def test_matches_the_audit_baseline_shape(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            summaries = []
            contacts_by_country: dict[str, list[ExtractedContact]] = {}
            for index in range(24):
                country_code = f'C{index:02d}'
                summaries.append(_summary(document_id=f'doc_{index:064x}', country_code=country_code, source_filename=f'{country_code}.docx'))
                (source_directory / f'{country_code}.docx').write_bytes(b'docx')
                contacts_by_country[country_code] = [] if index < 2 else [ExtractedContact(member_firm='Firm')]

            def fake_extract(path: Path, country: str):
                del country
                country_code = path.stem
                return contacts_by_country[country_code]
            with patch('app.services.admin_contacts.extract_contacts_from_docx', side_effect=fake_extract):
                report = bootstrap_legacy_contacts(source_directory=source_directory, client=object(), dry_run=False, document_lister=_fake_document_lister(summaries))
            self.assertEqual(report.documents_seen, 24)
            self.assertEqual(report.contacts_seeded, 22)
            self.assertEqual(report.zero_contact_documents, 2)

def _legal_chunk(document_id: str=DOCUMENT_ID) -> DocumentChunk:
    return DocumentChunk(document_id=document_id, chunk_id='chunk_' + 'c' * 64, country='United Kingdom', country_code='GB', legal_topic='Employment Contracts', document_type='overview', language='en', section='Employment Contracts', subsection=None, content='legal content', source_filename='GB.docx', source_format='docx', content_hash='hash-legal', reference_year=2026)

def _stale_docx_contact_chunk(document_id: str=DOCUMENT_ID) -> DocumentChunk:
    return DocumentChunk(document_id=document_id, chunk_id='chunk_' + 'd' * 64, country='United Kingdom', country_code='GB', legal_topic=None, document_type='overview', language='en', section='Employment Law Overview United Kingdom', subsection='Contact', content='Member firm: Stale DOCX Firm', source_filename='GB.docx', source_format='docx', content_hash='hash-stale', reference_year=2026)

class ApplyStructuredContactStateToChunksTests(unittest.TestCase):

    def test_no_sidecar_leaves_chunks_unchanged(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            chunks = [_legal_chunk(), _stale_docx_contact_chunk()]
            result = apply_structured_contact_state_to_chunks(chunks=list(chunks), document_id=DOCUMENT_ID, source_directory=source_directory)
            self.assertEqual(result, chunks)

    def test_existing_sidecar_replaces_the_stale_docx_contact_chunk(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=(_full_contact_record(member_firm='Admin Edited Firm'),)))
            chunks = [_legal_chunk(), _stale_docx_contact_chunk()]
            result = apply_structured_contact_state_to_chunks(chunks=list(chunks), document_id=DOCUMENT_ID, source_directory=source_directory)
            legal_chunks = [c for c in result if c.subsection != 'Contact']
            contact_chunks = [c for c in result if c.subsection == 'Contact']
            self.assertEqual(len(legal_chunks), 1)
            self.assertEqual(legal_chunks[0].content, 'legal content')
            self.assertEqual(len(contact_chunks), 1)
            self.assertIn('Admin Edited Firm', contact_chunks[0].content)
            self.assertNotIn('Stale DOCX Firm', contact_chunks[0].content)

    def test_existing_empty_sidecar_removes_the_docx_contact_chunk(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=()))
            chunks = [_legal_chunk(), _stale_docx_contact_chunk()]
            result = apply_structured_contact_state_to_chunks(chunks=list(chunks), document_id=DOCUMENT_ID, source_directory=source_directory)
            self.assertEqual(len(result), 1)
            self.assertIsNone(result[0].subsection)

class ReindexPreservesContactStateIntegrationTests(unittest.TestCase):
    """
    Real reindex_indexed_document, proving the wiring (not just the
    pure helper above) actually substitutes the Contact chunk during a
    genuine Refresh.
    """

    def test_refresh_uses_structured_state_not_stale_docx_text(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'GB.docx').write_bytes(b'docx')
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=(_full_contact_record(member_firm='Admin Edited Firm'),)))
            client = FakeContactOpenSearchClient(chunks={'existing-legal': {'document_id': DOCUMENT_ID, 'country_code': 'GB', 'country': 'United Kingdom', 'source_filename': 'GB.docx', 'reference_year': 2026, 'legal_topic': 'Employment Contracts', 'subsection': None, 'content': 'legal content'}})

            def chunk_builder(path: Path) -> list[DocumentChunk]:
                del path
                return [_legal_chunk(), _stale_docx_contact_chunk()]
            captured: dict[str, Any] = {}

            def document_indexer(*, chunks, client=None) -> DocumentIndexingResult:
                del client
                captured['chunks'] = list(chunks)
                return DocumentIndexingResult(index_alias='legal-documents', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
            reindex_indexed_document(document_id=DOCUMENT_ID, source_directory=source_directory, client=client, chunk_builder=chunk_builder, document_indexer=document_indexer)
            indexed_contact_chunks = [c for c in captured['chunks'] if c.subsection == 'Contact']
            self.assertEqual(len(indexed_contact_chunks), 1)
            self.assertIn('Admin Edited Firm', indexed_contact_chunks[0].content)
            self.assertNotIn('Stale DOCX Firm', indexed_contact_chunks[0].content)

class ReseedContactStateFromParsedContactsTests(unittest.TestCase):

    def test_reseed_overwrites_prior_state_entirely(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=(_full_contact_record(member_firm='Old Firm A'), _full_contact_record(member_firm='Old Firm B'), _full_contact_record(member_firm='Old Firm C'))))
            mark_admin_modified(source_directory, DOCUMENT_ID)
            reseed_contact_state_from_parsed_contacts(document_id=DOCUMENT_ID, country_code='GB', source_directory=source_directory, contacts=[ExtractedContact(member_firm='New Firm')])
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].member_firm, 'New Firm')
            self.assertFalse(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_reseed_with_zero_contacts_clears_prior_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            write_contact_state_atomic(source_directory, ContactState(document_id=DOCUMENT_ID, country_code='GB', contacts=(_full_contact_record(member_firm='Old Firm A'), _full_contact_record(member_firm='Old Firm B'))))
            reseed_contact_state_from_parsed_contacts(document_id=DOCUMENT_ID, country_code='GB', source_directory=source_directory, contacts=[])
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertIsNotNone(state)
            self.assertEqual(state.contacts, ())

class ParsedContactPhotoReseedTests(unittest.TestCase):
    """
    reseed_contact_state_from_parsed_contacts's OTHER branch: when
    docx_path is supplied, real corpus photo extraction/association is
    also exercised (see _build_photo_aware_contact_records) - a
    materially different code path than the two tests above, which
    never pass docx_path at all.
    """

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.source_directory = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def _require(self, filename: str) -> Path:
        path = _test_admin_contacts__SOURCE_ROOT / filename
        if not path.exists():
            self.skipTest(f'Corpus file unavailable: {path}')
        skip_if_already_canonicalized(self, path)
        return path

    def _photo_files(self) -> list[Path]:
        directory = self.source_directory / '.admin-state' / 'contact-photos'
        if not directory.exists():
            return []
        return sorted((path for path in directory.iterdir() if path.is_file()))

    def test_parsed_reseed_belgium_creates_two_contacts_and_two_photos(self) -> None:
        path = self._require('Labour and Employment Law in Belgium 2026.docx')
        parsed = extract_contacts_from_docx(path, country='Belgium')
        self.assertEqual(1, len(parsed))
        reseed_contact_state_from_parsed_contacts(document_id='doc-belgium', country_code='BE', source_directory=self.source_directory, contacts=parsed, docx_path=path)
        state = read_contact_state(self.source_directory, 'doc-belgium')
        self.assertIsNotNone(state)
        self.assertEqual(2, len(state.contacts))
        first, second = state.contacts
        self.assertEqual('Chris van Olmen', first.contact_person)
        self.assertEqual('chris.van.olmen@vow.be', first.email)
        self.assertEqual('Nicolas Simon', second.contact_person)
        self.assertEqual('nicolas.simon@vow.be', second.email)
        self.assertNotEqual(first.contact_id, second.contact_id)
        for record in (first, second):
            self.assertIsNotNone(record.photo_filename)
            self.assertIsNotNone(record.photo_content_type)
            self.assertIsNotNone(record.photo_sha256)
            self.assertTrue((self.source_directory / '.admin-state' / 'contact-photos' / record.photo_filename).is_file())
        expected_photos = extract_contact_photo_candidates(path)
        self.assertEqual(expected_photos[0].sha256, first.photo_sha256)
        self.assertEqual(expected_photos[1].sha256, second.photo_sha256)
        self.assertEqual(2, len(self._photo_files()))

    def test_parsed_reseed_france_remains_one_contact_without_photo(self) -> None:
        path = self._require('FR.docx')
        parsed = extract_contacts_from_docx(path, country='France')
        reseed_contact_state_from_parsed_contacts(document_id='doc-france', country_code='FR', source_directory=self.source_directory, contacts=parsed, docx_path=path)
        state = read_contact_state(self.source_directory, 'doc-france')
        self.assertEqual(1, len(state.contacts))
        contact = state.contacts[0]
        self.assertEqual('Caroline Scherrmann and Florence Bacquet', contact.contact_person)
        self.assertIsNone(contact.photo_filename)
        self.assertIsNone(contact.photo_content_type)
        self.assertIsNone(contact.photo_sha256)
        self.assertEqual([], self._photo_files())

    def test_parsed_reseed_indonesia_persists_the_valid_photo(self) -> None:
        path = self._require('ID.docx')
        parsed = extract_contacts_from_docx(path, country='Indonesia')
        expected = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(expected))
        self.assertEqual('image3.jpeg', expected[0].source_filename)
        reseed_contact_state_from_parsed_contacts(document_id='doc-indonesia', country_code='ID', source_directory=self.source_directory, contacts=parsed, docx_path=path)
        state = read_contact_state(self.source_directory, 'doc-indonesia')
        self.assertEqual(1, len(state.contacts))
        contact = state.contacts[0]
        self.assertEqual(expected[0].sha256, contact.photo_sha256)
        self.assertTrue((self.source_directory / '.admin-state' / 'contact-photos' / contact.photo_filename).is_file())

    def test_photo_extraction_failure_keeps_contacts_without_photo(self) -> None:
        path = self._require('FR.docx')
        parsed = extract_contacts_from_docx(path, country='France')
        with patch.object(admin_contacts, 'extract_contact_photo_candidates', side_effect=admin_contacts.ContactPhotoExtractionError('unsupported synthetic DOCX image package')):
            reseed_contact_state_from_parsed_contacts(document_id='doc-photo-fallback', country_code='FR', source_directory=self.source_directory, contacts=parsed, docx_path=path)
        state = read_contact_state(self.source_directory, 'doc-photo-fallback')
        self.assertIsNotNone(state)
        self.assertEqual(1, len(state.contacts))
        contact = state.contacts[0]
        self.assertEqual('Caroline Scherrmann and Florence Bacquet', contact.contact_person)
        self.assertIsNone(contact.photo_filename)
        self.assertIsNone(contact.photo_content_type)
        self.assertIsNone(contact.photo_sha256)
        self.assertEqual([], self._photo_files())

    def test_second_photo_write_failure_leaves_no_partial_seed(self) -> None:
        path = self._require('Labour and Employment Law in Belgium 2026.docx')
        parsed = extract_contacts_from_docx(path, country='Belgium')
        real_write_contact_photo_atomic = write_contact_photo_atomic
        calls = 0

        def fail_second(*args, **kwargs):
            nonlocal calls
            calls += 1
            if calls == 2:
                raise ContactPhotoStorageError('second photo boom')
            return real_write_contact_photo_atomic(*args, **kwargs)
        with patch.object(admin_contacts, 'write_contact_photo_atomic', side_effect=fail_second, create=True):
            with self.assertRaises(ContactPhotoStorageError):
                reseed_contact_state_from_parsed_contacts(document_id='doc-belgium', country_code='BE', source_directory=self.source_directory, contacts=parsed, docx_path=path)
        self.assertIsNone(read_contact_state(self.source_directory, 'doc-belgium'))
        self.assertEqual([], self._photo_files())

class UploadReplaceReseedsContactsIntegrationTests(unittest.TestCase):
    """
    Real safe_upload_and_index_document, proving a genuine
    upload/replace transaction really reseeds structured contact state
    from the newly-accepted DOCX (not merely the lower-level helper
    above in isolation).
    """

    def _run_upload(self, *, source_directory: Path, parsed_contacts: list[ExtractedContact], replace_existing: bool):
        from app.services.document_chunk_builder import DOCUMENT_FAMILY
        del DOCUMENT_FAMILY
        chunk = _legal_chunk()

        def chunk_builder(path: Path) -> list[DocumentChunk]:
            del path
            return [chunk]

        def country_document_lookup(country_code, client):
            del country_code, client
            return []

        def country_document_indexer(*, chunks, client=None):
            del client
            return DocumentIndexingResult(index_alias='legal-documents', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
        with patch('app.services.admin_document_replacement.extract_contacts_from_docx', return_value=parsed_contacts):
            import io
            return safe_upload_and_index_document(filename='GB.docx', file_stream=io.BytesIO(b'fake docx bytes'), source_directory=source_directory, processed_directory=source_directory / 'processed', maximum_bytes=10000000, country_confirmed=True, confirm_warnings=True, replace_existing=replace_existing, chunk_builder=chunk_builder, country_document_lookup=country_document_lookup, country_document_indexer=country_document_indexer)

    def test_fresh_upload_seeds_contact_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            response = self._run_upload(source_directory=source_directory, parsed_contacts=[ExtractedContact(member_firm='New Firm')], replace_existing=False)
            state = read_contact_state(source_directory, response.document_id)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].member_firm, 'New Firm')
            self.assertFalse(is_admin_modified_since_upload(source_directory, response.document_id))

class ReseedContactsFromCurrentDocxTests(unittest.TestCase):

    def test_confirmed_reseed_discards_admin_edits(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()
            with _patched_indexer(client):
                add_contact(document_id=DOCUMENT_ID, fields=_write_request(member_firm='Admin Edited Firm'), source_directory=source_directory, client=client)
            with patch('app.services.admin_contacts.extract_contacts_from_docx', return_value=[ExtractedContact(member_firm='Parsed DOCX Firm')]), _patched_indexer(client):
                response = reseed_contacts_from_current_docx(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(len(response.contacts), 1)
            self.assertEqual(response.contacts[0].member_firm, 'Parsed DOCX Firm')
            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(state.contacts[0].member_firm, 'Parsed DOCX Firm')
            self.assertFalse(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

class CurrentDocxPhotoReseedTests(unittest.TestCase):
    """
    Real _reseed_contacts_from_current_docx_locked (the private,
    lock-already-held entry point), against a REAL corpus document -
    proving the current-DOCX reseed path also carries real photo
    association through, complementing
    test_confirmed_reseed_discards_admin_edits above (which patches
    extract_contacts_from_docx directly and never touches a real
    photo)."""

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.source_directory = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_current_docx_reseed_belgium_seeds_two_photo_contacts(self) -> None:
        belgium_filename = 'Labour and Employment Law in Belgium 2026.docx'
        path = _test_admin_contacts__SOURCE_ROOT / belgium_filename
        if not path.exists():
            self.skipTest('Belgium corpus DOCX unavailable')
        skip_if_already_canonicalized(self, path)
        metadata = {'source_filename': belgium_filename, 'country': 'Belgium', 'country_code': 'BE'}
        with patch.object(admin_contacts, '_load_country_code_and_metadata', return_value=('BE', metadata)), patch.object(admin_contacts, 'resolve_document_source_path', return_value=SimpleNamespace(path=path)), patch.object(admin_contacts, '_document_metadata_for_chunks', return_value={}), patch.object(admin_contacts, 'build_contact_chunk_for_contacts', return_value=None), patch.object(admin_contacts, 'replace_document_contact_chunk'), patch.object(admin_contacts, 'is_admin_modified_since_upload', return_value=False), patch.object(admin_contacts, 'reset_admin_modified'):
            admin_contacts._reseed_contacts_from_current_docx_locked(validated_document_id='doc-current-belgium', source_directory=self.source_directory, opensearch_client=object())
        state = read_contact_state(self.source_directory, 'doc-current-belgium')
        self.assertIsNotNone(state)
        self.assertEqual(2, len(state.contacts))
        self.assertEqual(['Chris van Olmen', 'Nicolas Simon'], [contact.contact_person for contact in state.contacts])
        self.assertEqual(2, sum((contact.photo_filename is not None for contact in state.contacts)))



# ================================================================
# SOURCE: backend/tests/test_contact_documents.py
# ================================================================

import hashlib
import io
import re
import tempfile
import unittest
import xml.etree.ElementTree as ET
import zipfile
from pathlib import Path
from unittest.mock import patch
from docx import Document
from docx import Document as WordDocument
from lxml import etree
from app.services import contact_document_area as _contact_document_area
from app.services.contact_document_area import ContactAreaError, ContactPhotoPayload, rebuild_canonical_contact_table, resolve_untracked_contact_photo
from app.services.contact_document_photos import ContactDocumentPhotoError, add_contact_photo_to_document, add_new_contact_photo_to_document, remove_contact_photo_from_document, replace_contact_photo_in_document
from app.services.contact_people import associate_contact_photos
from app.services.contact_photos import ContactPhotoCandidate, extract_contact_photo_candidates
from app.services.docx_parser import CONTACT_TABLE_HIDDEN_MARKER, ExtractedContact, extract_contacts_from_docx, split_combined_legacy_contact
from app.services.document_contact_materializer import _find_all_contact_runs, _is_contact_related_block
from tests.support.documents import resolve_source_root
from tests.support.documents import make_png, require_corpus_copy, skip_if_already_canonicalized
_test_contact_documents__SOURCE_ROOT = resolve_source_root()

def _sha(data: bytes) -> str:
    return hashlib.sha256(data).hexdigest()
_test_contact_documents__VALID_PNG = make_png(183, 234, (200, 50, 50))

def _make_webp_vp8l(width: int, height: int) -> bytes:
    """A minimal, real, valid lossless (VP8L) WebP image - the format
    Pillow's default Image.save(format="WEBP") produces for a synthetic
    in-memory image, and the one this environment's own _WebpImageHeader
    parser (contact_document_area.py) must recognize for python-docx to
    embed it at all (python-docx 1.2.0 ships no WebP support whatsoever -
    confirmed directly)."""
    import struct
    packed = height - 1 << 14 | width - 1
    payload = bytes([47]) + packed.to_bytes(4, 'little') + bytes([136, 136, 8])
    chunk = b'VP8L' + struct.pack('<I', len(payload)) + payload
    if len(chunk) % 2:
        chunk += b'\x00'
    riff_payload = b'WEBP' + chunk
    return b'RIFF' + struct.pack('<I', len(riff_payload)) + riff_payload

def _make_webp_vp8(width: int, height: int) -> bytes:
    """The other real-world WebP sub-format (simple lossy VP8) - a
    different byte layout than VP8L, commonly produced by a browser or
    screenshot tool."""
    import struct
    frame_tag = b'\x00\x00\x00'
    start_code = b'\x9d\x01*'
    width_bytes = (width & 16383).to_bytes(2, 'little')
    height_bytes = (height & 16383).to_bytes(2, 'little')
    payload = frame_tag + start_code + width_bytes + height_bytes + b'\x00' * 4
    chunk = b'VP8 ' + struct.pack('<I', len(payload)) + payload
    if len(chunk) % 2:
        chunk += b'\x00'
    riff_payload = b'WEBP' + chunk
    return b'RIFF' + struct.pack('<I', len(riff_payload)) + riff_payload
_VALID_WEBP_LOSSLESS = _make_webp_vp8l(120, 160)
_VALID_WEBP_LOSSY = _make_webp_vp8(120, 160)
_WP_NS = '{http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing}'
_W_NS = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
_PIC_NS = '{http://schemas.openxmlformats.org/drawingml/2006/picture}'

class ContactDocumentAreaTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)

    def _require_copy(self, filename: str) -> Path:
        return require_corpus_copy(self, _test_contact_documents__SOURCE_ROOT, filename, Path(self.temp.name))

    def _structural_checks(self, path: Path) -> None:
        """Zip integrity, XML well-formedness, no dangling
        relationships, no [trash]/ parts, and no NEW floating contact
        shapes (wp:anchor) - the canonical table's own content is
        exclusively ordinary inline pictures/paragraphs."""
        with zipfile.ZipFile(path) as archive:
            self.assertIsNone(archive.testzip())
            names = set(archive.namelist())
            self.assertFalse([n for n in names if n.startswith('[trash]')], 'no [trash]/ parts may be left behind')
            document_xml = archive.read('word/document.xml').decode('utf-8')
            rels_xml = archive.read('word/_rels/document.xml.rels').decode('utf-8')
            content_types_xml = archive.read('[Content_Types].xml').decode('utf-8')
        for label, content in (('document.xml', document_xml), ('rels', rels_xml), ('content types', content_types_xml)):
            try:
                ET.fromstring(content)
            except ET.ParseError as error:
                self.fail(f'{label} is not well-formed XML: {error}')
        referenced_ids = set(re.findall('r:(?:id|embed)="([^"]+)"', document_xml))
        declared_ids = set(re.findall('Id="([^"]+)"', rels_xml))
        self.assertEqual(set(), referenced_ids - declared_ids, 'every r:id/r:embed referenced in document.xml must be declared in the relationships part')
        WordDocument(path)

    def _decorative_top_and_bottom_wrapped_shapes(self, document_xml: str) -> int:
        """Count floating shapes wrapped topAndBottom that carry
        neither real text nor a picture - contact_document_area.py's
        own removal target for the legacy reserved-space rectangle
        ("the tallest topAndBottom-wrapped shape with neither text nor
        a picture", per its module docstring). Checking this semantic
        invariant directly - rather than hardcoding an expected shape
        count - stays correct whether or not a given source document
        still happens to carry a decorative rectangle: real organic
        content drifts over time (a page's own title text box, for
        example, is also wrapped topAndBottom, has real text, and must
        never be mistaken for a removal candidate)."""
        root = ET.fromstring(document_xml)
        decorative = 0
        for anchor in root.iter(f'{_WP_NS}anchor'):
            if anchor.find(f'{_WP_NS}wrapTopAndBottom') is None:
                continue
            has_text = any(((node.text or '').strip() for node in anchor.iter(f'{_W_NS}t')))
            has_picture = anchor.find(f'.//{_PIC_NS}pic') is not None
            if not has_text and (not has_picture):
                decorative += 1
        return decorative

    def _table_xml(self, document_xml: str) -> str:
        marker_index = document_xml.find(CONTACT_TABLE_HIDDEN_MARKER)
        self.assertNotEqual(-1, marker_index, "the canonical table's own hidden marker must be present")
        table_start = document_xml.rfind('<w:tbl>', 0, marker_index)
        table_end = document_xml.find('</w:tbl>', marker_index)
        self.assertNotEqual(-1, table_start)
        self.assertNotEqual(-1, table_end)
        return document_xml[table_start:table_end + len('</w:tbl>')]

    def test_add_first_contact_with_photo_produces_coherent_block(self) -> None:
        """Canonicalizing AU's own organic contact area leaves no large
        empty band (Introduction follows the new table directly) and
        the sole contact keeps its own photo.

        The "no empty band" check asserts the semantic contract
        directly (no decorative, text-less/picture-less
        topAndBottom-wrapped shape survives, and no NEW one is
        introduced) rather than a hardcoded shape count: AU's own real
        content has drifted since this test was written and its
        current baseline no longer carries a separate removable
        reserved-space rectangle at all (verified directly against the
        real corpus - AU's only topAndBottom-wrapped shape today is
        its own title text box, which legitimately survives
        unchanged)."""
        path = self._require_copy('AU.docx')
        with zipfile.ZipFile(path) as archive:
            original_document_xml = archive.read('word/document.xml').decode('utf-8')
        original_wrap_count = original_document_xml.count('wrapTopAndBottom')
        original_contacts = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(1, len(original_contacts))
        michael = original_contacts[0]
        michael_photo = resolve_untracked_contact_photo(path, contact_person=michael.contact_person, country='Australia')
        self.assertIsNotNone(michael_photo)
        new_bytes = rebuild_canonical_contact_table(path, contacts=(michael,), photos=(michael_photo,), country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8')
        table_xml = self._table_xml(document_xml)
        self.assertEqual(0, table_xml.count('<wp:anchor'), 'the canonical table must contain no floating shapes at all - only ordinary inline content')
        self.assertEqual(0, self._decorative_top_and_bottom_wrapped_shapes(document_xml), 'no decorative, text-less, picture-less topAndBottom rectangle may survive canonicalization')
        self.assertLessEqual(document_xml.count('wrapTopAndBottom'), original_wrap_count, 'canonicalization must never introduce a new topAndBottom-wrapped shape')
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(1, len(reparsed))
        self.assertEqual('Michael Harmer', reparsed[0].contact_person)
        photos = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(photos))
        self.assertEqual(michael_photo.data, photos[0].data)

    def test_two_contacts_two_distinct_photos(self) -> None:
        """Two contacts, each with its own distinct photo, round-trip
        as two separate contacts with two separate photos - never
        merged or cross-associated."""
        path = self._require_copy('AU.docx')
        michael = extract_contacts_from_docx(path, country='Australia')[0]
        michael_photo = resolve_untracked_contact_photo(path, contact_person=michael.contact_person, country='Australia')
        jane = ExtractedContact(member_firm='Second Firm Pty Ltd', contact_person='Jane Secondary', email='jane.secondary@secondfirm.com.au', phone='+61 2 9000 0000', address='100 Test Street, Level 5', website='www.secondfirm.com.au')
        jane_photo = ContactPhotoPayload(data=_test_contact_documents__VALID_PNG, content_type='image/png')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(michael, jane), photos=(michael_photo, jane_photo), country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(2, len(reparsed))
        self.assertEqual('Michael Harmer', reparsed[0].contact_person)
        self.assertEqual('Jane Secondary', reparsed[1].contact_person)
        photos = extract_contact_photo_candidates(path)
        shas = {p.sha256 for p in photos}
        self.assertEqual(2, len(shas))
        self.assertIn(michael_photo and _sha(michael_photo.data), shas)
        self.assertIn(_sha(jane_photo.data), shas)

    def test_three_contacts_all_preserved_in_order(self) -> None:
        """Three contacts round-trip in the same order, with the same
        field values."""
        path = self._require_copy('AU.docx')
        michael = extract_contacts_from_docx(path, country='Australia')[0]
        jane = ExtractedContact(member_firm='Second Firm Pty Ltd', contact_person='Jane Secondary', email='jane.secondary@secondfirm.com.au')
        priya = ExtractedContact(member_firm='Third Firm Pty Ltd', contact_person='Priya Third', email='priya.third@thirdfirm.com.au')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(michael, jane, priya), photos=(None, None, None), country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(3, len(reparsed))
        self.assertEqual(['Michael Harmer', 'Jane Secondary', 'Priya Third'], [c.contact_person for c in reparsed])

    def test_contact_without_photo_produces_coherent_block(self) -> None:
        """A contact with no photo at all still produces a coherent
        block - an empty right cell, never a missing row or misaligned
        text."""
        path = self._require_copy('AU.docx')
        michael = extract_contacts_from_docx(path, country='Australia')[0]
        jane = ExtractedContact(member_firm='Second Firm Pty Ltd', contact_person='Jane Secondary', email='jane.secondary@secondfirm.com.au')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(michael, jane), photos=(None, None), country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(2, len(reparsed))
        photos = extract_contact_photo_candidates(path)
        self.assertEqual(0, len(photos))

    def test_delete_contact_b_leaves_a_intact(self) -> None:
        """Rebuilding from a contact list with B removed leaves B's
        text/photo gone and A's text/photo untouched."""
        path = self._require_copy('AU.docx')
        michael = extract_contacts_from_docx(path, country='Australia')[0]
        michael_photo = resolve_untracked_contact_photo(path, contact_person=michael.contact_person, country='Australia')
        jane = ExtractedContact(member_firm='Second Firm Pty Ltd', contact_person='Jane Secondary', email='jane.secondary@secondfirm.com.au')
        jane_photo = ContactPhotoPayload(data=_test_contact_documents__VALID_PNG, content_type='image/png')
        with_both = rebuild_canonical_contact_table(path, contacts=(michael, jane), photos=(michael_photo, jane_photo), country='Australia')
        path.write_bytes(with_both)
        after_delete = rebuild_canonical_contact_table(path, contacts=(michael,), photos=(michael_photo,), country='Australia')
        path.write_bytes(after_delete)
        self._structural_checks(path)
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8')
        self.assertNotIn('Jane Secondary', document_xml)
        self.assertNotIn('jane.secondary@secondfirm.com.au', document_xml)
        self.assertIn('Michael Harmer', document_xml)
        remaining_photos = {p.sha256 for p in extract_contact_photo_candidates(path)}
        self.assertNotIn(_sha(jane_photo.data), remaining_photos)
        self.assertIn(_sha(michael_photo.data), remaining_photos)

    def test_deleting_every_contact_leaves_no_table_and_no_band(self) -> None:
        """Rebuilding with zero contacts removes the canonical area
        entirely - no marker-only leftover table, no empty band."""
        path = self._require_copy('AU.docx')
        michael = extract_contacts_from_docx(path, country='Australia')[0]
        with_contact = rebuild_canonical_contact_table(path, contacts=(michael,), photos=(None,), country='Australia')
        path.write_bytes(with_contact)
        emptied = rebuild_canonical_contact_table(path, contacts=(), photos=(), country='Australia')
        path.write_bytes(emptied)
        self._structural_checks(path)
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8')
        self.assertNotIn(CONTACT_TABLE_HIDDEN_MARKER, document_xml)
        self.assertNotIn('<w:tbl>', document_xml)
        self.assertEqual(0, len(extract_contacts_from_docx(path)))

    def test_legal_content_unchanged(self) -> None:
        """Real legal section text (well past the contact area) is
        byte-identical after the rebuild."""
        path = self._require_copy('AU.docx')
        with zipfile.ZipFile(path) as archive:
            before = archive.read('word/document.xml').decode('utf-8')
        introduction_index = before.find('Introduction')
        self.assertNotEqual(-1, introduction_index)
        legal_tail_before = before[introduction_index:]
        michael = extract_contacts_from_docx(path, country='Australia')[0]
        new_bytes = rebuild_canonical_contact_table(path, contacts=(michael,), photos=(None,), country='Australia')
        path.write_bytes(new_bytes)
        with zipfile.ZipFile(path) as archive:
            after = archive.read('word/document.xml').decode('utf-8')
        introduction_index_after = after.find('Introduction')
        self.assertNotEqual(-1, introduction_index_after)
        legal_tail_after = after[introduction_index_after:]
        self.assertEqual(legal_tail_before, legal_tail_after, 'legal content from Introduction onward must be byte-for-byte unchanged')

    def test_no_matching_contact_area_still_produces_canonical_table(self) -> None:
        """A document with no legacy floating contact area at all
        (Portugal) still gets a real canonical table - never left
        ContactState-only."""
        path = self._require_copy('PT.docx')
        new_contact = ExtractedContact(member_firm='Someone New Lda', contact_person='Someone New', email='new@example.test', phone='+351 21 000 0000', address='Rua Nova 1', website='www.newfirm.test')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(new_contact,), photos=(None,), country='Portugal')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Portugal')
        self.assertEqual(1, len(reparsed))
        self.assertEqual('Someone New', reparsed[0].contact_person)

    def test_second_rebuild_replaces_rather_than_duplicates_table(self) -> None:
        """Rebuilding twice in a row replaces the one canonical table
        rather than stacking a second one alongside it."""
        path = self._require_copy('PT.docx')
        first = ExtractedContact(member_firm='First Firm Lda', contact_person='First Person', email='first@example.test')
        second_contact = ExtractedContact(member_firm='Second Firm Lda', contact_person='Second Person', email='second@example.test')
        path.write_bytes(rebuild_canonical_contact_table(path, contacts=(first,), photos=(None,), country='Portugal'))
        path.write_bytes(rebuild_canonical_contact_table(path, contacts=(first, second_contact), photos=(None, None), country='Portugal'))
        self._structural_checks(path)
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8')
        self.assertEqual(1, document_xml.count('<w:tbl>'))
        self.assertEqual(1, document_xml.count(CONTACT_TABLE_HIDDEN_MARKER))
        reparsed = extract_contacts_from_docx(path, country='Portugal')
        self.assertEqual(2, len(reparsed))

    def test_canada_legacy_floating_carrier_add_produces_single_zone(self) -> None:
        """Real Canada: Robert Bayne's own original contact area is a
        legacy floating-shape carrier, still un-canonicalized. Adding a
        second contact must leave EXACTLY one canonical table, with the
        legacy carrier entirely gone - never both side by side."""
        path = self._require_copy('CA.docx')
        baseline = extract_contacts_from_docx(path, country='Canada')
        self.assertEqual(1, len(baseline))
        self.assertEqual('Robert Bayne', baseline[0].contact_person)
        new_contact = ExtractedContact(member_firm='CRUD Regression LLP', contact_person='Regression Contact', email='contact@crud-regression.example', phone='+33 1 23 45 67 89', address='12 Test Street, Floor 3, 75008 Paris', website='www.crud-regression.example')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(*baseline, new_contact), photos=(None, None), country='Canada')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8')
        self.assertEqual(1, document_xml.count(CONTACT_TABLE_HIDDEN_MARKER))
        self.assertEqual(1, document_xml.count('<w:tbl>'))
        reparsed = extract_contacts_from_docx(path, country='Canada')
        self.assertEqual(2, len(reparsed))
        self.assertEqual('Robert Bayne', reparsed[0].contact_person)
        self.assertEqual('Regression Contact', reparsed[1].contact_person)

    def test_single_zone_invariant_rejects_undetected_legacy_carrier(self) -> None:
        """Adversarial: simulate a legacy floating-shape carrier that
        detection failed to remove (a canonical table inserted
        alongside it, exactly as if _remove_legacy_carrier_and_get_
        anchor had missed it) - the invariant must refuse to persist
        this, rather than silently producing the historical Argentina/
        Canada two-block defect for whatever country's shape geometry
        someday defeats detection."""
        path = self._require_copy('CA.docx')
        source_bytes = path.read_bytes()
        legacy_photo_relationship_ids = {candidate.relationship_id for candidate in extract_contact_photo_candidates(path) if candidate.relationship_id}
        document = WordDocument(path)
        new_contact = ExtractedContact(member_firm='Undetected Carrier Regression LLP', contact_person='Undetected Carrier Contact', email='undetected@example.test')
        table = _contact_document_area._build_canonical_table(document, contacts=(new_contact,), photos=(None,))
        document.element.body.insert(0, table._tbl)
        output = io.BytesIO()
        document.save(output)
        with self.assertRaises(ContactAreaError) as raised:
            _contact_document_area._assert_single_contact_zone(output.getvalue(), expected_table_count=1, legacy_photo_relationship_ids=legacy_photo_relationship_ids)
        self.assertIn('floating-shape', str(raised.exception))
        self.assertEqual(source_bytes, path.read_bytes(), 'the real corpus file must never be touched by this test')

    def test_single_zone_invariant_rejects_duplicated_canonical_table(self) -> None:
        """Adversarial: two canonical tables in the same document (a
        table-count detection miss) must also be refused, not just a
        leftover legacy shape."""
        path = self._require_copy('AU.docx')
        document = WordDocument(path)
        extra_contact = ExtractedContact(member_firm='Duplicate Table Regression LLP', contact_person='Duplicate Table Contact', email='duplicate@example.test')
        extra_table = _contact_document_area._build_canonical_table(document, contacts=(extra_contact,), photos=(None,))
        document.element.body.insert(0, extra_table._tbl)
        output = io.BytesIO()
        document.save(output)
        with self.assertRaises(ContactAreaError) as raised:
            _contact_document_area._assert_single_contact_zone(output.getvalue(), expected_table_count=1, legacy_photo_relationship_ids=set())
        self.assertIn('Expected exactly 1', str(raised.exception))

    def test_photo_with_crop_rectangle_is_refused_not_guessed(self) -> None:
        """resolve_untracked_contact_photo refuses (rather than
        blindly embeds) a legacy photo whose own source shape carries
        an OOXML crop rectangle this environment cannot reproduce."""
        path = self._require_copy('AU.docx')
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8')
        patched_xml = document_xml.replace('<a:blip r:embed="rId10"/>', '<a:blip r:embed="rId10"/><a:srcRect l="1000" t="1000" r="1000" b="1000"/>', 1)
        self.assertNotEqual(document_xml, patched_xml, 'sanity: the replacement must actually match something')
        from app.services.contact_document_photos import _rewrite_zip
        patched_bytes = _rewrite_zip(path.read_bytes(), replacements={'word/document.xml': patched_xml})
        path.write_bytes(patched_bytes)
        with self.assertRaises(ContactAreaError):
            resolve_untracked_contact_photo(path, contact_person='Michael Harmer', country='Australia')

    def test_bare_srcrect_with_no_offsets_is_not_treated_as_a_crop(self) -> None:
        """A bare <a:srcRect/> (no l/t/r/b attributes at all - OOXML
        defaults each to 0) declares no visible crop whatsoever and
        must NOT be refused - only a genuinely non-zero offset should
        be. Real Indonesia corpus content carries exactly this shape on
        its own contact portrait (see
        test_indonesia_untracked_photo_with_noop_crop_resolves for the
        real-corpus regression); this test isolates the same defect
        with a synthetic AU fixture so it does not depend on Indonesia's
        own content remaining unchanged."""
        path = self._require_copy('AU.docx')
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8')
        patched_xml = document_xml.replace('<a:blip r:embed="rId10"/>', '<a:blip r:embed="rId10"/><a:srcRect/>', 1)
        self.assertNotEqual(document_xml, patched_xml, 'sanity: the replacement must actually match something')
        from app.services.contact_document_photos import _rewrite_zip
        patched_bytes = _rewrite_zip(path.read_bytes(), replacements={'word/document.xml': patched_xml})
        path.write_bytes(patched_bytes)
        photo = resolve_untracked_contact_photo(path, contact_person='Michael Harmer', country='Australia')
        self.assertIsNotNone(photo)

    def test_indonesia_untracked_photo_with_noop_crop_resolves(self) -> None:
        """Real-corpus regression: Indonesia's own contact portrait
        (Marshall Situmorang) carries a bare, all-zero <a:srcRect/> -
        a Word-emitted no-op, not an actual crop. Before the offset-
        aware fix, resolve_untracked_contact_photo refused it outright,
        which would block ANY first Contact CRUD mutation on this
        document (Add/Edit/Delete all need to preserve an untracked
        contact's existing photo via this same path)."""
        path = self._require_copy('ID.docx')
        photo = resolve_untracked_contact_photo(path, contact_person='Marshall Situmorang', country='Indonesia')
        self.assertIsNotNone(photo)

    def test_each_of_the_six_fields_individually_empty_round_trips(self) -> None:
        """All six fields are individually optional. address/website-
        both-empty (France's own real shape) and member_firm-empty
        (IE/IN's own real shape) are already covered elsewhere; this
        closes the remaining combinations - contact_person and email
        individually empty - at the actual DOCX round-trip layer, not
        just the AdminContactWriteRequest validation layer (see
        AdminContactWriteRequestOptionalFieldTests in
        test_admin_contacts.py for that one)."""
        path = self._require_copy('AU.docx')
        full_fields = dict(member_firm='Complete Fields LLP', address='1 Complete Street', phone='+1 555 000 1111', website='www.complete-fields.example', contact_person='Complete Person', email='complete@example.test')
        for empty_field in full_fields:
            with self.subTest(empty_field=empty_field):
                fields = dict(full_fields)
                fields[empty_field] = None
                contact = ExtractedContact(**fields)
                new_bytes = rebuild_canonical_contact_table(path, contacts=(contact,), photos=(None,), country='Australia')
                path.write_bytes(new_bytes)
                reparsed = extract_contacts_from_docx(path, country='Australia')
                self.assertEqual(1, len(reparsed))
                for field_name, expected_value in fields.items():
                    self.assertEqual(expected_value, getattr(reparsed[0], field_name), f'field {field_name!r} did not round-trip when {empty_field!r} was empty')

    def test_multiple_contacts_with_different_missing_fields_and_a_duplicate(self) -> None:
        """Several contacts with DIFFERENT missing-field shapes,
        content deliberately chosen to fool a content-based classifier
        (a postal code, a phone-like substring embedded in an address,
        an apartment/unit number, a scheme-less URL, accented/unicode
        names and an apostrophe), and an exact duplicate pair - all in
        the SAME canonical table, at once. No field may shift into
        another contact's slot, and duplicates must survive unmerged."""
        path = self._require_copy('AU.docx')
        firm_only = ExtractedContact(member_firm='Firm Only LLP')
        person_email_only = ExtractedContact(contact_person='Person Only', email='person-only@example.test')
        fooling_shapes = ExtractedContact(member_firm="Müller & O'Connell Associés", address="Apt #4B, 12 Rue de l'Église, 75008 Paris, +33 1 23 45 67 89", phone='75008', website='fooling-shapes.example', contact_person="François O'Connell-Müller", email='françois@fooling-shapes.example')
        duplicate_a = ExtractedContact(member_firm='Duplicate LLP', contact_person='Duplicate Person', email='duplicate@example.test')
        duplicate_b = ExtractedContact(member_firm='Duplicate LLP', contact_person='Duplicate Person', email='duplicate@example.test')
        contacts = (firm_only, person_email_only, fooling_shapes, duplicate_a, duplicate_b)
        photos = (None,) * len(contacts)
        new_bytes = rebuild_canonical_contact_table(path, contacts=contacts, photos=photos, country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(list(contacts), reparsed)

    def test_accented_email_local_part_never_corrupted(self) -> None:
        """Real bug found while building the item C/D/E combinatorial
        test above: _EMAIL_PATTERN's local-part character class is
        ASCII-only, so an email with an accented local part (a
        genuinely legitimate address, e.g. a French admin's own name)
        was silently truncated on read-back, and the truncated
        fragment then failed to even look like an email, spilling into
        contact_person instead. contact_person/email are now tagged
        the same deterministic way as the left cell's four fields, so
        this no longer depends on _EMAIL_PATTERN at all for a
        canonical table this code itself wrote."""
        path = self._require_copy('AU.docx')
        contact = ExtractedContact(contact_person='François Müller', email='françois.müller@example.test')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(contact,), photos=(None,), country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(1, len(reparsed))
        self.assertEqual('François Müller', reparsed[0].contact_person)
        self.assertEqual('françois.müller@example.test', reparsed[0].email)

    def test_webp_contact_photo_can_be_embedded(self) -> None:
        """Real bug, reproduced directly against pre-fix code: python-
        docx 1.2.0 has no WebP header parser at all, so ANY Add+Photo/
        Edit+Photo/Replace-Photo mutation with a WebP file (a format
        this codebase already accepts and advertises - admin_contact_
        photos.py's own "Only JPEG, PNG and WebP images are accepted")
        failed with a raw UnrecognizedImageError wrapped as
        ContactAreaError. contact_document_area.py registers a small,
        dependency-free WebP header parser into python-docx's own
        image factory; this proves both real-world WebP sub-formats -
        lossless VP8L (what Pillow produces by default) and lossy
        simple VP8 (a different byte layout, common from a browser or
        screenshot tool) - now embed and round-trip correctly, with no
        image library added."""
        path = self._require_copy('AU.docx')
        baseline = extract_contacts_from_docx(path, country='Australia')
        webp_variants = {'lossless (VP8L)': _VALID_WEBP_LOSSLESS, 'lossy (VP8)': _VALID_WEBP_LOSSY}
        for variant_name, webp_bytes in webp_variants.items():
            with self.subTest(variant=variant_name):
                new_contact = ExtractedContact(member_firm='WebP Test Firm', contact_person=f'WebP Test Person ({variant_name})', email='webp@example.test')
                photo = ContactPhotoPayload(data=webp_bytes, content_type='image/webp')
                new_bytes = rebuild_canonical_contact_table(path, contacts=(*baseline, new_contact), photos=(*[None] * len(baseline), photo), country='Australia')
                path.write_bytes(new_bytes)
                self._structural_checks(path)
                candidates = extract_contact_photo_candidates(path)
                actual_shas = {c.sha256 for c in candidates}
                self.assertIn(_sha(webp_bytes), actual_shas)
                reparsed = extract_contacts_from_docx(path, country='Australia')
                self.assertEqual(len(baseline) + 1, len(reparsed))
                self.assertEqual(f'WebP Test Person ({variant_name})', reparsed[-1].contact_person)

    def test_multiple_emails_on_one_line_round_trip_preserved(self) -> None:
        """Regression for the real France defect found by the
        full-corpus mutation test: FR's actual contact carries two
        email addresses on a single comma-joined line
        ("scherrmann@flichy.com, bacquet@flichy.com"), because
        ExtractedContact/ContactState model email as one string field
        and the writer renders it verbatim as one line. The canonical
        table reader used to recover only the first address via a
        single regex .search(); the round-trip validator correctly
        caught the loss and refused the rebuild with a
        ContactAreaError, but the fix (using .findall() to recover
        every address on the line) must let the exact real France
        contact round-trip cleanly with both emails intact, in
        order."""
        path = self._require_copy('FR.docx')
        flichy = ExtractedContact(member_firm='Flichy Grangé Avocats', contact_person='Caroline Scherrmann and Florence Bacquet', email='scherrmann@flichy.com, bacquet@flichy.com', phone='+33 1 56 62 30 00')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(flichy,), photos=(None,), country='France')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='France')
        self.assertEqual(1, len(reparsed))
        self.assertEqual('scherrmann@flichy.com, bacquet@flichy.com', reparsed[0].email)

    def test_empty_firm_cell_phone_addition_round_trips(self) -> None:
        """Real cases: IE's real contact (Aoife Bradley) and IN's real
        contact (Avik Biswas) both have member_firm/address/phone/
        website all empty. Giving each a phone value must land it as
        phone - never misread as member_firm (which the writer would
        otherwise never have written in the first place, since it was
        empty to begin with). Confirmed against both real documents,
        since the fix is architectural (never specific to either
        country)."""
        cases = {'Ireland': ('IE.docx', 'Ireland', '+353 1 234 5678'), 'India': ('IN.docx', 'India', '+91 11 4567 8900')}
        for label, (filename, country, new_phone) in cases.items():
            with self.subTest(country=label):
                path = self._require_copy(filename)
                baseline = extract_contacts_from_docx(path, country=country)
                self.assertEqual(1, len(baseline))
                self.assertIsNone(baseline[0].member_firm)
                self.assertIsNone(baseline[0].phone)
                with_phone = ExtractedContact(contact_person=baseline[0].contact_person, email=baseline[0].email, phone=new_phone)
                new_bytes = rebuild_canonical_contact_table(path, contacts=(with_phone,), photos=(None,), country=country)
                path.write_bytes(new_bytes)
                self._structural_checks(path)
                reparsed = extract_contacts_from_docx(path, country=country)
                self.assertEqual(1, len(reparsed))
                self.assertIsNone(reparsed[0].member_firm)
                self.assertEqual(new_phone, reparsed[0].phone)

    def test_us_noop_rebuild_preserves_embedded_url_and_website(self) -> None:
        """Quirk C, real case: US's real address is a prose sentence
        that mentions "www.jacksonlewis.com." (with its own
        sentence-ending period), and website is the separate, clean
        "www.jacksonlewis.com". A no-op rebuild (identical fields, no
        edit at all) must reproduce BOTH exactly - the embedded
        mention must never be mistaken for the dedicated website
        line, and the dedicated website line must never pick up the
        sentence's trailing punctuation."""
        path = self._require_copy('US.docx')
        baseline = extract_contacts_from_docx(path, country='United States')
        if len(baseline) != 1 or not baseline[0].address or 'please see www.jacksonlewis.com.' not in baseline[0].address:
            self.skipTest("US.docx's real content no longer has the exact embedded-URL contact this test assumes (real corpus content has drifted since this test was written) - see test_docx_parser.ClassifyCanonicalFirmLinesTests.test_website_mentioned_inside_address_prose_is_not_extracted for permanent, corpus-independent coverage of this exact fix")
        new_bytes = rebuild_canonical_contact_table(path, contacts=tuple(baseline), photos=(None,), country='United States')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='United States')
        self.assertEqual(1, len(reparsed))
        self.assertEqual(baseline[0].address, reparsed[0].address)
        self.assertEqual('www.jacksonlewis.com', reparsed[0].website)

    def test_phone_with_trailing_annotation_round_trips(self) -> None:
        """Quirk A, real case: AU's real phone value, given a
        trailing annotation an Admin might legitimately type (e.g.
        "(mobile)"), must round-trip as ONE whole phone value -
        never split, with the annotation leaking into address."""
        path = self._require_copy('AU.docx')
        baseline = extract_contacts_from_docx(path, country='Australia')[0]
        annotated = ExtractedContact(member_firm=baseline.member_firm, contact_person=baseline.contact_person, email=baseline.email, phone=f'{baseline.phone} (mobile)', address=baseline.address, website=baseline.website)
        new_bytes = rebuild_canonical_contact_table(path, contacts=(annotated,), photos=(None,), country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(1, len(reparsed))
        self.assertEqual(f'{baseline.phone} (mobile)', reparsed[0].phone)
        self.assertEqual(baseline.address, reparsed[0].address)

    def test_full_field_validator_rejects_a_field_shift(self) -> None:
        """The strengthened _validate_canonical_table() must reject a
        rebuild whose re-parsed contact shows a field-shift (a phone
        value landing in member_firm) - not just a contact_person/
        email mismatch. Simulated via a patched extract_contacts_
        from_docx, since the real reader no longer produces this
        shift on its own (that's the fix)."""
        path = self._require_copy('AU.docx')
        baseline = extract_contacts_from_docx(path, country='Australia')[0]
        shifted = ExtractedContact(member_firm=baseline.phone, contact_person=baseline.contact_person, email=baseline.email, phone=None, address=baseline.address, website=baseline.website)
        with patch('app.services.contact_document_area.extract_contacts_from_docx', return_value=[shifted]):
            with self.assertRaises(ContactAreaError) as raised:
                rebuild_canonical_contact_table(path, contacts=(baseline,), photos=(None,), country='Australia')
        message = str(raised.exception)
        self.assertIn('member_firm', message)
        self.assertIn('phone', message)

    def _france_split_contacts(self) -> tuple[ExtractedContact, ExtractedContact]:
        path = self._require_copy('FR.docx')
        document = WordDocument(path)
        if any((table.rows and CONTACT_TABLE_HIDDEN_MARKER in table.rows[0].cells[0].text for table in document.tables)):
            self.skipTest('FR.docx has since been canonicalized by real Admin usage (real corpus content has drifted since this test was written) - see FranceSyntheticContactCrudTests for permanent, corpus-independent coverage of the same scenarios')
        combined = extract_contacts_from_docx(path, country='France')[0]
        split = split_combined_legacy_contact(combined)
        self.assertIsNotNone(split)
        self.assertEqual(2, len(split))
        return (split[0], split[1])

    def test_france_canonical_table_replaces_legacy_block_before_poc(self) -> None:
        """The canonical table must land at the EXACT location the
        legacy plain-paragraph block occupied - before "YOUR L&E
        GLOBAL POC" - never the document start, with the old legacy
        text fully removed and the POC/Jessica Stout block untouched."""
        path = self._require_copy('FR.docx')
        caroline, florence = self._france_split_contacts()
        new_bytes = rebuild_canonical_contact_table(path, contacts=(caroline, florence), photos=(None, None), country='France')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        document = WordDocument(path)
        full_text = '\n'.join((p.text for p in document.paragraphs))
        self.assertEqual(1, len(document.tables))
        self.assertEqual(0, full_text.count('Caroline Scherrmann and Florence Bacquet'), 'the old legacy combined-name paragraph must be gone')
        self.assertEqual(1, full_text.count('YOUR L&E GLOBAL POC'), 'the POC heading must survive exactly once')
        self.assertEqual(1, full_text.count('Jessica Stout'), 'Jessica Stout must survive - she is not a member-firm contact and must never be removed')
        body_children = list(document.element.body)
        table_position = next((i for i, child in enumerate(body_children) if child.tag.endswith('}tbl')))
        poc_position = next((i for i, child in enumerate(body_children) if child.tag.endswith('}p') and 'YOUR L&E GLOBAL POC' in ''.join((node.text or '' for node in child.iter() if node.tag.endswith('}t')))))
        self.assertLess(table_position, poc_position, "the canonical table must be inserted before the POC block, not at the document's start")
        self.assertGreater(table_position, 10, 'the table must not have landed at the very start of the document (the _default_insertion_anchor fallback must never fire when a real legacy contact area exists)')
        reparsed = extract_contacts_from_docx(path, country='France')
        self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet'], [c.contact_person for c in reparsed])

    def test_france_add_temporary_contact_then_delete_round_trips(self) -> None:
        """Caroline/Florence baseline -> add a temporary third contact
        -> full-field round-trip valid, in order -> delete the
        temporary contact -> back to exactly Caroline/Florence, with no
        ContactAreaError at any stage (the internal round-trip
        validator is never weakened or bypassed - the data is made to
        satisfy it)."""
        path = self._require_copy('FR.docx')
        caroline, florence = self._france_split_contacts()
        temporary = ExtractedContact(member_firm='Temp Firm', contact_person='Temporary Person', email='temp@example.com')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(caroline, florence), photos=(None, None), country='France')
        path.write_bytes(new_bytes)
        new_bytes = rebuild_canonical_contact_table(path, contacts=(caroline, florence, temporary), photos=(None, None, None), country='France')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='France')
        self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet', 'Temporary Person'], [c.contact_person for c in reparsed])
        new_bytes = rebuild_canonical_contact_table(path, contacts=(caroline, florence), photos=(None, None), country='France')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='France')
        self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet'], [c.contact_person for c in reparsed])

    def test_france_photo_stays_with_the_correct_split_contact(self) -> None:
        """A photo attached to only one of the two split France
        contacts must never migrate to the other."""
        path = self._require_copy('FR.docx')
        caroline, florence = self._france_split_contacts()
        caroline_photo = ContactPhotoPayload(data=_test_contact_documents__VALID_PNG, content_type='image/png')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(caroline, florence), photos=(caroline_photo, None), country='France')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        photos = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(photos))
        self.assertEqual(_sha(caroline_photo.data), photos[0].sha256)
        reparsed = extract_contacts_from_docx(path, country='France')
        self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet'], [c.contact_person for c in reparsed])

    def test_france_clearing_website_persists_as_cleared(self) -> None:
        """Clearing a previously-set field must round-trip as actually
        cleared, never silently retaining the stale value."""
        path = self._require_copy('FR.docx')
        caroline, florence = self._france_split_contacts()
        with_website = ExtractedContact(member_firm=caroline.member_firm, contact_person=caroline.contact_person, email=caroline.email, phone=caroline.phone, website='www.example.com')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(with_website, florence), photos=(None, None), country='France')
        path.write_bytes(new_bytes)
        reparsed = extract_contacts_from_docx(path, country='France')
        self.assertEqual('www.example.com', reparsed[0].website)
        cleared = ExtractedContact(member_firm=caroline.member_firm, contact_person=caroline.contact_person, email=caroline.email, phone=caroline.phone, website=None)
        new_bytes = rebuild_canonical_contact_table(path, contacts=(cleared, florence), photos=(None, None), country='France')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='France')
        self.assertIsNone(reparsed[0].website, 'a cleared field must round-trip as cleared, not stale')

    def test_belgium_combined_contact_splits_consistently_like_france(self) -> None:
        """Verified directly against the real corpus rather than
        assumed: Belgium's own real legacy contact ALSO names two
        people sharing one comma-joined email string ("Chris van Olmen
        and Nicolas Simon"), the identical shape to France's. The
        multi-person split is a general, country-agnostic rule (never a
        France-specific exception), so it correctly and safely applies
        here too - this is the consistency the split logic must have,
        not a special case to avoid. What must genuinely stay
        unaffected is the ALREADY-validated canonical table mechanism's
        own handling of an explicitly-given, already-distinct
        multi-contact list: split_combined_legacy_contact is only ever
        invoked once, at legacy bootstrap time - never during an
        ordinary rebuild/update round-trip - so two contacts already
        split (by bootstrap or given directly, as here) keep
        round-tripping as two distinct, unmerged contacts."""
        path = self._require_copy('Labour and Employment Law in Belgium 2026.docx')
        document = WordDocument(path)
        if any((table.rows and CONTACT_TABLE_HIDDEN_MARKER in table.rows[0].cells[0].text for table in document.tables)):
            self.skipTest("Belgium's document has since been canonicalized by real Admin usage (real corpus content has drifted since this test was written) - its combined legacy contact no longer exists in raw form to split")
        baseline = extract_contacts_from_docx(path, country='Belgium')
        self.assertEqual(1, len(baseline), "Belgium's real legacy contact is currently one combined record, the same shape as France's, prior to any split")
        split = split_combined_legacy_contact(baseline[0])
        self.assertIsNotNone(split)
        self.assertEqual(['Chris van Olmen', 'Nicolas Simon'], [c.contact_person for c in split])
        new_bytes = rebuild_canonical_contact_table(path, contacts=tuple(split), photos=tuple((None for _ in split)), country='Belgium')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Belgium')
        self.assertEqual(['Chris van Olmen', 'Nicolas Simon'], [c.contact_person for c in reparsed])
        for contact in reparsed:
            self.assertIsNone(split_combined_legacy_contact(contact), f'{contact.contact_person!r} is already a single person and must never be split further')

    def test_adversarial_address_phone_never_swaps(self) -> None:
        """The identical field-tag mechanism, confirmed independently
        against three real, differently-shaped adversarial addresses -
        never a country-specific patch."""
        cases = {'Brazil': ('Labour and Employment Law in Brazil 2026.docx', ExtractedContact(member_firm='Tozzini Freire', contact_person='Gabriela Lima', email='glima@tozzinifreire.com.br', address='Rua Borges Lagoa, 1328, 04038-904 São Paulo, +55 115 086 5000', phone='04038-904')), 'Romania': ('Labour and Employment Law in Romania 2026.docx', ExtractedContact(member_firm='Volonciu & Associates', contact_person='Magda Volonciu', email='magdavolonciu@volonciu.ro', address='No. 35 Alexandru Constantinescu Street, 2nd Floor, 011471 1st District Bucharest, +40 372 755 699', phone='011471 1')), 'Singapore': ('Employment Law Overview Singapore 2026.docx', ExtractedContact(member_firm='Clyde & Co Clasis', contact_person='Thomas Choo', email='thomas.choo@clydeco.com', address='12 Marina Boulevard | , Marina Bay Financial Centre Tower 3 | #30 - 03, 018982 Singapore, +65 654 465 00', phone='30 - 03'))}
        for country, (filename, adversarial) in cases.items():
            with self.subTest(country=country):
                path = self._require_copy(filename)
                new_bytes = rebuild_canonical_contact_table(path, contacts=(adversarial,), photos=(None,), country=country)
                path.write_bytes(new_bytes)
                self._structural_checks(path)
                reparsed = extract_contacts_from_docx(path, country=country)
                self.assertEqual(1, len(reparsed))
                self.assertEqual(adversarial.address, reparsed[0].address)
                self.assertEqual(adversarial.phone, reparsed[0].phone)

    def test_address_never_mistaken_for_website(self) -> None:
        """An address value that happens to contain a URL-shaped
        fragment must never be reclassified as the website field, and
        vice-versa - the tag alone decides, never content shape."""
        path = self._require_copy('AU.docx')
        adversarial = ExtractedContact(member_firm='Test Firm', contact_person='Test Person', email='test@example.com', address='123 www.looks-like-a-site.com Street', website='not-a-url-shaped-value')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(adversarial,), photos=(None,), country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(1, len(reparsed))
        self.assertEqual(adversarial.address, reparsed[0].address)
        self.assertEqual(adversarial.website, reparsed[0].website)

    def test_full_six_field_canonical_round_trip_is_deterministic(self) -> None:
        """Every one of the six fields, including deliberately
        adversarial/ambiguous shapes, must round-trip byte-for-byte
        through a canonical rebuild - the general contract this whole
        mechanism exists to guarantee, not just the reported
        countries."""
        path = self._require_copy('AU.docx')
        adversarial = ExtractedContact(member_firm='123-456-7890', contact_person='Test Person', email='test@example.com', phone='www.not-really-a-website.com', address='+1 555 000 0000 is not the phone field here', website='04038-904')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(adversarial,), photos=(None,), country='Australia')
        path.write_bytes(new_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(1, len(reparsed))
        self.assertEqual(adversarial, reparsed[0])

    def test_website_without_www_prefix_never_leaks_into_address(self) -> None:
        """Explains a real, previously-observed divergence: a global
        dry-run using a "clean" (already www.-prefixed) website value
        passed, while a real Admin typing a website WITHOUT that
        prefix in the live browser (e.g. "zhonglun.com" instead of
        "www.zhonglun.com" - a completely natural human edit) hit
        "field(s) changed: address, website", because the old website
        detector required a www./http:// prefix to recognize a line as
        the dedicated website field at all; without it, the value fell
        through into address instead, and website came back None.
        Confirmed by reproducing this exact error message against the
        real China and Spain contacts on the pre-fix code - the
        identical architectural fix, not a country-specific patch."""
        cases = {'China': ('CN.docx', 'China'), 'Spain': ('Labour and Employment Law in Spain 2026.docx', 'Spain')}
        for label, (filename, country) in cases.items():
            with self.subTest(country=label):
                path = self._require_copy(filename)
                baseline = extract_contacts_from_docx(path, country=country)[0]
                edited = ExtractedContact(member_firm=baseline.member_firm, contact_person=baseline.contact_person, email=baseline.email, phone=baseline.phone, address=baseline.address, website=baseline.website.replace('www.', ''))
                new_bytes = rebuild_canonical_contact_table(path, contacts=(edited,), photos=(None,), country=country)
                path.write_bytes(new_bytes)
                self._structural_checks(path)
                reparsed = extract_contacts_from_docx(path, country=country)
                self.assertEqual(1, len(reparsed))
                self.assertEqual(edited.address, reparsed[0].address)
                self.assertEqual(edited.website, reparsed[0].website)

    def test_old_untagged_canonical_table_still_reads_correctly(self) -> None:
        """An already-canonicalized real document, written by a
        version of this code before the field tags existed (verified
        directly: AR.docx's real production canonical table has none
        of them), must keep reading correctly via the same content-
        based fallback classification this reader has always used -
        this mechanism never breaks an already-deployed canonical
        table, it only makes NEW ones fully deterministic."""
        path = self._require_copy('AU.docx')
        contact = ExtractedContact(member_firm='Untagged Firm', contact_person='Untagged Person', email='untagged@example.com', phone='+1 555 010 0000', address='1 Untagged Street', website='www.untagged.example.com')
        new_bytes = rebuild_canonical_contact_table(path, contacts=(contact,), photos=(None,), country='Australia')
        path.write_bytes(new_bytes)
        from app.services.docx_parser import _FIELD_TAG_ADDRESS, _FIELD_TAG_CONTACT_PERSON, _FIELD_TAG_EMAIL, _FIELD_TAG_MEMBER_FIRM, _FIELD_TAG_PHONE, _FIELD_TAG_WEBSITE
        with zipfile.ZipFile(path) as archive:
            document_xml = archive.read('word/document.xml').decode('utf-8')
        for tag in (_FIELD_TAG_MEMBER_FIRM, _FIELD_TAG_ADDRESS, _FIELD_TAG_PHONE, _FIELD_TAG_WEBSITE, _FIELD_TAG_CONTACT_PERSON, _FIELD_TAG_EMAIL):
            self.assertIn(tag, document_xml, 'sanity: the tag must actually be present to strip')
            document_xml = document_xml.replace(tag, '')
        from app.services.contact_document_photos import _rewrite_zip
        untagged_bytes = _rewrite_zip(path.read_bytes(), replacements={'word/document.xml': document_xml})
        path.write_bytes(untagged_bytes)
        self._structural_checks(path)
        reparsed = extract_contacts_from_docx(path, country='Australia')
        self.assertEqual(1, len(reparsed))
        self.assertEqual(contact, reparsed[0])

def _build_synthetic_france_shaped_docx(path: Path) -> None:
    """
    A minimal, from-scratch DOCX matching France's own real structure
    exactly enough to exercise find_plain_paragraph_contact_block_
    bounds / split_combined_legacy_contact / rebuild_canonical_
    contact_table the same way real FR.docx used to, before real Admin
    usage canonicalized it in production (real corpus content drift -
    see ContactDocumentAreaTests._france_split_contacts).

    Shape: title -> padding body paragraphs (legal-looking filler, so
    the eventual canonical table lands well past the document start,
    matching the real "table must not land at the very start" check) ->
    the legacy plain-paragraph contact block (person / role+firm /
    comma-joined emails / phone) -> "YOUR L&E GLOBAL POC" -> Jessica
    Stout (the project-level POC, never a member-firm contact) ->
    disclaimer.
    """
    document = WordDocument()
    document.add_paragraph('Employment Law Overview - France')
    for index in range(12):
        document.add_paragraph(f'Filler legal paragraph number {index} with enough text to look like real body content.')
    document.add_paragraph('Caroline Scherrmann and Florence Bacquet')
    document.add_paragraph('Partners, Flichy Grangé Avocats')
    document.add_paragraph('scherrmann@flichy.com, bacquet@flichy.com')
    document.add_paragraph('+33 1 56 62 30 00')
    document.add_paragraph('YOUR L&E GLOBAL POC')
    document.add_paragraph('Jessica Stout')
    document.add_paragraph('jstout@leglobal.law')
    document.add_paragraph('Disclaimer: this publication is for informational purposes only and does not constitute legal advice.')
    document.save(path)

class FranceSyntheticContactCrudTests(unittest.TestCase):
    """
    Permanent, corpus-independent coverage of France's own special
    contact shape - a synthetic fixture built fresh in setUp, immune to
    real corpus content drift (unlike ContactDocumentAreaTests._france_
    split_contacts, which now skips because real FR.docx has since been
    canonicalized by genuine Admin usage in production).
    """

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.path = Path(self.temp.name) / 'france-synthetic.docx'
        _build_synthetic_france_shaped_docx(self.path)

    def _split_contacts(self) -> tuple[ExtractedContact, ExtractedContact]:
        combined = extract_contacts_from_docx(self.path, country='France')[0]
        split = split_combined_legacy_contact(combined)
        self.assertIsNotNone(split)
        self.assertEqual(2, len(split))
        return (split[0], split[1])

    def test_canonical_table_replaces_legacy_block_before_poc(self) -> None:
        caroline, florence = self._split_contacts()
        new_bytes = rebuild_canonical_contact_table(self.path, contacts=(caroline, florence), photos=(None, None), country='France')
        self.path.write_bytes(new_bytes)
        document = WordDocument(self.path)
        full_text = '\n'.join((p.text for p in document.paragraphs))
        self.assertEqual(1, len(document.tables))
        self.assertEqual(0, full_text.count('Caroline Scherrmann and Florence Bacquet'), 'the old legacy combined-name paragraph must be gone')
        self.assertEqual(1, full_text.count('YOUR L&E GLOBAL POC'))
        self.assertEqual(1, full_text.count('Jessica Stout'))
        body_children = list(document.element.body)
        table_position = next((i for i, child in enumerate(body_children) if child.tag.endswith('}tbl')))
        poc_position = next((i for i, child in enumerate(body_children) if child.tag.endswith('}p') and 'YOUR L&E GLOBAL POC' in ''.join((node.text or '' for node in child.iter() if node.tag.endswith('}t')))))
        self.assertLess(table_position, poc_position)
        self.assertGreater(table_position, 10)
        reparsed = extract_contacts_from_docx(self.path, country='France')
        self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet'], [c.contact_person for c in reparsed])

    def test_add_temporary_contact_then_delete_round_trips(self) -> None:
        caroline, florence = self._split_contacts()
        temporary = ExtractedContact(member_firm='Temp Firm', contact_person='Temporary Person', email='temp@example.com')
        self.path.write_bytes(rebuild_canonical_contact_table(self.path, contacts=(caroline, florence), photos=(None, None), country='France'))
        self.path.write_bytes(rebuild_canonical_contact_table(self.path, contacts=(caroline, florence, temporary), photos=(None, None, None), country='France'))
        reparsed = extract_contacts_from_docx(self.path, country='France')
        self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet', 'Temporary Person'], [c.contact_person for c in reparsed])
        self.path.write_bytes(rebuild_canonical_contact_table(self.path, contacts=(caroline, florence), photos=(None, None), country='France'))
        reparsed = extract_contacts_from_docx(self.path, country='France')
        self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet'], [c.contact_person for c in reparsed])

    def test_photo_stays_with_the_correct_split_contact(self) -> None:
        caroline, florence = self._split_contacts()
        caroline_photo = ContactPhotoPayload(data=_test_contact_documents__VALID_PNG, content_type='image/png')
        self.path.write_bytes(rebuild_canonical_contact_table(self.path, contacts=(caroline, florence), photos=(caroline_photo, None), country='France'))
        photos = extract_contact_photo_candidates(self.path)
        self.assertEqual(1, len(photos))
        self.assertEqual(_sha(caroline_photo.data), photos[0].sha256)
        reparsed = extract_contacts_from_docx(self.path, country='France')
        self.assertEqual(['Caroline Scherrmann', 'Florence Bacquet'], [c.contact_person for c in reparsed])

    def test_clearing_website_persists_as_cleared(self) -> None:
        caroline, florence = self._split_contacts()
        with_website = ExtractedContact(member_firm=caroline.member_firm, contact_person=caroline.contact_person, email=caroline.email, phone=caroline.phone, website='www.example.com')
        self.path.write_bytes(rebuild_canonical_contact_table(self.path, contacts=(with_website, florence), photos=(None, None), country='France'))
        reparsed = extract_contacts_from_docx(self.path, country='France')
        self.assertEqual('www.example.com', reparsed[0].website)
        cleared = ExtractedContact(member_firm=caroline.member_firm, contact_person=caroline.contact_person, email=caroline.email, phone=caroline.phone, website=None)
        self.path.write_bytes(rebuild_canonical_contact_table(self.path, contacts=(cleared, florence), photos=(None, None), country='France'))
        reparsed = extract_contacts_from_docx(self.path, country='France')
        self.assertIsNone(reparsed[0].website)
_LEGACY_VALID_JPEG = bytes.fromhex('ffd8ffe000104a46494600010100000100010000ffd9')
_LEGACY_VALID_PNG = bytes.fromhex('89504e470d0a1a0a0000000d49484452000000010000000108060000001f15c4890000000a49444154789c6360000002000100ffff03000006000557bfabd40000000049454e44ae426082')

class RealCorpusContactDocumentPhotoTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)

    def _require_copy(self, filename: str) -> Path:
        return require_corpus_copy(self, _test_contact_documents__SOURCE_ROOT, filename, Path(self.temp.name))

    def _unrelated_parts(self, path: Path) -> dict[str, bytes]:
        """Every zip part except document.xml/rels/content-types -
        the parts a photo mutation must never touch."""
        skip = {'word/document.xml', 'word/_rels/document.xml.rels', '[Content_Types].xml'}
        with zipfile.ZipFile(path) as archive:
            return {name: archive.read(name) for name in archive.namelist() if name not in skip}

    def test_belgium_replace_isolates_the_other_contacts_photo(self) -> None:
        path = self._require_copy('Labour and Employment Law in Belgium 2026.docx')
        original_candidates = extract_contact_photo_candidates(path)
        by_name = {c.source_filename: c for c in original_candidates}
        chris = by_name['image2.jpg']
        chris_sha = chris.sha256
        nicolas_sha = by_name['image1.png'].sha256
        original_paragraphs = len(Document(str(path)).paragraphs)
        unrelated_before = self._unrelated_parts(path)
        del unrelated_before[chris.media_path]
        new_bytes = replace_contact_photo_in_document(path, target_sha256=chris_sha, new_data=_LEGACY_VALID_JPEG, new_content_type='image/jpeg')
        out = Path(self.temp.name) / 'belgium_replaced.docx'
        out.write_bytes(new_bytes)
        new_shas = {c.sha256 for c in extract_contact_photo_candidates(out)}
        self.assertIn(hashlib.sha256(_LEGACY_VALID_JPEG).hexdigest(), new_shas)
        self.assertNotIn(chris_sha, new_shas)
        self.assertIn(nicolas_sha, new_shas)
        self.assertEqual(original_paragraphs, len(Document(str(out)).paragraphs))
        unrelated_after = self._unrelated_parts(out)
        del unrelated_after[chris.media_path]
        self.assertEqual(unrelated_before, unrelated_after)

    def test_belgium_remove_isolates_the_other_contacts_photo(self) -> None:
        path = self._require_copy('Labour and Employment Law in Belgium 2026.docx')
        skip_if_already_canonicalized(self, path)
        by_name = {c.source_filename: c.sha256 for c in extract_contact_photo_candidates(path)}
        chris_sha = by_name['image2.jpg']
        nicolas_sha = by_name['image1.png']
        original_paragraphs = len(Document(str(path)).paragraphs)
        new_bytes = remove_contact_photo_from_document(path, target_sha256=nicolas_sha)
        out = Path(self.temp.name) / 'belgium_removed.docx'
        out.write_bytes(new_bytes)
        remaining = extract_contact_photo_candidates(out)
        self.assertEqual(1, len(remaining))
        self.assertEqual(chris_sha, remaining[0].sha256)
        self.assertEqual(original_paragraphs, len(Document(str(out)).paragraphs))

    def test_belgium_add_into_a_shared_zone_fails_closed(self) -> None:
        """
        Chris and Nicolas share ONE combined "CONTACT PERSON" textbox
        zone, disambiguated only by their two photos' own separate
        geometry - never by two separate zones. Adding a photo for
        Nicolas after his own is removed has no safe, deterministic
        place to insert it (the zone still names Chris too), so this
        must fail closed rather than guess.
        """
        path = self._require_copy('Labour and Employment Law in Belgium 2026.docx')
        by_name = {c.source_filename: c.sha256 for c in extract_contact_photo_candidates(path)}
        nicolas_sha = by_name['image1.png']
        removed_bytes = remove_contact_photo_from_document(path, target_sha256=nicolas_sha)
        removed_path = Path(self.temp.name) / 'belgium_removed.docx'
        removed_path.write_bytes(removed_bytes)
        with self.assertRaises(ContactDocumentPhotoError):
            add_contact_photo_to_document(removed_path, contact_person='Nicolas Simon', new_data=_LEGACY_VALID_JPEG, new_content_type='image/jpeg', other_contact_persons=['Chris van Olmen'])

    def test_germany_add_to_a_genuinely_photo_less_contact(self) -> None:
        path = self._require_copy('DE.docx')
        original_candidates = extract_contact_photo_candidates(path)
        self.assertEqual(0, len(original_candidates), "Tobias Pusch is expected to have no existing photo in the real corpus - if this fails, the corpus changed and this test's premise needs revisiting.")
        original_paragraphs = len(Document(str(path)).paragraphs)
        new_bytes = add_contact_photo_to_document(path, contact_person='Tobias Pusch', new_data=_LEGACY_VALID_PNG, new_content_type='image/png', other_contact_persons=[])
        out = Path(self.temp.name) / 'de_added.docx'
        out.write_bytes(new_bytes)
        candidates = extract_contact_photo_candidates(out)
        self.assertEqual(1, len(candidates))
        self.assertEqual(hashlib.sha256(_LEGACY_VALID_PNG).hexdigest(), candidates[0].sha256)
        self.assertEqual(original_paragraphs, len(Document(str(out)).paragraphs))

    def test_argentina_replace_then_remove_single_contact(self) -> None:
        path = self._require_copy('AR.docx')
        original_candidates = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(original_candidates))
        target_sha = original_candidates[0].sha256
        original_paragraphs = len(Document(str(path)).paragraphs)
        replaced_bytes = replace_contact_photo_in_document(path, target_sha256=target_sha, new_data=_LEGACY_VALID_JPEG, new_content_type='image/jpeg')
        replaced_path = Path(self.temp.name) / 'ar_replaced.docx'
        replaced_path.write_bytes(replaced_bytes)
        replaced_candidates = extract_contact_photo_candidates(replaced_path)
        self.assertEqual(1, len(replaced_candidates))
        self.assertEqual(hashlib.sha256(_LEGACY_VALID_JPEG).hexdigest(), replaced_candidates[0].sha256)
        self.assertEqual(original_paragraphs, len(Document(str(replaced_path)).paragraphs))
        removed_bytes = remove_contact_photo_from_document(replaced_path, target_sha256=hashlib.sha256(_LEGACY_VALID_JPEG).hexdigest())
        removed_path = Path(self.temp.name) / 'ar_removed.docx'
        removed_path.write_bytes(removed_bytes)
        self.assertEqual(0, len(extract_contact_photo_candidates(removed_path)))
        self.assertEqual(original_paragraphs, len(Document(str(removed_path)).paragraphs))

    def test_replace_of_an_unlocatable_sha_fails_closed(self) -> None:
        path = self._require_copy('AR.docx')
        with self.assertRaises(ContactDocumentPhotoError):
            replace_contact_photo_in_document(path, target_sha256='0' * 64, new_data=_LEGACY_VALID_JPEG, new_content_type='image/jpeg')

    def test_remove_of_an_unlocatable_sha_fails_closed(self) -> None:
        path = self._require_copy('AR.docx')
        with self.assertRaises(ContactDocumentPhotoError):
            remove_contact_photo_from_document(path, target_sha256='0' * 64)

    def test_add_for_an_unknown_contact_person_fails_closed(self) -> None:
        path = self._require_copy('DE.docx')
        with self.assertRaises(ContactDocumentPhotoError):
            add_contact_photo_to_document(path, contact_person='Nobody Real', new_data=_LEGACY_VALID_PNG, new_content_type='image/png', other_contact_persons=[])

    def test_replace_of_the_existing_zone_now_uses_true_geometry(self) -> None:
        """
        Regression guard for a fixed bug: adding a photo for an
        EXISTING named zone must genuinely geometrically overlap that
        zone (reason "GEOMETRY"), never merely happen to be the sole
        remaining portrait in the whole document (reason
        "UNIQUE_PORTRAIT") - the fallback only ever works by
        coincidence and breaks the moment the document has more than
        one photo total, exactly the scenario a brand-new second
        contact's photo introduces.
        """
        path = self._require_copy('DE.docx')
        new_bytes = add_contact_photo_to_document(path, contact_person='Tobias Pusch', new_data=_LEGACY_VALID_PNG, new_content_type='image/png', other_contact_persons=[])
        out = Path(self.temp.name) / 'de_geometry.docx'
        out.write_bytes(new_bytes)
        candidates = extract_contact_photo_candidates(out)
        self.assertEqual(1, len(candidates))
        self.assertEqual('GEOMETRY', candidates[0].reason)

    def test_add_new_contact_photo_anchors_to_the_largest_zone_alongside_an_existing_photo(self) -> None:
        """
        The core blocker scenario this primitive exists for: a
        document that ALREADY has one contact's photo, and a genuinely
        brand-new second contact (whose name cannot possibly appear
        anywhere yet) also gets a photo. Both must round-trip as
        independently valid, GEOMETRY-reasoned candidates - never
        relying on the "exactly one remaining portrait" fallback,
        which cannot disambiguate once there is more than one
        unassociated photo.
        """
        path = self._require_copy('AR.docx')
        skip_if_already_canonicalized(self, path)
        original = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(original))
        original_sha = original[0].sha256
        original_paragraphs = len(Document(str(path)).paragraphs)
        new_bytes = add_new_contact_photo_to_document(path, new_data=_LEGACY_VALID_JPEG, new_content_type='image/jpeg')
        out = Path(self.temp.name) / 'ar_new_contact.docx'
        out.write_bytes(new_bytes)
        candidates = extract_contact_photo_candidates(out)
        self.assertEqual(2, len(candidates))
        for candidate in candidates:
            self.assertEqual('GEOMETRY', candidate.reason, 'must be a real geometric match, never a fallback that only works by coincidence')
        new_shas = {c.sha256 for c in candidates}
        self.assertIn(original_sha, new_shas)
        self.assertIn(hashlib.sha256(_LEGACY_VALID_JPEG).hexdigest(), new_shas)
        self.assertEqual(original_paragraphs, len(Document(str(out)).paragraphs))

    def test_add_new_contact_photo_into_belgiums_shared_zone_isolates_existing_photos(self) -> None:
        """
        Belgium's shared "CONTACT PERSON" zone (Chris + Nicolas) is
        the document's only (and therefore largest) CONTACT PERSON
        zone - a brand-new third contact's photo anchors there too,
        and must never disturb either of the two existing contacts'
        own photos.
        """
        path = self._require_copy('Labour and Employment Law in Belgium 2026.docx')
        skip_if_already_canonicalized(self, path)
        original_shas = {c.sha256 for c in extract_contact_photo_candidates(path)}
        self.assertEqual(2, len(original_shas))
        new_bytes = add_new_contact_photo_to_document(path, new_data=_LEGACY_VALID_PNG, new_content_type='image/png')
        out = Path(self.temp.name) / 'belgium_new_contact.docx'
        out.write_bytes(new_bytes)
        candidates = extract_contact_photo_candidates(out)
        new_shas = {c.sha256 for c in candidates}
        self.assertEqual(3, len(new_shas))
        self.assertTrue(original_shas.issubset(new_shas))
        self.assertIn(hashlib.sha256(_LEGACY_VALID_PNG).hexdigest(), new_shas)

    def test_add_new_contact_photo_fails_closed_with_no_contact_zone_at_all(self) -> None:
        """
        A document with genuinely zero "CONTACT PERSON" zones (no
        structural home at all) must fail closed rather than insert
        the photo at an arbitrary position such as the document's end.
        """
        path = self._require_copy('PT.docx')
        self.assertEqual(0, len(extract_contact_photo_candidates(path)))
        with self.assertRaises(ContactDocumentPhotoError):
            add_new_contact_photo_to_document(path, new_data=_LEGACY_VALID_JPEG, new_content_type='image/jpeg')

class RealCorpusContactPhotoTests(unittest.TestCase):

    def _require(self, filename: str) -> Path:
        path = _test_contact_documents__SOURCE_ROOT / filename
        if not path.exists():
            self.skipTest(f'Real corpus source unavailable: {path}')
        return path

    def test_belgium_has_exactly_two_contact_photos(self) -> None:
        path = self._require('Labour and Employment Law in Belgium 2026.docx')
        skip_if_already_canonicalized(self, path)
        photos = extract_contact_photo_candidates(path)
        self.assertEqual(2, len(photos))
        self.assertEqual(['image2.jpg', 'image1.png'], [photo.source_filename for photo in photos])
        self.assertEqual(['GEOMETRY', 'GEOMETRY'], [photo.reason for photo in photos])
        for photo in photos:
            self.assertTrue(photo.data)
            self.assertTrue(photo.sha256)
            self.assertTrue(photo.content_type.startswith('image/'))

    def test_ireland_rejects_the_wide_false_positive(self) -> None:
        path = self._require('IE.docx')
        photos = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(photos))
        self.assertEqual('image2.jpg', photos[0].source_filename)

    def test_indonesia_rejects_pagoda_and_logo(self) -> None:
        path = self._require('ID.docx')
        photos = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(photos))
        self.assertEqual('image3.jpeg', photos[0].source_filename)
        self.assertEqual('UNIQUE_PORTRAIT', photos[0].reason)

    def test_chile_has_no_contact_photo(self) -> None:
        path = self._require('CL.docx')
        self.assertEqual([], extract_contact_photo_candidates(path))

    def test_germany_has_no_contact_photo(self) -> None:
        path = self._require('DE.docx')
        self.assertEqual([], extract_contact_photo_candidates(path))

    def test_india_has_no_contact_photo(self) -> None:
        path = self._require('IN.docx')
        self.assertEqual([], extract_contact_photo_candidates(path))

    def test_france_has_no_contact_photo(self) -> None:
        path = self._require('FR.docx')
        skip_if_already_canonicalized(self, path)
        self.assertEqual([], extract_contact_photo_candidates(path))

    def test_result_is_deterministic(self) -> None:
        path = self._require('Labour and Employment Law in Belgium 2026.docx')
        first = extract_contact_photo_candidates(path)
        second = extract_contact_photo_candidates(path)
        self.assertEqual([(item.source_filename, item.sha256, item.reason) for item in first], [(item.source_filename, item.sha256, item.reason) for item in second])

    def test_candidate_is_an_immutable_value_object(self) -> None:
        candidate = ContactPhotoCandidate(source_filename='portrait.jpg', content_type='image/jpeg', data=b'photo', sha256='abc', reason='GEOMETRY')
        with self.assertRaises((AttributeError, TypeError)):
            candidate.source_filename = 'changed.jpg'

def _photo(name: str) -> ContactPhotoCandidate:
    return ContactPhotoCandidate(source_filename=name, content_type='image/jpeg', data=name.encode(), sha256='sha-' + name, reason='GEOMETRY')

class ContactPhotoAssociationUnitTests(unittest.TestCase):

    def test_single_contact_single_photo_is_associated(self) -> None:
        contact = ExtractedContact(member_firm='Firm', contact_person='Jane Doe', email='jane@example.com', phone='+1 111', address='Address', website='example.com')
        result = associate_contact_photos([contact], [_photo('jane.jpg')])
        self.assertEqual(1, len(result))
        self.assertEqual('Jane Doe', result[0].contact_person)
        self.assertEqual('jane@example.com', result[0].email)
        self.assertEqual('jane.jpg', result[0].photo.source_filename)

    def test_single_contact_without_photo_is_preserved(self) -> None:
        contact = ExtractedContact(member_firm='Firm', contact_person='Jane Doe', email='jane@example.com', phone='+1 111', address='Address', website='example.com')
        result = associate_contact_photos([contact], [])
        self.assertEqual(1, len(result))
        self.assertEqual('Jane Doe', result[0].contact_person)
        self.assertIsNone(result[0].photo)

    def test_two_existing_contacts_two_photos_are_mapped_in_order(self) -> None:
        contacts = [ExtractedContact(member_firm='Firm', contact_person='Person One', email='one@example.com'), ExtractedContact(member_firm='Firm', contact_person='Person Two', email='two@example.com')]
        result = associate_contact_photos(contacts, [_photo('one.jpg'), _photo('two.jpg')])
        self.assertEqual(2, len(result))
        self.assertEqual(('Person One', 'one@example.com', 'one.jpg'), (result[0].contact_person, result[0].email, result[0].photo.source_filename))
        self.assertEqual(('Person Two', 'two@example.com', 'two.jpg'), (result[1].contact_person, result[1].email, result[1].photo.source_filename))

    def test_combined_contact_is_split_only_when_photo_count_matches(self) -> None:
        contact = ExtractedContact(member_firm='Shared Firm', contact_person='Person One and Person Two', email='one@example.com, two@example.com', phone='+32 123', address='Shared address', website='firm.example')
        result = associate_contact_photos([contact], [_photo('one.jpg'), _photo('two.jpg')])
        self.assertEqual(2, len(result))
        self.assertEqual('Person One', result[0].contact_person)
        self.assertEqual('one@example.com', result[0].email)
        self.assertEqual('one.jpg', result[0].photo.source_filename)
        self.assertEqual('Person Two', result[1].contact_person)
        self.assertEqual('two@example.com', result[1].email)
        self.assertEqual('two.jpg', result[1].photo.source_filename)
        for item in result:
            self.assertEqual('Shared Firm', item.member_firm)
            self.assertEqual('+32 123', item.phone)
            self.assertEqual('Shared address', item.address)
            self.assertEqual('firm.example', item.website)

    def test_combined_contact_without_photos_is_not_split(self) -> None:
        contact = ExtractedContact(member_firm='Firm', contact_person='Person One and Person Two', email='one@example.com, two@example.com')
        result = associate_contact_photos([contact], [])
        self.assertEqual(1, len(result))
        self.assertEqual('Person One and Person Two', result[0].contact_person)
        self.assertEqual('one@example.com, two@example.com', result[0].email)
        self.assertIsNone(result[0].photo)

    def test_mismatched_person_email_photo_counts_never_guess(self) -> None:
        contact = ExtractedContact(member_firm='Firm', contact_person='Person One and Person Two', email='one@example.com')
        result = associate_contact_photos([contact], [_photo('one.jpg'), _photo('two.jpg')])
        self.assertEqual(1, len(result))
        self.assertEqual('Person One and Person Two', result[0].contact_person)
        self.assertEqual('one@example.com', result[0].email)
        self.assertIsNone(result[0].photo)

    def test_extra_photos_never_get_arbitrarily_assigned(self) -> None:
        contact = ExtractedContact(member_firm='Firm', contact_person='Jane Doe', email='jane@example.com')
        result = associate_contact_photos([contact], [_photo('a.jpg'), _photo('b.jpg')])
        self.assertEqual(1, len(result))
        self.assertIsNone(result[0].photo)

class RealCorpusContactPersonPhotoTests(unittest.TestCase):

    def _require(self, filename: str) -> Path:
        path = _test_contact_documents__SOURCE_ROOT / filename
        if not path.exists():
            self.skipTest(f'Real corpus source unavailable: {path}')
        return path

    def test_belgium_becomes_two_individual_contacts(self) -> None:
        path = self._require('Labour and Employment Law in Belgium 2026.docx')
        skip_if_already_canonicalized(self, path)
        parsed = extract_contacts_from_docx(path)
        photos = extract_contact_photo_candidates(path)
        self.assertEqual(1, len(parsed))
        self.assertEqual(2, len(photos))
        result = associate_contact_photos(parsed, photos)
        self.assertEqual(2, len(result))
        self.assertEqual('Chris van Olmen', result[0].contact_person)
        self.assertEqual('chris.van.olmen@vow.be', result[0].email)
        self.assertEqual('image2.jpg', result[0].photo.source_filename)
        self.assertEqual('Nicolas Simon', result[1].contact_person)
        self.assertEqual('nicolas.simon@vow.be', result[1].email)
        self.assertEqual('image1.png', result[1].photo.source_filename)
        for contact in result:
            self.assertEqual('Van Olmen & Wynant', contact.member_firm)
            self.assertEqual('+32 264 405 11', contact.phone)
            self.assertEqual('www.vow.be', contact.website)

    def test_france_remains_backward_compatible_without_photos(self) -> None:
        path = self._require('FR.docx')
        skip_if_already_canonicalized(self, path)
        parsed = extract_contacts_from_docx(path)
        photos = extract_contact_photo_candidates(path)
        self.assertEqual([], photos)
        result = associate_contact_photos(parsed, photos)
        self.assertEqual(1, len(result))
        self.assertEqual('Caroline Scherrmann and Florence Bacquet', result[0].contact_person)
        self.assertEqual('scherrmann@flichy.com, bacquet@flichy.com', result[0].email)
        self.assertIsNone(result[0].photo)

def _textbox_run_xml(lines: list[str], width_emu: int=1000000, height_emu: int=500000) -> str:
    inner_paragraphs = ''.join((f'<w:p><w:r><w:t>{line}</w:t></w:r></w:p>' for line in lines))
    return f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"><w:rPr><w:noProof/></w:rPr><w:drawing><wp:anchor><wp:extent cx="{width_emu}" cy="{height_emu}"/><wp:docPr id="1" name="Box"/><a:graphic><a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape"><wps:wsp><wps:txbx><w:txbxContent>{inner_paragraphs}</w:txbxContent></wps:txbx></wps:wsp></a:graphicData></a:graphic></wp:anchor></w:drawing></w:r>'

def _build_docx_with_contact_textbox_at(paragraphs: list[str], textbox_paragraph_index: int, lines: list[str] | None=None, width_emu: int=1000000, height_emu: int=500000) -> bytes:
    """
    A real, openable DOCX with a minimal but structurally faithful
    floating Contact text box (CONTACT PERSON marker + name + email,
    matching the real corpus's own <w:txbxContent> shape, wrapped in a
    DrawingML <w:drawing><wp:anchor> run so _extract_extent and the
    <w:r>/<w:txbxContent> nesting _find_run_span must see through both
    apply) injected into one specific paragraph.
    """
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    target_paragraph = document.paragraphs[textbox_paragraph_index]
    run_element = etree.fromstring(_textbox_run_xml(lines or ['CONTACT PERSON', 'Test Person', 'test.person@example.com'], width_emu=width_emu, height_emu=height_emu).encode('utf-8'))
    target_paragraph._p.append(run_element)
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()

def _document_xml_of(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
        return archive.read('word/document.xml').decode('utf-8')

class IsContactRelatedBlockTests(unittest.TestCase):

    def test_contact_person_block_is_classified_as_contact_related(self) -> None:
        self.assertTrue(_is_contact_related_block(['CONTACT PERSON', 'Jane Doe']))

    def test_firm_block_with_email_is_classified_as_contact_related(self) -> None:
        self.assertTrue(_is_contact_related_block(['SOME FIRM', 'info@example.com']))

    def test_firm_block_with_phone_is_classified_as_contact_related(self) -> None:
        self.assertTrue(_is_contact_related_block(['SOME FIRM', '+1 555 000 1234']))

    def test_plain_branding_block_is_not_contact_related(self) -> None:
        self.assertFalse(_is_contact_related_block(['www.leglobal.law']))

    def test_document_title_block_is_not_contact_related(self) -> None:
        self.assertFalse(_is_contact_related_block(['Employment Law Overview - Testland']))

class FindAllContactRunsTests(unittest.TestCase):
    """
    Regression coverage for two bugs found and fixed together: (1)
    _find_run_span originally took the first "</w:r>" after a
    txbxContent's start, which is actually an INNER run nested inside
    the box's own paragraphs, not the enclosing floating-shape run: it
    must instead balance nested <w:r>/</w:r> tags. (2) a floating
    shape's Choice and Fallback branches each produce their own
    <w:txbxContent> regex match for the SAME enclosing run; the second
    occurrence must be recognized as already-covered by the first
    match's span, not re-resolved (which - before the fix - picked one
    of the first branch's own inner runs as a bogus second "run").
    """

    def test_finds_single_contact_run_with_its_extent(self) -> None:
        docx_bytes = _build_docx_with_contact_textbox_at(['Title', 'Body'], textbox_paragraph_index=0, width_emu=1234000, height_emu=567000)
        runs = _find_all_contact_runs(_document_xml_of(docx_bytes))
        self.assertEqual(len(runs), 1)
        self.assertEqual(runs[0].width_emu, 1234000)
        self.assertEqual(runs[0].height_emu, 567000)

    def test_ignores_non_contact_textboxes(self) -> None:
        document = Document()
        document.add_paragraph('Title')
        document.add_paragraph('Body')
        branding_run = etree.fromstring(_textbox_run_xml(['www.leglobal.law']).encode('utf-8'))
        document.paragraphs[0]._p.append(branding_run)
        buffer = io.BytesIO()
        document.save(buffer)
        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))
        self.assertEqual(runs, [])

    def test_two_separate_contact_runs_both_found_and_largest_is_selectable(self) -> None:
        document = Document()
        document.add_paragraph('Title')
        small_run = etree.fromstring(_textbox_run_xml(['CONTACT PERSON', 'Jane Doe', 'jane@example.com'], width_emu=500000, height_emu=300000).encode('utf-8'))
        large_run = etree.fromstring(_textbox_run_xml(['SOME FIRM', 'firm@example.com'], width_emu=2000000, height_emu=1000000).encode('utf-8'))
        document.paragraphs[0]._p.append(small_run)
        document.paragraphs[0]._p.append(large_run)
        buffer = io.BytesIO()
        document.save(buffer)
        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))
        self.assertEqual(len(runs), 2)
        primary = max(runs, key=lambda run: run.width_emu * run.height_emu)
        self.assertEqual(primary.width_emu, 2000000)

    def test_choice_and_fallback_style_duplicate_txbxcontent_is_one_run(self) -> None:
        document = Document()
        document.add_paragraph('Title')
        inner = '<w:p><w:r><w:t>CONTACT PERSON</w:t></w:r></w:p><w:p><w:r><w:t>Jane Doe</w:t></w:r></w:p><w:p><w:r><w:t>jane@example.com</w:t></w:r></w:p>'
        duplicated_run_xml = f'<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"><mc:AlternateContent xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006"><mc:Choice Requires="wps"><w:txbxContent>{inner}</w:txbxContent></mc:Choice><mc:Fallback><w:txbxContent>{inner}</w:txbxContent></mc:Fallback></mc:AlternateContent></w:r>'
        run_element = etree.fromstring(duplicated_run_xml.encode('utf-8'))
        document.paragraphs[0]._p.append(run_element)
        buffer = io.BytesIO()
        document.save(buffer)
        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))
        self.assertEqual(len(runs), 1, 'Choice+Fallback duplicate content must resolve to one run, not one real run plus a bogus nested-run duplicate')
