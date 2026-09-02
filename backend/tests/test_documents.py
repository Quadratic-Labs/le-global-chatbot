"""Consolidated test module generated from validated domain owners."""

from __future__ import annotations



# ================================================================
# SOURCE: backend/tests/test_corpus_paths.py
# ================================================================

import os
import unittest
from pathlib import Path
from unittest.mock import patch
from tests.support.documents import DEFAULT_SOURCE_ROOT, resolve_source_root

class ResolveSourceRootTests(unittest.TestCase):

    def test_override_absent_returns_the_default_source_root(self) -> None:
        with patch.dict(os.environ, {}, clear=False):
            os.environ.pop('TEST_DOCUMENT_SOURCE_ROOT', None)
            self.assertEqual(resolve_source_root(), DEFAULT_SOURCE_ROOT)

    def test_override_present_returns_the_overridden_root(self) -> None:
        with patch.dict(os.environ, {'TEST_DOCUMENT_SOURCE_ROOT': '/var/tmp/some-sanitized-corpus'}):
            self.assertEqual(resolve_source_root(), Path('/var/tmp/some-sanitized-corpus'))

    def test_override_present_but_empty_still_uses_the_default(self) -> None:
        with patch.dict(os.environ, {'TEST_DOCUMENT_SOURCE_ROOT': ''}):
            self.assertEqual(resolve_source_root(), DEFAULT_SOURCE_ROOT)



# ================================================================
# SOURCE: backend/tests/test_document_chunk_builder.py
# ================================================================

import unittest
import zipfile
from pathlib import Path
from tempfile import TemporaryDirectory
from docx import Document
from app.core.country_registry import CountryMetadataMismatchError
from app.core.legal_taxonomy import get_canonical_legal_topic, normalize_topic
from app.services.document_chunk_builder import AmbiguousDocumentCountryError, DocumentMetadata, InvalidDocxFormatError, UndeterminableDocumentCountryError, UnknownLegalTopicError, build_document_chunks, build_document_chunks_from_docx, metadata_from_content, validate_docx_format
from app.services.document_chunk_builder import write_country_marker
from app.services.docx_parser import ParsedSection

def _test_document_chunk_builder__build_docx(directory: Path, title_lines: list[str], body_paragraphs: list[str] | None=None, filename: str='document.docx') -> Path:
    """
    A minimal real DOCX whose leading paragraphs are exactly
    `title_lines` - the title/cover area metadata_from_content scans
    - followed by any extra body content, saved under an arbitrary
    filename (never itself a source of metadata).
    """
    document = Document()
    for line in title_lines:
        document.add_paragraph(line)
    for paragraph in body_paragraphs or []:
        document.add_paragraph(paragraph)
    file_path = directory / filename
    document.save(file_path)
    return file_path
_DOCUMENT_XML_PATH = 'word/document.xml'

def _inject_body_xml_fragment(file_path: Path, xml_fragment: str) -> None:
    """
    Insert a raw XML fragment as the first child of <w:body> in an
    existing, already-saved DOCX - the only way to build a fixture
    containing structures python-docx has no high-level API for
    (a real Word text box's w:txbxContent, DrawingML a:t runs), since
    both are anchored drawings rather than ordinary paragraph content.
    """
    with zipfile.ZipFile(file_path) as archive:
        entries = {name: archive.read(name) for name in archive.namelist()}
    document_xml = entries[_DOCUMENT_XML_PATH].decode('utf-8')
    marker = '<w:body>'
    insertion_point = document_xml.index(marker) + len(marker)
    entries[_DOCUMENT_XML_PATH] = (document_xml[:insertion_point] + xml_fragment + document_xml[insertion_point:]).encode('utf-8')
    with zipfile.ZipFile(file_path, 'w') as archive:
        for name, data in entries.items():
            archive.writestr(name, data)

def _build_docx_with_text_box_title(directory: Path, text_box_text: str, body_paragraphs: list[str] | None=None, filename: str='document.docx') -> Path:
    """
    A minimal real DOCX whose only title-shaped text lives inside a
    Word text box (w:txbxContent) - invisible to python-docx's own
    paragraph iteration, exactly like 8 of the 10 real production
    documents this mission adds support for.
    """
    document = Document()
    for paragraph in body_paragraphs or ['Introduction']:
        document.add_paragraph(paragraph)
    file_path = directory / filename
    document.save(file_path)
    _inject_body_xml_fragment(file_path, f'<w:p><w:txbxContent><w:p><w:r><w:t>{text_box_text}</w:t></w:r></w:p></w:txbxContent></w:p>')
    return file_path

def _build_docx_with_drawingml_title(directory: Path, drawingml_text: str, body_paragraphs: list[str] | None=None, filename: str='document.docx') -> Path:
    """
    A minimal real DOCX whose only title-shaped text is a DrawingML
    (SmartArt/WordArt) a:t run - a third XML namespace, distinct from
    both ordinary paragraphs and text boxes.
    """
    document = Document()
    for paragraph in body_paragraphs or ['Introduction']:
        document.add_paragraph(paragraph)
    file_path = directory / filename
    document.save(file_path)
    _inject_body_xml_fragment(file_path, f'<w:p><w:r><a:t xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main">{drawingml_text}</a:t></w:r></w:p>')
    return file_path

def _build_docx_with_header_title(directory: Path, header_text: str, body_paragraphs: list[str] | None=None, filename: str='document.docx') -> Path:
    """A minimal real DOCX whose only title-shaped text is in a header."""
    document = Document()
    document.sections[0].header.paragraphs[0].text = header_text
    for paragraph in body_paragraphs or ['Introduction']:
        document.add_paragraph(paragraph)
    file_path = directory / filename
    document.save(file_path)
    return file_path

class LegalTaxonomyTests(unittest.TestCase):

    def test_normalizes_numbered_topic(self) -> None:
        self.assertEqual(normalize_topic(section='01. Hiring Practices', country='Spain'), 'Hiring Practices')

    def test_removes_country_suffix(self) -> None:
        self.assertEqual(normalize_topic(section='06. Social Media and Data Privacy in Spain', country='Spain'), 'Social Media and Data Privacy')

    def test_returns_canonical_topic(self) -> None:
        self.assertEqual(get_canonical_legal_topic(section='07. Termination of Employment Contracts', country='Spain'), 'Termination of Employment Contracts')

    def test_roman_numeral_prefix_is_stripped(self) -> None:
        self.assertEqual(get_canonical_legal_topic(section='II.  Hiring Practices', country='Taiwan'), 'Hiring Practices')

    def test_wages_and_work_hours_is_a_working_conditions_alias(self) -> None:
        self.assertEqual(get_canonical_legal_topic(section='03. Wages and Work Hours', country='United States'), 'Working Conditions')

    def test_termination_of_employment_without_contracts_is_an_alias(self) -> None:
        self.assertEqual(get_canonical_legal_topic(section='Termination of Employment', country='Philippines'), 'Termination of Employment Contracts')

class DocumentChunkBuilderTests(unittest.TestCase):

    def setUp(self) -> None:
        self.metadata = DocumentMetadata(country='Spain', country_code='ES', reference_year=2026, language='en', source_filename='Labour and Employment Law in Spain 2026.docx')

    def test_builds_overview_and_comparator_chunks(self) -> None:
        parsed_sections = [ParsedSection(section='Employment Law Overview Spain', subsection='Introduction', content='Overview content.'), ParsedSection(section='01. Hiring Practices', subsection='Requirement for Foreign Employees to Work', content='Hiring content.')]
        chunks = build_document_chunks(parsed_sections=parsed_sections, metadata=self.metadata)
        self.assertEqual(len(chunks), 2)
        overview_chunk = chunks[0]
        self.assertEqual(overview_chunk.document_type, 'overview')
        self.assertIsNone(overview_chunk.legal_topic)
        comparator_chunk = chunks[1]
        self.assertEqual(comparator_chunk.document_type, 'comparator')
        self.assertEqual(comparator_chunk.legal_topic, 'Hiring Practices')
        self.assertEqual(comparator_chunk.country, 'Spain')
        self.assertEqual(comparator_chunk.country_code, 'ES')
        self.assertEqual(comparator_chunk.reference_year, 2026)

    def test_ids_are_deterministic(self) -> None:
        original_sections = [ParsedSection(section='02. Employment Contracts', subsection='Notice Period', content='Original legal content.')]
        updated_sections = [ParsedSection(section='02. Employment Contracts', subsection='Notice Period', content='Updated legal content.')]
        original_chunks = build_document_chunks(parsed_sections=original_sections, metadata=self.metadata)
        repeated_chunks = build_document_chunks(parsed_sections=original_sections, metadata=self.metadata)
        updated_chunks = build_document_chunks(parsed_sections=updated_sections, metadata=self.metadata)
        self.assertEqual(original_chunks[0].document_id, repeated_chunks[0].document_id)
        self.assertEqual(original_chunks[0].chunk_id, repeated_chunks[0].chunk_id)
        self.assertEqual(original_chunks[0].chunk_id, updated_chunks[0].chunk_id)
        self.assertNotEqual(original_chunks[0].content_hash, updated_chunks[0].content_hash)

    def test_rejects_unknown_legal_topic(self) -> None:
        parsed_sections = [ParsedSection(section='12. Imaginary Legal Topic', subsection='Unknown subsection', content='Content that must not be indexed.')]
        with self.assertRaises(UnknownLegalTopicError):
            build_document_chunks(parsed_sections=parsed_sections, metadata=self.metadata)

    def test_extracts_metadata_from_content(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Labour and Employment Law in Spain 2026'], filename='final.docx')
            metadata = metadata_from_content(file_path=file_path, country_code='es')
        self.assertEqual(metadata.country, 'Spain')
        self.assertEqual(metadata.country_code, 'ES')
        self.assertEqual(metadata.reference_year, 2026)
        self.assertEqual(metadata.language, 'en')
        self.assertEqual(metadata.source_filename, 'final.docx')

    def test_rejects_content_with_no_identifiable_country(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Some random legal memo with no title structure.'])
            with self.assertRaises(UndeterminableDocumentCountryError):
                metadata_from_content(file_path=file_path)

    def test_title_with_leading_definite_article_resolves_country(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Labour and employment law in the Czech Republic'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Czech Republic')
        self.assertEqual(metadata.country_code, 'CZ')

    def test_uk_content_uses_canonical_country_name(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Labour and Employment Law in UK 2026'])
            metadata = metadata_from_content(file_path=file_path, country_code='GB')
        self.assertEqual(metadata.country, 'United Kingdom')
        self.assertEqual(metadata.country_code, 'GB')
        self.assertEqual(metadata.reference_year, 2026)

    def test_infers_country_code_from_content(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Labour and Employment Law in Sweden 2026'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Sweden')
        self.assertEqual(metadata.country_code, 'SE')
        self.assertEqual(metadata.reference_year, 2026)

    def test_accepts_overview_title_without_year(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Australia'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Australia')
        self.assertEqual(metadata.country_code, 'AU')
        self.assertIsNone(metadata.reference_year)

    def test_an_arbitrary_filename_never_affects_detection(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Canada 2026'], filename='Spain-template-used-for-Canada.docx')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Canada')
        self.assertEqual(metadata.country_code, 'CA')
        self.assertEqual(metadata.source_filename, 'Spain-template-used-for-Canada.docx')

    def test_rejects_content_country_code_mismatch(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Labour and Employment Law in Spain 2026'])
            with self.assertRaises(CountryMetadataMismatchError):
                metadata_from_content(file_path=file_path, country_code='GB')

class ContentMetadataFixtureTests(unittest.TestCase):
    """
    Mission "CONTINUATION PATCH 0.4.3", section 15 - the 8 mandatory
    filename/content/expected-outcome fixture cases A-H.
    """

    def test_case_a_plain_filename_with_title_country_and_year(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Canada 2026'], filename='final.docx')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Canada')
        self.assertEqual(metadata.country_code, 'CA')
        self.assertEqual(metadata.reference_year, 2026)

    def test_case_b_filename_names_a_different_country_than_content(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Labour and Employment Law in Canada 2025'], filename='Spain-final.docx')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Canada')
        self.assertEqual(metadata.country_code, 'CA')
        self.assertEqual(metadata.reference_year, 2025)
        self.assertEqual(metadata.source_filename, 'Spain-final.docx')

    def test_case_c_edited_replacement_filename_with_matching_content(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Canada 2026'], filename='Canada_2026-04-15-Employment-Law-Overview-EDITED.docx')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Canada')
        self.assertEqual(metadata.country_code, 'CA')
        self.assertEqual(metadata.reference_year, 2026)

    def test_case_d_generic_filename_with_no_year_in_content(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Spain'], filename='document_received_from_client.docx')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Spain')
        self.assertEqual(metadata.country_code, 'ES')
        self.assertIsNone(metadata.reference_year)

    def test_case_e_filename_names_a_country_but_content_names_none(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Some random legal memo with no title structure.'], filename='Peru-2026.docx')
            with self.assertRaises(UndeterminableDocumentCountryError):
                metadata_from_content(file_path=file_path)

    def test_case_f_ambiguous_cover_naming_two_countries_is_refused(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Spain 2026', 'Employment Law Overview Canada 2026'])
            with self.assertRaises(AmbiguousDocumentCountryError):
                metadata_from_content(file_path=file_path)

    def test_case_g_body_mentions_of_other_countries_are_ignored(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Canada 2026'], body_paragraphs=['This overview also references comparable rules in France, Germany, and Peru for context.', 'See also the equivalent Spain and Australia frameworks discussed elsewhere.'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Canada')
        self.assertEqual(metadata.country_code, 'CA')

    def test_case_h_content_year_wins_over_filename_year(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Canada 2025'], filename='Canada-2026-report.docx')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country_code, 'CA')
        self.assertEqual(metadata.reference_year, 2025)

def _valid_docx_entries(directory: Path) -> dict[str, bytes]:
    """The raw ZIP entries of one real, minimal, valid DOCX."""
    valid_path = directory / 'valid-reference.docx'
    document = Document()
    document.add_paragraph('Employment Law Overview Canada 2026')
    document.save(valid_path)
    with zipfile.ZipFile(valid_path) as archive:
        return {name: archive.read(name) for name in archive.namelist()}

class InvalidDocxFormatTests(unittest.TestCase):
    """
    Mission "CONTINUATION PATCH 0.4.3", section 16 - the .docx
    extension alone proves nothing: each of these must be refused by
    validate_docx_format before any content parsing is attempted.
    """

    def test_renamed_text_file_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = Path(directory) / 'renamed.docx'
            file_path.write_bytes(b'This is just plain text, not a docx at all.')
            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_invalid_zip_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = Path(directory) / 'invalid-zip.docx'
            file_path.write_bytes(b'PK\x03\x04' + b'\x00' * 50)
            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_empty_file_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = Path(directory) / 'empty.docx'
            file_path.write_bytes(b'')
            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_zip_missing_content_types_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            entries = _valid_docx_entries(Path(directory))
            file_path = Path(directory) / 'missing-content-types.docx'
            with zipfile.ZipFile(file_path, 'w') as archive:
                for name, data in entries.items():
                    if name == '[Content_Types].xml':
                        continue
                    archive.writestr(name, data)
            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_zip_missing_document_xml_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            entries = _valid_docx_entries(Path(directory))
            file_path = Path(directory) / 'missing-document-xml.docx'
            with zipfile.ZipFile(file_path, 'w') as archive:
                for name, data in entries.items():
                    if name == 'word/document.xml':
                        continue
                    archive.writestr(name, data)
            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

    def test_corrupted_document_xml_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            entries = dict(_valid_docx_entries(Path(directory)))
            entries['word/document.xml'] = b'<this is not valid xml'
            file_path = Path(directory) / 'corrupted.docx'
            with zipfile.ZipFile(file_path, 'w') as archive:
                for name, data in entries.items():
                    archive.writestr(name, data)
            with self.assertRaises(InvalidDocxFormatError):
                validate_docx_format(file_path)

class RealDocumentStructureTests(unittest.TestCase):
    """
    Mission "HOTFIX 0.4.4", Mission 2/2, section 11 - the real-world
    document structures the original 17-document corpus never used:
    title in a text box, DrawingML, header, a document with no
    standard cover, roman-numeral topic headings, decomposed Unicode,
    a law's own year ignored, filename ignored, and an ambiguous
    cover - each proven with a minimal synthetic fixture, mirroring
    the 10 real production documents this mission validates.
    """

    def test_title_in_a_text_box_is_found(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx_with_text_box_title(Path(directory), 'Employment Law Overview Chile')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Chile')
        self.assertEqual(metadata.country_code, 'CL')

    def test_title_in_a_text_box_with_en_dash_separator(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx_with_text_box_title(Path(directory), 'Employment Law Overview – Taiwan')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Taiwan')
        self.assertEqual(metadata.country_code, 'TW')

    def test_title_in_drawingml_is_found(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx_with_drawingml_title(Path(directory), 'Employment Law Overview Ireland')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Ireland')
        self.assertEqual(metadata.country_code, 'IE')

    def test_title_in_a_header_is_found(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _build_docx_with_header_title(Path(directory), 'Employment Law Overview Colombia')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Colombia')
        self.assertEqual(metadata.country_code, 'CO')

    def test_all_caps_title_with_no_separator_is_found(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['EMPLOYMENT LAW OVERVIEW PHILIPPINES'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Philippines')
        self.assertEqual(metadata.country_code, 'PH')

    def test_reversed_cover_country_then_plural_heading_with_year_range(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['FRANCE', 'EMPLOYMENT LAW OVERVIEWS 2025 - 2026'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'France')
        self.assertEqual(metadata.country_code, 'FR')

    def test_reversed_cover_country_then_heading_with_template_boilerplate(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['PORTUGAL', 'COUNTRY-SPECIFIC EMPLOYMENT LAW OVERVIEWS 2026 TEMPLATE'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Portugal')
        self.assertEqual(metadata.country_code, 'PT')

    def test_reversed_cover_fallback_never_silently_misdetects_an_unrelated_country(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Spain', 'Employment Law Overview Canada'])
            with self.assertRaises(AmbiguousDocumentCountryError):
                metadata_from_content(file_path=file_path)

    def test_document_with_no_standard_cover_is_refused(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Table of Contents', 'Some unrelated preamble text.'])
            with self.assertRaises(UndeterminableDocumentCountryError):
                metadata_from_content(file_path=file_path)

    def test_decomposed_unicode_turkiye_is_recognized(self) -> None:
        decomposed_turkiye = 'Türkiye'
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), [f'Employment Law Overview {decomposed_turkiye}'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Türkiye')
        self.assertEqual(metadata.country_code, 'TR')

    def test_a_laws_own_year_is_never_used_as_reference_year(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Chile'], body_paragraphs=['The Labour Code was originally enacted in 1984 and has been amended many times since.'])
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country_code, 'CL')
        self.assertIsNone(metadata.reference_year)

    def test_filename_country_and_year_are_ignored(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Colombia'], filename='Chile_2020_FINAL-EDITED.docx')
            metadata = metadata_from_content(file_path=file_path)
        self.assertEqual(metadata.country, 'Colombia')
        self.assertEqual(metadata.country_code, 'CO')
        self.assertIsNone(metadata.reference_year)
        self.assertEqual(metadata.source_filename, 'Chile_2020_FINAL-EDITED.docx')

    def test_ambiguous_cover_among_new_countries_is_refused(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Chile', 'Employment Law Overview Colombia'])
            with self.assertRaises(AmbiguousDocumentCountryError):
                metadata_from_content(file_path=file_path)

    def test_roman_numeral_topic_heading_is_recognized(self) -> None:
        parsed_sections = [ParsedSection(section='II.  Hiring Practices', subsection='1. Introduction', content='Hiring content.')]
        metadata = DocumentMetadata(country='Taiwan', country_code='TW', reference_year=None, language='en', source_filename='taiwan.docx')
        chunks = build_document_chunks(parsed_sections=parsed_sections, metadata=metadata)
        topics = {chunk.legal_topic for chunk in chunks if chunk.legal_topic is not None}
        self.assertIn('Hiring Practices', topics)

    def test_roman_numeral_heading_recognized_end_to_end_from_real_docx(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Taiwan'], body_paragraphs=['I. GENERAL OVERVIEW', '1. Introduction', 'Some overview content.', 'II.  Hiring Practices', 'Hiring content for Taiwan.'])
            chunks = build_document_chunks_from_docx(file_path=file_path, country_code='TW', language='en')
        topics = {chunk.legal_topic for chunk in chunks if chunk.legal_topic is not None}
        self.assertIn('Hiring Practices', topics)

    def test_usa_jurisdiction_suffix_heading_text_is_unchanged(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview United States'], body_paragraphs=['06. Social Media and Data Privacy in the USA', 'Content about social media rules.'])
            chunks = build_document_chunks_from_docx(file_path=file_path, country_code='US', language='en')
        matching_chunks = [chunk for chunk in chunks if chunk.legal_topic == 'Social Media and Data Privacy']
        self.assertTrue(matching_chunks)
        self.assertEqual(matching_chunks[0].section, '06. Social Media and Data Privacy in the USA')

class CountryMarkerPriorityTests(unittest.TestCase):
    """
    Mission "ORDER 8E-A1", section 11 - a valid DOCX-native country
    marker always takes priority over content detection, but never
    prevents content from being scanned for a reference_year, and
    never disturbs detection for a document that has no marker at all
    (every one of the 33 real production documents today).
    """

    def test_marker_resolves_a_document_with_no_detectable_country(self) -> None:
        with TemporaryDirectory() as directory:
            countryless = _test_document_chunk_builder__build_docx(Path(directory), ['Some heading with no recognizable country.'], body_paragraphs=['Body content.'])
            with self.assertRaises(UndeterminableDocumentCountryError):
                metadata_from_content(countryless)
            marked = Path(directory) / 'marked.docx'
            write_country_marker(countryless, marked, country_code='fr', country_name='France')
            metadata = metadata_from_content(marked)
            self.assertEqual(metadata.country_code, 'FR')
            self.assertEqual(metadata.country, 'France')

    def test_marker_wins_even_when_content_would_detect_differently(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Canada'])
            marked = Path(directory) / 'marked.docx'
            write_country_marker(file_path, marked, country_code='de', country_name='Germany')
            metadata = metadata_from_content(marked)
            self.assertEqual(metadata.country_code, 'DE')

    def test_marker_does_not_prevent_year_detection_from_content(self) -> None:
        with TemporaryDirectory() as directory:
            countryless_with_year = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Elbonia 2027'])
            marked = Path(directory) / 'marked.docx'
            write_country_marker(countryless_with_year, marked, country_code='jp', country_name='Japan')
            metadata = metadata_from_content(marked)
            self.assertEqual(metadata.country_code, 'JP')
            self.assertEqual(metadata.reference_year, 2027)

    def test_documents_without_a_marker_are_completely_unaffected(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Employment Law Overview Spain 2026'])
            metadata = metadata_from_content(file_path)
            self.assertEqual(metadata.country_code, 'ES')
            self.assertEqual(metadata.reference_year, 2026)

    def test_explicit_country_code_mismatch_against_a_marker_is_rejected(self) -> None:
        with TemporaryDirectory() as directory:
            file_path = _test_document_chunk_builder__build_docx(Path(directory), ['Some heading with no recognizable country.'])
            marked = Path(directory) / 'marked.docx'
            write_country_marker(file_path, marked, country_code='fr', country_name='France')
            with self.assertRaises(CountryMetadataMismatchError):
                metadata_from_content(marked, country_code='DE')

    def test_marker_survives_full_chunk_building_end_to_end(self) -> None:
        with TemporaryDirectory() as directory:
            countryless = _test_document_chunk_builder__build_docx(Path(directory), ['Some heading with no recognizable country.'], body_paragraphs=['01. Hiring Practices', 'Content about hiring.'])
            marked = Path(directory) / 'marked.docx'
            write_country_marker(countryless, marked, country_code='be', country_name='Belgium')
            chunks = build_document_chunks_from_docx(marked)
            self.assertTrue(chunks)
            self.assertTrue(all((chunk.country_code == 'BE' for chunk in chunks)))



# ================================================================
# SOURCE: backend/tests/test_document_indexer_pagination.py
# ================================================================

import bisect
import unittest
from typing import Any
from opensearchpy.exceptions import OpenSearchException
from app.services.document_indexer import _EXHAUSTIVE_FETCH_PAGE_SIZE, DocumentIndexingError, _fetch_all_chunks

class PaginatingFakeOpenSearch:
    """
    Genuinely paginates a fixed set of chunk_ids via search_after, on
    the same "chunk_id asc" sort _fetch_all_chunks itself requests -
    mirrors real OpenSearch 3.7's response shape exactly (hits.total,
    per-hit "sort" arrays), not an imagined one (mission "ORDER 3B",
    section 6).
    """

    def __init__(self, *, chunk_ids: list[str], reported_total: int | None=None) -> None:
        self._all_ids = sorted(chunk_ids)
        self._reported_total = reported_total if reported_total is not None else len(self._all_ids)
        self.search_calls = 0
        self.raise_on_call: int | None = None
        self.inject_duplicate_on_call: int | None = None
        self.force_empty_intermediate_page_on_call: int | None = None

    def search(self, *, index: str, body: dict[str, Any]) -> dict[str, Any]:
        del index
        self.search_calls += 1
        if self.raise_on_call == self.search_calls:
            raise OpenSearchException('simulated exhaustive fetch page failure')
        size = body['size']
        search_after = body.get('search_after')
        if search_after is None:
            start_index = 0
        else:
            cursor = search_after[0]
            start_index = bisect.bisect_right(self._all_ids, cursor)
        if self.force_empty_intermediate_page_on_call == self.search_calls:
            page_ids: list[str] = []
        else:
            page_ids = self._all_ids[start_index:start_index + size]
        hits = [{'_id': chunk_id, '_source': {'chunk_id': chunk_id, 'document_id': 'doc_' + 'a' * 64, 'country_code': 'ZZ'}, 'sort': [chunk_id]} for chunk_id in page_ids]
        if self.inject_duplicate_on_call == self.search_calls and hits:
            hits.append(dict(hits[0]))
        return {'hits': {'total': {'value': self._reported_total}, 'hits': hits}}

def _chunk_ids(count: int) -> list[str]:
    return [f'chunk_{i:06d}' for i in range(count)]

class FetchAllChunksPaginationTests(unittest.TestCase):

    def test_zero_hits(self) -> None:
        client = PaginatingFakeOpenSearch(chunk_ids=[])
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(result, [])
        self.assertEqual(client.search_calls, 1)

    def test_one_hit(self) -> None:
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(1))
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(len(result), 1)

    def test_page_size_minus_one(self) -> None:
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(_EXHAUSTIVE_FETCH_PAGE_SIZE - 1))
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(len(result), _EXHAUSTIVE_FETCH_PAGE_SIZE - 1)
        self.assertEqual(client.search_calls, 1)

    def test_exactly_page_size(self) -> None:
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(_EXHAUSTIVE_FETCH_PAGE_SIZE))
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(len(result), _EXHAUSTIVE_FETCH_PAGE_SIZE)
        self.assertEqual(client.search_calls, 2)

    def test_page_size_plus_one(self) -> None:
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(_EXHAUSTIVE_FETCH_PAGE_SIZE + 1))
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(len(result), _EXHAUSTIVE_FETCH_PAGE_SIZE + 1)
        self.assertEqual(client.search_calls, 2)

    def test_several_pages(self) -> None:
        count = _EXHAUSTIVE_FETCH_PAGE_SIZE * 3 + 250
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(count))
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(len(result), count)
        self.assertEqual({hit['_id'] for hit in result}, set(_chunk_ids(count)))

    def test_exactly_10000(self) -> None:
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(10000))
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(len(result), 10000)

    def test_10001_the_original_boundary(self) -> None:
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(10001))
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(len(result), 10001)

    def test_14083_the_real_mission_document(self) -> None:
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(14083))
        result = _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertEqual(len(result), 14083)
        self.assertEqual({hit['_id'] for hit in result}, set(_chunk_ids(14083)))

    def test_incoherent_intermediate_empty_page_is_rejected(self) -> None:
        count = _EXHAUSTIVE_FETCH_PAGE_SIZE + 500
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(count))
        client.force_empty_intermediate_page_on_call = 2
        with self.assertRaises(DocumentIndexingError) as context:
            _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertIn('did not exhaust', str(context.exception))

    def test_exception_on_page_2_propagates(self) -> None:
        count = _EXHAUSTIVE_FETCH_PAGE_SIZE + 10
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(count))
        client.raise_on_call = 2
        with self.assertRaises(DocumentIndexingError) as context:
            _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertIn('exhaustive chunk fetch failed', str(context.exception))

    def test_exception_on_final_page_propagates(self) -> None:
        count = _EXHAUSTIVE_FETCH_PAGE_SIZE * 2
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(count))
        client.raise_on_call = 3
        with self.assertRaises(DocumentIndexingError):
            _fetch_all_chunks(client=client, field='document_id', value='doc_x')

    def test_duplicate_hit_across_pages_is_rejected(self) -> None:
        count = _EXHAUSTIVE_FETCH_PAGE_SIZE + 10
        client = PaginatingFakeOpenSearch(chunk_ids=_chunk_ids(count))
        client.inject_duplicate_on_call = 2
        with self.assertRaises(DocumentIndexingError) as context:
            _fetch_all_chunks(client=client, field='document_id', value='doc_x')
        self.assertIn('same chunk twice across pages', str(context.exception))

    def test_scroll_contexts_not_applicable(self) -> None:
        import app.services.document_indexer as document_indexer_module
        source = document_indexer_module.__file__
        with open(source) as f:
            contents = f.read()
        self.assertNotIn('scroll', contents.lower())



# ================================================================
# SOURCE: backend/tests/test_document_mutation.py
# ================================================================

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from docx import Document
from app.services.docx_parser import locate_top_level_topics, parse_docx_sections
from app.services.document_mutation import InvalidSectionPositionError, LegalTopicAlreadyExistsError, LegalTopicNotFoundError, NoAnchorTopicError, insert_top_level_topic, normalize_topic_title, replace_top_level_topic

def _write_docx(path: Path, sections: list[tuple[str, str]]) -> None:
    document = Document()
    document.add_heading('Employment Law Overview United Kingdom', level=1)
    for heading, content in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(content)
    document.save(path)

def _topic_names(path: Path, country: str='United Kingdom') -> list[str]:
    return [location.legal_topic for location in locate_top_level_topics(Document(path), country=country)]

def _write_bold_only_docx(path: Path, sections: list[tuple[str, str]]) -> None:
    """
    A legacy-style document whose native topics use only direct bold
    run formatting - no Heading 1 style, no numbering - representative
    of the ~10/33 real corpus documents ORDER 8A-C's marker style
    exists to support.
    """
    document = Document()
    document.add_paragraph('Employment Law Overview United Kingdom')
    for heading, content in sections:
        heading_paragraph = document.add_paragraph()
        heading_paragraph.add_run(heading).bold = True
        document.add_paragraph(content)
    document.save(path)

class NormalizeTopicTitleTests(unittest.TestCase):

    def test_case_and_whitespace_insensitive(self) -> None:
        self.assertEqual(normalize_topic_title('Hiring Practices'), normalize_topic_title('hiring   PRACTICES'))

    def test_ignores_leading_numeric_prefix(self) -> None:
        self.assertEqual(normalize_topic_title('01. Hiring Practices'), normalize_topic_title('Hiring Practices'))

class ReplaceTopLevelTopicTests(unittest.TestCase):

    def test_replaces_content_and_preserves_heading(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'Original content.')])
            replace_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', legal_topic='Hiring Practices', new_content='New content.\n\nSecond paragraph.')
            paragraphs = [p.text for p in Document(output_path).paragraphs]
            self.assertIn('Hiring Practices', paragraphs)
            self.assertIn('New content.', paragraphs)
            self.assertIn('Second paragraph.', paragraphs)
            self.assertNotIn('Original content.', paragraphs)

    def test_unrelated_topics_are_byte_identical(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'HP content.'), ('Employment Contracts', 'EC content.')])
            replace_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', legal_topic='Hiring Practices', new_content='Edited HP content.')
            sections = parse_docx_sections(output_path, country='United Kingdom')
            ec_sections = [s for s in sections if s.section == 'Employment Contracts']
            self.assertEqual(len(ec_sections), 1)
            self.assertEqual(ec_sections[0].content, 'EC content.')

    def test_missing_topic_raises_not_found(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'Original content.')])
            with self.assertRaises(LegalTopicNotFoundError):
                replace_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', legal_topic='Employment Contracts', new_content='New content.')
            self.assertFalse(output_path.exists())

class InsertTopLevelTopicTests(unittest.TestCase):

    def test_insert_at_end(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'HP content.')])
            insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='Remote Working', content='Remote work content.', position='end')
            names = _topic_names(output_path)
            self.assertEqual(names, ['Hiring Practices', 'Remote Working'])

    def test_insert_at_beginning(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'HP content.')])
            insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='Remote Working', content='Remote work content.', position='beginning')
            names = _topic_names(output_path)
            self.assertEqual(names, ['Remote Working', 'Hiring Practices'])

    def test_insert_after_existing_topic(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'HP content.'), ('Employment Contracts', 'EC content.')])
            insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='Remote Working', content='Remote work content.', position='after:Hiring Practices')
            names = _topic_names(output_path)
            self.assertEqual(names, ['Hiring Practices', 'Remote Working', 'Employment Contracts'])

    def test_duplicate_title_is_rejected(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'HP content.')])
            with self.assertRaises(LegalTopicAlreadyExistsError):
                insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='hiring   PRACTICES', content='whatever', position='end')
            self.assertFalse(output_path.exists())

    def test_invalid_position_is_rejected(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'HP content.')])
            with self.assertRaises(InvalidSectionPositionError):
                insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='Remote Working', content='whatever', position='middle')

    def test_no_anchor_topic_raises(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            document = Document()
            document.add_paragraph('Just some plain text, no topics.')
            document.save(source_path)
            with self.assertRaises(NoAnchorTopicError):
                insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='Remote Working', content='whatever', position='end')

class AdminSectionStyleMarkerTests(unittest.TestCase):
    """
    ORDER 8A-C - the internal DOCX-native style marker: created once,
    reused afterward, works identically regardless of the surrounding
    document's own native heading convention.
    """

    def test_insert_on_bold_only_document(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_bold_only_docx(source_path, [('Hiring Practices', 'HP content.')])
            insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='Remote Working', content='Remote work content.', position='end')
            names = _topic_names(output_path)
            self.assertEqual(names, ['Hiring Practices', 'Remote Working'])
            sections = parse_docx_sections(output_path, country='United Kingdom')
            hiring = [s for s in sections if s.section == 'Hiring Practices']
            self.assertEqual(len(hiring), 1)
            self.assertEqual(hiring[0].content, 'HP content.')

    def test_marker_style_created_once_and_reused(self) -> None:
        from app.services.docx_parser import ADMIN_SECTION_STYLE_NAME
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            first_output = Path(temp_dir) / 'first.docx'
            second_output = Path(temp_dir) / 'second.docx'
            _write_bold_only_docx(source_path, [('Hiring Practices', 'HP content.')])
            insert_top_level_topic(file_path=source_path, output_path=first_output, country='United Kingdom', title='Remote Working', content='Remote content.', position='end')
            first_styles = [style.name for style in Document(first_output).styles]
            self.assertEqual(first_styles.count(ADMIN_SECTION_STYLE_NAME), 1)
            insert_top_level_topic(file_path=first_output, output_path=second_output, country='United Kingdom', title='Artificial Intelligence at Work', content='AI content.', position='end')
            second_styles = [style.name for style in Document(second_output).styles]
            self.assertEqual(second_styles.count(ADMIN_SECTION_STYLE_NAME), 1)
            names = _topic_names(second_output)
            self.assertEqual(names, ['Hiring Practices', 'Remote Working', 'Artificial Intelligence at Work'])

    def test_visual_formatting_derived_from_bold_only_anchor(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_bold_only_docx(source_path, [('Hiring Practices', 'HP content.')])
            insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='Remote Working', content='Remote content.', position='end')
            reloaded = Document(output_path)
            new_heading = next((p for p in reloaded.paragraphs if p.text == 'Remote Working'))
            self.assertTrue(new_heading.style.font.bold)
            self.assertIsNotNone(new_heading.style.font.size)

    def test_visual_formatting_derived_from_heading1_anchor(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            output_path = Path(temp_dir) / 'output.docx'
            _write_docx(source_path, [('Hiring Practices', 'HP content.')])
            insert_top_level_topic(file_path=source_path, output_path=output_path, country='United Kingdom', title='Remote Working', content='Remote content.', position='end')
            reloaded = Document(output_path)
            heading_one = reloaded.styles['Heading 1']
            new_heading = next((p for p in reloaded.paragraphs if p.text == 'Remote Working'))
            self.assertEqual(new_heading.style.font.bold, heading_one.font.bold)
            self.assertEqual(new_heading.style.font.size, heading_one.font.size)

    def test_after_another_custom_section(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            first_output = Path(temp_dir) / 'first.docx'
            second_output = Path(temp_dir) / 'second.docx'
            _write_bold_only_docx(source_path, [('Hiring Practices', 'HP content.')])
            insert_top_level_topic(file_path=source_path, output_path=first_output, country='United Kingdom', title='Custom A', content='A content.', position='end')
            insert_top_level_topic(file_path=first_output, output_path=second_output, country='United Kingdom', title='Custom B', content='B content.', position='after:Custom A')
            names = _topic_names(second_output)
            self.assertEqual(names, ['Hiring Practices', 'Custom A', 'Custom B'])

    def test_download_reparse_roundtrip_no_external_state(self) -> None:
        import shutil
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / 'source.docx'
            added_path = Path(temp_dir) / 'added.docx'
            downloaded_path = Path(temp_dir) / 'downloaded_copy.docx'
            _write_bold_only_docx(source_path, [('Hiring Practices', 'HP content.')])
            insert_top_level_topic(file_path=source_path, output_path=added_path, country='United Kingdom', title='Remote Working', content='Remote content. ROUNDTRIP-MARKER.', position='end')
            shutil.copy(added_path, downloaded_path)
            from app.services.document_chunk_builder import build_document_chunks_from_docx
            names = _topic_names(downloaded_path)
            self.assertIn('Remote Working', names)
            sections = parse_docx_sections(downloaded_path, country='United Kingdom')
            remote_sections = [s for s in sections if s.section == 'Remote Working']
            self.assertEqual(len(remote_sections), 1)
            self.assertIn('ROUNDTRIP-MARKER', remote_sections[0].content)
            reuploaded_chunks = build_document_chunks_from_docx(downloaded_path)
            remote_chunks = [c for c in reuploaded_chunks if c.legal_topic == 'Remote Working']
            self.assertEqual(len(remote_chunks), 1)
            self.assertEqual(remote_chunks[0].document_type, 'comparator')



# ================================================================
# SOURCE: backend/tests/test_document_source_resolver.py
# ================================================================

import os
import tempfile
import unittest
from pathlib import Path
from app.services.document_source_resolver import DocumentSourceConflictError, resolve_document_source_path
REAL_PRODUCTION_FILENAMES = ('Employment Law Overview Australia.docx', 'Employment Law Overview Peru 2026.docx', 'Employment Law Overview Singapore 2026.docx', 'Labour and Employment Law in Argentina 2026.docx', 'Labour and Employment Law in Belgium 2026.docx', 'Labour and Employment Law in Brazil 2026.docx', 'Labour and Employment Law in Italy 2026.docx', 'Labour and Employment Law in Japan 2026.docx', 'Labour and Employment Law in Poland 2026.docx', 'Labour and Employment Law in Romania 2026.docx', 'Labour and Employment Law in Spain 2026.docx', 'Labour and Employment Law in Sweden 2026.docx', 'Labour and Employment Law in Switzerland 2026.docx', 'Labour and Employment Law in UK 2026.docx', 'Labour and employment law in Czech Republic 2026.docx', 'Labour and employment law in Greece.docx', 'Labour and employment law in Mexico 2026.docx')

class ResolveViaHistoricalFilenameTests(unittest.TestCase):

    def test_resolves_via_source_filename_when_canonical_absent(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            legacy_path = source_root / 'Labour and Employment Law in Spain 2026.docx'
            legacy_path.write_bytes(b'legacy-docx-bytes')
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename=legacy_path.name)
        self.assertEqual(resolved.path, legacy_path)
        self.assertEqual(resolved.origin, 'source_filename')

    def test_resolves_via_canonical_when_source_filename_absent(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            canonical_path = source_root / 'CA.docx'
            canonical_path.write_bytes(b'new-docx-bytes')
            resolved = resolve_document_source_path(source_root=source_root, country_code='CA', source_filename='Canada_2026-04-15-EDITED.docx')
        self.assertEqual(resolved.path, canonical_path)
        self.assertEqual(resolved.origin, 'canonical')

    def test_missing_when_neither_exists(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='Labour and Employment Law in Spain 2026.docx')
        self.assertIsNone(resolved.path)
        self.assertEqual(resolved.origin, 'missing')

    def test_storage_filename_takes_priority_when_only_one_exists(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            legacy_path = source_root / 'Labour and Employment Law in Spain 2026.docx'
            legacy_path.write_bytes(b'legacy')
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', storage_filename='does-not-exist-on-disk.docx', source_filename=legacy_path.name)
        self.assertEqual(resolved.path, legacy_path)
        self.assertEqual(resolved.origin, 'source_filename')

    def test_the_17_real_production_filenames_all_resolve(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            for filename in REAL_PRODUCTION_FILENAMES:
                with self.subTest(filename=filename):
                    file_path = source_root / filename
                    file_path.write_bytes(b'legacy-docx-bytes')
                    resolved = resolve_document_source_path(source_root=source_root, country_code='XX', source_filename=filename)
                    self.assertEqual(resolved.path, file_path)
                    self.assertEqual(resolved.origin, 'source_filename')
                    file_path.unlink()

class ConflictDetectionTests(unittest.TestCase):

    def test_conflict_when_two_distinct_files_exist(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            legacy_path = source_root / 'Labour and Employment Law in Spain 2026.docx'
            legacy_path.write_bytes(b'legacy')
            canonical_path = source_root / 'ES.docx'
            canonical_path.write_bytes(b'canonical')
            with self.assertRaises(DocumentSourceConflictError) as context:
                resolve_document_source_path(source_root=source_root, country_code='ES', source_filename=legacy_path.name)
        self.assertEqual(set(context.exception.conflicting_paths), {legacy_path, canonical_path})

    def test_no_conflict_when_two_fields_name_the_same_file(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            path = source_root / 'ES.docx'
            path.write_bytes(b'content')
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='ES.docx')
        self.assertEqual(resolved.path, path)

class PathSecurityTests(unittest.TestCase):

    def test_rejects_path_traversal(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root) / 'source'
            source_root.mkdir()
            outside_path = Path(root) / 'secret.docx'
            outside_path.write_bytes(b'secret')
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='../secret.docx')
        self.assertIsNone(resolved.path)
        self.assertEqual(resolved.origin, 'missing')

    def test_rejects_forward_slash(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='folder/document.docx')
        self.assertIsNone(resolved.path)

    def test_rejects_backslash(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='folder\\document.docx')
        self.assertIsNone(resolved.path)

    def test_rejects_null_byte(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='document\x00.docx')
        self.assertIsNone(resolved.path)

    def test_rejects_non_docx_extension(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            path = source_root / 'document.pdf'
            path.write_bytes(b'content')
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='document.pdf')
        self.assertIsNone(resolved.path)

    def test_rejects_symlink_escaping_source_root(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root) / 'source'
            source_root.mkdir()
            outside_target = Path(root) / 'outside.docx'
            outside_target.write_bytes(b'outside-content')
            symlink_path = source_root / 'escape.docx'
            try:
                os.symlink(outside_target, symlink_path)
            except OSError:
                self.skipTest('Symlinks are not supported in this environment.')
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='escape.docx')
        self.assertIsNone(resolved.path)

    def test_accepts_internal_symlink(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_root = Path(root)
            real_target = source_root / 'real.docx'
            real_target.write_bytes(b'content')
            symlink_path = source_root / 'alias.docx'
            try:
                os.symlink(real_target, symlink_path)
            except OSError:
                self.skipTest('Symlinks are not supported in this environment.')
            resolved = resolve_document_source_path(source_root=source_root, country_code='ES', source_filename='alias.docx')
        self.assertEqual(resolved.path, real_target)



# ================================================================
# SOURCE: backend/tests/test_document_warnings.py
# ================================================================

import unittest
from app.core.legal_taxonomy import LEGAL_TOPICS
from app.models.document import DocumentChunk
from app.services.admin_document_replacement import CONTEXT_WARNING_CODE, EXPECTED_TOPICS_COUNT, STRUCTURE_WARNING_CODE, evaluate_topic_coverage, recognized_topics_for

def _chunk(*, legal_topic: str | None, document_type: str='comparator') -> DocumentChunk:
    return DocumentChunk(document_id='doc_' + 'a' * 64, chunk_id='chunk_' + 'a' * 64, country='Testland', country_code='ZZ', legal_topic=legal_topic, document_type=document_type, language='en', section=legal_topic or 'General', subsection=None, content='Some content.', source_filename='test.docx', source_format='docx', content_hash='hash')

def _chunks_with_topics(count: int) -> list[DocumentChunk]:
    overview = _chunk(legal_topic=None, document_type='overview')
    topic_chunks = [_chunk(legal_topic=topic) for topic in LEGAL_TOPICS[:count]]
    return [overview, *topic_chunks]

class TopicCoverageThresholdTests(unittest.TestCase):

    def test_expected_topics_count_is_eleven(self) -> None:
        self.assertEqual(EXPECTED_TOPICS_COUNT, 11)
        self.assertEqual(len(LEGAL_TOPICS), 11)

    def test_zero_recognized_topics_is_context_warning(self) -> None:
        chunks = _chunks_with_topics(0)
        warning = evaluate_topic_coverage(chunks)
        self.assertIsNotNone(warning)
        self.assertEqual(warning.code, CONTEXT_WARNING_CODE)
        self.assertEqual(warning.recognized_topics_count, 0)
        self.assertEqual(warning.expected_topics_count, 11)
        self.assertEqual(len(warning.missing_topics), 11)

    def test_one_to_five_recognized_topics_is_structure_warning(self) -> None:
        for count in (1, 2, 3, 4, 5):
            with self.subTest(count=count):
                chunks = _chunks_with_topics(count)
                warning = evaluate_topic_coverage(chunks)
                self.assertIsNotNone(warning)
                self.assertEqual(warning.code, STRUCTURE_WARNING_CODE)
                self.assertEqual(warning.recognized_topics_count, count)
                self.assertEqual(len(warning.missing_topics), 11 - count)

    def test_six_or_more_recognized_topics_is_no_warning(self) -> None:
        for count in (6, 7, 8, 9, 10, 11):
            with self.subTest(count=count):
                chunks = _chunks_with_topics(count)
                self.assertIsNone(evaluate_topic_coverage(chunks))

    def test_overview_chunks_never_count_as_recognized_topics(self) -> None:
        chunks = [_chunk(legal_topic=None, document_type='overview') for _ in range(20)]
        warning = evaluate_topic_coverage(chunks)
        self.assertIsNotNone(warning)
        self.assertEqual(warning.code, CONTEXT_WARNING_CODE)

    def test_duplicate_topic_chunks_count_once(self) -> None:
        chunks = [_chunk(legal_topic='Hiring Practices') for _ in range(5)]
        recognized = recognized_topics_for(chunks)
        self.assertEqual(recognized, ('Hiring Practices',))
        warning = evaluate_topic_coverage(chunks)
        self.assertIsNotNone(warning)
        self.assertEqual(warning.recognized_topics_count, 1)

    def test_missing_topics_are_exactly_the_complement(self) -> None:
        chunks = _chunks_with_topics(3)
        warning = evaluate_topic_coverage(chunks)
        self.assertIsNotNone(warning)
        self.assertEqual(set(warning.recognized_topics) | set(warning.missing_topics), set(LEGAL_TOPICS))
        self.assertEqual(set(warning.recognized_topics) & set(warning.missing_topics), set())



# ================================================================
# SOURCE: backend/tests/test_docx_country_marker.py
# ================================================================

import tempfile
import unittest
import zipfile
from pathlib import Path
from docx import Document
from app.services.document_chunk_builder import CountryMarker, InvalidCountryMarkerValueError, read_country_marker, write_country_marker

def _test_docx_country_marker__build_docx(directory: Path, paragraphs: list[str]) -> Path:
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    path = directory / 'document.docx'
    document.save(path)
    return path

class ReadWithNoMarkerTests(unittest.TestCase):

    def test_document_with_no_custom_properties_part_reads_as_none(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            path = _test_docx_country_marker__build_docx(Path(root), ['No marker here.'])
            self.assertIsNone(read_country_marker(path))

class WriteAndReadRoundtripTests(unittest.TestCase):

    def test_write_then_read_returns_the_same_marker(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Some content.'])
            destination = Path(root) / 'marked.docx'
            write_country_marker(source, destination, country_code='fr', country_name='France')
            marker = read_country_marker(destination)
            self.assertEqual(marker, CountryMarker(country_code='FR', country_name='France'))

    def test_marker_survives_a_full_python_docx_load_and_save(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Some content.'])
            marked = Path(root) / 'marked.docx'
            write_country_marker(source, marked, country_code='de', country_name='Germany')
            document = Document(marked)
            document.add_paragraph('An unrelated edit.')
            edited = Path(root) / 'edited.docx'
            document.save(edited)
            self.assertEqual(read_country_marker(edited), CountryMarker(country_code='DE', country_name='Germany'))

    def test_marked_document_is_still_a_valid_docx(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Body paragraph one.'])
            marked = Path(root) / 'marked.docx'
            write_country_marker(source, marked, country_code='jp', country_name='Japan')
            reopened = Document(marked)
            self.assertEqual([p.text for p in reopened.paragraphs], ['Body paragraph one.'])

class DeterminismTests(unittest.TestCase):

    def test_writing_the_same_marker_twice_is_byte_identical(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Content.'])
            first = Path(root) / 'first.docx'
            second = Path(root) / 'second.docx'
            write_country_marker(source, first, country_code='ca', country_name='Canada')
            write_country_marker(source, second, country_code='ca', country_name='Canada')
            self.assertEqual(first.read_bytes(), second.read_bytes())

    def test_re_embedding_on_an_already_marked_file_is_idempotent(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Content.'])
            once = Path(root) / 'once.docx'
            twice = Path(root) / 'twice.docx'
            write_country_marker(source, once, country_code='ca', country_name='Canada')
            write_country_marker(once, twice, country_code='ca', country_name='Canada')
            self.assertEqual(once.read_bytes(), twice.read_bytes())

    def test_changing_the_country_updates_the_value_not_the_pid(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Content.'])
            first = Path(root) / 'first.docx'
            changed = Path(root) / 'changed.docx'
            write_country_marker(source, first, country_code='fr', country_name='France')
            write_country_marker(first, changed, country_code='de', country_name='Germany')
            with zipfile.ZipFile(first) as archive:
                first_xml = archive.read('docProps/custom.xml')
            with zipfile.ZipFile(changed) as archive:
                changed_xml = archive.read('docProps/custom.xml')
            self.assertIn(b'pid="2"', first_xml)
            self.assertIn(b'pid="2"', changed_xml)
            self.assertEqual(read_country_marker(changed), CountryMarker(country_code='DE', country_name='Germany'))

class InvalidMarkerTests(unittest.TestCase):

    def test_write_rejects_an_unrecognized_country_code(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Content.'])
            with self.assertRaises(InvalidCountryMarkerValueError):
                write_country_marker(source, Path(root) / 'invalid.docx', country_code='ZZ', country_name='Nowhere')

    def test_read_safely_ignores_a_corrupted_marker_value(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Content.'])
            marked = Path(root) / 'marked.docx'
            write_country_marker(source, marked, country_code='fr', country_name='France')
            with zipfile.ZipFile(marked) as archive:
                contents = {name: archive.read(name) for name in archive.namelist()}
            contents['docProps/custom.xml'] = contents['docProps/custom.xml'].replace(b'>FR<', b'>NOT-A-COUNTRY<')
            corrupted = Path(root) / 'corrupted.docx'
            with zipfile.ZipFile(corrupted, 'w', zipfile.ZIP_DEFLATED) as archive:
                for name, data in contents.items():
                    archive.writestr(name, data)
            self.assertIsNone(read_country_marker(corrupted))

    def test_write_rejects_an_empty_country_name(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Content.'])
            with self.assertRaises(InvalidCountryMarkerValueError):
                write_country_marker(source, Path(root) / 'invalid.docx', country_code='FR', country_name='   ')

class ForeignCustomPropertiesPreservedTests(unittest.TestCase):

    def test_a_pre_existing_unrelated_custom_property_survives(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source = _test_docx_country_marker__build_docx(Path(root), ['Content.'])
            with_foreign_property = Path(root) / 'foreign.docx'
            with zipfile.ZipFile(source) as archive:
                contents = {name: archive.read(name) for name in archive.namelist()}
            custom_xml = '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\r\n<Properties xmlns="http://schemas.openxmlformats.org/officeDocument/2006/custom-properties" xmlns:vt="http://schemas.openxmlformats.org/officeDocument/2006/docPropsVTypes"><property fmtid="{D5CDD505-2E9C-101B-9397-08002B2CF9AE}" pid="2" name="SomeOtherTool"><vt:lpwstr>keep-me</vt:lpwstr></property></Properties>'.encode('utf-8')
            contents['docProps/custom.xml'] = custom_xml
            contents['[Content_Types].xml'] = contents['[Content_Types].xml'].replace(b'</Types>', b'<Override PartName="/docProps/custom.xml" ContentType="application/vnd.openxmlformats-officedocument.custom-properties+xml"/></Types>')
            contents['_rels/.rels'] = contents['_rels/.rels'].replace(b'</Relationships>', b'<Relationship Id="rId99" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/custom-properties" Target="docProps/custom.xml"/></Relationships>')
            with zipfile.ZipFile(with_foreign_property, 'w', zipfile.ZIP_DEFLATED) as archive:
                for name, data in contents.items():
                    archive.writestr(name, data)
            marked = Path(root) / 'marked-with-foreign.docx'
            write_country_marker(with_foreign_property, marked, country_code='pt', country_name='Portugal')
            with zipfile.ZipFile(marked) as archive:
                final_custom_xml = archive.read('docProps/custom.xml')
            self.assertIn(b'SomeOtherTool', final_custom_xml)
            self.assertIn(b'keep-me', final_custom_xml)
            self.assertEqual(read_country_marker(marked), CountryMarker(country_code='PT', country_name='Portugal'))
            Document(marked)



# ================================================================
# SOURCE: backend/tests/test_docx_parser.py
# ================================================================

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory
from unittest.mock import patch
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from app.core.legal_taxonomy import get_canonical_legal_topic
from app.core.subsection_taxonomy import get_subsection_topic_override
from app.services.docx_parser import ExtractedContact, build_contact_chunk_content, extract_contacts_from_docx, find_plain_paragraph_contact_block_bounds, parse_contact_blocks, parse_docx_sections, split_combined_legacy_contact, _classify_canonical_firm_lines

def _mark_as_numbered(paragraph: Paragraph) -> None:
    """Mark a paragraph as a numbered-list item in Word XML."""
    paragraph_properties = paragraph._p.get_or_add_pPr()
    numbering_properties = OxmlElement('w:numPr')
    indentation_level = OxmlElement('w:ilvl')
    indentation_level.set(qn('w:val'), '0')
    numbering_id = OxmlElement('w:numId')
    numbering_id.set(qn('w:val'), '1')
    numbering_properties.append(indentation_level)
    numbering_properties.append(numbering_id)
    paragraph_properties.append(numbering_properties)

def _mark_as_explicitly_unbolded(paragraph: Paragraph) -> None:
    """Apply the malformed formatting found in some L&E DOCX files."""
    paragraph_properties = paragraph._p.get_or_add_pPr()
    run_properties = paragraph_properties.find(qn('w:rPr'))
    if run_properties is None:
        run_properties = OxmlElement('w:rPr')
        paragraph_properties.append(run_properties)
    bold_property = run_properties.find(qn('w:b'))
    if bold_property is None:
        bold_property = OxmlElement('w:b')
        run_properties.append(bold_property)
    bold_property.set(qn('w:val'), '0')

class DocxParserTests(unittest.TestCase):

    def test_groups_content_by_headings(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'sample-legal-document.docx'
            document = Document()
            document.add_paragraph('Introductory legal information.')
            document.add_heading('Employment contracts', level=1)
            document.add_paragraph('General information about employment contracts.')
            document.add_heading('Probationary period', level=2)
            document.add_paragraph('The probationary period depends on local law.')
            document.save(file_path)
            sections = parse_docx_sections(file_path)
            self.assertEqual(len(sections), 3)
            self.assertEqual(sections[0].section, 'General')
            self.assertIsNone(sections[0].subsection)
            self.assertEqual(sections[0].content, 'Introductory legal information.')
            self.assertEqual(sections[1].section, 'Employment contracts')
            self.assertIsNone(sections[1].subsection)
            self.assertEqual(sections[1].content, 'General information about employment contracts.')
            self.assertEqual(sections[2].section, 'Employment contracts')
            self.assertEqual(sections[2].subsection, 'Probationary period')
            self.assertEqual(sections[2].content, 'The probationary period depends on local law.')

    def test_preserves_tables_in_document_order(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'legal-table-document.docx'
            document = Document()
            document.add_heading('Termination', level=1)
            document.add_paragraph('Rules applicable before the table.')
            table = document.add_table(rows=3, cols=3)
            table.cell(0, 0).text = 'Country'
            table.cell(0, 1).text = 'Notice period'
            table.cell(0, 2).text = 'Written notice'
            table.cell(1, 0).text = 'France'
            table.cell(1, 1).text = '30 days'
            table.cell(1, 2).text = 'Required'
            table.cell(2, 0).text = 'Belgium'
            table.cell(2, 1).text = '45 days'
            table.cell(2, 2).text = 'Required'
            document.add_paragraph('Additional rules after the table.')
            document.save(file_path)
            sections = parse_docx_sections(file_path)
            self.assertEqual(len(sections), 1)
            self.assertEqual(sections[0].content, 'Rules applicable before the table.\n\nCountry | Notice period | Written notice\nFrance | 30 days | Required\nBelgium | 45 days | Required\n\nAdditional rules after the table.')

    def test_ignores_decorative_table_content(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'decorative-table-document.docx'
            document = Document()
            document.add_heading('Working Conditions', level=1)
            document.add_table(rows=1, cols=2)
            document.add_heading('Salary', level=2)
            document.add_paragraph('Employees must receive their agreed salary.')
            document.save(file_path)
            sections = parse_docx_sections(file_path)
            self.assertEqual(len(sections), 1)
            self.assertEqual(sections[0].subsection, 'Salary')
            self.assertEqual(sections[0].content, 'Employees must receive their agreed salary.')

    def test_keeps_numbered_heading_as_content(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'numbered-heading-document.docx'
            document = Document()
            document.add_heading('Termination', level=1)
            document.add_heading('Remedies', level=2)
            document.add_paragraph('The employee may challenge the dismissal.')
            numbered_paragraph = document.add_paragraph('Pregnant employees are protected.')
            numbered_paragraph.style = 'Heading 2'
            _mark_as_numbered(numbered_paragraph)
            document.add_heading('Whistleblower Laws', level=2)
            document.add_paragraph('Whistleblowers receive legal protection.')
            document.save(file_path)
            sections = parse_docx_sections(file_path)
            self.assertEqual(len(sections), 2)
            self.assertEqual(sections[0].subsection, 'Remedies')
            self.assertIn('Pregnant employees are protected.', sections[0].content)
            self.assertEqual(sections[1].subsection, 'Whistleblower Laws')

    def test_keeps_unbolded_headings_as_content(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'unbolded-heading-document.docx'
            document = Document()
            document.add_heading('07. Termination of Employment Contracts', level=1)
            document.add_heading('Remedies for Wrongful Termination', level=2)
            document.add_paragraph('The employee may challenge the dismissal.')
            false_heading = document.add_paragraph('Objectively null and void termination: protected situations include:')
            false_heading.style = 'Heading 2'
            _mark_as_explicitly_unbolded(false_heading)
            numbered_item = document.add_paragraph('Pregnant employees are protected.')
            numbered_item.style = 'Heading 2'
            _mark_as_numbered(numbered_item)
            _mark_as_explicitly_unbolded(numbered_item)
            final_paragraph = document.add_paragraph('This protection means that the dismissal may be declared null and void.')
            final_paragraph.style = 'Heading 2'
            _mark_as_explicitly_unbolded(final_paragraph)
            document.add_heading('Whistleblower Laws', level=2)
            document.add_paragraph('Whistleblowers receive legal protection.')
            document.save(file_path)
            sections = parse_docx_sections(file_path)
            self.assertEqual(len(sections), 2)
            self.assertEqual(sections[0].subsection, 'Remedies for Wrongful Termination')
            self.assertIn('Objectively null and void termination:', sections[0].content)
            self.assertIn('Pregnant employees are protected.', sections[0].content)
            self.assertIn('This protection means that the dismissal', sections[0].content)
            self.assertEqual(sections[1].subsection, 'Whistleblower Laws')

    def test_accepts_heading_ending_with_colon(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'heading-with-colon-document.docx'
            document = Document()
            document.add_heading('Employment Contracts', level=1)
            document.add_heading('Notice periods:', level=2)
            document.add_paragraph('The statutory notice period is fifteen days.')
            document.save(file_path)
            sections = parse_docx_sections(file_path)
            self.assertEqual(len(sections), 1)
            self.assertEqual(sections[0].subsection, 'Notice periods:')
            self.assertEqual(sections[0].content, 'The statutory notice period is fifteen days.')

    def test_detects_legacy_numbered_bold_topic(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'legacy-document.docx'
            document = Document()
            document.add_paragraph('Introduction content.')
            topic = document.add_paragraph('Hiring practices in Greece')
            topic.style = 'List Paragraph'
            topic.runs[0].bold = True
            _mark_as_numbered(topic)
            document.add_paragraph('Foreign employees require permission.')
            document.save(file_path)
            sections = parse_docx_sections(file_path=file_path, country='Greece')
            self.assertEqual(len(sections), 2)
            self.assertEqual(sections[0].section, 'Employment Law Overview Greece')
            self.assertEqual(sections[0].content, 'Introduction content.')
            self.assertEqual(sections[1].section, 'Hiring practices in Greece')
            self.assertEqual(sections[1].content, 'Foreign employees require permission.')

    def test_detects_hybrid_bold_numbered_topic(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'hybrid-document.docx'
            document = Document()
            topic = document.add_paragraph('02. Employment Contracts')
            topic.runs[0].bold = True
            document.add_paragraph('Employment contracts may be open-ended.')
            document.save(file_path)
            sections = parse_docx_sections(file_path=file_path, country='Italy')
            self.assertEqual(len(sections), 1)
            self.assertEqual(sections[0].section, '02. Employment Contracts')
            self.assertEqual(sections[0].content, 'Employment contracts may be open-ended.')

    def test_keeps_unrecognized_heading_one_as_content(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'australia-heading-document.docx'
            document = Document()
            document.add_heading('07. Termination of Employment Contracts', level=1)
            document.add_paragraph('Termination content before the malformed heading.')
            false_heading = document.add_paragraph('Whistleblowers currently have legal protections.')
            false_heading.style = 'Heading 1'
            document.add_paragraph('Additional whistleblower content.')
            document.save(file_path)
            sections = parse_docx_sections(file_path=file_path, country='Australia')
            self.assertEqual(len(sections), 1)
            self.assertEqual(sections[0].section, '07. Termination of Employment Contracts')
            self.assertIn('Termination content before the malformed heading.', sections[0].content)
            self.assertIn('Whistleblowers currently have legal protections.', sections[0].content)
            self.assertIn('Additional whistleblower content.', sections[0].content)

    def test_keeps_heading_four_as_content(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'heading-four-document.docx'
            document = Document()
            document.add_heading('09. Transfer of Undertakings', level=1)
            body_paragraph = document.add_paragraph('The transfer does not automatically terminate employment.')
            body_paragraph.style = 'Heading 4'
            document.save(file_path)
            sections = parse_docx_sections(file_path=file_path, country='Brazil')
            self.assertEqual(len(sections), 1)
            self.assertEqual(sections[0].section, '09. Transfer of Undertakings')
            self.assertEqual(sections[0].content, 'The transfer does not automatically terminate employment.')

    def test_splits_working_conditions_subsections(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'working-conditions-document.docx'
            document = Document()
            document.add_heading('03. Working Conditions', level=1)
            for title, content in (('Overtime', 'Overtime legal content.'), ('Work Hours Record', 'Working time recording content.'), ('Paid Leave', 'Paid leave legal content.')):
                paragraph = document.add_paragraph(title)
                paragraph.runs[0].bold = True
                document.add_paragraph(content)
            document.save(file_path)
            sections = parse_docx_sections(file_path=file_path, country='Spain')
            self.assertEqual([section.subsection for section in sections], ['Overtime', 'Work Hours Record', 'Paid Leave'])
            self.assertEqual(sections[0].content, 'Overtime legal content.')

    def test_notice_of_termination_and_redundancy_pay_starts_new_section(self) -> None:
        """
        Regression test for a real L&E document defect.

        In the Australia source document, a bold "Notice of Termination
        and Redundancy Pay" paragraph appears inside the "Working
        Conditions" section (after "Overtime"), using the same plain
        bold-"Normal"-style formatting as ordinary emphasis elsewhere in
        the document. Before the fix, this heading was not recognized
        by any taxonomy rule, so its content silently stayed folded
        into the preceding Overtime subsection under the wrong legal
        topic - making it unreachable for questions about redundancy
        pay, which are filtered by legal topic.
        """
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'working-conditions-termination-document.docx'
            document = Document()
            document.add_heading('03. Working Conditions', level=1)
            overtime_heading = document.add_paragraph('Overtime')
            overtime_heading.runs[0].bold = True
            document.add_paragraph('Overtime is paid at a premium rate.')
            redundancy_heading = document.add_paragraph('Notice of Termination and Redundancy Pay')
            redundancy_heading.runs[0].bold = True
            document.add_paragraph("An employee is entitled to redundancy pay calculated using the employee's base rate of pay.")
            paid_leave_heading = document.add_paragraph('Paid Leave')
            paid_leave_heading.runs[0].bold = True
            document.add_paragraph('Employees accrue paid leave over each year of service.')
            document.save(file_path)
            sections = parse_docx_sections(file_path=file_path, country='Australia')
            overtime_section = next((section for section in sections if section.subsection == 'Overtime'))
            self.assertNotIn('redundancy pay', overtime_section.content.casefold())
            self.assertIn('premium rate', overtime_section.content)
            redundancy_section = next((section for section in sections if 'redundancy' in section.content.casefold()), None)
            self.assertIsNotNone(redundancy_section)
            self.assertEqual(redundancy_section.section, 'Notice of Termination and Redundancy Pay')
            self.assertIsNone(redundancy_section.subsection)
            self.assertIsNone(get_canonical_legal_topic(section=redundancy_section.section, country='Australia'))
            self.assertEqual(get_subsection_topic_override(redundancy_section.section), 'Termination of Employment Contracts')
            self.assertIn('base rate of pay', redundancy_section.content)
            self.assertLess(sections.index(overtime_section), sections.index(redundancy_section))
            paid_leave_section = next((section for section in sections if section.subsection == 'Paid Leave'))
            self.assertNotIn('redundancy', paid_leave_section.content.casefold())
            self.assertIn('accrue paid leave', paid_leave_section.content)
            self.assertEqual(get_canonical_legal_topic(section=paid_leave_section.section, country='Australia'), 'Working Conditions')
            self.assertLess(sections.index(redundancy_section), sections.index(paid_leave_section))

class CustomTopicRecognitionTests(unittest.TestCase):
    """
    ORDER 8A, section 12 - the parser must recognize a brand-new,
    non-taxonomy top-level legal topic an admin adds, using only the
    document's own real structure, generically (no topic name ever
    hardcoded) - while never promoting front matter, an ordinary
    subsection, or a stray list item that happens to share one weak
    structural signal with the document's own real topics.
    """

    def _build(self, directory: Path, blocks: list[tuple[str, str, int | None]]) -> Path:
        """
        blocks: (text, kind, heading_level) where kind is "heading" or
        "paragraph"; heading_level is the Word heading style level (or
        None for a plain paragraph).
        """
        file_path = directory / 'sample.docx'
        document = Document()
        for text, kind, heading_level in blocks:
            if kind == 'heading':
                document.add_heading(text, level=heading_level)
            else:
                document.add_paragraph(text)
        document.save(file_path)
        return file_path

    def test_custom_topic_recognized_after_a_confirmed_topic(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = self._build(Path(temp_dir), [('Employment Law Overview United Kingdom', 'heading', 1), ('Hiring Practices', 'heading', 1), ('Real hiring content.', 'paragraph', None), ('Remote Working', 'heading', 1), ('Employees may work remotely. MARKER.', 'paragraph', None)])
            sections = parse_docx_sections(file_path, country='United Kingdom')
            custom = [section for section in sections if section.is_custom_legal_topic]
            self.assertEqual(len(custom), 1)
            self.assertEqual(custom[0].section, 'Remote Working')
            self.assertIn('MARKER', custom[0].content)

    def test_custom_topic_not_recognized_before_any_confirmed_topic(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = self._build(Path(temp_dir), [('Introduction', 'heading', 1), ('Some introductory front matter text.', 'paragraph', None), ('Employment Law Overview United Kingdom', 'heading', 1), ('Hiring Practices', 'heading', 1), ('Real hiring content.', 'paragraph', None)])
            sections = parse_docx_sections(file_path, country='United Kingdom')
            self.assertFalse(any((section.is_custom_legal_topic for section in sections)))
            self.assertNotIn('Introduction', [section.section for section in sections])

    def test_custom_topic_unsupported_when_real_topics_use_bold_only(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / 'sample.docx'
            document = Document()
            document.add_heading('Employment Law Overview United Kingdom', level=1)
            bold_heading = document.add_paragraph()
            bold_heading.add_run('Hiring Practices').bold = True
            document.add_paragraph('Real hiring content.')
            bold_custom = document.add_paragraph()
            bold_custom.add_run('Remote Working').bold = True
            document.add_paragraph('Remote work content.')
            document.save(file_path)
            sections = parse_docx_sections(file_path, country='United Kingdom')
            self.assertFalse(any((section.is_custom_legal_topic for section in sections)))

    def test_custom_topic_requires_the_same_signal_as_confirmed_topics(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / 'sample.docx'
            document = Document()
            document.add_heading('Employment Law Overview United Kingdom', level=1)
            hiring_heading = document.add_heading('Hiring Practices', level=1)
            _mark_as_numbered(hiring_heading)
            document.add_paragraph('Real hiring content.')
            document.add_heading('Key Points', level=1)
            document.add_paragraph('This is an ordinary sub-heading, not numbered like the real topics.')
            document.save(file_path)
            sections = parse_docx_sections(file_path, country='United Kingdom')
            self.assertFalse(any((section.is_custom_legal_topic for section in sections)))
            self.assertNotIn('Key Points', [section.section for section in sections])

    def test_custom_topic_rejects_sentence_shaped_candidates(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / 'sample.docx'
            document = Document()
            document.add_heading('Employment Law Overview United Kingdom', level=1)
            hiring_heading = document.add_heading('Hiring Practices', level=1)
            _mark_as_numbered(hiring_heading)
            document.add_paragraph('Real hiring content.')
            list_item_heading = document.add_heading('the Corporations Act;', level=1)
            _mark_as_numbered(list_item_heading)
            document.add_paragraph('More content.')
            document.save(file_path)
            sections = parse_docx_sections(file_path, country='United Kingdom')
            self.assertFalse(any((section.is_custom_legal_topic for section in sections)))

    def test_generic_mode_never_recognizes_custom_legal_topics(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = self._build(Path(temp_dir), [('Some Section', 'heading', 1), ('Some content.', 'paragraph', None), ('Another Section', 'heading', 1), ('More content.', 'paragraph', None)])
            sections = parse_docx_sections(file_path)
            self.assertFalse(any((section.is_custom_legal_topic for section in sections)))

class AdminSectionMarkerRecognitionTests(unittest.TestCase):
    """
    ORDER 8A-C - the dedicated ADMIN-section marker style
    (ADMIN_SECTION_STYLE_NAME) is the sole, deterministic identity a
    reparse relies on for an admin-added top-level topic: it must work
    regardless of the surrounding document's own native convention
    (Heading 1, bold-only, or anything else), regardless of country,
    and regardless of document position - unlike the older, structural-
    signal-based custom-topic heuristic this supersedes for anything
    the ADMIN Add feature itself creates.
    """

    def _add_marker_heading(self, document: Document, title: str) -> None:
        from docx.enum.style import WD_STYLE_TYPE
        from app.services.docx_parser import ADMIN_SECTION_STYLE_NAME
        try:
            style = document.styles[ADMIN_SECTION_STYLE_NAME]
        except KeyError:
            style = document.styles.add_style(ADMIN_SECTION_STYLE_NAME, WD_STYLE_TYPE.PARAGRAPH)
            style.font.bold = True
        document.add_paragraph(title, style=style.name)

    def test_marker_recognized_in_bold_only_document(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / 'sample.docx'
            document = Document()
            document.add_paragraph('Employment Law Overview United Kingdom')
            hiring = document.add_paragraph()
            hiring.add_run('Hiring Practices').bold = True
            document.add_paragraph('Real hiring content.')
            self._add_marker_heading(document, 'Remote Working')
            document.add_paragraph('Remote work content. MARKER.')
            document.save(file_path)
            sections = parse_docx_sections(file_path, country='United Kingdom')
            custom = [s for s in sections if s.is_custom_legal_topic]
            self.assertEqual(len(custom), 1)
            self.assertEqual(custom[0].section, 'Remote Working')
            self.assertIn('MARKER', custom[0].content)
            hiring_sections = [s for s in sections if s.section == 'Hiring Practices']
            self.assertEqual(len(hiring_sections), 1)
            self.assertEqual(hiring_sections[0].content, 'Real hiring content.')

    def test_marker_recognized_with_no_native_topics_at_all(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / 'sample.docx'
            document = Document()
            document.add_paragraph('Just some plain front matter.')
            self._add_marker_heading(document, 'Custom Only Section')
            document.add_paragraph('Custom-only content.')
            document.save(file_path)
            sections = parse_docx_sections(file_path, country='United Kingdom')
            custom = [s for s in sections if s.is_custom_legal_topic]
            self.assertEqual(len(custom), 1)
            self.assertEqual(custom[0].section, 'Custom Only Section')

    def test_marker_style_alone_never_appears_on_native_content(self) -> None:
        with TemporaryDirectory() as temp_dir:
            file_path = Path(temp_dir) / 'sample.docx'
            document = Document()
            document.add_heading('Employment Law Overview United Kingdom', level=1)
            document.add_heading('Hiring Practices', level=1)
            document.add_paragraph('Real hiring content.')
            document.save(file_path)
            sections = parse_docx_sections(file_path, country='United Kingdom')
            self.assertFalse(any((s.is_custom_legal_topic for s in sections)))

class ContactBlockParsingTests(unittest.TestCase):
    """
    Tests for parse_contact_blocks() against a synthetic paragraph
    structure - the shape _extract_text_box_blocks() would have
    produced from a real DOCX, without needing one.
    """

    def test_firm_then_contact_person_are_paired(self) -> None:
        blocks = [['Example & Partners Advogados', 'Freedonia', '1 Example Street, 6th floor, 00000 Sample City', '+00 000 000 00', 'www.example-partners.test'], ['CONTACT PERSON', 'Alex Example', 'alex@example-partners.test']]
        contacts = parse_contact_blocks(blocks, country='Freedonia')
        self.assertEqual(len(contacts), 1)
        contact = contacts[0]
        self.assertEqual(contact.member_firm, 'Example & Partners Advogados')
        self.assertEqual(contact.contact_person, 'Alex Example')
        self.assertEqual(contact.email, 'alex@example-partners.test')
        self.assertEqual(contact.phone, '+00 000 000 00')
        self.assertEqual(contact.website, 'www.example-partners.test')
        self.assertNotIn('Freedonia', contact.address or '')

    def test_contact_person_before_firm_block_are_still_paired(self) -> None:
        blocks = [['CONTACT PERSON', 'Nicolás Grandi', 'ngrandi@allende.com'], ['Allende & Brea', 'Argentina', 'Torre IRSA, Maipú 1300', '+54 114 318 9984', 'www.allendebrea.com']]
        contacts = parse_contact_blocks(blocks, country='Argentina')
        self.assertEqual(len(contacts), 1)
        contact = contacts[0]
        self.assertEqual(contact.member_firm, 'Allende & Brea')
        self.assertEqual(contact.contact_person, 'Nicolás Grandi')
        self.assertEqual(contact.email, 'ngrandi@allende.com')

    def test_bare_website_alone_is_not_a_firm_block(self) -> None:
        blocks = [['www.leglobal.law'], ['Some Firm', 'Country', '123 Main Street', '+1 555 000 0000', 'www.somefirm.example'], ['CONTACT PERSON', 'Jane Doe', 'jane@somefirm.example']]
        contacts = parse_contact_blocks(blocks)
        self.assertEqual(len(contacts), 1)
        self.assertEqual(contacts[0].member_firm, 'Some Firm')

    def test_plural_contact_persons_marker_with_multiple_emails(self) -> None:
        blocks = [['Van Olmen & Wynant', 'Belgium', 'Avenue Louise 221, 1050 Brussels', '+32 264 405 11', 'www.vow.be'], ['CONTACT PERSONS', 'Chris van Olmen and Nicolas Simon', 'chris.van.olmen@vow.be', 'nicolas.simon@vow.be']]
        contacts = parse_contact_blocks(blocks, country='Belgium')
        self.assertEqual(len(contacts), 1)
        contact = contacts[0]
        self.assertEqual(contact.contact_person, 'Chris van Olmen and Nicolas Simon')
        self.assertIn('chris.van.olmen@vow.be', contact.email or '')
        self.assertIn('nicolas.simon@vow.be', contact.email or '')

    def test_postal_code_before_real_phone_does_not_win(self) -> None:
        blocks = [['Atsumi & Sakai', 'Japan', 'Fukoku Seimei Bldg., Reception: 16 F, 2-2-2 Uchisaiwaicho, Chiyoda-ku, 100-0011 Tokyo, +81 355 012 111', 'www.aplaw.jp/en/'], ['CONTACT PERSON', 'Tatsuo Yamashima', 'tatsuo.yamashima@aplaw.jp']]
        contacts = parse_contact_blocks(blocks, country='Japan')
        self.assertEqual(len(contacts), 1)
        contact = contacts[0]
        self.assertEqual(contact.phone, '+81 355 012 111')
        self.assertIn('100-0011 Tokyo', contact.address or '')
        self.assertNotIn('+81 355 012 111', contact.address or '')
        self.assertEqual(contact.member_firm, 'Atsumi & Sakai')
        self.assertEqual(contact.contact_person, 'Tatsuo Yamashima')

    def test_postal_code_line_before_local_phone_prefers_phone(self) -> None:
        blocks = [['Example Firm', '100-0011 Tokyo', '03 5501 2111']]
        contacts = parse_contact_blocks(blocks)
        self.assertEqual(len(contacts), 1)
        self.assertEqual(contacts[0].phone, '03 5501 2111')
        self.assertIn('100-0011 Tokyo', contacts[0].address or '')
        self.assertNotIn('03 5501 2111', contacts[0].address or '')

    def test_multiple_documents_worth_of_contacts_are_all_kept(self) -> None:
        blocks = [['Firm One', '123 Street', '+1 555 111 1111'], ['CONTACT PERSON', 'Person One', 'one@example.com'], ['Firm Two', '456 Avenue', '+1 555 222 2222'], ['CONTACT PERSON', 'Person Two', 'two@example.com']]
        contacts = parse_contact_blocks(blocks)
        self.assertEqual(len(contacts), 2)
        self.assertEqual([contact.email for contact in contacts], ['one@example.com', 'two@example.com'])

    def test_no_text_box_blocks_returns_no_contacts(self) -> None:
        self.assertEqual(parse_contact_blocks([]), [])

    def test_unmatched_firm_block_reports_only_its_own_fields(self) -> None:
        blocks = [['Only Firm', '42 Road', '+1 555 333 3333']]
        contacts = parse_contact_blocks(blocks)
        self.assertEqual(len(contacts), 1)
        self.assertEqual(contacts[0].member_firm, 'Only Firm')
        self.assertIsNone(contacts[0].contact_person)
        self.assertIsNone(contacts[0].email)

    def test_build_contact_chunk_content_omits_missing_fields(self) -> None:
        blocks = [['CONTACT PERSON', 'Jane Doe', 'jane@example.com']]
        contacts = parse_contact_blocks(blocks)
        content = build_contact_chunk_content(contacts)
        self.assertIn('Contact person: Jane Doe', content)
        self.assertIn('Email: jane@example.com', content)
        self.assertNotIn('Member firm', content)
        self.assertNotIn('Phone', content)
        self.assertNotIn('Address', content)
        self.assertNotIn('Website', content)

class PlainParagraphContactFallbackTests(unittest.TestCase):
    """Contract for contacts stored in ordinary DOCX paragraphs."""

    def test_france_plain_paragraph_contact_is_extracted(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'france-body-contact.docx'
            document = Document()
            document.add_paragraph('FRANCE')
            document.add_paragraph('EMPLOYMENT LAW OVERVIEWS 2025 - 2026')
            document.add_paragraph('FLICHY GRANGÉ AVOCATS')
            document.add_heading('I. GENERAL OVERVIEW', level=1)
            document.add_paragraph('Representative employment-law content.')
            document.add_paragraph('Caroline Scherrmann and Florence Bacquet')
            document.add_paragraph('Partners, Flichy Grangé Avocats')
            document.add_paragraph('scherrmann@flichy.com')
            document.add_paragraph('bacquet@flichy.com')
            document.add_paragraph('+33 1 56 62 30 00')
            document.add_paragraph('YOUR L&E GLOBAL POC')
            document.add_paragraph('For all inquiries related to this project, please contact Jessica Stout, International Business Development Executive at L&E Global, at jessica.stout@leglobal.law.')
            document.save(file_path)
            contacts = extract_contacts_from_docx(file_path, country='France')
            self.assertEqual(len(contacts), 1)
            contact = contacts[0]
            self.assertEqual(contact.member_firm, 'Flichy Grangé Avocats')
            self.assertEqual(contact.contact_person, 'Caroline Scherrmann and Florence Bacquet')
            self.assertEqual(contact.email, 'scherrmann@flichy.com, bacquet@flichy.com')
            self.assertEqual(contact.phone, '+33 1 56 62 30 00')
            self.assertIsNone(contact.address)
            self.assertIsNone(contact.website)
            rendered = ' '.join((value for value in (contact.member_firm, contact.contact_person, contact.email, contact.phone) if value))
            self.assertNotIn('Jessica Stout', rendered)
            self.assertNotIn('jessica.stout@leglobal.law', rendered)

    def test_generic_plain_paragraph_contact_is_extracted(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'generic-body-contact.docx'
            document = Document()
            document.add_paragraph('FREEDONIA')
            document.add_paragraph('Example Employment Law')
            document.add_heading('Employment Contracts', level=1)
            document.add_paragraph('Representative legal information.')
            document.add_paragraph('Alex Example and Sam Sample')
            document.add_paragraph('Partners, Example Employment Law')
            document.add_paragraph('alex@example-law.test')
            document.add_paragraph('sam@example-law.test')
            document.add_paragraph('+99 123 456 789')
            document.save(file_path)
            contacts = extract_contacts_from_docx(file_path, country='Freedonia')
            self.assertEqual(len(contacts), 1)
            contact = contacts[0]
            self.assertEqual(contact.member_firm, 'Example Employment Law')
            self.assertEqual(contact.contact_person, 'Alex Example and Sam Sample')
            self.assertEqual(contact.email, 'alex@example-law.test, sam@example-law.test')
            self.assertEqual(contact.phone, '+99 123 456 789')

    def test_project_poc_alone_is_not_member_firm_contact(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'project-poc-only.docx'
            document = Document()
            document.add_paragraph('YOUR L&E GLOBAL POC')
            document.add_paragraph('For all inquiries related to this project, please contact Jessica Stout, International Business Development Executive at L&E Global, at jessica.stout@leglobal.law.')
            document.save(file_path)
            self.assertEqual(extract_contacts_from_docx(file_path, country='France'), [])

    def test_firm_name_alone_is_not_a_contact(self) -> None:
        cases = (('Portugal', 'SRS LEGAL'), ('Taiwan', 'Lee and Li, Attorneys-at-Law'))
        for country, firm_name in cases:
            with self.subTest(country=country):
                with TemporaryDirectory() as temporary_directory:
                    file_path = Path(temporary_directory) / f'{country.lower()}-firm-only.docx'
                    document = Document()
                    document.add_paragraph(country.upper())
                    document.add_paragraph(firm_name)
                    document.add_heading('GENERAL OVERVIEW', level=1)
                    document.add_paragraph('Representative legal information.')
                    document.save(file_path)
                    self.assertEqual(extract_contacts_from_docx(file_path, country=country), [])

    def test_legal_reference_coordinates_do_not_become_contact(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'legal-references-only.docx'
            document = Document()
            document.add_paragraph('PHILIPPINES')
            document.add_heading('Employment Benefits', level=1)
            document.add_paragraph('More information may be obtained from the public labour authority at www.labour-authority.example.')
            document.add_paragraph('The authority may also be reached at +63 2 8123 4567 for public information.')
            document.save(file_path)
            self.assertEqual(extract_contacts_from_docx(file_path, country='Philippines'), [])

    def test_existing_text_box_contact_keeps_priority(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'legacy-textbox-priority.docx'
            document = Document()
            document.add_paragraph('Alternative Person')
            document.add_paragraph('Partner, Alternative Firm')
            document.add_paragraph('alternative@alternative.example')
            document.add_paragraph('+44 20 0000 0000')
            document.save(file_path)
            legacy_blocks = [['Stable Firm', 'Testland', '1 Existing Street', '+1 555 111 2222', 'www.stable.example'], ['CONTACT PERSON', 'Stable Person', 'stable@stable.example']]
            with patch('app.services.docx_parser.extract_text_box_blocks', return_value=legacy_blocks):
                contacts = extract_contacts_from_docx(file_path, country='Testland')
            self.assertEqual(len(contacts), 1)
            contact = contacts[0]
            self.assertEqual(contact.member_firm, 'Stable Firm')
            self.assertEqual(contact.contact_person, 'Stable Person')
            self.assertEqual(contact.email, 'stable@stable.example')
            self.assertEqual(contact.phone, '+1 555 111 2222')

def _build_france_shaped_document(file_path: Path) -> None:
    """The exact real France source layout (mission "FINAL CONTACT CRUD
    CLOSURE"): legal content, then a legacy member-firm contact block
    naming two people sharing one firm/phone, immediately before the
    L&E Global POC block - never a synthetic table-only stand-in."""
    document = Document()
    document.add_paragraph('FRANCE')
    document.add_paragraph('EMPLOYMENT LAW OVERVIEWS 2025 - 2026')
    document.add_paragraph('FLICHY GRANGÉ AVOCATS')
    document.add_heading('I. GENERAL OVERVIEW', level=1)
    document.add_paragraph('Representative employment-law content.')
    document.add_paragraph('Caroline Scherrmann and Florence Bacquet')
    document.add_paragraph('Partners, Flichy Grangé Avocats')
    document.add_paragraph('scherrmann@flichy.com')
    document.add_paragraph('bacquet@flichy.com')
    document.add_paragraph('+33 1 56 62 30 00')
    document.add_paragraph('YOUR L&E GLOBAL POC')
    document.add_paragraph('For all inquiries related to this project, please contact Jessica Stout, International Business Development Executive at L&E Global, at jessica.stout@leglobal.law.')
    document.add_paragraph('Disclaimer text follows here.')
    document.save(file_path)

class FindPlainParagraphContactBlockBoundsTests(unittest.TestCase):
    """Corpus-independent coverage of find_plain_paragraph_contact_
    block_bounds - the structural anchor-finder contact_document_
    area.py's canonicalizer uses so a France-style legacy contact
    area (ordinary body paragraphs, no floating shape) gets its
    canonical table replacement inserted at the SAME logical location,
    never silently falling back to the document's start."""

    def test_finds_the_exact_block_bounds_before_poc(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'france.docx'
            _build_france_shaped_document(file_path)
            document = Document(file_path)
            bounds = find_plain_paragraph_contact_block_bounds(document)
            self.assertIsNotNone(bounds)
            first_index, last_index = bounds
            self.assertEqual('Caroline Scherrmann and Florence Bacquet', document.paragraphs[first_index].text)
            self.assertEqual('+33 1 56 62 30 00', document.paragraphs[last_index].text)
            self.assertEqual('YOUR L&E GLOBAL POC', document.paragraphs[last_index + 1].text)

    def test_returns_none_when_no_plain_paragraph_block_exists(self) -> None:
        with TemporaryDirectory() as temporary_directory:
            file_path = Path(temporary_directory) / 'no-contact.docx'
            document = Document()
            document.add_paragraph('Just an ordinary legal document.')
            document.add_paragraph('With no contact information at all.')
            document.save(file_path)
            reopened = Document(file_path)
            self.assertIsNone(find_plain_paragraph_contact_block_bounds(reopened))

class SplitCombinedLegacyContactTests(unittest.TestCase):
    """Corpus-independent coverage of split_combined_legacy_contact -
    the narrow, deterministic normalization that turns a legacy
    contact naming multiple people into one ExtractedContact per
    person, only when the split is unambiguous."""

    def test_splits_two_people_sharing_one_firm_and_phone(self) -> None:
        combined = ExtractedContact(member_firm='Flichy Grangé Avocats', contact_person='Caroline Scherrmann and Florence Bacquet', email='scherrmann@flichy.com, bacquet@flichy.com', phone='+33 1 56 62 30 00')
        split = split_combined_legacy_contact(combined)
        self.assertIsNotNone(split)
        self.assertEqual(2, len(split))
        self.assertEqual('Caroline Scherrmann', split[0].contact_person)
        self.assertEqual('scherrmann@flichy.com', split[0].email)
        self.assertEqual('Florence Bacquet', split[1].contact_person)
        self.assertEqual('bacquet@flichy.com', split[1].email)
        for contact in split:
            self.assertEqual('Flichy Grangé Avocats', contact.member_firm)
            self.assertEqual('+33 1 56 62 30 00', contact.phone)

    def test_returns_none_for_a_single_person_contact(self) -> None:
        single = ExtractedContact(member_firm='Some Firm', contact_person='Jane Doe', email='jane@example.com')
        self.assertIsNone(split_combined_legacy_contact(single))

    def test_returns_none_when_person_and_email_counts_mismatch(self) -> None:
        mismatched = ExtractedContact(member_firm='Some Firm', contact_person='Jane Doe and John Roe', email='jane@example.com')
        self.assertIsNone(split_combined_legacy_contact(mismatched))

    def test_never_splits_a_firm_name_containing_and(self) -> None:
        """A firm name that happens to contain "and" is never at
        risk: only contact_person is ever split, member_firm is
        always copied through unchanged."""
        firm_with_and = ExtractedContact(member_firm='Smith and Jones LLP', contact_person='Jane Doe', email='jane@example.com')
        self.assertIsNone(split_combined_legacy_contact(firm_with_and))

    def test_splits_three_people(self) -> None:
        combined = ExtractedContact(member_firm='Big Firm LLP', contact_person='Alice Smith, Bob Jones and Carol White', email='alice@example.com, bob@example.com, carol@example.com')
        split = split_combined_legacy_contact(combined)
        self.assertIsNotNone(split)
        self.assertEqual(3, len(split))
        self.assertEqual(['Alice Smith', 'Bob Jones', 'Carol White'], [c.contact_person for c in split])
        self.assertEqual(['alice@example.com', 'bob@example.com', 'carol@example.com'], [c.email for c in split])

    def test_returns_none_without_email(self) -> None:
        no_email = ExtractedContact(member_firm='Some Firm', contact_person='Jane Doe and John Roe', email=None)
        self.assertIsNone(split_combined_legacy_contact(no_email))

class ClassifyCanonicalFirmLinesTests(unittest.TestCase):
    """
    Direct, corpus-independent coverage of
    _classify_canonical_firm_lines - the canonical table reader's
    field classifier, fixed to be semantic (content-based) rather
    than position-only. Corpus regressions for the real documents
    this was found against (IE/IN/US) live in
    test_contact_document_area.py; these tests pin the underlying
    classification rules precisely, for every input shape called out
    by the mission: clean/international/spaced/+-prefixed phones,
    a phone occupying the first line when member_firm is empty, a
    phone with a trailing annotation, and a website mentioned inside
    address prose vs. one on its own dedicated line.
    """

    def test_normal_case_all_four_fields_present(self) -> None:
        member_firm, phone, website, remaining = _classify_canonical_firm_lines(['HARMERS WORKPLACE LAWYERS', '31 Market Street, Level 27 St Martins Tower, NSW 2000 Sydney', '+61 292 674 322', 'WWW.HARMERS.COM.AU'])
        self.assertEqual('HARMERS WORKPLACE LAWYERS', member_firm)
        self.assertEqual('+61 292 674 322', phone)
        self.assertEqual('WWW.HARMERS.COM.AU', website)
        self.assertEqual(['31 Market Street, Level 27 St Martins Tower, NSW 2000 Sydney'], [['HARMERS WORKPLACE LAWYERS', '31 Market Street, Level 27 St Martins Tower, NSW 2000 Sydney', '+61 292 674 322', 'WWW.HARMERS.COM.AU'][i] for i in remaining])

    def test_clean_phone(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(['+1 5551234567'])
        self.assertEqual('+1 5551234567', phone)

    def test_international_phone(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(['+81 355 012 111'])
        self.assertEqual('+81 355 012 111', phone)

    def test_phone_with_spaces(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(['+46 852 206 500'])
        self.assertEqual('+46 852 206 500', phone)

    def test_phone_with_plus_prefix(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(['+353 1 234 5678'])
        self.assertEqual('+353 1 234 5678', phone)

    def test_phone_without_plus_prefix(self) -> None:
        _, phone, _, _ = _classify_canonical_firm_lines(['1 212 545 4050'])
        self.assertEqual('1 212 545 4050', phone)

    def test_phone_as_first_firm_side_line_no_member_firm(self) -> None:
        """Quirk B: a contact whose member_firm/address/website are
        all empty has its phone land on line 0 - it must still be
        recognized as phone, never misread as member_firm."""
        member_firm, phone, website, remaining = _classify_canonical_firm_lines(['+1 555 000 0000'])
        self.assertIsNone(member_firm)
        self.assertEqual('+1 555 000 0000', phone)
        self.assertIsNone(website)
        self.assertEqual([], remaining)

    def test_phone_with_trailing_annotation_preserved_whole(self) -> None:
        """Quirk A: a phone value with a trailing annotation must be
        preserved in its ENTIRETY as the phone field - never split,
        with the remainder leaking into address."""
        member_firm, phone, website, remaining = _classify_canonical_firm_lines(['Cederquist', 'Hovslagargatan 3, SE-111 96 Stockholm', '+46 852 206 500 (updated)', 'www.cederquist.se'])
        self.assertEqual('Cederquist', member_firm)
        self.assertEqual('+46 852 206 500 (updated)', phone)
        self.assertEqual('www.cederquist.se', website)
        self.assertEqual(['Hovslagargatan 3, SE-111 96 Stockholm'], [['Cederquist', 'Hovslagargatan 3, SE-111 96 Stockholm', '+46 852 206 500 (updated)', 'www.cederquist.se'][i] for i in remaining])

    def test_phone_as_first_line_with_annotation_combined(self) -> None:
        """Quirk A and Quirk B together: no member_firm AND a phone
        annotation on the same, first line."""
        member_firm, phone, website, remaining = _classify_canonical_firm_lines(['+353 1 234 5678 (mobile)'])
        self.assertIsNone(member_firm)
        self.assertEqual('+353 1 234 5678 (mobile)', phone)
        self.assertEqual([], remaining)

    def test_website_mentioned_inside_address_prose_is_not_extracted(self) -> None:
        """Quirk C: an address sentence that merely MENTIONS a URL
        (with its own sentence-ending punctuation attached) must never
        be misread as the website field - only a line that is, in its
        entirety, just a URL is the dedicated website field."""
        member_firm, phone, website, remaining = _classify_canonical_firm_lines(['Jackson Lewis PC', 'USA, 666 Third Avenue, 29th Floor, 10017 New York, Jackson Lewis has over 60 offices throughout the USA. , For information, please see www.jacksonlewis.com.', '1 212 545 4050', 'www.jacksonlewis.com'])
        self.assertEqual('Jackson Lewis PC', member_firm)
        self.assertEqual('1 212 545 4050', phone)
        self.assertEqual('www.jacksonlewis.com', website)
        remaining_lines = [['Jackson Lewis PC', 'USA, 666 Third Avenue, 29th Floor, 10017 New York, Jackson Lewis has over 60 offices throughout the USA. , For information, please see www.jacksonlewis.com.', '1 212 545 4050', 'www.jacksonlewis.com'][i] for i in remaining]
        self.assertEqual(1, len(remaining_lines))
        self.assertIn('please see www.jacksonlewis.com.', remaining_lines[0])

    def test_no_phone_or_website_present(self) -> None:
        member_firm, phone, website, remaining = _classify_canonical_firm_lines(['Some Firm Ltd', '123 Some Street'])
        self.assertEqual('Some Firm Ltd', member_firm)
        self.assertIsNone(phone)
        self.assertIsNone(website)
        self.assertEqual([1], remaining)

    def test_empty_firm_lines(self) -> None:
        member_firm, phone, website, remaining = _classify_canonical_firm_lines([])
        self.assertIsNone(member_firm)
        self.assertIsNone(phone)
        self.assertIsNone(website)
        self.assertEqual([], remaining)



# ================================================================
# SOURCE: backend/tests/test_legacy_subsections.py
# ================================================================

import tempfile
import unittest
from pathlib import Path
from docx import Document
from app.services.docx_parser import parse_docx_sections

class LegacySubsectionParserTests(unittest.TestCase):

    def _save_document(self, document: Document, directory: str, filename: str='test.docx') -> Path:
        file_path = Path(directory) / filename
        document.save(file_path)
        return file_path

    def test_detects_known_bold_legacy_subsection(self) -> None:
        document = Document()
        document.add_heading('01. Hiring Practices', level=1)
        subsection = document.add_paragraph('Requirement for Foreign Employees to Work')
        subsection.runs[0].bold = True
        document.add_paragraph('Foreign employees require a work permit.')
        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(document=document, directory=directory)
            sections = parse_docx_sections(file_path=file_path, country='Spain')
        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].subsection, 'Requirement for Foreign Employees to Work')
        self.assertEqual(sections[0].content, 'Foreign employees require a work permit.')

    def test_detects_known_heading_four_subsection(self) -> None:
        document = Document()
        document.add_heading('07. Termination of Employment Contracts', level=1)
        document.add_heading('Grounds for Termination', level=4)
        document.add_paragraph('Termination requires a valid legal ground.')
        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(document=document, directory=directory)
            sections = parse_docx_sections(file_path=file_path, country='Italy')
        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].subsection, 'Grounds for Termination')

    def test_keeps_unknown_bold_paragraph_as_content(self) -> None:
        document = Document()
        document.add_heading('07. Termination of Employment Contracts', level=1)
        unknown_title = document.add_paragraph('Immediate Termination')
        unknown_title.runs[0].bold = True
        document.add_paragraph('Immediate termination is possible in limited cases.')
        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(document=document, directory=directory)
            sections = parse_docx_sections(file_path=file_path, country='Czech Republic')
        self.assertEqual(len(sections), 1)
        self.assertIsNone(sections[0].subsection)
        self.assertEqual(sections[0].content, 'Immediate Termination\n\nImmediate termination is possible in limited cases.')

    def test_normalizes_known_typo_alias(self) -> None:
        document = Document()
        document.add_heading('07. Termination of Employment Contracts', level=1)
        subsection = document.add_paragraph('Whitsleblower Laws')
        subsection.runs[0].bold = True
        document.add_paragraph('Whistleblowers receive statutory protection.')
        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(document=document, directory=directory)
            sections = parse_docx_sections(file_path=file_path, country='Japan')
        self.assertEqual(sections[0].subsection, 'Whistleblower Laws')

    def test_scopes_same_subsection_to_each_topic(self) -> None:
        document = Document()
        document.add_heading('04. Anti-Discrimination Laws', level=1)
        first_remedies = document.add_paragraph('Remedies')
        first_remedies.runs[0].bold = True
        document.add_paragraph('Anti-discrimination remedies.')
        document.add_heading('05. Pay Equity Laws', level=1)
        second_remedies = document.add_paragraph('Remedies')
        second_remedies.runs[0].bold = True
        document.add_paragraph('Pay equity remedies.')
        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(document=document, directory=directory)
            sections = parse_docx_sections(file_path=file_path, country='Spain')
        self.assertEqual(len(sections), 2)
        self.assertEqual(sections[0].subsection, 'Remedies')
        self.assertEqual(sections[1].subsection, 'Remedies')
        self.assertIn('Anti-Discrimination', sections[0].section)
        self.assertIn('Pay Equity', sections[1].section)

    def test_generic_parser_keeps_custom_heading_two(self) -> None:
        document = Document()
        document.add_heading('Custom Main Section', level=1)
        document.add_heading('Custom Subsection', level=2)
        document.add_paragraph('Custom content.')
        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(document=document, directory=directory)
            sections = parse_docx_sections(file_path=file_path)
        self.assertEqual(len(sections), 1)
        self.assertEqual(sections[0].section, 'Custom Main Section')
        self.assertEqual(sections[0].subsection, 'Custom Subsection')
        self.assertEqual(sections[0].content, 'Custom content.')



# ================================================================
# SOURCE: backend/tests/test_section_splitter.py
# ================================================================

import unittest
from app.services.docx_parser import ParsedSection
from app.services.document_chunk_builder import SectionSplitterError, split_parsed_sections, split_text

class SectionSplitterTests(unittest.TestCase):

    def test_keeps_short_text_unchanged(self) -> None:
        text = 'A short legal paragraph.'
        self.assertEqual(split_text(text=text, max_chars=100), [text])

    def test_prefers_paragraph_boundaries(self) -> None:
        first_paragraph = 'A' * 60
        second_paragraph = 'B' * 60
        chunks = split_text(text=f'{first_paragraph}\n\n{second_paragraph}', max_chars=100)
        self.assertEqual(chunks, [first_paragraph, second_paragraph])

    def test_splits_oversized_paragraph_by_sentences(self) -> None:
        text = 'First sentence contains legal information. Second sentence contains additional information. Third sentence completes the legal explanation.'
        chunks = split_text(text=text, max_chars=100)
        self.assertGreater(len(chunks), 1)
        self.assertTrue(all((len(chunk) <= 100 for chunk in chunks)))
        self.assertIn('First sentence', chunks[0])

    def test_uses_whitespace_fallback(self) -> None:
        text = ' '.join((f'word{index}' for index in range(100)))
        chunks = split_text(text=text, max_chars=100)
        self.assertGreater(len(chunks), 1)
        self.assertTrue(all((len(chunk) <= 100 for chunk in chunks)))

    def test_preserves_section_metadata(self) -> None:
        parsed_section = ParsedSection(section='07. Termination of Employment Contracts', subsection='Grounds for Termination', content=' '.join(('termination' for _ in range(100))))
        split_sections = split_parsed_sections(parsed_sections=[parsed_section], max_chars=100)
        self.assertGreater(len(split_sections), 1)
        for split_section in split_sections:
            self.assertEqual(split_section.section, parsed_section.section)
            self.assertEqual(split_section.subsection, parsed_section.subsection)
            self.assertLessEqual(len(split_section.content), 100)

    def test_is_deterministic(self) -> None:
        text = '\n\n'.join(('A' * 75, 'B' * 75, 'C' * 75))
        first_result = split_text(text=text, max_chars=100)
        second_result = split_text(text=text, max_chars=100)
        self.assertEqual(first_result, second_result)

    def test_rejects_invalid_limit(self) -> None:
        with self.assertRaises(SectionSplitterError):
            split_text(text='Legal content.', max_chars=50)
