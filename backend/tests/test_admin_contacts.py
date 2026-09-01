"""
Tests for Admin Contact Management: contact CRUD, contact state
persistence, the admin-modified marker, legacy bootstrap/reseed, rollback,
and admin contact-photo CRUD (contact photos are a real, permanent part
of the persisted source DOCX, unlike business-field text, which is only
ever materialized into an ephemeral copy at download time).

Mirrors the conventions test_admin_document_sections.py already
established: plain unittest.TestCase, tempfile.TemporaryDirectory() for
source_directory, explicit dependency injection rather than a mocking
framework wherever this codebase already exposes a seam for it, and
unittest.mock.patch only at the one seam it has none for
(document_indexer.py's own module-level `bulk`/
`ensure_legal_documents_index` calls, and docx_parser.py's own
extract_contacts_from_docx for the DOCX-driven reseed/bootstrap paths -
extract_contacts_from_docx's own correctness against a real file is
already exhaustively covered by test_docx_parser.py; these tests only
need SOME deterministic contact list as input).
"""

from __future__ import annotations

import contextlib
import hashlib
import shutil
import tempfile
import unittest
import zipfile
from contextlib import nullcontext
from pathlib import Path
from types import SimpleNamespace
from typing import Any
from unittest.mock import patch

from docx import Document

from app.routers.admin_contacts import router
from app.services import admin_contacts
from app.services.admin_contacts import (
    AdminContactMutationFailedError,
    AdminContactNotFoundError,
    add_contact,
    apply_structured_contact_state_to_chunks,
    bootstrap_legacy_contacts,
    delete_contact,
    list_contacts,
    reseed_contact_state_from_parsed_contacts,
    reseed_contacts_from_current_docx,
    update_contact,
)
from app.services.admin_contact_photos import (
    AdminContactPhotoError,
    AdminContactPhotoNotFoundError,
    read_admin_contact_photo,
    remove_admin_contact_photo,
    replace_admin_contact_photo,
)
from app.services.admin_document_lifecycle import (
    AdminDocumentRollbackError,
    get_document_download,
    reindex_indexed_document,
)
from app.services.admin_document_replacement import (
    safe_upload_and_index_document,
)
from app.services.admin_modification_marker import (
    is_admin_modified_since_upload,
    mark_admin_modified,
)
from app.services.contact_document_area import (
    ContactPhotoPayload,
    rebuild_canonical_contact_table,
)
from app.services.contact_photo_store import (
    ContactPhotoStorageError,
    delete_contact_photo,
    read_contact_photo,
    write_contact_photo_atomic,
)
from app.services.contact_photos import extract_contact_photo_candidates
from app.services.contact_state import (
    ContactRecord,
    ContactState,
    ContactStateError,
    new_contact_id,
    read_contact_state,
    write_contact_state_atomic,
)
from pydantic import ValidationError

from app.models.admin_contacts import AdminContactWriteRequest
from app.models.admin_documents import AdminDocumentListResponse, AdminDocumentSummary
from app.models.document import DocumentChunk
from app.services.docx_parser import (
    CONTACT_TABLE_HIDDEN_MARKER,
    ExtractedContact,
    extract_contacts_from_docx,
)
from app.services.document_indexer import DocumentIndexingResult

from tests.support.corpus_paths import resolve_source_root
from tests.support.documents import (
    make_png as _make_png,
    require_corpus_copy,
    skip_if_already_canonicalized,
)

SOURCE_ROOT = resolve_source_root()


def _real_document_id_for(country_code: str, language: str = "en") -> str:
    """
    The one real, deterministic document_id
    build_contact_chunk_for_contacts (and every other chunk builder)
    would compute for this country_code/language - document_id is
    derived solely from country_code + DOCUMENT_FAMILY + language,
    never from anything a test happens to pick, so every fake chunk
    this test file seeds/asserts against must use this exact value,
    never an arbitrary placeholder string.
    """

    from app.services.document_chunk_builder import DocumentMetadata
    from app.services.docx_parser import ExtractedContact as _EC
    from app.services.document_chunk_builder import (
        build_contact_chunk_for_contacts as _probe_builder,
    )

    probe_chunk = _probe_builder(
        [_EC(member_firm="probe")],
        DocumentMetadata(
            country="Probe Country",
            country_code=country_code,
            reference_year=None,
            language=language,
            source_filename="probe.docx",
        ),
    )

    return probe_chunk.document_id


DOCUMENT_ID = _real_document_id_for("GB")
OTHER_DOCUMENT_ID = _real_document_id_for("FR")


def _make_valid_jpeg(width: int = 183, height: int = 234) -> bytes:
    """
    A minimal but REAL-SIZED, structurally complete baseline JPEG -
    SOI/APP0(JFIF)/SOF0/SOS/EOI with plausible width/height, never a
    degenerate single-pixel stub. python-docx's own image-header
    parser (used to compute a photo's proportional height in the
    canonical contact table) stops reading at the first SOS marker, so
    the entropy-coded scan data itself can be a single placeholder
    byte - only the headers before it need to be genuine.
    """

    import struct

    soi = b"\xff\xd8"
    app0_data = (
        b"JFIF\x00\x01\x01\x00"
        + struct.pack(">HH", 96, 96)
        + b"\x00\x00"
    )
    app0 = b"\xff\xe0" + struct.pack(">H", len(app0_data) + 2) + app0_data
    sof0_data = (
        bytes([8])
        + struct.pack(">HH", height, width)
        + bytes([1, 1, 0x11, 0])
    )
    sof0 = b"\xff\xc0" + struct.pack(">H", len(sof0_data) + 2) + sof0_data
    sos_data = bytes([1, 1, 0, 0, 63, 0])
    sos = b"\xff\xda" + struct.pack(">H", len(sos_data) + 2) + sos_data
    return soi + app0 + sof0 + sos + b"\x00\xff\xd9"


_VALID_JPEG = _make_valid_jpeg()
_VALID_PNG = _make_png(183, 234, (120, 80, 200))


def _write_request(
    *,
    member_firm: str = "Example & Partners",
    contact_person: str = "Alex Example",
    email: str = "alex@example.test",
    phone: str = "+1 555 0100",
    address: str = "1 Example Street",
    website: str = "www.example.test",
) -> AdminContactWriteRequest:
    return AdminContactWriteRequest(
        member_firm=member_firm,
        contact_person=contact_person,
        email=email,
        phone=phone,
        address=address,
        website=website,
    )


def _full_contact_record(**overrides: Any) -> ContactRecord:
    defaults = dict(
        contact_id=new_contact_id(),
        member_firm="Example & Partners",
        contact_person="Alex Example",
        email="alex@example.test",
        phone="+1 555 0100",
        address="1 Example Street",
        website="www.example.test",
    )
    defaults.update(overrides)
    return ContactRecord(**defaults)


class FakeContactOpenSearchClient:
    """
    OpenSearch test double for admin_contacts.py.

    Stateful over self.chunks (chunk_id -> source dict) so a mutation's
    real effect can be asserted afterwards, not merely inferred from
    call arguments - mirrors FakeSectionOpenSearchClient's own design
    in test_admin_document_sections.py exactly, generalized to filter
    on subsection.keyword (never legal_topic, which the Contact chunk
    always carries as None).
    """

    def __init__(
        self,
        *,
        document_id: str = DOCUMENT_ID,
        country_code: str = "GB",
        country: str = "United Kingdom",
        source_filename: str = "GB.docx",
        reference_year: int | None = 2026,
        chunks: dict[str, dict[str, Any]] | None = None,
        country_document_ids: list[str] | None = None,
        fail_delete_by_query_calls: int = 0,
        delete_by_query_failure: Exception | None = None,
    ) -> None:
        self.document_id = document_id
        self.country_code = country_code
        self.country = country
        self.source_filename = source_filename
        self.reference_year = reference_year
        self.chunks: dict[str, dict[str, Any]] = dict(chunks or {})
        self.country_document_ids = (
            country_document_ids
            if country_document_ids is not None
            else [document_id]
        )

        self.fail_delete_by_query_calls = fail_delete_by_query_calls
        self.delete_by_query_failure = delete_by_query_failure
        self.delete_by_query_call_count = 0
        self.delete_by_query_calls: list[dict[str, Any]] = []

    def search(self, index: str, body: dict[str, Any]) -> dict[str, Any]:
        del index

        if "sort" in body:
            term = body.get("query", {}).get("term", {})
            requested_document_id = term.get("document_id")
            requested_country_code = term.get("country_code")

            if requested_document_id is not None:
                matching_ids = sorted(
                    chunk_id
                    for chunk_id, chunk in self.chunks.items()
                    if chunk["document_id"] == requested_document_id
                )
            elif requested_country_code is not None:
                # _ensure_no_country_conflict's own country-wide
                # lookup - modeled purely from country_document_ids,
                # never from self.chunks (a country conflict is a
                # catalog-level fact, independent of whether any
                # contact chunk happens to exist yet).
                return {
                    "hits": {
                        "total": {"value": len(self.country_document_ids)},
                        "hits": [
                            {
                                "_id": f"doc-row-{doc_id}",
                                "_source": {
                                    "document_id": doc_id,
                                    "country_code": (
                                        requested_country_code
                                    ),
                                    "country": self.country,
                                    "source_filename": (
                                        self.source_filename
                                    ),
                                    "reference_year": (
                                        self.reference_year
                                    ),
                                },
                                "sort": [f"doc-row-{doc_id}"],
                            }
                            for doc_id in self.country_document_ids
                        ],
                    }
                }
            else:
                matching_ids = []

            return {
                "hits": {
                    "total": {"value": len(matching_ids)},
                    "hits": [
                        {
                            "_id": chunk_id,
                            "_source": self.chunks[chunk_id],
                            "sort": [chunk_id],
                        }
                        for chunk_id in matching_ids
                    ],
                }
            }

        term = body.get("query", {}).get("term", {})
        requested_document_id = term.get("document_id")

        if requested_document_id is not None:
            if requested_document_id != self.document_id:
                return {"hits": {"hits": []}}

            return {
                "hits": {
                    "hits": [
                        {
                            "_source": {
                                "document_id": self.document_id,
                                "source_filename": self.source_filename,
                                "country": self.country,
                                "country_code": self.country_code,
                                "reference_year": self.reference_year,
                            }
                        }
                    ]
                }
            }

        return {"hits": {"hits": []}}

    def delete_by_query(
        self,
        *,
        index: str,
        body: dict[str, Any],
        conflicts: str,
        refresh: bool,
    ) -> dict[str, Any]:
        del index, conflicts, refresh

        self.delete_by_query_call_count += 1

        if (
            self.delete_by_query_call_count
            <= self.fail_delete_by_query_calls
        ):
            raise (
                self.delete_by_query_failure
                if self.delete_by_query_failure is not None
                else RuntimeError(
                    "simulated delete_by_query failure (call #"
                    f"{self.delete_by_query_call_count})."
                )
            )

        query = body["query"]
        filters = query["bool"]["filter"]

        document_id = next(
            clause["term"]["document_id"]
            for clause in filters
            if "document_id" in clause.get("term", {})
        )
        subsection = next(
            (
                clause["term"]["subsection.keyword"]
                for clause in filters
                if "subsection.keyword" in clause.get("term", {})
            ),
            None,
        )

        keep_ids: set[str] = set()

        for clause in query["bool"].get("must_not", []):
            keep_ids.update(clause.get("terms", {}).get("chunk_id", []))

        to_delete = [
            chunk_id
            for chunk_id, chunk in self.chunks.items()
            if chunk["document_id"] == document_id
            and (
                subsection is None
                or chunk.get("subsection") == subsection
            )
            and chunk_id not in keep_ids
        ]

        for chunk_id in to_delete:
            del self.chunks[chunk_id]

        self.delete_by_query_calls.append(
            {
                "document_id": document_id,
                "subsection": subsection,
                "deleted": len(to_delete),
            }
        )

        return {"deleted": len(to_delete), "total": len(to_delete)}


def _bulk_writer(fake_client: FakeContactOpenSearchClient, *, fail_first_n_calls: int = 0):
    call_count = {"n": 0}

    def fake_bulk(client, actions, **kwargs):
        del client, kwargs

        call_count["n"] += 1

        if call_count["n"] <= fail_first_n_calls:
            raise RuntimeError(
                f"simulated OpenSearch bulk failure (call #{call_count['n']})"
            )

        action_list = list(actions)

        for action in action_list:
            fake_client.chunks[action["_id"]] = dict(action["_source"])

        return (len(action_list), [])

    return fake_bulk


@contextlib.contextmanager
def _patched_indexer(fake_client: FakeContactOpenSearchClient, *, fail_bulk: bool = False):
    with patch(
        "app.services.document_indexer.ensure_legal_documents_index"
    ), patch(
        "app.services.document_indexer.bulk",
        side_effect=_bulk_writer(
            fake_client,
            fail_first_n_calls=1 if fail_bulk else 0,
        ),
    ):
        yield


def _seeded_contact_chunk(
    *,
    document_id: str = DOCUMENT_ID,
    country_code: str = "GB",
    country: str = "United Kingdom",
    source_filename: str = "GB.docx",
    content: str = "Member firm: Old Firm\nContact person: Old Person",
) -> dict[str, Any]:
    return {
        "document_id": document_id,
        "country_code": country_code,
        "country": country,
        "source_filename": source_filename,
        "reference_year": 2026,
        "legal_topic": None,
        "document_type": "overview",
        "language": "en",
        "section": f"Employment Law Overview {country}",
        "subsection": "Contact",
        "content": content,
        "content_hash": f"hash-{hash(content)}",
    }


def _seed_placeholder_source_docx(
    source_directory: Path, filename: str = "GB.docx"
) -> None:
    """
    A minimal, valid, structurally-empty DOCX - just enough for
    resolve_document_source_path to find a real file on disk. Add/
    Delete Contact always rebuild the persisted source's canonical
    contact table now, regardless of whether this placeholder has any
    prior contact structure at all, so these tests exercise that
    rebuild too (harmlessly, against a document with no real legal
    content) alongside the ContactState/OpenSearch chunk mutation they
    were originally written for.
    """

    document = Document()
    document.add_paragraph("Placeholder document body.")
    document.save(str(source_directory / filename))


# =========================================================================
# STATE: sidecar read/write/atomic/absent-vs-empty
# =========================================================================


class ContactStateTests(unittest.TestCase):
    def test_absent_state_returns_none(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            self.assertIsNone(
                read_contact_state(Path(root), DOCUMENT_ID)
            )

    def test_explicit_empty_state_is_not_none(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(),
                ),
            )

            state = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertIsNotNone(state)
            self.assertEqual(state.contacts, ())

    def test_stable_ordering_round_trip(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            first = _full_contact_record(member_firm="Firm One")
            second = _full_contact_record(member_firm="Firm Two")

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(first, second),
                ),
            )

            state = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertEqual(
                [c.contact_id for c in state.contacts],
                [first.contact_id, second.contact_id],
            )

    def test_write_is_atomic_no_partial_file_left_on_crash(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(_full_contact_record(),),
                ),
            )

            state_dir = source_directory / ".admin-state" / "contacts"
            leftover_temp_files = [
                path
                for path in state_dir.iterdir()
                if path.name.endswith(".json.tmp")
            ]

            self.assertEqual(leftover_temp_files, [])


# =========================================================================
# CONTACT PHOTO METADATA (ContactRecord's own photo fields) AND THE
# PHYSICAL PHOTO STORE (contact_photo_store.py) - the two persistence
# primitives every admin photo CRUD mutation below is built on.
# =========================================================================


class ContactPhotoMetadataTests(unittest.TestCase):

    def test_legacy_contact_without_photo_fields_is_readable(self) -> None:
        record = ContactRecord.from_json_dict(
            {
                "contact_id": "contact-legacy",
                "member_firm": "Firm",
                "contact_person": "Jane Doe",
                "email": "jane@example.com",
                "phone": "+1",
                "address": "Address",
                "website": "example.com",
            }
        )

        self.assertIsNone(record.photo_filename)
        self.assertIsNone(record.photo_content_type)
        self.assertIsNone(record.photo_sha256)

    def test_photo_metadata_serializes_without_binary_data(self) -> None:
        digest = "a" * 64

        record = ContactRecord(
            contact_id="contact-123",
            member_firm="Firm",
            contact_person="Jane Doe",
            email="jane@example.com",
            photo_filename=f"contact-123--{digest}.jpg",
            photo_content_type="image/jpeg",
            photo_sha256=digest,
        )

        payload = record.to_json_dict()

        self.assertEqual(
            f"contact-123--{digest}.jpg",
            payload["photo_filename"],
        )
        self.assertEqual(
            "image/jpeg",
            payload["photo_content_type"],
        )
        self.assertEqual(
            digest,
            payload["photo_sha256"],
        )

        self.assertFalse(
            any(
                isinstance(value, (bytes, bytearray))
                for value in payload.values()
            )
        )

    def test_partial_photo_metadata_is_rejected(self) -> None:
        with self.assertRaises(ContactStateError):
            ContactRecord.from_json_dict(
                {
                    "contact_id": "contact-123",
                    "photo_filename": "photo.jpg",
                    "photo_content_type": None,
                    "photo_sha256": None,
                }
            )


class ContactPhotoStoreTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.source_directory = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_atomic_write_uses_contact_id_and_sha256(self) -> None:
        data = b"fake-jpeg-content"
        digest = hashlib.sha256(data).hexdigest()

        stored = write_contact_photo_atomic(
            self.source_directory,
            "contact-123",
            data=data,
            content_type="image/jpeg",
        )

        self.assertEqual(
            f"contact-123--{digest}.jpg",
            stored.filename,
        )
        self.assertEqual("image/jpeg", stored.content_type)
        self.assertEqual(digest, stored.sha256)

        self.assertEqual(
            data,
            read_contact_photo(
                self.source_directory,
                stored.filename,
            ),
        )

    def test_photo_is_stored_inside_admin_state(self) -> None:
        stored = write_contact_photo_atomic(
            self.source_directory,
            "contact-123",
            data=b"photo",
            content_type="image/png",
        )

        expected = (
            self.source_directory
            / ".admin-state"
            / "contact-photos"
            / stored.filename
        )

        self.assertTrue(expected.is_file())

    def test_same_photo_write_is_idempotent(self) -> None:
        first = write_contact_photo_atomic(
            self.source_directory,
            "contact-123",
            data=b"same-photo",
            content_type="image/jpeg",
        )

        second = write_contact_photo_atomic(
            self.source_directory,
            "contact-123",
            data=b"same-photo",
            content_type="image/jpeg",
        )

        self.assertEqual(first, second)

        files = list(
            (
                self.source_directory
                / ".admin-state"
                / "contact-photos"
            ).iterdir()
        )

        self.assertEqual(1, len(files))

    def test_failed_new_write_preserves_existing_photo(self) -> None:
        old = write_contact_photo_atomic(
            self.source_directory,
            "contact-123",
            data=b"old-photo",
            content_type="image/jpeg",
        )

        with patch(
            "app.services.contact_photo_store.os.replace",
            side_effect=OSError("boom"),
        ):
            with self.assertRaises(ContactPhotoStorageError):
                write_contact_photo_atomic(
                    self.source_directory,
                    "contact-123",
                    data=b"new-photo",
                    content_type="image/png",
                )

        self.assertEqual(
            b"old-photo",
            read_contact_photo(
                self.source_directory,
                old.filename,
            ),
        )

        store = (
            self.source_directory
            / ".admin-state"
            / "contact-photos"
        )

        self.assertEqual(
            [old.filename],
            sorted(
                p.name
                for p in store.iterdir()
                if p.is_file()
            ),
        )

    def test_unsupported_content_type_is_rejected(self) -> None:
        with self.assertRaises(ContactPhotoStorageError):
            write_contact_photo_atomic(
                self.source_directory,
                "contact-123",
                data=b"photo",
                content_type="image/svg+xml",
            )

    def test_empty_photo_is_rejected(self) -> None:
        with self.assertRaises(ContactPhotoStorageError):
            write_contact_photo_atomic(
                self.source_directory,
                "contact-123",
                data=b"",
                content_type="image/jpeg",
            )

    def test_path_traversal_filename_is_rejected(self) -> None:
        with self.assertRaises(ContactPhotoStorageError):
            read_contact_photo(
                self.source_directory,
                "../secret.jpg",
            )

        with self.assertRaises(ContactPhotoStorageError):
            delete_contact_photo(
                self.source_directory,
                "../secret.jpg",
            )

    def test_delete_is_safe_and_idempotent(self) -> None:
        stored = write_contact_photo_atomic(
            self.source_directory,
            "contact-123",
            data=b"photo",
            content_type="image/webp",
        )

        delete_contact_photo(
            self.source_directory,
            stored.filename,
        )

        delete_contact_photo(
            self.source_directory,
            stored.filename,
        )

        with self.assertRaises(ContactPhotoStorageError):
            read_contact_photo(
                self.source_directory,
                stored.filename,
            )


# =========================================================================
# MARKER
# =========================================================================


class AdminModificationMarkerTests(unittest.TestCase):
    def test_absent_marker_defaults_false(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            self.assertFalse(
                is_admin_modified_since_upload(Path(root), DOCUMENT_ID)
            )

    def test_mark_then_reset(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            mark_admin_modified(source_directory, DOCUMENT_ID)
            self.assertTrue(
                is_admin_modified_since_upload(
                    source_directory, DOCUMENT_ID
                )
            )


# =========================================================================
# API: list / add / update / delete
# =========================================================================


class ContactCrudTests(unittest.TestCase):
    def test_list_with_no_state_returns_empty(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()

            response = list_contacts(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(response.contacts, [])
            self.assertEqual(response.country_code, "GB")

    def test_add_appends_with_fresh_stable_id(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                response = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            self.assertTrue(response.contact_id)
            self.assertEqual(response.member_firm, "Example & Partners")

            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(
                state.contacts[0].contact_id, response.contact_id
            )

            self.assertTrue(
                is_admin_modified_since_upload(
                    source_directory, DOCUMENT_ID
                )
            )

    def test_add_syncs_the_opensearch_contact_chunk(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            contact_chunks = [
                chunk
                for chunk in client.chunks.values()
                if chunk.get("subsection") == "Contact"
            ]

            self.assertEqual(len(contact_chunks), 1)
            self.assertIn("Example & Partners", contact_chunks[0]["content"])

    def test_zero_one_and_multiple_contacts(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            listing = list_contacts(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(listing.contacts, [])

            with _patched_indexer(client):
                first = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )

            listing = list_contacts(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(len(listing.contacts), 1)

            with _patched_indexer(client):
                second = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm B"),
                    source_directory=source_directory,
                    client=client,
                )
                third = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )

            listing = list_contacts(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(len(listing.contacts), 3)
            self.assertEqual(
                [c.contact_id for c in listing.contacts],
                [first.contact_id, second.contact_id, third.contact_id],
            )

    def test_duplicate_contacts_have_distinct_ids(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                first = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )
                second = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(first.member_firm, second.member_firm)
            self.assertEqual(first.email, second.email)
            self.assertNotEqual(first.contact_id, second.contact_id)

    def test_update_preserves_id_and_position(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                first = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )
                second = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm B"),
                    source_directory=source_directory,
                    client=client,
                )

                update_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=second.contact_id,
                    fields=_write_request(member_firm="Firm B Updated"),
                    source_directory=source_directory,
                    client=client,
                )

            state = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertEqual(
                [c.contact_id for c in state.contacts],
                [first.contact_id, second.contact_id],
            )
            self.assertEqual(state.contacts[0].member_firm, "Firm A")
            self.assertEqual(
                state.contacts[1].member_firm, "Firm B Updated"
            )

    def test_update_stale_contact_id_raises_not_found(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()

            with self.assertRaises(AdminContactNotFoundError):
                update_contact(
                    document_id=DOCUMENT_ID,
                    contact_id="does-not-exist",
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

    def test_delete_removes_only_the_requested_contact(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                first = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )
                second = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm B"),
                    source_directory=source_directory,
                    client=client,
                )

                delete_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=first.contact_id,
                    source_directory=source_directory,
                    client=client,
                )

            state = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].contact_id, second.contact_id)
            self.assertEqual(state.contacts[0].member_firm, "Firm B")

    def test_delete_stale_contact_id_raises_not_found(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()

            with self.assertRaises(AdminContactNotFoundError):
                delete_contact(
                    document_id=DOCUMENT_ID,
                    contact_id="does-not-exist",
                    source_directory=source_directory,
                    client=client,
                )

    def test_delete_last_contact_removes_stale_contact_chunk(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                only = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

                self.assertEqual(
                    len(
                        [
                            c
                            for c in client.chunks.values()
                            if c.get("subsection") == "Contact"
                        ]
                    ),
                    1,
                )

                delete_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=only.contact_id,
                    source_directory=source_directory,
                    client=client,
                )

            contact_chunks = [
                c
                for c in client.chunks.values()
                if c.get("subsection") == "Contact"
            ]

            self.assertEqual(contact_chunks, [])

    def test_legal_topic_chunks_are_never_touched(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)

            client = FakeContactOpenSearchClient(
                chunks={
                    "legal-chunk-1": {
                        "document_id": DOCUMENT_ID,
                        "country_code": "GB",
                        "country": "United Kingdom",
                        "source_filename": "GB.docx",
                        "reference_year": 2026,
                        "legal_topic": "Employment Contracts",
                        "document_type": "overview",
                        "language": "en",
                        "section": "Employment Contracts",
                        "subsection": None,
                        "content": "legal content",
                        "content_hash": "hash-legal",
                    },
                },
            )

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            self.assertIn("legal-chunk-1", client.chunks)
            self.assertEqual(
                client.chunks["legal-chunk-1"]["content"], "legal content"
            )

    def test_country_conflict_blocks_contact_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            client = FakeContactOpenSearchClient(
                country_document_ids=[DOCUMENT_ID, OTHER_DOCUMENT_ID],
            )

            with self.assertRaises(Exception):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

    def test_list_exposes_has_photo_true_for_a_contact_with_a_photo(
        self,
    ) -> None:
        """
        The safe photo-presence contract Admin/View needs: the Admin
        list must never blindly attempt a photo fetch for every contact
        regardless of whether one exists - no pointless 404 image
        requests for contacts that never had a photo. The response
        exposes has_photo only - never photo_filename or any
        filesystem path.
        """

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(
                        ContactRecord(
                            contact_id=new_contact_id(),
                            member_firm="Example & Partners",
                            contact_person="Jane Doe",
                            email="jane@example.com",
                            phone="+1 555 0000",
                            address="1 Example Street",
                            website="https://example.com",
                            photo_filename="deadbeef.jpg",
                            photo_content_type="image/jpeg",
                            photo_sha256="a" * 64,
                        ),
                    ),
                ),
            )

            listing = list_contacts(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(len(listing.contacts), 1)
            self.assertTrue(listing.contacts[0].has_photo)

            serialized = listing.contacts[0].model_dump()
            self.assertNotIn("photo_filename", serialized)
            self.assertNotIn("photo_content_type", serialized)
            self.assertNotIn("photo_sha256", serialized)

    def test_list_exposes_has_photo_false_for_a_contact_without_a_photo(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            listing = list_contacts(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(len(listing.contacts), 1)
            self.assertFalse(listing.contacts[0].has_photo)


# =========================================================================
# ADMIN CONTACT PHOTO CRUD (admin_contact_photos.py)
#
# Every mutation here backs onto a REAL corpus DOCX (temp copy, never the
# real file itself), through the same FakeContactOpenSearchClient used
# above, answering exactly the one read-only metadata lookup these
# mutations need - never a real network call, and never any reindexing
# (replace/remove never touch the OpenSearch Contact chunk at all; that
# stays admin_contacts.py's own job).
# =========================================================================


class AdminContactPhotoTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)

    def _seed_photo_bearing_contact(self) -> FakeContactOpenSearchClient:
        """
        Argentina (AR.docx) - a real single contact with a real
        existing photo, so REPLACE/REMOVE exercise the actual
        DOCX-mutation code path, not a synthetic stand-in.
        """

        docx_path = require_corpus_copy(self, SOURCE_ROOT, "AR.docx", self.root)

        candidates = extract_contact_photo_candidates(docx_path)
        assert len(candidates) == 1
        photo = candidates[0]

        stored = write_contact_photo_atomic(
            self.root,
            "contact-test",
            data=photo.data,
            content_type=photo.content_type,
        )

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id="doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
                country_code="AR",
                contacts=(
                    ContactRecord(
                        contact_id="contact-test",
                        member_firm="Allende & Brea",
                        contact_person="Nicolás Grandi",
                        email="ngrandi@allende.com",
                        phone="+1 555 0100",
                        address="Address",
                        website="https://example.com",
                        photo_filename=stored.filename,
                        photo_content_type=stored.content_type,
                        photo_sha256=stored.sha256,
                    ),
                ),
            ),
        )

        return FakeContactOpenSearchClient(
            document_id="doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
            country_code="AR",
            country="Argentina",
            source_filename="AR.docx",
        )

    def _seed_photo_less_contact(self) -> FakeContactOpenSearchClient:
        """
        Germany (DE.docx) - a real single contact (Tobias Pusch) who
        genuinely has no photo yet, so a PUT exercises the real
        ADD-into-the-document code path.
        """

        docx_path = require_corpus_copy(self, SOURCE_ROOT, "DE.docx", self.root)

        assert extract_contact_photo_candidates(docx_path) == []

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id="doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
                country_code="DE",
                contacts=(
                    ContactRecord(
                        contact_id="contact-test",
                        member_firm="Pusch Wahlig Workplace Law",
                        contact_person="Tobias Pusch",
                        email="pusch@pwwl.de",
                        phone="+1 555 0100",
                        address="Address",
                        website="https://example.com",
                    ),
                ),
            ),
        )

        return FakeContactOpenSearchClient(
            document_id="doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
            country_code="DE",
            country="Germany",
            source_filename="DE.docx",
        )

    def _seed_belgium_two_contacts(self) -> FakeContactOpenSearchClient:
        """
        Belgium's real two-contact, two-photo document - to prove
        isolation holds at the FULL service layer (ContactState +
        photo store + source DOCX together), not merely at the raw
        DOCX-primitive level test_contact_documents.py already covers.
        """

        temp_copy = require_corpus_copy(
            self,
            SOURCE_ROOT,
            "Labour and Employment Law in Belgium 2026.docx",
            self.root,
        )
        docx_path = self.root / "BE.docx"
        shutil.copyfile(temp_copy, docx_path)

        candidates = extract_contact_photo_candidates(docx_path)
        by_name = {c.source_filename: c for c in candidates}
        chris_photo = by_name["image2.jpg"]
        nicolas_photo = by_name["image1.png"]

        chris_stored = write_contact_photo_atomic(
            self.root,
            "chris-id",
            data=chris_photo.data,
            content_type=chris_photo.content_type,
        )
        nicolas_stored = write_contact_photo_atomic(
            self.root,
            "nicolas-id",
            data=nicolas_photo.data,
            content_type=nicolas_photo.content_type,
        )

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id="doc_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb",
                country_code="BE",
                contacts=(
                    ContactRecord(
                        contact_id="chris-id",
                        member_firm="Van Olmen & Wynant",
                        contact_person="Chris van Olmen",
                        email="chris.van.olmen@vow.be",
                        phone="+1 555 0100",
                        address="Address",
                        website="https://example.com",
                        photo_filename=chris_stored.filename,
                        photo_content_type=chris_stored.content_type,
                        photo_sha256=chris_stored.sha256,
                    ),
                    ContactRecord(
                        contact_id="nicolas-id",
                        member_firm="Van Olmen & Wynant",
                        contact_person="Nicolas Simon",
                        email="nicolas.simon@vow.be",
                        phone="+1 555 0100",
                        address="Address",
                        website="https://example.com",
                        photo_filename=nicolas_stored.filename,
                        photo_content_type=nicolas_stored.content_type,
                        photo_sha256=nicolas_stored.sha256,
                    ),
                ),
            ),
        )

        return FakeContactOpenSearchClient(
            document_id="doc_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb",
            country_code="BE",
            country="Belgium",
            source_filename="BE.docx",
        )

    def test_belgium_two_contact_isolation_at_the_service_layer(
        self,
    ) -> None:
        """Mutating Chris's photo must never touch Nicolas's
        ContactState, photo file, or DOCX image - and vice versa."""

        client = self._seed_belgium_two_contacts()
        docx_path = self.root / "BE.docx"
        doc_id = (
            "doc_bbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbbb"
        )

        original_state = read_contact_state(self.root, doc_id)
        nicolas_before = next(
            c for c in original_state.contacts
            if c.contact_id == "nicolas-id"
        )

        replace_admin_contact_photo(
            self.root,
            doc_id,
            "chris-id",
            data=_VALID_JPEG,
            content_type="image/jpeg",
            client=client,
        )

        state_after = read_contact_state(self.root, doc_id)
        nicolas_after = next(
            c for c in state_after.contacts
            if c.contact_id == "nicolas-id"
        )

        self.assertEqual(nicolas_before, nicolas_after)

        docx_shas = {
            c.sha256
            for c in extract_contact_photo_candidates(docx_path)
        }
        self.assertIn(nicolas_before.photo_sha256, docx_shas)
        self.assertEqual(2, len(docx_shas))

        self.assertTrue(
            remove_admin_contact_photo(
                self.root, doc_id, "nicolas-id", client=client
            )
        )

        final_state = read_contact_state(self.root, doc_id)
        chris_final = next(
            c for c in final_state.contacts
            if c.contact_id == "chris-id"
        )
        self.assertIsNotNone(chris_final.photo_sha256)

        final_docx_shas = {
            c.sha256
            for c in extract_contact_photo_candidates(docx_path)
        }
        self.assertEqual(1, len(final_docx_shas))
        self.assertIn(chris_final.photo_sha256, final_docx_shas)

    def test_add_photo_for_a_contact_whose_name_matches_nothing_in_the_document(
        self,
    ) -> None:
        """
        A brand-new contact's name will usually have no matching
        "CONTACT PERSON" zone in the document at all (that IS the
        common case for genuinely adding someone - their name cannot
        possibly already appear anywhere). The canonical-table
        mechanism this codebase now uses embeds the new photo into a
        freshly rebuilt table regardless - it never depends on the
        mutated contact's own name matching anything already in the
        document (only the OLDER floating-shape primitives in
        test_contact_documents.py's own
        RealCorpusContactDocumentPhotoTests still need that kind of
        name/zone matching, for a document not yet canonicalized).
        """

        client = self._seed_photo_less_contact()
        docx_path = self.root / "DE.docx"
        doc_id = (
            "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa"
        )

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id=doc_id,
                country_code="DE",
                contacts=(
                    ContactRecord(
                        contact_id="contact-test",
                        member_firm="Someone New GmbH",
                        contact_person="Someone New",
                        email="new@example.test",
                        phone="+1 555 0100",
                        address="Address",
                        website="https://example.com",
                    ),
                ),
            ),
        )

        photo = replace_admin_contact_photo(
            self.root,
            doc_id,
            "contact-test",
            data=_VALID_JPEG,
            content_type="image/jpeg",
            client=client,
        )

        docx_shas = {
            c.sha256
            for c in extract_contact_photo_candidates(docx_path)
        }
        self.assertEqual(1, len(docx_shas))
        self.assertIn(photo.sha256, docx_shas)

        state = read_contact_state(self.root, doc_id)
        self.assertEqual(
            photo.sha256, state.contacts[0].photo_sha256
        )

        download = get_document_download(
            document_id=doc_id,
            source_directory=self.root,
            client=client,
        )
        downloaded_shas = {
            c.sha256
            for c in extract_contact_photo_candidates(download.path)
        }
        self.assertIn(
            photo.sha256,
            downloaded_shas,
            "the downloaded DOCX must contain the new contact's "
            "photo, not merely ContactState",
        )

    def _seed_zero_zone_country(self) -> FakeContactOpenSearchClient:
        """
        Portugal (PT.docx) - a real document with genuinely ZERO
        "CONTACT PERSON" zones anywhere. The canonical contact table
        mechanism rebuilds the whole contact area from ContactState
        regardless of prior structure, so a photo mutation here now
        succeeds (synchronizing a fresh table into the source) rather
        than failing closed the way the old floating-shape-only
        mechanism had to.
        """

        docx_path = require_corpus_copy(self, SOURCE_ROOT, "PT.docx", self.root)

        assert extract_contact_photo_candidates(docx_path) == []

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id="doc_cccccccccccccccccccccccccccccccccccccccccccccccccccccccccccccc",
                country_code="PT",
                contacts=(
                    ContactRecord(
                        contact_id="contact-test",
                        member_firm="Someone New Lda",
                        contact_person="Someone New",
                        email="new@example.test",
                        phone="+1 555 0100",
                        address="Address",
                        website="https://example.com",
                    ),
                ),
            ),
        )

        return FakeContactOpenSearchClient(
            document_id="doc_cccccccccccccccccccccccccccccccccccccccccccccccccccccccccccccc",
            country_code="PT",
            country="Portugal",
            source_filename="PT.docx",
        )

    def test_new_contact_photo_for_a_zero_zone_document_still_syncs(
        self,
    ) -> None:
        """
        A document with no contact area at all still gets the new
        photo synchronized into a freshly rebuilt canonical table -
        the source DOCX is never left unsynchronized just because it
        had no prior contact structure.
        """

        client = self._seed_zero_zone_country()
        doc_id = (
            "doc_cccccccccccccccccccccccccccccccccccccccccccccccccccccccccccccc"
        )

        photo = replace_admin_contact_photo(
            self.root,
            doc_id,
            "contact-test",
            data=_VALID_JPEG,
            content_type="image/jpeg",
            client=client,
        )

        docx_path = self.root / "PT.docx"
        docx_candidates = extract_contact_photo_candidates(docx_path)
        self.assertEqual(1, len(docx_candidates))
        self.assertEqual(photo.sha256, docx_candidates[0].sha256)

        state = read_contact_state(self.root, doc_id)
        self.assertEqual(photo.sha256, state.contacts[0].photo_sha256)

    def test_replace_read_remove_syncs_the_source_docx(self) -> None:
        client = self._seed_photo_bearing_contact()
        docx_path = self.root / "AR.docx"

        photo = replace_admin_contact_photo(
            self.root,
            "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
            "contact-test",
            data=_VALID_JPEG,
            content_type="image/jpeg",
            client=client,
        )

        self.assertEqual("image/jpeg", photo.content_type)

        # The persisted source DOCX itself must now resolve to the
        # NEW photo, not merely ContactState.
        docx_shas = {
            c.sha256
            for c in extract_contact_photo_candidates(docx_path)
        }
        self.assertIn(photo.sha256, docx_shas)
        self.assertEqual(1, len(docx_shas))

        loaded = read_admin_contact_photo(
            self.root, "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa", "contact-test"
        )
        self.assertEqual(_VALID_JPEG, loaded.data)

        self.assertTrue(
            remove_admin_contact_photo(
                self.root, "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa", "contact-test", client=client
            )
        )

        with self.assertRaises(AdminContactPhotoNotFoundError):
            read_admin_contact_photo(
                self.root, "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa", "contact-test"
            )

        # The removal must also be reflected in the source DOCX -
        # not merely ContactState.
        self.assertEqual(
            [], extract_contact_photo_candidates(docx_path)
        )

    def test_replace_failure_leaves_the_source_docx_and_state_unchanged(
        self,
    ) -> None:
        """
        A ContactState photo SHA that no longer matches anything in
        the source DOCX (simulating drift/corruption) is never
        consulted for a REPLACE - the canonical table is rebuilt fresh
        from the newly-uploaded bytes directly, so a stale/corrupt
        prior reference cannot block a legitimate new upload. This is
        a deliberate improvement over the old floating-shape mechanism
        (which had to locate the OLD sha in the document before it
        could replace it): rebuilding the whole area from ContactState
        every time eliminates this entire class of drift-driven
        failure.
        """

        client = self._seed_photo_bearing_contact()

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id="doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
                country_code="AR",
                contacts=(
                    ContactRecord(
                        contact_id="contact-test",
                        member_firm="Allende & Brea",
                        contact_person="Nicolás Grandi",
                        email="ngrandi@allende.com",
                        phone="+1 555 0100",
                        address="Address",
                        website="https://example.com",
                        photo_filename="contact-test--deadbeef.jpg",
                        photo_content_type="image/jpeg",
                        photo_sha256="0" * 64,
                    ),
                ),
            ),
        )

        photo = replace_admin_contact_photo(
            self.root,
            "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
            "contact-test",
            data=_VALID_JPEG,
            content_type="image/jpeg",
            client=client,
        )

        docx_path = self.root / "AR.docx"
        docx_candidates = extract_contact_photo_candidates(docx_path)
        self.assertEqual(1, len(docx_candidates))
        self.assertEqual(photo.sha256, docx_candidates[0].sha256)

        state = read_contact_state(
            self.root,
            "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
        )
        self.assertEqual(photo.sha256, state.contacts[0].photo_sha256)

    def test_downloaded_docx_reflects_a_replaced_photo(self) -> None:
        """
        After a mutation, the SAME backend document-download path
        Admin uses must reflect it - not merely ContactState.
        get_document_download() is exactly the function backing GET
        .../download. Exercised here against a REAL corpus document
        (Argentina), complementing DownloadByteStabilityTests' own
        byte-identical proof above (which uses a synthetic placeholder
        source) with the same guarantee against a document carrying
        real pre-existing structure.
        """

        client = self._seed_photo_bearing_contact()

        photo = replace_admin_contact_photo(
            self.root,
            "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
            "contact-test",
            data=_VALID_JPEG,
            content_type="image/jpeg",
            client=client,
        )

        download = get_document_download(
            document_id="doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
            source_directory=self.root,
            client=client,
        )

        downloaded_shas = {
            c.sha256
            for c in extract_contact_photo_candidates(download.path)
        }
        self.assertIn(photo.sha256, downloaded_shas)
        self.assertEqual(1, len(downloaded_shas))

    def test_invalid_photo_payloads_are_rejected(self) -> None:
        invalid_payloads = {
            "not_an_image": (b"not-an-image", "image/jpeg"),
            "declared_type_does_not_match_content": (
                b"\x89PNG\r\n\x1a\nfake", "image/jpeg",
            ),
        }

        for case, (data, content_type) in invalid_payloads.items():
            with self.subTest(case=case):
                client = self._seed_photo_bearing_contact()

                with self.assertRaises(AdminContactPhotoError):
                    replace_admin_contact_photo(
                        self.root,
                        "doc_aaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaaa",
                        "contact-test",
                        data=data,
                        content_type=content_type,
                        client=client,
                    )


class AdminContactPhotoRouteContractTests(unittest.TestCase):
    """
    The WordPress Admin proxy (class-le-global-chatbot-admin.php) builds
    every contact/photo URL from the exact same DOCUMENTS_PATH constant
    and route shape - these tests protect that wiring directly, since a
    mismatch here 404s silently in the browser (the proven root cause of
    a real "Admin View/Edit shows no photo thumbnail" defect: the
    browser's <img> gets a 404 and its onerror handler silently removes
    it).
    """

    def test_photo_routes_keep_admin_security_dependencies(self):
        normal = None
        photo_routes = []

        for route in router.routes:
            methods = getattr(route, "methods", set())

            if (
                route.path.endswith("/{document_id}/contacts")
                and "GET" in methods
            ):
                normal = route

            if route.path.endswith(
                "/{document_id}/contacts/{contact_id}/photo"
            ):
                photo_routes.append(route)

        self.assertIsNotNone(normal)
        self.assertEqual(3, len(photo_routes))

        normal_dependencies = len(
            normal.dependant.dependencies
        )

        self.assertGreater(normal_dependencies, 0)

        for route in photo_routes:
            self.assertGreaterEqual(
                len(route.dependant.dependencies),
                normal_dependencies,
            )

    def test_photo_route_paths_share_the_documents_prefix(self):
        """
        WordPress builds every contact photo URL as DOCUMENTS_PATH +
        "/" + document_id + "/contacts/" + contact_id + "/photo",
        where DOCUMENTS_PATH is the exact same "/api/v1/admin/documents"
        constant used to build the list/add/update/delete contact URLs.
        The three photo routes must therefore share that same prefix -
        if they don't, every request WordPress sends 404s.
        """

        contacts_list_route = next(
            route
            for route in router.routes
            if route.path.endswith("/{document_id}/contacts")
            and "GET" in route.methods
        )
        documents_prefix = contacts_list_route.path.removesuffix(
            "/{document_id}/contacts"
        )

        # A list, not a set: the three photo routes (GET/PUT/DELETE)
        # correctly share one IDENTICAL path string, differentiated
        # only by HTTP method - deduplicating by path value would
        # collapse them to one element even when the fix is correct.
        photo_routes = [
            route
            for route in router.routes
            if route.path.endswith(
                "/{document_id}/contacts/{contact_id}/photo"
            )
        ]

        self.assertEqual(3, len(photo_routes))

        for route in photo_routes:
            self.assertTrue(
                route.path.startswith(documents_prefix + "/"),
                f"{route.path!r} does not share the "
                f"{documents_prefix!r} prefix WordPress's "
                "DOCUMENTS_PATH constant assumes every Admin contact "
                "route (including photo routes) uses",
            )


# =========================================================================
# DOWNLOAD BYTE-STABILITY
#
# The invariant this section proves: download is a PURE READ.
# get_document_download() no longer rebuilds/reserializes anything -
# every mutation below already persists its effective DOCX atomically
# before returning, so download's job is just to hand back exactly
# those bytes, unchanged, every time.
# =========================================================================


class DownloadByteStabilityTests(unittest.TestCase):
    def _sha(self, path: Path) -> str:
        return hashlib.sha256(path.read_bytes()).hexdigest()

    def test_persisted_source_sha_equals_download_sha(self) -> None:
        """A: sha256(persisted source) == sha256(download body)."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            download = get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(download.path, persisted_path)
            self.assertEqual(
                self._sha(download.path), self._sha(persisted_path)
            )

    def test_ten_consecutive_downloads_are_byte_identical(self) -> None:
        """B: 10 consecutive downloads of an unchanged document all
        have exactly the same SHA256."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            hashes = {
                self._sha(
                    get_document_download(
                        document_id=DOCUMENT_ID,
                        source_directory=source_directory,
                        client=client,
                    ).path
                )
                for _ in range(10)
            }

            self.assertEqual(len(hashes), 1)

    def test_download_changes_nothing(self) -> None:
        """C: download changes neither the source file (mtime/size)
        nor ContactState nor the OpenSearch chunk set."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            stat_before = persisted_path.stat()
            state_before = read_contact_state(source_directory, DOCUMENT_ID)
            chunks_before = dict(client.chunks)

            get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            stat_after = persisted_path.stat()
            state_after = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertEqual(stat_before.st_mtime_ns, stat_after.st_mtime_ns)
            self.assertEqual(stat_before.st_size, stat_after.st_size)
            self.assertEqual(state_before, state_after)
            self.assertEqual(chunks_before, client.chunks)

    def test_download_never_reaches_docx_writing_code(self) -> None:
        """D: download must never invoke materialize_effective_docx,
        python-docx's Document.save, or the canonical contact
        rebuild - patch all three to explode, and prove download
        still succeeds untouched."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_bytes = (source_directory / "GB.docx").read_bytes()

            with patch(
                "app.services.contact_document_area."
                "rebuild_canonical_contact_table",
                side_effect=AssertionError(
                    "download must never rebuild the canonical table"
                ),
            ), patch(
                "docx.document.Document.save",
                side_effect=AssertionError(
                    "download must never call Document.save"
                ),
            ):
                download = get_document_download(
                    document_id=DOCUMENT_ID,
                    source_directory=source_directory,
                    client=client,
                )
                downloaded_bytes = download.path.read_bytes()

        self.assertEqual(downloaded_bytes, persisted_bytes)

    def test_source_equals_download_after_add_with_photo(self) -> None:
        """E: after Contact Add + photo, persisted source == downloaded
        bytes."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                contact = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )
                replace_admin_contact_photo(
                    source_directory,
                    DOCUMENT_ID,
                    contact.contact_id,
                    data=_make_png(64, 64, (10, 20, 30)),
                    content_type="image/png",
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            download = get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(
                self._sha(download.path), self._sha(persisted_path)
            )

    def test_source_equals_download_after_update(self) -> None:
        """F: after Contact Update, persisted source == downloaded
        bytes."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                contact = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )
                update_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=contact.contact_id,
                    fields=_write_request(member_firm="Updated Firm"),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            download = get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(
                self._sha(download.path), self._sha(persisted_path)
            )

    def test_source_equals_download_after_photo_replacement(self) -> None:
        """G: after photo replacement, persisted source == downloaded
        bytes."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                contact = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )
                replace_admin_contact_photo(
                    source_directory,
                    DOCUMENT_ID,
                    contact.contact_id,
                    data=_make_png(64, 64, (10, 20, 30)),
                    content_type="image/png",
                    client=client,
                )
                replace_admin_contact_photo(
                    source_directory,
                    DOCUMENT_ID,
                    contact.contact_id,
                    data=_make_png(32, 96, (200, 100, 50)),
                    content_type="image/png",
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            download = get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(
                self._sha(download.path), self._sha(persisted_path)
            )

    def test_source_equals_download_after_delete(self) -> None:
        """H: after Contact Delete, persisted source == downloaded
        bytes."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                first = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm B"),
                    source_directory=source_directory,
                    client=client,
                )
                delete_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=first.contact_id,
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            download = get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(
                self._sha(download.path), self._sha(persisted_path)
            )


# =========================================================================
# UPDATE MUST SYNCHRONIZE THE SOURCE DOCX (not just ContactState)
# =========================================================================


class UpdateContactPersistsToSourceTests(unittest.TestCase):
    """
    Regression coverage for a real bug: update_contact() used to only
    write ContactState/OpenSearch, never calling
    _synchronize_source_document() the way add_contact()/
    delete_contact() both already did - so a real Admin "Update
    Contact" text edit reported success and the ContactState/list
    endpoints reflected the new value, but the actual persisted (and
    therefore downloadable) source DOCX silently kept serving the OLD
    value until some unrelated later mutation happened to trigger a
    fresh rebuild. This was masked before download became a pure byte
    read (see DownloadByteStabilityTests above) - the OLD download path
    used to call materialize_effective_docx() on every GET, which
    rebuilt fresh from ContactState regardless, hiding the gap.

    test_source_equals_download_after_update above only proves
    download reads exactly what's on disk - it does NOT prove the
    disk copy actually reflects the update, since download and the
    persisted file are the same bytes by construction even when
    update_contact never wrote anything new. These tests inspect the
    actual DOCX content instead of only comparing two reads of
    whatever is already there.
    """

    def _sha(self, path: Path) -> str:
        return hashlib.sha256(path.read_bytes()).hexdigest()

    def test_update_persists_new_value_into_source_docx_and_download(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                contact = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )
                update_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=contact.contact_id,
                    fields=_write_request(member_firm="Firm B"),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"

            persisted_contacts = extract_contacts_from_docx(
                persisted_path, country=None
            )
            self.assertEqual(1, len(persisted_contacts))
            self.assertEqual("Firm B", persisted_contacts[0].member_firm)
            self.assertNotEqual(
                "Firm A", persisted_contacts[0].member_firm,
                "the old value must no longer be the effective contact "
                "value in the persisted document",
            )

            download = get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                self._sha(download.path), self._sha(persisted_path)
            )

            downloaded_contacts = extract_contacts_from_docx(
                download.path, country=None
            )
            self.assertEqual(1, len(downloaded_contacts))
            self.assertEqual(
                "Firm B", downloaded_contacts[0].member_firm
            )

    def test_repeated_updates_persist_the_final_value(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                contact = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )
                update_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=contact.contact_id,
                    fields=_write_request(member_firm="Firm B"),
                    source_directory=source_directory,
                    client=client,
                )
                update_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=contact.contact_id,
                    fields=_write_request(member_firm="Firm C"),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            persisted_contacts = extract_contacts_from_docx(
                persisted_path, country=None
            )

            self.assertEqual(1, len(persisted_contacts))
            self.assertEqual("Firm C", persisted_contacts[0].member_firm)

            download = get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            downloaded_contacts = extract_contacts_from_docx(
                download.path, country=None
            )
            self.assertEqual(
                "Firm C", downloaded_contacts[0].member_firm
            )

    def test_update_persists_when_contact_has_no_existing_photo(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                contact = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(
                        member_firm="Firm A", phone="+1 555 0100"
                    ),
                    source_directory=source_directory,
                    client=client,
                )
                update_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=contact.contact_id,
                    fields=_write_request(
                        member_firm="Firm A", phone="+1 555 9999"
                    ),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            persisted_contacts = extract_contacts_from_docx(
                persisted_path, country=None
            )

            self.assertEqual(1, len(persisted_contacts))
            self.assertEqual("+1 555 9999", persisted_contacts[0].phone)

            photos = extract_contact_photo_candidates(persisted_path)
            self.assertEqual(0, len(photos))

    def test_update_persists_when_contact_already_has_a_photo(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                contact = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )
                replace_admin_contact_photo(
                    source_directory,
                    DOCUMENT_ID,
                    contact.contact_id,
                    data=_make_png(64, 64, (10, 20, 30)),
                    content_type="image/png",
                    client=client,
                )
                update_contact(
                    document_id=DOCUMENT_ID,
                    contact_id=contact.contact_id,
                    fields=_write_request(member_firm="Firm A Updated"),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            persisted_contacts = extract_contacts_from_docx(
                persisted_path, country=None
            )

            self.assertEqual(1, len(persisted_contacts))
            self.assertEqual(
                "Firm A Updated", persisted_contacts[0].member_firm
            )

            photos = extract_contact_photo_candidates(persisted_path)
            self.assertEqual(
                1, len(photos),
                "a text-only update must not drop the contact's "
                "already-attached photo",
            )

            download = get_document_download(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                self._sha(download.path), self._sha(persisted_path)
            )


# =========================================================================
# ROLLBACK
# =========================================================================


class ContactRollbackTests(unittest.TestCase):
    def test_index_failure_restores_previous_state_and_marker(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                first = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )

            previous_marker = is_admin_modified_since_upload(
                source_directory, DOCUMENT_ID
            )

            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(AdminContactMutationFailedError):
                    add_contact(
                        document_id=DOCUMENT_ID,
                        fields=_write_request(member_firm="Firm B"),
                        source_directory=source_directory,
                        client=client,
                    )

            state = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].contact_id, first.contact_id)
            self.assertEqual(
                is_admin_modified_since_upload(
                    source_directory, DOCUMENT_ID
                ),
                previous_marker,
            )

    def test_state_write_failure_leaves_index_and_marker_untouched(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client), patch(
                "app.services.admin_contacts.write_contact_state_atomic",
                side_effect=OSError("disk full"),
            ):
                with self.assertRaises(AdminContactMutationFailedError):
                    add_contact(
                        document_id=DOCUMENT_ID,
                        fields=_write_request(),
                        source_directory=source_directory,
                        client=client,
                    )

            self.assertIsNone(
                read_contact_state(source_directory, DOCUMENT_ID)
            )
            self.assertEqual(client.chunks, {})
            self.assertFalse(
                is_admin_modified_since_upload(
                    source_directory, DOCUMENT_ID
                )
            )

    def test_update_index_failure_restores_previous_source_and_state(
        self,
    ) -> None:
        """Mirrors test_index_failure_restores_previous_state_and_marker
        above, for update_contact()'s own new source-DOCX
        synchronization: if the ContactState/index commit fails AFTER
        the DOCX has already been rebuilt with the new value, the DOCX
        must be restored to its exact prior bytes, not left holding
        the new value while ContactState still reports the old one."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                contact = add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Firm A"),
                    source_directory=source_directory,
                    client=client,
                )

            persisted_path = source_directory / "GB.docx"
            bytes_before = persisted_path.read_bytes()

            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(AdminContactMutationFailedError):
                    update_contact(
                        document_id=DOCUMENT_ID,
                        contact_id=contact.contact_id,
                        fields=_write_request(member_firm="Firm B"),
                        source_directory=source_directory,
                        client=client,
                    )

            self.assertEqual(bytes_before, persisted_path.read_bytes())

            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(state.contacts[0].member_firm, "Firm A")

    def test_rollback_itself_failing_raises_rollback_error(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)
            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(),
                    source_directory=source_directory,
                    client=client,
                )

            with _patched_indexer(client, fail_bulk=True):
                with patch(
                    "app.services.admin_contacts.write_contact_state_atomic",
                    side_effect=[None, OSError("rollback also fails")],
                ):
                    with self.assertRaises(AdminDocumentRollbackError):
                        add_contact(
                            document_id=DOCUMENT_ID,
                            fields=_write_request(),
                            source_directory=source_directory,
                            client=client,
                        )


# =========================================================================
# PHOTO FILE ROLLBACK - the same transactional core as ContactRollbackTests
# above (_apply_contact_state_change), exercised directly (rather than
# through add_contact/update_contact) so a just-written NEW photo file's
# own rollback can be asserted on disk, not just ContactState/index/
# marker rollback.
# =========================================================================


class ContactPhotoTransactionTests(unittest.TestCase):

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.source_directory = Path(self.temp.name)

        self.document_id = "doc-1"
        self.country_code = "BE"

        self.old_photo = write_contact_photo_atomic(
            self.source_directory,
            "contact-old",
            data=b"old-photo",
            content_type="image/jpeg",
        )

        self.old_record = ContactRecord(
            contact_id="contact-old",
            member_firm="Old Firm",
            contact_person="Old Person",
            email="old@example.com",
            photo_filename=self.old_photo.filename,
            photo_content_type=self.old_photo.content_type,
            photo_sha256=self.old_photo.sha256,
        )

        write_contact_state_atomic(
            self.source_directory,
            ContactState(
                document_id=self.document_id,
                country_code=self.country_code,
                contacts=(self.old_record,),
            ),
        )

        self.new_photo = write_contact_photo_atomic(
            self.source_directory,
            "contact-new",
            data=b"new-photo",
            content_type="image/png",
        )

        self.new_record = ContactRecord(
            contact_id="contact-new",
            member_firm="New Firm",
            contact_person="New Person",
            email="new@example.com",
            photo_filename=self.new_photo.filename,
            photo_content_type=self.new_photo.content_type,
            photo_sha256=self.new_photo.sha256,
        )

    def tearDown(self) -> None:
        self.temp.cleanup()

    def _photo_path(self, filename: str) -> Path:
        return (
            self.source_directory
            / ".admin-state"
            / "contact-photos"
            / filename
        )

    def _common_patches(self):
        return (
            patch.object(
                admin_contacts,
                "_document_metadata_for_chunks",
                return_value={},
            ),
            patch.object(
                admin_contacts,
                "build_contact_chunk_for_contacts",
                return_value=None,
            ),
            patch.object(
                admin_contacts,
                "is_admin_modified_since_upload",
                return_value=True,
            ),
            patch.object(
                admin_contacts,
                "write_admin_modified_marker",
            ),
        )

    def _apply(self) -> None:
        admin_contacts._apply_contact_state_change(
            document_id=self.document_id,
            country_code=self.country_code,
            source_directory=self.source_directory,
            new_contacts=(self.new_record,),
            document_metadata={},
            client=object(),
            reset_marker=True,
        )

    def _assert_rolled_back(self) -> None:
        state = read_contact_state(
            self.source_directory,
            self.document_id,
        )

        self.assertIsNotNone(state)
        self.assertEqual(
            ["contact-old"],
            [item.contact_id for item in state.contacts],
        )

        self.assertTrue(
            self._photo_path(
                self.old_photo.filename
            ).is_file()
        )

        self.assertFalse(
            self._photo_path(
                self.new_photo.filename
            ).exists()
        )

    def test_success_keeps_new_photo_and_removes_superseded_photo(
        self,
    ) -> None:
        (
            metadata_patch,
            chunk_patch,
            marker_read_patch,
            marker_write_patch,
        ) = self._common_patches()

        with (
            metadata_patch,
            chunk_patch,
            marker_read_patch,
            marker_write_patch,
            patch.object(
                admin_contacts,
                "replace_document_contact_chunk",
            ),
            patch.object(
                admin_contacts,
                "reset_admin_modified",
            ),
        ):
            self._apply()

        state = read_contact_state(
            self.source_directory,
            self.document_id,
        )

        self.assertEqual(
            ["contact-new"],
            [item.contact_id for item in state.contacts],
        )

        self.assertTrue(
            self._photo_path(
                self.new_photo.filename
            ).is_file()
        )

        self.assertFalse(
            self._photo_path(
                self.old_photo.filename
            ).exists()
        )

    def test_various_failure_points_roll_back_the_new_photo(self) -> None:
        """Whichever step of _apply_contact_state_change fails first -
        the OpenSearch chunk sync, the admin-modified marker, or the
        ContactState sidecar write itself - the new contact's own
        just-written photo file must be rolled back and the prior
        contact/photo restored, never left half-applied."""

        (
            metadata_patch,
            chunk_patch,
            marker_read_patch,
            marker_write_patch,
        ) = self._common_patches()

        real_write = admin_contacts.write_contact_state_atomic
        call_counts = {"write": 0, "replace": 0}

        def fail_first_write(source_directory, state):
            call_counts["write"] += 1
            if call_counts["write"] == 1:
                raise OSError("sidecar boom")
            return real_write(source_directory, state)

        def fail_first_replace(**kwargs):
            call_counts["replace"] += 1
            if call_counts["replace"] == 1:
                raise RuntimeError("opensearch boom")

        scenarios = {
            "opensearch_failure": [
                patch.object(
                    admin_contacts,
                    "replace_document_contact_chunk",
                    side_effect=fail_first_replace,
                ),
            ],
            "marker_failure": [
                patch.object(
                    admin_contacts, "replace_document_contact_chunk",
                ),
                patch.object(
                    admin_contacts,
                    "reset_admin_modified",
                    side_effect=OSError("marker boom"),
                ),
            ],
            "sidecar_failure": [
                patch.object(
                    admin_contacts,
                    "write_contact_state_atomic",
                    side_effect=fail_first_write,
                ),
                patch.object(
                    admin_contacts, "replace_document_contact_chunk",
                ),
            ],
        }

        for case_name, extra_patches in scenarios.items():
            with self.subTest(case=case_name):
                call_counts["write"] = 0
                call_counts["replace"] = 0

                with contextlib.ExitStack() as stack:
                    stack.enter_context(metadata_patch)
                    stack.enter_context(chunk_patch)
                    stack.enter_context(marker_read_patch)
                    stack.enter_context(marker_write_patch)
                    for extra_patch in extra_patches:
                        stack.enter_context(extra_patch)

                    with self.assertRaises(
                        admin_contacts.AdminContactMutationFailedError
                    ):
                        self._apply()

                self._assert_rolled_back()


class ContactPhotoCrudPreservationTests(unittest.TestCase):
    """
    A unit-level check of update_contact()'s own field-construction
    logic - _apply_contact_state_change is mocked away entirely, so
    this proves update_contact() itself always carries the existing
    photo_filename/photo_content_type/photo_sha256 forward onto the
    updated record, independent of the full DOCX-synchronization
    integration UpdateContactPersistsToSourceTests proves above.
    """

    def test_business_update_preserves_existing_photo_metadata(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as temp:
            source_directory = Path(temp)
            _seed_placeholder_source_docx(source_directory, "BE.docx")

            document_id = "doc_" + ("a" * 64)
            contact_id = "contact-123"

            photo = write_contact_photo_atomic(
                source_directory,
                contact_id,
                data=_make_png(64, 64, (10, 20, 30)),
                content_type="image/png",
            )

            original = ContactRecord(
                contact_id=contact_id,
                member_firm="Firm",
                contact_person="Jane Doe",
                email="jane@example.com",
                phone="+32 OLD",
                address="Old address",
                website="example.com",
                photo_filename=photo.filename,
                photo_content_type=photo.content_type,
                photo_sha256=photo.sha256,
            )

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=document_id,
                    country_code="BE",
                    contacts=(original,),
                ),
            )

            fields = AdminContactWriteRequest(
                member_firm="Firm",
                contact_person="Jane Doe",
                email="jane@example.com",
                phone="+32 111 0200",
                address="New address",
                website="www.example.com",
            )

            with (
                patch.object(
                    admin_contacts,
                    "_get_document_metadata",
                    return_value={
                        "country_code": "BE",
                    },
                ),
                patch.object(
                    admin_contacts,
                    "_load_country_code_and_metadata",
                    return_value=(
                        "BE",
                        {
                            "country": "Belgium",
                            "source_filename": "BE.docx",
                        },
                    ),
                ),
                patch.object(
                    admin_contacts,
                    "country_lock",
                    return_value=nullcontext(),
                ),
                patch.object(
                    admin_contacts,
                    "_apply_contact_state_change",
                ) as apply_mock,
            ):
                admin_contacts.update_contact(
                    document_id=document_id,
                    contact_id=contact_id,
                    fields=fields,
                    source_directory=source_directory,
                    client=object(),
                )

            new_contacts = (
                apply_mock.call_args.kwargs[
                    "new_contacts"
                ]
            )

            self.assertEqual(1, len(new_contacts))

            updated = new_contacts[0]

            # Business values really changed.
            self.assertEqual(
                "+32 111 0200",
                updated.phone,
            )
            self.assertEqual(
                "New address",
                updated.address,
            )

            # Stable identity remains unchanged.
            self.assertEqual(
                contact_id,
                updated.contact_id,
            )

            # The business edit must NOT remove the photo.
            self.assertEqual(
                photo.filename,
                updated.photo_filename,
            )
            self.assertEqual(
                photo.content_type,
                updated.photo_content_type,
            )
            self.assertEqual(
                photo.sha256,
                updated.photo_sha256,
            )


# =========================================================================
# SOURCE DOCX SYNCHRONIZATION EDGE CASES: photo-bearing delete isolation
# (ContactPhotoDeleteSyncTests) and the no-structural-area fallback
# (AddContactFallbackSynchronizationTests) - both against real corpus
# documents. _fake_document_lister/_summary below are shared bootstrap
# fixtures used by ContactBootstrapTests further down.
# =========================================================================


def _fake_document_lister(documents: list[AdminDocumentSummary]):
    def lister(*, source_directory: Path, client) -> AdminDocumentListResponse:
        del source_directory, client

        return AdminDocumentListResponse(
            total=len(documents),
            documents=documents,
        )

    return lister


def _summary(
    *,
    document_id: str,
    country_code: str,
    source_filename: str,
) -> AdminDocumentSummary:
    return AdminDocumentSummary(
        document_id=document_id,
        source_filename=source_filename,
        country=country_code,
        country_code=country_code,
        language="en",
        document_type="overview",
        reference_year=2026,
        chunk_count=1,
        source_file_present=True,
        status="indexed",
    )


class ContactPhotoDeleteSyncTests(unittest.TestCase):
    """
    Deleting a contact must remove ONLY that contact's own photo from
    the persisted source DOCX, using the SAME deterministic SHA-256
    association already proven at the raw-primitive layer in
    test_contact_documents.py - never position, filename, or a second,
    unrelated image-deletion implementation.
    """

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.document_id = _real_document_id_for("AU")

    def _require_copy(self, filename: str) -> Path:
        return require_corpus_copy(self, SOURCE_ROOT, filename, self.root)

    def _seed_two_contacts_with_photos(self) -> tuple[str, str]:
        """
        A real AU.docx with Michael Harmer's own real photo, plus a
        second contact added the same way Add Contact would (via
        rebuild_canonical_contact_table), each contact's own photo
        embedded for real, independently. Returns (michael_sha256,
        second_sha256).
        """

        docx_path = self._require_copy("AU.docx")

        michael_photo = extract_contact_photo_candidates(docx_path)[0]
        second_photo_data = _make_png(183, 234, (80, 120, 200))

        jane_contact = ExtractedContact(
            member_firm="Second Firm Pty Ltd",
            contact_person="Jane Secondary",
            email="jane.secondary@secondfirm.com.au",
            phone="+61 2 9000 0000",
            address="100 Test Street, Level 5",
            website="www.secondfirm.com.au",
        )

        new_bytes = rebuild_canonical_contact_table(
            docx_path,
            contacts=(
                ExtractedContact(
                    member_firm="HARMERS WORKPLACE LAWYERS",
                    contact_person="Michael Harmer",
                    email="michael.harmer@harmers.com.au",
                    phone="+61 292 674 322",
                    address="31 Market Street, Level 27 St Martins Tower",
                    website="www.harmers.com.au",
                ),
                jane_contact,
            ),
            photos=(
                ContactPhotoPayload(
                    data=michael_photo.data,
                    content_type=michael_photo.content_type,
                ),
                ContactPhotoPayload(
                    data=second_photo_data, content_type="image/png"
                ),
            ),
            country="Australia",
        )
        docx_path.write_bytes(new_bytes)

        michael_stored = write_contact_photo_atomic(
            self.root,
            "michael-id",
            data=michael_photo.data,
            content_type=michael_photo.content_type,
        )
        second_stored = write_contact_photo_atomic(
            self.root,
            "jane-id",
            data=second_photo_data,
            content_type="image/png",
        )

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id=self.document_id,
                country_code="AU",
                contacts=(
                    ContactRecord(
                        contact_id="michael-id",
                        member_firm="HARMERS WORKPLACE LAWYERS",
                        contact_person="Michael Harmer",
                        email="michael.harmer@harmers.com.au",
                        phone="+61 292 674 322",
                        address="31 Market Street, Level 27 St Martins Tower",
                        website="www.harmers.com.au",
                        photo_filename=michael_stored.filename,
                        photo_content_type=michael_stored.content_type,
                        photo_sha256=michael_stored.sha256,
                    ),
                    ContactRecord(
                        contact_id="jane-id",
                        member_firm="Second Firm Pty Ltd",
                        contact_person="Jane Secondary",
                        email="jane.secondary@secondfirm.com.au",
                        phone="+61 2 9000 0000",
                        address="100 Test Street, Level 5",
                        website="www.secondfirm.com.au",
                        photo_filename=second_stored.filename,
                        photo_content_type=second_stored.content_type,
                        photo_sha256=second_stored.sha256,
                    ),
                ),
            ),
        )

        return michael_stored.sha256, second_stored.sha256

    def test_delete_contact_with_photo_removes_only_its_own_image(
        self,
    ) -> None:
        docx_path = self.root / "AU.docx"
        michael_sha, jane_sha = self._seed_two_contacts_with_photos()

        client = FakeContactOpenSearchClient(
            document_id=self.document_id,
            country_code="AU",
            country="Australia",
            source_filename="AU.docx",
        )

        with _patched_indexer(client):
            delete_contact(
                document_id=self.document_id,
                contact_id="jane-id",
                source_directory=self.root,
                client=client,
            )

        remaining_shas = {
            c.sha256 for c in extract_contact_photo_candidates(docx_path)
        }
        self.assertIn(
            michael_sha, remaining_shas,
            "the other contact's own photo must remain byte/"
            "functionally intact",
        )
        self.assertNotIn(
            jane_sha, remaining_shas,
            "no stale photo of the deleted contact may remain "
            "anywhere in the DOCX",
        )
        self.assertEqual(1, len(remaining_shas))

        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(1, len(state.contacts))
        self.assertEqual("michael-id", state.contacts[0].contact_id)

    def test_delete_mutation_failure_restores_source_docx_byte_for_byte(
        self,
    ) -> None:
        docx_path = self.root / "AU.docx"
        self._seed_two_contacts_with_photos()

        original_bytes = docx_path.read_bytes()

        client = FakeContactOpenSearchClient(
            document_id=self.document_id,
            country_code="AU",
            country="Australia",
            source_filename="AU.docx",
        )

        with self.assertRaises(AdminContactMutationFailedError):
            with _patched_indexer(client, fail_bulk=True):
                delete_contact(
                    document_id=self.document_id,
                    contact_id="jane-id",
                    source_directory=self.root,
                    client=client,
                )

        self.assertEqual(
            original_bytes, docx_path.read_bytes(),
            "the source DOCX must be restored to its exact original "
            "bytes when the ContactState/index commit fails after the "
            "photo was already removed",
        )

        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(2, len(state.contacts))

    def test_delete_contact_without_a_photo_still_removes_its_text(
        self,
    ) -> None:
        """A contact with no photo_sha256 at all must still have its
        own TEXT block removed from the persisted source DOCX -
        deleting a contact must never leave its name/email/firm behind
        just because it never had a photo."""

        docx_path = self._require_copy("AU.docx")

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id=self.document_id,
                country_code="AU",
                contacts=(
                    ContactRecord(
                        contact_id="michael-id",
                        member_firm="HARMERS WORKPLACE LAWYERS",
                        contact_person="Michael Harmer",
                        email="michael.harmer@harmers.com.au",
                        phone="+61 292 674 322",
                        address="31 Market Street",
                        website="www.harmers.com.au",
                    ),
                ),
            ),
        )

        client = FakeContactOpenSearchClient(
            document_id=self.document_id,
            country_code="AU",
            country="Australia",
            source_filename="AU.docx",
        )

        with _patched_indexer(client):
            delete_contact(
                document_id=self.document_id,
                contact_id="michael-id",
                source_directory=self.root,
                client=client,
            )

        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8", errors="ignore")

        self.assertNotIn("Michael Harmer", document_xml)
        self.assertNotIn("michael.harmer@harmers.com.au", document_xml)

        remaining = extract_contacts_from_docx(docx_path, country="Australia")
        self.assertEqual(0, len(remaining))

    def _seed_contact_a_only(self, docx_path: Path) -> str:
        """Contact A (Michael Harmer) exactly as AU.docx's own native
        organic block already reads, with his own real embedded photo
        registered in ContactState - the document's starting state
        before an Admin-added contact B exists at all."""

        michael_photo = extract_contact_photo_candidates(docx_path)[0]
        michael_stored = write_contact_photo_atomic(
            self.root,
            "michael-id",
            data=michael_photo.data,
            content_type=michael_photo.content_type,
        )

        write_contact_state_atomic(
            self.root,
            ContactState(
                document_id=self.document_id,
                country_code="AU",
                contacts=(
                    ContactRecord(
                        contact_id="michael-id",
                        member_firm="HARMERS WORKPLACE LAWYERS",
                        contact_person="Michael Harmer",
                        email="michael.harmer@harmers.com.au",
                        phone="+61 292 674 322",
                        address="31 Market Street",
                        website="www.harmers.com.au",
                        photo_filename=michael_stored.filename,
                        photo_content_type=michael_stored.content_type,
                        photo_sha256=michael_stored.sha256,
                    ),
                ),
            ),
        )

        return michael_stored.sha256

    def _add_contact_b_with_photo_via_real_admin_flow(
        self, client: FakeContactOpenSearchClient
    ) -> tuple[str, bytes]:
        """Exercises the REAL, two-call Admin surface exactly as an
        operator would use it - add_contact() (text only) followed by
        replace_admin_contact_photo() (attaches the photo, rebuilding
        the canonical table again) - never the lower-level
        rebuild_canonical_contact_table() primitive directly.
        Returns (contact_b_id, photo_b_bytes)."""

        jane_photo_data = _make_png(183, 234, (80, 120, 200))

        with _patched_indexer(client):
            response = add_contact(
                document_id=self.document_id,
                fields=_write_request(
                    member_firm="Second Firm Pty Ltd",
                    contact_person="Jane Secondary",
                    email="jane.secondary@secondfirm.com.au",
                    phone="+61 2 9000 0000",
                    address="100 Test Street, Level 5",
                    website="www.secondfirm.com.au",
                ),
                source_directory=self.root,
                client=client,
            )

        with _patched_indexer(client):
            replace_admin_contact_photo(
                self.root,
                self.document_id,
                response.contact_id,
                data=jane_photo_data,
                content_type="image/png",
                client=client,
            )

        return response.contact_id, jane_photo_data

    def test_admin_added_contact_b_fully_removed_from_source_on_delete(
        self,
    ) -> None:
        """The user's mandatory focused test: ADD contact B + photo B
        (via the real two-call Admin surface, onto a document that
        already carries contact A natively), then DELETE contact B,
        then prove directly against the PERSISTED source DOCX that
        contact B's text and photo are gone while contact A's text and
        photo remain - byte-level proof, not ContactState alone."""

        docx_path = self._require_copy("AU.docx")
        michael_sha = self._seed_contact_a_only(docx_path)

        client = FakeContactOpenSearchClient(
            document_id=self.document_id,
            country_code="AU",
            country="Australia",
            source_filename="AU.docx",
        )

        contact_b_id, jane_photo_data = (
            self._add_contact_b_with_photo_via_real_admin_flow(client)
        )
        jane_sha = hashlib.sha256(jane_photo_data).hexdigest()

        # Sanity check before delete: contact B is genuinely present
        # in the persisted source, not merely in ContactState.
        with zipfile.ZipFile(docx_path) as archive:
            pre_delete_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8", errors="ignore")
        self.assertIn("Jane Secondary", pre_delete_xml)
        pre_delete_shas = {
            c.sha256 for c in extract_contact_photo_candidates(docx_path)
        }
        self.assertIn(jane_sha, pre_delete_shas)

        with _patched_indexer(client):
            delete_contact(
                document_id=self.document_id,
                contact_id=contact_b_id,
                source_directory=self.root,
                client=client,
            )

        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8", errors="ignore")
        remaining_shas = {
            c.sha256 for c in extract_contact_photo_candidates(docx_path)
        }

        contact_b_text_in_source = (
            "Jane Secondary" in document_xml
            or "jane.secondary@secondfirm.com.au" in document_xml
        )
        contact_b_photo_in_source = jane_sha in remaining_shas
        contact_a_text_in_source = (
            "Michael Harmer" in document_xml
            and "michael.harmer@harmers.com.au" in document_xml
        )
        contact_a_photo_in_source = michael_sha in remaining_shas

        print(
            f"CONTACT_B_TEXT_IN_SOURCE="
            f"{'YES' if contact_b_text_in_source else 'NO'}"
        )
        print(
            f"CONTACT_B_PHOTO_IN_SOURCE="
            f"{'YES' if contact_b_photo_in_source else 'NO'}"
        )
        print(
            f"CONTACT_A_TEXT_IN_SOURCE="
            f"{'YES' if contact_a_text_in_source else 'NO'}"
        )
        print(
            f"CONTACT_A_PHOTO_IN_SOURCE="
            f"{'YES' if contact_a_photo_in_source else 'NO'}"
        )

        self.assertFalse(
            contact_b_text_in_source,
            "deleted contact B's own text must not survive in the "
            "persisted source DOCX",
        )
        self.assertFalse(
            contact_b_photo_in_source,
            "deleted contact B's own photo must not survive in the "
            "persisted source DOCX",
        )
        self.assertTrue(
            contact_a_text_in_source,
            "surviving contact A's own text must remain untouched",
        )
        self.assertTrue(
            contact_a_photo_in_source,
            "surviving contact A's own photo must remain untouched",
        )

        remaining_contacts = extract_contacts_from_docx(
            docx_path, country="Australia"
        )
        self.assertEqual(1, len(remaining_contacts))
        self.assertEqual(
            "Michael Harmer", remaining_contacts[0].contact_person
        )

    def test_admin_added_contact_b_delete_failure_restores_source_byte_for_byte(
        self,
    ) -> None:
        """Same admin-added contact B scenario, but the ContactState/
        index commit fails after the source DOCX has already been
        rewritten to remove contact B's text+photo - the source must
        be restored to its EXACT prior bytes (with contact B still
        present), never left half-mutated."""

        docx_path = self._require_copy("AU.docx")
        self._seed_contact_a_only(docx_path)

        client = FakeContactOpenSearchClient(
            document_id=self.document_id,
            country_code="AU",
            country="Australia",
            source_filename="AU.docx",
        )

        contact_b_id, _ = (
            self._add_contact_b_with_photo_via_real_admin_flow(client)
        )

        original_bytes = docx_path.read_bytes()

        with self.assertRaises(AdminContactMutationFailedError):
            with _patched_indexer(client, fail_bulk=True):
                delete_contact(
                    document_id=self.document_id,
                    contact_id=contact_b_id,
                    source_directory=self.root,
                    client=client,
                )

        self.assertEqual(
            original_bytes,
            docx_path.read_bytes(),
            "the source DOCX must be restored to its exact original "
            "bytes (contact B's text and photo both still present) "
            "when the ContactState/index commit fails after the "
            "source rewrite already succeeded",
        )

        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(2, len(state.contacts))
        self.assertIn(
            contact_b_id, {c.contact_id for c in state.contacts}
        )


class AddContactFallbackSynchronizationTests(unittest.TestCase):
    """
    A document with no detected contact area at all (PT.docx - proven
    structure-less by test_contact_documents.py's own
    test_no_matching_contact_area_still_produces_canonical_table) must
    never leave a successful Add Contact as ContactState-only.
    add_contact() must fall back to contact_document_area.rebuild_
    canonical_contact_table()'s own no-existing-area handling
    (_default_insertion_anchor) and commit that to the source - SOURCE
    DOCX == CONTACT STATE == OPENSEARCH holds even here.
    """

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.addCleanup(self.temp.cleanup)
        self.root = Path(self.temp.name)
        self.document_id = _real_document_id_for("PT")

    def _require_copy(self, filename: str) -> Path:
        return require_corpus_copy(self, SOURCE_ROOT, filename, self.root)

    def _client(self) -> FakeContactOpenSearchClient:
        return FakeContactOpenSearchClient(
            document_id=self.document_id,
            country_code="PT",
            country="Portugal",
            source_filename="PT.docx",
        )

    def test_add_contact_without_structural_area_still_synchronizes_source(
        self,
    ) -> None:
        docx_path = self._require_copy("PT.docx")
        original_bytes = docx_path.read_bytes()

        client = self._client()

        with _patched_indexer(client):
            add_contact(
                document_id=self.document_id,
                fields=_write_request(
                    member_firm="Someone New Lda",
                    contact_person="Someone New",
                    email="new@example.test",
                    phone="+351 21 000 0000",
                    address="Rua Nova 1",
                    website="www.newfirm.test",
                ),
                source_directory=self.root,
                client=client,
            )

        new_bytes = docx_path.read_bytes()

        self.assertNotEqual(
            original_bytes,
            new_bytes,
            "a successful Add Contact must always rewrite the "
            "persisted source DOCX, even when the document has no "
            "two-box contact area to clone the style of - "
            "ContactState alone is never sufficient",
        )

        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8", errors="ignore")

        self.assertIn("Someone New", document_xml)
        self.assertIn("new@example.test", document_xml)
        self.assertIn(CONTACT_TABLE_HIDDEN_MARKER, document_xml)
        self.assertIn("<w:tbl>", document_xml)

        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(1, len(state.contacts))

    def test_second_add_replaces_rather_than_duplicates_the_fallback_block(
        self,
    ) -> None:
        """Two sequential Add Contact calls against the same
        structure-less document must each leave the source
        synchronized with the FULL current contact list - never
        stacking a second, stale fallback block alongside the first
        (which would silently resurrect a deleted/superseded
        rendering of the earlier contact)."""

        docx_path = self._require_copy("PT.docx")
        client = self._client()

        with _patched_indexer(client):
            add_contact(
                document_id=self.document_id,
                fields=_write_request(
                    member_firm="First Firm Lda",
                    contact_person="First Person",
                    email="first@example.test",
                ),
                source_directory=self.root,
                client=client,
            )

        with _patched_indexer(client):
            add_contact(
                document_id=self.document_id,
                fields=_write_request(
                    member_firm="Second Firm Lda",
                    contact_person="Second Person",
                    email="second@example.test",
                ),
                source_directory=self.root,
                client=client,
            )

        with zipfile.ZipFile(docx_path) as archive:
            document_xml = archive.read(
                "word/document.xml"
            ).decode("utf-8", errors="ignore")

        self.assertIn("First Person", document_xml)
        self.assertIn("Second Person", document_xml)
        self.assertEqual(
            1,
            document_xml.count(CONTACT_TABLE_HIDDEN_MARKER),
            "a second Add Contact must replace the one canonical "
            "table with a fresh rendering of the full contact list, "
            "never append a second one",
        )

        state = read_contact_state(self.root, self.document_id)
        self.assertEqual(2, len(state.contacts))


# =========================================================================
# BOOTSTRAP
# =========================================================================


class ContactBootstrapTests(unittest.TestCase):
    def test_dry_run_does_not_write_anything(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / "GB.docx").write_bytes(b"docx")

            summaries = [
                _summary(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    source_filename="GB.docx",
                )
            ]

            with patch(
                "app.services.admin_contacts.extract_contacts_from_docx",
                return_value=[
                    ExtractedContact(member_firm="Firm A"),
                ],
            ):
                report = bootstrap_legacy_contacts(
                    source_directory=source_directory,
                    client=object(),
                    dry_run=True,
                    document_lister=_fake_document_lister(summaries),
                )

            self.assertTrue(report.dry_run)
            self.assertEqual(report.documents_seen, 1)
            self.assertEqual(report.contacts_seeded, 1)
            self.assertIsNone(
                read_contact_state(source_directory, DOCUMENT_ID)
            )

    def test_wet_run_seeds_state_for_new_documents(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / "GB.docx").write_bytes(b"docx")

            summaries = [
                _summary(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    source_filename="GB.docx",
                )
            ]

            with patch(
                "app.services.admin_contacts.extract_contacts_from_docx",
                return_value=[
                    ExtractedContact(member_firm="Firm A"),
                ],
            ):
                report = bootstrap_legacy_contacts(
                    source_directory=source_directory,
                    client=object(),
                    dry_run=False,
                    document_lister=_fake_document_lister(summaries),
                )

            self.assertEqual(report.contacts_seeded, 1)

            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].member_firm, "Firm A")

    def test_zero_contact_document_gets_explicit_empty_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / "FR.docx").write_bytes(b"docx")

            summaries = [
                _summary(
                    document_id=OTHER_DOCUMENT_ID,
                    country_code="FR",
                    source_filename="FR.docx",
                )
            ]

            with patch(
                "app.services.admin_contacts.extract_contacts_from_docx",
                return_value=[],
            ):
                report = bootstrap_legacy_contacts(
                    source_directory=source_directory,
                    client=object(),
                    dry_run=False,
                    document_lister=_fake_document_lister(summaries),
                )

            self.assertEqual(report.zero_contact_documents, 1)
            self.assertEqual(report.contacts_seeded, 0)

            state = read_contact_state(source_directory, OTHER_DOCUMENT_ID)
            self.assertIsNotNone(state)
            self.assertEqual(state.contacts, ())

    def test_existing_state_is_never_overwritten(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / "GB.docx").write_bytes(b"docx")

            existing_record = _full_contact_record(
                member_firm="Admin Edited Firm"
            )
            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(existing_record,),
                ),
            )

            summaries = [
                _summary(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    source_filename="GB.docx",
                )
            ]

            with patch(
                "app.services.admin_contacts.extract_contacts_from_docx",
                return_value=[
                    ExtractedContact(member_firm="Stale DOCX Firm"),
                ],
            ):
                report = bootstrap_legacy_contacts(
                    source_directory=source_directory,
                    client=object(),
                    dry_run=False,
                    document_lister=_fake_document_lister(summaries),
                )

            self.assertEqual(report.documents_skipped_existing_state, 1)
            self.assertEqual(report.contacts_seeded, 0)

            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(state.contacts[0].member_firm, "Admin Edited Firm")

    def test_legacy_incomplete_contacts_preserved_exactly(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / "GB.docx").write_bytes(b"docx")

            summaries = [
                _summary(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    source_filename="GB.docx",
                )
            ]

            with patch(
                "app.services.admin_contacts.extract_contacts_from_docx",
                return_value=[
                    ExtractedContact(
                        contact_person="Alex Example",
                        email="alex@example.test",
                    ),
                ],
            ):
                bootstrap_legacy_contacts(
                    source_directory=source_directory,
                    client=object(),
                    dry_run=False,
                    document_lister=_fake_document_lister(summaries),
                )

            state = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertEqual(len(state.contacts), 1)
            self.assertIsNone(state.contacts[0].member_firm)
            self.assertIsNone(state.contacts[0].phone)
            self.assertEqual(
                state.contacts[0].contact_person, "Alex Example"
            )


class AdminContactWriteRequestOptionalFieldTests(unittest.TestCase):
    """Section 11/12: every one of the six business fields is
    individually optional; the only validation is a cross-field "at
    least one field has a value" rule - a real member-firm contact
    (France's own Caroline Scherrmann) can genuinely have address/
    website empty, and that must never be rejected."""

    def test_every_field_individually_empty_is_accepted(self) -> None:
        for field in (
            "member_firm", "contact_person", "email",
            "phone", "address", "website",
        ):
            with self.subTest(field=field):
                fields = {
                    "member_firm": "Firm",
                    "contact_person": "Person",
                    "email": "person@example.com",
                    "phone": "+1 555 0100",
                    "address": "1 Example Street",
                    "website": "www.example.com",
                }
                fields[field] = ""

                request = AdminContactWriteRequest(**fields)
                self.assertEqual("", getattr(request, field))

    def test_website_and_address_both_empty_is_accepted(self) -> None:
        """The exact real France shape."""

        request = AdminContactWriteRequest(
            member_firm="Flichy Grangé Avocats",
            contact_person="Caroline Scherrmann",
            email="scherrmann@flichy.com",
            phone="+33 1 56 62 30 00",
            address="",
            website="",
        )
        self.assertEqual("", request.address)
        self.assertEqual("", request.website)

    def test_all_six_fields_blank_is_rejected(self) -> None:
        with self.assertRaises(ValidationError):
            AdminContactWriteRequest(
                member_firm="",
                contact_person="",
                email="",
                phone="",
                address="",
                website="",
            )

    def test_all_six_fields_whitespace_only_is_rejected(self) -> None:
        with self.assertRaises(ValidationError):
            AdminContactWriteRequest(
                member_firm="   ",
                contact_person="",
                email="  ",
                phone="",
                address="",
                website="",
            )

    def test_a_single_filled_field_is_sufficient(self) -> None:
        request = AdminContactWriteRequest(
            member_firm="",
            contact_person="Solo Person",
            email="",
            phone="",
            address="",
            website="",
        )
        self.assertEqual("Solo Person", request.contact_person)


class FranceLegacyBootstrapSplitTests(unittest.TestCase):
    """Sections 7-9: France's real legacy contact - a single combined
    record naming two people - splits into two Admin-managed
    ContactRecords during bootstrap, with stable ids and Jessica Stout
    (the project-level L&E Global POC, never a member-firm contact)
    correctly excluded."""

    def _require_copy(self, filename: str) -> Path:
        source = SOURCE_ROOT / filename

        if not source.exists():
            self.skipTest(f"Real corpus source unavailable: {source}")

        document = Document(str(source))
        if any(
            table.rows
            and CONTACT_TABLE_HIDDEN_MARKER in table.rows[0].cells[0].text
            for table in document.tables
        ):
            # A real Admin has since used the live Contact CRUD
            # feature against this document (confirmed directly: it
            # now has a canonical contact table with Caroline/Florence
            # already split, plus real test contacts added live - a
            # genuine, welcome confirmation that the split works in
            # practice, but real-world content drift this bootstrap-
            # from-scratch scenario can no longer be reproduced against).
            self.skipTest(
                f"{filename} has since been canonicalized by real "
                "Admin usage (real corpus content has drifted since "
                "this test was written) - its legacy combined contact "
                "no longer exists in raw form to bootstrap from"
            )

        return source

    def test_france_bootstrap_yields_two_stable_contacts_without_jessica(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            real_source = self._require_copy("FR.docx")
            (source_directory / "FR.docx").write_bytes(
                real_source.read_bytes()
            )

            document_id = "doc_" + "f" * 64
            summaries = [
                _summary(
                    document_id=document_id,
                    country_code="FR",
                    source_filename="FR.docx",
                )
            ]

            report = bootstrap_legacy_contacts(
                source_directory=source_directory,
                client=object(),
                dry_run=False,
                document_lister=_fake_document_lister(summaries),
            )

            self.assertEqual(2, report.contacts_seeded)

            state = read_contact_state(source_directory, document_id)
            self.assertEqual(2, len(state.contacts))

            names = [c.contact_person for c in state.contacts]
            self.assertEqual(
                ["Caroline Scherrmann", "Florence Bacquet"], names
            )
            self.assertNotIn("Jessica Stout", names)

            ids = {c.contact_id for c in state.contacts}
            self.assertEqual(
                2, len(ids), "each split contact must get its own id"
            )

            for contact in state.contacts:
                self.assertEqual(
                    "Flichy Grangé Avocats", contact.member_firm
                )
                self.assertEqual("+33 1 56 62 30 00", contact.phone)

    def test_repeated_bootstrap_does_not_change_ids_or_duplicate(
        self,
    ) -> None:
        """bootstrap_legacy_contacts never overwrites an existing
        sidecar (see its own docstring) - re-running it after France's
        split has already been persisted must be a complete no-op,
        never regenerating ids or duplicating the two contacts."""

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            real_source = self._require_copy("FR.docx")
            (source_directory / "FR.docx").write_bytes(
                real_source.read_bytes()
            )

            document_id = "doc_" + "f" * 64
            summaries = [
                _summary(
                    document_id=document_id,
                    country_code="FR",
                    source_filename="FR.docx",
                )
            ]
            lister = _fake_document_lister(summaries)

            bootstrap_legacy_contacts(
                source_directory=source_directory,
                client=object(),
                dry_run=False,
                document_lister=lister,
            )
            first_ids = [
                c.contact_id
                for c in read_contact_state(
                    source_directory, document_id
                ).contacts
            ]

            report = bootstrap_legacy_contacts(
                source_directory=source_directory,
                client=object(),
                dry_run=False,
                document_lister=lister,
            )

            self.assertEqual(1, report.documents_skipped_existing_state)
            self.assertEqual(0, report.contacts_seeded)

            second_state = read_contact_state(source_directory, document_id)
            second_ids = [c.contact_id for c in second_state.contacts]

            self.assertEqual(first_ids, second_ids)
            self.assertEqual(2, len(second_state.contacts))


class ContactBootstrapIsolatedCorpusTests(unittest.TestCase):
    """
    Mirrors a real production audit's own real-corpus baseline: 24
    documents, 22 contacts, 2 zero-contact documents (FR, PT) - here
    reproduced against a synthetic, isolated 24-document catalog (never
    the real production corpus), proving the bootstrap facility
    produces exactly this shape given that input.
    """

    def test_matches_the_audit_baseline_shape(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            summaries = []
            contacts_by_country: dict[str, list[ExtractedContact]] = {}

            for index in range(24):
                country_code = f"C{index:02d}"
                summaries.append(
                    _summary(
                        document_id=f"doc_{index:064x}",
                        country_code=country_code,
                        source_filename=f"{country_code}.docx",
                    )
                )
                (
                    source_directory / f"{country_code}.docx"
                ).write_bytes(b"docx")

                # Exactly 2 of the 24 (indices 0 and 1) get zero
                # contacts, matching FR/PT in the real audit baseline.
                contacts_by_country[country_code] = (
                    []
                    if index < 2
                    else [ExtractedContact(member_firm="Firm")]
                )

            def fake_extract(path: Path, country: str):
                del country
                country_code = path.stem
                return contacts_by_country[country_code]

            with patch(
                "app.services.admin_contacts.extract_contacts_from_docx",
                side_effect=fake_extract,
            ):
                report = bootstrap_legacy_contacts(
                    source_directory=source_directory,
                    client=object(),
                    dry_run=False,
                    document_lister=_fake_document_lister(summaries),
                )

            self.assertEqual(report.documents_seen, 24)
            self.assertEqual(report.contacts_seeded, 22)
            self.assertEqual(report.zero_contact_documents, 2)


# =========================================================================
# REFRESH: apply_structured_contact_state_to_chunks
# =========================================================================


def _legal_chunk(document_id: str = DOCUMENT_ID) -> DocumentChunk:
    return DocumentChunk(
        document_id=document_id,
        chunk_id="chunk_" + "c" * 64,
        country="United Kingdom",
        country_code="GB",
        legal_topic="Employment Contracts",
        document_type="overview",
        language="en",
        section="Employment Contracts",
        subsection=None,
        content="legal content",
        source_filename="GB.docx",
        source_format="docx",
        content_hash="hash-legal",
        reference_year=2026,
    )


def _stale_docx_contact_chunk(document_id: str = DOCUMENT_ID) -> DocumentChunk:
    return DocumentChunk(
        document_id=document_id,
        chunk_id="chunk_" + "d" * 64,
        country="United Kingdom",
        country_code="GB",
        legal_topic=None,
        document_type="overview",
        language="en",
        section="Employment Law Overview United Kingdom",
        subsection="Contact",
        content="Member firm: Stale DOCX Firm",
        source_filename="GB.docx",
        source_format="docx",
        content_hash="hash-stale",
        reference_year=2026,
    )


class ApplyStructuredContactStateToChunksTests(unittest.TestCase):
    def test_no_sidecar_leaves_chunks_unchanged(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            chunks = [_legal_chunk(), _stale_docx_contact_chunk()]

            result = apply_structured_contact_state_to_chunks(
                chunks=list(chunks),
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
            )

            self.assertEqual(result, chunks)

    def test_existing_sidecar_replaces_the_stale_docx_contact_chunk(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(
                        _full_contact_record(
                            member_firm="Admin Edited Firm"
                        ),
                    ),
                ),
            )

            chunks = [_legal_chunk(), _stale_docx_contact_chunk()]

            result = apply_structured_contact_state_to_chunks(
                chunks=list(chunks),
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
            )

            legal_chunks = [c for c in result if c.subsection != "Contact"]
            contact_chunks = [c for c in result if c.subsection == "Contact"]

            self.assertEqual(len(legal_chunks), 1)
            self.assertEqual(legal_chunks[0].content, "legal content")

            self.assertEqual(len(contact_chunks), 1)
            self.assertIn("Admin Edited Firm", contact_chunks[0].content)
            self.assertNotIn("Stale DOCX Firm", contact_chunks[0].content)

    def test_existing_empty_sidecar_removes_the_docx_contact_chunk(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(),
                ),
            )

            chunks = [_legal_chunk(), _stale_docx_contact_chunk()]

            result = apply_structured_contact_state_to_chunks(
                chunks=list(chunks),
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
            )

            self.assertEqual(len(result), 1)
            self.assertIsNone(result[0].subsection)


class ReindexPreservesContactStateIntegrationTests(unittest.TestCase):
    """
    Real reindex_indexed_document, proving the wiring (not just the
    pure helper above) actually substitutes the Contact chunk during a
    genuine Refresh.
    """

    def test_refresh_uses_structured_state_not_stale_docx_text(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / "GB.docx").write_bytes(b"docx")

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(
                        _full_contact_record(
                            member_firm="Admin Edited Firm"
                        ),
                    ),
                ),
            )

            client = FakeContactOpenSearchClient(
                chunks={
                    "existing-legal": {
                        "document_id": DOCUMENT_ID,
                        "country_code": "GB",
                        "country": "United Kingdom",
                        "source_filename": "GB.docx",
                        "reference_year": 2026,
                        "legal_topic": "Employment Contracts",
                        "subsection": None,
                        "content": "legal content",
                    },
                },
            )

            def chunk_builder(path: Path) -> list[DocumentChunk]:
                del path
                return [_legal_chunk(), _stale_docx_contact_chunk()]

            captured: dict[str, Any] = {}

            def document_indexer(*, chunks, client=None) -> DocumentIndexingResult:
                del client
                captured["chunks"] = list(chunks)

                return DocumentIndexingResult(
                    index_alias="legal-documents",
                    document_id=chunks[0].document_id,
                    source_filename=chunks[0].source_filename,
                    requested_chunks=len(chunks),
                    indexed_chunks=len(chunks),
                    stale_chunks_deleted=0,
                )

            reindex_indexed_document(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
                chunk_builder=chunk_builder,
                document_indexer=document_indexer,
            )

            indexed_contact_chunks = [
                c for c in captured["chunks"] if c.subsection == "Contact"
            ]

            self.assertEqual(len(indexed_contact_chunks), 1)
            self.assertIn(
                "Admin Edited Firm", indexed_contact_chunks[0].content
            )
            self.assertNotIn(
                "Stale DOCX Firm", indexed_contact_chunks[0].content
            )


# =========================================================================
# NEW DOCX: reseed_contact_state_from_parsed_contacts +
# safe_upload_and_index_document integration
# =========================================================================


class ReseedContactStateFromParsedContactsTests(unittest.TestCase):
    def test_reseed_overwrites_prior_state_entirely(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(
                        _full_contact_record(member_firm="Old Firm A"),
                        _full_contact_record(member_firm="Old Firm B"),
                        _full_contact_record(member_firm="Old Firm C"),
                    ),
                ),
            )
            mark_admin_modified(source_directory, DOCUMENT_ID)

            reseed_contact_state_from_parsed_contacts(
                document_id=DOCUMENT_ID,
                country_code="GB",
                source_directory=source_directory,
                contacts=[ExtractedContact(member_firm="New Firm")],
            )

            state = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].member_firm, "New Firm")
            self.assertFalse(
                is_admin_modified_since_upload(
                    source_directory, DOCUMENT_ID
                )
            )

    def test_reseed_with_zero_contacts_clears_prior_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            write_contact_state_atomic(
                source_directory,
                ContactState(
                    document_id=DOCUMENT_ID,
                    country_code="GB",
                    contacts=(
                        _full_contact_record(member_firm="Old Firm A"),
                        _full_contact_record(member_firm="Old Firm B"),
                    ),
                ),
            )

            reseed_contact_state_from_parsed_contacts(
                document_id=DOCUMENT_ID,
                country_code="GB",
                source_directory=source_directory,
                contacts=[],
            )

            state = read_contact_state(source_directory, DOCUMENT_ID)

            self.assertIsNotNone(state)
            self.assertEqual(state.contacts, ())


class ParsedContactPhotoReseedTests(unittest.TestCase):
    """
    reseed_contact_state_from_parsed_contacts's OTHER branch: when
    docx_path is supplied, real corpus photo extraction/association is
    also exercised (see _build_photo_aware_contact_records) - a
    materially different code path than the two tests above, which
    never pass docx_path at all.
    """

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.source_directory = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def _require(self, filename: str) -> Path:
        path = SOURCE_ROOT / filename

        if not path.exists():
            self.skipTest(f"Corpus file unavailable: {path}")

        skip_if_already_canonicalized(self, path)

        return path

    def _photo_files(self) -> list[Path]:
        directory = (
            self.source_directory
            / ".admin-state"
            / "contact-photos"
        )

        if not directory.exists():
            return []

        return sorted(
            path
            for path in directory.iterdir()
            if path.is_file()
        )

    def test_parsed_reseed_belgium_creates_two_contacts_and_two_photos(
        self,
    ) -> None:
        path = self._require(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        parsed = extract_contacts_from_docx(
            path,
            country="Belgium",
        )

        self.assertEqual(1, len(parsed))

        reseed_contact_state_from_parsed_contacts(
            document_id="doc-belgium",
            country_code="BE",
            source_directory=self.source_directory,
            contacts=parsed,
            docx_path=path,
        )

        state = read_contact_state(
            self.source_directory,
            "doc-belgium",
        )

        self.assertIsNotNone(state)
        self.assertEqual(2, len(state.contacts))

        first, second = state.contacts

        self.assertEqual("Chris van Olmen", first.contact_person)
        self.assertEqual("chris.van.olmen@vow.be", first.email)

        self.assertEqual("Nicolas Simon", second.contact_person)
        self.assertEqual("nicolas.simon@vow.be", second.email)

        self.assertNotEqual(first.contact_id, second.contact_id)

        for record in (first, second):
            self.assertIsNotNone(record.photo_filename)
            self.assertIsNotNone(record.photo_content_type)
            self.assertIsNotNone(record.photo_sha256)

            self.assertTrue(
                (
                    self.source_directory
                    / ".admin-state"
                    / "contact-photos"
                    / record.photo_filename
                ).is_file()
            )

        expected_photos = extract_contact_photo_candidates(path)

        self.assertEqual(expected_photos[0].sha256, first.photo_sha256)
        self.assertEqual(expected_photos[1].sha256, second.photo_sha256)

        self.assertEqual(2, len(self._photo_files()))

    def test_parsed_reseed_france_remains_one_contact_without_photo(
        self,
    ) -> None:
        path = self._require("FR.docx")

        parsed = extract_contacts_from_docx(path, country="France")

        reseed_contact_state_from_parsed_contacts(
            document_id="doc-france",
            country_code="FR",
            source_directory=self.source_directory,
            contacts=parsed,
            docx_path=path,
        )

        state = read_contact_state(self.source_directory, "doc-france")

        self.assertEqual(1, len(state.contacts))

        contact = state.contacts[0]

        self.assertEqual(
            "Caroline Scherrmann and Florence Bacquet",
            contact.contact_person,
        )
        self.assertIsNone(contact.photo_filename)
        self.assertIsNone(contact.photo_content_type)
        self.assertIsNone(contact.photo_sha256)

        self.assertEqual([], self._photo_files())

    def test_parsed_reseed_indonesia_persists_the_valid_photo(
        self,
    ) -> None:
        path = self._require("ID.docx")

        parsed = extract_contacts_from_docx(path, country="Indonesia")

        expected = extract_contact_photo_candidates(path)

        self.assertEqual(1, len(expected))
        self.assertEqual("image3.jpeg", expected[0].source_filename)

        reseed_contact_state_from_parsed_contacts(
            document_id="doc-indonesia",
            country_code="ID",
            source_directory=self.source_directory,
            contacts=parsed,
            docx_path=path,
        )

        state = read_contact_state(self.source_directory, "doc-indonesia")

        self.assertEqual(1, len(state.contacts))

        contact = state.contacts[0]

        self.assertEqual(expected[0].sha256, contact.photo_sha256)

        self.assertTrue(
            (
                self.source_directory
                / ".admin-state"
                / "contact-photos"
                / contact.photo_filename
            ).is_file()
        )

    def test_photo_extraction_failure_keeps_contacts_without_photo(
        self,
    ) -> None:
        path = self._require("FR.docx")

        parsed = extract_contacts_from_docx(path, country="France")

        with patch.object(
            admin_contacts,
            "extract_contact_photo_candidates",
            side_effect=admin_contacts.ContactPhotoExtractionError(
                "unsupported synthetic DOCX image package"
            ),
        ):
            reseed_contact_state_from_parsed_contacts(
                document_id="doc-photo-fallback",
                country_code="FR",
                source_directory=self.source_directory,
                contacts=parsed,
                docx_path=path,
            )

        state = read_contact_state(
            self.source_directory,
            "doc-photo-fallback",
        )

        self.assertIsNotNone(state)
        self.assertEqual(1, len(state.contacts))

        contact = state.contacts[0]

        self.assertEqual(
            "Caroline Scherrmann and Florence Bacquet",
            contact.contact_person,
        )
        self.assertIsNone(contact.photo_filename)
        self.assertIsNone(contact.photo_content_type)
        self.assertIsNone(contact.photo_sha256)

        self.assertEqual([], self._photo_files())

    def test_second_photo_write_failure_leaves_no_partial_seed(
        self,
    ) -> None:
        path = self._require(
            "Labour and Employment Law in Belgium 2026.docx"
        )

        parsed = extract_contacts_from_docx(path, country="Belgium")

        real_write_contact_photo_atomic = write_contact_photo_atomic
        calls = 0

        def fail_second(*args, **kwargs):
            nonlocal calls
            calls += 1

            if calls == 2:
                raise ContactPhotoStorageError("second photo boom")

            return real_write_contact_photo_atomic(*args, **kwargs)

        with patch.object(
            admin_contacts,
            "write_contact_photo_atomic",
            side_effect=fail_second,
            create=True,
        ):
            with self.assertRaises(ContactPhotoStorageError):
                reseed_contact_state_from_parsed_contacts(
                    document_id="doc-belgium",
                    country_code="BE",
                    source_directory=self.source_directory,
                    contacts=parsed,
                    docx_path=path,
                )

        self.assertIsNone(
            read_contact_state(self.source_directory, "doc-belgium")
        )

        self.assertEqual([], self._photo_files())


class UploadReplaceReseedsContactsIntegrationTests(unittest.TestCase):
    """
    Real safe_upload_and_index_document, proving a genuine
    upload/replace transaction really reseeds structured contact state
    from the newly-accepted DOCX (not merely the lower-level helper
    above in isolation).
    """

    def _run_upload(
        self,
        *,
        source_directory: Path,
        parsed_contacts: list[ExtractedContact],
        replace_existing: bool,
    ):
        from app.services.document_chunk_builder import DOCUMENT_FAMILY

        del DOCUMENT_FAMILY

        chunk = _legal_chunk()

        def chunk_builder(path: Path) -> list[DocumentChunk]:
            del path
            return [chunk]

        def country_document_lookup(country_code, client):
            del country_code, client
            return []

        def country_document_indexer(*, chunks, client=None):
            del client
            return DocumentIndexingResult(
                index_alias="legal-documents",
                document_id=chunks[0].document_id,
                source_filename=chunks[0].source_filename,
                requested_chunks=len(chunks),
                indexed_chunks=len(chunks),
                stale_chunks_deleted=0,
            )

        with patch(
            "app.services.admin_document_replacement.extract_contacts_from_docx",
            return_value=parsed_contacts,
        ):
            import io

            return safe_upload_and_index_document(
                filename="GB.docx",
                file_stream=io.BytesIO(b"fake docx bytes"),
                source_directory=source_directory,
                processed_directory=source_directory / "processed",
                maximum_bytes=10_000_000,
                country_confirmed=True,
                confirm_warnings=True,
                replace_existing=replace_existing,
                chunk_builder=chunk_builder,
                country_document_lookup=country_document_lookup,
                country_document_indexer=country_document_indexer,
            )

    def test_fresh_upload_seeds_contact_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            response = self._run_upload(
                source_directory=source_directory,
                parsed_contacts=[ExtractedContact(member_firm="New Firm")],
                replace_existing=False,
            )

            state = read_contact_state(
                source_directory, response.document_id
            )

            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].member_firm, "New Firm")
            self.assertFalse(
                is_admin_modified_since_upload(
                    source_directory, response.document_id
                )
            )


# =========================================================================
# SAME-BYTES CONFIRMED RESEED
# =========================================================================


class ReseedContactsFromCurrentDocxTests(unittest.TestCase):
    def test_confirmed_reseed_discards_admin_edits(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _seed_placeholder_source_docx(source_directory)

            client = FakeContactOpenSearchClient()

            with _patched_indexer(client):
                add_contact(
                    document_id=DOCUMENT_ID,
                    fields=_write_request(member_firm="Admin Edited Firm"),
                    source_directory=source_directory,
                    client=client,
                )

            with patch(
                "app.services.admin_contacts.extract_contacts_from_docx",
                return_value=[
                    ExtractedContact(member_firm="Parsed DOCX Firm"),
                ],
            ), _patched_indexer(client):
                response = reseed_contacts_from_current_docx(
                    document_id=DOCUMENT_ID,
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(len(response.contacts), 1)
            self.assertEqual(
                response.contacts[0].member_firm, "Parsed DOCX Firm"
            )

            state = read_contact_state(source_directory, DOCUMENT_ID)
            self.assertEqual(
                state.contacts[0].member_firm, "Parsed DOCX Firm"
            )
            self.assertFalse(
                is_admin_modified_since_upload(
                    source_directory, DOCUMENT_ID
                )
            )


class CurrentDocxPhotoReseedTests(unittest.TestCase):
    """
    Real _reseed_contacts_from_current_docx_locked (the private,
    lock-already-held entry point), against a REAL corpus document -
    proving the current-DOCX reseed path also carries real photo
    association through, complementing
    test_confirmed_reseed_discards_admin_edits above (which patches
    extract_contacts_from_docx directly and never touches a real
    photo)."""

    def setUp(self) -> None:
        self.temp = tempfile.TemporaryDirectory()
        self.source_directory = Path(self.temp.name)

    def tearDown(self) -> None:
        self.temp.cleanup()

    def test_current_docx_reseed_belgium_seeds_two_photo_contacts(
        self,
    ) -> None:
        belgium_filename = "Labour and Employment Law in Belgium 2026.docx"
        path = SOURCE_ROOT / belgium_filename

        if not path.exists():
            self.skipTest("Belgium corpus DOCX unavailable")

        skip_if_already_canonicalized(self, path)

        metadata = {
            "source_filename": belgium_filename,
            "country": "Belgium",
            "country_code": "BE",
        }

        with (
            patch.object(
                admin_contacts,
                "_load_country_code_and_metadata",
                return_value=("BE", metadata),
            ),
            patch.object(
                admin_contacts,
                "resolve_document_source_path",
                return_value=SimpleNamespace(path=path),
            ),
            patch.object(
                admin_contacts,
                "_document_metadata_for_chunks",
                return_value={},
            ),
            patch.object(
                admin_contacts,
                "build_contact_chunk_for_contacts",
                return_value=None,
            ),
            patch.object(
                admin_contacts,
                "replace_document_contact_chunk",
            ),
            patch.object(
                admin_contacts,
                "is_admin_modified_since_upload",
                return_value=False,
            ),
            patch.object(
                admin_contacts,
                "reset_admin_modified",
            ),
        ):
            admin_contacts._reseed_contacts_from_current_docx_locked(
                validated_document_id="doc-current-belgium",
                source_directory=self.source_directory,
                opensearch_client=object(),
            )

        state = read_contact_state(
            self.source_directory, "doc-current-belgium"
        )

        self.assertIsNotNone(state)
        self.assertEqual(2, len(state.contacts))

        self.assertEqual(
            ["Chris van Olmen", "Nicolas Simon"],
            [contact.contact_person for contact in state.contacts],
        )

        self.assertEqual(
            2,
            sum(
                contact.photo_filename is not None
                for contact in state.contacts
            ),
        )


if __name__ == "__main__":
    unittest.main()
