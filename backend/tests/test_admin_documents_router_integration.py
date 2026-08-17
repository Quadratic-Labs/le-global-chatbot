"""
Router-level integration tests for the admin document HTTP contract
(mission "HOTFIX 0.4.9" review, sections 5/6/9).

These call the FastAPI route functions directly (upload_admin_document,
delete_admin_document) rather than through an ASGI TestClient - this
project has never used TestClient (see git history), and a plain
function call already exercises 100% of the route function's own
code, including every exception -> HTTPException mapping; only the
ASGI/HTTP transport itself is skipped, which is not what this suite
verifies. get_settings() and get_opensearch_client() are both
@lru_cache-decorated and read real environment variables / construct
a real OpenSearch client when uncached, so every test here clears
both caches before and after itself and points settings at a fresh
temp directory - never /data/documents/source, never a real cluster.
"""

from __future__ import annotations

import io
import os
import tempfile
import unittest
from pathlib import Path
from typing import Any
from unittest.mock import patch

from docx import Document
from fastapi import HTTPException, UploadFile

from app.core.config import get_settings
from app.routers import admin_document_lifecycle as lifecycle_router
from app.routers import admin_documents as documents_router
from app.services import admin_document_lifecycle as lifecycle_service
from app.services.contact_state import (
    ContactRecord,
    ContactState,
    write_contact_state_atomic,
)
from app.services.docx_parser import extract_contacts_from_docx
from tests.admin_invariants import (
    assert_chunk_count_matches,
    assert_no_orphan_chunks,
    assert_one_active_document_per_country,
    assert_one_active_source,
    assert_zero_mutation,
    real_source_entries,
)


_SETTINGS_ENV_KEYS = (
    "OPENSEARCH_URL",
    "OPENSEARCH_PASSWORD",
    "REDIS_URL",
    "DOCUMENT_SOURCE_DIR",
    "DOCUMENT_PROCESSED_DIR",
    "ADMIN_API_KEY",
    "API_ACCESS_KEY",
)


def _build_real_docx_bytes(country_line: str) -> bytes:
    document = Document()
    document.add_paragraph(country_line)
    document.add_paragraph("I. GENERAL OVERVIEW")
    document.add_paragraph("1. Introduction")
    document.add_paragraph("Overview content.")
    document.add_paragraph("II. Hiring Practices")
    document.add_paragraph("Hiring content.")
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


_AR_BYTES = _build_real_docx_bytes(
    "Labour and Employment Law in Argentina 2026"
)


class FakeOpenSearch:
    """
    In-memory OpenSearch double at chunk granularity - matches the
    real index's own granularity (one document has 1+ chunks, each a
    separate hit), correctly handling every query shape the admin
    upload/replace/delete/reindex code paths issue: term.document_id,
    term.country_code, and the bool/filter+must_not "keep these chunk
    ids" shape _delete_country_chunks uses.
    """

    def __init__(self) -> None:
        self.chunks: dict[str, dict[str, Any]] = {}

    def add(
        self,
        *,
        document_id: str,
        country_code: str,
        source_filename: str,
        chunk_id: str | None = None,
        language: str = "en",
        document_type: str = "overview",
    ) -> None:
        resolved_chunk_id = chunk_id or f"{document_id}-chunk-0"
        self.chunks[resolved_chunk_id] = {
            "document_id": document_id,
            "chunk_id": resolved_chunk_id,
            "country_code": country_code,
            "source_filename": source_filename,
            "country": "Argentina" if country_code == "AR" else "Other",
            "reference_year": 2026,
            "language": language,
            "document_type": document_type,
        }

    def document_ids_for_country(self, country_code: str) -> set[str]:
        return {
            c["document_id"]
            for c in self.chunks.values()
            if c["country_code"] == country_code
        }

    def search(self, *, index: str, body: dict[str, Any]) -> dict[str, Any]:
        del index

        if "query" not in body:
            # build_admin_document_catalog_body()'s aggregation shape
            # (list_indexed_documents/get_admin_document_stats) - one
            # bucket per distinct document_id, built from the real
            # in-memory chunks rather than a second, hardcoded fixture.
            by_document: dict[str, list[dict[str, Any]]] = {}

            for chunk in self.chunks.values():
                by_document.setdefault(
                    chunk["document_id"], []
                ).append(chunk)

            buckets = [
                {
                    "key": document_id,
                    "doc_count": len(chunks),
                    "metadata": {
                        "hits": {
                            "hits": [
                                {"_source": chunks[0]}
                            ]
                        }
                    },
                }
                for document_id, chunks in sorted(
                    by_document.items()
                )
            ]

            return {
                "aggregations": {
                    "documents": {"buckets": buckets}
                }
            }

        term = body["query"].get("term") or {}

        if "country_code" in term:
            # One hit per actual chunk, matching the real index's own
            # granularity - callers that want distinct documents
            # (e.g. lookup_existing_country_documents) already
            # deduplicate by document_id client-side; callers that
            # need the true per-document chunk count (e.g. mission
            # "HOTFIX 0.4.9" review 3's _snapshot_country_chunks-based
            # expected_chunks validation) require every chunk hit,
            # never one representative hit per document.
            hits = [
                c for c in self.chunks.values()
                if c["country_code"] == term["country_code"]
            ]
        elif "document_id" in term:
            hits = [
                c for c in self.chunks.values()
                if c["document_id"] == term["document_id"]
            ]
        else:
            hits = list(self.chunks.values())

        return {
            "hits": {
                "total": {"value": len(hits)},
                "hits": [
                    {
                        "_id": h["chunk_id"],
                        "_source": h,
                        # Mission "ORDER 3B": real OpenSearch 3.7
                        # includes a "sort" array per hit whenever the
                        # request itself carries a "sort" clause (as
                        # _fetch_all_chunks's search_after pagination
                        # always does) - never omitted here.
                        "sort": [h["chunk_id"]],
                    }
                    for h in sorted(
                        hits, key=lambda c: c["chunk_id"]
                    )
                ],
            }
        }

    def delete_by_query(self, **kwargs: Any) -> dict[str, Any]:
        query = kwargs["body"]["query"]
        term = query.get("term")
        deleted = 0

        if term and "document_id" in term:
            target = term["document_id"]
            for chunk_id in [
                cid for cid, c in self.chunks.items()
                if c["document_id"] == target
            ]:
                del self.chunks[chunk_id]
                deleted += 1
        elif term and "country_code" in term:
            country = term["country_code"]
            for chunk_id in [
                cid for cid, c in self.chunks.items()
                if c["country_code"] == country
            ]:
                del self.chunks[chunk_id]
                deleted += 1
        elif "bool" in query:
            country = query["bool"]["filter"][0]["term"]["country_code"]
            keep_ids = set(
                query["bool"]["must_not"][0]["terms"]["chunk_id"]
            )
            for chunk_id in [
                cid for cid, c in self.chunks.items()
                if c["country_code"] == country and cid not in keep_ids
            ]:
                del self.chunks[chunk_id]
                deleted += 1

        return {"deleted": deleted}


def _bulk_fake(client: FakeOpenSearch, actions, **kwargs):
    action_list = list(actions)
    for action in action_list:
        doc = action["_source"]
        client.add(
            document_id=doc["document_id"],
            chunk_id=doc["chunk_id"],
            country_code=doc["country_code"],
            source_filename=doc["source_filename"],
        )
    return (len(action_list), [])


def _make_upload_file(filename: str, content: bytes) -> UploadFile:
    return UploadFile(file=io.BytesIO(content), filename=filename)


class AdminRouterIntegrationTestCase(unittest.TestCase):
    """
    Base class: fresh temp source/processed dirs and a fresh
    FakeOpenSearch per test, get_settings()/get_opensearch_client()
    patched at every call site the admin code uses them from.
    """

    def setUp(self) -> None:
        self._tempdir = tempfile.TemporaryDirectory()
        root = Path(self._tempdir.name)
        self.source_dir = root / "source"
        self.processed_dir = root / "processed"
        self.source_dir.mkdir(parents=True)
        self.processed_dir.mkdir(parents=True)

        self._original_env = {
            key: os.environ.get(key) for key in _SETTINGS_ENV_KEYS
        }
        os.environ["OPENSEARCH_URL"] = "http://unused-in-this-test:9200"
        os.environ["OPENSEARCH_PASSWORD"] = "unused"
        os.environ["REDIS_URL"] = "redis://unused-in-this-test:6379/0"
        os.environ["DOCUMENT_SOURCE_DIR"] = str(self.source_dir)
        os.environ["DOCUMENT_PROCESSED_DIR"] = str(self.processed_dir)
        os.environ["ADMIN_API_KEY"] = "unused-admin-key"
        os.environ["API_ACCESS_KEY"] = "unused-api-key"
        get_settings.cache_clear()

        self.fake = FakeOpenSearch()

        self._patches = [
            patch(
                "app.services.admin_document_replacement."
                "get_opensearch_client",
                return_value=self.fake,
            ),
            patch(
                "app.services.admin_document_lifecycle."
                "get_opensearch_client",
                return_value=self.fake,
            ),
            patch(
                "app.services.admin_documents.get_opensearch_client",
                return_value=self.fake,
            ),
            patch(
                "app.services.document_indexer.get_opensearch_client",
                return_value=self.fake,
            ),
            patch(
                "app.services.document_indexer.bulk",
                side_effect=lambda client, actions, **kw: _bulk_fake(
                    self.fake, actions, **kw
                ),
            ),
            patch(
                "app.services.document_indexer."
                "ensure_legal_documents_index"
            ),
        ]
        for p in self._patches:
            p.start()

    def tearDown(self) -> None:
        for p in reversed(self._patches):
            p.stop()

        for key, value in self._original_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value
        get_settings.cache_clear()

        self._tempdir.cleanup()


class FreshUploadHttpContractTests(AdminRouterIntegrationTestCase):
    """Mission section 5 - fresh upload -> HTTP 201, then delete, then
    reupload under a different filename."""

    def test_fresh_upload_returns_201_uploaded_one_document_indexed(
        self,
    ) -> None:
        upload = _make_upload_file("Argentina.docx", _AR_BYTES)

        response = documents_router.upload_admin_document(
            file=upload,
            replace_existing=False,
            confirm_warnings=True,
            country_confirmed=True,
            selected_country_code=None,
        )

        self.assertEqual(response.status, "uploaded")
        self.assertEqual(response.country_code, "AR")
        self.assertGreater(response.indexed_chunks, 0)
        self.assertTrue((self.source_dir / "AR.docx").exists())
        self.assertEqual(
            len(self.fake.document_ids_for_country("AR")), 1
        )

    def test_upload_then_delete_then_reupload_different_filename(
        self,
    ) -> None:
        first = documents_router.upload_admin_document(
            file=_make_upload_file("Argentina.docx", _AR_BYTES),
            replace_existing=False,
            confirm_warnings=True,
            country_confirmed=True,
            selected_country_code=None,
        )
        self.assertEqual(first.status, "uploaded")

        delete_response = lifecycle_router.delete_admin_document(
            document_id=first.document_id
        )
        self.assertEqual(delete_response.status, "deleted")
        self.assertEqual(
            len(self.fake.document_ids_for_country("AR")), 0
        )
        self.assertFalse((self.source_dir / "AR.docx").exists())

        second = documents_router.upload_admin_document(
            file=_make_upload_file(
                "random-file-name.docx", _AR_BYTES
            ),
            replace_existing=False,
            confirm_warnings=True,
            country_confirmed=True,
            selected_country_code=None,
        )

        self.assertEqual(second.status, "uploaded")
        self.assertEqual(
            len(self.fake.document_ids_for_country("AR")), 1
        )
        self.assertTrue((self.source_dir / "AR.docx").exists())


class ExistingCountryHttpContractTests(AdminRouterIntegrationTestCase):
    """Mission section 6 - existing country always yields 409 with the
    structured detail, regardless of the uploaded filename."""

    def setUp(self) -> None:
        super().setUp()
        self.fake.add(
            document_id="doc_" + "a" * 64,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        (self.source_dir / "AR.docx").write_bytes(b"existing-ar-bytes")

    def test_same_filename_returns_409_with_structured_detail(
        self,
    ) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(
                file=_make_upload_file("Argentina.docx", _AR_BYTES),
                replace_existing=False,
                confirm_warnings=True,
                country_confirmed=True,
                selected_country_code=None,
            )

        error = context.exception
        self.assertEqual(error.status_code, 409)
        self.assertEqual(
            error.detail["code"], "document_replacement_required"
        )
        self.assertEqual(error.detail["country_code"], "AR")
        # No mutation.
        self.assertEqual(
            (self.source_dir / "AR.docx").read_bytes(),
            b"existing-ar-bytes",
        )
        self.assertEqual(
            len(self.fake.document_ids_for_country("AR")), 1
        )

    def test_completely_unrelated_filename_still_returns_409(
        self,
    ) -> None:
        # Same content (AR), a filename with no relation whatsoever
        # to the country or to the existing source's own name.
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(
                file=_make_upload_file(
                    "Legal-update-final-v7.docx", _AR_BYTES
                ),
                replace_existing=False,
                confirm_warnings=True,
                country_confirmed=True,
                selected_country_code=None,
            )

        error = context.exception
        self.assertEqual(error.status_code, 409)
        self.assertEqual(
            error.detail["code"], "document_replacement_required"
        )
        self.assertEqual(
            (self.source_dir / "AR.docx").read_bytes(),
            b"existing-ar-bytes",
        )


class DeleteHttpContractTests(AdminRouterIntegrationTestCase):
    """Mission section 9 - DELETE HTTP contract: duplicates (2 and 3
    ids), the last document, source missing, and an unknown id."""

    def test_delete_one_of_two_duplicates_defers_and_keeps_the_other(
        self,
    ) -> None:
        first_id = "doc_" + "a" * 64
        second_id = "doc_" + "b" * 64
        self.fake.add(
            document_id=first_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        self.fake.add(
            document_id=second_id,
            country_code="AR",
            source_filename="Argentina.docx",
            chunk_id=f"{second_id}-chunk-0",
        )
        (self.source_dir / "AR.docx").write_bytes(b"canonical")
        (self.source_dir / "Argentina.docx").write_bytes(b"legacy")

        response = lifecycle_router.delete_admin_document(
            document_id=first_id
        )

        self.assertEqual(response.status, "deleted")
        self.assertTrue(response.source_cleanup_deferred)
        self.assertFalse(response.source_file_deleted)
        self.assertEqual(
            self.fake.document_ids_for_country("AR"), {second_id}
        )
        self.assertTrue((self.source_dir / "AR.docx").exists())
        self.assertTrue((self.source_dir / "Argentina.docx").exists())

    def test_delete_one_of_three_duplicates_defers_and_keeps_the_others(
        self,
    ) -> None:
        ids = ["doc_" + letter * 64 for letter in ("a", "b", "c")]
        for document_id in ids:
            self.fake.add(
                document_id=document_id,
                country_code="AR",
                source_filename="Argentina.docx",
                chunk_id=f"{document_id}-chunk-0",
            )
        (self.source_dir / "AR.docx").write_bytes(b"canonical")

        response = lifecycle_router.delete_admin_document(
            document_id=ids[0]
        )

        self.assertEqual(response.status, "deleted")
        self.assertTrue(response.source_cleanup_deferred)
        self.assertEqual(
            self.fake.document_ids_for_country("AR"), set(ids[1:])
        )
        self.assertTrue((self.source_dir / "AR.docx").exists())

    def test_delete_of_the_last_document_cleans_up_candidate_files(
        self,
    ) -> None:
        only_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=only_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        (self.source_dir / "AR.docx").write_bytes(b"canonical")
        (self.source_dir / "Argentina.docx").write_bytes(b"legacy")

        response = lifecycle_router.delete_admin_document(
            document_id=only_id
        )

        self.assertEqual(response.status, "deleted")
        self.assertFalse(response.source_cleanup_deferred)
        self.assertTrue(response.source_file_deleted)
        self.assertEqual(real_source_entries(self.source_dir), [])
        self.assertEqual(
            self.fake.document_ids_for_country("AR"), set()
        )

    def test_delete_with_source_missing_still_succeeds(self) -> None:
        only_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=only_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        # No physical file at all - the source is already missing.

        response = lifecycle_router.delete_admin_document(
            document_id=only_id
        )

        self.assertEqual(response.status, "deleted")
        self.assertFalse(response.source_file_deleted)
        self.assertFalse(response.source_cleanup_deferred)

    def test_delete_unknown_document_id_returns_404(self) -> None:
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.delete_admin_document(
                document_id="doc_" + "f" * 64
            )

        self.assertEqual(context.exception.status_code, 404)

        # Mission "ORDER 2": the detail must be the structured
        # {"code", "message", "operation", "document_id"} contract,
        # not a bare string a caller has to parse by hand.
        detail = context.exception.detail
        self.assertEqual(detail["code"], "document_not_found")
        self.assertEqual(detail["operation"], "delete")
        self.assertEqual(detail["document_id"], "doc_" + "f" * 64)
        self.assertTrue(detail["message"])


class ReindexHttpContractTests(AdminRouterIntegrationTestCase):
    """
    Mission "ORDER 2": every reindex failure path must return the
    structured {"code", "message", "operation", "document_id"}
    contract instead of one fixed, generic sentence - this is what
    let Czech Republic's real UndeterminableDocumentCountryError
    reach WordPress/the logs as a generic 502 with no actionable
    detail (see backend/app/core/country_registry.py's leading-
    definite-article fallback for the actual Czech fix; this class
    only covers the observability contract, not that fix).
    """

    def test_reindex_unknown_document_id_returns_structured_404(
        self,
    ) -> None:
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.reindex_admin_document(
                document_id="doc_" + "f" * 64
            )

        self.assertEqual(context.exception.status_code, 404)
        detail = context.exception.detail
        self.assertEqual(detail["code"], "document_not_found")
        self.assertEqual(detail["operation"], "reindex")
        self.assertEqual(detail["document_id"], "doc_" + "f" * 64)

    def test_reindex_with_source_missing_returns_structured_409(
        self,
    ) -> None:
        document_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        # No physical file at all, canonical or historical.

        with self.assertRaises(HTTPException) as context:
            lifecycle_router.reindex_admin_document(
                document_id=document_id
            )

        self.assertEqual(context.exception.status_code, 409)
        detail = context.exception.detail
        self.assertEqual(detail["code"], "source_missing")
        self.assertEqual(detail["operation"], "reindex")
        self.assertEqual(detail["document_id"], document_id)
        self.assertIn("missing", detail["message"].casefold())

    def test_reindex_unparseable_source_returns_structured_502(
        self,
    ) -> None:
        # A source DOCX whose content names no supported country at
        # all - the same class of chunk_builder failure Czech
        # Republic's real front matter hit (an unregistered phrasing
        # of a real country), reproduced generically here rather than
        # depending on any one country's exact wording.
        document_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        (self.source_dir / "Argentina.docx").write_bytes(
            _build_real_docx_bytes(
                "Some random legal memo with no title structure."
            )
        )

        with self.assertRaises(HTTPException) as context:
            lifecycle_router.reindex_admin_document(
                document_id=document_id
            )

        self.assertEqual(context.exception.status_code, 502)
        detail = context.exception.detail
        self.assertEqual(detail["code"], "document_reindex_failed")
        self.assertEqual(detail["operation"], "reindex")
        self.assertEqual(detail["document_id"], document_id)
        # The real underlying reason must survive to the client, not
        # just a fixed "could not be reindexed" sentence.
        self.assertIn(
            "supported country", detail["message"].casefold()
        )

    def test_reindex_with_country_conflict_returns_structured_409(
        self,
    ) -> None:
        # ORDER 8A, section 23 - a country with more than one active
        # document_id must refuse Reindex with a structured
        # country_document_conflict, never an unmapped raw exception.
        document_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        self.fake.add(
            document_id="doc_" + "b" * 64,
            country_code="AR",
            source_filename="Argentina-legacy.docx",
            chunk_id="doc_" + "b" * 64 + "-chunk-0",
        )

        with self.assertRaises(HTTPException) as context:
            lifecycle_router.reindex_admin_document(
                document_id=document_id
            )

        self.assertEqual(context.exception.status_code, 409)
        detail = context.exception.detail
        self.assertEqual(detail["code"], "country_document_conflict")
        self.assertEqual(detail["operation"], "reindex")
        self.assertEqual(detail["country_code"], "AR")


class TechnicalValidationHttpContractTests(AdminRouterIntegrationTestCase):
    """Mission "ORDER 3", sections 8/9/27 - every technical upload
    failure returns its own specific, structured code."""

    def test_wrong_extension_returns_invalid_document_type(
        self,
    ) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(
                file=_make_upload_file(
                    "not-a-docx.txt", _AR_BYTES
                ),
                replace_existing=False,
                confirm_warnings=True,
                country_confirmed=True,
                selected_country_code=None,
            )

        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(
            context.exception.detail["code"],
            "invalid_document_type",
        )

    def test_empty_file_returns_document_empty(self) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(
                file=_make_upload_file("Argentina.docx", b""),
                replace_existing=False,
                confirm_warnings=True,
                country_confirmed=True,
                selected_country_code=None,
            )

        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(
            context.exception.detail["code"], "document_empty"
        )

    def test_oversized_file_returns_413_document_too_large(
        self,
    ) -> None:
        os.environ["DOCUMENT_UPLOAD_MAX_BYTES"] = "100"
        get_settings.cache_clear()
        try:
            with self.assertRaises(HTTPException) as context:
                documents_router.upload_admin_document(
                    file=_make_upload_file(
                        "Argentina.docx", _AR_BYTES
                    ),
                    replace_existing=False,
                    confirm_warnings=True,
                    country_confirmed=True,
                    selected_country_code=None,
                )
        finally:
            os.environ.pop("DOCUMENT_UPLOAD_MAX_BYTES", None)
            get_settings.cache_clear()

        self.assertEqual(context.exception.status_code, 413)
        detail = context.exception.detail
        self.assertEqual(detail["code"], "document_too_large")
        self.assertEqual(detail["max_bytes"], 100)
        self.assertIn("max_mb", detail)

    def test_corrupt_zip_returns_document_corrupt(self) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(
                file=_make_upload_file(
                    "Argentina.docx", b"not a real docx file at all"
                ),
                replace_existing=False,
                confirm_warnings=True,
                country_confirmed=True,
                selected_country_code=None,
            )

        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(
            context.exception.detail["code"], "document_corrupt"
        )

    def test_no_identifiable_country_returns_selection_required(
        self,
    ) -> None:
        # Mission "ORDER 8E-A1", section 8: an otherwise-processable
        # DOCX with no identifiable country is no longer a hard
        # 422 failure - it is a 409 SELECT_COUNTRY decision carrying
        # the allowed country list (superseded from "ORDER 3"'s
        # original document_country_undetermined/422 expectation).
        no_country_bytes = _build_real_docx_bytes(
            "Some random legal memo with no title structure."
        )

        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(
                file=_make_upload_file(
                    "mystery.docx", no_country_bytes
                ),
                replace_existing=False,
                confirm_warnings=True,
                country_confirmed=True,
                selected_country_code=None,
            )

        self.assertEqual(context.exception.status_code, 409)
        detail = context.exception.detail
        self.assertEqual(
            detail["code"],
            "document_country_selection_required",
        )
        self.assertTrue(detail["allowed_countries"])


class DownloadHttpContractTests(AdminRouterIntegrationTestCase):
    """Mission "ORDER 3", section 25 - GET .../download."""

    def test_download_returns_the_real_source_bytes(self) -> None:
        document_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        real_bytes = b"the-real-argentina-source-bytes"
        (self.source_dir / "AR.docx").write_bytes(real_bytes)

        response = lifecycle_router.download_admin_document(
            document_id=document_id
        )

        self.assertEqual(
            Path(response.path).read_bytes(), real_bytes
        )
        self.assertEqual(
            response.filename, "Argentina.docx"
        )

    def test_download_unknown_document_id_returns_404(self) -> None:
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.download_admin_document(
                document_id="doc_" + "f" * 64
            )

        self.assertEqual(context.exception.status_code, 404)
        self.assertEqual(
            context.exception.detail["code"], "document_not_found"
        )

    def test_download_with_missing_source_returns_409(self) -> None:
        document_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        # No physical file at all.

        with self.assertRaises(HTTPException) as context:
            lifecycle_router.download_admin_document(
                document_id=document_id
            )

        self.assertEqual(context.exception.status_code, 409)
        self.assertEqual(
            context.exception.detail["code"], "source_missing"
        )

    def test_download_with_source_conflict_returns_409(self) -> None:
        document_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        # Both the canonical AND the historical name exist - two
        # distinct real files resolve for the same document.
        (self.source_dir / "AR.docx").write_bytes(b"canonical")
        (self.source_dir / "Argentina.docx").write_bytes(b"legacy")

        with self.assertRaises(HTTPException) as context:
            lifecycle_router.download_admin_document(
                document_id=document_id
            )

        self.assertEqual(context.exception.status_code, 409)
        self.assertEqual(
            context.exception.detail["code"], "source_conflict"
        )

    def test_download_with_no_contact_state_returns_source_unmaterialized(
        self,
    ) -> None:
        # No sidecar ever written for this document_id - matches every
        # existing test/document above, and confirms the fallback path
        # is unchanged (mission "ORDER 8G-B2.1").
        document_id = "doc_" + "a" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        real_bytes = _build_real_docx_bytes("Argentina overview")
        (self.source_dir / "AR.docx").write_bytes(real_bytes)

        download = lifecycle_service.get_document_download(
            document_id=document_id,
            source_directory=self.source_dir,
        )

        self.assertFalse(download.contacts_materialized)
        self.assertIsNone(download.cleanup_path)
        self.assertEqual(
            download.path.read_bytes(), real_bytes
        )

    def test_download_with_contact_state_returns_materialized_effective_docx(
        self,
    ) -> None:
        document_id = "doc_" + "b" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        (self.source_dir / "AR.docx").write_bytes(
            _build_real_docx_bytes("Argentina overview")
        )

        write_contact_state_atomic(
            self.source_dir,
            ContactState(
                document_id=document_id,
                country_code="AR",
                contacts=(
                    ContactRecord(
                        contact_id="contact-1",
                        member_firm="CURRENT FIRM",
                        contact_person="Current Person",
                        email="current@example.com",
                        phone="+1 555 0100",
                        address="1 Current Street",
                        website="www.current.example",
                    ),
                ),
            ),
        )

        download = lifecycle_service.get_document_download(
            document_id=document_id,
            source_directory=self.source_dir,
        )

        try:
            self.assertTrue(download.contacts_materialized)
            self.assertIsNotNone(download.cleanup_path)
            self.assertTrue(download.cleanup_path.exists())

            reparsed = extract_contacts_from_docx(download.path)
            self.assertEqual(len(reparsed), 1)
            self.assertEqual(
                reparsed[0].member_firm, "CURRENT FIRM"
            )
            self.assertEqual(
                reparsed[0].email, "current@example.com"
            )

            # source itself is never touched
            self.assertNotEqual(
                download.path, self.source_dir / "AR.docx"
            )
        finally:
            if download.cleanup_path is not None:
                download.cleanup_path.unlink(missing_ok=True)

    def test_materialization_failure_leaves_no_orphan_temp_file(
        self,
    ) -> None:
        document_id = "doc_" + "c" * 64
        self.fake.add(
            document_id=document_id,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        (self.source_dir / "AR.docx").write_bytes(
            _build_real_docx_bytes("Argentina overview")
        )

        write_contact_state_atomic(
            self.source_dir,
            ContactState(
                document_id=document_id,
                country_code="AR",
                contacts=(
                    ContactRecord(
                        contact_id="contact-1",
                        member_firm="CURRENT FIRM",
                    ),
                ),
            ),
        )

        temp_dir_before = {
            entry.name
            for entry in Path(tempfile.gettempdir()).iterdir()
            if entry.suffix == ".docx"
        }

        with patch(
            "app.services.admin_document_lifecycle."
            "materialize_effective_docx",
            side_effect=RuntimeError("boom"),
        ):
            with self.assertRaises(RuntimeError):
                lifecycle_service.get_document_download(
                    document_id=document_id,
                    source_directory=self.source_dir,
                )

        temp_dir_after = {
            entry.name
            for entry in Path(tempfile.gettempdir()).iterdir()
            if entry.suffix == ".docx"
        }

        self.assertEqual(
            temp_dir_before,
            temp_dir_after,
            "a failed materialization must never leave a temp file behind",
        )


class StatsHttpContractTests(AdminRouterIntegrationTestCase):
    """Mission "ORDER 3", section 26 - GET .../documents/stats."""

    def test_stats_reflect_the_real_catalog_exactly(self) -> None:
        self.fake.add(
            document_id="doc_" + "a" * 64,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        self.fake.add(
            document_id="doc_" + "b" * 64,
            country_code="BR",
            source_filename="Brazil.docx",
        )
        self.fake.add(
            document_id="doc_" + "c" * 64,
            country_code="CO",
            source_filename="Colombia.docx",
            chunk_id="doc_" + "c" * 64 + "-chunk-0",
        )
        # AR and BR have a real source on disk; CO does not - a
        # deliberately mixed status set, and each country_code has
        # exactly one document (avoids resolve_document_source_path's
        # own canonical-name fallback complicating the scenario).
        (self.source_dir / "AR.docx").write_bytes(b"argentina")
        (self.source_dir / "BR.docx").write_bytes(b"brazil")

        stats = documents_router.get_admin_document_stats_route()

        self.assertEqual(stats.total_documents, 3)
        self.assertEqual(stats.total_countries, 3)
        self.assertEqual(
            stats.status_counts.get("indexed"), 2
        )
        self.assertEqual(
            stats.status_counts.get("indexed_source_missing"), 1
        )


class SharedInvariantHelperUsageTests(AdminRouterIntegrationTestCase):
    """Mission "ORDER 3", section 7 - demonstrates the shared,
    reusable invariant helpers (admin_invariants.py) against a real
    catalog listing, rather than each caller re-deriving these checks."""

    def test_clean_catalog_satisfies_every_generic_invariant(
        self,
    ) -> None:
        first = documents_router.upload_admin_document(
            file=_make_upload_file("Argentina.docx", _AR_BYTES),
            replace_existing=False,
            confirm_warnings=True,
            country_confirmed=True,
            selected_country_code=None,
        )

        catalog = documents_router.get_admin_documents()
        documents = [d.model_dump() for d in catalog.documents]

        assert_one_active_document_per_country(documents, "AR")
        assert_one_active_source(documents, "AR")

        real_count = self.fake.search(
            index="unused",
            body={"query": {"term": {"country_code": "AR"}}},
        )["hits"]["total"]["value"]

        assert_chunk_count_matches(
            first.indexed_chunks, real_count, country_code="AR"
        )

        total_catalog_chunks = sum(
            d["chunk_count"] for d in documents
        )
        total_real_chunks = len(self.fake.chunks)
        assert_no_orphan_chunks(
            total_catalog_chunks, total_real_chunks
        )

    def test_rejected_replacement_is_zero_mutation_via_shared_helper(
        self,
    ) -> None:
        self.fake.add(
            document_id="doc_" + "a" * 64,
            country_code="AR",
            source_filename="Argentina.docx",
        )
        (self.source_dir / "AR.docx").write_bytes(b"existing-ar-bytes")

        before = [
            d.model_dump()
            for d in documents_router.get_admin_documents().documents
        ]

        with self.assertRaises(HTTPException):
            documents_router.upload_admin_document(
                file=_make_upload_file(
                    "Legal-update-final-v7.docx", _AR_BYTES
                ),
                replace_existing=False,
                confirm_warnings=True,
                country_confirmed=True,
                selected_country_code=None,
            )

        after = [
            d.model_dump()
            for d in documents_router.get_admin_documents().documents
        ]

        assert_zero_mutation(before, after)


class RestoreEndpointRemovedTests(unittest.TestCase):
    """
    ORDER 8A, section 5: the ORDER 7C "Restore from document" endpoint
    is now moot (Edit itself mutates the current DOCX) and must be
    fully removed from the runtime - never merely hidden.
    """

    def test_no_restore_route_is_registered(self) -> None:
        restore_paths = [
            route.path
            for route in lifecycle_router.router.routes
            if route.path.endswith("/restore")
        ]

        self.assertEqual(restore_paths, [])

    def test_add_section_route_is_registered_instead(self) -> None:
        add_section_routes = [
            route
            for route in lifecycle_router.router.routes
            if route.path
            == "/api/v1/admin/documents/{document_id}/sections"
            and "POST" in route.methods
        ]

        self.assertEqual(len(add_section_routes), 1)

    def test_restore_effective_section_no_longer_exists(self) -> None:
        import app.services.admin_document_sections as sections_service

        self.assertFalse(
            hasattr(sections_service, "restore_effective_section")
        )


class IdenticalButAdminModifiedHttpContractTests(
    AdminRouterIntegrationTestCase
):
    """
    Mission "ORDER 8G-B2", section 14 - the router's own new exception
    -> HTTP mapping for AdminDocumentIdenticalButAdminModifiedError.
    Deliberately does not extend FakeOpenSearch's own query shapes for
    this (its delete_by_query double is hard-coded to the country-wide
    cleanup shape, not the Contact-chunk-scoped one) - the service-level
    behavior this maps is already exhaustively covered in
    test_admin_document_replacement.py; this proves only the NEW
    router-level mapping itself.
    """

    def test_maps_to_409_with_structured_detail(self) -> None:
        from app.services.admin_document_replacement import (
            AdminDocumentIdenticalButAdminModifiedError,
        )

        def _raise(**kwargs):
            del kwargs
            raise AdminDocumentIdenticalButAdminModifiedError(
                country="Argentina",
                country_code="AR",
                document_id="doc_" + "a" * 64,
            )

        with patch(
            "app.routers.admin_documents.safe_upload_and_index_document",
            side_effect=_raise,
        ):
            with self.assertRaises(HTTPException) as context:
                documents_router.upload_admin_document(
                    file=_make_upload_file("Argentina.docx", _AR_BYTES),
                    replace_existing=False,
                    confirm_warnings=True,
                    country_confirmed=True,
                    selected_country_code=None,
                )

        error = context.exception
        self.assertEqual(error.status_code, 409)
        self.assertEqual(
            error.detail["code"],
            "document_identical_but_admin_modified",
        )
        self.assertTrue(error.detail["admin_modified"])
        self.assertEqual(
            error.detail["document_id"], "doc_" + "a" * 64
        )
