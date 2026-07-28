import tempfile
import unittest
from pathlib import Path

from docx import Document

from app.services.docx_parser import (
    parse_docx_sections,
)


class LegacySubsectionParserTests(
    unittest.TestCase
):
    def _save_document(
        self,
        document: Document,
        directory: str,
        filename: str = "test.docx",
    ) -> Path:
        file_path = Path(
            directory
        ) / filename

        document.save(
            file_path
        )

        return file_path

    def test_detects_known_bold_legacy_subsection(
        self,
    ) -> None:
        document = Document()

        document.add_heading(
            "01. Hiring Practices",
            level=1,
        )

        subsection = document.add_paragraph(
            "Requirement for Foreign Employees to Work"
        )

        subsection.runs[0].bold = True

        document.add_paragraph(
            "Foreign employees require a work permit."
        )

        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(
                document=document,
                directory=directory,
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Spain",
            )

        self.assertEqual(
            len(sections),
            1,
        )

        self.assertEqual(
            sections[0].subsection,
            "Requirement for Foreign Employees to Work",
        )

        self.assertEqual(
            sections[0].content,
            "Foreign employees require a work permit.",
        )

    def test_detects_known_heading_four_subsection(
        self,
    ) -> None:
        document = Document()

        document.add_heading(
            "07. Termination of Employment Contracts",
            level=1,
        )

        document.add_heading(
            "Grounds for Termination",
            level=4,
        )

        document.add_paragraph(
            "Termination requires a valid legal ground."
        )

        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(
                document=document,
                directory=directory,
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Italy",
            )

        self.assertEqual(
            len(sections),
            1,
        )

        self.assertEqual(
            sections[0].subsection,
            "Grounds for Termination",
        )

    def test_keeps_unknown_bold_paragraph_as_content(
        self,
    ) -> None:
        document = Document()

        document.add_heading(
            "07. Termination of Employment Contracts",
            level=1,
        )

        unknown_title = document.add_paragraph(
            "Immediate Termination"
        )

        unknown_title.runs[0].bold = True

        document.add_paragraph(
            "Immediate termination is possible in limited cases."
        )

        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(
                document=document,
                directory=directory,
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Czech Republic",
            )

        self.assertEqual(
            len(sections),
            1,
        )

        self.assertIsNone(
            sections[0].subsection
        )

        self.assertEqual(
            sections[0].content,
            (
                "Immediate Termination\n\n"
                "Immediate termination is possible "
                "in limited cases."
            ),
        )

    def test_normalizes_known_typo_alias(
        self,
    ) -> None:
        document = Document()

        document.add_heading(
            "07. Termination of Employment Contracts",
            level=1,
        )

        subsection = document.add_paragraph(
            "Whitsleblower Laws"
        )

        subsection.runs[0].bold = True

        document.add_paragraph(
            "Whistleblowers receive statutory protection."
        )

        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(
                document=document,
                directory=directory,
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Japan",
            )

        self.assertEqual(
            sections[0].subsection,
            "Whistleblower Laws",
        )

    def test_scopes_same_subsection_to_each_topic(
        self,
    ) -> None:
        document = Document()

        document.add_heading(
            "04. Anti-Discrimination Laws",
            level=1,
        )

        first_remedies = document.add_paragraph(
            "Remedies"
        )

        first_remedies.runs[0].bold = True

        document.add_paragraph(
            "Anti-discrimination remedies."
        )

        document.add_heading(
            "05. Pay Equity Laws",
            level=1,
        )

        second_remedies = document.add_paragraph(
            "Remedies"
        )

        second_remedies.runs[0].bold = True

        document.add_paragraph(
            "Pay equity remedies."
        )

        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(
                document=document,
                directory=directory,
            )

            sections = parse_docx_sections(
                file_path=file_path,
                country="Spain",
            )

        self.assertEqual(
            len(sections),
            2,
        )

        self.assertEqual(
            sections[0].subsection,
            "Remedies",
        )

        self.assertEqual(
            sections[1].subsection,
            "Remedies",
        )

        self.assertIn(
            "Anti-Discrimination",
            sections[0].section,
        )

        self.assertIn(
            "Pay Equity",
            sections[1].section,
        )

    def test_generic_parser_keeps_custom_heading_two(
        self,
    ) -> None:
        document = Document()

        document.add_heading(
            "Custom Main Section",
            level=1,
        )

        document.add_heading(
            "Custom Subsection",
            level=2,
        )

        document.add_paragraph(
            "Custom content."
        )

        with tempfile.TemporaryDirectory() as directory:
            file_path = self._save_document(
                document=document,
                directory=directory,
            )

            sections = parse_docx_sections(
                file_path=file_path,
            )

        self.assertEqual(
            len(sections),
            1,
        )

        self.assertEqual(
            sections[0].section,
            "Custom Main Section",
        )

        self.assertEqual(
            sections[0].subsection,
            "Custom Subsection",
        )

        self.assertEqual(
            sections[0].content,
            "Custom content.",
        )


if __name__ == "__main__":
    unittest.main()