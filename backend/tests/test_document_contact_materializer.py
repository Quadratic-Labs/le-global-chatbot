"""
Targeted tests for the legacy floating-shape Contact-box detection
helpers in app.services.document_contact_materializer.

This module used to ALSO test materialize_effective_docx() (the
per-download ephemeral DOCX rebuild) - that function was removed
(mission "DOCX HARDENING", 2026-08-24): every Admin mutation already
persists its effective result into the source DOCX atomically before
returning success, so download is now a pure byte read of the
persisted source (see test_admin_documents_router_integration.py's
DownloadHttpContractTests and
backend/integration_tests/docx_contact_mutation_matrix.py). What
remains here - _is_contact_related_block/_find_all_contact_runs - is
still live: contact_document_photos.py uses it to detect whether a
document's Contact area is still the OLD floating-shape format.
"""

from __future__ import annotations

import io
import unittest
import zipfile

from lxml import etree

from docx import Document

from app.services.document_contact_materializer import (
    _find_all_contact_runs,
    _is_contact_related_block,
)


def _textbox_run_xml(
    lines: list[str],
    width_emu: int = 1000000,
    height_emu: int = 500000,
) -> str:
    inner_paragraphs = "".join(
        f"<w:p><w:r><w:t>{line}</w:t></w:r></w:p>" for line in lines
    )
    return (
        '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<w:rPr><w:noProof/></w:rPr>"
        "<w:drawing><wp:anchor>"
        f'<wp:extent cx="{width_emu}" cy="{height_emu}"/>'
        '<wp:docPr id="1" name="Box"/>'
        "<a:graphic>"
        '<a:graphicData uri="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        "<wps:wsp>"
        f"<wps:txbx><w:txbxContent>{inner_paragraphs}</w:txbxContent></wps:txbx>"
        "</wps:wsp></a:graphicData></a:graphic>"
        "</wp:anchor></w:drawing></w:r>"
    )


def _build_docx_with_contact_textbox_at(
    paragraphs: list[str],
    textbox_paragraph_index: int,
    lines: list[str] | None = None,
    width_emu: int = 1000000,
    height_emu: int = 500000,
) -> bytes:
    """
    A real, openable DOCX with a minimal but structurally faithful
    floating Contact text box (CONTACT PERSON marker + name + email,
    matching the real corpus's own <w:txbxContent> shape, wrapped in a
    DrawingML <w:drawing><wp:anchor> run so _extract_extent and the
    <w:r>/<w:txbxContent> nesting _find_run_span must see through both
    apply) injected into one specific paragraph.
    """

    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)

    target_paragraph = document.paragraphs[textbox_paragraph_index]

    run_element = etree.fromstring(
        _textbox_run_xml(
            lines
            or ["CONTACT PERSON", "Test Person", "test.person@example.com"],
            width_emu=width_emu,
            height_emu=height_emu,
        ).encode("utf-8")
    )
    target_paragraph._p.append(run_element)

    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _document_xml_of(docx_bytes: bytes) -> str:
    with zipfile.ZipFile(io.BytesIO(docx_bytes)) as archive:
        return archive.read("word/document.xml").decode("utf-8")


class IsContactRelatedBlockTests(unittest.TestCase):
    def test_contact_person_block_is_classified_as_contact_related(
        self,
    ) -> None:
        self.assertTrue(
            _is_contact_related_block(["CONTACT PERSON", "Jane Doe"])
        )

    def test_firm_block_with_email_is_classified_as_contact_related(
        self,
    ) -> None:
        self.assertTrue(
            _is_contact_related_block(["SOME FIRM", "info@example.com"])
        )

    def test_firm_block_with_phone_is_classified_as_contact_related(
        self,
    ) -> None:
        self.assertTrue(
            _is_contact_related_block(["SOME FIRM", "+1 555 000 1234"])
        )

    def test_plain_branding_block_is_not_contact_related(self) -> None:
        self.assertFalse(_is_contact_related_block(["www.leglobal.law"]))

    def test_document_title_block_is_not_contact_related(self) -> None:
        self.assertFalse(
            _is_contact_related_block(
                ["Employment Law Overview - Testland"]
            )
        )


class FindAllContactRunsTests(unittest.TestCase):
    """
    Regression coverage for two bugs found and fixed this mission: (1)
    _find_run_span originally took the first "</w:r>" after a
    txbxContent's start, which is actually an INNER run nested inside
    the box's own paragraphs, not the enclosing floating-shape run: it
    must instead balance nested <w:r>/</w:r> tags. (2) a floating
    shape's Choice and Fallback branches each produce their own
    <w:txbxContent> regex match for the SAME enclosing run; the second
    occurrence must be recognized as already-covered by the first
    match's span, not re-resolved (which - before the fix - picked one
    of the first branch's own inner runs as a bogus second "run").
    """

    def test_finds_single_contact_run_with_its_extent(self) -> None:
        docx_bytes = _build_docx_with_contact_textbox_at(
            ["Title", "Body"],
            textbox_paragraph_index=0,
            width_emu=1234000,
            height_emu=567000,
        )

        runs = _find_all_contact_runs(_document_xml_of(docx_bytes))

        self.assertEqual(len(runs), 1)
        self.assertEqual(runs[0].width_emu, 1234000)
        self.assertEqual(runs[0].height_emu, 567000)

    def test_ignores_non_contact_textboxes(self) -> None:
        document = Document()
        document.add_paragraph("Title")
        document.add_paragraph("Body")

        branding_run = etree.fromstring(
            _textbox_run_xml(["www.leglobal.law"]).encode("utf-8")
        )
        document.paragraphs[0]._p.append(branding_run)

        buffer = io.BytesIO()
        document.save(buffer)

        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))
        self.assertEqual(runs, [])

    def test_two_separate_contact_runs_both_found_and_largest_is_selectable(
        self,
    ) -> None:
        document = Document()
        document.add_paragraph("Title")

        small_run = etree.fromstring(
            _textbox_run_xml(
                ["CONTACT PERSON", "Jane Doe", "jane@example.com"],
                width_emu=500000,
                height_emu=300000,
            ).encode("utf-8")
        )
        large_run = etree.fromstring(
            _textbox_run_xml(
                ["SOME FIRM", "firm@example.com"],
                width_emu=2000000,
                height_emu=1000000,
            ).encode("utf-8")
        )
        document.paragraphs[0]._p.append(small_run)
        document.paragraphs[0]._p.append(large_run)

        buffer = io.BytesIO()
        document.save(buffer)

        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))

        self.assertEqual(len(runs), 2)
        primary = max(runs, key=lambda run: run.width_emu * run.height_emu)
        self.assertEqual(primary.width_emu, 2000000)

    def test_choice_and_fallback_style_duplicate_txbxcontent_is_one_run(
        self,
    ) -> None:
        # Mirrors a real floating shape's mc:Choice/mc:Fallback pair:
        # the SAME visible box content appears twice inside one <w:r>,
        # each occurrence separately matched by the txbxContent regex.
        # Must resolve to exactly one run, not two.
        document = Document()
        document.add_paragraph("Title")

        inner = (
            "<w:p><w:r><w:t>CONTACT PERSON</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>Jane Doe</w:t></w:r></w:p>"
            "<w:p><w:r><w:t>jane@example.com</w:t></w:r></w:p>"
        )
        duplicated_run_xml = (
            '<w:r xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
            "<mc:AlternateContent "
            'xmlns:mc="http://schemas.openxmlformats.org/markup-compatibility/2006">'
            '<mc:Choice Requires="wps">'
            f"<w:txbxContent>{inner}</w:txbxContent>"
            "</mc:Choice>"
            "<mc:Fallback>"
            f"<w:txbxContent>{inner}</w:txbxContent>"
            "</mc:Fallback>"
            "</mc:AlternateContent>"
            "</w:r>"
        )
        run_element = etree.fromstring(duplicated_run_xml.encode("utf-8"))
        document.paragraphs[0]._p.append(run_element)

        buffer = io.BytesIO()
        document.save(buffer)

        runs = _find_all_contact_runs(_document_xml_of(buffer.getvalue()))
        self.assertEqual(
            len(runs), 1,
            "Choice+Fallback duplicate content must resolve to one run, "
            "not one real run plus a bogus nested-run duplicate",
        )


if __name__ == "__main__":
    unittest.main()
