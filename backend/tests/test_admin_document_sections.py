"""
Tests for admin section editing - ORDER 8A architecture.

The CURRENT DOCX is the unique source of truth: Edit and Add physically
mutate a copy of the source DOCX, validate it by a full reparse, and
only then atomically apply both the targeted OpenSearch chunks and the
new current source file. There is no more "effective section" override
layer - list/get read the real, current DOCX directly.

Mirrors the FakeOpenSearchClient style already established elsewhere in
this codebase: plain unittest.TestCase, tempfile.TemporaryDirectory()
for source_directory, explicit dependency injection rather than a
mocking framework, except for the one seam admin_document_sections.py
has none for (document_indexer.py's own module-level `bulk`/
`ensure_legal_documents_index` calls) - unittest.mock.patch there only,
never a broader mocking framework.
"""

from __future__ import annotations

import contextlib
import tempfile
import unittest
from pathlib import Path
from typing import Any
from unittest.mock import patch

from docx import Document
from opensearchpy.exceptions import OpenSearchException

from app.services.admin_document_lifecycle import (
    AdminDocumentCountryConflictError,
    AdminDocumentRollbackError,
    reindex_indexed_document,
)
from app.services.admin_document_sections import (
    AdminDocumentSectionAlreadyExistsError,
    AdminDocumentSectionInvalidError,
    AdminDocumentSectionLastRemainingError,
    AdminDocumentSectionNotFoundError,
    AdminDocumentSectionPositionError,
    AdminDocumentSectionUpdateFailedError,
    add_new_section,
    delete_section,
    get_effective_section,
    list_effective_sections,
    section_id_for_legal_topic,
    update_effective_section,
)
from app.services.document_chunk_builder import (
    DocumentMetadata,
    build_document_chunks,
)
from app.services.docx_parser import ParsedSection


def _real_document_id_for(
    country_code: str,
    language: str = "en",
) -> str:
    """
    The one real, deterministic document_id build_document_chunks
    would compute for this country_code/language - document_id is
    derived solely from country_code + DOCUMENT_FAMILY + language,
    never from anything a test happens to pick.
    """

    probe_chunks = build_document_chunks(
        [
            ParsedSection(
                section="Employment Contracts",
                subsection=None,
                content="probe content",
            ),
        ],
        DocumentMetadata(
            country="United Kingdom",
            country_code=country_code,
            reference_year=None,
            language=language,
            source_filename="probe.docx",
        ),
    )

    return probe_chunks[0].document_id


DOCUMENT_ID = _real_document_id_for("GB")
OTHER_COUNTRY_DOCUMENT_ID = _real_document_id_for("FR")

EMPLOYMENT_CONTRACTS_SECTION_ID = section_id_for_legal_topic(
    "Employment Contracts"
)
HIRING_PRACTICES_SECTION_ID = section_id_for_legal_topic(
    "Hiring Practices"
)


def _write_docx(
    path: Path,
    sections: list[tuple[str, str]],
) -> None:
    """
    Build one minimal, real DOCX with a Heading 1 (a genuine L&E
    legal-topic name) followed by one content paragraph, per entry -
    every entry uses the exact same structural signal (Heading 1, no
    numbering/bold/prefix), so a custom (non-taxonomy) topic added
    after at least one real one is recognized on reparse too.

    Starts with a real "Employment Law Overview United Kingdom" H1
    overview heading, exactly like every real corpus document has -
    besides letting metadata_from_content auto-detect the country from
    the document's own content (required for ORDER 8A's mandatory
    reparse-validation), it is also what flips past_front_matter for
    the custom-topic-recognition gate: a custom section added at
    "beginning" lands right after this heading, never before it, the
    same as in a real document.
    """

    document = Document()
    document.add_heading(
        "Employment Law Overview United Kingdom", level=1
    )

    for heading, content in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(content)

    document.save(path)


def _write_bold_only_docx(
    path: Path,
    sections: list[tuple[str, str]],
) -> None:
    """
    ORDER 8A-C - a legacy-style DOCX whose native topics use only
    direct bold run formatting, no Heading 1 style, no numbering -
    representative of the ~10/33 real corpus documents that could not
    support Add before the internal DOCX-native style marker existed.
    """

    document = Document()
    document.add_paragraph("Employment Law Overview United Kingdom")

    for heading, content in sections:
        heading_paragraph = document.add_paragraph()
        heading_paragraph.add_run(heading).bold = True
        document.add_paragraph(content)

    document.save(path)


def _seed_chunk(
    *,
    document_id: str,
    legal_topic: str,
    content: str,
    country_code: str = "GB",
    country: str = "United Kingdom",
    source_filename: str = "GB.docx",
    reference_year: int | None = 2026,
) -> dict[str, Any]:
    """One minimal OpenSearch-shaped chunk source dict for the fake."""

    return {
        "document_id": document_id,
        "country_code": country_code,
        "country": country,
        "source_filename": source_filename,
        "reference_year": reference_year,
        "legal_topic": legal_topic,
        "content": content,
        "content_hash": f"hash-{hash(content)}",
    }


class FakeSectionOpenSearchClient:
    """
    OpenSearch test double for admin_document_sections.py.

    Stateful over self.chunks (chunk_id -> source dict), so a real
    edit's effect (a stale chunk removed, a new/overwritten chunk
    present) can be asserted afterwards, not merely inferred from a
    mock's own call arguments. Paired at the call site with `bulk`
    patched to write real entries into this same self.chunks (see
    _bulk_writer/_patched_indexer below).
    """

    def __init__(
        self,
        *,
        document_id: str = DOCUMENT_ID,
        country_code: str = "GB",
        country: str = "United Kingdom",
        source_filename: str = "GB.docx",
        reference_year: int | None = 2026,
        chunks: dict[str, dict[str, Any]] | None = None,
        fail_delete_by_query_calls: int = 0,
        delete_by_query_failure: Exception | None = None,
        fail_snapshot_search: bool = False,
        snapshot_search_failure: Exception | None = None,
    ) -> None:
        self.document_id = document_id
        self.country_code = country_code
        self.country = country
        self.source_filename = source_filename
        self.reference_year = reference_year
        self.chunks: dict[str, dict[str, Any]] = dict(chunks or {})
        self.delete_by_query_calls: list[dict[str, Any]] = []

        self.fail_delete_by_query_calls = fail_delete_by_query_calls
        self.delete_by_query_failure = delete_by_query_failure
        self.delete_by_query_call_count = 0

        self.fail_snapshot_search = fail_snapshot_search
        self.snapshot_search_failure = snapshot_search_failure

    def search(
        self,
        index: str,
        body: dict[str, Any],
    ) -> dict[str, Any]:
        del index

        term = body.get("query", {}).get("term", {})
        requested_document_id = term.get("document_id")
        requested_country_code = term.get("country_code")

        if "sort" in body:
            if self.fail_snapshot_search:
                raise (
                    self.snapshot_search_failure
                    if self.snapshot_search_failure is not None
                    else RuntimeError(
                        "simulated snapshot search failure"
                    )
                )

            # _fetch_all_chunks' own exhaustive-snapshot shape:
            # track_total_hits + sort on chunk_id (search_after
            # pagination), scoped to EITHER document_id (per-document
            # snapshot) OR country_code (country-wide lookup, used by
            # _ensure_no_country_conflict). This fake never has enough
            # chunks to need a second page.
            if requested_document_id is not None:
                matching_ids = sorted(
                    chunk_id
                    for chunk_id, chunk in self.chunks.items()
                    if chunk["document_id"] == requested_document_id
                )
            elif requested_country_code is not None:
                matching_ids = sorted(
                    chunk_id
                    for chunk_id, chunk in self.chunks.items()
                    if chunk["country_code"] == requested_country_code
                )
            else:
                matching_ids = []

            return {
                "hits": {
                    "total": {"value": len(matching_ids)},
                    "hits": [
                        {
                            "_id": chunk_id,
                            "_source": self.chunks[chunk_id],
                            "sort": [chunk_id],
                        }
                        for chunk_id in matching_ids
                    ],
                }
            }

        if requested_document_id is not None:
            # _get_document_metadata's own shape: size 1, term on
            # document_id, no aggregation/sort.
            if requested_document_id != self.document_id:
                return {"hits": {"hits": []}}

            return {
                "hits": {
                    "hits": [
                        {
                            "_source": {
                                "document_id": self.document_id,
                                "source_filename": (
                                    self.source_filename
                                ),
                                "country": self.country,
                                "country_code": self.country_code,
                                "reference_year": self.reference_year,
                            }
                        }
                    ]
                }
            }

        return {"hits": {"hits": []}}

    def delete_by_query(
        self,
        *,
        index: str,
        body: dict[str, Any],
        conflicts: str,
        refresh: bool,
    ) -> dict[str, Any]:
        del index, conflicts, refresh

        self.delete_by_query_call_count += 1

        if self.delete_by_query_call_count <= self.fail_delete_by_query_calls:
            raise (
                self.delete_by_query_failure
                if self.delete_by_query_failure is not None
                else RuntimeError(
                    "simulated delete_by_query failure (call #"
                    f"{self.delete_by_query_call_count})."
                )
            )

        # document_indexer.py's _delete_chunks_except always wraps its
        # filters in "bool"/"filter" (document_id alone for a whole-
        # document scope, or document_id+legal_topic for a section-
        # scoped one); admin_document_lifecycle.py's own direct calls
        # (reindex's previous-document/stray-chunk cleanup) instead
        # send a flat top-level "term". Both real shapes are handled
        # here, since this same fake now backs both callers.
        query = body["query"]
        keep_ids: set[str] = set()

        if "bool" in query:
            filters = query["bool"]["filter"]

            document_id = next(
                clause["term"]["document_id"]
                for clause in filters
                if "document_id" in clause.get("term", {})
            )
            legal_topic = next(
                (
                    clause["term"]["legal_topic"]
                    for clause in filters
                    if "legal_topic" in clause.get("term", {})
                ),
                None,
            )

            for clause in query["bool"].get("must_not", []):
                keep_ids.update(
                    clause.get("terms", {}).get("chunk_id", [])
                )

        else:
            document_id = query["term"]["document_id"]
            legal_topic = None

        to_delete = [
            chunk_id
            for chunk_id, chunk in self.chunks.items()
            if chunk["document_id"] == document_id
            and (
                legal_topic is None
                or chunk.get("legal_topic") == legal_topic
            )
            and chunk_id not in keep_ids
        ]

        for chunk_id in to_delete:
            del self.chunks[chunk_id]

        self.delete_by_query_calls.append(
            {
                "document_id": document_id,
                "legal_topic": legal_topic,
                "deleted": len(to_delete),
            }
        )

        return {"deleted": len(to_delete), "total": len(to_delete)}


class MustNotBeCalledClient:
    """
    A client double that fails any test relying on it - proves a
    code path really does short-circuit before ever touching
    OpenSearch (e.g. an invalid section_id, or empty content).
    """

    def search(self, index: str, body: dict[str, Any]) -> dict[str, Any]:
        raise AssertionError(
            "OpenSearch search() must not be called for this path."
        )

    def delete_by_query(self, **kwargs: Any) -> dict[str, Any]:
        raise AssertionError(
            "OpenSearch delete_by_query() must not be called for "
            "this path."
        )


def _bulk_writer(
    fake_client: FakeSectionOpenSearchClient,
    *,
    fail_first_n_calls: int = 0,
):
    """
    A `bulk()` side_effect that writes every action into fake_client's
    own self.chunks, so replace_document_section_chunks' effect on the
    fake's state is real, not merely a recorded call.

    fail_first_n_calls fails only the first N invocations (the initial
    indexing attempt) while leaving later ones (a rollback's own
    re-indexing) genuinely succeeding - mirrors
    fail_delete_by_query_calls' own counter-based semantics.
    """

    call_count = {"n": 0}

    def fake_bulk(client, actions, **kwargs):
        del client, kwargs

        call_count["n"] += 1

        if call_count["n"] <= fail_first_n_calls:
            raise RuntimeError(
                f"simulated OpenSearch bulk failure (call #{call_count['n']})"
            )

        action_list = list(actions)

        for action in action_list:
            fake_client.chunks[action["_id"]] = dict(action["_source"])

        return (len(action_list), [])

    return fake_bulk


@contextlib.contextmanager
def _patched_indexer(
    fake_client: FakeSectionOpenSearchClient,
    *,
    fail_bulk: bool = False,
):
    """
    Patch the two document_indexer.py internals
    replace_document_section_chunks calls directly - the one seam
    admin_document_sections.py has none for.

    fail_bulk fails only the FIRST bulk call (the initial indexing
    attempt) - a subsequent rollback's own re-indexing call still
    succeeds, exactly like a real, isolated OpenSearch write failure
    would (never every future call forever).
    """

    with patch(
        "app.services.document_indexer.ensure_legal_documents_index"
    ), patch(
        "app.services.document_indexer.bulk",
        side_effect=_bulk_writer(
            fake_client,
            fail_first_n_calls=1 if fail_bulk else 0,
        ),
    ):
        yield


def _seeded_client(
    *,
    topics: list[tuple[str, str]],
    document_id: str = DOCUMENT_ID,
    country_code: str = "GB",
    country: str = "United Kingdom",
    source_filename: str = "GB.docx",
) -> FakeSectionOpenSearchClient:
    return FakeSectionOpenSearchClient(
        document_id=document_id,
        country_code=country_code,
        country=country,
        source_filename=source_filename,
        chunks={
            f"chunk-seed-{index}": _seed_chunk(
                document_id=document_id,
                legal_topic=topic,
                content=content,
                country_code=country_code,
                country=country,
                source_filename=source_filename,
            )
            for index, (topic, content) in enumerate(topics)
        },
    )


class AdminDocumentSectionListTests(unittest.TestCase):
    def test_lists_every_real_topic_in_the_current_docx(self) -> None:
        # ORDER 8A, section 6 - never a fixed list of the 11 taxonomy
        # topics, never derived from OpenSearch: only whatever the
        # CURRENT DOCX actually contains right now.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "stale opensearch text")]
            )

            response = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(response.document_id, DOCUMENT_ID)
            self.assertEqual(
                [section.legal_topic for section in response.sections],
                ["Employment Contracts", "Hiring Practices"],
            )
            self.assertEqual(
                {section.section_id for section in response.sections},
                {
                    EMPLOYMENT_CONTRACTS_SECTION_ID,
                    HIRING_PRACTICES_SECTION_ID,
                },
            )


class AdminDocumentSectionGetTests(unittest.TestCase):
    def test_content_always_comes_from_the_current_docx(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [
                    (
                        "Employment Contracts",
                        "DOCX text for employment contracts.",
                    ),
                ],
            )

            client = _seeded_client(
                topics=[
                    (
                        "Employment Contracts",
                        "irrelevant stale chunk text",
                    )
                ]
            )

            response = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                source_directory=source_directory,
                client=client,
            )

            self.assertEqual(
                response.content,
                "DOCX text for employment contracts.",
            )
            self.assertEqual(response.country_code, "GB")
            self.assertEqual(response.country_name, "United Kingdom")
            self.assertEqual(
                response.legal_topic, "Employment Contracts"
            )

    def test_invalid_section_id_is_not_found_without_touching_opensearch(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [("Employment Contracts", "EC content.")],
            )

            with self.assertRaises(
                AdminDocumentSectionNotFoundError
            ) as context:
                get_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id="not-a-real-topic-slug",
                    source_directory=source_directory,
                    client=_seeded_client(
                        topics=[("Employment Contracts", "x")]
                    ),
                )

            self.assertEqual(
                context.exception.to_detail()["code"],
                "document_section_not_found",
            )


class AdminDocumentSectionEditTests(unittest.TestCase):
    """
    ORDER 8A, sections 8/14-18: Edit really modifies the current DOCX,
    validated by a full reparse, applied atomically to both OpenSearch
    (targeted to this one legal_topic) and the source file.
    """

    def test_edit_writes_the_current_docx_and_removes_old_content(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                response = update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="  New effective content.  ",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                response.legal_topic, "Employment Contracts"
            )
            self.assertEqual(response.indexed_chunks, 1)

            # The DOCX itself was really modified - old text gone, new
            # text present, heading preserved.
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertIn("Employment Contracts", paragraphs)
            self.assertIn("New effective content.", paragraphs)
            self.assertNotIn("Original DOCX text.", paragraphs)

            # OpenSearch holds exactly the new content, targeted only
            # to this legal_topic.
            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Employment Contracts"
            ]
            self.assertEqual(len(remaining), 1)
            self.assertEqual(
                remaining[0]["content"], "New effective content."
            )

            # Download (the current DOCX) reflects the edit too.
            reparsed = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                reparsed.content, "New effective content."
            )

    def test_unrelated_topics_are_exactly_unchanged(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "Original EC text."),
                    ("Hiring Practices", "Original HP text."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="Edited EC content.",
                    source_directory=source_directory,
                    client=client,
                )

            # Hiring Practices' own chunk survives completely
            # untouched - _delete_stale_section_chunks is scoped to
            # legal_topic as well as document_id.
            hp_chunks = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Hiring Practices"
            ]
            self.assertEqual(len(hp_chunks), 1)
            self.assertEqual(hp_chunks[0]["content"], "placeholder HP")

            # The DOCX's own Hiring Practices paragraph is untouched.
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertIn("Original HP text.", paragraphs)

    def test_reindex_after_edit_reproduces_the_same_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [("Employment Contracts", "Original DOCX text.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="Edited content survives reindex.",
                    source_directory=source_directory,
                    client=client,
                )

            with patch(
                "app.services.document_indexer.ensure_legal_documents_index"
            ), patch(
                "app.services.document_indexer.bulk",
                side_effect=_bulk_writer(client),
            ):
                reindex_indexed_document(
                    document_id=DOCUMENT_ID,
                    source_directory=source_directory,
                    client=client,
                )

            reparsed = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                reparsed.content, "Edited content survives reindex."
            )

    def test_second_edit_fully_overwrites_first_no_history(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [("Employment Contracts", "Original DOCX text.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="First edit content.",
                    source_directory=source_directory,
                    client=client,
                )

                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="Second edit content - final.",
                    source_directory=source_directory,
                    client=client,
                )

            response = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                response.content, "Second edit content - final."
            )

            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Employment Contracts"
            ]
            self.assertEqual(len(remaining), 1)

    def test_empty_content_is_invalid_with_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )

            original_bytes = source_path.read_bytes()

            with self.assertRaises(
                AdminDocumentSectionInvalidError
            ) as context:
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="   \n\t  ",
                    source_directory=source_directory,
                    client=MustNotBeCalledClient(),
                )

            self.assertEqual(
                context.exception.to_detail()["code"],
                "document_section_invalid",
            )
            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

    def test_unknown_section_is_not_found_with_zero_mutation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [("Employment Contracts", "Original DOCX text.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with self.assertRaises(AdminDocumentSectionNotFoundError):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=HIRING_PRACTICES_SECTION_ID,
                    new_content="Some content.",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(len(client.chunks), 1)


class AdminDocumentSectionAddTests(unittest.TestCase):
    """ORDER 8A, sections 9-13: adding a brand-new top-level topic."""

    def test_add_at_end_is_editable_and_retrievable(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "EC content.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                response = add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="Employees may work remotely. MARKER.",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(response.legal_topic, "Remote Working")
            self.assertEqual(
                response.section_id,
                section_id_for_legal_topic("Remote Working"),
            )

            listing = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                [s.legal_topic for s in listing.sections],
                ["Employment Contracts", "Remote Working"],
            )

            fetched = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=section_id_for_legal_topic("Remote Working"),
                source_directory=source_directory,
                client=client,
            )
            self.assertIn("MARKER", fetched.content)

            remote_chunks = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Remote Working"
            ]
            self.assertEqual(len(remote_chunks), 1)
            self.assertIn("MARKER", remote_chunks[0]["content"])

    def test_add_at_beginning_and_after_existing_section(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder"),
                    ("Hiring Practices", "placeholder"),
                ]
            )

            with _patched_indexer(client):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Artificial Intelligence at Work",
                    content="AI disclosure rules.",
                    position="beginning",
                    source_directory=source_directory,
                    client=client,
                )

            listing = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                listing.sections[0].legal_topic,
                "Artificial Intelligence at Work",
            )

            with _patched_indexer(client):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="Remote work rules.",
                    position=f"after:{EMPLOYMENT_CONTRACTS_SECTION_ID}",
                    source_directory=source_directory,
                    client=client,
                )

            listing_after = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            names = [s.legal_topic for s in listing_after.sections]
            self.assertEqual(
                names,
                [
                    "Artificial Intelligence at Work",
                    "Employment Contracts",
                    "Remote Working",
                    "Hiring Practices",
                ],
            )

    def test_duplicate_title_is_rejected_with_zero_mutation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "EC content.")],
            )

            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with self.assertRaises(
                AdminDocumentSectionAlreadyExistsError
            ) as context:
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="employment   CONTRACTS",
                    content="whatever",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                context.exception.to_detail()["code"],
                "section_already_exists",
            )
            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )
            self.assertEqual(len(client.chunks), 1)

    def test_invalid_position_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [("Employment Contracts", "EC content.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with self.assertRaises(AdminDocumentSectionPositionError):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="content",
                    position="middle",
                    source_directory=source_directory,
                    client=client,
                )

    def test_custom_section_survives_full_reindex(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [("Employment Contracts", "EC content.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="Remote work rules. MARKER-REMOTE.",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )

            with patch(
                "app.services.document_indexer.ensure_legal_documents_index"
            ), patch(
                "app.services.document_indexer.bulk",
                side_effect=_bulk_writer(client),
            ):
                reindex_indexed_document(
                    document_id=DOCUMENT_ID,
                    source_directory=source_directory,
                    client=client,
                )

            listing = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertIn(
                "Remote Working",
                [s.legal_topic for s in listing.sections],
            )

            remote_chunks = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Remote Working"
            ]
            self.assertEqual(len(remote_chunks), 1)
            self.assertIn(
                "MARKER-REMOTE", remote_chunks[0]["content"]
            )


class AdminDocumentSectionTransactionTests(unittest.TestCase):
    """
    ORDER 8A, sections 15-18: the exact transaction order and rollback
    matrix. Every failure scenario must leave either the fully-applied
    new state, or exactly the pre-operation state - never a mix.
    """

    def test_zero_mutation_when_target_topic_does_not_exist(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "EC content.")],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with self.assertRaises(AdminDocumentSectionNotFoundError):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=HIRING_PRACTICES_SECTION_ID,
                    new_content="should never apply",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )
            self.assertEqual(len(client.chunks), 1)

            leftover_temps = list(
                source_directory.glob(".*tmp.docx")
            )
            self.assertEqual(leftover_temps, [])

    def test_opensearch_bulk_failure_leaves_source_and_index_untouched(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(
                    AdminDocumentSectionUpdateFailedError
                ):
                    update_effective_section(
                        document_id=DOCUMENT_ID,
                        section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                        new_content="should roll back",
                        source_directory=source_directory,
                        client=client,
                    )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )
            self.assertEqual(
                client.chunks["chunk-seed-0"]["content"],
                "placeholder",
            )

            leftover_temps = list(
                source_directory.glob(".*tmp.docx")
            )
            self.assertEqual(leftover_temps, [])

    def test_stale_delete_failure_rolls_back_to_pre_edit_snapshot(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )
            client.fail_delete_by_query_calls = 1

            with _patched_indexer(client):
                with self.assertRaises(
                    AdminDocumentSectionUpdateFailedError
                ):
                    update_effective_section(
                        document_id=DOCUMENT_ID,
                        section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                        new_content="should roll back",
                        source_directory=source_directory,
                        client=client,
                    )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Employment Contracts"
            ]
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]["content"], "placeholder")

    def test_source_replace_failure_rolls_back_opensearch(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                with patch(
                    "app.services.admin_document_sections.os.replace",
                    side_effect=OSError("simulated disk failure"),
                ):
                    with self.assertRaises(
                        AdminDocumentSectionUpdateFailedError
                    ):
                        update_effective_section(
                            document_id=DOCUMENT_ID,
                            section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                            new_content="should roll back",
                            source_directory=source_directory,
                            client=client,
                        )

            # Source untouched (os.replace never really ran).
            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

            # OpenSearch rolled back to the pre-edit snapshot.
            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Employment Contracts"
            ]
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]["content"], "placeholder")

            leftover_temps = list(
                source_directory.glob(".*tmp.docx")
            )
            self.assertEqual(leftover_temps, [])

    def test_rollback_failure_after_source_replace_failure_is_surfaced(
        self,
    ) -> None:
        # When BOTH the atomic source replace AND the OpenSearch
        # rollback it triggers fail, the transaction must surface a
        # structured rollback_failed error - never a false SUCCESS,
        # and never a raw/unwrapped exception.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                with patch(
                    "app.services.admin_document_sections.os.replace",
                    side_effect=OSError("simulated disk failure"),
                ), patch(
                    "app.services.admin_document_sections."
                    "_restore_section_snapshot",
                    side_effect=RuntimeError(
                        "simulated rollback failure"
                    ),
                ):
                    with self.assertRaises(
                        AdminDocumentRollbackError
                    ):
                        update_effective_section(
                            document_id=DOCUMENT_ID,
                            section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                            new_content="should surface rollback failure",
                            source_directory=source_directory,
                            client=client,
                        )


class AdminDocumentCountryInvariantTests(unittest.TestCase):
    """ORDER 8A, sections 22-23: one active document per country."""

    def test_country_conflict_blocks_edit_with_zero_mutation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            # Seed a second, distinct active document_id for the SAME
            # country_code - a duplicate-country conflict state.
            client.chunks["chunk-conflict"] = _seed_chunk(
                document_id="doc_" + "c" * 64,
                legal_topic="Hiring Practices",
                content="legacy duplicate content",
                country_code="GB",
                country="United Kingdom",
                source_filename="GB-legacy.docx",
            )

            with self.assertRaises(AdminDocumentCountryConflictError):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="should never apply",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

    def test_country_conflict_blocks_add_with_zero_mutation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )
            client.chunks["chunk-conflict"] = _seed_chunk(
                document_id="doc_" + "c" * 64,
                legal_topic="Hiring Practices",
                content="legacy duplicate content",
                country_code="GB",
                country="United Kingdom",
                source_filename="GB-legacy.docx",
            )

            with self.assertRaises(AdminDocumentCountryConflictError):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="content",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

    def test_no_conflict_when_only_one_active_document_exists(
        self,
    ) -> None:
        # Sanity check: the conflict check itself must never
        # false-positive on the document's own single, real chunk set.
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [("Employment Contracts", "Original DOCX text.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="applies normally",
                    source_directory=source_directory,
                    client=client,
                )

            response = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(response.content, "applies normally")

    def test_country_lookup_failure_is_never_a_raw_exception(
        self,
    ) -> None:
        # A transient failure while checking for a country conflict
        # must surface through the same structured error contract as
        # every other failure point - never an unwrapped exception
        # type the router has no mapping for.
        from app.services.admin_document_lifecycle import (
            AdminDocumentLifecycleError,
        )

        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [("Employment Contracts", "Original DOCX text.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with patch(
                "app.services.admin_document_lifecycle."
                "lookup_existing_country_documents",
                side_effect=RuntimeError("simulated transient failure"),
            ):
                with self.assertRaises(AdminDocumentLifecycleError):
                    update_effective_section(
                        document_id=DOCUMENT_ID,
                        section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                        new_content="should not apply",
                        source_directory=source_directory,
                        client=client,
                    )


class AdminSectionBoldOnlyDocumentAddTests(unittest.TestCase):
    """
    ORDER 8A-C main objective: Add must work end-to-end through the
    full service layer on a document whose native topics use only
    bold formatting - previously unsupported before the internal
    DOCX-native style marker existed.
    """

    def test_add_list_get_edit_on_bold_only_document(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_bold_only_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            client = _seeded_client(
                topics=[("Hiring Practices", "placeholder")]
            )

            with _patched_indexer(client):
                add_result = add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="Employees may work remotely. MARKER.",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(add_result.legal_topic, "Remote Working")

            listing = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                [s.legal_topic for s in listing.sections],
                ["Hiring Practices", "Remote Working"],
            )

            remote_section_id = section_id_for_legal_topic(
                "Remote Working"
            )

            fetched = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=remote_section_id,
                source_directory=source_directory,
                client=client,
            )
            self.assertIn("MARKER", fetched.content)

            with _patched_indexer(client):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=remote_section_id,
                    new_content="Overwritten remote content.",
                    source_directory=source_directory,
                    client=client,
                )

            fetched_after_edit = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=remote_section_id,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                fetched_after_edit.content, "Overwritten remote content."
            )
            self.assertNotIn("MARKER", fetched_after_edit.content)

            # native topic on the bold-only document remains untouched
            hiring_fetched = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=HIRING_PRACTICES_SECTION_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(hiring_fetched.content, "HP content.")

    def test_reindex_preserves_custom_section_on_bold_only_document(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_bold_only_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            client = _seeded_client(
                topics=[("Hiring Practices", "placeholder")]
            )

            with _patched_indexer(client):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="Remote content. MARKER.",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )

            with patch(
                "app.services.document_indexer.ensure_legal_documents_index"
            ), patch(
                "app.services.document_indexer.bulk",
                side_effect=_bulk_writer(client),
            ):
                reindex_indexed_document(
                    document_id=DOCUMENT_ID,
                    source_directory=source_directory,
                    client=client,
                )

            listing = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertIn(
                "Remote Working",
                [s.legal_topic for s in listing.sections],
            )

    def test_duplicate_custom_title_rejected_on_bold_only_document(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_bold_only_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            client = _seeded_client(
                topics=[("Hiring Practices", "placeholder")]
            )

            with self.assertRaises(
                AdminDocumentSectionAlreadyExistsError
            ):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="hiring   PRACTICES",
                    content="whatever",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )


class AdminDocumentSectionRenameTests(unittest.TestCase):
    """
    Mission "ORDER 8G-A", sections 2-5: Rename extends the same Edit
    Section transaction - an omitted or effectively-unchanged title is
    a normal content-only edit, never a fake rename; a genuine title
    change re-validates the reparsed document (old title gone, new
    title exactly once, topic count/other topics unchanged) before
    touching OpenSearch, then removes the old topic's chunks once the
    new ones are safely indexed, atomically replaces the source, and
    verifies both invariants (new topic present, old topic absent).
    """

    def test_native_canonical_section_rename_with_new_content(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            # A second, untouched canonical topic must remain so the
            # reparse can still LEARN the document's own heading-level-1
            # structural signal (docx_parser._learn_custom_topic_signal_
            # requirement) to recognize the renamed heading afterward -
            # a document with only one canonical topic being renamed
            # away has nothing left to learn from by construction.
            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "Original DOCX text."),
                    ("Hiring Practices", "Unrelated HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                response = update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="Renamed section content.",
                    new_title="Remote Work Equipment Requirements",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                response.legal_topic,
                "Remote Work Equipment Requirements",
            )
            self.assertEqual(
                response.section_id,
                section_id_for_legal_topic(
                    "Remote Work Equipment Requirements"
                ),
            )

            paragraphs = [
                p.text for p in Document(source_path).paragraphs
            ]
            self.assertIn(
                "Remote Work Equipment Requirements", paragraphs
            )
            self.assertNotIn("Employment Contracts", paragraphs)
            self.assertIn("Renamed section content.", paragraphs)

            remaining_new = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"]
                == "Remote Work Equipment Requirements"
            ]
            self.assertEqual(len(remaining_new), 1)

            remaining_old = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Employment Contracts"
            ]
            self.assertEqual(remaining_old, [])

    def test_custom_section_rename_title_only_content_preserved(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "EC content.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder EC")]
            )

            with _patched_indexer(client):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="Original custom content, unchanged.",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )

                response = update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=section_id_for_legal_topic(
                        "Remote Working"
                    ),
                    new_content="Original custom content, unchanged.",
                    new_title="Remote Working Policy",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                response.legal_topic, "Remote Working Policy"
            )

            reparsed = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=section_id_for_legal_topic(
                    "Remote Working Policy"
                ),
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(
                reparsed.content,
                "Original custom content, unchanged.",
            )

    def test_duplicate_title_is_rejected_with_zero_mutation(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with self.assertRaises(
                AdminDocumentSectionAlreadyExistsError
            ) as context:
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="whatever",
                    new_title="Hiring Practices",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                context.exception.to_detail()["operation"],
                "section_update",
            )
            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )
            self.assertEqual(len(client.chunks), 2)

    def test_normalized_duplicate_title_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with self.assertRaises(
                AdminDocumentSectionAlreadyExistsError
            ):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="whatever",
                    new_title="  hiring   PRACTICES  ",
                    source_directory=source_directory,
                    client=client,
                )

    def test_unchanged_effective_title_is_a_normal_edit_not_a_rename(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "Original DOCX text.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with _patched_indexer(client):
                response = update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="Edited content only.",
                    new_title="  employment   contracts  ",
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                response.legal_topic, "Employment Contracts"
            )
            self.assertEqual(
                response.section_id,
                EMPLOYMENT_CONTRACTS_SECTION_ID,
            )

            paragraphs = [
                p.text for p in Document(source_path).paragraphs
            ]
            self.assertIn("Employment Contracts", paragraphs)

    def test_stale_section_id_after_rename_is_not_found(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "Unrelated HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="Renamed away.",
                    new_title="Something Else Entirely",
                    source_directory=source_directory,
                    client=client,
                )

                with self.assertRaises(
                    AdminDocumentSectionNotFoundError
                ):
                    update_effective_section(
                        document_id=DOCUMENT_ID,
                        section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                        new_content="should not apply",
                        source_directory=source_directory,
                        client=client,
                    )

    def test_other_topics_unchanged_by_rename(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="Renamed content.",
                    new_title="Something Else Entirely",
                    source_directory=source_directory,
                    client=client,
                )

            hp_chunks = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Hiring Practices"
            ]
            self.assertEqual(len(hp_chunks), 1)
            self.assertEqual(hp_chunks[0]["content"], "placeholder HP")

            paragraphs = [
                p.text for p in Document(source_path).paragraphs
            ]
            self.assertIn("HP content.", paragraphs)

    def test_new_topic_index_failure_rolls_back_rename(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "Original DOCX text."),
                    ("Hiring Practices", "Unrelated HP content."),
                ],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(
                    AdminDocumentSectionUpdateFailedError
                ):
                    update_effective_section(
                        document_id=DOCUMENT_ID,
                        section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                        new_content="should roll back",
                        new_title="Something Else Entirely",
                        source_directory=source_directory,
                        client=client,
                    )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )
            self.assertEqual(
                client.chunks["chunk-seed-0"]["legal_topic"],
                "Employment Contracts",
            )
            self.assertEqual(
                client.chunks["chunk-seed-0"]["content"],
                "placeholder",
            )

            leftover_temps = list(
                source_directory.glob(".*tmp.docx")
            )
            self.assertEqual(leftover_temps, [])

    def test_old_topic_cleanup_failure_rolls_back_new_topic_too(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "Original DOCX text."),
                    ("Hiring Practices", "Unrelated HP content."),
                ],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            # Fail specifically the SECOND delete_by_query call - the
            # first belongs to replace_document_section_chunks' own
            # internal staleness cleanup for the NEW topic (which must
            # succeed for this scenario to actually reach the OLD-
            # topic cleanup step this test targets).
            real_delete_by_query = client.delete_by_query
            call_count = {"n": 0}

            def delete_by_query_fail_second(**kwargs: Any) -> Any:
                call_count["n"] += 1

                if call_count["n"] == 2:
                    raise OpenSearchException(
                        "simulated old-topic cleanup failure"
                    )

                return real_delete_by_query(**kwargs)

            client.delete_by_query = delete_by_query_fail_second

            with _patched_indexer(client):
                with self.assertRaises(
                    AdminDocumentSectionUpdateFailedError
                ):
                    update_effective_section(
                        document_id=DOCUMENT_ID,
                        section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                        new_content="should roll back",
                        new_title="Something Else Entirely",
                        source_directory=source_directory,
                        client=client,
                    )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Employment Contracts"
            ]
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]["content"], "placeholder")

            leftover_new = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Something Else Entirely"
            ]
            self.assertEqual(leftover_new, [])

    def test_docx_replace_failure_rolls_back_rename_entirely(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "Original DOCX text."),
                    ("Hiring Practices", "Unrelated HP content."),
                ],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                with patch(
                    "app.services.admin_document_sections.os.replace",
                    side_effect=OSError("simulated disk failure"),
                ):
                    with self.assertRaises(
                        AdminDocumentSectionUpdateFailedError
                    ):
                        update_effective_section(
                            document_id=DOCUMENT_ID,
                            section_id=(
                                EMPLOYMENT_CONTRACTS_SECTION_ID
                            ),
                            new_content="should roll back",
                            new_title="Something Else Entirely",
                            source_directory=source_directory,
                            client=client,
                        )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Employment Contracts"
            ]
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]["content"], "placeholder")

            leftover_new = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Something Else Entirely"
            ]
            self.assertEqual(leftover_new, [])

            leftover_temps = list(
                source_directory.glob(".*tmp.docx")
            )
            self.assertEqual(leftover_temps, [])

    def test_live_topic_vocabulary_updates_immediately_after_rename(
        self,
    ) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [
                    ("Employment Contracts", "Original DOCX text."),
                    ("Hiring Practices", "Unrelated HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                update_effective_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    new_content="Renamed content.",
                    new_title="Something Else Entirely",
                    source_directory=source_directory,
                    client=client,
                )

            listing = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            listed_topics = [s.legal_topic for s in listing.sections]

            self.assertIn("Something Else Entirely", listed_topics)
            self.assertNotIn("Employment Contracts", listed_topics)


class AdminDocumentSectionDeleteTests(unittest.TestCase):
    """
    Mission "ORDER 8G-A", sections 6-8: Delete mirrors Edit/Rename's
    exact transaction shape for one already-known legal_topic (no new
    chunks are ever built) - lock, snapshot, mutate a temp copy,
    reparse-validate (target gone, every other topic unchanged),
    delete the target's chunks, atomically replace the source, verify.
    Blocks deleting the document's last remaining usable section.
    """

    def test_custom_section_delete(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "EC content.")],
            )

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder EC")]
            )

            with _patched_indexer(client):
                add_new_section(
                    document_id=DOCUMENT_ID,
                    title="Remote Working",
                    content="Custom content.",
                    position="end",
                    source_directory=source_directory,
                    client=client,
                )

                response = delete_section(
                    document_id=DOCUMENT_ID,
                    section_id=section_id_for_legal_topic(
                        "Remote Working"
                    ),
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(response.legal_topic, "Remote Working")

            paragraphs = [
                p.text for p in Document(source_path).paragraphs
            ]
            self.assertNotIn("Remote Working", paragraphs)
            self.assertNotIn("Custom content.", paragraphs)
            self.assertIn("EC content.", paragraphs)

            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Remote Working"
            ]
            self.assertEqual(remaining, [])

    def test_native_canonical_section_delete(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                response = delete_section(
                    document_id=DOCUMENT_ID,
                    section_id=HIRING_PRACTICES_SECTION_ID,
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(response.legal_topic, "Hiring Practices")

            paragraphs = [
                p.text for p in Document(source_path).paragraphs
            ]
            self.assertNotIn("Hiring Practices", paragraphs)
            self.assertNotIn("HP content.", paragraphs)
            self.assertIn("Employment Contracts", paragraphs)
            self.assertIn("EC content.", paragraphs)

            hp_chunks = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Hiring Practices"
            ]
            self.assertEqual(hp_chunks, [])

            ec_chunks = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Employment Contracts"
            ]
            self.assertEqual(len(ec_chunks), 1)

    def test_last_remaining_section_delete_is_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [("Employment Contracts", "EC content.")],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[("Employment Contracts", "placeholder")]
            )

            with self.assertRaises(
                AdminDocumentSectionLastRemainingError
            ) as context:
                delete_section(
                    document_id=DOCUMENT_ID,
                    section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                    source_directory=source_directory,
                    client=client,
                )

            self.assertEqual(
                context.exception.to_detail()["code"],
                "section_is_last_remaining",
            )
            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )
            self.assertEqual(len(client.chunks), 1)

    def test_other_topics_unchanged_by_delete(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                delete_section(
                    document_id=DOCUMENT_ID,
                    section_id=HIRING_PRACTICES_SECTION_ID,
                    source_directory=source_directory,
                    client=client,
                )

            # get_effective_section reads the CURRENT DOCX directly,
            # never OpenSearch - "placeholder EC" is only the fake
            # index's own seed value.
            reparsed = get_effective_section(
                document_id=DOCUMENT_ID,
                section_id=EMPLOYMENT_CONTRACTS_SECTION_ID,
                source_directory=source_directory,
                client=client,
            )
            self.assertEqual(reparsed.content, "EC content.")

    def test_opensearch_delete_failure_rolls_back(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )
            client.fail_delete_by_query_calls = 1
            client.delete_by_query_failure = OpenSearchException(
                "simulated delete failure"
            )

            with _patched_indexer(client):
                with self.assertRaises(
                    AdminDocumentSectionUpdateFailedError
                ):
                    delete_section(
                        document_id=DOCUMENT_ID,
                        section_id=HIRING_PRACTICES_SECTION_ID,
                        source_directory=source_directory,
                        client=client,
                    )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Hiring Practices"
            ]
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]["content"], "placeholder HP")

    def test_docx_replace_failure_rolls_back_delete(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / "GB.docx"

            _write_docx(
                source_path,
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )
            original_bytes = source_path.read_bytes()

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                with patch(
                    "app.services.admin_document_sections.os.replace",
                    side_effect=OSError("simulated disk failure"),
                ):
                    with self.assertRaises(
                        AdminDocumentSectionUpdateFailedError
                    ):
                        delete_section(
                            document_id=DOCUMENT_ID,
                            section_id=HIRING_PRACTICES_SECTION_ID,
                            source_directory=source_directory,
                            client=client,
                        )

            self.assertEqual(
                source_path.read_bytes(), original_bytes
            )

            remaining = [
                chunk
                for chunk in client.chunks.values()
                if chunk["legal_topic"] == "Hiring Practices"
            ]
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]["content"], "placeholder HP")

            leftover_temps = list(
                source_directory.glob(".*tmp.docx")
            )
            self.assertEqual(leftover_temps, [])

    def test_live_topic_vocabulary_removes_deleted_title(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)

            _write_docx(
                source_directory / "GB.docx",
                [
                    ("Employment Contracts", "EC content."),
                    ("Hiring Practices", "HP content."),
                ],
            )

            client = _seeded_client(
                topics=[
                    ("Employment Contracts", "placeholder EC"),
                    ("Hiring Practices", "placeholder HP"),
                ]
            )

            with _patched_indexer(client):
                delete_section(
                    document_id=DOCUMENT_ID,
                    section_id=HIRING_PRACTICES_SECTION_ID,
                    source_directory=source_directory,
                    client=client,
                )

            listing = list_effective_sections(
                document_id=DOCUMENT_ID,
                source_directory=source_directory,
                client=client,
            )
            listed_topics = [s.legal_topic for s in listing.sections]

            self.assertNotIn("Hiring Practices", listed_topics)
            self.assertIn("Employment Contracts", listed_topics)


if __name__ == "__main__":
    unittest.main()
