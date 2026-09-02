"""Consolidated test module generated from validated domain owners."""
from __future__ import annotations
import tempfile
import unittest
from io import BytesIO
from pathlib import Path
from typing import Any
from app.core.admin_country_policy import ADMIN_ALLOWED_COUNTRY_CODES, is_admin_country_allowed
from app.core.country_registry import COUNTRIES, canonical_country_name
from app.models.document import DocumentChunk
from app.security.admin import admin_key_matches
from app.services.admin_document_replacement import InvalidDocumentUploadError, _sanitize_filename, list_indexed_documents
from app.services.admin_document_replacement import AdminDocumentAlreadyCurrentError, AdminDocumentCountryConfirmationRequiredError, AdminDocumentCountryConflictReviewRequiredError, AdminDocumentCountryNotAllowedError, AdminDocumentCountrySelectionInvalidError, AdminDocumentCountrySelectionRequiredError, AdminDocumentIdenticalButAdminModifiedError, AdminDocumentReplacementRequiredError, AdminDocumentWarningConfirmationRequiredError, ExistingCountryDocument, safe_upload_and_index_document
from app.services.admin_document_replacement import DocumentCountryUndeterminedError
from app.services.admin_document_lifecycle import AdminDocumentNotFoundError, AdminDocumentSourceConflictError, AdminDocumentSourceMissingError, delete_indexed_document, reindex_indexed_document
from app.services.document_chunk_builder import read_country_marker
from app.services.document_chunk_builder import DOCUMENT_FAMILY, AmbiguousDocumentCountryError, UndeterminableDocumentCountryError, build_document_chunks_from_docx, metadata_from_content
from app.services.document_indexer import DocumentIndexingError, DocumentIndexingResult
from app.services.document_section_state import SectionEdit, SectionEditState, read_section_edit_state, section_id_for_legal_topic as _test_admin_documents__section_id_for_legal_topic, write_section_edit_state_atomic
from tests.support.documents import real_source_entries
from tests.support.opensearch import FakeAdminOpenSearch
from docx import Document as DocxDocument
from unittest.mock import patch

def _build_chunk(filename: str) -> DocumentChunk:
    """Build one valid document chunk."""
    return DocumentChunk(document_id='document-1', chunk_id='chunk-1', country='United Kingdom', country_code='GB', legal_topic='Employment Contracts', document_type='comparator', language='en', section='Employment Contracts', subsection='Notice Period', content='One week of notice may apply.', source_filename=filename, source_format='docx', content_hash='content-hash', reference_year=2026)

def _catalog_fake(*, chunk_count: int=1) -> FakeAdminOpenSearch:
    """
    Seed a shared FakeAdminOpenSearch with exactly one indexed
    document ("document-1", GB, "UK 2026.docx") reporting chunk_count
    chunks - the fixed single-document catalog shape every test below
    exercises. Deliberately keyed by document-1/GB/"UK 2026.docx",
    never the on-disk storage filename, so a mismatch between
    OpenSearch's own display source_filename and the physical
    {COUNTRY_CODE}.docx storage path is exactly what these tests
    prove list_indexed_documents resolves through.
    """
    fake = FakeAdminOpenSearch()
    for index in range(chunk_count):
        fake.add(document_id='document-1', country_code='GB', country='United Kingdom', source_filename='UK 2026.docx', chunk_id=f'document-1-chunk-{index}', document_type='comparator', reference_year=2026)
    return fake

class AdminDocumentTests(unittest.TestCase):
    """Tests for administration services."""

    def test_admin_key_matching(self) -> None:
        self.assertTrue(admin_key_matches('admin-secret', 'admin-secret'))
        self.assertFalse(admin_key_matches('wrong', 'admin-secret'))
        self.assertFalse(admin_key_matches(None, 'admin-secret'))

    def test_indexed_documents_are_listed(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'GB.docx').write_bytes(b'document')
            response = list_indexed_documents(source_directory=source_directory, client=_catalog_fake(chunk_count=41))
            self.assertEqual(response.total, 1)
            self.assertEqual(response.documents[0].country_code, 'GB')
            self.assertEqual(response.documents[0].source_filename, 'UK 2026.docx')
            self.assertEqual(response.documents[0].chunk_count, 41)
            self.assertEqual(response.documents[0].status, 'indexed')

    def test_indexed_document_missing_from_disk_is_flagged(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            response = list_indexed_documents(source_directory=source_directory, client=_catalog_fake())
            self.assertEqual(response.documents[0].status, 'indexed_source_missing')

    def test_legacy_document_with_historical_filename_shows_source_available(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'UK 2026.docx').write_bytes(b'legacy-document-bytes')
            response = list_indexed_documents(source_directory=source_directory, client=_catalog_fake())
            self.assertEqual(response.documents[0].status, 'indexed')
            self.assertTrue(response.documents[0].source_file_present)
            self.assertFalse((source_directory / 'GB.docx').exists())

    def test_conflicting_sources_are_flagged_not_guessed(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'UK 2026.docx').write_bytes(b'legacy-bytes')
            (source_directory / 'GB.docx').write_bytes(b'canonical-bytes')
            response = list_indexed_documents(source_directory=source_directory, client=_catalog_fake())
            self.assertEqual(response.documents[0].status, 'indexed_source_conflict')
            self.assertFalse(response.documents[0].source_file_present)

class FilenameAcceptanceTests(unittest.TestCase):
    """
    Arbitrary safe DOCX filenames are accepted verbatim: no business
    naming format is required at all, only basic safety checks.
    """
    ACCEPTED_FILENAMES = ('Canada_2026-04-15-Employment-Law-Overview-EDITED.docx', 'final.docx', 'Canada final version.docx', 'document_received_from_client.docx', 'Version corrigée (3).DOCX', 'fichier client été 2026.docx', 'Spain-template-used-for-Canada.docx')

    def test_all_example_filenames_are_accepted_verbatim(self) -> None:
        for filename in self.ACCEPTED_FILENAMES:
            with self.subTest(filename=filename):
                self.assertEqual(_sanitize_filename(filename), filename)

class FilenameRejectionTests(unittest.TestCase):
    """
    Only safety properties are checked, never a business naming
    format: these filenames must still be rejected for real safety
    reasons (path traversal, wrong extension, null byte, empty name).
    """
    REJECTED_FILENAMES = ('../../document.docx', '../document.docx', 'folder/document.docx', 'folder\\document.docx', 'document.pdf', 'document.docx.exe', 'document.docm', '', 'document\x00.docx')

    def test_all_example_filenames_are_rejected(self) -> None:
        for filename in self.REJECTED_FILENAMES:
            with self.subTest(filename=filename):
                with self.assertRaises(InvalidDocumentUploadError):
                    _sanitize_filename(filename)

    def test_case_insensitive_docx_extension_is_still_accepted(self) -> None:
        self.assertEqual(_sanitize_filename('Report.DOCX'), 'Report.DOCX')
CANADA_DOCUMENT_ID = 'doc_' + 'd' * 64

def _build_canada_chunk(*, filename: str, reference_year: int) -> DocumentChunk:
    """
    One Canada/employment-law-overview chunk carrying the same fixed
    document_id regardless of year or filename - exactly how the real
    country_code+family identity scheme (document_chunk_builder.
    _build_document_id) behaves once a country is already active.
    """
    return DocumentChunk(document_id=CANADA_DOCUMENT_ID, chunk_id='chunk-ca-1', country='Canada', country_code='CA', legal_topic='Employment Contracts', document_type='comparator', language='en', section='Employment Contracts', subsection='Notice Period', content='Notice content.', source_filename=filename, source_format='docx', content_hash='content-hash', reference_year=reference_year)
SPAIN_DOCUMENT_ID = 'doc_' + 'e' * 64

def _build_spain_chunk(*, filename: str) -> DocumentChunk:
    """One Spain/employment-law-overview chunk, fixed document_id."""
    return DocumentChunk(document_id=SPAIN_DOCUMENT_ID, chunk_id='chunk-es-1', country='Spain', country_code='ES', legal_topic='Employment Contracts', document_type='comparator', language='en', section='Employment Contracts', subsection='Notice Period', content='Notice content.', source_filename=filename, source_format='docx', content_hash='content-hash', reference_year=2027)

class LegacyReplacementAndRollbackTests(unittest.TestCase):
    """
    Replace and Rollback for a
    country whose currently active document is still stored under its
    historical filename, never {COUNTRY_CODE}.docx.
    """
    LEGACY_FILENAME = 'Labour and Employment Law in Spain 2026.docx'

def _real_docx_bytes(paragraphs: list[str]) -> bytes:
    """A minimal real DOCX (valid zip/OOXML), for tests that exercise
    the actual marker-writing/parsing pipeline rather than a fake
    chunk_builder."""
    document = DocxDocument()
    for text in paragraphs:
        document.add_paragraph(text)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()

def _fake_indexer(*, chunks, client=None):
    del client
    return DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
AU_OLD_ID = 'doc_' + 'a' * 64
AU_NEW_ID = 'doc_' + 'b' * 64

def _build_au_chunk(filename: str) -> DocumentChunk:
    return DocumentChunk(document_id=AU_NEW_ID, chunk_id='chunk-au-new-1', country='Australia', country_code='AU', legal_topic='Employment Contracts', document_type='comparator', language='en', section='Employment Contracts', subsection='Trial Period', content='Australian probation content.', source_filename=filename, source_format='docx', content_hash='new-content-hash', reference_year=2026)

def _existing_documents(country_code: str, client=None) -> list[ExistingCountryDocument]:
    del client
    assert country_code == 'AU'
    return [ExistingCountryDocument(document_id=AU_OLD_ID, source_filename='Employment Law Overview Australia.docx', country='Australia', country_code='AU', reference_year=None)]

class SafeCountryReplacementTests(unittest.TestCase):

    def test_existing_country_requires_confirmation_without_changes(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            source_path = source_directory / 'Employment Law Overview Australia.docx'
            source_path.write_bytes(b'old-australia')
            indexer_called = False

            def country_indexer(**kwargs):
                nonlocal indexer_called
                del kwargs
                indexer_called = True
                raise AssertionError('indexer must not run')
            with self.assertRaises(AdminDocumentReplacementRequiredError) as context:
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'new-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=country_indexer)
            self.assertEqual(context.exception.country_code, 'AU')
            self.assertFalse(indexer_called)
            self.assertEqual(source_path.read_bytes(), b'old-australia')
            self.assertFalse((source_directory / 'AU.docx').exists())

    def test_replace_existing_never_bypasses_a_country_conflict(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            legacy_path = source_directory / 'Employment Law Overview Australia.docx'
            canonical_path = source_directory / 'AU.docx'
            legacy_path.write_bytes(b'legacy-australia')
            canonical_path.write_bytes(b'other-australia')
            existing = [ExistingCountryDocument(document_id=AU_OLD_ID, source_filename=legacy_path.name, country='Australia', country_code='AU', reference_year=None), ExistingCountryDocument(document_id='doc_' + 'c' * 64, source_filename=legacy_path.name, country='Australia', country_code='AU', reference_year=None)]

            def lookup(country_code: str, client=None):
                del client
                self.assertEqual(country_code, 'AU')
                return existing

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError('indexer must not run')
            with self.assertRaises(AdminDocumentCountryConflictReviewRequiredError) as context:
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'new-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, replace_existing=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lookup, country_document_indexer=indexer_must_not_be_called)
            self.assertEqual(context.exception.country_code, 'AU')
            self.assertEqual(sorted((candidate.document_id for candidate in context.exception.candidates)), sorted([AU_OLD_ID, 'doc_' + 'c' * 64]))
            self.assertEqual(legacy_path.read_bytes(), b'legacy-australia')
            self.assertEqual(canonical_path.read_bytes(), b'other-australia')

    def test_identical_single_document_is_not_reindexed(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'same-australia')

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError('the indexer must not be called for an already-current document')
            with self.assertRaises(AdminDocumentAlreadyCurrentError):
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'same-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=indexer_must_not_be_called)

    def test_fresh_country_is_staged_and_parsed_exactly_once(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            chunk_builder_calls = []

            def counting_chunk_builder(path: Path):
                chunk_builder_calls.append(path)
                return [_build_au_chunk(path.name)]
            stream = BytesIO(b'new-argentina')
            original_read = stream.read
            read_calls = []

            def counting_read(*args, **kwargs):
                read_calls.append(1)
                return original_read(*args, **kwargs)
            stream.read = counting_read
            response = safe_upload_and_index_document(filename='Argentina.docx', file_stream=stream, source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=counting_chunk_builder, country_document_lookup=lambda code, client: [], country_document_indexer=lambda *, chunks, client=None: DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0))
            self.assertEqual(len(chunk_builder_calls), 1, 'chunk_builder must run exactly once for a fresh upload - never a second, redundant parse.')
            self.assertLessEqual(len(read_calls), 2)
            self.assertEqual(response.status, 'uploaded')
            self.assertEqual(response.document_id, AU_NEW_ID)
            self.assertEqual(response.country_code, 'AU')
            self.assertFalse(response.replaced_source_file)
            self.assertEqual(response.replaced_document_ids, [])
            self.assertTrue((source_directory / 'AU.docx').exists())
            self.assertEqual((source_directory / 'AU.docx').read_bytes(), b'new-argentina')

    def test_deleted_then_reuploaded_country_succeeds(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'

            def indexer(*, chunks, client=None):
                return DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
            first = safe_upload_and_index_document(filename='Argentina.docx', file_stream=BytesIO(b'first-argentina'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: [], country_document_indexer=indexer)
            self.assertEqual(first.status, 'uploaded')
            (source_directory / 'AU.docx').unlink()
            second = safe_upload_and_index_document(filename='random-name-completely-different.docx', file_stream=BytesIO(b'second-argentina'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: [], country_document_indexer=indexer)
            self.assertEqual(second.status, 'uploaded')
            self.assertEqual([path.name for path in real_source_entries(source_directory)], ['AU.docx'])
            self.assertEqual((source_directory / 'AU.docx').read_bytes(), b'second-argentina')

class FilenameNeverDecidesIdentityTests(unittest.TestCase):
    """
    The country_code detected from a DOCX's own content is the only
    identity that matters; the
    uploaded filename is display information only. Three completely
    different filenames carrying the exact same detected country must
    all resolve to the exact same existing-country/409 outcome.
    """

    def test_same_filename_as_the_existing_document_still_requires_confirmation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'old-australia')
            with self.assertRaises(AdminDocumentReplacementRequiredError):
                safe_upload_and_index_document(filename='Employment Law Overview Australia.docx', file_stream=BytesIO(b'new-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents)

    def test_unrelated_filenames_all_detect_the_same_country(self) -> None:
        for filename in ('Argentina.docx', 'random-file-name.docx', 'Legal-update-final-v7.docx'):
            with self.subTest(filename=filename):
                with tempfile.TemporaryDirectory() as root:
                    source_directory = Path(root) / 'source'
                    processed_directory = Path(root) / 'processed'
                    response = safe_upload_and_index_document(filename=filename, file_stream=BytesIO(b'argentina-bytes'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: [], country_document_indexer=lambda *, chunks, client=None: DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0))
                    self.assertEqual(response.status, 'uploaded')
                    self.assertEqual(response.country_code, 'AU')
                    self.assertTrue((source_directory / 'AU.docx').exists())

class ConfirmedReplacementDuplicateCountTests(unittest.TestCase):
    """
    The ordinary upload/replace decision must refuse to collapse a
    country conflict itself,
    however many duplicate legacy document_ids and source files it
    started with (tested here for 3) - it always defers to the
    dedicated conflict-resolution API instead, never collapsing a
    conflict to one document on its own.
    """

    def test_confirmed_replace_with_three_duplicate_ids_requires_conflict_review(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            legacy_path = source_directory / 'Employment Law Overview Australia.docx'
            canonical_path = source_directory / 'AU.docx'
            legacy_path.write_bytes(b'legacy-australia')
            canonical_path.write_bytes(b'other-australia')
            existing = [ExistingCountryDocument(document_id='doc_' + letter * 64, source_filename=legacy_path.name, country='Australia', country_code='AU', reference_year=None) for letter in ('a', 'c', 'd')]

            def lookup(country_code: str, client=None):
                del client
                self.assertEqual(country_code, 'AU')
                return existing

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError('indexer must not run')
            with self.assertRaises(AdminDocumentCountryConflictReviewRequiredError) as context:
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'new-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, replace_existing=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lookup, country_document_indexer=indexer_must_not_be_called)
            self.assertEqual(sorted((candidate.document_id for candidate in context.exception.candidates)), sorted((document.document_id for document in existing)))
            self.assertEqual(sorted((path.name for path in real_source_entries(source_directory))), ['AU.docx', 'Employment Law Overview Australia.docx'])

    def test_confirmed_replace_with_source_missing(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            existing = [ExistingCountryDocument(document_id=AU_OLD_ID, source_filename='Employment Law Overview Australia.docx', country='Australia', country_code='AU', reference_year=None)]

            def indexer(*, chunks, client=None):
                return DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
            response = safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'new-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, replace_existing=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: existing, country_document_indexer=indexer)
            self.assertEqual(response.status, 'replaced')
            self.assertFalse(response.replaced_source_file)
            self.assertEqual(response.replaced_document_ids, [AU_OLD_ID])
            self.assertEqual((source_directory / 'AU.docx').read_bytes(), b'new-australia')

class IdenticalFileAmbiguousCountryTests(unittest.TestCase):
    """
    A country left with several document_ids (a pre-existing
    conflict/duplicate state) must never
    be told "document_already_current", even if the newly uploaded
    bytes happen to match one of the candidate files exactly: the
    country's own state is inconsistent, and normalizing it now always
    requires the dedicated conflict-resolution review, never a plain
    confirmed replacement.
    """

    def test_identical_bytes_but_ambiguous_country_still_requires_conflict_review(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            legacy_path = source_directory / 'Employment Law Overview Australia.docx'
            canonical_path = source_directory / 'AU.docx'
            legacy_path.write_bytes(b'same-australia')
            canonical_path.write_bytes(b'other-australia')
            existing = [ExistingCountryDocument(document_id=AU_OLD_ID, source_filename=legacy_path.name, country='Australia', country_code='AU', reference_year=None), ExistingCountryDocument(document_id='doc_' + 'c' * 64, source_filename=legacy_path.name, country='Australia', country_code='AU', reference_year=None)]
            with self.assertRaises(AdminDocumentCountryConflictReviewRequiredError):
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'same-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: existing)
            self.assertEqual(legacy_path.read_bytes(), b'same-australia')
            self.assertEqual(canonical_path.read_bytes(), b'other-australia')

def _build_chunks_with_topic_count(country_code: str, country: str, document_id: str, filename: str, topic_count: int) -> list[DocumentChunk]:
    from app.core.legal_taxonomy import LEGAL_TOPICS
    overview = DocumentChunk(document_id=document_id, chunk_id='chunk-overview', country=country, country_code=country_code, legal_topic=None, document_type='overview', language='en', section='General', subsection=None, content='Overview content.', source_filename=filename, source_format='docx', content_hash='overview-hash')
    topic_chunks = [DocumentChunk(document_id=document_id, chunk_id=f'chunk-topic-{index}', country=country, country_code=country_code, legal_topic=topic, document_type='comparator', language='en', section=topic, subsection=None, content=f'{topic} content.', source_filename=filename, source_format='docx', content_hash=f'hash-{index}') for index, topic in enumerate(LEGAL_TOPICS[:topic_count])]
    return [overview, *topic_chunks]

def _no_existing_documents(country_code: str, client=None) -> list[ExistingCountryDocument]:
    del country_code, client
    return []

class WarningConfirmationRequiredTests(unittest.TestCase):
    """14 - the warning confirmation
    gate, and its combination with a pending replacement decision."""

    def test_zero_recognized_topics_is_context_warning_end_to_end(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            with self.assertRaises(AdminDocumentWarningConfirmationRequiredError) as context:
                safe_upload_and_index_document(filename='Chile.docx', file_stream=BytesIO(b'chile-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=False, chunk_builder=lambda path: _build_chunks_with_topic_count('CL', 'Chile', 'doc_' + 'c' * 64, path.name, 0), country_document_lookup=_no_existing_documents, country_document_indexer=lambda **kwargs: (_ for _ in ()).throw(AssertionError('indexer must not run')))
            self.assertEqual(len(context.exception.warnings), 1)
            self.assertEqual(context.exception.warnings[0].code, 'context_warning')
            self.assertEqual(context.exception.warnings[0].recognized_topics_count, 0)

            def country_indexer(*, chunks, client):
                del client
                return DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
            response = safe_upload_and_index_document(filename='Chile.docx', file_stream=BytesIO(b'chile-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=lambda path: _build_chunks_with_topic_count('CL', 'Chile', 'doc_' + 'c' * 64, path.name, 0), country_document_lookup=_no_existing_documents, country_document_indexer=country_indexer)
            self.assertEqual(response.status, 'uploaded')

    def test_fresh_country_with_thin_coverage_requires_confirmation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            indexer_called = False

            def country_indexer(**kwargs):
                nonlocal indexer_called
                del kwargs
                indexer_called = True
                raise AssertionError('indexer must not run')
            with self.assertRaises(AdminDocumentWarningConfirmationRequiredError) as context:
                safe_upload_and_index_document(filename='Chile.docx', file_stream=BytesIO(b'chile-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=False, chunk_builder=lambda path: _build_chunks_with_topic_count('CL', 'Chile', 'doc_' + 'c' * 64, path.name, 2), country_document_lookup=_no_existing_documents, country_document_indexer=country_indexer)
            error = context.exception
            self.assertEqual(error.country_code, 'CL')
            self.assertFalse(error.replacement_required)
            self.assertEqual(error.existing_document_ids, ())
            self.assertEqual(len(error.warnings), 1)
            self.assertEqual(error.warnings[0].recognized_topics_count, 2)
            self.assertFalse(indexer_called)
            self.assertEqual(list(real_source_entries(source_directory)), [])
            detail = error.to_detail()
            self.assertEqual(detail['code'], 'document_warning_confirmation_required')
            self.assertEqual(detail['operation'], 'upload')
            self.assertEqual(detail['country_code'], 'CL')
            self.assertEqual(detail['country_name'], 'Chile')
            self.assertFalse(detail['replacement_required'])
            self.assertEqual(detail['existing_document_ids'], [])
            self.assertEqual(len(detail['warnings']), 1)
            warning_detail = detail['warnings'][0]
            self.assertEqual(warning_detail['recognized_topics_count'], 2)
            self.assertEqual(warning_detail['expected_topics_count'], 11)
            self.assertEqual(len(warning_detail['missing_topics']), 9)

    def test_confirm_warnings_true_proceeds_to_index(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            indexed_chunks = []

            def country_indexer(*, chunks, client):
                del client
                indexed_chunks.extend(chunks)
                return DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
            response = safe_upload_and_index_document(filename='Chile.docx', file_stream=BytesIO(b'chile-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=lambda path: _build_chunks_with_topic_count('CL', 'Chile', 'doc_' + 'c' * 64, path.name, 2), country_document_lookup=_no_existing_documents, country_document_indexer=country_indexer)
            self.assertEqual(response.status, 'uploaded')
            self.assertEqual(response.country_code, 'CL')
            self.assertEqual(len(indexed_chunks), 3)

    def test_six_recognized_topics_never_needs_confirmation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'

            def country_indexer(*, chunks, client):
                del client
                return DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
            response = safe_upload_and_index_document(filename='Chile.docx', file_stream=BytesIO(b'chile-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=False, chunk_builder=lambda path: _build_chunks_with_topic_count('CL', 'Chile', 'doc_' + 'c' * 64, path.name, 6), country_document_lookup=_no_existing_documents, country_document_indexer=country_indexer)
            self.assertEqual(response.status, 'uploaded')

    def test_warning_and_replacement_both_pending_are_combined(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'AU.docx').write_bytes(b'existing-australia-bytes')
            with self.assertRaises(AdminDocumentWarningConfirmationRequiredError) as context:
                safe_upload_and_index_document(filename='Australia-new.docx', file_stream=BytesIO(b'new-different-australia'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, replace_existing=False, confirm_warnings=False, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=lambda **kwargs: (_ for _ in ()).throw(AssertionError('indexer must not run')))
            error = context.exception
            self.assertTrue(error.replacement_required)
            self.assertEqual(error.existing_document_ids, (AU_OLD_ID,))
            self.assertEqual(len(error.warnings), 1)
            self.assertEqual((source_directory / 'AU.docx').read_bytes(), b'existing-australia-bytes')

    def test_warnings_confirmed_falls_through_to_plain_replacement(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'AU.docx').write_bytes(b'existing-australia-bytes')
            with self.assertRaises(AdminDocumentReplacementRequiredError):
                safe_upload_and_index_document(filename='Australia-new.docx', file_stream=BytesIO(b'new-different-australia'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, replace_existing=False, confirm_warnings=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=lambda **kwargs: (_ for _ in ()).throw(AssertionError('indexer must not run')))

    def test_identical_bytes_wins_over_warnings(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'AU.docx').write_bytes(b'identical-bytes')

            def existing_thin_document(country_code, client=None):
                del client
                assert country_code == 'AU'
                return [ExistingCountryDocument(document_id=AU_OLD_ID, source_filename='AU.docx', country='Australia', country_code='AU', reference_year=None)]
            with self.assertRaises(AdminDocumentAlreadyCurrentError):
                safe_upload_and_index_document(filename='Australia-again.docx', file_stream=BytesIO(b'identical-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=False, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=existing_thin_document, country_document_indexer=lambda **kwargs: (_ for _ in ()).throw(AssertionError('indexer must not run')))

class AdminUploadAllowlistTests(unittest.TestCase):
    """
    The allowlist check runs after country detection/normalization
    but strictly before any mutation - no
    source commit, no OpenSearch write, no country_lock even acquired
    (the check sits inside the still-open staging TemporaryDirectory,
    before country_lock is entered) - so a rejected upload leaves
    everything exactly as it was.

    A perfectly-detected-but-disallowed country (Tunisia) must raise a
    distinct error/code (document_country_not_allowed) from an
    undeterminable one (document_country_undetermined) - the two must
    never be conflated.
    """

    def _reject_if_called(self, **kwargs):
        del kwargs
        raise AssertionError('indexer must not run')

    def test_disallowed_but_registered_country_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            with self.assertRaises(AdminDocumentCountryNotAllowedError) as context:
                safe_upload_and_index_document(filename='Tunisia.docx', file_stream=BytesIO(b'tunisia-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=lambda path: _build_chunks_with_topic_count('TN', 'Tunisia', 'doc_' + 'd' * 64, path.name, 2), country_document_lookup=_no_existing_documents, country_document_indexer=self._reject_if_called)
            detail = context.exception.to_detail()
            self.assertEqual(detail['code'], 'document_country_not_allowed')
            self.assertEqual(context.exception.country_code, 'TN')
            self.assertEqual(list(source_directory.iterdir()), [])
            self.assertEqual(list(processed_directory.iterdir()), [])

    def test_undetermined_country_is_a_different_error_than_not_allowed(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'

            def undeterminable_builder(path):
                del path
                raise UndeterminableDocumentCountryError('no country could be resolved from content')
            with self.assertRaises(AdminDocumentCountrySelectionRequiredError) as context:
                safe_upload_and_index_document(filename='Unknown.docx', file_stream=BytesIO(b'unknown-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=undeterminable_builder, country_document_lookup=_no_existing_documents, country_document_indexer=self._reject_if_called)
            self.assertNotIsInstance(context.exception, AdminDocumentCountryNotAllowedError)
            self.assertNotIsInstance(context.exception, DocumentCountryUndeterminedError)
            self.assertTrue(context.exception.allowed_countries)

    def test_every_one_of_the_34_allowed_codes_passes_the_check(self) -> None:
        for index, code in enumerate(sorted(ADMIN_ALLOWED_COUNTRY_CODES)):
            with self.subTest(code=code):
                with tempfile.TemporaryDirectory() as root:
                    source_directory = Path(root) / 'source'
                    processed_directory = Path(root) / 'processed'
                    country_name = canonical_country_name(code)
                    document_id = 'doc_' + format(index, '064x')
                    indexed: dict[str, object] = {}

                    def indexer(*, chunks, client=None):
                        del client
                        indexed['chunks'] = chunks
                        return DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
                    safe_upload_and_index_document(filename=f'{code}.docx', file_stream=BytesIO(f'{code}-bytes'.encode()), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=lambda path, _code=code, _name=country_name, _doc_id=document_id: _build_chunks_with_topic_count(_code, _name, _doc_id, path.name, 2), country_document_lookup=_no_existing_documents, country_document_indexer=indexer)
                    self.assertIn('chunks', indexed)

class SectionEditReplacementIntegrationTests(unittest.TestCase):
    """
    A CONFIRMED replace is a full country document reset: every
    persisted section edit belonging to
    the document(s) just replaced must be gone afterwards, so it can
    never silently reapply to the new DOCX. Never on a cancelled
    replace (AdminDocumentReplacementRequiredError) or an identical-
    bytes no-op (AdminDocumentAlreadyCurrentError) - both raise before
    safe_upload_and_index_document ever reaches that point.
    """

    def test_confirmed_replace_wipes_old_section_edit_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'old-australia')
            write_section_edit_state_atomic(source_directory, SectionEditState(document_id=AU_OLD_ID, country_code='AU', sections={_test_admin_documents__section_id_for_legal_topic('Employment Contracts'): SectionEdit(legal_topic='Employment Contracts', section='Employment Contracts', subsection=None, content='An edit that must not reapply to the replaced document.')}))
            self.assertIsNotNone(read_section_edit_state(source_directory, AU_OLD_ID))

            def indexer(*, chunks, client=None):
                del client
                return DocumentIndexingResult(index_alias='legal-documents-v1', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=1)
            response = safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'new-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, replace_existing=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=indexer)
            self.assertEqual(response.status, 'replaced')
            self.assertIsNone(read_section_edit_state(source_directory, AU_OLD_ID))

    def test_cancelled_replace_leaves_section_edit_state_untouched(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'old-australia')
            write_section_edit_state_atomic(source_directory, SectionEditState(document_id=AU_OLD_ID, country_code='AU', sections={_test_admin_documents__section_id_for_legal_topic('Employment Contracts'): SectionEdit(legal_topic='Employment Contracts', section='Employment Contracts', subsection=None, content='Must survive an unconfirmed replace.')}))

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError('indexer must not run')
            with self.assertRaises(AdminDocumentReplacementRequiredError):
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'new-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=indexer_must_not_be_called)
            after = read_section_edit_state(source_directory, AU_OLD_ID)
            self.assertIsNotNone(after)
            self.assertEqual(after.sections[_test_admin_documents__section_id_for_legal_topic('Employment Contracts')].content, 'Must survive an unconfirmed replace.')

    def test_identical_upload_leaves_section_edit_state_untouched(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'same-australia')
            write_section_edit_state_atomic(source_directory, SectionEditState(document_id=AU_OLD_ID, country_code='AU', sections={_test_admin_documents__section_id_for_legal_topic('Employment Contracts'): SectionEdit(legal_topic='Employment Contracts', section='Employment Contracts', subsection=None, content='Must survive an identical no-op upload.')}))

            def indexer_must_not_be_called(**kwargs):
                del kwargs
                raise AssertionError('the indexer must not be called for an already-current document')
            with self.assertRaises(AdminDocumentAlreadyCurrentError):
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'same-australia'), source_directory=source_directory, confirm_warnings=True, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=indexer_must_not_be_called)
            self.assertIsNotNone(read_section_edit_state(source_directory, AU_OLD_ID))

class CountryConfirmationGateTests(unittest.TestCase):
    """
    A detected country (from content or a marker) must never, by
    itself, cause any mutation:
    country_confirmed=True is required first, checked before the
    content-warning gate and before the existing-country/conflict
    checks, one decision at a time.
    """

    def test_detected_country_requires_confirmation_first(self) -> None:
        indexer_called = False

        def indexer_must_not_be_called(**kwargs):
            nonlocal indexer_called
            del kwargs
            indexer_called = True
            raise AssertionError('indexer must not run')
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            with self.assertRaises(AdminDocumentCountryConfirmationRequiredError) as context:
                safe_upload_and_index_document(filename='Australia.docx', file_stream=BytesIO(b'australia-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, confirm_warnings=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: (_ for _ in ()).throw(AssertionError('existing-country lookup must not run before country is confirmed')), country_document_indexer=indexer_must_not_be_called)
            self.assertEqual(context.exception.country_code, 'AU')
            self.assertEqual(context.exception.country, 'Australia')
            self.assertEqual(context.exception.detection_source, 'content')
            self.assertFalse(indexer_called)
            detail = context.exception.to_detail()
            self.assertEqual(detail['code'], 'document_country_confirmation_required')
            self.assertEqual(detail['country_code'], 'AU')
            self.assertEqual(detail['detection_source'], 'content')
            codes = {option['code'] for option in detail['allowed_countries']}
            self.assertIn('FR', codes)
            self.assertIn('AU', codes)
            self.assertEqual(list(source_directory.iterdir()) if source_directory.exists() else [], [])
            self.assertEqual(list(processed_directory.iterdir()) if processed_directory.exists() else [], [])

    def test_country_confirmed_true_proceeds(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            response = safe_upload_and_index_document(filename='Australia.docx', file_stream=BytesIO(b'australia-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, confirm_warnings=True, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: [], country_document_indexer=_fake_indexer)
            self.assertEqual(response.status, 'uploaded')
            self.assertEqual(response.country_code, 'AU')

    def test_confirmation_gate_precedes_the_content_warning_gate(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            with self.assertRaises(AdminDocumentCountryConfirmationRequiredError):
                safe_upload_and_index_document(filename='Chile.docx', file_stream=BytesIO(b'chile-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, confirm_warnings=False, chunk_builder=lambda path: _build_chunks_with_topic_count('CL', 'Chile', 'doc_' + 'c' * 64, path.name, 0), country_document_lookup=_no_existing_documents, country_document_indexer=lambda **kwargs: (_ for _ in ()).throw(AssertionError('indexer must not run')))

    def test_confirmation_gate_precedes_conflict_review(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            with self.assertRaises(AdminDocumentCountryConfirmationRequiredError):
                safe_upload_and_index_document(filename='Australia.docx', file_stream=BytesIO(b'australia-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, confirm_warnings=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: (_ for _ in ()).throw(AssertionError('must never reach the existing-country lookup before confirmation')))

class CountrySelectionGateTests(unittest.TestCase):
    """
    An otherwise-processable DOCX with no identifiable country is a
    SELECT_COUNTRY decision, never a
    hard failure; a manually-selected country is validated server-side
    against the admin allowlist before anything else happens.
    """

    def test_undetermined_or_ambiguous_country_returns_selection_required(self) -> None:

        def undeterminable_builder(path):
            del path
            raise UndeterminableDocumentCountryError('no country could be resolved from content')

        def ambiguous_builder(path):
            del path
            raise AmbiguousDocumentCountryError('more than one country found in the cover')
        cases = {'undetermined': (undeterminable_builder, 'Unknown.docx'), 'ambiguous': (ambiguous_builder, 'Ambiguous.docx')}
        for kind, (chunk_builder, filename) in cases.items():
            with self.subTest(kind=kind):
                with tempfile.TemporaryDirectory() as root:
                    source_directory = Path(root) / 'source'
                    processed_directory = Path(root) / 'processed'
                    with self.assertRaises(AdminDocumentCountrySelectionRequiredError) as context:
                        safe_upload_and_index_document(filename=filename, file_stream=BytesIO(b'unknown-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=chunk_builder, country_document_lookup=_no_existing_documents)
                    detail = context.exception.to_detail()
                    self.assertEqual(detail['code'], 'document_country_selection_required')
                    codes = {option['code'] for option in detail['allowed_countries']}
                    self.assertIn('FR', codes)
                    self.assertIn('AU', codes)

    def test_invalid_manual_selection_is_rejected_with_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            with self.assertRaises(AdminDocumentCountrySelectionInvalidError) as context:
                safe_upload_and_index_document(filename='Unknown.docx', file_stream=BytesIO(b'not-a-real-docx'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, confirm_warnings=True, selected_country_code='ZZ', chunk_builder=lambda path: (_ for _ in ()).throw(AssertionError('chunk_builder must not run for an invalid manual selection')))
            detail = context.exception.to_detail()
            self.assertEqual(detail['code'], 'document_country_selection_invalid')
            self.assertEqual(detail['country_code'], 'ZZ')
            self.assertEqual(list(source_directory.iterdir()) if source_directory.exists() else [], [])

    def test_unregistered_manual_selection_code_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            with self.assertRaises(AdminDocumentCountrySelectionInvalidError):
                safe_upload_and_index_document(filename='Unknown.docx', file_stream=BytesIO(b'not-a-real-docx'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, confirm_warnings=True, selected_country_code='XX')

class ManualCountryDocxRoundtripTests(unittest.TestCase):
    """
    The full manual-country invariant: a countryless DOCX, once the
    Admin selects a country,
    must have that choice persisted INSIDE the DOCX itself (a DOCX-
    native marker), surviving Download, a fresh independent parse, a
    full Reindex, and a later re-upload - with no external state
    anywhere.
    """

    def test_selecting_a_country_embeds_the_marker_and_uploads(self) -> None:
        countryless_bytes = _real_docx_bytes(['Some heading with no recognizable country.', '01. Hiring Practices', 'Content about hiring rules and probation periods.'])
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            built_chunks = []

            def capturing_chunk_builder(path):
                chunks = build_document_chunks_from_docx(path)
                built_chunks.extend(chunks)
                return chunks
            response = safe_upload_and_index_document(filename='mystery.docx', file_stream=BytesIO(countryless_bytes), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000000, confirm_warnings=True, selected_country_code='fr', chunk_builder=capturing_chunk_builder, country_document_lookup=lambda code, client: [], country_document_indexer=_fake_indexer)
            self.assertEqual(response.status, 'uploaded')
            self.assertEqual(response.country_code, 'FR')
            stored_path = source_directory / 'FR.docx'
            self.assertTrue(stored_path.exists())
            marker = read_country_marker(stored_path)
            self.assertIsNotNone(marker)
            self.assertEqual(marker.country_code, 'FR')
            self.assertEqual(response.source_filename, 'mystery.docx')
            self.assertTrue(built_chunks)
            self.assertTrue(all((chunk.source_filename == 'mystery.docx' for chunk in built_chunks)), [chunk.source_filename for chunk in built_chunks])

    def test_full_manual_country_roundtrip_invariant(self) -> None:
        countryless_bytes = _real_docx_bytes(['Some heading with no recognizable country.', '01. Hiring Practices', 'Content about hiring rules and probation periods.'])
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            first = safe_upload_and_index_document(filename='mystery.docx', file_stream=BytesIO(countryless_bytes), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000000, confirm_warnings=True, selected_country_code='fr', chunk_builder=build_document_chunks_from_docx, country_document_lookup=lambda code, client: [], country_document_indexer=_fake_indexer)
            self.assertEqual(first.status, 'uploaded')
            stored_path = source_directory / 'FR.docx'
            downloaded_bytes = stored_path.read_bytes()
            fresh_metadata = metadata_from_content(stored_path)
            self.assertEqual(fresh_metadata.country_code, 'FR')
            reindexed_chunks = build_document_chunks_from_docx(stored_path)
            self.assertTrue(reindexed_chunks)
            self.assertTrue(all((chunk.country_code == 'FR' for chunk in reindexed_chunks)))
            existing = [ExistingCountryDocument(document_id=first.document_id, source_filename='FR.docx', country='France', country_code='FR', reference_year=None)]
            with self.assertRaises(AdminDocumentCountryConfirmationRequiredError) as context:
                safe_upload_and_index_document(filename='mystery.docx', file_stream=BytesIO(downloaded_bytes), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000000, confirm_warnings=True, chunk_builder=build_document_chunks_from_docx, country_document_lookup=lambda code, client: existing)
            self.assertEqual(context.exception.country_code, 'FR')
            self.assertEqual(context.exception.detection_source, 'marker')

    def test_manual_selection_marker_embed_never_touches_disk_on_a_later_hard_failure(self) -> None:
        countryless_bytes = _real_docx_bytes(['Some heading with no recognizable country.'])
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            with self.assertRaises(AdminDocumentWarningConfirmationRequiredError):
                safe_upload_and_index_document(filename='mystery.docx', file_stream=BytesIO(countryless_bytes), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000000, confirm_warnings=False, selected_country_code='fr', chunk_builder=build_document_chunks_from_docx, country_document_lookup=lambda code, client: [], country_document_indexer=lambda **kwargs: (_ for _ in ()).throw(AssertionError('indexer must not run')))
            self.assertEqual(list(real_source_entries(source_directory)), [])

class CountryConflictReviewCandidateTests(unittest.TestCase):
    """
    The conflict review exposes only safe, business-facing
    per-candidate fields (filename, year,
    last-updated, file size); document_id is present only as an
    internal identity, never required for the Admin to read.
    """

    def test_candidates_expose_only_safe_fields_with_real_file_stats(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            first_path = source_directory / 'Australia-2024.docx'
            second_path = source_directory / 'Australia-legacy-v2.docx'
            first_path.write_bytes(b'first-australia-content')
            second_path.write_bytes(b'second-australia-content-longer')
            existing = [ExistingCountryDocument(document_id=AU_OLD_ID, source_filename=first_path.name, country='Australia', country_code='AU', reference_year=2024), ExistingCountryDocument(document_id='doc_' + 'c' * 64, source_filename=second_path.name, country='Australia', country_code='AU', reference_year=2026)]
            with self.assertRaises(AdminDocumentCountryConflictReviewRequiredError) as context:
                safe_upload_and_index_document(filename='Australia-new.docx', file_stream=BytesIO(b'new-australia'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, confirm_warnings=True, country_confirmed=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=lambda code, client: existing)
            detail = context.exception.to_detail()
            self.assertEqual(detail['code'], 'document_country_conflict_review_required')
            candidates = {candidate['source_filename']: candidate for candidate in detail['candidates']}
            self.assertEqual(set(candidates), {'Australia-2024.docx', 'Australia-legacy-v2.docx'})
            self.assertEqual(candidates['Australia-2024.docx']['reference_year'], 2024)
            self.assertEqual(candidates['Australia-2024.docx']['source_bytes'], len(b'first-australia-content'))
            self.assertIsNotNone(candidates['Australia-2024.docx']['updated_at'])
            self.assertIn('document_id', candidates['Australia-legacy-v2.docx'])

class _FakeReseedMetadataClient:
    """
    Minimal OpenSearch double for the "identical bytes, explicit
    confirmed reseed" tests below - just enough for
    reseed_contacts_from_current_docx's own metadata fetch
    (_get_document_metadata's plain, non-"sort" shape) plus a
    country-wide conflict check (_ensure_no_country_conflict's own
    "sort" shape). Never a full FakeCountryOpenSearch/FakeSection
    OpenSearchClient - those exist for OTHER tests' own different
    metadata shapes.
    """

    def __init__(self, *, document_id: str, country_code: str='AU', country: str='Australia', source_filename: str='Employment Law Overview Australia.docx', reference_year: int | None=None) -> None:
        self.document_id = document_id
        self.country_code = country_code
        self.country = country
        self.source_filename = source_filename
        self.reference_year = reference_year

    def search(self, index, body):
        del index
        if 'sort' in body:
            return {'hits': {'total': {'value': 0}, 'hits': []}}
        term = body.get('query', {}).get('term', {})
        if term.get('document_id') != self.document_id:
            return {'hits': {'hits': []}}
        return {'hits': {'hits': [{'_source': {'document_id': self.document_id, 'source_filename': self.source_filename, 'country': self.country, 'country_code': self.country_code, 'reference_year': self.reference_year}}]}}

class AdminModifiedReplacementWarningTests(unittest.TestCase):
    """
    The admin_modified-aware replacement warning integration - the
    admin-modified/reseed decision behavior that layers on top of the
    ordinary Contact CRUD/state tests.
    """

    def test_replacement_flow_admin_modified_flag_matches_marker_state(self) -> None:
        from app.services.document_section_state import mark_admin_modified
        for marker_set in (False, True):
            with self.subTest(marker_set=marker_set):
                with tempfile.TemporaryDirectory() as root:
                    source_directory = Path(root) / 'source'
                    processed_directory = Path(root) / 'processed'
                    source_directory.mkdir(parents=True)
                    (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'legacy-australia')
                    if marker_set:
                        mark_admin_modified(source_directory, AU_OLD_ID)
                    with self.assertRaises(AdminDocumentReplacementRequiredError) as context:
                        safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'new-australia-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents)
                    self.assertEqual(context.exception.to_detail()['admin_modified'], marker_set)

    def test_dirty_document_with_topic_warning_flags_admin_modified_too(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'legacy-australia')
            from app.services.document_section_state import mark_admin_modified
            mark_admin_modified(source_directory, AU_OLD_ID)
            with self.assertRaises(AdminDocumentWarningConfirmationRequiredError) as context:
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'new-australia-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=False, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents)
            error = context.exception
            self.assertTrue(error.replacement_required)
            self.assertTrue(error.to_detail()['admin_modified'])

    def test_identical_bytes_clean_still_raises_already_current(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'same-australia')
            with self.assertRaises(AdminDocumentAlreadyCurrentError):
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'same-australia'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents)

    def test_identical_bytes_dirty_requires_explicit_reseed_confirmation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'same-australia')
            from app.services.document_section_state import mark_admin_modified
            mark_admin_modified(source_directory, AU_OLD_ID)
            with self.assertRaises(AdminDocumentIdenticalButAdminModifiedError) as context:
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'same-australia'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents)
            self.assertEqual(context.exception.to_detail()['document_id'], AU_OLD_ID)

    def test_identical_bytes_dirty_cancel_leaves_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'same-australia')
            from app.services.document_section_state import mark_admin_modified
            from app.services.contact_state import ContactState, read_contact_state, write_contact_state_atomic
            from app.services.contact_state import ContactRecord
            mark_admin_modified(source_directory, AU_OLD_ID)
            write_contact_state_atomic(source_directory, ContactState(document_id=AU_OLD_ID, country_code='AU', contacts=(ContactRecord(contact_id='admin-edit-1', member_firm='Admin Edited Firm'),)))
            with self.assertRaises(AdminDocumentIdenticalButAdminModifiedError):
                safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'same-australia'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents)
            state = read_contact_state(source_directory, AU_OLD_ID)
            self.assertEqual(state.contacts[0].member_firm, 'Admin Edited Firm')

    def test_identical_bytes_dirty_confirmed_reseeds_contacts(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'same-australia')
            from app.services.document_section_state import is_admin_modified_since_upload, mark_admin_modified
            from app.services.contact_state import ContactRecord, ContactState, read_contact_state, write_contact_state_atomic
            from app.services.docx_parser import ExtractedContact
            mark_admin_modified(source_directory, AU_OLD_ID)
            write_contact_state_atomic(source_directory, ContactState(document_id=AU_OLD_ID, country_code='AU', contacts=(ContactRecord(contact_id='admin-edit-1', member_firm='Admin Edited Firm'),)))
            with patch('app.services.admin_contacts.extract_contacts_from_docx', return_value=[ExtractedContact(member_firm='Parsed DOCX Firm')]), patch('app.services.document_indexer.ensure_legal_documents_index'), patch('app.services.document_indexer.bulk', return_value=(1, [])), patch('app.services.document_indexer._snapshot_document_chunks', return_value=[]), patch('app.services.document_indexer._delete_chunks_except', return_value=0):
                response = safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'same-australia'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, confirm_contact_reseed=True, client=_FakeReseedMetadataClient(document_id=AU_OLD_ID), chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents)
            self.assertEqual(response.status, 'contacts_reseeded')
            self.assertEqual(response.contact_count, 1)
            state = read_contact_state(source_directory, AU_OLD_ID)
            self.assertEqual(state.contacts[0].member_firm, 'Parsed DOCX Firm')
            self.assertFalse(is_admin_modified_since_upload(source_directory, AU_OLD_ID))

    def test_confirmed_different_docx_resets_marker_and_reseeds(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'legacy-australia')
            from app.services.document_section_state import is_admin_modified_since_upload, mark_admin_modified
            from app.services.contact_state import ContactRecord, ContactState, read_contact_state, write_contact_state_atomic
            from app.services.docx_parser import ExtractedContact
            mark_admin_modified(source_directory, AU_NEW_ID)
            write_contact_state_atomic(source_directory, ContactState(document_id=AU_NEW_ID, country_code='AU', contacts=(ContactRecord(contact_id='admin-edit-1', member_firm='Old Admin Contact'), ContactRecord(contact_id='admin-edit-2', member_firm='Old Admin Contact 2'))))
            with patch('app.services.admin_document_replacement.extract_contacts_from_docx', return_value=[ExtractedContact(member_firm='New DOCX Firm')]):
                response = safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'brand-new-australia-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, replace_existing=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=_fake_indexer)
            self.assertEqual(response.contact_count, 1)
            state = read_contact_state(source_directory, AU_NEW_ID)
            self.assertEqual(len(state.contacts), 1)
            self.assertEqual(state.contacts[0].member_firm, 'New DOCX Firm')
            self.assertFalse(is_admin_modified_since_upload(source_directory, AU_NEW_ID))

    def test_confirmed_different_docx_with_zero_contacts_is_explicit_empty(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir(parents=True)
            (source_directory / 'Employment Law Overview Australia.docx').write_bytes(b'legacy-australia')
            from app.services.contact_state import ContactRecord, ContactState, read_contact_state, write_contact_state_atomic
            write_contact_state_atomic(source_directory, ContactState(document_id=AU_NEW_ID, country_code='AU', contacts=(ContactRecord(contact_id='admin-edit-1', member_firm='Old Admin Contact'),)))
            with patch('app.services.admin_document_replacement.extract_contacts_from_docx', return_value=[]):
                response = safe_upload_and_index_document(filename='Australia 2026.docx', file_stream=BytesIO(b'brand-new-australia-bytes'), source_directory=source_directory, processed_directory=processed_directory, maximum_bytes=1000, country_confirmed=True, confirm_warnings=True, replace_existing=True, chunk_builder=lambda path: [_build_au_chunk(path.name)], country_document_lookup=_existing_documents, country_document_indexer=_fake_indexer)
            self.assertEqual(response.contact_count, 0)
            state = read_contact_state(source_directory, AU_NEW_ID)
            self.assertIsNotNone(state)
            self.assertEqual(state.contacts, ())
_EXPECTED_ALLOWED_CODES = frozenset({'AR', 'AU', 'BE', 'BR', 'CA', 'CL', 'CN', 'CO', 'CZ', 'FR', 'DE', 'GR', 'ID', 'IE', 'IT', 'IN', 'JP', 'MX', 'NL', 'NO', 'PE', 'PH', 'PL', 'PT', 'RO', 'SG', 'SK', 'ES', 'SE', 'CH', 'TW', 'TR', 'GB', 'US'})

class AllowlistContentTests(unittest.TestCase):

    def test_exactly_the_34_client_mandated_codes(self) -> None:
        self.assertEqual(ADMIN_ALLOWED_COUNTRY_CODES, _EXPECTED_ALLOWED_CODES)
        self.assertEqual(len(ADMIN_ALLOWED_COUNTRY_CODES), 34)

    def test_every_allowed_code_is_a_real_registry_entry(self) -> None:
        registry_codes = {country.code for country in COUNTRIES}
        for code in ADMIN_ALLOWED_COUNTRY_CODES:
            with self.subTest(code=code):
                self.assertIn(code, registry_codes)

class IsAdminCountryAllowedTests(unittest.TestCase):

    def test_every_allowed_code_returns_true(self) -> None:
        for code in ADMIN_ALLOWED_COUNTRY_CODES:
            with self.subTest(code=code):
                self.assertTrue(is_admin_country_allowed(code))

    def test_registered_but_not_allowed_country_returns_false(self) -> None:
        self.assertFalse(is_admin_country_allowed('TN'))

    def test_entirely_unregistered_country_returns_false(self) -> None:
        self.assertFalse(is_admin_country_allowed('ZZ'))

    def test_is_case_and_whitespace_insensitive(self) -> None:
        self.assertTrue(is_admin_country_allowed(' fr '))
        self.assertTrue(is_admin_country_allowed('Fr'))
        self.assertFalse(is_admin_country_allowed(' tn '))

    def test_slovakia_identity_and_policy(self) -> None:
        self.assertTrue(is_admin_country_allowed('SK'))
        self.assertIn('SK', ADMIN_ALLOWED_COUNTRY_CODES)
OLD_DOCUMENT_ID = 'doc_' + 'a' * 64
NEW_DOCUMENT_ID = 'doc_' + 'b' * 64

def _build_lifecycle_chunk(*, document_id: str, source_filename: str) -> DocumentChunk:
    """Build one valid test chunk."""
    return DocumentChunk(document_id=document_id, chunk_id='chunk_' + 'c' * 64, country='United Kingdom', country_code='GB', legal_topic='Employment Contracts', document_type='comparator', language='en', section='Employment Contracts', subsection='Notice Period', content='One week of notice may apply.', source_filename=source_filename, source_format='docx', content_hash='content-hash', reference_year=2026)

class FakeLifecycleOpenSearchClient:
    """
    OpenSearch test double for lifecycle operations.

    country_document_ids models every document_id currently active
    for this country, as delete_indexed_document's own country-level
    lookup would see it - defaulting to [OLD_DOCUMENT_ID] alone, which
    is exactly the single-document assumption every pre-existing test
    in this file already made before that lookup existed. A test
    that needs to exercise several documents
    sharing one country (a real Australia-shaped duplicate) passes
    its own explicit list instead.
    """

    def __init__(self, *, document_exists: bool=True, source_filename: str='UK 2026.docx', country_document_ids: list[str] | None=None, country_source_filenames: dict[str, str] | None=None) -> None:
        self.document_exists = document_exists
        self.source_filename = source_filename
        self.country_document_ids = country_document_ids if country_document_ids is not None else [OLD_DOCUMENT_ID]
        self.country_source_filenames = country_source_filenames or {}
        self.deleted_document_ids: list[str] = []

    def search(self, index: str, body: dict[str, Any]) -> dict[str, Any]:
        del index
        term = body['query']['term']
        if 'country_code' in term:
            sorted_document_ids = sorted(self.country_document_ids)
            return {'hits': {'total': {'value': len(sorted_document_ids)}, 'hits': [{'_id': f'{document_id}-snapshot-chunk', '_source': {'document_id': document_id, 'chunk_id': f'{document_id}-snapshot-chunk', 'source_filename': self.country_source_filenames.get(document_id, self.source_filename), 'country': 'United Kingdom', 'country_code': 'GB', 'reference_year': 2026}, 'sort': [f'{document_id}-snapshot-chunk']} for document_id in sorted_document_ids]}}
        requested_document_id = term['document_id']
        if not self.document_exists:
            return {'hits': {'hits': []}}
        return {'hits': {'hits': [{'_source': {'document_id': requested_document_id, 'source_filename': self.source_filename, 'country': 'United Kingdom', 'country_code': 'GB', 'reference_year': 2026}}]}}

    def delete_by_query(self, *, index: str, body: dict[str, Any], conflicts: str, refresh: bool) -> dict[str, Any]:
        del index
        del conflicts
        del refresh
        document_id = body['query']['term']['document_id']
        self.deleted_document_ids.append(document_id)
        return {'deleted': 1}

class BackupInspectingOpenSearchClient(FakeLifecycleOpenSearchClient):
    """
    Records on-disk backup state at the moment chunks are deleted.

    This is the only point in delete_indexed_document where the
    source DOCX has already been moved but the operation has not
    yet completed - the right moment to observe where the backup
    was actually created.
    """

    def __init__(self, *, source_directory: Path, processed_directory: Path, **kwargs: Any) -> None:
        super().__init__(**kwargs)
        self.source_directory = source_directory
        self.processed_directory = processed_directory
        self.backup_path_at_delete_time: Path | None = None
        self.processed_directory_entries_at_delete_time: list[str] = []

    def delete_by_query(self, *, index: str, body: dict[str, Any], conflicts: str, refresh: bool) -> dict[str, Any]:
        backups = [path for path in real_source_entries(self.source_directory) if path.name.startswith('.delete-backup-')]
        if backups:
            self.backup_path_at_delete_time = backups[0]
        if self.processed_directory.exists():
            self.processed_directory_entries_at_delete_time = [path.name for path in self.processed_directory.iterdir()]
        return super().delete_by_query(index=index, body=body, conflicts=conflicts, refresh=refresh)

class AdminDocumentLifecycleTests(unittest.TestCase):
    """Tests for the ordinary, successful reindex and delete paths."""

    def test_reindex_existing_document(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_filename = 'UK 2026.docx'
            (source_directory / 'GB.docx').write_bytes(b'docx')
            client = FakeLifecycleOpenSearchClient(source_filename=source_filename)

            def chunk_builder(path: Path) -> list[DocumentChunk]:
                return [_build_lifecycle_chunk(document_id=OLD_DOCUMENT_ID, source_filename=path.name)]

            def document_indexer(*, chunks, client=None) -> DocumentIndexingResult:
                del client
                return DocumentIndexingResult(index_alias='legal-documents', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=1, indexed_chunks=1, stale_chunks_deleted=0)
            response = reindex_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, client=client, chunk_builder=chunk_builder, document_indexer=document_indexer)
            self.assertEqual(response.status, 'reindexed')
            self.assertFalse(response.document_id_changed)
            self.assertEqual(response.indexed_chunks, 1)
            self.assertEqual(client.deleted_document_ids, [])

    def test_changed_document_id_removes_previous_version(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_filename = 'UK 2026.docx'
            (source_directory / 'GB.docx').write_bytes(b'docx')
            client = FakeLifecycleOpenSearchClient(source_filename=source_filename)

            def chunk_builder(path: Path) -> list[DocumentChunk]:
                return [_build_lifecycle_chunk(document_id=NEW_DOCUMENT_ID, source_filename=path.name)]

            def document_indexer(*, chunks, client=None) -> DocumentIndexingResult:
                del client
                return DocumentIndexingResult(index_alias='legal-documents', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=1, indexed_chunks=1, stale_chunks_deleted=0)
            response = reindex_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, client=client, chunk_builder=chunk_builder, document_indexer=document_indexer)
            self.assertTrue(response.document_id_changed)
            self.assertEqual(response.document_id, NEW_DOCUMENT_ID)
            self.assertEqual(response.previous_chunks_deleted, 1)
            self.assertEqual(client.deleted_document_ids, [OLD_DOCUMENT_ID])

    def test_reindex_rejects_missing_source(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            with self.assertRaises(AdminDocumentSourceMissingError):
                reindex_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=Path(root), client=FakeLifecycleOpenSearchClient())

    def test_delete_removes_chunks_and_source(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir()
            source_filename = 'UK 2026.docx'
            source_path = source_directory / 'GB.docx'
            source_path.write_bytes(b'docx')
            client = FakeLifecycleOpenSearchClient(source_filename=source_filename)
            response = delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, processed_directory=processed_directory, client=client)
            self.assertEqual(response.status, 'deleted')
            self.assertEqual(response.deleted_chunks, 1)
            self.assertTrue(response.source_file_deleted)
            self.assertFalse(source_path.exists())
            self.assertEqual(client.deleted_document_ids, [OLD_DOCUMENT_ID])
            self.assertEqual(list(real_source_entries(source_directory)), [])

    def test_delete_backup_is_created_next_to_source_not_processed(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir()
            source_filename = 'UK 2026.docx'
            source_path = source_directory / 'GB.docx'
            source_path.write_bytes(b'docx')
            client = BackupInspectingOpenSearchClient(source_directory=source_directory, processed_directory=processed_directory, source_filename=source_filename)
            delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, processed_directory=processed_directory, client=client)
            self.assertIsNotNone(client.backup_path_at_delete_time)
            self.assertEqual(client.backup_path_at_delete_time.parent, source_directory)
            self.assertEqual(client.processed_directory_entries_at_delete_time, [])

    def test_delete_backup_path_does_not_end_with_docx(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir()
            source_filename = 'UK 2026.docx'
            source_path = source_directory / 'GB.docx'
            source_path.write_bytes(b'docx')
            client = BackupInspectingOpenSearchClient(source_directory=source_directory, processed_directory=processed_directory, source_filename=source_filename)
            delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, processed_directory=processed_directory, client=client)
            self.assertIsNotNone(client.backup_path_at_delete_time)
            self.assertFalse(client.backup_path_at_delete_time.name.endswith('.docx'))

    def test_unknown_document_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            with self.assertRaises(AdminDocumentNotFoundError):
                delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=Path(root) / 'source', processed_directory=Path(root) / 'processed', client=FakeLifecycleOpenSearchClient(document_exists=False))

class LegacySourceResolutionTests(unittest.TestCase):
    """
    Reindex and Delete for a
    document indexed before country-keyed storage existed, still
    physically stored under its own historical filename.
    """
    LEGACY_FILENAME = 'Labour and Employment Law in UK 2026.docx'

    def test_reindex_opens_the_exact_historical_file(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            legacy_path = source_directory / self.LEGACY_FILENAME
            legacy_path.write_bytes(b'legacy-docx-bytes')
            client = FakeLifecycleOpenSearchClient(source_filename=self.LEGACY_FILENAME)
            opened_paths: list[Path] = []

            def chunk_builder(path: Path) -> list[DocumentChunk]:
                opened_paths.append(path)
                return [_build_lifecycle_chunk(document_id=OLD_DOCUMENT_ID, source_filename=path.name)]

            def document_indexer(*, chunks, client=None) -> DocumentIndexingResult:
                del client
                return DocumentIndexingResult(index_alias='legal-documents', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=1, indexed_chunks=1, stale_chunks_deleted=0)
            response = reindex_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, client=client, chunk_builder=chunk_builder, document_indexer=document_indexer)
            self.assertEqual(response.status, 'reindexed')
            self.assertEqual(opened_paths, [legacy_path])
            self.assertFalse((source_directory / 'GB.docx').exists())

    def test_reindex_refuses_on_source_conflict(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / self.LEGACY_FILENAME).write_bytes(b'legacy-bytes')
            (source_directory / 'GB.docx').write_bytes(b'canonical-bytes')
            client = FakeLifecycleOpenSearchClient(source_filename=self.LEGACY_FILENAME)
            with self.assertRaises(AdminDocumentSourceConflictError):
                reindex_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual((source_directory / self.LEGACY_FILENAME).read_bytes(), b'legacy-bytes')
            self.assertEqual((source_directory / 'GB.docx').read_bytes(), b'canonical-bytes')

    def test_delete_targets_only_the_historical_file(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir()
            legacy_path = source_directory / self.LEGACY_FILENAME
            legacy_path.write_bytes(b'legacy-docx-bytes')
            decoy_path = source_directory / 'Labour and Employment Law in Spain 2026.docx'
            decoy_path.write_bytes(b'unrelated-spain-bytes')
            client = FakeLifecycleOpenSearchClient(source_filename=self.LEGACY_FILENAME)
            response = delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, processed_directory=processed_directory, client=client)
            self.assertEqual(response.status, 'deleted')
            self.assertTrue(response.source_file_deleted)
            self.assertFalse(legacy_path.exists())
            self.assertTrue(decoy_path.exists())
            self.assertEqual(decoy_path.read_bytes(), b'unrelated-spain-bytes')

    def test_delete_of_the_last_document_retires_every_conflicting_file(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir()
            (source_directory / self.LEGACY_FILENAME).write_bytes(b'legacy-bytes')
            (source_directory / 'GB.docx').write_bytes(b'canonical-bytes')
            client = FakeLifecycleOpenSearchClient(source_filename=self.LEGACY_FILENAME)
            response = delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, processed_directory=processed_directory, client=client)
            self.assertEqual(response.status, 'deleted')
            self.assertTrue(response.source_file_deleted)
            self.assertFalse(response.source_cleanup_deferred)
            self.assertEqual(client.deleted_document_ids, [OLD_DOCUMENT_ID])
            self.assertEqual(list(real_source_entries(source_directory)), [])

    def test_delete_of_one_duplicate_defers_file_cleanup(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir()
            (source_directory / self.LEGACY_FILENAME).write_bytes(b'legacy-bytes')
            (source_directory / 'GB.docx').write_bytes(b'canonical-bytes')
            sibling_document_id = 'doc_' + 'e' * 64
            client = FakeLifecycleOpenSearchClient(source_filename=self.LEGACY_FILENAME, country_document_ids=[OLD_DOCUMENT_ID, sibling_document_id])
            response = delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, processed_directory=processed_directory, client=client)
            self.assertEqual(response.status, 'deleted')
            self.assertFalse(response.source_file_deleted)
            self.assertTrue(response.source_cleanup_deferred)
            self.assertEqual(client.deleted_document_ids, [OLD_DOCUMENT_ID])
            self.assertEqual((source_directory / self.LEGACY_FILENAME).read_bytes(), b'legacy-bytes')
            self.assertEqual((source_directory / 'GB.docx').read_bytes(), b'canonical-bytes')

    def test_delete_of_one_of_three_duplicates_defers_file_cleanup(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir()
            (source_directory / self.LEGACY_FILENAME).write_bytes(b'legacy-bytes')
            (source_directory / 'GB.docx').write_bytes(b'canonical-bytes')
            sibling_ids = ['doc_' + 'e' * 64, 'doc_' + 'f' * 64]
            client = FakeLifecycleOpenSearchClient(source_filename=self.LEGACY_FILENAME, country_document_ids=[OLD_DOCUMENT_ID, *sibling_ids])
            response = delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, processed_directory=processed_directory, client=client)
            self.assertEqual(response.status, 'deleted')
            self.assertFalse(response.source_file_deleted)
            self.assertTrue(response.source_cleanup_deferred)
            self.assertEqual(client.deleted_document_ids, [OLD_DOCUMENT_ID])
            self.assertEqual((source_directory / self.LEGACY_FILENAME).read_bytes(), b'legacy-bytes')
            self.assertEqual((source_directory / 'GB.docx').read_bytes(), b'canonical-bytes')

    def test_deleting_down_to_the_last_of_three_then_retires_files(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root) / 'source'
            processed_directory = Path(root) / 'processed'
            source_directory.mkdir()
            (source_directory / self.LEGACY_FILENAME).write_bytes(b'legacy-bytes')
            (source_directory / 'GB.docx').write_bytes(b'canonical-bytes')
            client = FakeLifecycleOpenSearchClient(source_filename=self.LEGACY_FILENAME, country_document_ids=[OLD_DOCUMENT_ID])
            response = delete_indexed_document(document_id=OLD_DOCUMENT_ID, source_directory=source_directory, processed_directory=processed_directory, client=client)
            self.assertEqual(response.status, 'deleted')
            self.assertTrue(response.source_file_deleted)
            self.assertFalse(response.source_cleanup_deferred)
            self.assertEqual(list(real_source_entries(source_directory)), [])

class SectionEditReindexIntegrationTests(unittest.TestCase):
    """
    The current DOCX is the unique source of truth for Reindex too - a
    legacy .admin-state override file left over on disk from an older
    architecture must never be applied; every topic's fresh,
    DOCX-derived content wins, including the one a stale override
    file claims to have edited.
    """

    def test_reindex_ignores_legacy_persisted_edit_state(self) -> None:
        """
        The current DOCX is the unique source of truth - Reindex must
        NEVER apply a legacy .admin-state override, even if one is
        still present on disk from an older architecture. Every
        topic's fresh, DOCX-derived content wins, including the one a
        stale override file claims to have edited.
        """
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            (source_directory / 'GB.docx').write_bytes(b'docx-bytes')
            document_id = OLD_DOCUMENT_ID

            def fresh_chunk_builder(path: Path) -> list[DocumentChunk]:
                del path
                return [DocumentChunk(document_id=document_id, chunk_id='chunk-fresh-ec', country='United Kingdom', country_code='GB', legal_topic='Employment Contracts', document_type='comparator', language='en', section='Employment Contracts', subsection=None, content='FRESH DOCX content - must be replaced by the persisted edit.', source_filename='UK 2026.docx', source_format='docx', content_hash='fresh-ec-hash', reference_year=2026), DocumentChunk(document_id=document_id, chunk_id='chunk-fresh-hp', country='United Kingdom', country_code='GB', legal_topic='Hiring Practices', document_type='comparator', language='en', section='Hiring Practices', subsection=None, content='FRESH DOCX content - unaffected, no persisted edit exists for this topic.', source_filename='UK 2026.docx', source_format='docx', content_hash='fresh-hp-hash', reference_year=2026)]
            write_section_edit_state_atomic(source_directory, SectionEditState(document_id=document_id, country_code='GB', sections={_test_admin_documents__section_id_for_legal_topic('Employment Contracts'): SectionEdit(legal_topic='Employment Contracts', section='Employment Contracts', subsection=None, content='EDITED content - must survive Reindex.')}))
            captured_chunks: list[list[DocumentChunk]] = []

            def spy_document_indexer(*, chunks, client=None) -> DocumentIndexingResult:
                del client
                captured_chunks.append(chunks)
                return DocumentIndexingResult(index_alias='legal-documents', document_id=chunks[0].document_id, source_filename=chunks[0].source_filename, requested_chunks=len(chunks), indexed_chunks=len(chunks), stale_chunks_deleted=0)
            client = FakeLifecycleOpenSearchClient(source_filename='UK 2026.docx', country_document_ids=[document_id])
            response = reindex_indexed_document(document_id=document_id, source_directory=source_directory, client=client, chunk_builder=fresh_chunk_builder, document_indexer=spy_document_indexer)
            self.assertEqual(response.status, 'reindexed')
            self.assertFalse(response.document_id_changed)
            self.assertEqual(len(captured_chunks), 1)
            by_topic = {chunk.legal_topic: chunk for chunk in captured_chunks[0]}
            self.assertEqual(by_topic['Employment Contracts'].content, 'FRESH DOCX content - must be replaced by the persisted edit.')
            self.assertEqual(by_topic['Hiring Practices'].content, 'FRESH DOCX content - unaffected, no persisted edit exists for this topic.')
            legacy_state = read_section_edit_state(source_directory, document_id)
            self.assertIsNotNone(legacy_state)
            self.assertIn(_test_admin_documents__section_id_for_legal_topic('Employment Contracts'), legacy_state.sections)
import asyncio
import hashlib
import io
import json
import os
import tempfile
import unittest
import uuid
from pathlib import Path
from typing import Any
from unittest.mock import patch
from docx import Document
from fastapi import HTTPException, UploadFile
from app.core.config import get_settings
from app.routers import admin_document_conflicts as conflicts_router
from app.routers import admin_document_lifecycle as lifecycle_router
from app.routers import admin_documents as documents_router
from app.services import admin_document_lifecycle as lifecycle_service
from app.services.contact_state import ContactRecord, ContactState, write_contact_state_atomic
from tests.support.documents import assert_chunk_count_matches, assert_no_orphan_chunks, assert_one_active_document_per_country, assert_one_active_source, assert_zero_mutation, real_source_entries
from tests.support.opensearch import FakeAdminOpenSearch, bulk_writer_for
_SETTINGS_ENV_KEYS = ('OPENSEARCH_URL', 'OPENSEARCH_PASSWORD', 'REDIS_URL', 'DOCUMENT_SOURCE_DIR', 'DOCUMENT_PROCESSED_DIR', 'ADMIN_API_KEY', 'API_ACCESS_KEY')

def _build_real_docx_bytes(country_line: str) -> bytes:
    document = Document()
    document.add_paragraph(country_line)
    document.add_paragraph('I. GENERAL OVERVIEW')
    document.add_paragraph('1. Introduction')
    document.add_paragraph('Overview content.')
    document.add_paragraph('II. Hiring Practices')
    document.add_paragraph('Hiring content.')
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
_AR_BYTES = _build_real_docx_bytes('Labour and Employment Law in Argentina 2026')

def _make_upload_file(filename: str, content: bytes) -> UploadFile:
    return UploadFile(file=io.BytesIO(content), filename=filename)

class AdminRouterIntegrationTestCase(unittest.TestCase):
    """
    Base class: fresh temp source/processed dirs and a fresh
    FakeAdminOpenSearch per test, get_settings()/get_opensearch_client()
    patched at every call site the admin code uses them from.
    """

    def setUp(self) -> None:
        self._tempdir = tempfile.TemporaryDirectory()
        root = Path(self._tempdir.name)
        self.source_dir = root / 'source'
        self.processed_dir = root / 'processed'
        self.source_dir.mkdir(parents=True)
        self.processed_dir.mkdir(parents=True)
        self._original_env = {key: os.environ.get(key) for key in _SETTINGS_ENV_KEYS}
        os.environ['OPENSEARCH_URL'] = 'http://unused-in-this-test:9200'
        os.environ['OPENSEARCH_PASSWORD'] = 'unused'
        os.environ['REDIS_URL'] = 'redis://unused-in-this-test:6379/0'
        os.environ['DOCUMENT_SOURCE_DIR'] = str(self.source_dir)
        os.environ['DOCUMENT_PROCESSED_DIR'] = str(self.processed_dir)
        os.environ['ADMIN_API_KEY'] = 'unused-admin-key'
        os.environ['API_ACCESS_KEY'] = 'unused-api-key'
        get_settings.cache_clear()
        self.fake = FakeAdminOpenSearch()
        self._patches = [patch('app.services.admin_document_replacement.get_opensearch_client', return_value=self.fake), patch('app.services.admin_document_lifecycle.get_opensearch_client', return_value=self.fake), patch('app.services.admin_document_replacement.get_opensearch_client', return_value=self.fake), patch('app.services.document_indexer.get_opensearch_client', return_value=self.fake), patch('app.services.document_indexer.bulk', side_effect=bulk_writer_for(self.fake)), patch('app.services.document_indexer.ensure_legal_documents_index')]
        for p in self._patches:
            p.start()

    def tearDown(self) -> None:
        for p in reversed(self._patches):
            p.stop()
        for key, value in self._original_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value
        get_settings.cache_clear()
        self._tempdir.cleanup()

class FreshUploadHttpContractTests(AdminRouterIntegrationTestCase):
    """Fresh upload -> HTTP 201, then delete, then reupload under a
    different filename."""

    def test_fresh_upload_returns_201_uploaded_one_document_indexed(self) -> None:
        upload = _make_upload_file('Argentina.docx', _AR_BYTES)
        response = documents_router.upload_admin_document(file=upload, replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        self.assertEqual(response.status, 'uploaded')
        self.assertEqual(response.country_code, 'AR')
        self.assertGreater(response.indexed_chunks, 0)
        self.assertTrue((self.source_dir / 'AR.docx').exists())
        self.assertEqual(len(self.fake.document_ids_for_country('AR')), 1)

    def test_upload_then_delete_then_reupload_different_filename(self) -> None:
        first = documents_router.upload_admin_document(file=_make_upload_file('Argentina.docx', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        self.assertEqual(first.status, 'uploaded')
        delete_response = lifecycle_router.delete_admin_document(document_id=first.document_id)
        self.assertEqual(delete_response.status, 'deleted')
        self.assertEqual(len(self.fake.document_ids_for_country('AR')), 0)
        self.assertFalse((self.source_dir / 'AR.docx').exists())
        second = documents_router.upload_admin_document(file=_make_upload_file('random-file-name.docx', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        self.assertEqual(second.status, 'uploaded')
        self.assertEqual(len(self.fake.document_ids_for_country('AR')), 1)
        self.assertTrue((self.source_dir / 'AR.docx').exists())

class ExistingCountryHttpContractTests(AdminRouterIntegrationTestCase):
    """An existing country always yields 409 with the structured detail,
    regardless of the uploaded filename."""

    def setUp(self) -> None:
        super().setUp()
        self.fake.add(document_id='doc_' + 'a' * 64, country_code='AR', source_filename='Argentina.docx')
        (self.source_dir / 'AR.docx').write_bytes(b'existing-ar-bytes')

    def test_same_filename_returns_409_with_structured_detail(self) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(file=_make_upload_file('Argentina.docx', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        error = context.exception
        self.assertEqual(error.status_code, 409)
        self.assertEqual(error.detail['code'], 'document_replacement_required')
        self.assertEqual(error.detail['country_code'], 'AR')
        self.assertEqual((self.source_dir / 'AR.docx').read_bytes(), b'existing-ar-bytes')
        self.assertEqual(len(self.fake.document_ids_for_country('AR')), 1)

    def test_completely_unrelated_filename_still_returns_409(self) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(file=_make_upload_file('Legal-update-final-v7.docx', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        error = context.exception
        self.assertEqual(error.status_code, 409)
        self.assertEqual(error.detail['code'], 'document_replacement_required')
        self.assertEqual((self.source_dir / 'AR.docx').read_bytes(), b'existing-ar-bytes')

class DeleteHttpContractTests(AdminRouterIntegrationTestCase):
    """DELETE HTTP contract: duplicates (2 and 3 ids), the last document,
    source missing, and an unknown id."""

    def test_delete_one_of_two_duplicates_defers_and_keeps_the_other(self) -> None:
        first_id = 'doc_' + 'a' * 64
        second_id = 'doc_' + 'b' * 64
        self.fake.add(document_id=first_id, country_code='AR', source_filename='Argentina.docx')
        self.fake.add(document_id=second_id, country_code='AR', source_filename='Argentina.docx', chunk_id=f'{second_id}-chunk-0')
        (self.source_dir / 'AR.docx').write_bytes(b'canonical')
        (self.source_dir / 'Argentina.docx').write_bytes(b'legacy')
        response = lifecycle_router.delete_admin_document(document_id=first_id)
        self.assertEqual(response.status, 'deleted')
        self.assertTrue(response.source_cleanup_deferred)
        self.assertFalse(response.source_file_deleted)
        self.assertEqual(self.fake.document_ids_for_country('AR'), {second_id})
        self.assertTrue((self.source_dir / 'AR.docx').exists())
        self.assertTrue((self.source_dir / 'Argentina.docx').exists())

    def test_delete_one_of_three_duplicates_defers_and_keeps_the_others(self) -> None:
        ids = ['doc_' + letter * 64 for letter in ('a', 'b', 'c')]
        for document_id in ids:
            self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx', chunk_id=f'{document_id}-chunk-0')
        (self.source_dir / 'AR.docx').write_bytes(b'canonical')
        response = lifecycle_router.delete_admin_document(document_id=ids[0])
        self.assertEqual(response.status, 'deleted')
        self.assertTrue(response.source_cleanup_deferred)
        self.assertEqual(self.fake.document_ids_for_country('AR'), set(ids[1:]))
        self.assertTrue((self.source_dir / 'AR.docx').exists())

    def test_delete_of_the_last_document_cleans_up_candidate_files(self) -> None:
        only_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=only_id, country_code='AR', source_filename='Argentina.docx')
        (self.source_dir / 'AR.docx').write_bytes(b'canonical')
        (self.source_dir / 'Argentina.docx').write_bytes(b'legacy')
        response = lifecycle_router.delete_admin_document(document_id=only_id)
        self.assertEqual(response.status, 'deleted')
        self.assertFalse(response.source_cleanup_deferred)
        self.assertTrue(response.source_file_deleted)
        self.assertEqual(real_source_entries(self.source_dir), [])
        self.assertEqual(self.fake.document_ids_for_country('AR'), set())

    def test_delete_with_source_missing_still_succeeds(self) -> None:
        only_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=only_id, country_code='AR', source_filename='Argentina.docx')
        response = lifecycle_router.delete_admin_document(document_id=only_id)
        self.assertEqual(response.status, 'deleted')
        self.assertFalse(response.source_file_deleted)
        self.assertFalse(response.source_cleanup_deferred)

    def test_delete_unknown_document_id_returns_404(self) -> None:
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.delete_admin_document(document_id='doc_' + 'f' * 64)
        self.assertEqual(context.exception.status_code, 404)
        detail = context.exception.detail
        self.assertEqual(detail['code'], 'document_not_found')
        self.assertEqual(detail['operation'], 'delete')
        self.assertEqual(detail['document_id'], 'doc_' + 'f' * 64)
        self.assertTrue(detail['message'])

class ReindexHttpContractTests(AdminRouterIntegrationTestCase):
    """
    Every reindex failure path must return the structured
    {"code", "message", "operation", "document_id"}
    contract instead of one fixed, generic sentence - this is what
    let Czech Republic's real UndeterminableDocumentCountryError
    reach WordPress/the logs as a generic 502 with no actionable
    detail (see backend/app/core/country_registry.py's leading-
    definite-article fallback for the actual Czech fix; this class
    only covers the observability contract, not that fix).
    """

    def test_reindex_unknown_document_id_returns_structured_404(self) -> None:
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.reindex_admin_document(document_id='doc_' + 'f' * 64)
        self.assertEqual(context.exception.status_code, 404)
        detail = context.exception.detail
        self.assertEqual(detail['code'], 'document_not_found')
        self.assertEqual(detail['operation'], 'reindex')
        self.assertEqual(detail['document_id'], 'doc_' + 'f' * 64)

    def test_reindex_with_source_missing_returns_structured_409(self) -> None:
        document_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.reindex_admin_document(document_id=document_id)
        self.assertEqual(context.exception.status_code, 409)
        detail = context.exception.detail
        self.assertEqual(detail['code'], 'source_missing')
        self.assertEqual(detail['operation'], 'reindex')
        self.assertEqual(detail['document_id'], document_id)
        self.assertIn('missing', detail['message'].casefold())

    def test_reindex_unparseable_source_returns_structured_502(self) -> None:
        document_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        (self.source_dir / 'Argentina.docx').write_bytes(_build_real_docx_bytes('Some random legal memo with no title structure.'))
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.reindex_admin_document(document_id=document_id)
        self.assertEqual(context.exception.status_code, 502)
        detail = context.exception.detail
        self.assertEqual(detail['code'], 'document_reindex_failed')
        self.assertEqual(detail['operation'], 'reindex')
        self.assertEqual(detail['document_id'], document_id)
        self.assertIn('supported country', detail['message'].casefold())

    def test_reindex_with_country_conflict_returns_structured_409(self) -> None:
        document_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        self.fake.add(document_id='doc_' + 'b' * 64, country_code='AR', source_filename='Argentina-legacy.docx', chunk_id='doc_' + 'b' * 64 + '-chunk-0')
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.reindex_admin_document(document_id=document_id)
        self.assertEqual(context.exception.status_code, 409)
        detail = context.exception.detail
        self.assertEqual(detail['code'], 'country_document_conflict')
        self.assertEqual(detail['operation'], 'reindex')
        self.assertEqual(detail['country_code'], 'AR')

class TechnicalValidationHttpContractTests(AdminRouterIntegrationTestCase):
    """Every technical upload failure returns its own specific,
    structured code."""

    def test_wrong_extension_returns_invalid_document_type(self) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(file=_make_upload_file('not-a-docx.txt', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail['code'], 'invalid_document_type')

    def test_empty_file_returns_document_empty(self) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(file=_make_upload_file('Argentina.docx', b''), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail['code'], 'document_empty')

    def test_oversized_file_returns_413_document_too_large(self) -> None:
        os.environ['DOCUMENT_UPLOAD_MAX_BYTES'] = '100'
        get_settings.cache_clear()
        try:
            with self.assertRaises(HTTPException) as context:
                documents_router.upload_admin_document(file=_make_upload_file('Argentina.docx', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        finally:
            os.environ.pop('DOCUMENT_UPLOAD_MAX_BYTES', None)
            get_settings.cache_clear()
        self.assertEqual(context.exception.status_code, 413)
        detail = context.exception.detail
        self.assertEqual(detail['code'], 'document_too_large')
        self.assertEqual(detail['max_bytes'], 100)
        self.assertIn('max_mb', detail)

    def test_corrupt_zip_returns_document_corrupt(self) -> None:
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(file=_make_upload_file('Argentina.docx', b'not a real docx file at all'), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail['code'], 'document_corrupt')

    def test_no_identifiable_country_returns_selection_required(self) -> None:
        no_country_bytes = _build_real_docx_bytes('Some random legal memo with no title structure.')
        with self.assertRaises(HTTPException) as context:
            documents_router.upload_admin_document(file=_make_upload_file('mystery.docx', no_country_bytes), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        self.assertEqual(context.exception.status_code, 409)
        detail = context.exception.detail
        self.assertEqual(detail['code'], 'document_country_selection_required')
        self.assertTrue(detail['allowed_countries'])

class DownloadHttpContractTests(AdminRouterIntegrationTestCase):
    """GET .../download."""

    def test_download_returns_the_real_source_bytes(self) -> None:
        document_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        real_bytes = b'the-real-argentina-source-bytes'
        (self.source_dir / 'AR.docx').write_bytes(real_bytes)
        response = lifecycle_router.download_admin_document(document_id=document_id)
        self.assertEqual(Path(response.path).read_bytes(), real_bytes)
        self.assertEqual(response.filename, 'Argentina.docx')

    def test_download_unknown_document_id_returns_404(self) -> None:
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.download_admin_document(document_id='doc_' + 'f' * 64)
        self.assertEqual(context.exception.status_code, 404)
        self.assertEqual(context.exception.detail['code'], 'document_not_found')

    def test_download_with_missing_source_returns_409(self) -> None:
        document_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.download_admin_document(document_id=document_id)
        self.assertEqual(context.exception.status_code, 409)
        self.assertEqual(context.exception.detail['code'], 'source_missing')

    def test_download_with_source_conflict_returns_409(self) -> None:
        document_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        (self.source_dir / 'AR.docx').write_bytes(b'canonical')
        (self.source_dir / 'Argentina.docx').write_bytes(b'legacy')
        with self.assertRaises(HTTPException) as context:
            lifecycle_router.download_admin_document(document_id=document_id)
        self.assertEqual(context.exception.status_code, 409)
        self.assertEqual(context.exception.detail['code'], 'source_conflict')

    def test_download_with_no_contact_state_returns_source_unchanged(self) -> None:
        document_id = 'doc_' + 'a' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        real_bytes = _build_real_docx_bytes('Argentina overview')
        (self.source_dir / 'AR.docx').write_bytes(real_bytes)
        download = lifecycle_service.get_document_download(document_id=document_id, source_directory=self.source_dir)
        self.assertEqual(download.path, self.source_dir / 'AR.docx')
        self.assertEqual(download.path.read_bytes(), real_bytes)

    def test_download_with_contact_state_still_returns_persisted_source_unchanged(self) -> None:
        """
        Download is a pure read of the persisted source, whether or
        not ContactState exists -
        it must NEVER rebuild/reserialize a "materialized effective"
        copy. Every Admin mutation already writes its effective result
        into the source atomically before returning success, so a
        real ContactState here changes nothing about what download
        returns: it is not even read.
        """
        document_id = 'doc_' + 'b' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        real_bytes = _build_real_docx_bytes('Argentina overview')
        (self.source_dir / 'AR.docx').write_bytes(real_bytes)
        write_contact_state_atomic(self.source_dir, ContactState(document_id=document_id, country_code='AR', contacts=(ContactRecord(contact_id='contact-1', member_firm='CURRENT FIRM', contact_person='Current Person', email='current@example.com', phone='+1 555 0100', address='1 Current Street', website='www.current.example'),)))
        self.assertFalse(hasattr(lifecycle_service, 'read_contact_state'), 'get_document_download() must not even import read_contact_state - a real ContactState sidecar must have zero effect on what download returns')
        download = lifecycle_service.get_document_download(document_id=document_id, source_directory=self.source_dir)
        self.assertEqual(download.path, self.source_dir / 'AR.docx')
        self.assertEqual(download.path.read_bytes(), real_bytes)

    def test_repeated_downloads_of_unchanged_document_are_byte_identical(self) -> None:
        """
        The exact regression this hardening exists for: the OLD
        materializer produced a different SHA256 on every single call
        for the same unchanged document. A pure byte read cannot do
        that.
        """
        document_id = 'doc_' + 'd' * 64
        self.fake.add(document_id=document_id, country_code='AR', source_filename='Argentina.docx')
        real_bytes = _build_real_docx_bytes('Argentina overview')
        (self.source_dir / 'AR.docx').write_bytes(real_bytes)
        write_contact_state_atomic(self.source_dir, ContactState(document_id=document_id, country_code='AR', contacts=(ContactRecord(contact_id='contact-1', member_firm='CURRENT FIRM'),)))
        hashes = set()
        for _ in range(10):
            download = lifecycle_service.get_document_download(document_id=document_id, source_directory=self.source_dir)
            hashes.add(hashlib.sha256(download.path.read_bytes()).hexdigest())
        self.assertEqual(len(hashes), 1, '10 consecutive downloads of an unchanged document must all have exactly the same SHA256')

    def test_download_does_not_import_document_contact_materializer(self) -> None:
        """
        Structural guard against reintroducing document materialization
        into the download endpoint: the service module must not import
        materialize_effective_docx at all - not "import it but never
        call it", genuinely absent, so a future re-addition of the
        import itself fails this test before anyone even wires it back
        into get_document_download().
        """
        import inspect
        source = inspect.getsource(lifecycle_service)
        self.assertNotIn('materialize_effective_docx', source, 'get_document_download() must never reference materialize_effective_docx - download is a pure byte read of the persisted source')
        self.assertFalse(hasattr(lifecycle_service, 'materialize_effective_docx'), 'materialize_effective_docx must not be importable from the download service module')

class StatsHttpContractTests(AdminRouterIntegrationTestCase):
    """GET .../documents/stats."""

    def test_stats_reflect_the_real_catalog_exactly(self) -> None:
        self.fake.add(document_id='doc_' + 'a' * 64, country_code='AR', source_filename='Argentina.docx')
        self.fake.add(document_id='doc_' + 'b' * 64, country_code='BR', source_filename='Brazil.docx')
        self.fake.add(document_id='doc_' + 'c' * 64, country_code='CO', source_filename='Colombia.docx', chunk_id='doc_' + 'c' * 64 + '-chunk-0')
        (self.source_dir / 'AR.docx').write_bytes(b'argentina')
        (self.source_dir / 'BR.docx').write_bytes(b'brazil')
        stats = documents_router.get_admin_document_stats_route()
        self.assertEqual(stats.total_documents, 3)
        self.assertEqual(stats.total_countries, 3)
        self.assertEqual(stats.status_counts.get('indexed'), 2)
        self.assertEqual(stats.status_counts.get('indexed_source_missing'), 1)

class SharedInvariantHelperUsageTests(AdminRouterIntegrationTestCase):
    """Demonstrates the shared, reusable invariant helpers
    (admin_invariants.py) against a real catalog listing, rather than
    each caller re-deriving these checks."""

    def test_clean_catalog_satisfies_every_generic_invariant(self) -> None:
        first = documents_router.upload_admin_document(file=_make_upload_file('Argentina.docx', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        catalog = documents_router.get_admin_documents()
        documents = [d.model_dump() for d in catalog.documents]
        assert_one_active_document_per_country(documents, 'AR')
        assert_one_active_source(documents, 'AR')
        real_count = self.fake.search(index='unused', body={'query': {'term': {'country_code': 'AR'}}})['hits']['total']['value']
        assert_chunk_count_matches(first.indexed_chunks, real_count, country_code='AR')
        total_catalog_chunks = sum((d['chunk_count'] for d in documents))
        total_real_chunks = len(self.fake.chunks)
        assert_no_orphan_chunks(total_catalog_chunks, total_real_chunks)

    def test_rejected_replacement_is_zero_mutation_via_shared_helper(self) -> None:
        self.fake.add(document_id='doc_' + 'a' * 64, country_code='AR', source_filename='Argentina.docx')
        (self.source_dir / 'AR.docx').write_bytes(b'existing-ar-bytes')
        before = [d.model_dump() for d in documents_router.get_admin_documents().documents]
        with self.assertRaises(HTTPException):
            documents_router.upload_admin_document(file=_make_upload_file('Legal-update-final-v7.docx', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        after = [d.model_dump() for d in documents_router.get_admin_documents().documents]
        assert_zero_mutation(before, after)

class RestoreEndpointRemovedTests(unittest.TestCase):
    """
    The old "Restore from document" endpoint is now moot (Edit itself
    mutates the current DOCX) and must be fully removed from the
    runtime - never merely hidden.
    """

    def test_no_restore_route_is_registered(self) -> None:
        restore_paths = [route.path for route in lifecycle_router.router.routes if route.path.endswith('/restore')]
        self.assertEqual(restore_paths, [])

    def test_add_section_route_is_registered_instead(self) -> None:
        add_section_routes = [route for route in lifecycle_router.router.routes if route.path == '/api/v1/admin/documents/{document_id}/sections' and 'POST' in route.methods]
        self.assertEqual(len(add_section_routes), 1)

    def test_restore_effective_section_no_longer_exists(self) -> None:
        import app.services.admin_document_sections as sections_service
        self.assertFalse(hasattr(sections_service, 'restore_effective_section'))

class IdenticalButAdminModifiedHttpContractTests(AdminRouterIntegrationTestCase):
    """
    The router's own exception -> HTTP mapping for
    AdminDocumentIdenticalButAdminModifiedError.
    Deliberately does not extend FakeAdminOpenSearch's own query shapes for
    this (its delete_by_query double is hard-coded to the country-wide
    cleanup shape, not the Contact-chunk-scoped one) - the service-level
    behavior this maps is already exhaustively covered in
    test_admin_document_replacement.py; this proves only the NEW
    router-level mapping itself.
    """

    def test_maps_to_409_with_structured_detail(self) -> None:
        from app.services.admin_document_replacement import AdminDocumentIdenticalButAdminModifiedError

        def _raise(**kwargs):
            del kwargs
            raise AdminDocumentIdenticalButAdminModifiedError(country='Argentina', country_code='AR', document_id='doc_' + 'a' * 64)
        with patch('app.routers.admin_documents.safe_upload_and_index_document', side_effect=_raise):
            with self.assertRaises(HTTPException) as context:
                documents_router.upload_admin_document(file=_make_upload_file('Argentina.docx', _AR_BYTES), replace_existing=False, confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        error = context.exception
        self.assertEqual(error.status_code, 409)
        self.assertEqual(error.detail['code'], 'document_identical_but_admin_modified')
        self.assertTrue(error.detail['admin_modified'])
        self.assertEqual(error.detail['document_id'], 'doc_' + 'a' * 64)

class ConflictResolutionRouterTests(AdminRouterIntegrationTestCase):
    """
    Router-level HTTP contract tests for the two country-conflict
    endpoints - reuses this file's own FakeAdminOpenSearch double and
    AdminRouterIntegrationTestCase base rather than a second, parallel
    fixture.
    """

    def setUp(self) -> None:
        super().setUp()
        self._extra_patches = [patch('app.services.admin_document_replacement.get_opensearch_client', return_value=self.fake)]
        for p in self._extra_patches:
            p.start()

    def tearDown(self) -> None:
        for p in reversed(self._extra_patches):
            p.stop()
        super().tearDown()

    def _seed_conflict(self) -> None:
        self.fake.add(document_id='doc_' + 'a' * 64, country_code='AR', source_filename='Argentina-old.docx')
        self.fake.add(document_id='doc_' + 'b' * 64, country_code='AR', source_filename='Argentina-new.docx', chunk_id='doc_' + 'b' * 64 + '-chunk-0')
        (self.source_dir / 'Argentina-old.docx').write_bytes(b'old content')
        (self.source_dir / 'Argentina-new.docx').write_bytes(b'new content, different')

    def test_review_returns_candidates_and_auto_dedup_flag(self) -> None:
        self._seed_conflict()
        response = conflicts_router.get_country_conflict_review(country_code='AR')
        self.assertEqual(response.country_code, 'AR')
        self.assertEqual(len(response.candidates), 2)
        self.assertFalse(response.auto_deduplicate_available)

    def test_review_404s_when_country_is_not_conflicted(self) -> None:
        self.fake.add(document_id='doc_' + 'a' * 64, country_code='AR', source_filename='Argentina.docx')
        with self.assertRaises(HTTPException) as context:
            conflicts_router.get_country_conflict_review(country_code='AR')
        self.assertEqual(context.exception.status_code, 404)
        self.assertEqual(context.exception.detail['code'], 'country_conflict_not_found')

    def test_choose_document_resolves_via_http_contract(self) -> None:
        self._seed_conflict()
        response = conflicts_router.resolve_country_conflict(country_code='AR', resolution_mode='CHOOSE_DOCUMENT', keep_document_id='doc_' + 'b' * 64, file=None)
        self.assertEqual(response.kept_document_id, 'doc_' + 'b' * 64)
        self.assertEqual(response.removed_document_ids, ['doc_' + 'a' * 64])
        self.assertFalse((self.source_dir / 'Argentina-old.docx').exists())
        self.assertTrue((self.source_dir / 'Argentina-new.docx').exists())

    def test_choose_document_stale_id_returns_422(self) -> None:
        self._seed_conflict()
        with self.assertRaises(HTTPException) as context:
            conflicts_router.resolve_country_conflict(country_code='AR', resolution_mode='CHOOSE_DOCUMENT', keep_document_id='doc_totally_unrelated', file=None)
        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail['code'], 'country_conflict_resolution_invalid')

    def test_replace_with_document_requires_a_file(self) -> None:
        self._seed_conflict()
        with self.assertRaises(HTTPException) as context:
            conflicts_router.resolve_country_conflict(country_code='AR', resolution_mode='REPLACE_WITH_DOCUMENT', file=None)
        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail['code'], 'document_required')

    def test_replace_with_document_rejects_a_mismatched_country(self) -> None:
        self._seed_conflict()
        document = Document()
        document.add_paragraph('Labour and Employment Law in Chile 2026')
        buffer = io.BytesIO()
        document.save(buffer)
        with self.assertRaises(HTTPException) as context:
            conflicts_router.resolve_country_conflict(country_code='AR', resolution_mode='REPLACE_WITH_DOCUMENT', file=_make_upload_file('chile.docx', buffer.getvalue()), confirm_warnings=True, country_confirmed=True, selected_country_code=None)
        self.assertEqual(context.exception.status_code, 422)
        self.assertEqual(context.exception.detail['code'], 'document_unexpected_country')
        self.assertTrue((self.source_dir / 'Argentina-old.docx').exists())
        self.assertTrue((self.source_dir / 'Argentina-new.docx').exists())

class ASGIResponse:
    """A minimal stand-in for httpx.Response - just what these tests
    need: status_code, header lookup, raw body, and .json()."""

    def __init__(self, status_code: int, headers: list[tuple[bytes, bytes]], body: bytes) -> None:
        self.status_code = status_code
        self._headers = headers
        self.content = body

    def header(self, name: str) -> str | None:
        target = name.lower().encode()
        for key, value in self._headers:
            if key.lower() == target:
                return value.decode()
        return None

    def json(self) -> Any:
        return json.loads(self.content)

async def _asgi_request(app: Any, method: str, path: str, headers: dict[str, str] | None=None, body: bytes=b'') -> ASGIResponse:
    """
    Drive the real ASGI application callable directly with a
    synthetic in-memory scope/receive/send - the same protocol
    httpx.ASGITransport itself implements, without needing httpx.
    """
    request_headers = [(key.lower().encode(), value.encode()) for key, value in (headers or {}).items()]
    scope = {'type': 'http', 'asgi': {'version': '3.0', 'spec_version': '2.3'}, 'http_version': '1.1', 'method': method, 'path': path, 'raw_path': path.encode(), 'query_string': b'', 'root_path': '', 'headers': request_headers, 'client': ('127.0.0.1', 12345), 'server': ('127.0.0.1', 80), 'scheme': 'http', 'state': {}}
    body_delivered = {'done': False}

    async def receive() -> dict[str, Any]:
        if body_delivered['done']:
            return {'type': 'http.disconnect'}
        body_delivered['done'] = True
        return {'type': 'http.request', 'body': body, 'more_body': False}
    collected: dict[str, Any] = {'status': None, 'headers': [], 'body': b''}

    async def send(message: dict[str, Any]) -> None:
        if message['type'] == 'http.response.start':
            collected['status'] = message['status']
            collected['headers'] = message['headers']
        elif message['type'] == 'http.response.body':
            collected['body'] += message.get('body', b'')
    await app(scope, receive, send)
    return ASGIResponse(collected['status'], collected['headers'], collected['body'])

def asgi_request(app: Any, method: str, path: str, headers: dict[str, str] | None=None, body: bytes=b'') -> ASGIResponse:
    return asyncio.run(_asgi_request(app, method, path, headers=headers, body=body))

def build_multipart_body(*, file_field: str, filename: str, file_content: bytes, extra_fields: dict[str, str] | None=None, include_file: bool=True) -> tuple[bytes, str]:
    """Build a real multipart/form-data body byte-for-byte, the same
    shape a browser or the WordPress PHP proxy would send."""
    boundary = f'----ASGITestBoundary{uuid.uuid4().hex}'
    lines: list[bytes] = []
    for name, value in (extra_fields or {}).items():
        lines.append(f'--{boundary}'.encode())
        lines.append(f'Content-Disposition: form-data; name="{name}"'.encode())
        lines.append(b'')
        lines.append(value.encode())
    if include_file:
        lines.append(f'--{boundary}'.encode())
        lines.append(f'Content-Disposition: form-data; name="{file_field}"; filename="{filename}"'.encode())
        lines.append(b'Content-Type: application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        lines.append(b'')
        lines.append(file_content)
    lines.append(f'--{boundary}--'.encode())
    lines.append(b'')
    body = b'\r\n'.join(lines)
    content_type = f'multipart/form-data; boundary={boundary}'
    return (body, content_type)

def _build_minimal_docx_bytes(country_line: str) -> bytes:
    document = Document()
    document.add_paragraph(country_line)
    document.add_paragraph('I. GENERAL OVERVIEW')
    document.add_paragraph('1. Introduction')
    document.add_paragraph('Overview content.')
    document.add_paragraph('II. Hiring Practices')
    document.add_paragraph('Hiring content.')
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()
_CL_DOCX_BYTES = _build_minimal_docx_bytes('Labour and Employment Law in Chile 2026')
TEST_API_ACCESS_KEY = 'test-only-asgi-api-access-key'
TEST_ADMIN_API_KEY = 'test-only-asgi-admin-api-key'

class FakeRedis:
    """Always allows the request through - a real rate-limit backend
    is unrelated to what these tests verify."""

    def eval(self, script: str, numkeys: int, *args: Any) -> list[int]:
        del script, numkeys, args
        return [1, 60]

class AdminAsgiTestCase(unittest.TestCase):
    """Fresh temp dirs, fake OpenSearch, fake Redis, real test API/
    admin keys - the real FastAPI ASGI app, imported fresh under
    these settings."""

    def setUp(self) -> None:
        self._tempdir = tempfile.TemporaryDirectory()
        root = Path(self._tempdir.name)
        self.source_dir = root / 'source'
        self.processed_dir = root / 'processed'
        self.source_dir.mkdir(parents=True)
        self.processed_dir.mkdir(parents=True)
        self._original_env = {key: os.environ.get(key) for key in _SETTINGS_ENV_KEYS}
        os.environ['OPENSEARCH_URL'] = 'http://unused-in-asgi-tests:9200'
        os.environ['OPENSEARCH_PASSWORD'] = 'unused'
        os.environ['REDIS_URL'] = 'redis://unused-in-asgi-tests:6379/0'
        os.environ['DOCUMENT_SOURCE_DIR'] = str(self.source_dir)
        os.environ['DOCUMENT_PROCESSED_DIR'] = str(self.processed_dir)
        os.environ['ADMIN_API_KEY'] = TEST_ADMIN_API_KEY
        os.environ['API_ACCESS_KEY'] = TEST_API_ACCESS_KEY
        get_settings.cache_clear()
        self.fake_opensearch = FakeAdminOpenSearch()
        self.fake_redis = FakeRedis()
        self._patches = [patch('app.services.admin_document_replacement.get_opensearch_client', return_value=self.fake_opensearch), patch('app.services.admin_document_lifecycle.get_opensearch_client', return_value=self.fake_opensearch), patch('app.services.admin_document_replacement.get_opensearch_client', return_value=self.fake_opensearch), patch('app.services.document_indexer.get_opensearch_client', return_value=self.fake_opensearch), patch('app.services.document_indexer.bulk', side_effect=bulk_writer_for(self.fake_opensearch)), patch('app.services.document_indexer.ensure_legal_documents_index'), patch('app.middleware.api_protection.get_redis_client', return_value=self.fake_redis)]
        for one_patch in self._patches:
            one_patch.start()
        import importlib
        import app.main as main_module
        importlib.reload(main_module)
        self.app = main_module.app

    def tearDown(self) -> None:
        for one_patch in reversed(self._patches):
            one_patch.stop()
        for key, value in self._original_env.items():
            if value is None:
                os.environ.pop(key, None)
            else:
                os.environ[key] = value
        get_settings.cache_clear()
        self._tempdir.cleanup()

    def _auth_headers(self) -> dict[str, str]:
        return {'X-API-Key': TEST_API_ACCESS_KEY, 'X-Admin-Key': TEST_ADMIN_API_KEY}

class FreshUploadAsgiTests(AdminAsgiTestCase):

    def test_asgi_fresh_multipart_upload_returns_201(self) -> None:
        body, content_type = build_multipart_body(file_field='file', filename='Chile.docx', file_content=_CL_DOCX_BYTES, extra_fields={'replace_existing': 'false', 'confirm_warnings': 'true', 'country_confirmed': 'true'})
        response = asgi_request(self.app, 'POST', '/api/v1/admin/documents', headers={**self._auth_headers(), 'Content-Type': content_type, 'Content-Length': str(len(body))}, body=body)
        self.assertEqual(response.status_code, 201)
        payload = response.json()
        self.assertEqual(payload['status'], 'uploaded')
        self.assertEqual(payload['country_code'], 'CL')
        self.assertEqual(len(self.fake_opensearch.document_ids_for_country('CL')), 1)

class ExistingCountryAsgiTests(AdminAsgiTestCase):

    def test_asgi_existing_country_returns_structured_409(self) -> None:
        self.fake_opensearch.add(document_id='doc_' + 'a' * 64, country_code='CL', source_filename='Chile.docx')
        (self.source_dir / 'CL.docx').write_bytes(b'existing-cl-bytes')
        before_ids = set(self.fake_opensearch.document_ids_for_country('CL'))
        body, content_type = build_multipart_body(file_field='file', filename='Chile.docx', file_content=_CL_DOCX_BYTES, extra_fields={'replace_existing': 'false', 'confirm_warnings': 'true', 'country_confirmed': 'true'})
        response = asgi_request(self.app, 'POST', '/api/v1/admin/documents', headers={**self._auth_headers(), 'Content-Type': content_type, 'Content-Length': str(len(body))}, body=body)
        self.assertEqual(response.status_code, 409)
        payload = response.json()
        self.assertEqual(payload['detail']['code'], 'document_replacement_required')
        self.assertEqual(payload['detail']['country_code'], 'CL')
        self.assertEqual(set(self.fake_opensearch.document_ids_for_country('CL')), before_ids)
        self.assertEqual((self.source_dir / 'CL.docx').read_bytes(), b'existing-cl-bytes')

class FastApiValidationAsgiTests(AdminAsgiTestCase):

    def test_asgi_missing_file_returns_fastapi_detail_list(self) -> None:
        body, content_type = build_multipart_body(file_field='file', filename='unused.docx', file_content=b'unused', extra_fields={'replace_existing': 'false', 'confirm_warnings': 'true'}, include_file=False)
        response = asgi_request(self.app, 'POST', '/api/v1/admin/documents', headers={**self._auth_headers(), 'Content-Type': content_type, 'Content-Length': str(len(body))}, body=body)
        self.assertEqual(response.status_code, 422)
        payload = response.json()
        self.assertIsInstance(payload['detail'], list)
        self.assertGreater(len(payload['detail']), 0)
        self.assertIn('msg', payload['detail'][0])
        self.assertTrue(isinstance(payload['detail'][0]['msg'], str) and payload['detail'][0]['msg'])

class AuthAsgiTests(AdminAsgiTestCase):

    def test_asgi_rejects_every_incomplete_or_wrong_key_combination(self) -> None:
        cases = {'missing_api_key': {}, 'wrong_api_key': {'X-API-Key': 'wrong-key'}, 'missing_admin_key': {'X-API-Key': TEST_API_ACCESS_KEY}, 'wrong_admin_key': {'X-API-Key': TEST_API_ACCESS_KEY, 'X-Admin-Key': 'wrong-admin-key'}}
        for kind, headers in cases.items():
            with self.subTest(kind=kind):
                response = asgi_request(self.app, 'GET', '/api/v1/admin/documents', headers=headers)
                self.assertEqual(response.status_code, 401)

    def test_asgi_both_keys_correct_reaches_the_route(self) -> None:
        response = asgi_request(self.app, 'GET', '/api/v1/admin/documents', headers=self._auth_headers())
        self.assertEqual(response.status_code, 200)
        payload = response.json()
        self.assertEqual(payload['total'], 0)
import contextlib
import tempfile
import unittest
from pathlib import Path
from typing import Any
from unittest.mock import patch
from docx import Document
from opensearchpy.exceptions import OpenSearchException
from app.services.document_section_state import is_admin_modified_since_upload
from app.services.admin_document_lifecycle import AdminDocumentCountryConflictError, AdminDocumentRollbackError, reindex_indexed_document
from app.services.admin_document_sections import AdminDocumentSectionAlreadyExistsError, AdminDocumentSectionInvalidError, AdminDocumentSectionLastRemainingError, AdminDocumentSectionNotFoundError, AdminDocumentSectionPositionError, AdminDocumentSectionUpdateFailedError, add_new_section, delete_section, get_effective_section, list_effective_sections, section_id_for_legal_topic as _test_admin_sections__section_id_for_legal_topic, update_effective_section
from app.services.document_chunk_builder import DocumentMetadata, build_document_chunks
from app.services.docx_parser import ParsedSection

def _real_document_id_for(country_code: str, language: str='en') -> str:
    """
    The one real, deterministic document_id build_document_chunks
    would compute for this country_code/language - document_id is
    derived solely from country_code + DOCUMENT_FAMILY + language,
    never from anything a test happens to pick.
    """
    probe_chunks = build_document_chunks([ParsedSection(section='Employment Contracts', subsection=None, content='probe content')], DocumentMetadata(country='United Kingdom', country_code=country_code, reference_year=None, language=language, source_filename='probe.docx'))
    return probe_chunks[0].document_id
DOCUMENT_ID = _real_document_id_for('GB')
OTHER_COUNTRY_DOCUMENT_ID = _real_document_id_for('FR')
EMPLOYMENT_CONTRACTS_SECTION_ID = _test_admin_sections__section_id_for_legal_topic('Employment Contracts')
HIRING_PRACTICES_SECTION_ID = _test_admin_sections__section_id_for_legal_topic('Hiring Practices')

def _write_docx(path: Path, sections: list[tuple[str, str]]) -> None:
    """
    Build one minimal, real DOCX with a Heading 1 (a genuine L&E
    legal-topic name) followed by one content paragraph, per entry -
    every entry uses the exact same structural signal (Heading 1, no
    numbering/bold/prefix), so a custom (non-taxonomy) topic added
    after at least one real one is recognized on reparse too.

    Starts with a real "Employment Law Overview United Kingdom" H1
    overview heading, exactly like every real corpus document has -
    besides letting metadata_from_content auto-detect the country from
    the document's own content (required for the mandatory
    reparse-validation every mutation runs), it is also what flips
    past_front_matter for
    the custom-topic-recognition gate: a custom section added at
    "beginning" lands right after this heading, never before it, the
    same as in a real document.
    """
    document = Document()
    document.add_heading('Employment Law Overview United Kingdom', level=1)
    for heading, content in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(content)
    document.save(path)

def _write_bold_only_docx(path: Path, sections: list[tuple[str, str]]) -> None:
    """
    A legacy-style DOCX whose native topics use only direct bold
    run formatting, no Heading 1 style, no numbering -
    representative of the ~10/33 real corpus documents that could not
    support Add before the internal DOCX-native style marker existed.
    """
    document = Document()
    document.add_paragraph('Employment Law Overview United Kingdom')
    for heading, content in sections:
        heading_paragraph = document.add_paragraph()
        heading_paragraph.add_run(heading).bold = True
        document.add_paragraph(content)
    document.save(path)

def _seed_chunk(*, document_id: str, legal_topic: str, content: str, country_code: str='GB', country: str='United Kingdom', source_filename: str='GB.docx', reference_year: int | None=2026) -> dict[str, Any]:
    """One minimal OpenSearch-shaped chunk source dict for the fake."""
    return {'document_id': document_id, 'country_code': country_code, 'country': country, 'source_filename': source_filename, 'reference_year': reference_year, 'legal_topic': legal_topic, 'content': content, 'content_hash': f'hash-{hash(content)}'}

class FakeSectionOpenSearchClient:
    """
    OpenSearch test double for admin_document_sections.py.

    Stateful over self.chunks (chunk_id -> source dict), so a real
    edit's effect (a stale chunk removed, a new/overwritten chunk
    present) can be asserted afterwards, not merely inferred from a
    mock's own call arguments. Paired at the call site with `bulk`
    patched to write real entries into this same self.chunks (see
    _bulk_writer/_patched_indexer below).
    """

    def __init__(self, *, document_id: str=DOCUMENT_ID, country_code: str='GB', country: str='United Kingdom', source_filename: str='GB.docx', reference_year: int | None=2026, chunks: dict[str, dict[str, Any]] | None=None, fail_delete_by_query_calls: int=0, delete_by_query_failure: Exception | None=None, fail_snapshot_search: bool=False, snapshot_search_failure: Exception | None=None) -> None:
        self.document_id = document_id
        self.country_code = country_code
        self.country = country
        self.source_filename = source_filename
        self.reference_year = reference_year
        self.chunks: dict[str, dict[str, Any]] = dict(chunks or {})
        self.delete_by_query_calls: list[dict[str, Any]] = []
        self.fail_delete_by_query_calls = fail_delete_by_query_calls
        self.delete_by_query_failure = delete_by_query_failure
        self.delete_by_query_call_count = 0
        self.fail_snapshot_search = fail_snapshot_search
        self.snapshot_search_failure = snapshot_search_failure

    def search(self, index: str, body: dict[str, Any]) -> dict[str, Any]:
        del index
        term = body.get('query', {}).get('term', {})
        requested_document_id = term.get('document_id')
        requested_country_code = term.get('country_code')
        if 'sort' in body:
            if self.fail_snapshot_search:
                raise self.snapshot_search_failure if self.snapshot_search_failure is not None else RuntimeError('simulated snapshot search failure')
            if requested_document_id is not None:
                matching_ids = sorted((chunk_id for chunk_id, chunk in self.chunks.items() if chunk['document_id'] == requested_document_id))
            elif requested_country_code is not None:
                matching_ids = sorted((chunk_id for chunk_id, chunk in self.chunks.items() if chunk['country_code'] == requested_country_code))
            else:
                matching_ids = []
            return {'hits': {'total': {'value': len(matching_ids)}, 'hits': [{'_id': chunk_id, '_source': self.chunks[chunk_id], 'sort': [chunk_id]} for chunk_id in matching_ids]}}
        if requested_document_id is not None:
            if requested_document_id != self.document_id:
                return {'hits': {'hits': []}}
            return {'hits': {'hits': [{'_source': {'document_id': self.document_id, 'source_filename': self.source_filename, 'country': self.country, 'country_code': self.country_code, 'reference_year': self.reference_year}}]}}
        return {'hits': {'hits': []}}

    def delete_by_query(self, *, index: str, body: dict[str, Any], conflicts: str, refresh: bool) -> dict[str, Any]:
        del index, conflicts, refresh
        self.delete_by_query_call_count += 1
        if self.delete_by_query_call_count <= self.fail_delete_by_query_calls:
            raise self.delete_by_query_failure if self.delete_by_query_failure is not None else RuntimeError(f'simulated delete_by_query failure (call #{self.delete_by_query_call_count}).')
        query = body['query']
        keep_ids: set[str] = set()
        if 'bool' in query:
            filters = query['bool']['filter']
            document_id = next((clause['term']['document_id'] for clause in filters if 'document_id' in clause.get('term', {})))
            legal_topic = next((clause['term']['legal_topic'] for clause in filters if 'legal_topic' in clause.get('term', {})), None)
            for clause in query['bool'].get('must_not', []):
                keep_ids.update(clause.get('terms', {}).get('chunk_id', []))
        else:
            document_id = query['term']['document_id']
            legal_topic = None
        to_delete = [chunk_id for chunk_id, chunk in self.chunks.items() if chunk['document_id'] == document_id and (legal_topic is None or chunk.get('legal_topic') == legal_topic) and (chunk_id not in keep_ids)]
        for chunk_id in to_delete:
            del self.chunks[chunk_id]
        self.delete_by_query_calls.append({'document_id': document_id, 'legal_topic': legal_topic, 'deleted': len(to_delete)})
        return {'deleted': len(to_delete), 'total': len(to_delete)}

class MustNotBeCalledClient:
    """
    A client double that fails any test relying on it - proves a
    code path really does short-circuit before ever touching
    OpenSearch (e.g. an invalid section_id, or empty content).
    """

    def search(self, index: str, body: dict[str, Any]) -> dict[str, Any]:
        raise AssertionError('OpenSearch search() must not be called for this path.')

    def delete_by_query(self, **kwargs: Any) -> dict[str, Any]:
        raise AssertionError('OpenSearch delete_by_query() must not be called for this path.')

def _bulk_writer(fake_client: FakeSectionOpenSearchClient, *, fail_first_n_calls: int=0):
    """
    A `bulk()` side_effect that writes every action into fake_client's
    own self.chunks, so replace_document_section_chunks' effect on the
    fake's state is real, not merely a recorded call.

    fail_first_n_calls fails only the first N invocations (the initial
    indexing attempt) while leaving later ones (a rollback's own
    re-indexing) genuinely succeeding - mirrors
    fail_delete_by_query_calls' own counter-based semantics.
    """
    call_count = {'n': 0}

    def fake_bulk(client, actions, **kwargs):
        del client, kwargs
        call_count['n'] += 1
        if call_count['n'] <= fail_first_n_calls:
            raise RuntimeError(f"simulated OpenSearch bulk failure (call #{call_count['n']})")
        action_list = list(actions)
        for action in action_list:
            fake_client.chunks[action['_id']] = dict(action['_source'])
        return (len(action_list), [])
    return fake_bulk

@contextlib.contextmanager
def _patched_indexer(fake_client: FakeSectionOpenSearchClient, *, fail_bulk: bool=False):
    """
    Patch the two document_indexer.py internals
    replace_document_section_chunks calls directly - the one seam
    admin_document_sections.py has none for.

    fail_bulk fails only the FIRST bulk call (the initial indexing
    attempt) - a subsequent rollback's own re-indexing call still
    succeeds, exactly like a real, isolated OpenSearch write failure
    would (never every future call forever).
    """
    with patch('app.services.document_indexer.ensure_legal_documents_index'), patch('app.services.document_indexer.bulk', side_effect=_bulk_writer(fake_client, fail_first_n_calls=1 if fail_bulk else 0)):
        yield

def _seeded_client(*, topics: list[tuple[str, str]], document_id: str=DOCUMENT_ID, country_code: str='GB', country: str='United Kingdom', source_filename: str='GB.docx') -> FakeSectionOpenSearchClient:
    return FakeSectionOpenSearchClient(document_id=document_id, country_code=country_code, country=country, source_filename=source_filename, chunks={f'chunk-seed-{index}': _seed_chunk(document_id=document_id, legal_topic=topic, content=content, country_code=country_code, country=country, source_filename=source_filename) for index, (topic, content) in enumerate(topics)})

class AdminDocumentSectionListTests(unittest.TestCase):

    def test_lists_every_real_topic_in_the_current_docx(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'stale opensearch text')])
            response = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(response.document_id, DOCUMENT_ID)
            self.assertEqual([section.legal_topic for section in response.sections], ['Employment Contracts', 'Hiring Practices'])
            self.assertEqual({section.section_id for section in response.sections}, {EMPLOYMENT_CONTRACTS_SECTION_ID, HIRING_PRACTICES_SECTION_ID})

class AdminDocumentSectionGetTests(unittest.TestCase):

    def test_content_always_comes_from_the_current_docx(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'DOCX text for employment contracts.')])
            client = _seeded_client(topics=[('Employment Contracts', 'irrelevant stale chunk text')])
            response = get_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(response.content, 'DOCX text for employment contracts.')
            self.assertEqual(response.country_code, 'GB')
            self.assertEqual(response.country_name, 'United Kingdom')
            self.assertEqual(response.legal_topic, 'Employment Contracts')

    def test_invalid_section_id_is_not_found_without_touching_opensearch(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.')])
            with self.assertRaises(AdminDocumentSectionNotFoundError) as context:
                get_effective_section(document_id=DOCUMENT_ID, section_id='not-a-real-topic-slug', source_directory=source_directory, client=_seeded_client(topics=[('Employment Contracts', 'x')]))
            self.assertEqual(context.exception.to_detail()['code'], 'document_section_not_found')

class AdminDocumentSectionEditTests(unittest.TestCase):
    """
    Edit really modifies the current DOCX, validated by a full
    reparse, applied atomically to both OpenSearch
    (targeted to this one legal_topic) and the source file.
    """

    def test_edit_writes_the_current_docx_and_removes_old_content(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                response = update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='  New effective content.  ', source_directory=source_directory, client=client)
            self.assertEqual(response.legal_topic, 'Employment Contracts')
            self.assertEqual(response.indexed_chunks, 1)
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertIn('Employment Contracts', paragraphs)
            self.assertIn('New effective content.', paragraphs)
            self.assertNotIn('Original DOCX text.', paragraphs)
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Employment Contracts']
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]['content'], 'New effective content.')
            reparsed = get_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(reparsed.content, 'New effective content.')

    def test_unrelated_topics_are_exactly_unchanged(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original EC text.'), ('Hiring Practices', 'Original HP text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Edited EC content.', source_directory=source_directory, client=client)
            hp_chunks = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Hiring Practices']
            self.assertEqual(len(hp_chunks), 1)
            self.assertEqual(hp_chunks[0]['content'], 'placeholder HP')
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertIn('Original HP text.', paragraphs)

    def test_reindex_after_edit_reproduces_the_same_state(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Edited content survives reindex.', source_directory=source_directory, client=client)
            with patch('app.services.document_indexer.ensure_legal_documents_index'), patch('app.services.document_indexer.bulk', side_effect=_bulk_writer(client)):
                reindex_indexed_document(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            reparsed = get_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(reparsed.content, 'Edited content survives reindex.')

    def test_second_edit_fully_overwrites_first_no_history(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='First edit content.', source_directory=source_directory, client=client)
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Second edit content - final.', source_directory=source_directory, client=client)
            response = get_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(response.content, 'Second edit content - final.')
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Employment Contracts']
            self.assertEqual(len(remaining), 1)

    def test_empty_content_is_invalid_with_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            original_bytes = source_path.read_bytes()
            with self.assertRaises(AdminDocumentSectionInvalidError) as context:
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='   \n\t  ', source_directory=source_directory, client=MustNotBeCalledClient())
            self.assertEqual(context.exception.to_detail()['code'], 'document_section_invalid')
            self.assertEqual(source_path.read_bytes(), original_bytes)

    def test_unknown_section_is_not_found_with_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with self.assertRaises(AdminDocumentSectionNotFoundError):
                update_effective_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, new_content='Some content.', source_directory=source_directory, client=client)
            self.assertEqual(len(client.chunks), 1)

class AdminDocumentSectionAddTests(unittest.TestCase):
    """Adding a brand-new top-level topic."""

    def test_add_at_end_is_editable_and_retrievable(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                response = add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='Employees may work remotely. MARKER.', position='end', source_directory=source_directory, client=client)
            self.assertEqual(response.legal_topic, 'Remote Working')
            self.assertEqual(response.section_id, _test_admin_sections__section_id_for_legal_topic('Remote Working'))
            listing = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual([s.legal_topic for s in listing.sections], ['Employment Contracts', 'Remote Working'])
            fetched = get_effective_section(document_id=DOCUMENT_ID, section_id=_test_admin_sections__section_id_for_legal_topic('Remote Working'), source_directory=source_directory, client=client)
            self.assertIn('MARKER', fetched.content)
            remote_chunks = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Remote Working']
            self.assertEqual(len(remote_chunks), 1)
            self.assertIn('MARKER', remote_chunks[0]['content'])

    def test_add_at_beginning_and_after_existing_section(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder')])
            with _patched_indexer(client):
                add_new_section(document_id=DOCUMENT_ID, title='Artificial Intelligence at Work', content='AI disclosure rules.', position='beginning', source_directory=source_directory, client=client)
            listing = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual(listing.sections[0].legal_topic, 'Artificial Intelligence at Work')
            with _patched_indexer(client):
                add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='Remote work rules.', position=f'after:{EMPLOYMENT_CONTRACTS_SECTION_ID}', source_directory=source_directory, client=client)
            listing_after = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            names = [s.legal_topic for s in listing_after.sections]
            self.assertEqual(names, ['Artificial Intelligence at Work', 'Employment Contracts', 'Remote Working', 'Hiring Practices'])

    def test_duplicate_title_is_rejected_with_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with self.assertRaises(AdminDocumentSectionAlreadyExistsError) as context:
                add_new_section(document_id=DOCUMENT_ID, title='employment   CONTRACTS', content='whatever', position='end', source_directory=source_directory, client=client)
            self.assertEqual(context.exception.to_detail()['code'], 'section_already_exists')
            self.assertEqual(source_path.read_bytes(), original_bytes)
            self.assertEqual(len(client.chunks), 1)

    def test_invalid_position_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with self.assertRaises(AdminDocumentSectionPositionError):
                add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='content', position='middle', source_directory=source_directory, client=client)

    def test_custom_section_survives_full_reindex(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='Remote work rules. MARKER-REMOTE.', position='end', source_directory=source_directory, client=client)
            with patch('app.services.document_indexer.ensure_legal_documents_index'), patch('app.services.document_indexer.bulk', side_effect=_bulk_writer(client)):
                reindex_indexed_document(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            listing = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertIn('Remote Working', [s.legal_topic for s in listing.sections])
            remote_chunks = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Remote Working']
            self.assertEqual(len(remote_chunks), 1)
            self.assertIn('MARKER-REMOTE', remote_chunks[0]['content'])

class AdminDocumentSectionTransactionTests(unittest.TestCase):
    """
    The exact transaction order and rollback matrix. Every failure
    scenario must leave either the fully-applied
    new state, or exactly the pre-operation state - never a mix.
    """

    def test_zero_mutation_when_target_topic_does_not_exist(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with self.assertRaises(AdminDocumentSectionNotFoundError):
                update_effective_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, new_content='should never apply', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            self.assertEqual(len(client.chunks), 1)
            leftover_temps = list(source_directory.glob('.*tmp.docx'))
            self.assertEqual(leftover_temps, [])

    def test_opensearch_bulk_failure_leaves_source_and_index_untouched(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                    update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should roll back', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            self.assertEqual(client.chunks['chunk-seed-0']['content'], 'placeholder')
            leftover_temps = list(source_directory.glob('.*tmp.docx'))
            self.assertEqual(leftover_temps, [])

    def test_stale_delete_failure_rolls_back_to_pre_edit_snapshot(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            client.fail_delete_by_query_calls = 1
            with _patched_indexer(client):
                with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                    update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should roll back', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Employment Contracts']
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]['content'], 'placeholder')

    def test_source_replace_failure_rolls_back_opensearch(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                with patch('app.services.admin_document_sections.os.replace', side_effect=OSError('simulated disk failure')):
                    with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                        update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should roll back', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Employment Contracts']
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]['content'], 'placeholder')
            leftover_temps = list(source_directory.glob('.*tmp.docx'))
            self.assertEqual(leftover_temps, [])

    def test_rollback_failure_after_source_replace_failure_is_surfaced(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                with patch('app.services.admin_document_sections.os.replace', side_effect=OSError('simulated disk failure')), patch('app.services.admin_document_sections._restore_section_snapshot', side_effect=RuntimeError('simulated rollback failure')):
                    with self.assertRaises(AdminDocumentRollbackError):
                        update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should surface rollback failure', source_directory=source_directory, client=client)

class AdminDocumentCountryInvariantTests(unittest.TestCase):
    """One active document per country."""

    def test_country_conflict_blocks_edit_with_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            client.chunks['chunk-conflict'] = _seed_chunk(document_id='doc_' + 'c' * 64, legal_topic='Hiring Practices', content='legacy duplicate content', country_code='GB', country='United Kingdom', source_filename='GB-legacy.docx')
            with self.assertRaises(AdminDocumentCountryConflictError):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should never apply', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)

    def test_country_conflict_blocks_add_with_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            client.chunks['chunk-conflict'] = _seed_chunk(document_id='doc_' + 'c' * 64, legal_topic='Hiring Practices', content='legacy duplicate content', country_code='GB', country='United Kingdom', source_filename='GB-legacy.docx')
            with self.assertRaises(AdminDocumentCountryConflictError):
                add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='content', position='end', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)

    def test_no_conflict_when_only_one_active_document_exists(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='applies normally', source_directory=source_directory, client=client)
            response = get_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(response.content, 'applies normally')

    def test_country_lookup_failure_is_never_a_raw_exception(self) -> None:
        from app.services.admin_document_lifecycle import AdminDocumentLifecycleError
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with patch('app.services.admin_document_lifecycle.lookup_existing_country_documents', side_effect=RuntimeError('simulated transient failure')):
                with self.assertRaises(AdminDocumentLifecycleError):
                    update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should not apply', source_directory=source_directory, client=client)

class AdminSectionBoldOnlyDocumentAddTests(unittest.TestCase):
    """
    Add must work end-to-end through the full service layer on a
    document whose native topics use only
    bold formatting - previously unsupported before the internal
    DOCX-native style marker existed.
    """

    def test_add_list_get_edit_on_bold_only_document(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_bold_only_docx(source_path, [('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Hiring Practices', 'placeholder')])
            with _patched_indexer(client):
                add_result = add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='Employees may work remotely. MARKER.', position='end', source_directory=source_directory, client=client)
            self.assertEqual(add_result.legal_topic, 'Remote Working')
            listing = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertEqual([s.legal_topic for s in listing.sections], ['Hiring Practices', 'Remote Working'])
            remote_section_id = _test_admin_sections__section_id_for_legal_topic('Remote Working')
            fetched = get_effective_section(document_id=DOCUMENT_ID, section_id=remote_section_id, source_directory=source_directory, client=client)
            self.assertIn('MARKER', fetched.content)
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=remote_section_id, new_content='Overwritten remote content.', source_directory=source_directory, client=client)
            fetched_after_edit = get_effective_section(document_id=DOCUMENT_ID, section_id=remote_section_id, source_directory=source_directory, client=client)
            self.assertEqual(fetched_after_edit.content, 'Overwritten remote content.')
            self.assertNotIn('MARKER', fetched_after_edit.content)
            hiring_fetched = get_effective_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(hiring_fetched.content, 'HP content.')

    def test_reindex_preserves_custom_section_on_bold_only_document(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_bold_only_docx(source_path, [('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Hiring Practices', 'placeholder')])
            with _patched_indexer(client):
                add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='Remote content. MARKER.', position='end', source_directory=source_directory, client=client)
            with patch('app.services.document_indexer.ensure_legal_documents_index'), patch('app.services.document_indexer.bulk', side_effect=_bulk_writer(client)):
                reindex_indexed_document(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            listing = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            self.assertIn('Remote Working', [s.legal_topic for s in listing.sections])

    def test_duplicate_custom_title_rejected_on_bold_only_document(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_bold_only_docx(source_path, [('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Hiring Practices', 'placeholder')])
            with self.assertRaises(AdminDocumentSectionAlreadyExistsError):
                add_new_section(document_id=DOCUMENT_ID, title='hiring   PRACTICES', content='whatever', position='end', source_directory=source_directory, client=client)

class AdminDocumentSectionRenameTests(unittest.TestCase):
    """
    5: Rename extends the same Edit
    Section transaction - an omitted or effectively-unchanged title is
    a normal content-only edit, never a fake rename; a genuine title
    change re-validates the reparsed document (old title gone, new
    title exactly once, topic count/other topics unchanged) before
    touching OpenSearch, then removes the old topic's chunks once the
    new ones are safely indexed, atomically replaces the source, and
    verifies both invariants (new topic present, old topic absent).
    """

    def test_native_canonical_section_rename_with_new_content(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.'), ('Hiring Practices', 'Unrelated HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                response = update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Renamed section content.', new_title='Remote Work Equipment Requirements', source_directory=source_directory, client=client)
            self.assertEqual(response.legal_topic, 'Remote Work Equipment Requirements')
            self.assertEqual(response.section_id, _test_admin_sections__section_id_for_legal_topic('Remote Work Equipment Requirements'))
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertIn('Remote Work Equipment Requirements', paragraphs)
            self.assertNotIn('Employment Contracts', paragraphs)
            self.assertIn('Renamed section content.', paragraphs)
            remaining_new = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Remote Work Equipment Requirements']
            self.assertEqual(len(remaining_new), 1)
            remaining_old = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Employment Contracts']
            self.assertEqual(remaining_old, [])

    def test_custom_section_rename_title_only_content_preserved(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC')])
            with _patched_indexer(client):
                add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='Original custom content, unchanged.', position='end', source_directory=source_directory, client=client)
                response = update_effective_section(document_id=DOCUMENT_ID, section_id=_test_admin_sections__section_id_for_legal_topic('Remote Working'), new_content='Original custom content, unchanged.', new_title='Remote Working Policy', source_directory=source_directory, client=client)
            self.assertEqual(response.legal_topic, 'Remote Working Policy')
            reparsed = get_effective_section(document_id=DOCUMENT_ID, section_id=_test_admin_sections__section_id_for_legal_topic('Remote Working Policy'), source_directory=source_directory, client=client)
            self.assertEqual(reparsed.content, 'Original custom content, unchanged.')

    def test_duplicate_title_is_rejected_with_zero_mutation(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            with self.assertRaises(AdminDocumentSectionAlreadyExistsError) as context:
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='whatever', new_title='Hiring Practices', source_directory=source_directory, client=client)
            self.assertEqual(context.exception.to_detail()['operation'], 'section_update')
            self.assertEqual(source_path.read_bytes(), original_bytes)
            self.assertEqual(len(client.chunks), 2)

    def test_normalized_duplicate_title_is_rejected(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            with self.assertRaises(AdminDocumentSectionAlreadyExistsError):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='whatever', new_title='  hiring   PRACTICES  ', source_directory=source_directory, client=client)

    def test_unchanged_effective_title_is_a_normal_edit_not_a_rename(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                response = update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Edited content only.', new_title='  employment   contracts  ', source_directory=source_directory, client=client)
            self.assertEqual(response.legal_topic, 'Employment Contracts')
            self.assertEqual(response.section_id, EMPLOYMENT_CONTRACTS_SECTION_ID)
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertIn('Employment Contracts', paragraphs)

    def test_stale_section_id_after_rename_is_not_found(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'Unrelated HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Renamed away.', new_title='Something Else Entirely', source_directory=source_directory, client=client)
                with self.assertRaises(AdminDocumentSectionNotFoundError):
                    update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should not apply', source_directory=source_directory, client=client)

    def test_other_topics_unchanged_by_rename(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Renamed content.', new_title='Something Else Entirely', source_directory=source_directory, client=client)
            hp_chunks = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Hiring Practices']
            self.assertEqual(len(hp_chunks), 1)
            self.assertEqual(hp_chunks[0]['content'], 'placeholder HP')
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertIn('HP content.', paragraphs)

    def test_new_topic_index_failure_rolls_back_rename(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.'), ('Hiring Practices', 'Unrelated HP content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                    update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should roll back', new_title='Something Else Entirely', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            self.assertEqual(client.chunks['chunk-seed-0']['legal_topic'], 'Employment Contracts')
            self.assertEqual(client.chunks['chunk-seed-0']['content'], 'placeholder')
            leftover_temps = list(source_directory.glob('.*tmp.docx'))
            self.assertEqual(leftover_temps, [])

    def test_old_topic_cleanup_failure_rolls_back_new_topic_too(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.'), ('Hiring Practices', 'Unrelated HP content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            real_delete_by_query = client.delete_by_query
            call_count = {'n': 0}

            def delete_by_query_fail_second(**kwargs: Any) -> Any:
                call_count['n'] += 1
                if call_count['n'] == 2:
                    raise OpenSearchException('simulated old-topic cleanup failure')
                return real_delete_by_query(**kwargs)
            client.delete_by_query = delete_by_query_fail_second
            with _patched_indexer(client):
                with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                    update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should roll back', new_title='Something Else Entirely', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Employment Contracts']
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]['content'], 'placeholder')
            leftover_new = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Something Else Entirely']
            self.assertEqual(leftover_new, [])

    def test_docx_replace_failure_rolls_back_rename_entirely(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.'), ('Hiring Practices', 'Unrelated HP content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                with patch('app.services.admin_document_sections.os.replace', side_effect=OSError('simulated disk failure')):
                    with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                        update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should roll back', new_title='Something Else Entirely', source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Employment Contracts']
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]['content'], 'placeholder')
            leftover_new = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Something Else Entirely']
            self.assertEqual(leftover_new, [])
            leftover_temps = list(source_directory.glob('.*tmp.docx'))
            self.assertEqual(leftover_temps, [])

    def test_live_topic_vocabulary_updates_immediately_after_rename(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'Original DOCX text.'), ('Hiring Practices', 'Unrelated HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Renamed content.', new_title='Something Else Entirely', source_directory=source_directory, client=client)
            listing = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            listed_topics = [s.legal_topic for s in listing.sections]
            self.assertIn('Something Else Entirely', listed_topics)
            self.assertNotIn('Employment Contracts', listed_topics)

class AdminDocumentSectionDeleteTests(unittest.TestCase):
    """
    8: Delete mirrors Edit/Rename's
    exact transaction shape for one already-known legal_topic (no new
    chunks are ever built) - lock, snapshot, mutate a temp copy,
    reparse-validate (target gone, every other topic unchanged),
    delete the target's chunks, atomically replace the source, verify.
    Blocks deleting the document's last remaining usable section.
    """

    def test_custom_section_delete(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC')])
            with _patched_indexer(client):
                add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='Custom content.', position='end', source_directory=source_directory, client=client)
                response = delete_section(document_id=DOCUMENT_ID, section_id=_test_admin_sections__section_id_for_legal_topic('Remote Working'), source_directory=source_directory, client=client)
            self.assertEqual(response.legal_topic, 'Remote Working')
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertNotIn('Remote Working', paragraphs)
            self.assertNotIn('Custom content.', paragraphs)
            self.assertIn('EC content.', paragraphs)
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Remote Working']
            self.assertEqual(remaining, [])

    def test_native_canonical_section_delete(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                response = delete_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(response.legal_topic, 'Hiring Practices')
            paragraphs = [p.text for p in Document(source_path).paragraphs]
            self.assertNotIn('Hiring Practices', paragraphs)
            self.assertNotIn('HP content.', paragraphs)
            self.assertIn('Employment Contracts', paragraphs)
            self.assertIn('EC content.', paragraphs)
            hp_chunks = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Hiring Practices']
            self.assertEqual(hp_chunks, [])
            ec_chunks = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Employment Contracts']
            self.assertEqual(len(ec_chunks), 1)

    def test_last_remaining_section_delete_is_blocked(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with self.assertRaises(AdminDocumentSectionLastRemainingError) as context:
                delete_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(context.exception.to_detail()['code'], 'section_is_last_remaining')
            self.assertEqual(source_path.read_bytes(), original_bytes)
            self.assertEqual(len(client.chunks), 1)

    def test_other_topics_unchanged_by_delete(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                delete_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, source_directory=source_directory, client=client)
            reparsed = get_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(reparsed.content, 'EC content.')

    def test_opensearch_delete_failure_rolls_back(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            client.fail_delete_by_query_calls = 1
            client.delete_by_query_failure = OpenSearchException('simulated delete failure')
            with _patched_indexer(client):
                with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                    delete_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Hiring Practices']
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]['content'], 'placeholder HP')

    def test_docx_replace_failure_rolls_back_delete(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            original_bytes = source_path.read_bytes()
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                with patch('app.services.admin_document_sections.os.replace', side_effect=OSError('simulated disk failure')):
                    with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                        delete_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, source_directory=source_directory, client=client)
            self.assertEqual(source_path.read_bytes(), original_bytes)
            remaining = [chunk for chunk in client.chunks.values() if chunk['legal_topic'] == 'Hiring Practices']
            self.assertEqual(len(remaining), 1)
            self.assertEqual(remaining[0]['content'], 'placeholder HP')
            leftover_temps = list(source_directory.glob('.*tmp.docx'))
            self.assertEqual(leftover_temps, [])

    def test_live_topic_vocabulary_removes_deleted_title(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            _write_docx(source_directory / 'GB.docx', [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder EC'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                delete_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, source_directory=source_directory, client=client)
            listing = list_effective_sections(document_id=DOCUMENT_ID, source_directory=source_directory, client=client)
            listed_topics = [s.legal_topic for s in listing.sections]
            self.assertNotIn('Hiring Practices', listed_topics)
            self.assertIn('Employment Contracts', listed_topics)

class AdminModifiedMarkerHookTests(unittest.TestCase):
    """
    The generic admin-modified-since-upload marker (shared with Admin
    Contact mutations) now also marks dirty on a successful Section
    mutation, and stays untouched on a failed one. Does not retest
    Section business behavior itself (already covered exhaustively by
    the classes above).
    """

    def test_edit_marks_admin_modified(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            self.assertFalse(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='New effective content.', source_directory=source_directory, client=client)
            self.assertTrue(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_rename_marks_admin_modified(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.'), ('Hiring Practices', 'Unrelated HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='Renamed section content.', new_title='Remote Work Equipment Requirements', source_directory=source_directory, client=client)
            self.assertTrue(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_add_marks_admin_modified(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client):
                add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='Employees may work remotely.', position='end', source_directory=source_directory, client=client)
            self.assertTrue(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_delete_marks_admin_modified(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            with _patched_indexer(client):
                delete_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, source_directory=source_directory, client=client)
            self.assertTrue(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_failed_edit_does_not_mark_admin_modified(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'Original DOCX text.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                    update_effective_section(document_id=DOCUMENT_ID, section_id=EMPLOYMENT_CONTRACTS_SECTION_ID, new_content='should roll back', source_directory=source_directory, client=client)
            self.assertFalse(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_failed_add_does_not_mark_admin_modified(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder')])
            with _patched_indexer(client, fail_bulk=True):
                with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                    add_new_section(document_id=DOCUMENT_ID, title='Remote Working', content='should roll back', position='end', source_directory=source_directory, client=client)
            self.assertFalse(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))

    def test_failed_delete_does_not_mark_admin_modified(self) -> None:
        with tempfile.TemporaryDirectory() as root:
            source_directory = Path(root)
            source_path = source_directory / 'GB.docx'
            _write_docx(source_path, [('Employment Contracts', 'EC content.'), ('Hiring Practices', 'HP content.')])
            client = _seeded_client(topics=[('Employment Contracts', 'placeholder'), ('Hiring Practices', 'placeholder HP')])
            client.fail_delete_by_query_calls = 1
            client.delete_by_query_failure = OpenSearchException('simulated delete failure')
            with _patched_indexer(client):
                with self.assertRaises(AdminDocumentSectionUpdateFailedError):
                    delete_section(document_id=DOCUMENT_ID, section_id=HIRING_PRACTICES_SECTION_ID, source_directory=source_directory, client=client)
            self.assertFalse(is_admin_modified_since_upload(source_directory, DOCUMENT_ID))
