"""
Tests for app.services.document_mutation - ORDER 8A, section 7: the
reusable primitive that physically replaces or inserts a top-level
legal topic in a copy of a source DOCX.
"""

import unittest
from pathlib import Path
from tempfile import TemporaryDirectory

from docx import Document

from app.services.docx_parser import locate_top_level_topics, parse_docx_sections
from app.services.document_mutation import (
    InvalidSectionPositionError,
    LegalTopicAlreadyExistsError,
    LegalTopicNotFoundError,
    NoAnchorTopicError,
    insert_top_level_topic,
    normalize_topic_title,
    replace_top_level_topic,
)


def _write_docx(path: Path, sections: list[tuple[str, str]]) -> None:
    document = Document()
    document.add_heading(
        "Employment Law Overview United Kingdom", level=1
    )

    for heading, content in sections:
        document.add_heading(heading, level=1)
        document.add_paragraph(content)

    document.save(path)


def _topic_names(path: Path, country: str = "United Kingdom") -> list[str]:
    return [
        location.legal_topic
        for location in locate_top_level_topics(
            Document(path), country=country
        )
    ]


def _write_bold_only_docx(
    path: Path, sections: list[tuple[str, str]]
) -> None:
    """
    A legacy-style document whose native topics use only direct bold
    run formatting - no Heading 1 style, no numbering - representative
    of the ~10/33 real corpus documents ORDER 8A-C's marker style
    exists to support.
    """

    document = Document()
    document.add_paragraph("Employment Law Overview United Kingdom")

    for heading, content in sections:
        heading_paragraph = document.add_paragraph()
        heading_paragraph.add_run(heading).bold = True
        document.add_paragraph(content)

    document.save(path)


class NormalizeTopicTitleTests(unittest.TestCase):
    def test_case_and_whitespace_insensitive(self) -> None:
        self.assertEqual(
            normalize_topic_title("Hiring Practices"),
            normalize_topic_title("hiring   PRACTICES"),
        )

    def test_ignores_leading_numeric_prefix(self) -> None:
        self.assertEqual(
            normalize_topic_title("01. Hiring Practices"),
            normalize_topic_title("Hiring Practices"),
        )


class ReplaceTopLevelTopicTests(unittest.TestCase):
    def test_replaces_content_and_preserves_heading(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [("Hiring Practices", "Original content.")],
            )

            replace_top_level_topic(
                file_path=source_path,
                output_path=output_path,
                country="United Kingdom",
                legal_topic="Hiring Practices",
                new_content="New content.\n\nSecond paragraph.",
            )

            paragraphs = [
                p.text for p in Document(output_path).paragraphs
            ]
            self.assertIn("Hiring Practices", paragraphs)
            self.assertIn("New content.", paragraphs)
            self.assertIn("Second paragraph.", paragraphs)
            self.assertNotIn("Original content.", paragraphs)

    def test_unrelated_topics_are_byte_identical(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [
                    ("Hiring Practices", "HP content."),
                    ("Employment Contracts", "EC content."),
                ],
            )

            replace_top_level_topic(
                file_path=source_path,
                output_path=output_path,
                country="United Kingdom",
                legal_topic="Hiring Practices",
                new_content="Edited HP content.",
            )

            sections = parse_docx_sections(
                output_path, country="United Kingdom"
            )
            ec_sections = [
                s for s in sections if s.section == "Employment Contracts"
            ]
            self.assertEqual(len(ec_sections), 1)
            self.assertEqual(ec_sections[0].content, "EC content.")

    def test_missing_topic_raises_not_found(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [("Hiring Practices", "Original content.")],
            )

            with self.assertRaises(LegalTopicNotFoundError):
                replace_top_level_topic(
                    file_path=source_path,
                    output_path=output_path,
                    country="United Kingdom",
                    legal_topic="Employment Contracts",
                    new_content="New content.",
                )

            self.assertFalse(output_path.exists())


class InsertTopLevelTopicTests(unittest.TestCase):
    def test_insert_at_end(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=output_path,
                country="United Kingdom",
                title="Remote Working",
                content="Remote work content.",
                position="end",
            )

            names = _topic_names(output_path)
            self.assertEqual(
                names, ["Hiring Practices", "Remote Working"]
            )

    def test_insert_at_beginning(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=output_path,
                country="United Kingdom",
                title="Remote Working",
                content="Remote work content.",
                position="beginning",
            )

            names = _topic_names(output_path)
            self.assertEqual(
                names, ["Remote Working", "Hiring Practices"]
            )

    def test_insert_after_existing_topic(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [
                    ("Hiring Practices", "HP content."),
                    ("Employment Contracts", "EC content."),
                ],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=output_path,
                country="United Kingdom",
                title="Remote Working",
                content="Remote work content.",
                position="after:Hiring Practices",
            )

            names = _topic_names(output_path)
            self.assertEqual(
                names,
                [
                    "Hiring Practices",
                    "Remote Working",
                    "Employment Contracts",
                ],
            )

    def test_duplicate_title_is_rejected(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            with self.assertRaises(LegalTopicAlreadyExistsError):
                insert_top_level_topic(
                    file_path=source_path,
                    output_path=output_path,
                    country="United Kingdom",
                    title="hiring   PRACTICES",
                    content="whatever",
                    position="end",
                )

            self.assertFalse(output_path.exists())

    def test_invalid_position_is_rejected(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            with self.assertRaises(InvalidSectionPositionError):
                insert_top_level_topic(
                    file_path=source_path,
                    output_path=output_path,
                    country="United Kingdom",
                    title="Remote Working",
                    content="whatever",
                    position="middle",
                )

    def test_no_anchor_topic_raises(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            document = Document()
            document.add_paragraph("Just some plain text, no topics.")
            document.save(source_path)

            with self.assertRaises(NoAnchorTopicError):
                insert_top_level_topic(
                    file_path=source_path,
                    output_path=output_path,
                    country="United Kingdom",
                    title="Remote Working",
                    content="whatever",
                    position="end",
                )


class AdminSectionStyleMarkerTests(unittest.TestCase):
    """
    ORDER 8A-C - the internal DOCX-native style marker: created once,
    reused afterward, works identically regardless of the surrounding
    document's own native heading convention.
    """

    def test_insert_on_bold_only_document(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_bold_only_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=output_path,
                country="United Kingdom",
                title="Remote Working",
                content="Remote work content.",
                position="end",
            )

            names = _topic_names(output_path)
            self.assertEqual(
                names, ["Hiring Practices", "Remote Working"]
            )

            # native topic's own content is untouched
            sections = parse_docx_sections(
                output_path, country="United Kingdom"
            )
            hiring = [
                s for s in sections if s.section == "Hiring Practices"
            ]
            self.assertEqual(len(hiring), 1)
            self.assertEqual(hiring[0].content, "HP content.")

    def test_marker_style_created_once_and_reused(self) -> None:
        from app.services.docx_parser import ADMIN_SECTION_STYLE_NAME

        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            first_output = Path(temp_dir) / "first.docx"
            second_output = Path(temp_dir) / "second.docx"

            _write_bold_only_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=first_output,
                country="United Kingdom",
                title="Remote Working",
                content="Remote content.",
                position="end",
            )

            first_styles = [
                style.name for style in Document(first_output).styles
            ]
            self.assertEqual(
                first_styles.count(ADMIN_SECTION_STYLE_NAME), 1
            )

            insert_top_level_topic(
                file_path=first_output,
                output_path=second_output,
                country="United Kingdom",
                title="Artificial Intelligence at Work",
                content="AI content.",
                position="end",
            )

            second_styles = [
                style.name for style in Document(second_output).styles
            ]
            # still exactly one definition of the marker style - the
            # second Add reused it rather than creating a duplicate.
            self.assertEqual(
                second_styles.count(ADMIN_SECTION_STYLE_NAME), 1
            )

            names = _topic_names(second_output)
            self.assertEqual(
                names,
                [
                    "Hiring Practices",
                    "Remote Working",
                    "Artificial Intelligence at Work",
                ],
            )

    def test_visual_formatting_derived_from_bold_only_anchor(
        self,
    ) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_bold_only_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=output_path,
                country="United Kingdom",
                title="Remote Working",
                content="Remote content.",
                position="end",
            )

            reloaded = Document(output_path)
            new_heading = next(
                p
                for p in reloaded.paragraphs
                if p.text == "Remote Working"
            )
            # Cosmetic formatting is reasonable (bold, non-trivial
            # size) - the identity check itself never depends on this.
            self.assertTrue(new_heading.style.font.bold)
            self.assertIsNotNone(new_heading.style.font.size)

    def test_visual_formatting_derived_from_heading1_anchor(
        self,
    ) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            output_path = Path(temp_dir) / "output.docx"

            _write_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=output_path,
                country="United Kingdom",
                title="Remote Working",
                content="Remote content.",
                position="end",
            )

            reloaded = Document(output_path)
            heading_one = reloaded.styles["Heading 1"]
            new_heading = next(
                p
                for p in reloaded.paragraphs
                if p.text == "Remote Working"
            )

            self.assertEqual(
                new_heading.style.font.bold,
                heading_one.font.bold,
            )
            self.assertEqual(
                new_heading.style.font.size,
                heading_one.font.size,
            )

    def test_after_another_custom_section(self) -> None:
        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            first_output = Path(temp_dir) / "first.docx"
            second_output = Path(temp_dir) / "second.docx"

            _write_bold_only_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=first_output,
                country="United Kingdom",
                title="Custom A",
                content="A content.",
                position="end",
            )

            insert_top_level_topic(
                file_path=first_output,
                output_path=second_output,
                country="United Kingdom",
                title="Custom B",
                content="B content.",
                position="after:Custom A",
            )

            names = _topic_names(second_output)
            self.assertEqual(
                names,
                ["Hiring Practices", "Custom A", "Custom B"],
            )

    def test_download_reparse_roundtrip_no_external_state(self) -> None:
        # ORDER 8A-C, section 7: a custom section must still be
        # detected after the DOCX is saved, copied ("downloaded"),
        # and reparsed completely standalone - no external state file
        # of any kind is ever consulted.
        import shutil

        with TemporaryDirectory() as temp_dir:
            source_path = Path(temp_dir) / "source.docx"
            added_path = Path(temp_dir) / "added.docx"
            downloaded_path = Path(temp_dir) / "downloaded_copy.docx"

            _write_bold_only_docx(
                source_path,
                [("Hiring Practices", "HP content.")],
            )

            insert_top_level_topic(
                file_path=source_path,
                output_path=added_path,
                country="United Kingdom",
                title="Remote Working",
                content="Remote content. ROUNDTRIP-MARKER.",
                position="end",
            )

            shutil.copy(added_path, downloaded_path)

            from app.services.document_chunk_builder import (
                build_document_chunks_from_docx,
            )

            names = _topic_names(downloaded_path)
            self.assertIn("Remote Working", names)

            sections = parse_docx_sections(
                downloaded_path, country="United Kingdom"
            )
            remote_sections = [
                s for s in sections if s.section == "Remote Working"
            ]
            self.assertEqual(len(remote_sections), 1)
            self.assertIn(
                "ROUNDTRIP-MARKER", remote_sections[0].content
            )

            # simulate a confirmed re-upload of this exact file
            reuploaded_chunks = build_document_chunks_from_docx(
                downloaded_path
            )
            remote_chunks = [
                c
                for c in reuploaded_chunks
                if c.legal_topic == "Remote Working"
            ]
            self.assertEqual(len(remote_chunks), 1)
            self.assertEqual(remote_chunks[0].document_type, "comparator")


if __name__ == "__main__":
    unittest.main()
