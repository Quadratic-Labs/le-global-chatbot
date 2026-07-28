from __future__ import annotations

import json
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Final

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.core.legal_taxonomy import (
    get_canonical_legal_topic,
    is_overview_section,
)


_FILENAME_PATTERNS: Final[tuple[re.Pattern[str], ...]] = (
    re.compile(
        r"^Labour and Employment Law in\s+"
        r"(?P<country>.+?)"
        r"(?:\s+(?:19|20)\d{2})?$",
        re.IGNORECASE,
    ),
    re.compile(
        r"^Employment Law Overview(?:\s*-\s*)?\s+"
        r"(?P<country>.+?)"
        r"(?:\s+(?:19|20)\d{2})?$",
        re.IGNORECASE,
    ),
)

_COPY_SUFFIX_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"\s*\(\d+\)$"
)

_HEADING_STYLE_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"^(?:heading|titre)\s*([1-9])$",
    re.IGNORECASE,
)

_FALSE_XML_VALUES: Final[frozenset[str]] = frozenset(
    {
        "0",
        "false",
        "off",
    }
)

_COUNTRY_ALIASES: Final[dict[str, str]] = {
    "uk": "United Kingdom",
    "united kingdom": "United Kingdom",
}

_MAX_UNMATCHED_PER_DOCUMENT: Final[int] = 50


def _normalize_text(value: str) -> str:
    """Normalize whitespace and common Word punctuation."""

    return " ".join(
        value
        .replace("\xa0", " ")
        .replace("’", "'")
        .replace("–", "-")
        .replace("—", "-")
        .split()
    )


def _label_key(value: str) -> str:
    """Create a stable comparison key for a subsection label."""

    normalized = _normalize_text(
        value
    ).strip()

    normalized = normalized.rstrip(
        ":|¦= "
    )

    return normalized.casefold()


def _country_from_filename(
    file_path: Path,
) -> str:
    """Extract the country from the supported L&E filenames."""

    original_stem = file_path.stem.strip()

    cleaned_stem = _COPY_SUFFIX_PATTERN.sub(
        "",
        original_stem,
    ).strip()

    for pattern in _FILENAME_PATTERNS:
        match = pattern.fullmatch(
            cleaned_stem
        )

        if match is None:
            continue

        country = _normalize_text(
            match.group(
                "country"
            )
        ).strip(" -_")

        return _COUNTRY_ALIASES.get(
            country.casefold(),
            country,
        )

    raise ValueError(
        "Unable to detect country from filename: "
        f"{file_path.name}"
    )


def _get_heading_level(
    paragraph: Paragraph,
) -> int | None:
    """Return the effective Word heading level."""

    style_name = (
        paragraph.style.name
        if paragraph.style is not None
        else ""
    )

    style_match = _HEADING_STYLE_PATTERN.match(
        style_name.strip()
    )

    if style_match:
        return int(
            style_match.group(1)
        )

    paragraph_properties = paragraph._p.pPr

    if (
        paragraph_properties is not None
        and paragraph_properties.outlineLvl is not None
    ):
        return (
            int(
                paragraph_properties.outlineLvl.val
            )
            + 1
        )

    return None


def _has_numbering(
    paragraph: Paragraph,
) -> bool:
    """Return whether the paragraph belongs to a Word list."""

    paragraph_properties = paragraph._p.pPr

    return (
        paragraph_properties is not None
        and paragraph_properties.numPr is not None
    )


def _has_explicit_bold_text(
    paragraph: Paragraph,
) -> bool:
    """Return whether visible text is explicitly formatted in bold."""

    return any(
        run.bold is True
        for run in paragraph.runs
        if run.text.strip()
    )


def _is_explicitly_unbolded(
    paragraph: Paragraph,
) -> bool:
    """Detect a heading explicitly overridden with bold=false."""

    paragraph_properties = paragraph._p.pPr

    if paragraph_properties is not None:
        run_properties = paragraph_properties.find(
            qn("w:rPr")
        )

        if run_properties is not None:
            bold_property = run_properties.find(
                qn("w:b")
            )

            if bold_property is not None:
                value = bold_property.get(
                    qn("w:val")
                )

                if (
                    value is not None
                    and value.lower() in _FALSE_XML_VALUES
                ):
                    return True

    visible_runs = [
        run
        for run in paragraph.runs
        if run.text.strip()
    ]

    return (
        bool(visible_runs)
        and all(
            run.bold is False
            for run in visible_runs
        )
    )


def _classify_profile(
    document: Any,
    country: str,
) -> str:
    """Classify a document using its 11 legal-topic paragraphs."""

    heading_one_topics: set[str] = set()

    all_topics: set[str] = set()

    for paragraph in document.paragraphs:
        text = _normalize_text(
            paragraph.text
        )

        if not text:
            continue

        legal_topic = get_canonical_legal_topic(
            section=text,
            country=country,
        )

        if legal_topic is None:
            continue

        all_topics.add(
            legal_topic
        )

        if _get_heading_level(
            paragraph
        ) == 1:
            heading_one_topics.add(
                legal_topic
            )

    if len(all_topics) < 11:
        return "incomplete"

    if len(heading_one_topics) >= 10:
        return "standard_heading"

    if not heading_one_topics:
        return "bold_legacy"

    return "hybrid"


def _collect_reference_labels(
    documents: list[dict[str, Any]],
) -> dict[str, dict[str, Any]]:
    """
    Collect reliable Heading 2 labels from standard documents.

    These labels become evidence for identifying equivalent bold
    subsection labels in legacy and hybrid documents.
    """

    references: defaultdict[
        str,
        dict[str, Any],
    ] = defaultdict(
        lambda: {
            "count": 0,
            "variants": Counter(),
            "countries": set(),
        }
    )

    for document_info in documents:
        if document_info["profile"] != "standard_heading":
            continue

        document = document_info["document"]

        country = document_info["country"]

        for paragraph in document.paragraphs:
            text = _normalize_text(
                paragraph.text
            )

            if not text:
                continue

            if _get_heading_level(
                paragraph
            ) != 2:
                continue

            if _has_numbering(
                paragraph
            ):
                continue

            if _is_explicitly_unbolded(
                paragraph
            ):
                continue

            if get_canonical_legal_topic(
                section=text,
                country=country,
            ) is not None:
                continue

            if is_overview_section(
                section=text,
                country=country,
            ):
                continue

            key = _label_key(
                text
            )

            if not key:
                continue

            reference = references[
                key
            ]

            reference["count"] += 1

            reference["variants"][
                text
            ] += 1

            reference["countries"].add(
                country
            )

    return {
        key: {
            "count": value["count"],
            "variants": dict(
                value["variants"].most_common()
            ),
            "countries": sorted(
                value["countries"]
            ),
        }
        for key, value in references.items()
    }


def _collect_document_candidates(
    document_info: dict[str, Any],
    reference_labels: dict[str, dict[str, Any]],
) -> dict[str, Any]:
    """Find potential subsection labels in a legacy or hybrid document."""

    document = document_info["document"]

    country = document_info["country"]

    current_topic: str | None = None

    matched_candidates: list[dict[str, Any]] = []

    unmatched_candidates: list[dict[str, Any]] = []

    for paragraph_index, paragraph in enumerate(
        document.paragraphs
    ):
        text = _normalize_text(
            paragraph.text
        )

        if not text:
            continue

        legal_topic = get_canonical_legal_topic(
            section=text,
            country=country,
        )

        if legal_topic is not None:
            current_topic = legal_topic
            continue

        if is_overview_section(
            section=text,
            country=country,
        ):
            current_topic = None
            continue

        heading_level = _get_heading_level(
            paragraph
        )

        explicitly_bold = _has_explicit_bold_text(
            paragraph
        )

        if not (
            explicitly_bold
            or heading_level == 2
        ):
            continue

        if len(text) > 220:
            continue

        key = _label_key(
            text
        )

        candidate = {
            "paragraph_index": paragraph_index,
            "topic": current_topic or "Overview",
            "text": text,
            "style": (
                paragraph.style.name
                if paragraph.style is not None
                else ""
            ),
            "heading_level": heading_level,
            "numbered": _has_numbering(
                paragraph
            ),
            "explicitly_bold": explicitly_bold,
            "reference_match": key in reference_labels,
            "reference": reference_labels.get(
                key
            ),
        }

        if key in reference_labels:
            matched_candidates.append(
                candidate
            )
        else:
            unmatched_candidates.append(
                candidate
            )

    return {
        "source_filename": document_info[
            "file_path"
        ].name,
        "country": country,
        "profile": document_info["profile"],
        "matched_count": len(
            matched_candidates
        ),
        "unmatched_count": len(
            unmatched_candidates
        ),
        "matched_candidates": matched_candidates,
        "unmatched_candidates": unmatched_candidates,
    }


def _write_text_report(
    output_path: Path,
    reference_labels: dict[str, dict[str, Any]],
    reports: list[dict[str, Any]],
) -> None:
    """Write a readable subsection audit report."""

    lines = [
        "L&E subsection candidate audit",
        "=" * 80,
        (
            "Reliable subsection labels collected from "
            f"standard documents: {len(reference_labels)}"
        ),
        "",
    ]

    for report in reports:
        lines.extend(
            [
                "-" * 80,
                report["source_filename"],
                (
                    f"Country: {report['country']} | "
                    f"Profile: {report['profile']}"
                ),
                (
                    "Candidates matching standard labels: "
                    f"{report['matched_count']}"
                ),
                (
                    "Other bold candidates: "
                    f"{report['unmatched_count']}"
                ),
                "",
                "MATCHED CANDIDATES",
            ]
        )

        for candidate in report[
            "matched_candidates"
        ]:
            lines.append(
                (
                    f"  p#{candidate['paragraph_index']} "
                    f"[{candidate['topic']}] "
                    f"{candidate['text']}"
                )
            )

        lines.extend(
            [
                "",
                "UNMATCHED BOLD CANDIDATES",
            ]
        )

        unmatched_candidates = report[
            "unmatched_candidates"
        ][
            :_MAX_UNMATCHED_PER_DOCUMENT
        ]

        for candidate in unmatched_candidates:
            lines.append(
                (
                    f"  p#{candidate['paragraph_index']} "
                    f"[{candidate['topic']}] "
                    f"{candidate['text']}"
                )
            )

        remaining = (
            report["unmatched_count"]
            - len(unmatched_candidates)
        )

        if remaining > 0:
            lines.append(
                f"  ... {remaining} additional candidates in JSON"
            )

        lines.append(
            ""
        )

    output_path.write_text(
        "\n".join(
            lines
        ),
        encoding="utf-8",
    )


def main() -> None:
    if len(sys.argv) != 3:
        raise SystemExit(
            "Usage: python -m scripts.audit_subsection_candidates "
            "<source-directory> <output-directory>"
        )

    source_directory = Path(
        sys.argv[1]
    ).resolve()

    output_directory = Path(
        sys.argv[2]
    ).resolve()

    if not source_directory.is_dir():
        raise NotADirectoryError(
            "Source directory not found: "
            f"{source_directory}"
        )

    output_directory.mkdir(
        parents=True,
        exist_ok=True,
    )

    document_infos: list[dict[str, Any]] = []

    for file_path in sorted(
        source_directory.glob(
            "*.docx"
        ),
        key=lambda path: path.name.casefold(),
    ):
        country = _country_from_filename(
            file_path
        )

        document = Document(
            file_path
        )

        profile = _classify_profile(
            document=document,
            country=country,
        )

        document_infos.append(
            {
                "file_path": file_path,
                "country": country,
                "profile": profile,
                "document": document,
            }
        )

    reference_labels = _collect_reference_labels(
        document_infos
    )

    reports = [
        _collect_document_candidates(
            document_info=document_info,
            reference_labels=reference_labels,
        )
        for document_info in document_infos
        if document_info["profile"]
        in {
            "bold_legacy",
            "hybrid",
        }
    ]

    json_output = (
        output_directory
        / "subsection-candidates-audit.json"
    )

    text_output = (
        output_directory
        / "subsection-candidates-audit.txt"
    )

    json_output.write_text(
        json.dumps(
            {
                "reference_labels": reference_labels,
                "documents": reports,
            },
            ensure_ascii=False,
            indent=2,
        ),
        encoding="utf-8",
    )

    _write_text_report(
        output_path=text_output,
        reference_labels=reference_labels,
        reports=reports,
    )

    print(
        f"Documents analysed: {len(document_infos)}"
    )
    print(
        f"Reference subsection labels: {len(reference_labels)}"
    )
    print(
        f"Legacy/hybrid documents: {len(reports)}"
    )
    print(
        f"Text report: {text_output}"
    )
    print(
        f"JSON report: {json_output}"
    )


if __name__ == "__main__":
    main()