from __future__ import annotations

import hashlib
import json
import re
import sys
from collections import Counter, defaultdict
from pathlib import Path
from typing import Any, Final

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from app.core.legal_taxonomy import (
    LEGAL_TOPICS,
    _TOPIC_ALIASES,
    _TOPIC_PREFIX_PATTERN,
    _normalize_text,
)
from app.services.document_chunk_builder import (
    metadata_from_content,
)
from app.services.docx_parser import parse_docx_sections


_HEADING_STYLE_PATTERN: Final[re.Pattern[str]] = re.compile(
    r"^(?:heading|titre)\s*([1-9])$",
    re.IGNORECASE,
)

_DECORATIVE_TEXTS: Final[frozenset[str]] = frozenset({"|", "¦", "="})
_FALSE_XML_VALUES: Final[frozenset[str]] = frozenset(
    {"0", "false", "off"}
)


def _metadata_from_content(file_path: Path) -> dict[str, Any]:
    """
    Label each audit report with the same content-derived metadata the
    real ingestion pipeline would extract - the filename is never
    consulted (mission "CONTINUATION PATCH 0.4.3", section 5). Never
    raises: a document this audit script cannot label is still worth
    auditing for its structural issues, so failures degrade to
    "unknown" fields rather than aborting the run.
    """

    try:
        metadata = metadata_from_content(file_path=file_path)
    except Exception as error:
        return {
            "country": None,
            "country_code": None,
            "reference_year": None,
            "content_metadata_detected": False,
            "content_metadata_error": f"{type(error).__name__}: {error}",
        }

    return {
        "country": metadata.country,
        "country_code": metadata.country_code,
        "reference_year": metadata.reference_year,
        "content_metadata_detected": True,
        "content_metadata_error": None,
    }


def _file_sha256(file_path: Path) -> str:
    hasher = hashlib.sha256()
    with file_path.open("rb") as source_file:
        for block in iter(lambda: source_file.read(1024 * 1024), b""):
            hasher.update(block)
    return hasher.hexdigest()


def _get_heading_level(paragraph: Paragraph) -> int | None:
    style_name = paragraph.style.name if paragraph.style is not None else ""
    style_match = _HEADING_STYLE_PATTERN.match(style_name.strip())
    if style_match:
        return int(style_match.group(1))

    paragraph_properties = paragraph._p.pPr
    if (
        paragraph_properties is not None
        and paragraph_properties.outlineLvl is not None
    ):
        return int(paragraph_properties.outlineLvl.val) + 1

    return None


def _has_numbering(paragraph: Paragraph) -> bool:
    paragraph_properties = paragraph._p.pPr
    return (
        paragraph_properties is not None
        and paragraph_properties.numPr is not None
    )


def _is_explicitly_unbolded(paragraph: Paragraph) -> bool:
    paragraph_properties = paragraph._p.pPr

    if paragraph_properties is not None:
        run_properties = paragraph_properties.find(qn("w:rPr"))
        if run_properties is not None:
            bold_property = run_properties.find(qn("w:b"))
            if bold_property is not None:
                value = bold_property.get(qn("w:val"))
                if value is not None and value.lower() in _FALSE_XML_VALUES:
                    return True

    text_runs = [run for run in paragraph.runs if run.text.strip()]
    return bool(text_runs) and all(run.bold is False for run in text_runs)


def _explicit_bold_run_count(paragraph: Paragraph) -> int:
    return sum(
        1
        for run in paragraph.runs
        if run.text.strip() and run.bold is True
    )


def _normalize_topic_candidate(text: str) -> str:
    normalized = _TOPIC_PREFIX_PATTERN.sub("", _normalize_text(text))
    return normalized.strip(" |¦=-")


def _resolve_topic(text: str) -> str | None:
    candidate = _normalize_topic_candidate(text).casefold()

    for topic, aliases in _TOPIC_ALIASES.items():
        for alias in aliases:
            if candidate == alias or candidate.startswith(f"{alias} in "):
                return topic

    return None


def _is_overview_heading(text: str) -> bool:
    normalized = _normalize_text(text).casefold()
    return normalized.startswith("employment law overview") or normalized.startswith(
        "labour and employment law overview"
    )


def _candidate_score(
    paragraph: Paragraph,
    text: str,
    heading_level: int | None,
) -> int:
    score = 0

    if heading_level == 1:
        score += 100
    elif heading_level is not None:
        score += max(10, 60 - heading_level * 10)

    style_name = (
        paragraph.style.name.casefold()
        if paragraph.style is not None
        else ""
    )
    if style_name == "list paragraph":
        score += 30

    if re.match(r"^\s*(?:[|¦=]+\s*)?\d{1,2}\s*[.)]", text):
        score += 20

    score += min(_explicit_bold_run_count(paragraph) * 5, 20)

    if _has_numbering(paragraph):
        score -= 15
    if _is_explicitly_unbolded(paragraph):
        score -= 25

    return score


def _find_topic_candidates(
    document: Any,
) -> tuple[dict[str, dict[str, Any]], list[dict[str, Any]]]:
    candidates_by_topic: defaultdict[str, list[dict[str, Any]]] = defaultdict(list)

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        text = _normalize_text(paragraph.text)
        if not text:
            continue

        topic = _resolve_topic(text)
        if topic is None:
            continue

        heading_level = _get_heading_level(paragraph)
        candidate = {
            "topic": topic,
            "paragraph_index": paragraph_index,
            "text": text,
            "style": paragraph.style.name if paragraph.style is not None else "",
            "heading_level": heading_level,
            "numbered": _has_numbering(paragraph),
            "explicitly_unbolded": _is_explicitly_unbolded(paragraph),
            "explicit_bold_runs": _explicit_bold_run_count(paragraph),
            "score": _candidate_score(paragraph, text, heading_level),
        }
        candidates_by_topic[topic].append(candidate)

    selected: dict[str, dict[str, Any]] = {}
    duplicates: list[dict[str, Any]] = []

    for topic, candidates in candidates_by_topic.items():
        ordered = sorted(
            candidates,
            key=lambda candidate: (
                -candidate["score"],
                candidate["paragraph_index"],
            ),
        )
        selected[topic] = ordered[0]
        duplicates.extend(ordered[1:])

    return selected, duplicates


def _classify_profile(selected_topics: dict[str, dict[str, Any]]) -> str:
    if len(selected_topics) < len(LEGAL_TOPICS):
        return "incomplete"

    heading_one_count = sum(
        candidate["heading_level"] == 1
        for candidate in selected_topics.values()
    )

    if heading_one_count >= 10:
        return "standard_heading"
    if heading_one_count == 0:
        return "bold_legacy"
    return "hybrid"


def _find_suspicious_paragraphs(
    document: Any,
    selected_topics: dict[str, dict[str, Any]],
) -> list[dict[str, Any]]:
    selected_indexes = {
        candidate["paragraph_index"]
        for candidate in selected_topics.values()
    }
    suspicious: list[dict[str, Any]] = []

    for paragraph_index, paragraph in enumerate(document.paragraphs):
        heading_level = _get_heading_level(paragraph)
        if heading_level is None:
            continue

        text = _normalize_text(paragraph.text)
        style_name = paragraph.style.name if paragraph.style is not None else ""

        if not text:
            suspicious.append(
                {
                    "paragraph_index": paragraph_index,
                    "text": "",
                    "style": style_name,
                    "heading_level": heading_level,
                    "reason": "empty_heading",
                }
            )
            continue

        if paragraph_index in selected_indexes or _is_overview_heading(text):
            continue

        if heading_level == 1:
            suspicious.append(
                {
                    "paragraph_index": paragraph_index,
                    "text": text,
                    "style": style_name,
                    "heading_level": heading_level,
                    "reason": "unrecognized_heading_1",
                }
            )
            continue

        reasons: list[str] = []
        if _has_numbering(paragraph):
            reasons.append("numbered_heading")
        if _is_explicitly_unbolded(paragraph):
            reasons.append("explicitly_unbolded_heading")

        if reasons:
            suspicious.append(
                {
                    "paragraph_index": paragraph_index,
                    "text": text,
                    "style": style_name,
                    "heading_level": heading_level,
                    "reason": ",".join(reasons),
                }
            )

    return suspicious


def _audit_current_parser(file_path: Path) -> dict[str, Any]:
    try:
        sections = parse_docx_sections(file_path)
    except Exception as error:  # diagnostic script: keep auditing other files
        return {
            "status": "error",
            "error": f"{type(error).__name__}: {error}",
        }

    resolved_topics = {
        topic
        for section in sections
        if (topic := _resolve_topic(section.section)) is not None
    }
    main_sections = sorted({section.section for section in sections})
    unknown_main_sections = sorted(
        {
            section.section
            for section in sections
            if section.section != "General"
            and not _is_overview_heading(section.section)
            and _resolve_topic(section.section) is None
        }
    )
    long_sections = sorted(
        (
            {
                "section": section.section,
                "subsection": section.subsection,
                "content_length": len(section.content),
            }
            for section in sections
            if len(section.content) > 6000
        ),
        key=lambda item: item["content_length"],
        reverse=True,
    )

    return {
        "status": "ok",
        "estimated_chunks": len(sections),
        "main_sections": main_sections,
        "general_section_used": "General" in main_sections,
        "legal_topics_reached": len(resolved_topics),
        "unknown_main_sections": unknown_main_sections,
        "max_content_length": max(
            (len(section.content) for section in sections),
            default=0,
        ),
        "sections_over_6000_chars": long_sections,
    }


def _audit_file(file_path: Path) -> dict[str, Any]:
    document = Document(file_path)
    metadata = _metadata_from_content(file_path)
    selected_topics, duplicate_candidates = _find_topic_candidates(document)
    suspicious = _find_suspicious_paragraphs(document, selected_topics)

    heading_counts = Counter(
        level
        for paragraph in document.paragraphs
        if (level := _get_heading_level(paragraph)) is not None
    )
    style_counts = Counter(
        paragraph.style.name if paragraph.style is not None else "<none>"
        for paragraph in document.paragraphs
    )

    return {
        "source_filename": file_path.name,
        "file_sha256": _file_sha256(file_path),
        **metadata,
        "profile": _classify_profile(selected_topics),
        "paragraph_count": len(document.paragraphs),
        "table_count": len(document.tables),
        "decorative_paragraph_count": sum(
            _normalize_text(paragraph.text) in _DECORATIVE_TEXTS
            for paragraph in document.paragraphs
        ),
        "heading_counts": {
            str(level): count
            for level, count in sorted(heading_counts.items())
        },
        "top_styles": dict(style_counts.most_common(10)),
        "legal_topics_found": len(selected_topics),
        "legal_topics_missing": [
            topic for topic in LEGAL_TOPICS if topic not in selected_topics
        ],
        "topic_candidates": [
            selected_topics[topic]
            for topic in LEGAL_TOPICS
            if topic in selected_topics
        ],
        "duplicate_topic_candidates": duplicate_candidates,
        "suspicious_paragraph_count": len(suspicious),
        "suspicious_paragraphs": suspicious,
        "current_parser": _audit_current_parser(file_path),
    }


def _duplicate_groups(reports: list[dict[str, Any]]) -> list[dict[str, Any]]:
    files_by_hash: defaultdict[str, list[str]] = defaultdict(list)
    for report in reports:
        files_by_hash[report["file_sha256"]].append(report["source_filename"])

    return [
        {"file_sha256": file_hash, "files": sorted(file_names)}
        for file_hash, file_names in files_by_hash.items()
        if len(file_names) > 1
    ]


def _build_text_report(
    reports: list[dict[str, Any]],
    duplicate_groups: list[dict[str, Any]],
) -> str:
    profile_counts = Counter(report["profile"] for report in reports)
    lines = [
        "L&E DOCX corpus audit",
        "=" * 80,
        f"Files analysed: {len(reports)}",
        f"Profiles: {dict(sorted(profile_counts.items()))}",
        f"Exact duplicate groups: {len(duplicate_groups)}",
        "",
    ]

    if duplicate_groups:
        lines.append("Exact duplicate files:")
        for group in duplicate_groups:
            lines.append("- " + " | ".join(group["files"]))
        lines.append("")

    for report in reports:
        parser = report["current_parser"]
        lines.extend(
            [
                "-" * 80,
                report["source_filename"],
                (
                    "Metadata: "
                    f"country={report['country'] or '<unknown>'}, "
                    f"code={report['country_code'] or '<unknown>'}, "
                    f"year={report['reference_year'] or '<missing>'}"
                ),
                f"Profile: {report['profile']}",
                f"Raw topic detection: {report['legal_topics_found']}/11",
                (
                    "Missing topics: "
                    + (", ".join(report["legal_topics_missing"]) or "none")
                ),
                (
                    "Paragraphs / tables / decorative separators: "
                    f"{report['paragraph_count']} / {report['table_count']} / "
                    f"{report['decorative_paragraph_count']}"
                ),
                f"Heading counts: {report['heading_counts']}",
                f"Suspicious paragraphs: {report['suspicious_paragraph_count']}",
            ]
        )

        if parser["status"] == "ok":
            lines.extend(
                [
                    f"Current parser chunks: {parser['estimated_chunks']}",
                    f"Current parser topic coverage: {parser['legal_topics_reached']}/11",
                    f"Current parser max content length: {parser['max_content_length']}",
                    (
                        "Current parser unknown main sections: "
                        + (", ".join(parser["unknown_main_sections"]) or "none")
                    ),
                ]
            )
        else:
            lines.append(f"Current parser error: {parser['error']}")

        for item in report["suspicious_paragraphs"][:10]:
            lines.append(
                "  * "
                f"p#{item['paragraph_index']} "
                f"[{item['style']}; H{item['heading_level']}] "
                f"{item['reason']}: {item['text'][:180]}"
            )

        if report["suspicious_paragraph_count"] > 10:
            lines.append("  * ... see JSON for the full list")

        lines.append("")

    return "\n".join(lines)


def main() -> None:
    if len(sys.argv) not in {2, 3}:
        raise SystemExit(
            "Usage: python -m scripts.audit_docx_corpus "
            "<source-directory> [output-directory]"
        )

    source_directory = Path(sys.argv[1]).resolve()
    output_directory = (
        Path(sys.argv[2]).resolve()
        if len(sys.argv) == 3
        else source_directory.parent / "processed"
    )

    if not source_directory.is_dir():
        raise NotADirectoryError(f"Source directory not found: {source_directory}")

    docx_files = sorted(
        source_directory.glob("*.docx"),
        key=lambda path: path.name.casefold(),
    )
    if not docx_files:
        raise FileNotFoundError(f"No DOCX files found in: {source_directory}")

    reports = [_audit_file(file_path) for file_path in docx_files]
    duplicate_groups = _duplicate_groups(reports)
    profile_counts = Counter(report["profile"] for report in reports)

    output_data = {
        "summary": {
            "files_analysed": len(reports),
            "profiles": dict(sorted(profile_counts.items())),
            "exact_duplicate_groups": len(duplicate_groups),
        },
        "duplicate_groups": duplicate_groups,
        "documents": reports,
    }

    output_directory.mkdir(parents=True, exist_ok=True)
    json_path = output_directory / "docx-corpus-audit.json"
    text_path = output_directory / "docx-corpus-audit.txt"

    json_path.write_text(
        json.dumps(output_data, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    text_path.write_text(
        _build_text_report(reports, duplicate_groups),
        encoding="utf-8",
    )

    print(f"Files analysed: {len(reports)}")
    print(f"Profiles: {dict(sorted(profile_counts.items()))}")
    print(f"Exact duplicate groups: {len(duplicate_groups)}")
    print(f"JSON report: {json_path}")
    print(f"Text report: {text_path}")


if __name__ == "__main__":
    main()
